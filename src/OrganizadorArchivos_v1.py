import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import os
import shutil
import re
import random
import math

import pandas as pd # type: ignore
import threading
from docx import Document # type: ignore
from docx.shared import Inches # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
import json
from send2trash import send2trash # type: ignore
import fitz  # type: ignore # PyMuPDF
from PIL import Image, ImageDraw, ImageFont # type: ignore
import io
from zipfile import ZipFile, ZIP_DEFLATED
import string
from pdf2docx import Converter # type: ignore # type: ignore
from docx2pdf import convert # type: ignore
import fnmatch
import openpyxl  # type: ignore
import xml.etree.ElementTree as ET
import requests # type: ignore # type: ignore
import subprocess
import win32com.client # type: ignore
# EL 'import pyperclip' SE MUEVE DENTRO DE LA CLASE PARA MANEJAR EL ERROR

# === INICIO NUEVA FUNCIONALIDAD OVIDA ===
import time
import urllib.parse
import base64
# Las importaciones de Selenium se manejan dentro de un try/except para no romper la app si no están instaladas
# === FIN NUEVA FUNCIONALIDAD OVIDA ===


class OrganizadorArchivosApp:
    def __init__(self, root):
        self.root = root
        
        # ### CORRECCIÓN: La importación ahora está dentro del try/except ###
        try:
            import pyperclip # type: ignore # type: ignore
        except ImportError:
            messagebox.showwarning("Dependencia Faltante", 
                                   "La librería 'pyperclip' no está instalada.\n\n"
                                   "La función de 'Copiar al portapapeles' en el editor de datos no estará disponible.\n\n"
                                   "Puedes instalarla con: pip install pyperclip",
                                   parent=self.root)
            self.pyperclip = None
            self.pyperclip_available = False
        else:
            self.pyperclip = pyperclip
            self.pyperclip_available = True

        # === INICIO NUEVA FUNCIONALIDAD OVIDA ===
        # Comprobar dependencias de Selenium
        try:
            from selenium import webdriver # type: ignore
            from selenium.webdriver.common.by import By # type: ignore
            from selenium.webdriver.support.ui import WebDriverWait, Select # type: ignore
            from selenium.webdriver.support import expected_conditions as EC # type: ignore
            from selenium.webdriver.chrome.service import Service # type: ignore
            from selenium.webdriver.common.keys import Keys # type: ignore
            from webdriver_manager.chrome import ChromeDriverManager # type: ignore
            self.selenium_available = True
            self.selenium_deps = {
                "webdriver": webdriver, "By": By, "WebDriverWait": WebDriverWait,
                "Select": Select, "EC": EC, "Service": Service, "Keys": Keys,
                "ChromeDriverManager": ChromeDriverManager
            }
        except ImportError:
            self.selenium_available = False
        # === FIN NUEVA FUNCIONALIDAD OVIDA ===

        style = ttk.Style(self.root)
        style.theme_use("vista")

        self.root.title("📂 Organizador de Archivos v5.7 (Mapeo Parcial)")
        self.root.geometry("1100x850")
       
        # Variables de la Tab 1 (Búsqueda)
        self.source_path = tk.StringVar()
        self.search_pattern = tk.StringVar()
        self.filter_type = tk.StringVar(value="extensión")
        self.include_folders = tk.BooleanVar(value=True)
        self.element_type = tk.StringVar(value="archivos")
        self.accion = tk.StringVar(value="copiar a carpeta")
        self.resultados = []

        # --- Variables para el editor de datos (Tab 4) ---
        self.data_editor_path = tk.StringVar()
        self.data_tree = None
        self.save_button = None
        self.parsed_data = None
        self.current_file_path = None
        self.current_data_type = None
        self.is_data_modified = False
        self.embedded_xml_docs = {}
        self.tooltip_window = None 
        self.xml_element_map = {}
        
        # --- Variables para la búsqueda en el Treeview ---
        self.search_data_var = tk.StringVar()
        self.search_results_iids = []

        self.data_context_menu = None
        
        self.create_widgets()
        self._on_filter_type_change()

    def _create_progress_window(self, title, max_value):
        """Crea y devuelve una ventana emergente con una barra de progreso."""
        progress_win = tk.Toplevel(self.root)
        progress_win.title(title)
        progress_win.geometry("400x120")
        progress_win.resizable(False, False)
        progress_win.transient(self.root)
        progress_win.grab_set()

        win_x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 200
        win_y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 60
        progress_win.geometry(f'+{win_x}+{win_y}')

        progress_label = ttk.Label(progress_win, text="Iniciando...", padding=(10, 10))
        progress_label.pack()

        progress_bar = ttk.Progressbar(
            progress_win,
            orient='horizontal',
            length=360,
            mode='determinate',
            maximum=max_value,
            value=0
        )
        progress_bar.pack(pady=10, padx=20)
        
        return progress_win, progress_bar, progress_label
    
    def _on_filter_type_change(self, event=None):
        """Actualiza la etiqueta del patrón según el tipo de filtro seleccionado."""
        filtro = self.filter_type.get()
        if filtro == "Patrón con comodín (*, ?)":
            self.pattern_label.config(text="Patrón (ej: prueba.docx):")
        else:
            self.pattern_label.config(text="Patrón:")

    def create_widgets(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(pady=10, padx=10, fill="both", expand=True)
        
          # ================== TAB 1: Búsqueda y Acciones ==================

        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="🔎 Búsqueda y Acciones")
        tab1.columnconfigure(0, weight=1)
        tab1.columnconfigure(1, weight=1)
        tab1.rowconfigure(2, weight=1)

        origen_frame = ttk.LabelFrame(tab1, text="📁 Carpeta de origen", padding=10)
        origen_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        origen_frame.columnconfigure(0, weight=1)
        ttk.Entry(origen_frame, textvariable=self.source_path).grid(row=0, column=0, sticky="ew", padx=(0, 5))
        ttk.Button(origen_frame, text="Examinar", command=self.seleccionar_directorio).grid(row=0, column=1)

        filtro_frame = ttk.LabelFrame(tab1, text="🔎 Criterios de búsqueda", padding=10)
        filtro_frame.grid(row=1, column=0, sticky="nsew", pady=5, padx=5)
        ttk.Label(filtro_frame, text="Buscar por:").grid(row=0, column=0, sticky="w")
        
        filtro_combobox = ttk.Combobox(filtro_frame, textvariable=self.filter_type, 
                                     values=["extensión", "nombre", "Patrón con comodín (*, ?)", "expresión regular"], 
                                     state="readonly")
        filtro_combobox.grid(row=0, column=1, sticky="ew", padx=5)
        filtro_combobox.bind("<<ComboboxSelected>>", self._on_filter_type_change)

        self.pattern_label = ttk.Label(filtro_frame, text="Patrón:")
        self.pattern_label.grid(row=1, column=0, sticky="w", pady=(5, 0))
        ttk.Entry(filtro_frame, textvariable=self.search_pattern).grid(row=1, column=1, sticky="ew", padx=5, pady=(5, 0))
        
        ttk.Checkbutton(filtro_frame, text="Incluir subcarpetas", variable=self.include_folders).grid(row=2, column=0, columnspan=2, sticky="w", pady=(5, 0))
        ttk.Label(filtro_frame, text="Tipo de elemento:").grid(row=3, column=0, sticky="w", pady=(5, 0))
        ttk.Combobox(filtro_frame, textvariable=self.element_type, values=["archivos", "carpetas", "todos"], state="readonly").grid(row=3, column=1, sticky="ew", padx=5, pady=(5, 0))

        accion_frame = ttk.LabelFrame(tab1, text="🛠️ Acción a realizar", padding=10)
        accion_frame.grid(row=1, column=1, sticky="nsew", pady=5, padx=5)
        acciones_principales = ["Copiar a carpeta", "Mover a carpeta", "Modificar nombre", "Editar texto", "Comprimir archivos en ZIP", "Comprimir carpetas individualmente"]
        for i, texto in enumerate(acciones_principales):
            ttk.Radiobutton(accion_frame, text=texto, variable=self.accion, value=texto.lower()).grid(row=i, column=0, sticky="w")

        resultado_frame = ttk.LabelFrame(tab1, text="📄 Archivos encontrados", padding=10)
        resultado_frame.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=5, padx=5)
        resultado_frame.columnconfigure(0, weight=1)
        resultado_frame.rowconfigure(0, weight=1)
        self.tree = ttk.Treeview(resultado_frame, columns=("ruta", "fecha"), show="headings")
        self.tree.heading("ruta", text="Ruta completa")
        self.tree.heading("fecha", text="Fecha")
        self.tree.column("ruta", width=600)
        self.tree.column("fecha", width=150)
        self.tree.grid(row=0, column=0, sticky="nsew")

        acciones_inf_frame = ttk.Frame(tab1, padding=5)
        acciones_inf_frame.grid(row=3, column=0, columnspan=2, sticky="ew", padx=5)
        ttk.Button(acciones_inf_frame, text="Buscar archivos", command=self.buscar_archivos).pack(side="left", padx=(0, 10))
        ttk.Button(acciones_inf_frame, text="Ejecutar acción", command=self.ejecutar_accion).pack(side="left", padx=(0, 10))
        ttk.Button(acciones_inf_frame, text="Limpiar", command=self.limpiar).pack(side="left")
        ttk.Button(acciones_inf_frame, text="🗑️ Eliminar resultados", command=self.eliminar_resultados).pack(side="left", padx=(10, 0))
        ttk.Button(acciones_inf_frame, text="🚪 Salir", command=self.root.quit).pack(side="right")
        
        # ================== TAB 2: Acciones Automatizadas ==================

        tab2 = ttk.Frame(notebook)
        notebook.add(tab2, text="⚙️ Acciones Automatizadas")

        auto_col1 = ttk.Frame(tab2, padding=10)
        auto_col1.pack(side="left", fill="y", anchor="n", padx=5)
        auto_col2 = ttk.Frame(tab2, padding=10)
        auto_col2.pack(side="left", fill="y", anchor="n", padx=5)
        auto_col3 = ttk.Frame(tab2, padding=10)
        auto_col3.pack(side="left", fill="y", anchor="n", padx=5)
        
        accion_personalizada_frame = ttk.LabelFrame(auto_col1, text="Unificación y División Especial", padding=10)
        accion_personalizada_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(accion_personalizada_frame, text="🗂️ Unificar PDF por Carpeta", command=self.accion_unificar_por_carpeta).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="🖼️ Unificar JPG por Carpeta", command=self.accion_unificar_jpg_por_carpeta).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="🖼️ Unificar PNG por Carpeta", command=self.accion_unificar_png_por_carpeta).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="📄 Unificar DOCX por Carpeta", command=self.accion_unificar_docx_por_carpeta).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="🧷 Unificar PDFs (Selección Manual)", command=self.accion_unificar_pdfs).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="✂️ Dividir PDF en Páginas (Manual)", command=self.accion_dividir_pdf_en_paginas).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="✂️ Dividir PDFs Masivamente (por Carpeta)", command=self.accion_dividir_pdfs_masivamente).pack(anchor="w", pady=2, fill="x")

        organizacion_frame = ttk.LabelFrame(auto_col1, text="Organización de Archivo", padding=10)
        organizacion_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(organizacion_frame, text="📥 Organizar Facturas (FEOV)", command=self.accion_organizar_facturas_por_pdf).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📂➡️📁 Mover Archivos por Coincidencia de Nombre", command=self.accion_mover_archivos_por_coincidencia_nombre).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="🗺️ Copiar Archivos (Mapeo Subcarpetas)", command=self.accion_copiar_archivos_desde_mapeo_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📜 Copiar Archivos de Raíz (Mapeo Parcial)", command=self.accion_copiar_archivos_desde_raiz_mapeo_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📤 Consolidar Archivos de Subcarpetas", command=self.accion_consolidar_archivos_subcarpetas).pack(anchor="w", pady=2, fill="x")

        analisis_frame = ttk.LabelFrame(auto_col3, text="Análisis de Archivos", padding=10)
        analisis_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(analisis_frame, text="📊 Análisis de Carpetas a Excel", command=self.exportar_lista).pack(anchor="w", pady=2, fill="x")
        ttk.Button(analisis_frame, text="📊 Análisis de Historia Clínica", command=self.accion_analisis_historia_clinica).pack(anchor="w", pady=2, fill="x")
        ttk.Button(analisis_frame, text="📊 Análisis Autorización Nueva EPS", command=self.accion_analisis_autorizacion_nueva_eps).pack(anchor="w", pady=2, fill="x")
        ttk.Button(analisis_frame, text="📊 Análisis Cargue Sanitas", command=self.accion_analisis_cargue_sanitas).pack(anchor="w", pady=2, fill="x")
        ttk.Button(analisis_frame, text="📊 Análisis Retefuente y Ica", command=self.accion_leer_pdf_retefuente).pack(anchor="w", pady=2, fill="x")
        
        renombrado_excel_frame = ttk.LabelFrame(auto_col2, text="Modificación y Renombrado con Excel", padding=10)
        renombrado_excel_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(renombrado_excel_frame, text="📤 Exportar para renombrar", command=self.exportar_para_renombrar).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="📥 Aplicar renombrado desde Excel", command=self.importar_y_aplicar_renombrado).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="🏷️ Añadir Sufijo desde Excel", command=self.accion_anadir_sufijo_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✍️ Modificar Autorización DOCX (Excel)", command=self.accion_autorizacion_docx_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✍️ Modificar Régimen DOCX (Excel)", command=self.accion_regimen_docx_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✍️ Modificar DOCX Completo (Excel)", command=self.accion_modificar_docx_completo_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="🖋️ Firmar DOCX con Imagen (por Carpeta)", command=self.accion_firmar_docx_con_imagen).pack(anchor="w", pady=2, fill="x")

        masivo_frame = ttk.LabelFrame(auto_col2, text="Conversión Masiva (desde resultados)", padding=10)
        masivo_frame.pack(fill="x", pady=5, anchor="n")

        ttk.Button(masivo_frame, text="PDF → DOCX", command=self.accion_pdf_a_docx_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="JPG → PDF", command=self.accion_jpg_a_pdf_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="DOCX → PDF", command=self.accion_docx_a_pdf_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="PDF → JPG", command=self.accion_pdf_a_jpg_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="PNG → JPG", command=self.accion_png_a_jpg_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="TXT → JSON (Renombrar)", command=self.accion_txt_a_json_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="📄 PDF → PDF (Escala de Grises)", command=self.accion_pdf_a_escala_grises_masivo_desde_resultados).pack(anchor="w", pady=2, fill="x")

        renombrado_excel_frame = ttk.LabelFrame(auto_col3, text="Creación de Archivos", padding=10)
        renombrado_excel_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(renombrado_excel_frame, text="📂 Crear Carpetas desde Excel", command=self.accion_crear_carpetas_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="⬇️ Descargar Firmas (URL/Excel)", command=self.accion_descargar_firmas_url_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="📤 Copiar Archivo a Subcarpetas", command=self.accion_copiar_archivo_a_subcarpetas).pack(anchor="w", pady=2, fill="x")
          # === INICIO NUEVA FUNCIONALIDAD OVIDA ===
        ttk.Button(renombrado_excel_frame, text="⬇️ Descargar Historias Hospitalización (OVIDA)", command=self.accion_descargar_historias_hospitalizacion_ovida).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✒️ Crear Firma Digital desde Nombre de Carpeta", command=self.accion_crear_firma_digital).pack(anchor="w", pady=2, fill="x")

        # ================== TAB 3: Conversión Individual ==================

        tab3 = ttk.Frame(notebook)
        notebook.add(tab3, text="🔄 Conversión de Archivos Individuales")
        
        conversion_frame = ttk.LabelFrame(tab3, text="Seleccionar archivos y convertir", padding=20)
        conversion_frame.pack(pady=20, padx=20)

        ttk.Button(conversion_frame, text="PDF → DOCX", command=self.accion_pdf_a_docx).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="JPG → PDF", command=self.accion_jpg_a_pdf).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="DOCX → PDF", command=self.accion_docx_a_pdf).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="PDF → JPG", command=self.accion_pdf_a_jpg).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="PNG → JPG", command=self.accion_png_a_jpg).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="TXT → JSON (Renombrar)", command=self.accion_txt_a_json).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="📄 PDF → PDF (Escala de Grises)", command=self.accion_pdf_a_escala_grises_individual).pack(anchor="w", pady=5, fill="x")
            
        # ================== TAB 4: Visor y Editor de Datos ==================
        tab4 = ttk.Frame(notebook)
        notebook.add(tab4, text="📄 Visor (JSON/XML)")
        self._setup_data_editor_tab(tab4) 

        # ================== TAB 5: RIPS ==================
        tab5 = ttk.Frame(notebook)
        notebook.add(tab5, text="RIPS")
        
        rips_frame = ttk.LabelFrame(tab5, text="Convertidor", padding=20)
        rips_frame.pack(pady=20, padx=20, fill="x")
        
        ttk.Button(rips_frame, text="JSON a XLSX (Individual)", command=self.accion_convertir_json_rips_a_xlsx).pack(anchor="w", pady=5, fill="x") 
        ttk.Button(rips_frame, text="XLSX a JSON (Individual)", command=self.accion_convertir_xlsx_rips_a_json).pack(anchor="w", pady=5, fill="x") 
        ttk.Button(rips_frame, text="JSON Evento a XLSX (Masivo)", command=self.accion_json_evento_a_xlsx_masivo).pack(anchor="w", pady=5, fill="x") 
        ttk.Button(rips_frame, text="XLSX Evento a JSONs (Masivo)", command=self.accion_xlsx_evento_a_json_masivo).pack(anchor="w", pady=5, fill="x") 
        
 

        # ================== Log de Actividad (al final) ==================

        log_frame = ttk.LabelFrame(self.root, text="🧾 Registro de actividad", padding=10)
        log_frame.pack(side="bottom", fill="x", padx=10, pady=(0,10))
        self.log_text = tk.Text(log_frame, height=8, wrap="word", state="disabled")
        self.log_text.pack(fill="both", expand=True)
        log_frame.columnconfigure(0, weight=1)
        
    def log(self, mensaje):
        self.log_text.config(state="normal")
        self.log_text.insert(tk.END, f"{mensaje}\n")
        self.log_text.config(state="disabled")
        self.log_text.see(tk.END)

    def seleccionar_directorio(self):
        ruta = filedialog.askdirectory(parent=self.root)
        if ruta:
            self.source_path.set(ruta)

    def buscar_archivos(self):
        self.tree.delete(*self.tree.get_children())
        carpeta = self.source_path.get()
        patron = self.search_pattern.get()
        filtro = self.filter_type.get()
        incluir_sub = self.include_folders.get()
        tipo = self.element_type.get()
        self.resultados = []

        if not carpeta or not os.path.isdir(carpeta):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida.", parent=self.root)
            return

        for root, dirs, files in os.walk(carpeta):
            if not incluir_sub and root != carpeta:
                continue

            elementos = []
            if tipo == 'carpetas':
                elementos = dirs
            elif tipo == 'archivos':
                elementos = files
            else:
                elementos = dirs + files

            for nombre in elementos:
                ruta = os.path.join(root, nombre)
                es_carpeta = os.path.isdir(ruta)
                match = False

                if not patron:
                    match = True
                elif filtro == "extensión" and not es_carpeta:
                    match = nombre.lower().endswith(patron.lower())
                elif filtro == "nombre":
                    # Si el patrón empieza por punto, asumimos que busca por extensión (o sufijo)
                    if patron.startswith('.'):
                        match = nombre.lower().endswith(patron.lower())
                    else:
                        match = patron.lower() in nombre.lower()
                
                elif filtro == "Patrón con comodín (*, ?)":
                    try:
                        user_pattern = patron.lower()
                        effective_pattern = ''
                        if '*' in user_pattern or '?' in user_pattern:
                            effective_pattern = user_pattern
                        else:
                            if '.' in user_pattern:
                                name_part, ext_part = user_pattern.rsplit('.', 1)
                                effective_pattern = f"*{name_part}*.{ext_part}"
                            else:
                                effective_pattern = f"*{user_pattern}*"
                        match = fnmatch.fnmatch(nombre.lower(), effective_pattern)
                    except Exception as e:
                        self.log(f"Error en patrón con comodín: {e}")
                        match = False

                elif filtro == "expresión regular":
                    try:
                        match = bool(re.search(patron, nombre, re.IGNORECASE))
                    except re.error as e:
                        self.log(f"Error en expresión regular: {e}")
                        match = False
                
                if match:
                    try:
                        fecha_mod = os.path.getmtime(ruta)
                        fecha = pd.to_datetime(fecha_mod, unit='s').strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        fecha = "N/A"
                    self.tree.insert('', 'end', values=(ruta, fecha))
                    self.resultados.append(ruta)

        self.log(f"{len(self.resultados)} resultados encontrados.")

    def mostrar_dialogo_modificar_nombre(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("🏷️ Modificar Nombres - Opciones Avanzadas")
        dialog.geometry("520x680") 
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        opciones = {
            'sustituir_activo': tk.BooleanVar(), 'sustituir_buscar': tk.StringVar(), 'sustituir_reemplazar': tk.StringVar(),
            'prefijo_activo': tk.BooleanVar(), 'prefijo_texto': tk.StringVar(),
            'sufijo_activo': tk.BooleanVar(), 'sufijo_texto': tk.StringVar(),
            'renombrar_activo': tk.BooleanVar(), 'renombrar_texto': tk.StringVar(),
            'limpiar_id_activo': tk.BooleanVar(),
            'resultado': None
        }

        main_frame = ttk.Frame(dialog, padding=20)
        main_frame.pack(fill='both', expand=True)
        
        renombrar_frame = ttk.LabelFrame(main_frame, text="📝 Renombrado completo (ignora las demás opciones)", padding=10)
        renombrar_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(renombrar_frame, text="Activar renombrado completo", variable=opciones['renombrar_activo']).pack(anchor='w')
        ttk.Label(renombrar_frame, text="Nuevo nombre (sin extensión):").pack(anchor='w', pady=(5, 0))
        ttk.Entry(renombrar_frame, textvariable=opciones['renombrar_texto']).pack(fill='x', pady=2)

        sustituir_frame = ttk.LabelFrame(main_frame, text="🔄 Sustituir texto", padding=10)
        sustituir_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(sustituir_frame, text="Activar sustitución", variable=opciones['sustituir_activo']).pack(anchor='w')
        ttk.Label(sustituir_frame, text="Buscar:").pack(anchor='w', pady=(5, 0))
        ttk.Entry(sustituir_frame, textvariable=opciones['sustituir_buscar']).pack(fill='x', pady=2)
        ttk.Label(sustituir_frame, text="Reemplazar con:").pack(anchor='w')
        ttk.Entry(sustituir_frame, textvariable=opciones['sustituir_reemplazar']).pack(fill='x', pady=2)

        limpieza_frame = ttk.LabelFrame(main_frame, text="🧹 Limpieza Especial (FEOV)", padding=10)
        limpieza_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(limpieza_frame, text="Eliminar '_ID<números>_A' del nombre", variable=opciones['limpiar_id_activo']).pack(anchor='w', pady=5)

        prefijo_frame = ttk.LabelFrame(main_frame, text="⬅️ Añadir al inicio", padding=10)
        prefijo_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(prefijo_frame, text="Añadir prefijo", variable=opciones['prefijo_activo']).pack(anchor='w')
        ttk.Entry(prefijo_frame, textvariable=opciones['prefijo_texto']).pack(fill='x', pady=2)

        sufijo_frame = ttk.LabelFrame(main_frame, text="➡️ Añadir al final (antes de extensión)", padding=10)
        sufijo_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(sufijo_frame, text="Añadir sufijo", variable=opciones['sufijo_activo']).pack(anchor='w')
        ttk.Entry(sufijo_frame, textvariable=opciones['sufijo_texto']).pack(fill='x', pady=2)

        def aplicar():
            if not (opciones['sustituir_activo'].get() or opciones['prefijo_activo'].get() or opciones['sufijo_activo'].get() or opciones['renombrar_activo'].get() or opciones['limpiar_id_activo'].get()):
                messagebox.showwarning("Advertencia", "Debes activar y rellenar al menos una opción.", parent=dialog)
                return
            opciones['resultado'] = opciones
            dialog.destroy()

        botones_frame = ttk.Frame(main_frame)
        botones_frame.pack(fill='x', pady=20)
        ttk.Button(botones_frame, text="✅ Ejecutar Cambios", command=aplicar).pack(side="right", padx=5)
        ttk.Button(botones_frame, text="❌ Cancelar", command=dialog.destroy).pack(side="right")

        self.root.wait_window(dialog)
        return opciones['resultado']

    def ejecutar_accion(self):
        accion = self.accion.get()
        total_resultados = len(self.resultados)
        if total_resultados == 0:
            self.log("No hay archivos seleccionados para ejecutar la acción.")
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar una acción sobre ellos.", parent=self.root)
            return

        progress_win, progress_bar, progress_label = None, None, None

        try:
            if accion in ["copiar a carpeta", "mover a carpeta"]:
                action_word = "copiar" if accion == "copiar a carpeta" else "mover"
                destino = filedialog.askdirectory(parent=self.root, title=f'Seleccionar carpeta destino para {action_word}')
                if not destino: return
                progress_win, progress_bar, progress_label = self._create_progress_window(f"{action_word.capitalize()}ndo archivos...", total_resultados)
                for i, ruta in enumerate(self.resultados):
                    nombre_archivo = os.path.basename(ruta)
                    progress_label.config(text=f"{action_word.capitalize()}ndo: {nombre_archivo}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    try:
                        if accion == "copiar a carpeta":
                            shutil.copy2(ruta, destino)
                        else:
                            shutil.move(ruta, destino)
                        self.log(f'{action_word.capitalize()}do: {nombre_archivo} a {destino}')
                    except Exception as e:
                        self.log(f'Error al {action_word} {ruta}: {e}')

            elif accion == "modificar nombre":
                opciones = self.mostrar_dialogo_modificar_nombre()
                if not opciones:
                    self.log("Renombrado cancelado por el usuario.")
                    return

                procesados, errores = 0, 0
                progress_win, progress_bar, progress_label = self._create_progress_window("Renombrando archivos...", total_resultados)
                for i, archivo in enumerate(self.resultados):
                    nombre_completo = os.path.basename(archivo)
                    progress_label.config(text=f"Renombrando: {nombre_completo}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    try:
                        carpeta, _ = os.path.split(archivo)
                        nombre_base, extension = os.path.splitext(nombre_completo)
                        nuevo_nombre = nombre_base

                        if opciones['renombrar_activo'].get() and opciones['renombrar_texto'].get():
                            nuevo_nombre = opciones['renombrar_texto'].get()
                        else:
                            if opciones['sustituir_activo'].get() and opciones['sustituir_buscar'].get():
                                nuevo_nombre = nuevo_nombre.replace(opciones['sustituir_buscar'].get(), opciones['sustituir_reemplazar'].get())
                            
                            if opciones['limpiar_id_activo'].get():
                                nuevo_nombre = re.sub(r'_ID\d+_A', '', nuevo_nombre, flags=re.IGNORECASE)
                                                        
                            if opciones['prefijo_activo'].get() and opciones['prefijo_texto'].get():
                                nuevo_nombre = opciones['prefijo_texto'].get() + nuevo_nombre
                            if opciones['sufijo_activo'].get() and opciones['sufijo_texto'].get():
                                nuevo_nombre = nuevo_nombre + opciones['sufijo_texto'].get()

                        if nuevo_nombre != nombre_base:
                            nueva_ruta = os.path.join(carpeta, nuevo_nombre + extension)
                            if os.path.exists(nueva_ruta):
                                self.log(f"⚠️ Error: ya existe un archivo con el nombre '{os.path.basename(nueva_ruta)}'")
                                errores += 1
                                continue
                            os.rename(archivo, nueva_ruta)
                            self.log(f"✅ Renombrado: {nombre_completo} → {os.path.basename(os.path.basename(nueva_ruta))}")
                            procesados += 1
                    except Exception as e:
                        self.log(f"❌ Error al renombrar {archivo}: {e}")
                        errores += 1
                self.log(f"--- Renombrado finalizado. Procesados: {procesados}, Errores: {errores} ---")
                if procesados > 0: self.buscar_archivos()


            elif accion == "editar texto":
                buscar = simpledialog.askstring('Buscar en contenido', 'Texto a buscar:', parent=self.root)
                if buscar is None: return
                reemplazar = simpledialog.askstring('Reemplazar con', 'Nuevo texto:', parent=self.root)
                if reemplazar is None: return
                
                progress_win, progress_bar, progress_label = self._create_progress_window("Editando contenido...", total_resultados)
                for i, archivo in enumerate(self.resultados):
                    nombre_archivo = os.path.basename(archivo)
                    progress_label.config(text=f"Analizando: {nombre_archivo}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    try:
                        ext = os.path.splitext(archivo)[1].lower()
                        if ext in ['.txt', '.json', '.xml', '.csv', '.html', '.md', '.log']:
                            with open(archivo, 'r', encoding='utf-8', errors='ignore') as f:
                                contenido = f.read()
                            if buscar in contenido:
                                contenido_nuevo = contenido.replace(buscar, reemplazar)
                                with open(archivo, 'w', encoding='utf-8') as f:
                                    f.write(contenido_nuevo)
                                self.log(f'Texto modificado en: {archivo}')
                        elif ext == '.docx':
                            doc = Document(archivo)
                            modificado = False
                            for p in doc.paragraphs:
                                if buscar in p.text:
                                    p.text = p.text.replace(buscar, reemplazar)
                                    modificado = True
                            if modificado:
                                doc.save(archivo)
                                self.log(f'Texto modificado en DOCX: {archivo}')
                    except Exception as e:
                        self.log(f'❌ Error al editar {archivo}: {e}')

            elif accion == "comprimir archivos en zip":
                nombre = simpledialog.askstring("Nombre ZIP", "Nombre del archivo ZIP (sin .zip):", parent=self.root)
                if not nombre: return
                carpeta_destino = filedialog.askdirectory(parent=self.root, title="Selecciona carpeta donde guardar el ZIP")
                if not carpeta_destino: return
                ruta_zip = os.path.join(carpeta_destino, nombre + ".zip")
                progress_win, progress_bar, progress_label = self._create_progress_window("Comprimiendo en ZIP...", total_resultados)
                with ZipFile(ruta_zip, "w", ZIP_DEFLATED) as zipf:
                    for i, archivo in enumerate(self.resultados):
                        nombre_archivo = os.path.basename(archivo)
                        progress_label.config(text=f"Añadiendo: {nombre_archivo}")
                        progress_bar['value'] = i + 1
                        self.root.update_idletasks()
                        if os.path.isdir(archivo):
                            for root_dir, _, files_in_dir in os.walk(archivo):
                                for f in files_in_dir:
                                    abs_path = os.path.join(root_dir, f)
                                    arcname = os.path.relpath(abs_path, os.path.dirname(archivo))
                                    zipf.write(abs_path, arcname)
                        else:
                            zipf.write(archivo, os.path.basename(archivo))
                self.log(f"✅ ZIP creado en: {ruta_zip}")

            elif accion == "comprimir carpetas individualmente":
                carpetas_a_comprimir = [r for r in self.resultados if os.path.isdir(r)]
                if not carpetas_a_comprimir:
                    messagebox.showinfo("Información", "La selección no contiene carpetas para comprimir.", parent=self.root)
                    return
                
                carpeta_destino = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta destino para los ZIPs")
                if not carpeta_destino: return
                progress_win, progress_bar, progress_label = self._create_progress_window("Comprimiendo carpetas...", len(carpetas_a_comprimir))
                for i, ruta in enumerate(carpetas_a_comprimir):
                    nombre_zip = os.path.basename(ruta.rstrip(os.sep))
                    progress_label.config(text=f"Comprimiendo: {nombre_zip}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    ruta_zip = os.path.join(carpeta_destino, nombre_zip + ".zip")
                    try:
                        with ZipFile(ruta_zip, "w", ZIP_DEFLATED) as zipf:
                            for root_dir, _, files_in_dir in os.walk(ruta):
                                for f in files_in_dir:
                                    abs_path = os.path.join(root_dir, f)
                                    arcname = os.path.relpath(abs_path, start=ruta)
                                    zipf.write(abs_path, arcname)
                        self.log(f"✅ Carpeta comprimida: {nombre_zip}.zip")
                    except Exception as e:
                        self.log(f"❌ Error al comprimir {ruta}: {e}")
        
        finally:
            if progress_win:
                progress_win.destroy()

    def limpiar(self):
        self.tree.delete(*self.tree.get_children())
        self.log_text.config(state="normal")
        self.log_text.delete("1.0", tk.END)
        self.log_text.config(state="disabled")
        self.source_path.set("")
        self.search_pattern.set("")
        self.resultados = []
        self.log("Campos y resultados limpiados.")

    def eliminar_resultados(self):
        if not self.resultados:
            self.log("No hay nada que eliminar.")
            return
        confirm = messagebox.askyesno("Confirmar eliminación", f"¿Seguro que quieres enviar {len(self.resultados)} elementos a la papelera?", parent=self.root)
        if not confirm: return

        eliminados, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Eliminando...", len(self.resultados))
        try:
            for i, ruta in enumerate(self.resultados):
                nombre_archivo = os.path.basename(ruta)
                progress_label.config(text=f"Eliminando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                try:
                    ruta_normalizada = os.path.normpath(ruta)
                    send2trash(ruta_normalizada)
                    eliminados += 1
                except Exception as e:
                    self.log(f"❌ Error al eliminar {ruta}: {e}")
                    errores += 1
        finally:
            progress_win.destroy()
        
        self.log(f"🗑️ Enviados a la papelera: {eliminados}. Errores: {errores}.")
        if eliminados > 0: self.buscar_archivos()

    def _convertir_masivo(self, extension_filtro, output_ext, conversion_function, title="Convirtiendo archivos..."):
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Busca archivos primero para usar la conversión masiva.", parent=self.root)
            return
        
        archivos_filtrados = [ruta for ruta in self.resultados if any(ruta.lower().endswith(ext) for ext in (extension_filtro if isinstance(extension_filtro, (tuple, list)) else (extension_filtro,)))]
        if not archivos_filtrados:
             messagebox.showinfo("Información", f"No se encontraron archivos con la extensión '{extension_filtro}' en los resultados.", parent=self.root)
             return
        
        progress_win, progress_bar, progress_label = self._create_progress_window(title, len(archivos_filtrados))
        try:
            for i, ruta in enumerate(archivos_filtrados):
                nombre_base = os.path.splitext(os.path.basename(ruta))[0]
                salida = os.path.join(os.path.dirname(ruta), f"{nombre_base}.{output_ext}")

                progress_label.config(text=f"Convirtiendo: {os.path.basename(ruta)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                conversion_function(ruta, salida)
        finally:
            progress_win.destroy()
            self.log(f"--- Conversión masiva finalizada. ---")

    def importar_y_aplicar_renombrado(self):
        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar Excel editado", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            df = pd.read_excel(archivo_excel)
            if not all(col in df.columns for col in ["Ruta actual", "Nuevo nombre"]):
                self.log("❌ El Excel debe contener las columnas 'Ruta actual' y 'Nuevo nombre'.")
                return
            
            cambios, errores = 0, 0
            progress_win, progress_bar, progress_label = self._create_progress_window("Renombrando desde Excel...", len(df))
            try:
                for index, fila in df.iterrows():
                    ruta_actual = fila["Ruta actual"]
                    nuevo_nombre = fila["Nuevo nombre"]

                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()

                    if pd.isna(nuevo_nombre) or pd.isna(ruta_actual) or os.path.basename(ruta_actual) == nuevo_nombre: continue
                    
                    progress_label.config(text=f"Procesando: {os.path.basename(str(ruta_actual))}")
                    nueva_ruta = os.path.join(os.path.dirname(ruta_actual), str(nuevo_nombre))
                    try:
                        if os.path.exists(ruta_actual):
                            if os.path.exists(nueva_ruta):
                                self.log(f"⚠️ Error: ya existe un archivo con el nombre '{nuevo_nombre}'")
                                errores += 1
                                continue
                            os.rename(ruta_actual, nueva_ruta)
                            self.log(f"✅ Renombrado: {os.path.basename(ruta_actual)} → {nuevo_nombre}")
                            cambios += 1
                        else:
                             self.log(f"⚠️ No se encontró: {ruta_actual}")
                    except Exception as e:
                        self.log(f"❌ Error al renombrar {ruta_actual}: {e}")
                        errores += 1
            finally:
                progress_win.destroy()

            self.log(f"Renombrado desde Excel finalizado. Cambios: {cambios}, Errores: {errores}")
            if cambios > 0: self.buscar_archivos()
        except Exception as e:
            self.log(f"Error al procesar el Excel: {e}")

    def _convertir_archivos(self, title, filetypes, output_ext, conversion_function):
        archivos = filedialog.askopenfilenames(parent=self.root, title=title, filetypes=filetypes)
        if not archivos: return
        
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Convirtiendo a .{output_ext}", len(archivos))
        try:
            for i, archivo in enumerate(archivos):
                nombre_base = os.path.splitext(os.path.basename(archivo))[0]
                salida = os.path.join(os.path.dirname(archivo), f"{nombre_base}.{output_ext}")
                
                progress_label.config(text=f"Convirtiendo: {os.path.basename(archivo)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                conversion_function(archivo, salida)
        finally:
            progress_win.destroy()
            self.log("--- Conversión finalizada. ---")
    
    def accion_txt_a_json_masivo(self): self._convertir_masivo(".txt", "json", self._txt_a_json_worker, "Renombrando TXT a JSON...")

    def accion_pdf_a_docx_masivo(self): self._convertir_masivo(".pdf", "docx", self._pdf_a_docx_worker, "Convirtiendo PDF a DOCX...")
    def accion_jpg_a_pdf_masivo(self): self._convertir_masivo((".jpg",".jpeg"), "pdf", self._jpg_a_pdf_worker, "Convirtiendo JPG a PDF...")
    def accion_docx_a_pdf_masivo(self): self._convertir_masivo(".docx", "pdf", self._docx_a_pdf_worker, "Convirtiendo DOCX a PDF...")
    def accion_pdf_a_jpg_masivo(self): self._convertir_masivo(".pdf", "jpg", self._pdf_a_jpg_worker, "Convirtiendo PDF a JPG...")
    def accion_png_a_jpg_masivo(self): self._convertir_masivo(".png", "jpg", self._png_a_jpg_worker, "Convirtiendo PNG a JPG...")


    def accion_pdf_a_docx(self): self._convertir_archivos("Seleccionar PDF(s)", [("PDF", "*.pdf")], "docx", self._pdf_a_docx_worker)
    def accion_jpg_a_pdf(self): self._convertir_archivos("Seleccionar JPG(s)", [("JPG", "*.jpg;*.jpeg")], "pdf", self._jpg_a_pdf_worker)
    def accion_docx_a_pdf(self): self._convertir_archivos("Seleccionar DOCX", [("Word", "*.docx")], "pdf", self._docx_a_pdf_worker)
    def accion_pdf_a_jpg(self): self._convertir_archivos("Seleccionar PDF(s)", [("PDF", "*.pdf")], "jpg", self._pdf_a_jpg_worker)
    def accion_png_a_jpg(self): self._convertir_archivos("Seleccionar PNG(s)", [("PNG", "*.png")], "jpg", self._png_a_jpg_worker)
    def accion_convertir_json_rips_a_xlsx(self): self._convertir_archivos("Seleccionar JSON RIPS", [("JSON", "*.json")], "xlsx", self._json_rips_a_xlsx_worker)
    def accion_convertir_xlsx_rips_a_json(self): self._convertir_archivos("Seleccionar Excel RIPS", [("Excel", "*.xlsx")], "json", self._xlsx_rips_a_json_worker)

    def accion_json_evento_a_xlsx_masivo(self):
        carpeta_origen = filedialog.askdirectory(title="Seleccionar carpeta con JSONs de eventos")
        if not carpeta_origen: return
        
        archivo_salida = filedialog.asksaveasfilename(
            title="Guardar Excel consolidado como...",
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")]
        )
        if not archivo_salida: return

        # Buscar JSONs recursivamente
        archivos_json = []
        for root, dirs, files in os.walk(carpeta_origen):
            for file in files:
                if file.lower().endswith(".json"):
                    archivos_json.append(os.path.join(root, file))
        
        if not archivos_json:
            self.log("❌ No se encontraron archivos JSON en la carpeta seleccionada.")
            return
            
        self.log(f"--- Iniciando consolidación de {len(archivos_json)} JSONs a Excel ---")
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Consolidando JSONs...", len(archivos_json))
        
        try:
            # Listas globales para acumular todos los registros de todos los archivos
            todas_consultas = []
            todos_procedimientos = []
            todos_otros_servicios = []
            
            errores = 0
            
            for i, ruta_json in enumerate(archivos_json):
                nombre_archivo = os.path.basename(ruta_json)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                try:
                    with open(ruta_json, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    # Extraer datos usando la estructura RIPS (Usuarios -> Servicios)
                    usuarios_lista = data.get("usuarios", []) if isinstance(data, dict) else []
                    
                    for usuario in usuarios_lista:
                        base_info = {
                            "archivo_origen": nombre_archivo, # Columna extra para identificar origen
                            "tipo_documento_usuario": usuario.get("tipoDocumentoIdentificacion"),
                            "documento_usuario": usuario.get("numDocumentoIdentificacion"),
                            "tipo_usuario": usuario.get("tipoUsuario"),
                            "fecha_nacimiento": usuario.get("fechaNacimiento"),
                            "sexo": usuario.get("codSexo"),
                            "pais_residencia": usuario.get("codPaisResidencia"),
                            "municipio_residencia": usuario.get("codMunicipioResidencia"),
                            "zona_residencia": usuario.get("codZonaTerritorialResidencia"),
                            "incapacidad": usuario.get("incapacidad"),
                            "consecutivo_usuario": usuario.get("consecutivo"),
                            "pais_origen": usuario.get("codPaisOrigen")
                        }

                        servicios = usuario.get("servicios", {})

                        for consulta in servicios.get("consultas", []):
                            todas_consultas.append({**base_info, **consulta})

                        for procedimiento in servicios.get("procedimientos", []):
                            todos_procedimientos.append({**base_info, **procedimiento})

                        for otro in servicios.get("otrosServicios", []):
                            todos_otros_servicios.append({**base_info, **otro})
                        
                except Exception as e:
                    self.log(f"❌ Error leyendo {nombre_archivo}: {e}")
                    errores += 1
            
            # Crear DataFrames y guardar
            if todas_consultas or todos_procedimientos or todos_otros_servicios:
                with pd.ExcelWriter(archivo_salida, engine="openpyxl") as writer:
                    if todas_consultas:
                        pd.DataFrame(todas_consultas).to_excel(writer, sheet_name="Consultas", index=False)
                    if todos_procedimientos:
                        pd.DataFrame(todos_procedimientos).to_excel(writer, sheet_name="Procedimientos", index=False)
                    if todos_otros_servicios:
                        pd.DataFrame(todos_otros_servicios).to_excel(writer, sheet_name="OtrosServicios", index=False)
                    
                    if not (todas_consultas or todos_procedimientos or todos_otros_servicios):
                         pd.DataFrame().to_excel(writer, sheet_name="Vacio", index=False)

                self.log(f"✅ Excel consolidado RIPS creado: {archivo_salida}")
                total_reg = len(todas_consultas) + len(todos_procedimientos) + len(todos_otros_servicios)
                messagebox.showinfo("Éxito", f"Consolidación completada.\nTotal registros: {total_reg}\nErrores de lectura: {errores}")
            else:
                self.log("⚠️ No se encontraron datos RIPS válidos para exportar.")
                messagebox.showwarning("Advertencia", "No se generaron datos. Verifica que los JSON tengan la estructura correcta (usuarios -> servicios).")
                
        except Exception as e:
             self.log(f"❌ Error general: {e}")
             messagebox.showerror("Error", str(e))
        finally:
            progress_win.destroy()

    def accion_xlsx_evento_a_json_masivo(self):
        archivo_excel = filedialog.askopenfilename(
            title="Seleccionar Excel consolidado RIPS (Eventos)",
            filetypes=[("Excel", "*.xlsx")]
        )
        if not archivo_excel: return
        
        carpeta_destino = filedialog.askdirectory(title="Seleccionar carpeta de destino para los JSONs")
        if not carpeta_destino: return
        
        self.log(f"--- Iniciando desconsolidación de Excel a JSONs ---")
        
        try:
            xls = pd.ExcelFile(archivo_excel)
            
            # Leer todas las hojas disponibles
            df_consultas = pd.DataFrame()
            df_procedimientos = pd.DataFrame()
            df_otros = pd.DataFrame()
            
            if "Consultas" in xls.sheet_names:
                df_consultas = pd.read_excel(xls, sheet_name="Consultas")
            if "Procedimientos" in xls.sheet_names:
                df_procedimientos = pd.read_excel(xls, sheet_name="Procedimientos")
            if "OtrosServicios" in xls.sheet_names:
                df_otros = pd.read_excel(xls, sheet_name="OtrosServicios")
                
            # Convertir NaN a None en todos los DFs (usando astype(object) para asegurar que acepte None)
            df_consultas = df_consultas.astype(object).where(pd.notnull(df_consultas), None)
            df_procedimientos = df_procedimientos.astype(object).where(pd.notnull(df_procedimientos), None)
            df_otros = df_otros.astype(object).where(pd.notnull(df_otros), None)

            # Obtener lista de archivos únicos
            archivos_unicos = set()
            if "archivo_origen" in df_consultas.columns:
                archivos_unicos.update(df_consultas["archivo_origen"].dropna().unique())
            if "archivo_origen" in df_procedimientos.columns:
                archivos_unicos.update(df_procedimientos["archivo_origen"].dropna().unique())
            if "archivo_origen" in df_otros.columns:
                archivos_unicos.update(df_otros["archivo_origen"].dropna().unique())
            
            if not archivos_unicos:
                self.log("❌ No se encontró la columna 'archivo_origen' o está vacía. Verifica que sea un Excel generado por la opción masiva.")
                return

            progress_win, progress_bar, progress_label = self._create_progress_window("Generando JSONs...", len(archivos_unicos))
            
            errores = 0
            generados = 0
            
            for i, nombre_archivo in enumerate(archivos_unicos):
                progress_label.config(text=f"Generando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                try:
                    usuarios_dict = {}
                    
                    # Función para procesar un DF filtrado
                    def procesar_df(df_origen, clave_servicio):
                        if df_origen.empty or "archivo_origen" not in df_origen.columns:
                            return
                            
                        # Filtrar por el archivo actual
                        df_filtrado = df_origen[df_origen["archivo_origen"] == nombre_archivo]
                        
                        for _, row in df_filtrado.iterrows():
                            # Clave de usuario
                            td = str(row.get("tipo_documento_usuario", ""))
                            doc = str(row.get("documento_usuario", ""))
                            user_key = (td, doc)
                            
                            if user_key not in usuarios_dict:
                                usuarios_dict[user_key] = {
                                    "tipoDocumentoIdentificacion": row.get("tipo_documento_usuario"),
                                    "numDocumentoIdentificacion": row.get("documento_usuario"),
                                    "tipoUsuario": row.get("tipo_usuario"),
                                    "fechaNacimiento": row.get("fecha_nacimiento"), 
                                    "codSexo": row.get("sexo"),
                                    "codPaisResidencia": row.get("pais_residencia"),
                                    "codMunicipioResidencia": row.get("municipio_residencia"),
                                    "codZonaTerritorialResidencia": row.get("zona_residencia"),
                                    "incapacidad": row.get("incapacidad"),
                                    "consecutivo": row.get("consecutivo_usuario"),
                                    "codPaisOrigen": row.get("pais_origen"),
                                    "servicios": {
                                        "consultas": [],
                                        "procedimientos": [],
                                        "otrosServicios": []
                                    }
                                }
                            
                            # Datos del servicio
                            servicio_data = row.to_dict()
                            # Eliminar columnas base y archivo_origen
                            keys_to_remove = [
                                "tipo_documento_usuario", "documento_usuario", "tipo_usuario", 
                                "fecha_nacimiento", "sexo", "pais_residencia", "municipio_residencia", 
                                "zona_residencia", "incapacidad", "consecutivo_usuario", "pais_origen",
                                "archivo_origen"
                            ]
                            for k in keys_to_remove:
                                servicio_data.pop(k, None)
                                
                            # Agregar si tiene datos válidos
                            if any(v is not None for v in servicio_data.values()):
                                usuarios_dict[user_key]["servicios"][clave_servicio].append(servicio_data)

                    # Procesar los 3 DFs
                    procesar_df(df_consultas, "consultas")
                    procesar_df(df_procedimientos, "procedimientos")
                    procesar_df(df_otros, "otrosServicios")
                    
                    # Guardar JSON
                    resultado_final = { "usuarios": list(usuarios_dict.values()) }
                    ruta_salida = os.path.join(carpeta_destino, nombre_archivo)
                    
                    if not ruta_salida.lower().endswith(".json"):
                        ruta_salida += ".json"

                    with open(ruta_salida, 'w', encoding='utf-8') as f:
                        json.dump(resultado_final, f, ensure_ascii=False, indent=4)
                    
                    generados += 1
                    
                except Exception as e:
                    self.log(f"❌ Error generando {nombre_archivo}: {e}")
                    errores += 1
            
            progress_win.destroy()
            messagebox.showinfo("Proceso Completado", f"Se generaron {generados} archivos JSON.\nErrores: {errores}")
            
        except Exception as e:
            self.log(f"❌ Error leyendo Excel: {e}")
            messagebox.showerror("Error", str(e))



    def _txt_a_json_worker(self, entrada, salida):
        """Worker para renombrar un archivo de .txt a .json."""
        try:
            if os.path.exists(salida):
                self.log(f"⚠️ Error: El archivo '{os.path.basename(salida)}' ya existe. Se omite el renombrado.")
                return
            os.rename(entrada, salida)
            self.log(f"✅ Renombrado: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error renombrando {os.path.basename(entrada)}: {e}")

    def _json_rips_a_xlsx_worker(self, entrada, salida):
        try:
            with open(entrada, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Inicializar listas
            consultas = []
            procedimientos = []
            otros_servicios = []

            # Recorrer los usuarios
            for usuario in data.get("usuarios", []):
                # Separar el documento del usuario con nombre único para evitar conflictos
                base_info = {
                    "tipo_documento_usuario": usuario.get("tipoDocumentoIdentificacion"),
                    "documento_usuario": usuario.get("numDocumentoIdentificacion"),
                    "tipo_usuario": usuario.get("tipoUsuario"),
                    "fecha_nacimiento": usuario.get("fechaNacimiento"),
                    "sexo": usuario.get("codSexo"),
                    "pais_residencia": usuario.get("codPaisResidencia"),
                    "municipio_residencia": usuario.get("codMunicipioResidencia"),
                    "zona_residencia": usuario.get("codZonaTerritorialResidencia"),
                    "incapacidad": usuario.get("incapacidad"),
                    "consecutivo_usuario": usuario.get("consecutivo"),
                    "pais_origen": usuario.get("codPaisOrigen")
                }

                servicios = usuario.get("servicios", {})

                for consulta in servicios.get("consultas", []):
                    consultas.append({**base_info, **consulta})

                for procedimiento in servicios.get("procedimientos", []):
                    procedimientos.append({**base_info, **procedimiento})

                for otro in servicios.get("otrosServicios", []):
                    otros_servicios.append({**base_info, **otro})

            # Crear DataFrames
            df_consultas = pd.DataFrame(consultas)
            df_procedimientos = pd.DataFrame(procedimientos)
            df_otros = pd.DataFrame(otros_servicios)

            # Guardar a Excel con varias hojas
            with pd.ExcelWriter(salida, engine="openpyxl") as writer:
                if not df_consultas.empty:
                    df_consultas.to_excel(writer, sheet_name="Consultas", index=False)
                if not df_procedimientos.empty:
                    df_procedimientos.to_excel(writer, sheet_name="Procedimientos", index=False)
                if not df_otros.empty:
                    df_otros.to_excel(writer, sheet_name="OtrosServicios", index=False)
                
                # Si todos están vacíos, crear una hoja vacía
                if df_consultas.empty and df_procedimientos.empty and df_otros.empty:
                     pd.DataFrame().to_excel(writer, sheet_name="Vacio", index=False)

            self.log(f"✅ Convertido RIPS: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo RIPS {os.path.basename(entrada)}: {e}")

    def _xlsx_rips_a_json_worker(self, entrada, salida):
        try:
            # Leer el archivo Excel
            xls = pd.ExcelFile(entrada)
            
            usuarios_dict = {}
            
            # Función auxiliar para procesar cada hoja
            def procesar_hoja(nombre_hoja, clave_servicio):
                if nombre_hoja in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=nombre_hoja)
                    # Convertir NaN a None (null en JSON) asegurando tipo object
                    df = df.astype(object).where(pd.notnull(df), None)
                    
                    for _, row in df.iterrows():
                        # Clave única para el usuario (usando tipo y número de documento)
                        # Asegurar que sean strings para la clave, manejando posibles None
                        td = str(row.get("tipo_documento_usuario", ""))
                        doc = str(row.get("documento_usuario", ""))
                        user_key = (td, doc)
                        
                        if user_key not in usuarios_dict:
                            # Crear estructura base del usuario si no existe
                            usuarios_dict[user_key] = {
                                "tipoDocumentoIdentificacion": row.get("tipo_documento_usuario"),
                                "numDocumentoIdentificacion": row.get("documento_usuario"),
                                "tipoUsuario": row.get("tipo_usuario"),
                                "fechaNacimiento": row.get("fecha_nacimiento"), 
                                "codSexo": row.get("sexo"),
                                "codPaisResidencia": row.get("pais_residencia"),
                                "codMunicipioResidencia": row.get("municipio_residencia"),
                                "codZonaTerritorialResidencia": row.get("zona_residencia"),
                                "incapacidad": row.get("incapacidad"),
                                "consecutivo": row.get("consecutivo_usuario"),
                                "codPaisOrigen": row.get("pais_origen"),
                                "servicios": {
                                    "consultas": [],
                                    "procedimientos": [],
                                    "otrosServicios": []
                                }
                            }
                        
                        # Extraer datos del servicio (eliminando info del usuario para no duplicar)
                        servicio_data = row.to_dict()
                        # Eliminar claves de usuario base
                        keys_to_remove = [
                            "tipo_documento_usuario", "documento_usuario", "tipo_usuario", 
                            "fecha_nacimiento", "sexo", "pais_residencia", "municipio_residencia", 
                            "zona_residencia", "incapacidad", "consecutivo_usuario", "pais_origen"
                        ]
                        for k in keys_to_remove:
                            servicio_data.pop(k, None)
                            
                        # Agregar a la lista correspondiente si no es una fila vacía (chequeo básico)
                        # A veces pandas lee filas vacías si el excel está sucio
                        if any(servicio_data.values()):
                            usuarios_dict[user_key]["servicios"][clave_servicio].append(servicio_data)

            # Procesar las hojas
            procesar_hoja("Consultas", "consultas")
            procesar_hoja("Procedimientos", "procedimientos")
            procesar_hoja("OtrosServicios", "otrosServicios")
            
            # Construir el JSON final
            resultado_final = {
                "usuarios": list(usuarios_dict.values())
            }
            
            # Guardar JSON
            with open(salida, 'w', encoding='utf-8') as f:
                json.dump(resultado_final, f, ensure_ascii=False, indent=4)
                
            self.log(f"✅ Convertido RIPS Inverso: {os.path.basename(entrada)} → {os.path.basename(salida)}")
            
        except Exception as e:
            self.log(f"❌ Error convirtiendo Excel a RIPS JSON {os.path.basename(entrada)}: {e}")



    def _pdf_a_docx_worker(self, entrada, salida):
        try:
            cv = Converter(entrada)
            cv.convert(salida, start=0, end=None)
            cv.close()
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")

    def _jpg_a_pdf_worker(self, entrada, salida):
        try:
            img = Image.open(entrada)
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            img.save(salida, "PDF", resolution=100.0)
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")
            
    def _docx_a_pdf_worker(self, entrada, salida):
        """
        Convierte DOCX a PDF usando automatización COM directa.
        Esta es la versión robusta que funciona correctamente en el .exe compilado.
        """
        word = None
        doc = None
        try:
            in_file_abs = os.path.abspath(entrada)
            out_file_abs = os.path.abspath(salida)

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(in_file_abs)
            
            wdFormatPDF = 17
            doc.SaveAs(out_file_abs, FileFormat=wdFormatPDF)
            
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")

        except Exception as e:
            error_msg = str(e)
            if "com_error" in error_msg:
                 error_msg += ("\n (Asegúrese de que Microsoft Word esté instalado y no tenga diálogos abiertos que bloqueen la automatización).")
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {error_msg}")
        
        finally:
            if doc:
                doc.Close(False)
            if word:
                word.Quit()
            
    def _pdf_a_jpg_worker(self, entrada, salida):
        try:
            doc = fitz.open(entrada)
            for i, page in enumerate(doc):
                pix = page.get_pixmap()
                if len(doc) == 1:
                    nombre_salida = salida
                else:
                    base, ext = os.path.splitext(salida)
                    nombre_salida = f"{base}_p{i+1}{ext}"
                pix.save(nombre_salida)
            doc.close()
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)} (y páginas adicionales si aplica)")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")
        
    def _png_a_jpg_worker(self, entrada, salida):
        try:
            img = Image.open(entrada).convert("RGB")
            img.save(salida, "jpeg")
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")
        
    def accion_txt_a_json(self):
        archivos = filedialog.askopenfilenames(parent=self.root, title="Seleccionar archivos TXT para renombrar a JSON", filetypes=[("TXT", "*.txt")])
        if not archivos: return
        for archivo in archivos:
            try:
                base, _ = os.path.splitext(archivo)
                os.rename(archivo, base + ".json")
                self.log(f"Renombrado: {os.path.basename(archivo)} → {os.path.basename(base)}.json")
            except Exception as e:
                self.log(f"Error renombrando {os.path.basename(archivo)}: {e}")

    def accion_unificar_por_carpeta(self):
        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta principal con subcarpetas")
        if not carpeta_base: return
        
        nombre_final_base = simpledialog.askstring(
            "Nombre del archivo final", 
            "Nombre para los PDF combinados (sin .pdf).\nEste nombre se usará para el archivo resultante en CADA subcarpeta.", 
            parent=self.root
        )
        if not nombre_final_base: return
        
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
        if not subcarpetas:
            self.log("No se encontraron subcarpetas.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Unificando PDFs por carpeta...", len(subcarpetas))
        try:
            for i, carpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(carpeta)
                progress_label.config(text=f"Procesando: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                archivos_pdf_a_procesar = []
                for num_pdf in range(1, 11):
                    nombre_archivo_buscado = f"{num_pdf}.pdf"
                    ruta_archivo_buscado = os.path.join(carpeta, nombre_archivo_buscado)
                    if os.path.exists(ruta_archivo_buscado):
                        archivos_pdf_a_procesar.append(ruta_archivo_buscado)

                if not archivos_pdf_a_procesar:
                    self.log(f"No se encontraron PDFs numerados (1-10) en '{nombre_subcarpeta}', saltando.")
                    continue
                
                nombre_final = f"{nombre_final_base}.pdf"
                ruta_salida = os.path.join(carpeta, nombre_final)
                
                doc_final = fitz.open()
                
                for ruta_pdf in archivos_pdf_a_procesar:
                    try:
                        with fitz.open(ruta_pdf) as doc_origen:
                            for page_origen in doc_origen:
                                pix = page_origen.get_pixmap(dpi=300, colorspace=fitz.csGRAY)
                                pagina_nueva = doc_final.new_page(width=pix.width, height=pix.height)
                                pagina_nueva.insert_image(pagina_nueva.rect, pixmap=pix)
                    except Exception as e:
                        self.log(f"Error procesando '{os.path.basename(ruta_pdf)}' en '{nombre_subcarpeta}': {e}")
                
                if len(doc_final) > 0:
                    doc_final.save(ruta_salida, garbage=4, deflate=True)
                    self.log(f"✅ PDF unificado '{nombre_final}' creado en la carpeta '{nombre_subcarpeta}'.")
                doc_final.close()
        finally:
            if progress_win:
                progress_win.destroy()
    
    def _unificar_imagenes_por_carpeta(self, tipo_imagen):
        """
        Función genérica para unificar imágenes (JPG o PNG) de subcarpetas en archivos PDF.
        Un PDF por cada subcarpeta procesada.
        """
        # Mapeos para configurar la función según el tipo de imagen
        title_map = {
            "JPG": "Seleccionar carpeta principal con subcarpetas de JPGs",
            "PNG": "Seleccionar carpeta principal con subcarpetas de PNGs"
        }
        ext_map = {
            "JPG": ['.jpg', '.jpeg'],
            "PNG": ['.png']
        }

        base_path = filedialog.askdirectory(
            parent=self.root,
            title=title_map.get(tipo_imagen, "Seleccionar carpeta principal")
        )
        if not base_path:
            self.log(f"Operación de unificación de {tipo_imagen} cancelada.")
            return

        try:
            subfolders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not subfolders:
            self.log("No se encontraron subcarpetas para procesar.")
            messagebox.showinfo("Sin Carpetas", "No se encontraron subcarpetas en la ruta seleccionada.", parent=self.root)
            return

        self.log(f"--- Iniciando unificación de {tipo_imagen} para {len(subfolders)} carpetas ---")
        
        pdfs_creados, errores_conversion = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Unificando {tipo_imagen}...", len(subfolders))

        try:
            for i, folder_name in enumerate(subfolders):
                folder_path = os.path.join(base_path, folder_name)
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                self.log(f"-> Procesando carpeta: '{folder_name}'")
                
                image_files = []
                for file_name in os.listdir(folder_path):
                    if any(file_name.lower().endswith(ext) for ext in ext_map[tipo_imagen]):
                        try:
                            # Extraer número del nombre del archivo para ordenar
                            num = int(os.path.splitext(file_name)[0])
                            image_files.append((num, os.path.join(folder_path, file_name)))
                        except (ValueError, IndexError):
                            self.log(f"  -> Omitido: El archivo '{file_name}' no sigue el formato numérico esperado.")
                            continue
                
                if not image_files:
                    self.log(f"  -> No se encontraron archivos {tipo_imagen} con nombre numérico en '{folder_name}'.")
                    continue

                # Ordenar imágenes por el número extraído
                image_files.sort()

                # Convertir a PDF
                pdf_path = os.path.join(base_path, f"{folder_name}.pdf")
                
                try:
                    images_to_convert = []
                    first_image_path = image_files[0][1]
                    
                    # Abrir la primera imagen para usarla como base
                    with Image.open(first_image_path) as img:
                        # Si es RGBA (como algunos PNGs), convertir a RGB
                        if img.mode == 'RGBA':
                            img = img.convert('RGB')
                        
                        # Abrir el resto de las imágenes y asegurarse de que estén en modo RGB
                        for _, img_path in image_files[1:]:
                            with Image.open(img_path) as other_img:
                                if other_img.mode == 'RGBA':
                                    images_to_convert.append(other_img.convert('RGB'))
                                else:
                                    images_to_convert.append(other_img.copy())
                        
                        # Guardar el PDF
                        img.save(pdf_path, "PDF", resolution=100.0, save_all=True, append_images=images_to_convert)

                    self.log(f"  ✅ PDF creado: '{os.path.basename(pdf_path)}'")
                    pdfs_creados += 1

                except Exception as e:
                    self.log(f"  ❌ Error al convertir imágenes en '{folder_name}': {e}")
                    errores_conversion += 1
        
        finally:
            if progress_win:
                progress_win.destroy()

        resumen = (f"Proceso de unificación de {tipo_imagen} finalizado.\n\n"
                   f"- PDFs creados: {pdfs_creados}\n"
                   f"- Carpetas con errores: {errores_conversion}")
        self.log(f"--- {resumen.replace(chr(10)*2, ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)

    def accion_unificar_jpg_por_carpeta(self):
        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta principal con subcarpetas de JPGs")
        if not carpeta_base: return
        
        nombre_final_base = simpledialog.askstring(
            "Nombre del archivo final",
            "Nombre para el PDF combinado (sin .pdf).\nEste nombre se usará para el archivo resultante en CADA subcarpeta.",
            parent=self.root
        )
        if not nombre_final_base: return
        
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
        if not subcarpetas:
            self.log("No se encontraron subcarpetas.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Unificando JPGs por carpeta...", len(subcarpetas))
        try:
            for i, carpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(carpeta)
                progress_label.config(text=f"Procesando: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                archivos_img_a_procesar = []
                for num_img in range(1, 11):
                    ruta_encontrada = None
                    for ext in ['.jpg', '.jpeg']:
                        nombre_archivo = f"{num_img}{ext}"
                        ruta_archivo = os.path.join(carpeta, nombre_archivo)
                        if os.path.exists(ruta_archivo):
                            ruta_encontrada = ruta_archivo
                            break
                    if ruta_encontrada:
                        archivos_img_a_procesar.append(ruta_encontrada)

                if not archivos_img_a_procesar:
                    self.log(f"No se encontraron JPGs/JPEGs numerados (1-10) en '{nombre_subcarpeta}', saltando.")
                    continue

                lista_imagenes_procesadas = []
                for ruta_img in archivos_img_a_procesar:
                    try:
                        img = Image.open(ruta_img)
                        img_gray = img.convert('L') 
                        lista_imagenes_procesadas.append(img_gray)
                    except Exception as e:
                        self.log(f"Error abriendo imagen '{os.path.basename(ruta_img)}': {e}")
                
                if lista_imagenes_procesadas:
                    nombre_pdf = f"{nombre_final_base}.pdf"
                    ruta_salida = os.path.join(carpeta, nombre_pdf)
                    
                    lista_imagenes_procesadas[0].save(
                        ruta_salida, 
                        save_all=True, 
                        append_images=lista_imagenes_procesadas[1:], 
                        resolution=300.0
                    )
                    self.log(f"✅ PDF '{nombre_pdf}' creado a partir de JPGs en '{nombre_subcarpeta}'.")
        finally:
            if progress_win:
                progress_win.destroy()

    def accion_unificar_png_por_carpeta(self):
        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta principal con subcarpetas de PNGs")
        if not carpeta_base: return
        
        nombre_final_base = simpledialog.askstring(
            "Nombre del archivo final",
            "Nombre para el PDF combinado (sin .pdf).\nEste nombre se usará para el archivo resultante en CADA subcarpeta.",
            parent=self.root
        )
        if not nombre_final_base: return
        
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
        if not subcarpetas:
            self.log("No se encontraron subcarpetas.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Unificando PNGs por carpeta...", len(subcarpetas))
        try:
            for i, carpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(carpeta)
                progress_label.config(text=f"Procesando: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                archivos_img_a_procesar = []
                for num_img in range(1, 11):
                    nombre_archivo = f"{num_img}.png"
                    ruta_archivo = os.path.join(carpeta, nombre_archivo)
                    if os.path.exists(ruta_archivo):
                        archivos_img_a_procesar.append(ruta_archivo)

                if not archivos_img_a_procesar:
                    self.log(f"No se encontraron PNGs numerados (1-10) en '{nombre_subcarpeta}', saltando.")
                    continue

                lista_imagenes_procesadas = []
                for ruta_img in archivos_img_a_procesar:
                    try:
                        img = Image.open(ruta_img)
                        # Convertir a RGB si es necesario (para PNG con transparencia)
                        if img.mode in ('RGBA', 'LA'):
                            background = Image.new('RGB', img.size, (255, 255, 255))
                            if img.mode == 'RGBA':
                                background.paste(img, mask=img.split()[3])
                            else:
                                background.paste(img, mask=img.split()[1])
                            img = background
                        elif img.mode != 'RGB':
                            img = img.convert('RGB')
                        lista_imagenes_procesadas.append(img)
                    except Exception as e:
                        self.log(f"Error abriendo imagen '{os.path.basename(ruta_img)}': {e}")
                
                if lista_imagenes_procesadas:
                    nombre_pdf = f"{nombre_final_base}.pdf"
                    ruta_salida = os.path.join(carpeta, nombre_pdf)
                    
                    lista_imagenes_procesadas[0].save(
                        ruta_salida, 
                        save_all=True, 
                        append_images=lista_imagenes_procesadas[1:], 
                        resolution=300.0
                    )
                    self.log(f"✅ PDF '{nombre_pdf}' creado a partir de PNGs en '{nombre_subcarpeta}'.")
        finally:
            if progress_win:
                progress_win.destroy()

    def accion_unificar_docx_por_carpeta(self):
        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta principal con subcarpetas de DOCX")
        if not carpeta_base: return
        
        nombre_final_base = simpledialog.askstring(
            "Nombre del archivo final",
            "Nombre para el PDF combinado (sin .pdf).\nEste nombre se usará para el archivo resultante en CADA subcarpeta.",
            parent=self.root
        )
        if not nombre_final_base: return
        
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
        if not subcarpetas:
            self.log("No se encontraron subcarpetas.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Unificando DOCX por carpeta...", len(subcarpetas))
        try:
            for i, carpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(carpeta)
                progress_label.config(text=f"Procesando: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                archivos_docx_a_procesar = []
                for num_doc in range(1, 11):
                    nombre_archivo = f"{num_doc}.docx"
                    ruta_archivo = os.path.join(carpeta, nombre_archivo)
                    if os.path.exists(ruta_archivo):
                        archivos_docx_a_procesar.append(ruta_archivo)

                if not archivos_docx_a_procesar:
                    self.log(f"No se encontraron DOCX numerados (1-10) en '{nombre_subcarpeta}', saltando.")
                    continue

                # Convertir cada DOCX a PDF temporal y unificar
                pdfs_temporales = []
                try:
                    for ruta_docx in archivos_docx_a_procesar:
                        nombre_temp_pdf = os.path.splitext(os.path.basename(ruta_docx))[0] + "_temp.pdf"
                        ruta_temp_pdf = os.path.join(carpeta, nombre_temp_pdf)
                        
                        try:
                            # Usar docx2pdf para convertir
                            convert(ruta_docx, ruta_temp_pdf)
                            if os.path.exists(ruta_temp_pdf):
                                pdfs_temporales.append(ruta_temp_pdf)
                            else:
                                self.log(f"❌ Falló conversión de {os.path.basename(ruta_docx)}")
                        except Exception as e:
                            self.log(f"❌ Error al convertir {os.path.basename(ruta_docx)}: {e}")

                    if pdfs_temporales:
                        nombre_pdf_final = f"{nombre_final_base}.pdf"
                        ruta_salida = os.path.join(carpeta, nombre_pdf_final)
                        
                        doc_final = fitz.open()
                        for pdf_temp in pdfs_temporales:
                            try:
                                with fitz.open(pdf_temp) as doc_temp:
                                    doc_final.insert_pdf(doc_temp)
                            except Exception as e:
                                self.log(f"Error al unir {os.path.basename(pdf_temp)}: {e}")
                        
                        doc_final.save(ruta_salida)
                        doc_final.close()
                        self.log(f"✅ PDF '{nombre_pdf_final}' creado en '{nombre_subcarpeta}'.")

                        # Limpiar temporales
                        for pdf_temp in pdfs_temporales:
                            try:
                                os.remove(pdf_temp)
                            except:
                                pass

                except Exception as e:
                    self.log(f"Error procesando carpeta '{nombre_subcarpeta}': {e}")

        finally:
            if progress_win:
                progress_win.destroy()

    
    def accion_unificar_pdfs(self):
        archivos = filedialog.askopenfilenames(
            parent=self.root,
            title="Seleccionar PDFs a unificar",
            filetypes=[("Archivos PDF", "*.pdf")]
        )
        if len(archivos) < 2:
            self.log("Selecciona al menos dos archivos para unificar.")
            return

        archivo_map = {os.path.basename(p): p for p in archivos}
        orden_dialog = tk.Toplevel(self.root)
        orden_dialog.title("📚 Ordenar PDFs antes de unificar")
        orden_dialog.geometry("600x500")
        orden_dialog.transient(self.root)
        orden_dialog.grab_set()

        list_frame = ttk.Frame(orden_dialog)
        list_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical")
        orden_listbox = tk.Listbox(list_frame, selectmode=tk.SINGLE, yscrollcommand=scrollbar.set)
        scrollbar.config(command=orden_listbox.yview)

        scrollbar.pack(side="right", fill="y")
        orden_listbox.pack(side="left", fill="both", expand=True)

        try:
            import natsort # type: ignore
            sorted_archivos = natsort.natsorted(archivos)
        except ImportError:
            sorted_archivos = sorted(archivos)

        for archivo in sorted_archivos:
            orden_listbox.insert(tk.END, os.path.basename(archivo))
        
        botones_frame = ttk.Frame(orden_dialog)
        botones_frame.pack(pady=10)

        def mover_arriba():
            seleccion = orden_listbox.curselection()
            if not seleccion or seleccion[0] == 0: return
            i = seleccion[0]
            texto = orden_listbox.get(i)
            orden_listbox.delete(i)
            orden_listbox.insert(i - 1, texto)
            orden_listbox.select_set(i - 1)
            orden_listbox.see(i - 1)

        def mover_abajo():
            seleccion = orden_listbox.curselection()
            if not seleccion or seleccion[0] == orden_listbox.size() - 1: return
            i = seleccion[0]
            texto = orden_listbox.get(i)
            orden_listbox.delete(i)
            orden_listbox.insert(i + 1, texto)
            orden_listbox.select_set(i + 1)
            orden_listbox.see(i + 1)

        def unificar():
            ordered_basenames = orden_listbox.get(0, tk.END)
            if not ordered_basenames:
                messagebox.showwarning("Advertencia", "No hay archivos para unificar.", parent=orden_dialog)
                return

            ruta_guardado = filedialog.asksaveasfilename(
                parent=self.root,
                defaultextension=".pdf", 
                filetypes=[("Archivo PDF", "*.pdf")], 
                title="Guardar PDF unificado como"
            )
            if not ruta_guardado: return
            orden_dialog.destroy()
            self.log("Iniciando unificación manual de PDFs...")
            doc_final = fitz.open()

            progress_win, progress_bar, progress_label = self._create_progress_window("Unificando PDFs...", len(ordered_basenames))
            try:
                for i, basename in enumerate(ordered_basenames):
                    progress_label.config(text=f"Procesando: {basename}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    
                    ruta = archivo_map.get(basename)
                    if not ruta:
                        self.log(f"⚠️ No se encontró la ruta para {basename}, saltando.")
                        continue
                    
                    try:
                        with fitz.open(ruta) as doc_origen:
                           doc_final.insert_pdf(doc_origen)
                        self.log(f"📄 Procesado: {os.path.basename(ruta)}")
                    except Exception as e:
                        self.log(f"❌ Error al procesar {ruta}: {e}")
                
                if len(doc_final) > 0:
                    try:
                        doc_final.save(ruta_guardado, garbage=4, deflate=True)
                        self.log(f"✅ PDF combinado guardado en: {ruta_guardado}")
                    except Exception as e:
                        self.log(f"❌ Error al guardar el PDF final: {e}")
                else:
                    self.log("❌ No se generaron páginas. El PDF final no fue guardado.")
                doc_final.close()
            finally:
                if progress_win:
                    progress_win.destroy()

        ttk.Button(botones_frame, text="⬆️ Arriba", command=mover_arriba).pack(side="left", padx=5)
        ttk.Button(botones_frame, text="⬇️ Abajo", command=mover_abajo).pack(side="left", padx=5)
        ttk.Button(botones_frame, text="✅ Unificar y guardar", command=unificar).pack(side="left", padx=15)
        ttk.Button(botones_frame, text="❌ Cancelar", command=orden_dialog.destroy).pack(side="right", padx=5)
        
        self.root.wait_window(orden_dialog)
    
    def accion_dividir_pdf_en_paginas(self):
        archivos_pdf = filedialog.askopenfilenames(
            parent=self.root,
            title="Seleccionar PDF(s) para dividir en páginas",
            filetypes=[("Archivos PDF", "*.pdf")]
        )
        if not archivos_pdf:
            self.log("Operación de división de PDF cancelada.")
            return

        self.log(f"--- Iniciando división de {len(archivos_pdf)} PDF(s) en páginas individuales ---")
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Dividiendo PDFs...", len(archivos_pdf))
        try:
            for i, ruta_pdf_original in enumerate(archivos_pdf):
                nombre_base_original = os.path.splitext(os.path.basename(ruta_pdf_original))[0]
                directorio_origen = os.path.dirname(ruta_pdf_original)

                progress_label.config(text=f"Procesando: {nombre_base_original}.pdf")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                ruta_carpeta_salida = os.path.join(directorio_origen, nombre_base_original)
                os.makedirs(ruta_carpeta_salida, exist_ok=True)

                try:
                    with fitz.open(ruta_pdf_original) as doc_origen:
                        num_paginas = len(doc_origen)
                        if num_paginas == 0:
                            self.log(f"⚠️ El archivo '{nombre_base_original}.pdf' no tiene páginas y será omitido.")
                            continue

                        for num_pagina in range(num_paginas):
                            doc_nuevo = fitz.open()
                            doc_nuevo.insert_pdf(doc_origen, from_page=num_pagina, to_page=num_pagina)
                            
                            nombre_pagina_nueva = f"{num_pagina + 1}.pdf"
                            ruta_pagina_nueva = os.path.join(ruta_carpeta_salida, nombre_pagina_nueva)
                            
                            doc_nuevo.save(ruta_pagina_nueva)
                            doc_nuevo.close()
                        
                        self.log(f"✅ Dividido '{nombre_base_original}.pdf' en {num_paginas} páginas en la carpeta '{nombre_base_original}'")
                
                except Exception as e:
                    self.log(f"❌ Error al procesar '{nombre_base_original}.pdf': {e}")
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log("--- Proceso de división finalizado. ---")
        messagebox.showinfo("Proceso Completado", "La división de los archivos PDF ha finalizado. Revise el registro de actividad para más detalles.", parent=self.root)

    def accion_dividir_pdfs_masivamente(self):
        carpeta_base = filedialog.askdirectory(
            parent=self.root,
            title="Seleccionar carpeta principal para buscar y dividir PDFs"
        )
        if not carpeta_base:
            self.log("Operación de división masiva cancelada.")
            return

        self.log(f"--- Iniciando búsqueda y división masiva de PDFs en: {carpeta_base} ---")

        pdfs_a_procesar = []
        for root, _, files in os.walk(carpeta_base):
            for file in files:
                if file.lower().endswith('.pdf'):
                    pdfs_a_procesar.append(os.path.join(root, file))

        if not pdfs_a_procesar:
            self.log("No se encontraron archivos PDF en la carpeta seleccionada y sus subdirectorios.")
            messagebox.showinfo("Sin Archivos", "No se encontraron archivos PDF para procesar.", parent=self.root)
            return
            
        self.log(f"Se encontraron {len(pdfs_a_procesar)} PDFs para dividir.")
        
        progress_win, progress_bar, progress_label = self._create_progress_window(
            "Dividiendo PDFs masivamente...", len(pdfs_a_procesar)
        )
        
        try:
            for i, ruta_pdf_original in enumerate(pdfs_a_procesar):
                nombre_base_original = os.path.splitext(os.path.basename(ruta_pdf_original))[0]
                directorio_origen = os.path.dirname(ruta_pdf_original)

                progress_label.config(text=f"Procesando: {os.path.basename(ruta_pdf_original)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                ruta_carpeta_salida = os.path.join(directorio_origen, nombre_base_original)
                os.makedirs(ruta_carpeta_salida, exist_ok=True)

                try:
                    with fitz.open(ruta_pdf_original) as doc_origen:
                        num_paginas = len(doc_origen)
                        if num_paginas == 0:
                            self.log(f"⚠️ El archivo '{os.path.basename(ruta_pdf_original)}' no tiene páginas, omitido.")
                            continue

                        for num_pagina in range(num_paginas):
                            doc_nuevo = fitz.open()
                            doc_nuevo.insert_pdf(doc_origen, from_page=num_pagina, to_page=num_pagina)
                            
                            nombre_pagina_nueva = f"{num_pagina + 1}.pdf"
                            ruta_pagina_nueva = os.path.join(ruta_carpeta_salida, nombre_pagina_nueva)
                            
                            doc_nuevo.save(ruta_pagina_nueva)
                            doc_nuevo.close()
                        
                        self.log(f"✅ Dividido '{os.path.basename(ruta_pdf_original)}' en {num_paginas} páginas en la carpeta '{os.path.basename(ruta_carpeta_salida)}'")
                
                except Exception as e:
                    self.log(f"❌ Error al procesar '{os.path.basename(ruta_pdf_original)}': {e}")
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log("--- Proceso de división masiva finalizado. ---")
        messagebox.showinfo("Proceso Completado", "La división masiva de PDFs ha finalizado. Revise el registro de actividad para más detalles.", parent=self.root)

    def _pdf_a_escala_grises_worker(self, entrada, salida):
        es_sobrescritura = (os.path.normpath(entrada) == os.path.normpath(salida))
        ruta_temporal = entrada + "._tmp_grayscale" if es_sobrescritura else None

        try:
            with fitz.open(entrada) as doc_origen:
                if doc_origen.is_encrypted:
                    self.log(f"⚠️ El archivo '{os.path.basename(entrada)}' está encriptado y no se puede procesar.")
                    return
                
                doc_final = fitz.open()
                for page in doc_origen:
                    pix = page.get_pixmap(dpi=300, colorspace=fitz.csGRAY)
                    pagina_nueva = doc_final.new_page(width=pix.width, height=pix.height)
                    pagina_nueva.insert_image(pagina_nueva.rect, pixmap=pix)
                
                ruta_guardado = ruta_temporal if es_sobrescritura else salida
                doc_final.save(ruta_guardado, garbage=4, deflate=True)
                doc_final.close()
            
            if es_sobrescritura:
                os.replace(ruta_temporal, entrada)
                self.log(f"✅ Reemplazado (escala de grises): {os.path.basename(entrada)}")
            else:
                self.log(f"✅ Convertido a escala de grises: {os.path.basename(entrada)} → {os.path.basename(salida)}")

        except Exception as e:
            self.log(f"❌ Error convirtiendo a escala de grises '{os.path.basename(entrada)}': {e}")
        finally:
            if ruta_temporal and os.path.exists(ruta_temporal):
                os.remove(ruta_temporal)

    def accion_pdf_a_escala_grises_individual(self):
        archivos = filedialog.askopenfilenames(
            parent=self.root,
            title="Seleccionar PDF(s) para convertir a escala de grises (REEMPLAZARÁ ORIGINALES)",
            filetypes=[("Archivos PDF", "*.pdf")]
        )
        if not archivos:
            self.log("Operación de conversión a escala de grises cancelada.")
            return
            
        confirm = messagebox.askyesno(
            "Confirmar Sobrescritura",
            f"¿Estás seguro de que quieres convertir {len(archivos)} archivo(s) a escala de grises?\n\n"
            "¡ADVERTENCIA! Esta acción REEMPLAZARÁ los archivos originales y no se puede deshacer.",
            parent=self.root
        )
        if not confirm:
            self.log("Sobrescritura de archivos cancelada por el usuario.")
            return

        total_files = len(archivos)
        self.log(f"--- Iniciando conversión (con reemplazo) a escala de grises para {total_files} archivo(s) ---")
        
        progress_win, progress_bar, progress_label = self._create_progress_window(
            "Convirtiendo y reemplazando...", total_files
        )
        try:
            for i, entrada in enumerate(archivos):
                progress_label.config(text=f"Convirtiendo: {os.path.basename(entrada)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                self._pdf_a_escala_grises_worker(entrada, entrada)
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log("--- Proceso de conversión a escala de grises finalizado. ---")
        messagebox.showinfo("Proceso Completado", "Conversión a escala de grises (con reemplazo) finalizada. Revise el registro.", parent=self.root)
    
    def accion_pdf_a_escala_grises_masivo_desde_resultados(self):
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Busca archivos primero para usar esta conversión masiva.", parent=self.root)
            return

        archivos_filtrados = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        
        if not archivos_filtrados:
            messagebox.showinfo("Información", "No se encontraron archivos PDF en los resultados para convertir.", parent=self.root)
            return

        confirm = messagebox.askyesno(
            "Confirmar Sobrescritura desde Resultados",
            f"¿Estás seguro de que quieres convertir los {len(archivos_filtrados)} PDFs encontrados a escala de grises?\n\n"
            "¡ADVERTENCIA! Esta acción REEMPLAZARÁ los archivos originales y no se puede deshacer.",
            parent=self.root
        )
        if not confirm:
            self.log("Sobrescritura desde resultados cancelada por el usuario.")
            return

        self.log(f"--- Iniciando conversión (con reemplazo) a escala de grises para {len(archivos_filtrados)} archivo(s) desde los resultados ---")
        progress_win, progress_bar, progress_label = self._create_progress_window(
            "Convirtiendo y reemplazando...", len(archivos_filtrados)
        )
        try:
            for i, entrada in enumerate(archivos_filtrados):
                progress_label.config(text=f"Convirtiendo: {os.path.basename(entrada)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                self._pdf_a_escala_grises_worker(entrada, entrada)
        finally:
            if progress_win:
                progress_win.destroy()

        self.log("--- Proceso de conversión desde resultados finalizado. ---")
        messagebox.showinfo("Proceso Completado", "Conversión a escala de grises (con reemplazo) desde resultados finalizada.", parent=self.root)

    def accion_organizar_facturas_por_pdf(self):
        messagebox.showinfo(
            "Información - Organizar Facturas FEOV",
            "Esta función organiza archivos de facturas en dos pasos:\n\n"
            "1. Primero, selecciona la carpeta 'DESTINO' que contiene las subcarpetas finales (ej: una carpeta llamada '12345' que a su vez contiene un PDF '...FEOV12345.pdf').\n\n"
            "2. Luego, selecciona la carpeta 'ORIGEN' que contiene los archivos desorganizados que quieres mover.\n\n"
            "El programa moverá los archivos del 'ORIGEN' a la carpeta 'DESTINO' correcta si su nombre contiene el número de la factura FEOV.",
            parent=self.root
        )
        carpeta_destinos = filedialog.askdirectory(parent=self.root, title="PASO 1: Selecciona la carpeta principal de DESTINOS")
        if not carpeta_destinos:
            self.log("Operación cancelada. No se seleccionó carpeta de destinos.")
            return

        carpeta_origen_archivos = filedialog.askdirectory(parent=self.root, title="PASO 2: Selecciona la carpeta de ORIGEN con los archivos a mover")
        if not carpeta_origen_archivos:
            self.log("Operación cancelada. No se seleccionó carpeta de origen de archivos.")
            return

        self.log("Iniciando organización de facturas FEOV...")
        regex = re.compile(r'FEOV(\d+)', re.IGNORECASE)
        
        destinos_map = {}
        self.log("--- Analizando carpetas de destino para encontrar números de factura ---")
        
        try:
            list_carpetas_destino = [d for d in os.listdir(carpeta_destinos) if os.path.isdir(os.path.join(carpeta_destinos, d))]
        except FileNotFoundError:
            self.log(f"❌ Error: La carpeta de destinos '{carpeta_destinos}' no fue encontrada.")
            return

        for nombre_carpeta_destino in list_carpetas_destino:
            ruta_carpeta_destino = os.path.join(carpeta_destinos, nombre_carpeta_destino)
            for archivo in os.listdir(ruta_carpeta_destino):
                if archivo.lower().endswith('.pdf'):
                    match = regex.search(archivo)
                    if match:
                        numero_factura = match.group(1)
                        destinos_map[numero_factura] = ruta_carpeta_destino
                        self.log(f"  Encontrado destino para factura '{numero_factura}' en carpeta: '{nombre_carpeta_destino}'")
                        break 

        if not destinos_map:
            self.log("⚠️ No se encontraron PDFs con el patrón 'FEOV<numero>' en las carpetas de destino. Proceso detenido.")
            messagebox.showwarning("Sin Destinos", "No se encontró ningún PDF con el formato 'FEOV...' en las subcarpetas de la carpeta de destino seleccionada.", parent=self.root)
            return

        self.log(f"--- Buscando archivos para mover en '{carpeta_origen_archivos}' ---")
        movidos, errores, conflictos = 0, 0, 0

        total_files_to_scan = sum(len(files) for _, _, files in os.walk(carpeta_origen_archivos))
        if total_files_to_scan == 0:
            self.log("No se encontraron archivos en la carpeta de origen para procesar.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Organizando facturas...", total_files_to_scan)
        file_count = 0
        try:
            for root, _, files in os.walk(carpeta_origen_archivos):
                for file_to_move in files:
                    file_count += 1
                    progress_label.config(text=f"Escaneando: {file_to_move}")
                    progress_bar['value'] = file_count
                    self.root.update_idletasks()

                    for numero_factura, ruta_destino_final in destinos_map.items():
                        if numero_factura in file_to_move:
                            try:
                                ruta_origen_archivo = os.path.join(root, file_to_move)
                                ruta_final_archivo = os.path.join(ruta_destino_final, file_to_move)

                                if os.path.exists(ruta_final_archivo):
                                    self.log(f"  ⚠️ Conflicto: El archivo '{file_to_move}' ya existe en el destino. Se omite.")
                                    conflictos += 1
                                else:
                                    shutil.move(ruta_origen_archivo, ruta_destino_final)
                                    self.log(f"  ✅ Movido '{file_to_move}' a la carpeta '{os.path.basename(ruta_destino_final)}'")
                                    movidos += 1
                                break 
                            except Exception as e:
                                self.log(f"  ❌ Error moviendo '{file_to_move}': {e}")
                                errores += 1
                                break 
        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (f"Proceso de organización finalizado.\n\n"
                       f"- Archivos movidos: {movidos}\n"
                       f"- Conflictos (ya existían): {conflictos}\n"
                       f"- Errores: {errores}")
        self.log(f"--- Organización finalizada. Movidos: {movidos}, Conflictos: {conflictos}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def accion_crear_carpetas_desde_excel(self):
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel", filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
        except Exception as e:
            messagebox.showerror("Error de Archivo", f"No se pudo leer el archivo Excel:\n{e}", parent=self.root)
            return

        hoja_seleccionada = simpledialog.askstring(
            "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
            parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
        )
        if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return
        
        respuesta_filtro = messagebox.askyesnocancel(
            "Usar Filtros de Excel",
            "¿Desea crear carpetas solo para las filas VISIBLES (filtradas)?\n\n"
            "- 'Sí': Solo usará las filas que no están ocultas por un filtro.\n"
            "- 'No': Usará TODAS las filas de la hoja, ignorando cualquier filtro.\n"
            "- 'Cancelar': Abortará la operación.",
            parent=self.root
        )

        if respuesta_filtro is None:
            self.log("Operación cancelada por el usuario.")
            return

        usar_filas_visibles = respuesta_filtro

        nombres_carpetas_raw = []
        indice_columna = -1

        try:
            df_temp = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df_temp.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"
            
            letra_columna = simpledialog.askstring(
                "Seleccionar Columna", f"Ingrese la LETRA de la columna que contiene los nombres para las carpetas.\n(Opciones: {letras_rango})",
                parent=self.root
            )
            if not letra_columna or letra_columna.upper() not in letras_disponibles: return
            indice_columna = ord(letra_columna.upper()) - ord('A')
            
            if usar_filas_visibles:
                self.log("Leyendo solo filas visibles (filtradas) desde el Excel...")
                wb = openpyxl.load_workbook(archivo_excel, data_only=True)
                ws = wb[hoja_seleccionada]
                
                for i in range(2, ws.max_row + 1):
                    if not ws.row_dimensions[i].hidden:
                        valor_celda = ws.cell(row=i, column=indice_columna + 1).value
                        if valor_celda:
                            nombres_carpetas_raw.append(str(valor_celda))
            else:
                self.log("Leyendo todas las filas desde el Excel...")
                nombres_carpetas_raw = df_temp.iloc[:, indice_columna].dropna().astype(str).tolist()

        except Exception as e:
            messagebox.showerror("Error de Lectura", f"No se pudo leer la hoja '{hoja_seleccionada}':\n{e}", parent=self.root)
            return

        if not nombres_carpetas_raw:
            messagebox.showinfo("Sin Datos", "No se encontraron nombres de carpetas para crear con la configuración seleccionada.", parent=self.root)
            self.log("No se encontraron nombres en la columna y hoja especificadas.")
            return

        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta donde crear las nuevas carpetas")
        if not carpeta_base: return
        
        creadas, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Creando carpetas...", len(nombres_carpetas_raw))
        try:
            for i, nombre in enumerate(nombres_carpetas_raw):
                nombre_base = "".join(c for c in nombre if c.isalnum() or c in " _-").rstrip()
                progress_label.config(text=f"Procesando: {nombre_base}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                if not nombre_base: continue

                ruta_final = os.path.join(carpeta_base, nombre_base)
                
                if os.path.exists(ruta_final):
                    contador = 2
                    nombre_consecutivo = f"{nombre_base} ({contador})"
                    ruta_final = os.path.join(carpeta_base, nombre_consecutivo)
                    
                    while os.path.exists(ruta_final):
                        contador += 1
                        nombre_consecutivo = f"{nombre_base} ({contador})"
                        ruta_final = os.path.join(carpeta_base, nombre_consecutivo)
                    
                    self.log(f"⚠️ El nombre '{nombre_base}' ya existía. Se creará como '{os.path.basename(ruta_final)}'")
                
                try:
                    os.makedirs(ruta_final)
                    self.log(f"✅ Creada: {os.path.basename(ruta_final)}")
                    creadas += 1
                except Exception as e:
                    self.log(f"❌ Error creando '{os.path.basename(ruta_final)}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log(f"--- Proceso finalizado. Carpetas creadas: {creadas}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", f"Creación de carpetas finalizada.\n\n- Carpetas creadas: {creadas}\n- Errores: {errores}", parent=self.root)

    def accion_mover_archivos_por_coincidencia_nombre(self):
        carpeta_base = self.source_path.get()
        if not carpeta_base or not os.path.isdir(carpeta_base):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida que contenga los archivos y las carpetas de destino.", parent=self.root)
            return

        self.log("--- Iniciando el movimiento de archivos a carpetas por coincidencia de nombre ---")

        try:
            elementos = os.listdir(carpeta_base)
            archivos = [os.path.join(carpeta_base, e) for e in elementos if os.path.isfile(os.path.join(carpeta_base, e))]
            carpetas = [os.path.join(carpeta_base, e) for e in elementos if os.path.isdir(os.path.join(carpeta_base, e))]
        except Exception as e:
            self.log(f"❌ Error al leer el contenido de la carpeta: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta:\n{e}", parent=self.root)
            return

        if not archivos or not carpetas:
            self.log("No se encontraron suficientes archivos o carpetas para procesar en la ruta de origen.")
            return

        movidos, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Moviendo archivos...", len(archivos))
        try:
            for i, ruta_archivo in enumerate(archivos):
                nombre_archivo = os.path.basename(ruta_archivo)
                progress_label.config(text=f"Verificando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                for ruta_carpeta in carpetas:
                    nombre_carpeta = os.path.basename(ruta_carpeta)

                    if nombre_carpeta.lower() in nombre_archivo.lower():
                        try:
                            shutil.move(ruta_archivo, ruta_carpeta)
                            self.log(f"✅ Movido: '{nombre_archivo}' → a la carpeta '{nombre_carpeta}'")
                            movidos += 1
                            break
                        except Exception as e:
                            self.log(f"❌ Error al mover '{nombre_archivo}': {e}")
                            errores += 1
                            break
        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (f"Proceso finalizado.\n\n"
                       f"- Archivos movidos: {movidos}\n"
                       f"- Errores: {errores}")
        self.log(f"--- Proceso Finalizado. Movidos: {movidos}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def _prompt_for_excel_mapping_details(self):
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel con el mapeo", filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return None, None, None

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return None, None, None

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_origen = simpledialog.askstring("Columna de Origen", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETAS DE ORIGEN.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_origen or letra_col_origen.upper() not in letras_disponibles: return None, None, None

            letra_col_destino = simpledialog.askstring("Columna de Destino", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETAS DE DESTINO.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_destino or letra_col_destino.upper() not in letras_disponibles: return None, None, None

            col_origen_idx = ord(letra_col_origen.upper()) - ord('A')
            col_destino_idx = ord(letra_col_destino.upper()) - ord('A')

            return df, df.columns[col_origen_idx], df.columns[col_destino_idx]

        except Exception as e:
            self.log(f"❌ Error al procesar el archivo Excel: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al leer el Excel: {e}", parent=self.root)
            return None, None, None
            
    def accion_copiar_archivos_desde_mapeo_excel(self):
        self.log("--- Iniciando copia de archivos basada en mapeo de Excel (por subcarpetas) ---")
        
        df, col_origen, col_destino = self._prompt_for_excel_mapping_details()
        if df is None:
            self.log("Operación cancelada durante la selección del Excel.")
            return

        ruta_origen_base = filedialog.askdirectory(parent=self.root, title="PASO 1: Seleccione la CARPETA BASE que contiene las carpetas de Origen")
        if not ruta_origen_base:
            self.log("Operación cancelada. No se seleccionó carpeta base de origen.")
            return
        
        ruta_destino_base = filedialog.askdirectory(parent=self.root, title="PASO 2: Seleccione la CARPETA BASE que contiene las carpetas de Destino")
        if not ruta_destino_base:
            self.log("Operación cancelada. No se seleccionó carpeta base de destino.")
            return

        copiados, errores, conflictos, carpetas_no_encontradas = 0, 0, 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Copiando archivos según mapeo...", len(df))

        try:
            for index, fila in df.iterrows():
                progress_bar['value'] = index + 1
                
                if pd.isna(fila[col_origen]) or pd.isna(fila[col_destino]):
                    continue
                nombre_carpeta_origen = str(fila[col_origen]).strip()
                nombre_carpeta_destino = str(fila[col_destino]).strip()

                progress_label.config(text=f"Mapeando: {nombre_carpeta_origen} -> {nombre_carpeta_destino}")
                self.root.update_idletasks()
                
                ruta_origen_completa = os.path.join(ruta_origen_base, nombre_carpeta_origen)
                ruta_destino_completa = os.path.join(ruta_destino_base, nombre_carpeta_destino)

                if not os.path.isdir(ruta_origen_completa):
                    self.log(f"⚠️  Origen no encontrado: '{ruta_origen_completa}'. Fila {index+2} omitida.")
                    carpetas_no_encontradas += 1
                    continue
                if not os.path.isdir(ruta_destino_completa):
                    self.log(f"⚠️  Destino no encontrado: '{ruta_destino_completa}'. Fila {index+2} omitida.")
                    carpetas_no_encontradas += 1
                    continue
                
                for nombre_archivo in os.listdir(ruta_origen_completa):
                    ruta_archivo_origen = os.path.join(ruta_origen_completa, nombre_archivo)
                    if os.path.isfile(ruta_archivo_origen):
                        ruta_archivo_destino = os.path.join(ruta_destino_completa, nombre_archivo)
                        
                        if os.path.exists(ruta_archivo_destino):
                            self.log(f"  -> Conflicto: '{nombre_archivo}' ya existe en el destino. Omitido.")
                            conflictos += 1
                            continue
                        
                        try:
                            shutil.copy2(ruta_archivo_origen, ruta_archivo_destino)
                            self.log(f"  -> Copiado: '{nombre_archivo}' a la carpeta '{nombre_carpeta_destino}'")
                            copiados += 1
                        except Exception as e:
                            self.log(f"  -> ❌ Error al copiar '{nombre_archivo}': {e}")
                            errores += 1
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen_msg = (
            f"Proceso de copia por mapeo finalizado.\n\n"
            f"- Archivos copiados exitosamente: {copiados}\n"
            f"- Conflictos (archivos ya existentes): {conflictos}\n"
            f"- Carpetas no encontradas: {carpetas_no_encontradas}\n"
            f"- Errores de copia: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def _prompt_for_excel_file_mapping_details(self):
        """Pide al usuario los detalles para el mapeo Excel de IDENTIFICADOR a CARPETA."""
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel con el mapeo (Identificador -> Carpeta)",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return None, None, None

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return None, None, None

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_archivo = simpledialog.askstring("Columna de Identificadores", f"Ingrese la LETRA de la columna con los IDENTIFICADORES DE ARCHIVO (ej: FEOV123).\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_archivo or letra_col_archivo.upper() not in letras_disponibles: return None, None, None

            letra_col_carpeta = simpledialog.askstring("Columna de Carpetas Destino", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETA DE DESTINO.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return None, None, None

            col_archivo_idx = ord(letra_col_archivo.upper()) - ord('A')
            col_carpeta_idx = ord(letra_col_carpeta.upper()) - ord('A')

            return df, df.columns[col_archivo_idx], df.columns[col_carpeta_idx]

        except Exception as e:
            self.log(f"❌ Error al procesar el archivo Excel: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al leer el Excel: {e}", parent=self.root)
            return None, None, None

    def accion_copiar_archivos_desde_raiz_mapeo_excel(self):
        """Copia archivos desde la RAÍZ de una carpeta a subcarpetas, usando coincidencia parcial de nombres desde un Excel."""
        self.log("--- Iniciando copia de archivos desde la RAÍZ (Mapeo Parcial) ---")
        
        df, col_identificador, col_carpeta = self._prompt_for_excel_file_mapping_details()
        if df is None:
            self.log("Operación cancelada durante la selección del Excel.")
            return

        ruta_origen_raiz = filedialog.askdirectory(parent=self.root, title="PASO 1: Seleccione la CARPETA RAÍZ que contiene los archivos de Origen")
        if not ruta_origen_raiz:
            self.log("Operación cancelada. No se seleccionó carpeta raíz de origen.")
            return
        
        ruta_destino_base = filedialog.askdirectory(parent=self.root, title="PASO 2: Seleccione la CARPETA BASE donde se encuentran las carpetas de Destino")
        if not ruta_destino_base:
            self.log("Operación cancelada. No se seleccionó carpeta base de destino.")
            return

        try:
            archivos_en_origen = [f for f in os.listdir(ruta_origen_raiz) if os.path.isfile(os.path.join(ruta_origen_raiz, f))]
            self.log(f"Se encontraron {len(archivos_en_origen)} archivos en la carpeta de origen para escanear.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo leer la carpeta de origen: {e}", parent=self.root)
            self.log(f"❌ Error fatal al leer la carpeta de origen '{ruta_origen_raiz}': {e}")
            return

        copiados, errores, conflictos, no_encontrados = 0, 0, 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Copiando por mapeo parcial...", len(df))

        try:
            for index, fila in df.iterrows():
                progress_bar['value'] = index + 1
                
                if pd.isna(fila[col_identificador]) or pd.isna(fila[col_carpeta]):
                    continue
                
                identificador_archivo = str(fila[col_identificador]).strip()
                nombre_carpeta_destino = str(fila[col_carpeta]).strip()

                if not identificador_archivo or not nombre_carpeta_destino:
                    continue

                progress_label.config(text=f"Buscando por: {identificador_archivo}")
                self.root.update_idletasks()
                
                ruta_archivo_encontrado = None
                nombre_archivo_real = ""
                
                for f_name in archivos_en_origen:
                    # <<< INICIO DE LA MODIFICACIÓN: Búsqueda case-insensitive >>>
                    if identificador_archivo.lower() in f_name.lower():
                    # <<< FIN DE LA MODIFICACIÓN >>>
                        nombre_archivo_real = f_name
                        ruta_archivo_encontrado = os.path.join(ruta_origen_raiz, f_name)
                        break 

                if not ruta_archivo_encontrado:
                    self.log(f"⚠️ No se encontró archivo que contenga '{identificador_archivo}' en la raíz. Fila {index+2} omitida.")
                    no_encontrados += 1
                    continue
                
                self.log(f"  -> Coincidencia para '{identificador_archivo}': Se usará el archivo '{nombre_archivo_real}'")
                
                ruta_carpeta_destino_completa = os.path.join(ruta_destino_base, nombre_carpeta_destino)
                
                try:
                    os.makedirs(ruta_carpeta_destino_completa, exist_ok=True)
                except Exception as e:
                    self.log(f"❌ Error creando carpeta destino '{nombre_carpeta_destino}': {e}. Fila {index+2} omitida.")
                    errores += 1
                    continue

                ruta_archivo_destino = os.path.join(ruta_carpeta_destino_completa, nombre_archivo_real)

                if os.path.exists(ruta_archivo_destino):
                    self.log(f"  -> Conflicto: '{nombre_archivo_real}' ya existe en '{nombre_carpeta_destino}'. Omitido.")
                    conflictos += 1
                    continue
                
                try:
                    shutil.copy2(ruta_archivo_encontrado, ruta_archivo_destino)
                    self.log(f"  -> ✅ Copiado: '{nombre_archivo_real}' a la carpeta '{nombre_carpeta_destino}'")
                    copiados += 1
                except Exception as e:
                    self.log(f"  -> ❌ Error al copiar '{nombre_archivo_real}': {e}")
                    errores += 1
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen_msg = (
            f"Proceso de copia por mapeo parcial finalizado.\n\n"
            f"- Archivos copiados exitosamente: {copiados}\n"
            f"- Identificadores no encontrados: {no_encontrados}\n"
            f"- Conflictos (archivos ya existentes): {conflictos}\n"
            f"- Errores de copia/creación de carpetas: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
    
    def accion_copiar_archivo_a_subcarpetas(self):
        self.log("--- Iniciando copia de un archivo a múltiples subcarpetas ---")

        archivo_a_copiar = filedialog.askopenfilename(
            parent=self.root,
            title="PASO 1: Selecciona el ARCHIVO que quieres copiar"
        )
        if not archivo_a_copiar:
            self.log("Operación cancelada. No se seleccionó ningún archivo.")
            return

        carpeta_destino_base = filedialog.askdirectory(
            parent=self.root,
            title="PASO 2: Selecciona la CARPETA que contiene las subcarpetas de destino"
        )
        if not carpeta_destino_base:
            self.log("Operación cancelada. No se seleccionó la carpeta de destino.")
            return

        try:
            subcarpetas = [os.path.join(carpeta_destino_base, d) for d in os.listdir(carpeta_destino_base) if os.path.isdir(os.path.join(carpeta_destino_base, d))]
        except Exception as e:
            self.log(f"❌ Error al leer las subcarpetas de '{carpeta_destino_base}': {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta de destino:\n{e}", parent=self.root)
            return
        
        if not subcarpetas:
            self.log("No se encontraron subcarpetas en la ruta seleccionada.")
            messagebox.showinfo("Sin Carpetas", "No se encontraron subcarpetas en la ruta de destino para realizar la copia.", parent=self.root)
            return

        self.log(f"Se copiará el archivo '{os.path.basename(archivo_a_copiar)}' en {len(subcarpetas)} carpetas.")

        copiados, conflictos, errores = 0, 0, 0
        nombre_archivo_origen = os.path.basename(archivo_a_copiar)
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Copiando '{nombre_archivo_origen}'...", len(subcarpetas))

        try:
            for i, ruta_subcarpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(ruta_subcarpeta)
                progress_label.config(text=f"Copiando a: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                ruta_destino_final = os.path.join(ruta_subcarpeta, nombre_archivo_origen)

                if os.path.exists(ruta_destino_final):
                    self.log(f"  ⚠️ Conflicto: '{nombre_archivo_origen}' ya existe en '{nombre_subcarpeta}'. Omitido.")
                    conflictos += 1
                    continue

                try:
                    shutil.copy2(archivo_a_copiar, ruta_destino_final)
                    self.log(f"  ✅ Copiado a '{nombre_subcarpeta}'")
                    copiados += 1
                except Exception as e:
                    self.log(f"  ❌ Error al copiar a '{nombre_subcarpeta}': {e}")
                    errores += 1

        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (
            f"Proceso de copia finalizado.\n\n"
            f"- Archivo copiado exitosamente en: {copiados} carpetas\n"
            f"- Conflictos (archivos ya existentes): {conflictos}\n"
            f"- Errores de copia: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def _prompt_for_url_download_details(self):
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel con los datos de descarga", filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return None, None, None

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return None, None, None

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_id = simpledialog.askstring("Columna de ID", f"Ingrese la LETRA de la columna con los NÚMEROS para la URL.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_id or letra_col_id.upper() not in letras_disponibles: return None, None, None

            letra_col_carpeta = simpledialog.askstring("Columna de Carpeta", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETA de destino.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return None, None, None

            col_id_idx = ord(letra_col_id.upper()) - ord('A')
            col_carpeta_idx = ord(letra_col_carpeta.upper()) - ord('A')

            return df, df.columns[col_id_idx], df.columns[col_carpeta_idx]

        except Exception as e:
            self.log(f"❌ Error al procesar el archivo Excel: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al leer el Excel: {e}", parent=self.root)
            return None, None, None

    def accion_descargar_firmas_url_excel(self):
        self.log("--- Iniciando descarga de firmas desde URL por Excel ---")
        
        df, col_id, col_carpeta = self._prompt_for_url_download_details()
        if df is None:
            self.log("Operación cancelada durante la selección del Excel.")
            return

        base_path = filedialog.askdirectory(parent=self.root, title="Seleccione la CARPETA BASE donde se crearán las carpetas de destino")
        if not base_path:
            self.log("Operación cancelada. No se seleccionó carpeta base de destino.")
            return
        
        base_url = "https://oportunidaddevida.com/opvcitas/admisionescall/firmas/"
        descargados, no_encontrados, errores_red = 0, 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Descargando firmas...", len(df))

        try:
            for index, row in df.iterrows():
                progress_bar['value'] = index + 1
                
                if pd.isna(row[col_id]) or pd.isna(row[col_carpeta]):
                    self.log(f"⚠️ Fila {index+2} omitida: contiene datos vacíos.")
                    continue
                
                id_firma = str(row[col_id]).strip()
                nombre_carpeta = str(row[col_carpeta]).strip()
                
                progress_label.config(text=f"Procesando ID: {id_firma}")
                self.root.update_idletasks()

                url_completa = f"{base_url}{id_firma}.png"
                ruta_carpeta_destino = os.path.join(base_path, nombre_carpeta)
                os.makedirs(ruta_carpeta_destino, exist_ok=True)
                
                try:
                    response = requests.get(url_completa, stream=True, timeout=15)
                    
                    if response.status_code == 200:
                        ruta_salida_jpg = os.path.join(ruta_carpeta_destino, "firma.jpg")
                        
                        img = Image.open(io.BytesIO(response.content))
                        img_rgb = img.convert('RGB')
                        img_rgb.save(ruta_salida_jpg, 'JPEG')
                        
                        self.log(f"✅ Firma descargada para ID {id_firma} en carpeta '{nombre_carpeta}'")
                        descargados += 1
                    
                    else:
                        ruta_txt_error = os.path.join(ruta_carpeta_destino, "no tiene firma.txt")
                        with open(ruta_txt_error, 'w', encoding='utf-8') as f:
                            f.write(f"No se encontró firma en la URL: {url_completa}\n")
                            f.write(f"Código de estado: {response.status_code}\n")
                        
                        self.log(f"⚠️ No se encontró firma en URL para ID {id_firma} (Estado: {response.status_code}). Se creó TXT.")
                        no_encontrados += 1

                except requests.exceptions.RequestException as e:
                    self.log(f"❌ Error de red para ID {id_firma}: {e}")
                    errores_red += 1

        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (
            f"Proceso de descarga de firmas finalizado.\n\n"
            f"- Firmas descargadas y guardadas: {descargados}\n"
            f"- Firmas no encontradas (404, etc.): {no_encontrados}\n"
            f"- Errores de red (timeout, etc.): {errores_red}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    # === INICIO NUEVA FUNCIONALIDAD OVIDA ===


    def accion_descargar_historias_hospitalizacion_ovida(self):
        """Orquesta la descarga de historias de hospitalización desde OVIDA."""
        if not self.selenium_available:
            messagebox.showerror(
                "Dependencias Faltantes",
                "Esta función requiere 'selenium' y 'webdriver-manager'.\n\n"
                "Instálalos con: pip install selenium webdriver-manager",
                parent=self.root
            )
            self.log("❌ Faltan dependencias de Selenium para la descarga de hospitalización.")
            return

        self.log("--- Iniciando descarga de Historias de Hospitalización (OVIDA) ---")
        
        # 1. Pedir archivo Excel y mapeo
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar Excel con datos de hospitalización",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel:
            self.log("Operación cancelada. No se seleccionó archivo Excel.")
            return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja a usar.\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return
            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            # Pedir mapeo de columnas
            letra_col_estudio = simpledialog.askstring("Columna 'Nro ESTUDIO'", f"Ingrese la letra de la columna 'Nro ESTUDIO'.\n(Opciones: {letras_rango})", parent=self.root)
            letra_col_ingreso = simpledialog.askstring("Columna 'FECHA DE INGRESO'", f"Ingrese la letra de la columna 'FECHA DE INGRESO'.\n(Opciones: {letras_rango})", parent=self.root)
            letra_col_egreso = simpledialog.askstring("Columna 'FECHA EGRESO'", f"Ingrese la letra de la columna 'FECHA EGRESO'.\n(Opciones: {letras_rango})", parent=self.root)
            letra_col_carpeta = simpledialog.askstring("Columna 'NOMBRE DE LAS CARPETAS'", f"Ingrese la letra de la columna 'NOMBRE DE LAS CARPETAS'.\n(Opciones: {letras_rango})", parent=self.root)

            if not all([letra_col_estudio, letra_col_ingreso, letra_col_egreso, letra_col_carpeta]):
                self.log("Operación cancelada, no se especificaron todas las columnas.")
                return
            
            col_map = {
                'estudio': df.columns[ord(letra_col_estudio.upper()) - ord('A')],
                'ingreso': df.columns[ord(letra_col_ingreso.upper()) - ord('A')],
                'egreso': df.columns[ord(letra_col_egreso.upper()) - ord('A')],
                'carpeta': df.columns[ord(letra_col_carpeta.upper()) - ord('A')],
            }

        except Exception as e:
            self.log(f"❌ Error procesando el archivo Excel: {e}")
            messagebox.showerror("Error de Excel", f"No se pudo leer el archivo o las columnas.\n\n{e}", parent=self.root)
            return

        # 2. Pedir carpeta de guardado
        save_path = filedialog.askdirectory(parent=self.root, title="Seleccione la carpeta BASE para guardar las historias")
        if not save_path:
            self.log("Operación cancelada. No se seleccionó carpeta de guardado.")
            return

        # 3. Iniciar Selenium y esperar login
        webdriver = self.selenium_deps["webdriver"]
        Service = self.selenium_deps["Service"]
        ChromeDriverManager = self.selenium_deps["ChromeDriverManager"]
        driver = None
        
        try:
            self.log("Abriendo navegador. Por favor, inicie sesión en OVIDA.")
            # Iniciar navegador visible
            options = webdriver.ChromeOptions()
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)
            driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")

            # Pausar y esperar que el usuario confirme
            if not messagebox.askokcancel(
                "Acción Requerida: Iniciar Sesión",
                "Se ha abierto una ventana de Chrome.\n\n"
                "1. Inicie sesión en la plataforma OVIDA.\n"
                "2. Seleccione el punto de atención si es necesario.\n\n"
                "Cuando esté en la página principal, haga clic en 'Aceptar' para comenzar las descargas.",
                parent=self.root):
                self.log("Operación cancelada por el usuario durante el inicio de sesión.")
                driver.quit()
                return

            self.log("Inicio de sesión confirmado por el usuario. Empezando descargas...")
            
            # 4. Bucle de procesamiento
            descargados, errores, conflictos = 0, 0, 0
            total_filas = len(df)
            
            progress_win, progress_bar, progress_label = self._create_progress_window("Descargando Historias...", total_filas)

            for index, row in df.iterrows():
                progress_bar['value'] = index + 1
                
                # Extraer y limpiar datos de la fila
                try:
                    nro_estudio = str(int(row[col_map['estudio']])).strip()
                    # Convertir fechas
                    fecha_ingreso_dt = pd.to_datetime(row[col_map['ingreso']])
                    fecha_ingreso = fecha_ingreso_dt.strftime('%Y/%m/%d')
                    fecha_egreso_dt = pd.to_datetime(row[col_map['egreso']])
                    fecha_egreso = fecha_egreso_dt.strftime('%Y/%m/%d')
                    nombre_carpeta = str(row[col_map['carpeta']]).strip()
                except Exception as ex:
                    self.log(f"⚠️ Fila {index+2} omitida por datos inválidos o vacíos (Estudio, Fechas, Carpeta). Error: {ex}")
                    errores += 1
                    continue
                
                if not all([nro_estudio, fecha_ingreso, fecha_egreso, nombre_carpeta]):
                    self.log(f"⚠️ Fila {index+2} omitida por tener celdas vacías.")
                    errores += 1
                    continue

                progress_label.config(text=f"Procesando estudio: {nro_estudio}")
                self.root.update_idletasks()

                # Construir URL
                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                params = {
                    'estudio': nro_estudio, 'fecha_inicio': fecha_ingreso, 'fecha_fin': fecha_egreso,
                    'verHC': 1, 'verEvo': 1, 'verPar': 1, 'ImprimirOrdenamiento': 1,
                    'ImprimirNotasPcte': 0, 'ImprimirSolOrdenesExt': 1, 'ImprimirGraficasHC': 1,
                    'ImprimirFormatos': 1, 'ImprimirRegistroAdmon': 1, 'ImprimirNovedad': 0,
                    'ImprimirRecomendaciones': 0, 'ImprimirDescripcionQX': 0, 'ImprimirNotasEnfermeria': 1,
                    'ImprimirSignosVitales': 0, 'ImprimirLog': 0, 'ImprimirEpicrisisSinHC': 0
                }
                
                full_url = f"{base_url}?{urllib.parse.urlencode(params)}"

                # Crear carpeta destino y verificar si el archivo ya existe
                dest_folder = os.path.join(save_path, nombre_carpeta)
                os.makedirs(dest_folder, exist_ok=True)
                final_file_path = os.path.join(dest_folder, f"HC_{nro_estudio}.pdf")

                if os.path.exists(final_file_path):
                    self.log(f"  -> Conflicto: El archivo '{os.path.basename(final_file_path)}' ya existe en '{nombre_carpeta}'. Omitido.")
                    conflictos += 1
                    continue

                # Descargar
                try:
                    self.log(f"  > Navegando para descargar estudio {nro_estudio}")
                    driver.get(full_url)
                    time.sleep(2) # Pequeña espera para asegurar que la página se renderiza
                    
                    pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {
                        "landscape": False, "printBackground": True,
                        "paperWidth": 8.5, "paperHeight": 11,
                        "marginTop": 0.4, "marginBottom": 0.4, "marginLeft": 0.4, "marginRight": 0.4
                    })
                    
                    pdf_data = base64.b64decode(pdf_b64['data'])
                    
                    with open(final_file_path, 'wb') as f:
                        f.write(pdf_data)
                    
                    self.log(f"  ✅ Descargado: '{os.path.basename(final_file_path)}' en carpeta '{nombre_carpeta}'")
                    descargados += 1

                except Exception as e:
                    self.log(f"  ❌ Error al descargar el estudio {nro_estudio}: {str(e).split(chr(10))[0]}")
                    errores += 1
            
            # Finalización
            progress_win.destroy()
            resumen_msg = (f"Proceso finalizado.\n\n"
                           f"- Historias descargadas: {descargados}\n"
                           f"- Errores: {errores}\n"
                           f"- Conflictos (ya existían): {conflictos}")
            self.log(f"--- {resumen_msg.replace(chr(10)*2, ' ').replace(chr(10), ' | ')} ---")
            messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

        except Exception as e:
            self.log(f"❌ Ocurrió un error inesperado durante el proceso: {e}")
            messagebox.showerror("Error Crítico", f"Ha ocurrido un error:\n\n{e}", parent=self.root)
        finally:
            if driver:
                driver.quit()
                self.log("Navegador cerrado.")

    def _proceso_descarga_en_hilo(self, df, col_tid, col_num, col_estudio, download_path):
        """Maneja el bucle de descarga y la barra de progreso."""
        descargados = 0
        errores = 0
        total_pacientes = len(df)
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Descargando Historias Clínicas...", total_pacientes)
        
        for index, row in df.iterrows():
            progress_bar['value'] = index + 1
            
            tipo_id = str(row[col_tid]).strip()
            numero_id = str(row[col_num]).strip()
            nro_estudio = str(row[col_estudio]).strip()
            
            progress_label.config(text=f"Procesando: {numero_id}")
            self.root.update_idletasks()

            if not all([tipo_id, numero_id, nro_estudio]):
                self.log(f"⚠️ Fila {index+2} omitida por datos faltantes.")
                errores += 1
                continue
            
            self.log(f"--- Intentando descargar historia para ID: {numero_id}, Estudio: {nro_estudio} ---")
            
            if self._worker_descargar_historia(tipo_id, numero_id, nro_estudio, download_path):
                descargados += 1
            else:
                errores += 1
        
        progress_win.destroy()
        resumen_msg = (
            f"Proceso de descarga de historias finalizado.\n\n"
            f"- Historias descargadas: {descargados}\n"
            f"- Registros con error: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)*2, ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)


    def _worker_descargar_historia(self, tipo_id, numero_id, nro_estudio, download_path):
        """Realiza la automatización web para un solo paciente."""
        # Extraer las dependencias de Selenium
        webdriver = self.selenium_deps["webdriver"]
        By = self.selenium_deps["By"]
        WebDriverWait = self.selenium_deps["WebDriverWait"]
        Select = self.selenium_deps["Select"]
        EC = self.selenium_deps["EC"]
        Service = self.selenium_deps["Service"]
        Keys = self.selenium_deps["Keys"]
        ChromeDriverManager = self.selenium_deps["ChromeDriverManager"]
        
        # Configurar opciones de Chrome para descarga automática
        options = webdriver.ChromeOptions()
        prefs = {
            "download.default_directory": download_path,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True  # Importante para forzar la descarga
        }
        options.add_experimental_option("prefs", prefs)
        options.add_argument("--headless")  # Ejecutar en segundo plano
        options.add_argument("--log-level=3") # Suprimir logs de consola de selenium
        
        driver = None
        try:
            # Instalar/usar el driver de Chrome automáticamente
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)
            wait = WebDriverWait(driver, 20)  # Aumentar tiempo de espera a 20s

            # PASO 2: Ingresar al link
            url = "https://ovidazs.siesacloud.com/ZeusSalud/ips/App/Vistas/Hc/Historia.php"
            driver.get(url)

            # Seleccionar tipo de documento
            tipo_doc_select_el = wait.until(EC.presence_of_element_located((By.ID, "tipo-documento")))
            select = Select(tipo_doc_select_el)
            select.select_by_value(tipo_id) # Usar el valor (ej: 'CC') en lugar del texto visible
            
            # Ingresar número y presionar Enter
            num_doc_input = driver.find_element(By.ID, "documento")
            num_doc_input.send_keys(numero_id)
            num_doc_input.send_keys(Keys.RETURN)
            self.log("  > Datos de paciente enviados.")

            # PASO 3: Desplegar panel derecho y seleccionar estudio
            panel_button = wait.until(EC.element_to_be_clickable((By.ID, "iconoFrameDer")))
            panel_button.click()
            
            # Esperar y hacer clic en el número de estudio
            estudio_link = wait.until(EC.element_to_be_clickable((By.XPATH, f"//div[@class='letraDisplay' and contains(text(), '{nro_estudio}')]")))
            estudio_link.click()
            self.log(f"  > Estudio '{nro_estudio}' seleccionado.")
            
            # PASO 4: Pestaña Imprimir y opciones
            imprimir_tab = wait.until(EC.element_to_be_clickable((By.XPATH, "//div[contains(@class, 'tabsOrdinarios')]/ul/li/a[contains(@onclick, 'Imprimir()')]")))
            imprimir_tab.click()

            # Cambiar al iframe del modal de impresión
            iframe = wait.until(EC.presence_of_element_located((By.XPATH, "//iframe[contains(@src, 'reportes/impresionHC.php')]")))
            driver.switch_to.frame(iframe)

            # Clic en botón 'General'
            general_button = wait.until(EC.element_to_be_clickable((By.NAME, "opcion")))
            general_button.click()
            
            # Marcar checkbox 'Imp. Notas de enfermería'
            # El input no tiene id, lo buscamos por el texto de su label
            notas_enf_check = wait.until(EC.element_to_be_clickable((By.XPATH, "//td[contains(text(), 'Imp. Notas de enfermería')]/input[@type='checkbox']")))
            if not notas_enf_check.is_selected():
                 notas_enf_check.click()
            self.log("  > Opciones de impresión seleccionadas.")

            # Preparar para la descarga
            files_before = os.listdir(download_path)
            
            # Clic en el icono de imprimir
            imprimir_icon = wait.until(EC.element_to_be_clickable((By.XPATH, "//img[@title='Imprimir']")))
            imprimir_icon.click()
            
            # Esperar a que la descarga finalice
            self.log("  > Descarga iniciada. Esperando finalización...")
            time_waited = 0
            download_completed = False
            while time_waited < 60: # Esperar máximo 60 segundos por la descarga
                files_after = os.listdir(download_path)
                new_files = [f for f in files_after if f not in files_before and not f.endswith('.crdownload')]
                if new_files:
                    downloaded_file = new_files[0]
                    # Renombrar el archivo
                    os.rename(
                        os.path.join(download_path, downloaded_file),
                        os.path.join(download_path, f"{numero_id}.pdf")
                    )
                    self.log(f"✅ Descarga completada y renombrada a '{numero_id}.pdf'")
                    download_completed = True
                    break
                time.sleep(1)
                time_waited += 1

            if not download_completed:
                self.log(f"❌ Error: La descarga no se completó en 60 segundos para el ID {numero_id}.")
                return False

            return True

        except Exception as e:
            # Captura cualquier error de Selenium (ej: elemento no encontrado) o del sistema de archivos
            self.log(f"❌ Error procesando ID {numero_id}: {type(e).__name__} - {str(e).split(chr(10))[0]}")
            return False
        
        finally:
            # Asegurarse de cerrar el navegador sin importar lo que pase
            if driver:
                driver.quit()
    # === FIN NUEVA FUNCIONALIDAD OVIDA ===

    def accion_consolidar_archivos_subcarpetas(self):
        base_path = filedialog.askdirectory(parent=self.root, title="Seleccionar la carpeta base que contiene las carpetas a procesar")
        if not base_path:
            self.log("Operación de consolidación cancelada.")
            return

        try:
            main_folders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not main_folders:
            self.log("No se encontraron carpetas para procesar en la ruta seleccionada.")
            messagebox.showinfo("Sin Carpetas", "No se encontraron carpetas en la ruta seleccionada para procesar.", parent=self.root)
            return

        self.log(f"--- Iniciando consolidación de archivos para {len(main_folders)} carpetas en '{base_path}' ---")
        
        copiados, conflictos, errores = 0, 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Consolidando carpetas...", len(main_folders))

        try:
            for i, folder_name in enumerate(main_folders):
                main_folder_path = os.path.join(base_path, folder_name)
                
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                self.log(f"-> Procesando carpeta: '{folder_name}'")
                
                for sub_root, _, files in os.walk(main_folder_path):
                    if sub_root == main_folder_path:
                        continue

                    for file_name in files:
                        source_path = os.path.join(sub_root, file_name)
                        dest_path = os.path.join(main_folder_path, file_name)

                        try:
                            if os.path.exists(dest_path):
                                self.log(f"  ⚠️ Conflicto: El archivo '{file_name}' ya existe en la raíz de '{folder_name}'. Se omite la copia.")
                                conflictos += 1
                                continue
                            
                            shutil.copy2(source_path, dest_path)
                            self.log(f"  ✅ Copiado: '{file_name}' a la carpeta '{folder_name}'")
                            copiados += 1
                        
                        except Exception as e:
                            self.log(f"  ❌ Error al copiar '{file_name}': {e}")
                            errores += 1
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen = (f"Proceso de consolidación finalizado.\n\n"
                   f"- Archivos copiados: {copiados}\n"
                   f"- Conflictos (archivos omitidos): {conflictos}\n"
                   f"- Errores: {errores}")
        self.log(f"--- {resumen.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)
    
    def _crear_firma_estilizada(self, texto):
        """
        Crea una firma digital estilizada sin usar fuentes tipográficas.
        Convierte cada letra del texto en un trazo manuscrito único.
        """
        
        # Dimensiones de la imagen
        width = max(400, len(texto) * 40)
        height = 150
        
        # Crear imagen base con fondo blanco
        image = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(image)
        
        # Convertir cada letra del texto en un trazo manuscrito
        colores = ['black', 'gray', 'darkgray']
        
        # Procesar cada letra del nombre
        for i, letra in enumerate(texto):
            if letra.isspace():
                continue
                
            # Posición basada en la letra (más espaciado)
            x_base = 30 + (i * (width - 60) // len(texto))
            y_centro = height // 2
            
            # Crear trazo único basado en la letra
            # Usar el código ASCII de la letra para determinar el estilo
            ascii_val = ord(letra.upper()) if letra.isalpha() else ord('A')
            
            # Determinar la forma del trazo basada en la letra
            if letra.upper() in 'AEIOU':
                # Vocales: líneas curvas y abiertas
                self._dibujar_trazo_vocal(draw, x_base, y_centro, ascii_val, colores)
            elif letra.upper() in 'BCDFG':
                # Consonantes duras: líneas fuertes y angulares
                self._dibujar_trazo_consonante_dura(draw, x_base, y_centro, ascii_val, colores)
            else:
                # Otras letras: líneas fluidas
                self._dibujar_trazo_generico(draw, x_base, y_centro, ascii_val, colores)
        
        return image
    
    def _dibujar_trazo_vocal(self, draw, x_base, y_centro, ascii_val, colores):
        """Dibuja un trazo característico de vocal"""
        color = random.choice(colores)
        grosor = random.randint(2, 3)
        
        # Crear arco o curva
        altura_arco = 20 + (ascii_val % 15)
        puntos = []
        
        for i in range(20):
            angulo = (i / 19.0) * math.pi
            x = x_base + i * 2
            y = y_centro - math.sin(angulo) * altura_arco + random.randint(-2, 2)
            puntos.append((x, y))
        
        # Dibujar curva
        for i in range(len(puntos) - 1):
            draw.line([puntos[i], puntos[i + 1]], fill=color, width=grosor)
    
    def _dibujar_trazo_consonante_dura(self, draw, x_base, y_centro, ascii_val, colores):
        """Dibuja un trazo característico de consonante dura"""
        color = random.choice(colores)
        grosor = random.randint(3, 4)
        
        # Línea con ángulos y cambios de dirección
        x = x_base
        y = y_centro + random.randint(-10, 10)
        
        # Línea ascendente
        draw.line([(x, y), (x + 15, y - 20)], fill=color, width=grosor)
        # Línea horizontal
        draw.line([(x + 15, y - 20), (x + 30, y - 15)], fill=color, width=grosor)
        # Línea descendente
        draw.line([(x + 30, y - 15), (x + 40, y + 10)], fill=color, width=grosor)
    
    def _dibujar_trazo_generico(self, draw, x_base, y_centro, ascii_val, colores):
        """Dibuja un trazo genérico fluido"""
        color = random.choice(colores)
        grosor = random.randint(2, 3)
        
        # Crear línea ondulada
        puntos = []
        for i in range(30):
            x = x_base + i * 1.5
            # Onda basada en el valor ASCII
            onda = math.sin((x - x_base) * 0.2 + ascii_val * 0.1) * 15
            y = y_centro + onda + random.randint(-3, 3)
            puntos.append((x, y))
        
        # Dibujar línea ondulada
        for i in range(len(puntos) - 1):
            draw.line([puntos[i], puntos[i + 1]], fill=color, width=grosor)
        
        # Añadir algunos puntos extra para dar textura
        for _ in range(random.randint(2, 4)):
            punto_x = random.randint(x_base, x_base + 40)
            punto_y = y_centro + random.randint(-5, 5)
            draw.ellipse([punto_x - 1, punto_y - 1, punto_x + 1, punto_y + 1], fill=color)
        
        # Añadir algunos detalles decorativos adicionales
        # Pequeñas líneas adicionales
        for _ in range(random.randint(1, 3)):
            x_rand = random.randint(20, width - 20)
            y_rand = random.randint(height // 3, 2 * height // 3)
            longitud = random.randint(10, 25)
            angulo = random.uniform(-0.5, 0.5)
            
            x_fin = x_rand + int(longitud * math.cos(angulo))
            y_fin = y_rand + int(longitud * math.sin(angulo))
            
            draw.line([(x_rand, y_rand), (x_fin, y_fin)], fill=random.choice(colores), width=random.randint(1, 2))
        
        # Añadir un toque final con una línea más prominente
        y_linea_principal = height // 2 + random.randint(-10, 10)
        x_inicio_principal = random.randint(25, 45)
        x_fin_principal = width - random.randint(25, 45)
        
        # Crear la línea principal más larga
        puntos_principales = []
        for x in range(x_inicio_principal, x_fin_principal, 3):
            onda = math.sin((x - x_inicio_principal) * 0.08) * 8
            y = y_linea_principal + onda + random.randint(-2, 2)
            puntos_principales.append((x, y))
        
        # Dibujar línea principal
        if len(puntos_principales) > 1:
            for i in range(len(puntos_principales) - 1):
                draw.line([puntos_principales[i], puntos_principales[i + 1]], fill='black', width=3)
        
        return image
    
    def accion_crear_firma_digital(self):
        base_path = filedialog.askdirectory(
            parent=self.root,
            title="Seleccionar carpeta base con las carpetas a procesar"
        )
        if not base_path:
            self.log("Operación de creación de firma cancelada.")
            return

        font_path = filedialog.askopenfilename(
            parent=self.root,
            title="Selecciona un archivo de fuente para la firma (ej: .ttf, .otf)",
            filetypes=[("Archivos de Fuente", "*.ttf *.otf"), ("Todos los archivos", "*.*")]
        )
        if not font_path:
            self.log("No se seleccionó una fuente. Operación cancelada.")
            return

        font_size = simpledialog.askinteger(
            "Tamaño de Fuente", "Introduce el tamaño de la fuente para la firma:",
            parent=self.root, initialvalue=60, minvalue=10, maxvalue=500
        )
        if not font_size:
            self.log("No se especificó un tamaño de fuente. Operación cancelada.")
            return

        try:
            font = ImageFont.truetype(font_path, font_size)
        except IOError:
            self.log(f"❌ Error: No se pudo cargar el archivo de fuente '{font_path}'.")
            messagebox.showerror("Error de Fuente", "El archivo seleccionado no pudo ser cargado como una fuente.", parent=self.root)
            return

        try:
            folders_to_process = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not folders_to_process:
            self.log("No se encontraron carpetas para procesar.")
            return

        self.log(f"--- Iniciando creación de firmas digitales para {len(folders_to_process)} carpetas ---")
        self.log(f"Usando fuente: {os.path.basename(font_path)}")
        
        creadas, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Creando firmas...", len(folders_to_process))

        try:
            for i, folder_name in enumerate(folders_to_process):
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                try:
                    text_to_draw = folder_name
                    temp_img = Image.new('RGB', (1, 1))
                    draw_temp = ImageDraw.Draw(temp_img)
                    bbox = draw_temp.textbbox((0, 0), text_to_draw, font=font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]

                    padding = 20
                    img_width = text_width + (2 * padding)
                    img_height = text_height + (2 * padding)
                    
                    image = Image.new('RGB', (img_width, img_height), color='white')
                    draw = ImageDraw.Draw(image)
                    draw.text((padding, padding), text_to_draw, font=font, fill='black')
                    
                    main_folder_path = os.path.join(base_path, folder_name)
                    tipografia_folder = os.path.join(main_folder_path, "tipografia")
                    output_path = os.path.join(tipografia_folder, "firma.jpg")

                    os.makedirs(tipografia_folder, exist_ok=True)
                    
                    image.save(output_path, 'JPEG', quality=95)
                    
                    self.log(f"✅ Firma creada para '{folder_name}' en '{output_path}'")
                    creadas += 1

                except Exception as e:
                    self.log(f"❌ Error creando firma en '{folder_name}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen = f"Proceso de creación de firmas finalizado.\n\n- Firmas creadas: {creadas}\n- Errores: {errores}"
        self.log(f"--- {resumen.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)
    
    def accion_firmar_docx_con_imagen(self):
        base_path = filedialog.askdirectory(
            parent=self.root,
            title="Seleccionar carpeta base con las carpetas a procesar"
        )
        if not base_path:
            self.log("Operación de firma de DOCX cancelada.")
            return

        docx_filename = simpledialog.askstring(
            "Nombre del Documento", "Introduce el nombre del archivo DOCX a firmar:",
            parent=self.root, initialvalue="plantilla.docx"
        )
        if not docx_filename: return

        signature_filename = simpledialog.askstring(
            "Nombre de la Firma", "Introduce el nombre del archivo de imagen de la firma (JPG, PNG):",
            parent=self.root, initialvalue="firma.jpg"
        )
        if not signature_filename: return

        try:
            folders_to_process = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not folders_to_process:
            self.log("No se encontraron carpetas para procesar.")
            return

        self.log(f"--- Iniciando reemplazo de firma en DOCX para {len(folders_to_process)} carpetas ---")
        
        procesados, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Modificando documentos DOCX...", len(folders_to_process))

        try:
            for i, folder_name in enumerate(folders_to_process):
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                folder_path = os.path.join(base_path, folder_name)
                docx_path = os.path.join(folder_path, docx_filename)
                
                signature_path = None
                possible_sig_paths = [
                    os.path.join(folder_path, signature_filename),
                    os.path.join(folder_path, "tipografia", signature_filename)
                ]
                for path in possible_sig_paths:
                    if os.path.exists(path):
                        signature_path = path
                        break

                if not os.path.exists(docx_path):
                    self.log(f"⚠️ DOCX '{docx_filename}' no encontrado en '{folder_name}'. Se omite.")
                    errores += 1
                    continue
                
                if not signature_path:
                    self.log(f"⚠️ Firma '{signature_filename}' no encontrada en '{folder_name}' o sus subcarpetas. Se omite.")
                    errores += 1
                    continue
                
                if self._reemplazar_firma_en_docx_worker(docx_path, signature_path):
                    procesados += 1
                else:
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        resumen = f"Proceso de firma finalizado.\n\n- Documentos DOCX modificados: {procesados}\n- Errores/Omitidos: {errores}"
        self.log(f"--- {resumen.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)
    
    def _reemplazar_firma_en_docx_worker(self, docx_path, signature_path):
        try:
            doc = Document(docx_path)
            anchor_text = "Firma de Aceptacion"
            signature_p_index = -1

            for i, p in enumerate(doc.paragraphs):
                if anchor_text.lower() in p.text.lower():
                    target_index = i + 1
                    if target_index < len(doc.paragraphs):
                        signature_p_index = target_index
                    break
            
            if signature_p_index == -1:
                self.log(f"  ❌ ERROR: No se encontró el ancla '{anchor_text}' o el párrafo de la firma en '{os.path.basename(docx_path)}'.")
                return False

            signature_p = doc.paragraphs[signature_p_index]
            
            p_element = signature_p._p
            p_element.clear_content()

            run = signature_p.add_run()
            try:
                run.add_picture(signature_path, width=Inches(1.5))
                signature_p.alignment = WD_ALIGN_PARAGRAPH.LEFT 

            except FileNotFoundError:
                self.log(f"  ❌ ERROR: El archivo de imagen de firma no fue encontrado: {signature_path}")
                return False
            except Exception as img_e:
                self.log(f"  ❌ ERROR: No se pudo insertar la imagen '{os.path.basename(signature_path)}': {img_e}")
                return False

            doc.save(docx_path)
            self.log(f"  ✅ Firma reemplazada exitosamente en: '{os.path.basename(docx_path)}'")
            return True

        except Exception as e:
            self.log(f"❌ ERROR CRÍTICO al modificar '{os.path.basename(docx_path)}': {e}")
            return False
    
    def exportar_lista(self):
        if not self.source_path.get() or not os.path.isdir(self.source_path.get()):
            messagebox.showwarning("Advertencia", "Selecciona una carpeta de origen válida primero.", parent=self.root)
            return
        
        ruta_guardado = filedialog.asksaveasfilename(parent=self.root, defaultextension=".xlsx", filetypes=[("Archivo Excel", "*.xlsx")], title="Guardar análisis como")
        if not ruta_guardado: return
        
        self.log("Iniciando análisis de carpetas...")
        base_path = self.source_path.get()
        
        detalle_data = []
        resumen_data = []

        top_level_folders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]

        for top_folder in top_level_folders:
            top_folder_path = os.path.join(base_path, top_folder)
            total_files_in_top_folder = 0
            has_any_files = False

            for root, dirs, files in os.walk(top_folder_path):
                total_files_in_top_folder += len(files)
                if files:
                    has_any_files = True
                    for file in files:
                        detalle_data.append({'nombre carpeta': top_folder, 'archivos que contiene': file})
            
            if not has_any_files:
                if not any(os.scandir(top_folder_path)):
                    detalle_data.append({'nombre carpeta': top_folder, 'archivos que contiene': 'carpeta vacia'})
            
            resumen_data.append({'Carpeta': top_folder, 'Cantidad de archivos': total_files_in_top_folder})

        if not detalle_data and not resumen_data:
            self.log("No se encontraron carpetas o archivos para analizar.")
            messagebox.showinfo("Análisis Vacío", "La carpeta seleccionada no contiene subcarpetas para analizar.")
            return

        df_detalle = pd.DataFrame(detalle_data, columns=['nombre carpeta', 'archivos que contiene'])
        df_resumen = pd.DataFrame(resumen_data)
        
        total_row = pd.DataFrame([{'Carpeta': 'TOTAL DE CARPETAS', 'Cantidad de archivos': len(top_level_folders)}])
        df_resumen = pd.concat([df_resumen, total_row], ignore_index=True)

        try:
            with pd.ExcelWriter(ruta_guardado, engine='xlsxwriter') as writer:
                df_detalle.to_excel(writer, sheet_name='Detalle', index=False)
                df_resumen.to_excel(writer, sheet_name='Resumen', index=False)
            
            self.log(f"✅ Análisis de carpetas exportado a: {ruta_guardado}")
            os.startfile(ruta_guardado)
            messagebox.showinfo("Éxito", f"La lista de archivos y carpetas se ha guardado en:\n{ruta_guardado}", parent=self.root)
        except Exception as e:
            self.log(f"❌ Error al exportar a Excel: {e}")
            messagebox.showerror("Error", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)
    
    def _extract_field(self, text, pattern, default="No encontrado"):
        """Busca un patrón regex en un texto y devuelve el primer grupo de captura limpio."""
        match = pattern.search(text)
        if match:
            # Limpiar el texto extraído: quitar espacios al inicio/final y reemplazar saltos de línea y espacios múltiples por uno solo.
            return " ".join(match.group(1).strip().split())
        return default

    def accion_analisis_historia_clinica(self):
        """
        Analiza masivamente los archivos PDF de los resultados de búsqueda para extraer
        datos de historias clínicas y exportarlos a un archivo Excel.
        """
        self.log("--- Iniciando Análisis de Historia Clínica desde PDFs ---")
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar el análisis.", parent=self.root)
            self.log("Análisis cancelado: no hay archivos en los resultados.")
            return

        archivos_pdf = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        if not archivos_pdf:
            messagebox.showinfo("Sin PDFs", "No se encontraron archivos PDF en los resultados de la búsqueda.", parent=self.root)
            self.log("Análisis cancelado: no se encontraron PDFs en los resultados.")
            return

        ruta_guardado = filedialog.asksaveasfilename(
            parent=self.root,
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx")],
            title="Guardar Análisis de Historias Clínicas como"
        )
        if not ruta_guardado:
            self.log("Análisis cancelado por el usuario (no se seleccionó archivo de guardado).")
            return

        # --- INICIO DE LA MODIFICACIÓN ---
        # Regex corregidos para ser más precisos y robustos.
        patterns = {
            # Captura solo hasta el final de la línea para evitar incluir "Fecha".
            'Paciente': re.compile(r"Paciente:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            # Mantiene el patrón original que funciona bien.
            'Estrato': re.compile(r"Estrato:\s*(.*?)\s*Municipio:", re.IGNORECASE | re.DOTALL),
            # Captura solo hasta el final de la línea para evitar incluir "Ocupación".
            'Contrato': re.compile(r"Contrato:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            # Captura el texto entre "DATOS HISTORIA CLÍNICA" y "¿ES VICTIMA...", manejando saltos de línea.
            # También es más flexible con los acentos.
            'HISTORIA CLÍNICA': re.compile(r"DATOS HISTORIA CL[IÍ]NICA\s*(.*?)\s*¿ES V[IÍ]CTIMA DE VIOLENCIA\?", re.IGNORECASE | re.DOTALL)
        }
        
        extracted_data = []
        procesados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Analizando Historias Clínicas...", len(archivos_pdf))
        
        try:
            for i, pdf_path in enumerate(archivos_pdf):
                nombre_archivo = os.path.basename(pdf_path)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    full_text = ""
                    with fitz.open(pdf_path) as doc:
                        for page in doc:
                            full_text += page.get_text("text") + "\n" # Añadir newline para asegurar delimitadores
                    
                    record = {'Archivo': nombre_archivo}
                    for key, pattern in patterns.items():
                        record[key] = self._extract_field(full_text, pattern)

                    extracted_data.append(record)
                    self.log(f"✅ Analizado: {nombre_archivo}")
                    procesados += 1

                except Exception as e:
                    self.log(f"❌ Error al procesar el archivo '{nombre_archivo}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        if not extracted_data:
            self.log("No se pudo extraer información de ningún archivo. No se generó el Excel.")
            messagebox.showinfo("Sin Datos", "No se pudo extraer datos de los archivos PDF seleccionados.", parent=self.root)
            return

        try:
            # Orden de columnas actualizado para coincidir con las claves del diccionario.
            column_order = ['Archivo', 'Paciente', 'Estrato', 'Contrato', 'HISTORIA CLÍNICA']
            df = pd.DataFrame(extracted_data, columns=column_order)
            df.to_excel(ruta_guardado, index=False)
            self.log(f"--- Análisis completado. Procesados: {procesados}, Errores: {errores} ---")
            self.log(f"📊 Datos guardados exitosamente en: {ruta_guardado}")
            messagebox.showinfo("Proceso Completado", f"El análisis ha finalizado.\n\n- Archivos procesados: {procesados}\n- Errores: {errores}\n\nEl archivo Excel ha sido guardado.", parent=self.root)
            os.startfile(ruta_guardado)
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo Excel: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)
        # --- FIN DE LA MODIFICACIÓN ---

    ### NUEVO ###
    def accion_analisis_autorizacion_nueva_eps(self):
        """
        Analiza masivamente archivos PDF de autorizaciones de Nueva EPS desde los resultados 
        de búsqueda, extrae datos clave y los exporta a un archivo Excel.
        """
        self.log("--- Iniciando Análisis de Autorización Nueva EPS desde PDFs ---")
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar el análisis.", parent=self.root)
            self.log("Análisis cancelado: no hay archivos en los resultados.")
            return

        archivos_pdf = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        if not archivos_pdf:
            messagebox.showinfo("Sin PDFs", "No se encontraron archivos PDF en los resultados de la búsqueda.", parent=self.root)
            self.log("Análisis cancelado: no se encontraron PDFs en los resultados.")
            return

        ruta_guardado = filedialog.asksaveasfilename(
            parent=self.root,
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx")],
            title="Guardar Análisis de Autorizaciones Nueva EPS como"
        )
        if not ruta_guardado:
            self.log("Análisis cancelado por el usuario (no se seleccionó archivo de guardado).")
            return

        # Expresiones regulares para cada campo solicitado
        patterns = {
            'Afiliado': re.compile(r"Afiliado:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            'N° Autorización': re.compile(r"N° Autorización:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            'Autorizada el': re.compile(r"Autorizada el:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            'Descripción Servicio': re.compile(r"Descripción Servicio\s*\n\s*\d+\s+\d+\s+(.*?)(?:\n|$)", re.IGNORECASE | re.DOTALL),
            # --- INICIO DE LA MODIFICACIÓN ---
            # Se generaliza el patrón para que capture cualquier línea que comience con "Afiliado Cancela" o "Afiliado No Cancela".
            'Info de Pago': re.compile(r"(Afiliado (?:No )?Cancela.*?)(?:\n|$)", re.IGNORECASE)
            # --- FIN DE LA MODIFICACIÓN ---
        }
        
        extracted_data = []
        procesados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Analizando Autorizaciones...", len(archivos_pdf))
        
        try:
            for i, pdf_path in enumerate(archivos_pdf):
                nombre_archivo = os.path.basename(pdf_path)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    full_text = ""
                    with fitz.open(pdf_path) as doc:
                        for page in doc:
                            full_text += page.get_text("text") + "\n"
                    
                    record = {'Archivo': nombre_archivo}
                    for key, pattern in patterns.items():
                        record[key] = self._extract_field(full_text, pattern)

                    extracted_data.append(record)
                    self.log(f"✅ Analizado: {nombre_archivo}")
                    procesados += 1

                except Exception as e:
                    self.log(f"❌ Error al procesar el archivo '{nombre_archivo}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        if not extracted_data:
            self.log("No se pudo extraer información de ningún archivo. No se generó el Excel.")
            messagebox.showinfo("Sin Datos", "No se pudo extraer datos de los archivos PDF seleccionados.", parent=self.root)
            return

        try:
            column_order = ['Archivo', 'Afiliado', 'N° Autorización', 'Autorizada el', 'Descripción Servicio', 'Info de Pago']
            df = pd.DataFrame(extracted_data, columns=column_order)
            
            df.to_excel(ruta_guardado, index=False)
            
            self.log(f"--- Análisis de autorizaciones completado. Procesados: {procesados}, Errores: {errores} ---")
            self.log(f"📊 Datos guardados exitosamente en: {ruta_guardado}")
            messagebox.showinfo("Proceso Completado", f"El análisis ha finalizado.\n\n- Archivos procesados: {procesados}\n- Errores: {errores}\n\nEl archivo Excel ha sido guardado.", parent=self.root)
            os.startfile(ruta_guardado)
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo Excel: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)
    ### FIN NUEVO ###

    def accion_analisis_cargue_sanitas(self):
        """
        Analiza masivamente archivos PDF del cargue de Sanitas desde los resultados 
        de búsqueda, extrae datos clave (FEOV, Estado, Fecha) y los exporta a un archivo Excel.
        """
        self.log("--- Iniciando Análisis de Cargue Sanitas desde PDFs ---")
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar el análisis.", parent=self.root)
            self.log("Análisis cancelado: no hay archivos en los resultados.")
            return

        archivos_pdf = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        if not archivos_pdf:
            messagebox.showinfo("Sin PDFs", "No se encontraron archivos PDF en los resultados de la búsqueda.", parent=self.root)
            self.log("Análisis cancelado: no se encontraron PDFs en los resultados.")
            return

        ruta_guardado = filedialog.asksaveasfilename(
            parent=self.root,
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx")],
            title="Guardar Análisis de Cargue Sanitas como"
        )
        if not ruta_guardado:
            self.log("Análisis cancelado por el usuario (no se seleccionó archivo de guardado).")
            return

        # Expresiones regulares para los campos resaltados en amarillo en la imagen:
        # 1. Factura (FEOVXXX): Busca la sigla FEOV seguida de números.
        # 2. Estado: Busca la palabra que sigue a "Estado" y está antes del final de la línea o de la fecha.
        # 3. Fecha: Busca la fecha y hora combinada.
        patterns = {
            'Factura (FEOV)': re.compile(r"FEOV(\d+)", re.IGNORECASE),
        #'Estado': re.compile(r"Estado\s*(.*?)(?:\s*-\s*\d{1,2}\s+\w+\s+\d{4})", re.IGNORECASE), # Captura lo que sigue a 'Estado' hasta la fecha o fin de línea
            'Fecha y hora de cargue': re.compile(r"(\d{1,2}\s+\w+\s+\d{4}\s*-\s*\d{1,2}:\d{2})", re.IGNORECASE) # Captura la fecha y hora
        }
        
        extracted_data = []
        procesados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Analizando Cargue Sanitas...", len(archivos_pdf))
        
        try:
            for i, pdf_path in enumerate(archivos_pdf):
                nombre_archivo = os.path.basename(pdf_path)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    full_text = ""
                    with fitz.open(pdf_path) as doc:
                        for page in doc:
                            full_text += page.get_text("text") + "\n"
                    
                    record = {'Archivo': nombre_archivo}
                    for key, pattern in patterns.items():
                        record[key] = self._extract_field(full_text, pattern)

                    # Extraer el valor numérico de la factura si el patrón 'FEOV' lo encuentra
                    factura_match = patterns['Factura (FEOV)'].search(full_text)
                    if factura_match:
                        record['Factura (FEOV)'] = factura_match.group(1) # Guardar solo los números

                    extracted_data.append(record)
                    self.log(f"✅ Analizado: {nombre_archivo}")
                    procesados += 1

                except Exception as e:
                    self.log(f"❌ Error al procesar el archivo '{nombre_archivo}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        if not extracted_data:
            self.log("No se pudo extraer información de ningún archivo. No se generó el Excel.")
            messagebox.showinfo("Sin Datos", "No se pudo extraer datos de los archivos PDF seleccionados.", parent=self.root)
            return

        try:
            # Definir el orden de las columnas para el DataFrame
            column_order = ['Archivo', 'Factura (FEOV)', 'Fecha y hora de cargue']
            df = pd.DataFrame(extracted_data, columns=column_order)
            
            # Limpiar el valor de 'Estado' para quitar posibles espacios extra o saltos de línea
            if 'Estado' in df.columns:
                df['Estado'] = df['Estado'].apply(lambda x: " ".join(str(x).strip().split()) if pd.notna(x) else "No encontrado")
            
            df.to_excel(ruta_guardado, index=False)
            
            self.log(f"--- Análisis de Cargue Sanitas completado. Procesados: {procesados}, Errores: {errores} ---")
            self.log(f"📊 Datos guardados exitosamente en: {ruta_guardado}")
            messagebox.showinfo("Proceso Completado", f"El análisis ha finalizado.\n\n- Archivos procesados: {procesados}\n- Errores: {errores}\n\nEl archivo Excel ha sido guardado.", parent=self.root)
            os.startfile(ruta_guardado)
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo Excel: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)


    def exportar_para_renombrar(self):
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Busca archivos primero para poder exportarlos.", parent=self.root)
            return
        
        ruta_guardado = filedialog.asksaveasfilename(parent=self.root, defaultextension=".xlsx", filetypes=[("Archivo Excel", "*.xlsx")], title="Guardar para renombrar")
        if not ruta_guardado: return
        
        data = []
        for ruta in self.resultados:
            data.append({"Ruta actual": ruta, "Nuevo nombre": os.path.basename(ruta)})
        
        df = pd.DataFrame(data)
        try:
            df.to_excel(ruta_guardado, index=False)
            self.log(f"✅ Lista para renombrar exportada a: {ruta_guardado}")
            messagebox.showinfo("Éxito", f"El archivo para renombrar se ha guardado en:\n{ruta_guardado}\n\nEdita la columna 'Nuevo nombre' y luego usa la opción 'Aplicar renombrado desde Excel'.", parent=self.root)
        except Exception as e:
            self.log(f"❌ Error al exportar: {e}")

    def accion_anadir_sufijo_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida primero.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(
            parent=self.root,
            title="Seleccionar archivo Excel con los sufijos",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja",
                f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root,
                initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas:
                self.log("Operación cancelada o nombre de hoja no válido.")
                return

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_carpeta = simpledialog.askstring(
                "Columna de Carpetas",
                f"Ingrese la LETRA de la columna que contiene los NOMBRES DE CARPETA.\n(Opciones: {letras_rango})",
                parent=self.root
            )
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles:
                self.log(f"Letra de columna para carpetas no válida: '{letra_col_carpeta}'.")
                return
            
            letra_col_sufijo = simpledialog.askstring(
                "Columna de Sufijos",
                f"Ingrese la LETRA de la columna que contiene el SUFIJO (ej: _FEOV12345).\n(Opciones: {letras_rango})",
                parent=self.root
            )
            if not letra_col_sufijo or letra_col_sufijo.upper() not in letras_disponibles:
                self.log(f"Letra de columna para sufijos no válida: '{letra_col_sufijo}'.")
                return

            indice_col_carpeta = ord(letra_col_carpeta.upper()) - ord('A')
            indice_col_sufijo = ord(letra_col_sufijo.upper()) - ord('A')
            
            self.log(f"--- Iniciando proceso de añadir sufijos desde Excel ---")
            carpetas_procesadas, archivos_renombrados, errores = 0, 0, 0
            
            progress_win, progress_bar, progress_label = self._create_progress_window("Añadiendo sufijos...", len(df))
            try:
                for index, fila in df.iterrows():
                    nombre_carpeta = fila.iloc[indice_col_carpeta]
                    sufijo = fila.iloc[indice_col_sufijo]

                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()

                    if pd.isna(nombre_carpeta) or pd.isna(sufijo): continue 

                    nombre_carpeta = str(nombre_carpeta).strip()
                    sufijo = str(sufijo).strip()
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")

                    if not nombre_carpeta or not sufijo: continue 

                    ruta_carpeta = os.path.join(carpeta_origen, nombre_carpeta)

                    if os.path.isdir(ruta_carpeta):
                        self.log(f"Procesando carpeta: '{nombre_carpeta}'")
                        carpetas_procesadas += 1
                        for nombre_archivo in os.listdir(ruta_carpeta):
                            ruta_archivo_actual = os.path.join(ruta_carpeta, nombre_archivo)
                            if os.path.isfile(ruta_archivo_actual):
                                nombre_base, extension = os.path.splitext(nombre_archivo)
                                nuevo_nombre_archivo = f"{nombre_base}{sufijo}{extension}"
                                nueva_ruta_archivo = os.path.join(ruta_carpeta, nuevo_nombre_archivo)

                                if ruta_archivo_actual == nueva_ruta_archivo: continue
                                
                                if os.path.exists(nueva_ruta_archivo):
                                    self.log(f"⚠️  ERROR: El archivo '{nuevo_nombre_archivo}' ya existe. No se renombró '{nombre_archivo}'.")
                                    errores += 1
                                    continue
                                
                                try:
                                    os.rename(ruta_archivo_actual, nueva_ruta_archivo)
                                    self.log(f"  ✅ Renombrado: '{nombre_archivo}' → '{nuevo_nombre_archivo}'")
                                    archivos_renombrados += 1
                                except Exception as e:
                                    self.log(f"  ❌ ERROR al renombrar '{nombre_archivo}': {e}")
                                    errores += 1
                    else:
                        self.log(f"⚠️  ADVERTENCIA: No se encontró la carpeta '{nombre_carpeta}' en la ruta de origen.")
                        errores += 1
            finally:
                if progress_win:
                    progress_win.destroy()
            
            self.log(f"--- Proceso Finalizado. Carpetas procesadas: {carpetas_procesadas}, Archivos renombrados: {archivos_renombrados}, Errores/Advertencias: {errores} ---")
            messagebox.showinfo("Proceso Completado", f"Proceso de renombrado con sufijo finalizado.\n\n- Archivos renombrados: {archivos_renombrados}\n- Carpetas procesadas: {carpetas_procesadas}\n- Errores: {errores}\n\nRevise el registro de actividad para más detalles.", parent=self.root)

        except Exception as e:
            self.log(f"❌ Ocurrió un error general: {e}")
            messagebox.showerror("Error", f"Ocurrió un error inesperado al procesar el archivo Excel:\n{e}", parent=self.root)

    def accion_autorizacion_docx_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona la carpeta de origen que contiene todas las subcarpetas.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar archivo Excel con los datos", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring("Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}", parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else "")
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_carpeta = simpledialog.askstring("Columna de Carpetas", f"Ingrese la LETRA de la columna que contiene los NOMBRES DE CARPETA.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return
            
            letra_col_auth = simpledialog.askstring("Columna de Autorización", f"Ingrese la LETRA de la columna que contiene el nuevo NÚMERO DE AUTORIZACIÓN.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_auth or letra_col_auth.upper() not in letras_disponibles: return

            indice_col_carpeta = ord(letra_col_carpeta.upper()) - ord('A')
            indice_col_auth = ord(letra_col_auth.upper()) - ord('A')
            
            self.log(f"--- Iniciando modificación de Autorización en DOCX (por carpeta) ---")
            
            modificados, errores_carpeta, errores_docx, errores_proceso = 0, 0, 0, 0
            docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)

            progress_win, progress_bar, progress_label = self._create_progress_window("Modificando DOCX...", len(df))
            try:
                for index, fila in df.iterrows():
                    nombre_carpeta = fila.iloc[indice_col_carpeta]
                    nueva_autorizacion = fila.iloc[indice_col_auth]
                    
                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()
                    
                    if pd.isna(nombre_carpeta) or pd.isna(nueva_autorizacion): continue

                    nombre_carpeta = str(nombre_carpeta).strip()
                    nueva_autorizacion = str(int(nueva_autorizacion)) if isinstance(nueva_autorizacion, (float, int)) else str(nueva_autorizacion).strip()
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")

                    if not nombre_carpeta or not nueva_autorizacion: continue
                    
                    ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
                    if not os.path.isdir(ruta_carpeta_especifica):
                        self.log(f"⚠️ ADVERTENCIA: No se encontró la carpeta '{nombre_carpeta}' en la ruta de origen.")
                        errores_carpeta += 1
                        continue
                    
                    ruta_docx_encontrada = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if docx_pattern.match(f)), None)

                    if not ruta_docx_encontrada:
                        self.log(f"❌ ERROR: No se encontró un archivo DOCX con el patrón 'CRC_..._FEOV.docx' dentro de la carpeta '{nombre_carpeta}'.")
                        errores_docx += 1
                        continue

                    try:
                        doc = Document(ruta_docx_encontrada)
                        fue_modificado = False
                        for p in doc.paragraphs:
                            if "AUTORIZACION:" in p.text.upper():
                                p.text = re.sub(r'(AUTORIZACION:)\s*.*', r'\1 ' + str(nueva_autorizacion), p.text, flags=re.IGNORECASE)
                                fue_modificado = True
                                break
                        
                        if fue_modificado:
                            doc.save(ruta_docx_encontrada)
                            self.log(f"✅ Modificado '{os.path.basename(ruta_docx_encontrada)}' en carpeta '{nombre_carpeta}' con autorización: '{nueva_autorizacion}'")
                            modificados += 1
                        else:
                            self.log(f"⚠️ ADVERTENCIA: Se encontró el archivo '{os.path.basename(ruta_docx_encontrada)}' pero no la línea 'AUTORIZACION:'.")
                            errores_proceso += 1
                    except Exception as e:
                        self.log(f"❌ ERROR CRÍTICO al procesar el archivo '{os.path.basename(ruta_docx_encontrada)}': {e}")
                        errores_proceso += 1
            finally:
                if progress_win:
                    progress_win.destroy()

            resumen_msg = (f"Proceso de modificación de autorización finalizado.\n\n" f"- Modificados: {modificados}\n" f"- Carpetas no encontradas: {errores_carpeta}\n" f"- DOCX no encontrados: {errores_docx}\n" f"- Errores de procesamiento: {errores_proceso}")
            self.log(f"--- Proceso Finalizado. Modificados: {modificados}, Errores: {errores_carpeta}/{errores_docx}/{errores_proceso} ---")
            messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
        except Exception as e:
            self.log(f"❌ Ocurrió un error general: {e}")
            messagebox.showerror("Error", f"Ocurrió un error inesperado al procesar el archivo Excel:\n{e}", parent=self.root)

    def accion_regimen_docx_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona la carpeta de origen que contiene todas las subcarpetas.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar archivo Excel con los datos", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring("Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}", parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else "")
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_carpeta = simpledialog.askstring("Columna de Carpetas", f"Ingrese la LETRA de la columna que contiene los NOMBRES DE CARPETA.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return
            
            letra_col_regimen = simpledialog.askstring("Columna de Régimen", f"Ingrese la LETRA de la columna que contiene el nuevo RÉGIMEN (ej: CONTRIBUTIVO).\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_regimen or letra_col_regimen.upper() not in letras_disponibles: return

            indice_col_carpeta = ord(letra_col_carpeta.upper()) - ord('A')
            indice_col_regimen = ord(letra_col_regimen.upper()) - ord('A')
            
            self.log(f"--- Iniciando modificación de Régimen en DOCX (por carpeta) ---")
            
            modificados, errores_carpeta, errores_docx, errores_proceso = 0, 0, 0, 0
            docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)

            progress_win, progress_bar, progress_label = self._create_progress_window("Modificando DOCX...", len(df))
            try:
                for index, fila in df.iterrows():
                    nombre_carpeta = fila.iloc[indice_col_carpeta]
                    nuevo_regimen = fila.iloc[indice_col_regimen]
                    
                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()
                    
                    if pd.isna(nombre_carpeta) or pd.isna(nuevo_regimen): continue

                    nombre_carpeta = str(nombre_carpeta).strip()
                    nuevo_regimen = str(nuevo_regimen).strip()
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")

                    if not nombre_carpeta or not nuevo_regimen: continue
                    
                    ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
                    if not os.path.isdir(ruta_carpeta_especifica):
                        self.log(f"⚠️ ADVERTENCIA: No se encontró la carpeta '{nombre_carpeta}' en la ruta de origen.")
                        errores_carpeta += 1
                        continue
                    
                    ruta_docx_encontrada = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if docx_pattern.match(f)), None)

                    if not ruta_docx_encontrada:
                        self.log(f"❌ ERROR: No se encontró un archivo DOCX con el patrón 'CRC_..._FEOV.docx' dentro de la carpeta '{nombre_carpeta}'.")
                        errores_docx += 1
                        continue

                    try:
                        doc = Document(ruta_docx_encontrada)
                        fue_modificado = False
                        for p in doc.paragraphs:
                            if "REGIMEN:" in p.text.upper():
                                p.text = re.sub(r'(REGIMEN:)\s*.*', r'\1 ' + str(nuevo_regimen), p.text, flags=re.IGNORECASE)
                                fue_modificado = True
                                break
                        
                        if fue_modificado:
                            doc.save(ruta_docx_encontrada)
                            self.log(f"✅ Modificado '{os.path.basename(ruta_docx_encontrada)}' en carpeta '{nombre_carpeta}' con régimen: '{nuevo_regimen}'")
                            modificados += 1
                        else:
                            self.log(f"⚠️ ADVERTENCIA: Se encontró el archivo '{os.path.basename(ruta_docx_encontrada)}' pero no la línea 'REGIMEN:'.")
                            errores_proceso += 1

                    except Exception as e:
                        self.log(f"❌ ERROR CRÍTICO al procesar el archivo '{os.path.basename(ruta_docx_encontrada)}': {e}")
                        errores_proceso += 1
            finally:
                if progress_win:
                    progress_win.destroy()

            resumen_msg = (f"Proceso de modificación de régimen finalizado.\n\n" f"- Modificados: {modificados}\n" f"- Carpetas no encontradas: {errores_carpeta}\n" f"- DOCX no encontrados: {errores_docx}\n" f"- Errores de procesamiento: {errores_proceso}")
            self.log(f"--- Proceso Finalizado. Modificados: {modificados}, Errores: {errores_carpeta}/{errores_docx}/{errores_proceso} ---")
            messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
        except Exception as e:
            self.log(f"❌ Ocurrió un error general: {e}")
            messagebox.showerror("Error", f"Ocurrió un error inesperado al procesar el archivo Excel:\n{e}", parent=self.root)
    
    def accion_modificar_docx_completo_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida que contenga las subcarpetas de los pacientes.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar archivo Excel con los datos", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring("Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}", parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else "")
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return
            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
        except Exception as e:
            self.log(f"❌ Error al leer el archivo Excel: {e}")
            messagebox.showerror("Error de Lectura", f"No se pudo leer el archivo o la hoja de Excel:\n{e}", parent=self.root)
            return
        
        column_map, missing_cols = self._create_column_map_from_headers(df)
        
        if not column_map:
            error_msg = ("El archivo Excel no tiene el formato correcto.\n\n"
                         "Faltan las siguientes columnas obligatorias:\n- " + 
                         "\n- ".join(missing_cols))
            self.log(f"❌ Error de formato en Excel: {error_msg.replace(chr(10), ' ')}")
            messagebox.showerror("Error en Encabezados", error_msg, parent=self.root)
            return

        self.log(f"--- Iniciando modificación completa de DOCX desde Excel (Modo Automático) ---")
        self.log(f"Columnas encontradas y mapeadas correctamente desde la hoja '{hoja_seleccionada}'.")
        modificados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Modificando DOCX...", len(df))
        try:
            for index, row in df.iterrows():
                progress_bar['value'] = index + 1
                datos = {}
                try:
                    datos = {key: str(row[col_name]).strip() if pd.notna(row[col_name]) else "" for key, col_name in column_map.items()}
                    nombre_carpeta = datos.get('folder')

                    if not nombre_carpeta:
                        self.log(f"⚠️ Fila {index+2} omitida: El nombre de la carpeta está vacío.")
                        continue
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")
                    self.root.update_idletasks()

                    ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
                    if not os.path.isdir(ruta_carpeta_especifica):
                        self.log(f"⚠️ No se encontró la carpeta '{nombre_carpeta}' (Fila {index+2}).")
                        errores += 1
                        continue
                    
                    ruta_docx = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if f.lower().endswith('.docx') and 'plantilla' in f.lower()), None)
                    if not ruta_docx:
                        self.log(f"❌ No se encontró archivo DOCX con 'plantilla' en el nombre dentro de la carpeta '{nombre_carpeta}'.")
                        errores += 1
                        continue
                    
                    doc = Document(ruta_docx)
                    
                    for p in doc.paragraphs:
                        if "Santiago de Cali, " in p.text:
                            p.text =(f"Santiago de Cali,  {datos['date']}")
                            continue

                        if "Yo " in p.text and "identificado con" in p.text:
                            p.text = (f"Yo {datos['full_name']} identificado con {datos['doc_type']}, "
                                      f"Numero {datos['doc_num']} en calidad de paciente, doy fé y acepto el "
                                      f"servicio de {datos['service']} brindado por la IPS OPORTUNIDAD DE VIDA S.A.S"
                                      f"   "
                                      f"   ")
                            continue

                        if "EPS:" in p.text:
                            p.text = re.sub(r'(EPS:)\s*.*', r'\1 ' + datos['eps'], p.text, count=1)
                        if "TIPO SERVICIO:" in p.text:
                            p.text = re.sub(r'(TIPO SERVICIO:)\s*.*', r'\1 ' + datos['tipo_servicio'], p.text, count=1)
                        if "REGIMEN:" in p.text:
                            p.text = re.sub(r'(REGIMEN:)\s*.*', r'\1 ' + datos['regimen'], p.text, count=1)
                        if "CATEGORIA:" in p.text:
                            p.text = re.sub(r'(CATEGORIA:)\s*.*', r'\1 ' + datos['categoria'], p.text, count=1)
                        if "VALOR CUOTA MODERADORA:" in p.text:
                            p.text = re.sub(r'(VALOR CUOTA MODERADORA:)\s*.*', r'\1 ' + datos['cuota'], p.text, count=1)
                        if "AUTORIZACION:" in p.text:
                            p.text = re.sub(r'(AUTORIZACION:)\s*.*', r'\1 ' + datos['auth'], p.text, count=1)
                        if "Fecha de Atención:" in p.text:
                            p.text = re.sub(r'(Fecha de Atención:)\s*.*', r'\1 ' + datos['fecha_atencion'], p.text, count=1)
                        if "Fecha de Finalización:" in p.text:
                            p.text = re.sub(r'(Fecha de Finalización:)\s*.*', r'\1 ' + datos['fecha_fin'], p.text, count=1)
                            continue

                    signature_line_index = -1
                    for i, p in enumerate(doc.paragraphs):
                        if "FIRMA DE ACEPTACION" in p.text.upper():
                            signature_line_index = i
                            break 

                    if signature_line_index != -1 and signature_line_index + 2 < len(doc.paragraphs):
                        name_paragraph = doc.paragraphs[signature_line_index + 2]
                        name_paragraph.text = datos['full_name'].upper()
                    else:
                        self.log(f"⚠️ No se pudo encontrar la posición del nombre para la firma en '{os.path.basename(ruta_docx)}'.")

                    doc.save(ruta_docx)
                    self.log(f"✅ Modificado: {os.path.basename(ruta_docx)} en carpeta '{nombre_carpeta}'")
                    modificados += 1

                except Exception as e:
                    self.log(f"❌ ERROR CRÍTICO al procesar la fila {index+2} (Carpeta: {datos.get('folder', 'N/A')}): {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (f"Proceso de modificación de DOCX finalizado.\n\n" f"- Archivos modificados: {modificados}\n" f"- Errores/Advertencias: {errores}")
        self.log(f"--- Proceso Finalizado. Modificados: {modificados}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
    
    def _create_column_map_from_headers(self, df):
        """
        Valida que el DataFrame contenga los encabezados necesarios y devuelve el mapa de columnas.
        Si faltan columnas, devuelve None y una lista de las columnas faltantes.
        """
        required_map = {
            'folder': 'Nombre Carpeta',
            'date': 'Ciudad y Fecha',
            'full_name': 'Nombre Completo',
            'doc_type': 'Tipo Documento',
            'doc_num': 'Numero Documento',
            'service': 'Servicio',
            'eps': 'EPS',
            'tipo_servicio': 'Tipo Servicio',
            'regimen': 'Regimen',
            'categoria': 'Categoria',
            'cuota': 'Valor Cuota Moderadora',
            'auth': 'Numero Autorizacion',
            'fecha_atencion': 'Fecha y Hora Atencion',
            'fecha_fin': 'Fecha Finalizacion'
        }

        excel_headers = df.columns
        missing_cols = [header for header in required_map.values() if header not in excel_headers]

        if missing_cols:
            return None, missing_cols

        return required_map, []

    #================================================================================#
    #========== INICIO DE LA SECCIÓN MODIFICADA: VISOR Y EDITOR JSON / XML ==========
    #================================================================================#
    
    def _setup_data_editor_tab(self, tab):
        """Configura la interfaz de usuario para la pestaña del editor de datos."""
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(2, weight=1)

        # --- Frame para controles de archivo ---
        top_frame = ttk.Frame(tab, padding=(10, 10))
        top_frame.grid(row=0, column=0, sticky="ew")
        
        ttk.Button(top_frame, text="📂 Abrir Archivo (JSON/XML)", command=self._open_data_file).pack(side="left", padx=(0, 10))
        self.save_button = ttk.Button(top_frame, text="💾 Guardar Cambios", command=self._save_data_file, state="disabled")
        self.save_button.pack(side="left")
        
        # --- Frame para búsqueda ---
        search_frame = ttk.Frame(tab, padding=(10, 0, 10, 10))
        search_frame.grid(row=1, column=0, sticky="ew")
        ttk.Label(search_frame, text="Buscar:").pack(side="left")
        search_entry = ttk.Entry(search_frame, textvariable=self.search_data_var)
        search_entry.pack(side="left", fill="x", expand=True, padx=5)
        search_entry.bind("<Return>", self._search_in_tree)
        ttk.Button(search_frame, text="Buscar y Resaltar", command=self._search_in_tree).pack(side="left")
        ttk.Button(search_frame, text="Limpiar", command=self._clear_search_highlight).pack(side="left", padx=5)

        # --- Frame para el Treeview ---
        tree_frame = ttk.LabelFrame(tab, text="Estructura del Archivo", padding=10)
        tree_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.data_tree = ttk.Treeview(tree_frame, columns=("value",), show="tree headings")
        self.data_tree.heading("#0", text="Clave / Tag")
        self.data_tree.heading("value", text="Valor")
        self.data_tree.column("#0", width=350, stretch=tk.NO)
        self.data_tree.column("value", width=500, stretch=tk.YES)
        
        ysb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.data_tree.yview)
        xsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.data_tree.xview)
        self.data_tree.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        
        self.data_tree.grid(row=0, column=0, sticky="nsew")
        ysb.grid(row=0, column=1, sticky="ns")
        xsb.grid(row=1, column=0, sticky="ew")

        self._configure_tree_tags_and_menu()
        self.data_tree.bind("<Double-1>", self._on_tree_double_click)
        self.data_tree.bind("<Button-3>", self._show_context_menu)
        self.data_tree.bind("<Motion>", self._on_tree_motion)
        self.data_tree.bind("<Leave>", self._on_tree_leave)

    def _configure_tree_tags_and_menu(self):
        """Configura los estilos de color para el Treeview y el menú contextual."""
        self.data_tree.tag_configure('string', foreground='green')
        self.data_tree.tag_configure('number', foreground='blue')
        self.data_tree.tag_configure('boolean', foreground='purple')
        self.data_tree.tag_configure('null', foreground='gray')
        self.data_tree.tag_configure('dict', foreground='#a52a2a') # Brown
        self.data_tree.tag_configure('list', foreground='#a52a2a')
        self.data_tree.tag_configure('xml_tag', foreground='#00008B') # DarkBlue
        self.data_tree.tag_configure('xml_attr', foreground='#8B008B') # DarkMagenta
        
        self.data_tree.tag_configure('search_hit', background='yellow')

        self.data_context_menu = tk.Menu(self.root, tearoff=0)

    def _open_data_file(self):
        path = filedialog.askopenfilename(
            parent=self.root,
            title="Seleccionar archivo JSON o XML",
            filetypes=[
                ("Todos los soportados", "*.json *.xml"),
                ("Archivos JSON", "*.json"),
                ("Archivos XML", "*.xml"),
                ("Todos los archivos", "*.*")
            ]
        )
        if not path:
            return

        self._reset_editor_state()
        self.root.title(f"📂 Organizador de Archivos - Editando: {os.path.basename(path)}")
        
        try:
            if path.lower().endswith(".json"):
                self.current_data_type = 'json'
                with open(path, 'r', encoding='utf-8') as f:
                    self.parsed_data = json.load(f)
                self._populate_tree_from_json(self.parsed_data, '')
                self.log(f"✅ Archivo JSON '{os.path.basename(path)}' cargado.")
            
            elif path.lower().endswith(".xml"):
                self.current_data_type = 'xml'
                tree = ET.parse(path)
                self.parsed_data = tree.getroot()
                self._populate_tree_from_xml(self.parsed_data, '')
                self.log(f"✅ Archivo XML '{os.path.basename(path)}' cargado.")
            else:
                raise ValueError("Tipo de archivo no soportado.")
            self.current_file_path = path
        
        except Exception as e:
            self.log(f"❌ Error abriendo o procesando el archivo: {e}")
            messagebox.showerror("Error de Archivo", f"No se pudo cargar o procesar el archivo.\n\n{e}", parent=self.root)
            self._reset_editor_state()

    def _reset_editor_state(self):
        if self.data_tree:
            self.data_tree.delete(*self.data_tree.get_children())
        self.current_file_path = None
        self.parsed_data = None
        self.current_data_type = None
        self.is_data_modified = False
        self.embedded_xml_docs.clear()
        self.xml_element_map.clear()
        self.search_data_var.set("")
        self._clear_search_highlight()
        if self.save_button:
            self.save_button.config(state="disabled")
        self.root.title("📂 Organizador de Archivos v5.7 (Mapeo Parcial)")
    
    def _clean_tag(self, tag):
        """Elimina el namespace {....} de una etiqueta XML para hacerla más legible."""
        return re.sub(r'\{.*?\}', '', tag)
        
    def _get_value_tag(self, value):
        """Devuelve el nombre del tag de color según el tipo de dato."""
        if isinstance(value, str): return 'string'
        if isinstance(value, (int, float)): return 'number'
        if isinstance(value, bool): return 'boolean'
        if value is None: return 'null'
        return ''

    def _format_value(self, value):
        """Formatea un valor para mostrarlo en el Treeview (ej: añade comillas a strings)."""
        if isinstance(value, str):
            return f'"{value}"' # Añadir comillas a las cadenas de texto
        return value # Devolver otros tipos (números, booleanos, None) tal cual

    def _populate_tree_from_xml(self, element, parent_iid):
        tag = self._clean_tag(element.tag)
        
        if element.text and element.text.strip().startswith('<?xml'):
            node_id = self.data_tree.insert(parent_iid, 'end', text=tag, open=True, tags=('xml_tag',))
            self.xml_element_map[node_id] = element
            try:
                cdata_node_id = self.data_tree.insert(node_id, 'end', text="[Contenido XML Adjunto]", open=True, tags=('xml_tag',))
                inner_root = ET.fromstring(element.text.strip())
                self.embedded_xml_docs[cdata_node_id] = (element, inner_root)
                self.xml_element_map[cdata_node_id] = inner_root
                self._populate_tree_from_xml(inner_root, cdata_node_id)
            except ET.ParseError as e:
                err_node_id = self.data_tree.insert(node_id, 'end', text="[Error de parseo]", values=(str(e),), tags=('null',))
                self.xml_element_map[err_node_id] = None
            return

        value_text = element.text.strip() if element.text and element.text.strip() else ""
        node_id = self.data_tree.insert(parent_iid, 'end', text=tag, values=(value_text,), tags=('xml_tag', self._get_value_tag(value_text)))
        self.xml_element_map[node_id] = element

        if element.attrib:
            self.data_tree.item(node_id, open=True)
            attr_parent_id = self.data_tree.insert(node_id, 'end', text="@Atributos", open=True, tags=('xml_attr',))
            self.xml_element_map[attr_parent_id] = element
            for key, value in element.attrib.items():
                attr_node_id = self.data_tree.insert(attr_parent_id, 'end', text=f"@{key}", values=(value,), tags=('xml_attr', self._get_value_tag(value)))
                self.xml_element_map[attr_node_id] = (element, key)

        for child in element:
            self._populate_tree_from_xml(child, node_id)

    def _populate_tree_from_json(self, data, parent_iid):
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, dict):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"{{}} {key}", values=('{...}'), tags=('dict',))
                    self._populate_tree_from_json(value, node)
                elif isinstance(value, list):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"[] {key}", values=('[...]'), tags=('list',))
                    self._populate_tree_from_json(value, node)
                else:
                    self.data_tree.insert(parent_iid, 'end', text=key, values=(self._format_value(value),), tags=(self._get_value_tag(value),))
        elif isinstance(data, list):
            for index, item in enumerate(data):
                key = f"[{index}]"
                if isinstance(item, dict):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"{{}} {key}", values=('{...}'), tags=('dict',))
                    self._populate_tree_from_json(item, node)
                elif isinstance(item, list):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"[] {key}", values=('[...]'), tags=('list',))
                    self._populate_tree_from_json(item, node)
                else:
                    self.data_tree.insert(parent_iid, 'end', text=key, values=(self._format_value(item),), tags=(self._get_value_tag(item),))
    
    def _get_node_path(self, item_id):
        path = []
        parent = item_id
        while parent:
            text = self.data_tree.item(parent, "text")
            if text.startswith('{} ') or text.startswith('[] '):
                text = text[3:]
            path.insert(0, text)
            parent = self.data_tree.parent(parent)
        return path

    def _on_tree_double_click(self, event):
        item_id = self.data_tree.focus()
        if not item_id or self.data_tree.get_children(item_id):
            is_parent_of_attributes = self.data_tree.item(item_id, "text") == "@Atributos"
            if not is_parent_of_attributes:
                return

        if self.current_data_type == 'json':
            path = self._get_node_path(item_id)
            raw_value = self.data_tree.item(item_id, "values")[0]
            if isinstance(raw_value, str) and raw_value.startswith('"') and raw_value.endswith('"'):
                current_value = raw_value[1:-1]
            else:
                current_value = raw_value

            new_value_str = simpledialog.askstring("Editar Valor JSON", f"Nuevo valor para '{path[-1]}':", initialvalue=current_value, parent=self.root)

            if new_value_str is not None:
                try:
                    if new_value_str.lower() == 'true': typed_value = True
                    elif new_value_str.lower() == 'false': typed_value = False
                    elif new_value_str.lower() == 'null': typed_value = None
                    else:
                        try: typed_value = int(new_value_str)
                        except ValueError:
                            try: typed_value = float(new_value_str)
                            except ValueError: typed_value = new_value_str
                    
                    data_node = self.parsed_data
                    for i, key_or_index_str in enumerate(path):
                        if i == len(path) - 1:
                            if isinstance(data_node, dict):
                                data_node[key_or_index_str] = typed_value
                            elif isinstance(data_node, list):
                                index = int(re.search(r'\[(\d+)\]', key_or_index_str).group(1))
                                data_node[index] = typed_value
                        else:
                            if isinstance(data_node, dict):
                                data_node = data_node[key_or_index_str]
                            elif isinstance(data_node, list):
                                index = int(re.search(r'\[(\d+)\]', key_or_index_str).group(1))
                                data_node = data_node[index]
                    
                    self.data_tree.item(item_id, values=(self._format_value(typed_value),), tags=(self._get_value_tag(typed_value),))
                    self.log(f"📝 Valor JSON actualizado para '{'/'.join(path)}'. No olvides guardar.")
                    self.is_data_modified = True
                    self.save_button.config(state="normal")
                except Exception as e:
                    self.log(f"❌ Error al actualizar el valor JSON: {e}")

        elif self.current_data_type == 'xml':
            target = self.xml_element_map.get(item_id)
            if target is None:
                return

            current_value = self.data_tree.item(item_id, "values")[0]
            
            if isinstance(target, tuple):
                element, key = target
                new_value = simpledialog.askstring(f"Editar Atributo XML", f"Nuevo valor para el atributo '{key}':", initialvalue=current_value, parent=self.root)
                if new_value is not None:
                    element.set(key, new_value)
                    self.data_tree.item(item_id, values=(new_value,), tags=('xml_attr', self._get_value_tag(new_value)))
                    self.log(f"📝 Atributo '{key}' actualizado. No olvides guardar.")
                    self.is_data_modified = True
                    self.save_button.config(state="normal")

            elif isinstance(target, ET.Element):
                element = target
                new_value = simpledialog.askstring(f"Editar Valor de Tag XML", f"Nuevo valor para <{self._clean_tag(element.tag)}>:", initialvalue=current_value, parent=self.root)
                if new_value is not None:
                    element.text = new_value
                    self.data_tree.item(item_id, values=(new_value,), tags=('xml_tag', self._get_value_tag(new_value)))
                    self.log(f"📝 Valor del tag '<{self._clean_tag(element.tag)}>' actualizado. No olvides guardar.")
                    self.is_data_modified = True
                    self.save_button.config(state="normal")
            
    def _save_data_file(self):
        if not self.is_data_modified or not self.current_file_path:
            return
        
        warning_msg = f"¿Guardar cambios en '{os.path.basename(self.current_file_path)}'?"
        if self.current_data_type == 'xml':
            warning_msg += "\n\nADVERTENCIA: Guardar un XML firmado lo INVALIDARÁ digitalmente. ¿Continuar?"

        if not messagebox.askyesno("Confirmar Guardado", warning_msg, parent=self.root):
            return

        try:
            if self.current_data_type == 'xml':
                for parent_element, inner_root in self.embedded_xml_docs.values():
                    inner_xml_string = ET.tostring(inner_root, encoding='unicode', method='xml')
                    parent_element.text = f'{inner_xml_string}'
                tree = ET.ElementTree(self.parsed_data)
                tree.write(self.current_file_path, encoding='utf-8', xml_declaration=True)
            
            elif self.current_data_type == 'json':
                with open(self.current_file_path, 'w', encoding='utf-8') as f:
                    json.dump(self.parsed_data, f, indent=4, ensure_ascii=False)

            self.log(f"💾 Cambios guardados en '{os.path.basename(self.current_file_path)}'.")
            self.is_data_modified = False
            self.save_button.config(state="disabled")
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo: {e}")

    def _show_context_menu(self, event):
        """Muestra el menú contextual en la posición del cursor."""
        self.data_context_menu.delete(0, 'end')
        item_id = self.data_tree.identify_row(event.y)
        if not item_id: return
        
        self.data_tree.selection_set(item_id)
        
        item_text = self.data_tree.item(item_id, 'text')
        is_parent = bool(self.data_tree.get_children(item_id))
        is_attribute_parent = item_text == '@Atributos'
        is_root = not self.data_tree.parent(item_id)
        
        if self.pyperclip_available:
            self.data_context_menu.add_command(label="Copiar Clave/Tag", command=lambda: self._copy_to_clipboard(item_id, 'key'))
            if not is_parent or (is_parent and self.current_data_type == 'xml' and not is_attribute_parent):
                 self.data_context_menu.add_command(label="Copiar Valor", command=lambda: self._copy_to_clipboard(item_id, 'value'))
            self.data_context_menu.add_command(label="Copiar Ruta", command=lambda: self._copy_to_clipboard(item_id, 'path'))
            self.data_context_menu.add_separator()

        if not is_parent and not is_attribute_parent:
             self.data_context_menu.add_command(label="Editar Valor...", command=lambda: self._on_tree_double_click(None))
        
        self.data_context_menu.add_command(label="Añadir/Eliminar (En desarrollo)", state='disabled')

        if not is_root and not is_attribute_parent:
            self.data_context_menu.add_command(label="Eliminar", command=lambda: self._delete_item(item_id), foreground='red')

        if is_parent:
            self.data_context_menu.add_separator()
            self.data_context_menu.add_command(label="Expandir Todo", command=lambda: self._toggle_expand(item_id, True))
            self.data_context_menu.add_command(label="Contraer Todo", command=lambda: self._toggle_expand(item_id, False))
        
        self.data_context_menu.post(event.x_root, event.y_root)

    def _copy_to_clipboard(self, item_id, what):
        """Copia la información seleccionada al portapapeles."""
        if not self.pyperclip_available:
            self.log("La función de copiar no está disponible. Instala 'pyperclip'.")
            return
        try:
            if what == 'key':
                text = self.data_tree.item(item_id, 'text')
                if text.startswith('{} ') or text.startswith('[] '):
                    text = text[3:]
                self.pyperclip.copy(text)
                self.log(f"📋 Clave '{text}' copiada al portapapeles.")
            elif what == 'value':
                val_tuple = self.data_tree.item(item_id, 'values')
                if val_tuple:
                    val = str(val_tuple[0])
                    if val.startswith('"') and val.endswith('"'):
                        val = val[1:-1]
                    self.pyperclip.copy(val)
                    self.log(f"📋 Valor '{val}' copiado al portapapeles.")
            elif what == 'path':
                path_list = self._get_node_path(item_id)
                path_str = ".".join(p.replace(" ", "") for p in path_list)
                self.pyperclip.copy(path_str)
                self.log(f"📋 Ruta '{path_str}' copiada al portapapeles.")
        except Exception as e:
            self.log(f"❌ Error al copiar al portapapeles: {e}")
            messagebox.showerror("Error de Portapapeles", f"No se pudo copiar el texto.\n{e}", parent=self.root)

    def _toggle_expand(self, item_id, expand=True):
        """Expande o contrae recursivamente todos los hijos de un nodo."""
        self.data_tree.item(item_id, open=expand)
        for child_id in self.data_tree.get_children(item_id):
            self._toggle_expand(child_id, expand)

    def _delete_item(self, item_id):
        """Elimina un item tanto del modelo de datos como del Treeview."""
        if not messagebox.askyesno("Confirmar Eliminación", f"¿Seguro que quieres eliminar el elemento '{self.data_tree.item(item_id, 'text')}' y todo su contenido?", parent=self.root):
            return
        
        self.log(f"🔴 Funcionalidad de eliminación pendiente de implementación completa en el modelo de datos.")
        messagebox.showinfo("En Desarrollo", "La eliminación directa de nodos está en desarrollo.\n\nComo alternativa, puedes editar el valor a 'null' (JSON) o un texto vacío y guardar el archivo.", parent=self.root)
    
    def _search_in_tree(self, event=None):
        self._clear_search_highlight()
        search_term = self.search_data_var.get().lower()
        if not search_term: return

        self.search_results_iids = []
        self._collect_search_results(search_term, '')
        
        if not self.search_results_iids:
            self.log(f"Búsqueda: No se encontraron resultados para '{search_term}'.")
            return

        self.log(f"Búsqueda: Se encontraron {len(self.search_results_iids)} coincidencias para '{search_term}'.")
        
        first_item = None
        for item_id in self.search_results_iids:
            if not first_item:
                first_item = item_id
            
            current_tags = list(self.data_tree.item(item_id, 'tags'))
            if 'search_hit' not in current_tags:
                current_tags.append('search_hit')
                self.data_tree.item(item_id, tags=tuple(current_tags))

        if first_item:
            self._expand_to_item(first_item)
            self.data_tree.selection_set(first_item)
            self.data_tree.focus(first_item)
            self.data_tree.see(first_item)

    def _clear_search_highlight(self):
        """Elimina el resaltado de búsqueda de todos los items."""
        if not self.search_results_iids: return
            
        for item_id in self.search_results_iids:
            current_tags = list(self.data_tree.item(item_id, 'tags'))
            if 'search_hit' in current_tags:
                current_tags.remove('search_hit')
                self.data_tree.item(item_id, tags=tuple(current_tags))
        self.search_results_iids = []
        self.log("Resaltado de búsqueda limpiado.")

    def _collect_search_results(self, term, parent_iid):
        """Busca recursivamente y añade IIDs a la lista de resultados."""
        for item_id in self.data_tree.get_children(parent_iid):
            item_text = self.data_tree.item(item_id, "text").lower()
            values_tuple = self.data_tree.item(item_id, "values")
            item_value = str(values_tuple[0]).lower() if values_tuple else ""
            
            if term in item_text or term in item_value:
                self.search_results_iids.append(item_id)
            
            if self.data_tree.get_children(item_id):
                self._collect_search_results(term, item_id)
            
    def _expand_to_item(self, item_id):
        """Expande todos los nodos padre de un item para hacerlo visible."""
        parent = self.data_tree.parent(item_id)
        while parent:
            self.data_tree.item(parent, open=True)
            parent = self.data_tree.parent(parent)
            
    def _on_tree_motion(self, event):
        """Muestra un tooltip si el texto del valor es más largo que la columna."""
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

        item_id = self.data_tree.identify_row(event.y)
        if not item_id or self.data_tree.identify_column(event.x) != "#2": # #2 es la columna 'value'
            return
        
        try:
            values_tuple = self.data_tree.item(item_id, "values")
            if not values_tuple or not values_tuple[0]: return
            
            full_text = str(values_tuple[0])
            
            bbox = self.data_tree.bbox(item_id, column="value")
            if not bbox: return

            column_width = self.data_tree.column("value", "width")
            
            # Simple check if the text is likely wider than the column
            if len(full_text) * 7 > column_width: # Estimación, 7px por caracter
                self.tooltip_window = tk.Toplevel(self.root)
                self.tooltip_window.wm_overrideredirect(True)
                self.tooltip_window.wm_geometry(f"+{event.x_root + 15}+{event.y_root + 10}")
                
                label = ttk.Label(self.tooltip_window, text=full_text,
                                  background="#ffffe0", relief="solid", borderwidth=1,
                                  wraplength=600, padding=5)
                label.pack()
        except (IndexError, AttributeError):
            pass

    def _on_tree_leave(self, event):
        """Oculta el tooltip cuando el cursor sale del Treeview."""
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

    def accion_leer_pdf_retefuente(self):
        """
        Lee archivos PDF seleccionados y extrae 'RAZON SOCIAL O APELLIDOS...'
        de CADA PÁGINA para exportarlos a un archivo Excel.
        """
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos PDF de Retefuente",
            filetypes=[("Archivos PDF", "*.pdf")],
            parent=self.root
        )
        if not archivos:
            return

        resultados_datos = []
        progress_win, progress_bar, progress_label = self._create_progress_window("Leyendo PDFs...", len(archivos))

        try:
            for i, ruta_pdf in enumerate(archivos):
                nombre_archivo = os.path.basename(ruta_pdf)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    with fitz.open(ruta_pdf) as doc:
                        for num_pagina, page in enumerate(doc, start=1):
                            # --- LÓGICA MEJORADA: Búsqueda por Bloques (Coordenadas) ---
                            blocks = page.get_text("blocks")
                            blocks.sort(key=lambda b: b[1]) # Ordenar verticalmente
                            
                            label_block = None
                            nit_label_block = None
                            nombre_encontrado = "NO ENCONTRADO"
                            nit_encontrado = "NO ENCONTRADO"
                            
                            # 1. Encontrar etiquetas clave
                            for b in blocks:
                                text_clean = " ".join(b[4].split()).upper()
                                if "PRACTICO LA RETENCION" in text_clean:
                                    label_block = b
                                    break
                            
                            # Buscar etiqueta NIT (preferiblemente alineada con la etiqueta principal)
                            if label_block:
                                lx0, ly0, lx1, ly1 = label_block[:4]
                                for b in blocks:
                                    bx0, by0 = b[:2]
                                    text_clean = " ".join(b[4].split()).upper()
                                    # Debe estar a la derecha y misma altura aprox
                                    if bx0 > lx0 and abs(by0 - ly0) < 30:
                                        if "NIT" in text_clean or "C.C." in text_clean:
                                            nit_label_block = b
                                            break
                            
                            # Si no se encontró NIT alineado, buscar genéricamente
                            if not nit_label_block:
                                for b in blocks:
                                    text_clean = " ".join(b[4].split()).upper()
                                    if "NIT." in text_clean and "C.C." in text_clean:
                                        nit_label_block = b
                                        break

                            # 2. Extraer NOMBRE (Debajo de label_block)
                            if label_block:
                                lx0, ly0, lx1, ly1 = label_block[:4]
                                candidates = []
                                for b in blocks:
                                    if b == label_block: continue
                                    bx0, by0 = b[:2]
                                    if by0 > ly0 and abs(bx0 - lx0) < 100: 
                                        candidates.append(b)
                                candidates.sort(key=lambda b: b[1])
                                
                                for cand in candidates:
                                    text_cand = cand[4].strip()
                                    upper_cand = text_cand.upper()
                                    if not text_cand: continue
                                    if "DIRECCION" in upper_cand: break
                                    if "NIT" in upper_cand or "C.C." in upper_cand: continue
                                    nombre_encontrado = " ".join(text_cand.split())
                                    break
                                    
                            # 3. Extraer NIT (Debajo de nit_label_block)
                            if nit_label_block:
                                nx0, ny0 = nit_label_block[:2]
                                nit_candidates = []
                                for b in blocks:
                                    if b == nit_label_block: continue
                                    bx0, by0 = b[:2]
                                    # Debajo y alineado horizontalmente (margen 80px)
                                    if by0 > ny0 and abs(bx0 - nx0) < 80:
                                        nit_candidates.append(b)
                                nit_candidates.sort(key=lambda b: b[1])
                                
                                for cand in nit_candidates:
                                    text_cand = cand[4].strip()
                                    upper_cand = text_cand.upper()
                                    if not text_cand: continue
                                    if "CIUDAD" in upper_cand: break
                                    nit_encontrado = " ".join(text_cand.split())
                                    break
                            
                            # Fallback básico para Nombre
                            if nombre_encontrado == "NO ENCONTRADO":
                                full_text = page.get_text("text")
                                lines = [l.strip() for l in full_text.split('\n') if l.strip()]
                                for idx, line in enumerate(lines):
                                    if "PRACTICO LA RETENCION" in line.upper():
                                        if idx + 1 < len(lines):
                                            potential = lines[idx+1]
                                            if "NIT" not in potential.upper() and "DIRECCION" not in potential.upper():
                                                 nombre_encontrado = potential
                                        break
                            
                            # --- LIMPIEZA FINAL Y SEPARACIÓN ---
                            # Si el nombre contiene el NIT al final (ej: "JUAN PEREZ 123456"), separarlo
                            if nombre_encontrado != "NO ENCONTRADO":
                                # Regex: Busca texto seguido de un número largo al final (mínimo 6 dígitos)
                                match_mix = re.search(r'^(.*?)(\d{6,}[\d\s]*)$', nombre_encontrado)
                                if match_mix:
                                    nombre_limpio = match_mix.group(1).strip()
                                    nit_extraido = match_mix.group(2).replace(" ", "").strip()
                                    
                                    # Actualizamos nombre
                                    nombre_encontrado = nombre_limpio
                                    
                                    # Si no teníamos NIT o era erróneo, usamos el extraído del nombre
                                    if nit_encontrado == "NO ENCONTRADO" or not any(c.isdigit() for c in nit_encontrado):
                                        nit_encontrado = nit_extraido
                                    # Si el NIT actual parece contener texto (es una copia del nombre), lo sobrescribimos
                                    elif any(c.isalpha() for c in nit_encontrado):
                                         nit_encontrado = nit_extraido

                            # Agregar extensión .pdf al nombre para facilitar renombramiento futuro
                            if nombre_encontrado != "NO ENCONTRADO" and not nombre_encontrado.lower().endswith('.pdf'):
                                nombre_encontrado += ".pdf"

                            resultados_datos.append({
                                "Archivo": nombre_archivo,
                                "Página": num_pagina,
                                "RAZON SOCIAL / NOMBRE": nombre_encontrado,
                                "NIT / C.C.": nit_encontrado
                            })
                            
                    self.log(f"✅ Leído: {nombre_archivo} ({doc.page_count} páginas)")
                    
                except Exception as e:
                    self.log(f"❌ Error leyendo {nombre_archivo}: {e}")
                    resultados_datos.append({
                        "Archivo": nombre_archivo,
                        "Página": "Error",
                        "RAZON SOCIAL / NOMBRE": f"ERROR: {str(e)}"
                    })

        finally:
            if progress_win:
                progress_win.destroy()

        if resultados_datos:
            save_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                title="Guardar reporte Retefuente",
                initialfile="Reporte_Retefuente_Paginas.xlsx",
                parent=self.root
            )
            if save_path:
                try:
                    df = pd.DataFrame(resultados_datos)
                    df.to_excel(save_path, index=False)
                    messagebox.showinfo("Éxito", f"Reporte guardado en:\n{save_path}", parent=self.root)
                    self.log(f"Reporte Retefuente guardado en: {save_path}")
                    os.startfile(save_path)
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo guardar el Excel: {e}", parent=self.root)

if __name__ == "__main__":
    root = tk.Tk()
    app = OrganizadorArchivosApp(root)
    root.mainloop()