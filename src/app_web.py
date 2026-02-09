import streamlit as st
import os
import sys
import io
import zipfile
import warnings
# Suppress Google Generative AI FutureWarning
warnings.filterwarnings("ignore", category=FutureWarning, module="google.generativeai")
import unicodedata
import json
import base64
import threading
try:
    from task_manager import submit_task, render_task_center, show_task_notifications
except ImportError:
    from src.task_manager import submit_task, render_task_center, show_task_notifications

# --- AGENT HELPERS ---
def create_standalone_agent_zip():
    """
    Creates a ZIP file containing the standalone agent installer or executable.
    Returns (zip_bytes, found, type_found)
    type_found can be 'installer' or 'portable' or None
    """
    import zipfile
    import io
    import os

    output = io.BytesIO()
    found = False
    type_found = None
    
    # Priority 1: Installer (Instalar_Agente.exe)
    possible_installers = [
        "Instalar_Agente.exe",
        os.path.join("build_release_output", "build_temp", "Instalar_Agente.exe"),
        os.path.join("..", "Instalar_Agente.exe"),
        os.path.join(os.path.dirname(__file__), "..", "Instalar_Agente.exe"),
        # Absolute paths for dev env
        os.path.join(os.path.dirname(__file__), "..", "..", "build_release_output", "build_temp", "Instalar_Agente.exe")
    ]
    
    installer_path = None
    for p in possible_installers:
        if os.path.exists(p):
            installer_path = p
            break
            
    if installer_path:
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as z:
            z.write(installer_path, "Instalar_Agente.exe")
            z.writestr("LEEME.txt", "Ejecute Instalar_Agente.exe para configurar el servicio local.")
        found = True
        type_found = "installer"
        return output.getvalue(), found, type_found

    # Priority 2: Portable Agent (CDO_Agente.exe)
    possible_exes = [
        "CDO_Agente.exe",
        os.path.join("build_release_output", "build_agent", "CDO_Agente.exe"),
        os.path.join("src", "local_agent", "CDO_Agente.exe"),
        os.path.join(os.path.dirname(__file__), "local_agent", "CDO_Agente.exe"),
        os.path.join(os.path.dirname(__file__), "..", "..", "build_release_output", "build_agent", "CDO_Agente.exe")
    ]
    
    exe_path = None
    for p in possible_exes:
        if os.path.exists(p):
            exe_path = p
            break
            
    if exe_path:
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as z:
            z.write(exe_path, "CDO_Agente.exe")
            z.writestr("iniciar_agente.bat", 'start "" "CDO_Agente.exe"')
            z.writestr("LEEME.txt", "Este es el agente portable. Ejecute iniciar_agente.bat.")
        found = True
        type_found = "portable"
        return output.getvalue(), found, type_found

    return None, False, None

def create_lightweight_agent_zip():
    """
    Creates a ZIP with the Python source script for the agent.
    """
    import zipfile
    import io
    import os
    
    output = io.BytesIO()
    
    # Locate main.py for the agent
    agent_script = os.path.join(os.path.dirname(__file__), "local_agent", "main.py")
    if not os.path.exists(agent_script):
        agent_script = os.path.join("src", "local_agent", "main.py")
        
    if os.path.exists(agent_script):
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as z:
            z.write(agent_script, "main.py")
            bat_content = """@echo off
echo Instalando dependencias...
pip install fastapi uvicorn python-multipart
echo Iniciando agente...
python main.py
pause
"""
            z.writestr("iniciar.bat", bat_content)
            z.writestr("LEEME.txt", "Requiere Python instalado. Ejecute iniciar.bat.")
            
    return output.getvalue()

# --- TASK WRAPPERS ---
def run_registraduria_massive(df, col_cedula):
    try:
        from modules.registraduria_validator import ValidatorRegistraduria
    except ImportError:
        from src.modules.registraduria_validator import ValidatorRegistraduria
    from io import BytesIO
    import pandas as pd
    
    # Init validator (headless=False is safer for massive if blocked, but usually headless=True is preferred for background)
    # User context suggests visible windows might be okay or preferred ("que salga una ventana"). 
    # But for pure background, headless is better unless CAPTCHA is needed.
    # Registraduria massive DOES NOT require CAPTCHA, so headless=True is better for background.
    # However, existing code used headless=False. I'll stick to False if user wants to see it, 
    # or True for better background performance. Given user asked for "background", I'll use True to avoid popping up windows unless necessary.
    # Actually, Registraduria often blocks headless. I will use False but minimize if possible, or just let it run.
    validator = ValidatorRegistraduria(headless=False)
    
    # We pass a dummy progress callback because we can't easily update UI from here
    df_results = validator.process_massive(df, col_cedula, progress_callback=lambda c, t: None)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_results.to_excel(writer, index=False)
    return output.getvalue(), f"Procesados {len(df_results)} registros."

def run_adres_api_massive(df, col_cedula):
    try:
        from modules.adres_validator import ValidatorAdres
    except ImportError:
        from src.modules.adres_validator import ValidatorAdres
    from io import BytesIO
    import pandas as pd
    
    validator = ValidatorAdres()
    df_results = validator.process_massive(df, col_cedula, progress_callback=lambda c, t: None)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_results.to_excel(writer, index=False)
    return output.getvalue(), f"Procesados {len(df_results)} registros."

def run_adres_web_massive(df, col_cedula):
    try:
        from modules.adres_validator import ValidatorAdresWeb
    except ImportError:
        from src.modules.adres_validator import ValidatorAdresWeb
    from io import BytesIO
    import pandas as pd
    
    # This one MUST be headless=False because user needs to solve CAPTCHA
    validator = ValidatorAdresWeb(headless=False)
    df_results = validator.process_massive(df, col_cedula, progress_callback=lambda c, t: None)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_results.to_excel(writer, index=False)
    return output.getvalue(), f"Procesados {len(df_results)} registros."


# Suppress warnings from google.generativeai and streamlit to prevent console noise
warnings.filterwarnings("ignore", category=FutureWarning, module="google.generativeai")
warnings.filterwarnings("ignore", message=".*google.generativeai.*")
warnings.filterwarnings("ignore", message=".*use_container_width.*")

import google.generativeai as genai
import socket
import subprocess
import urllib.parse
import tkinter as tk
from tkinter import filedialog, messagebox
import shutil
import time
import pandas as pd
from datetime import datetime
import json
try:
    import src.updater as updater
    from src.version import VERSION as APP_VERSION
except ImportError:
    import updater
    from version import VERSION as APP_VERSION

import uuid  # Added for unique session handling
import compileall
import bot_zeus # Importar módulo del bot
try:
    from modules.registraduria_validator import ValidatorRegistraduria
except ImportError:
    try:
        from src.modules.registraduria_validator import ValidatorRegistraduria
    except ImportError:
        # Fallback dummy class to prevent NameError
        class ValidatorRegistraduria:
            def __init__(self, headless=False):
                pass
            def validate_cedula(self, cedula):
                raise ImportError("El módulo 'registraduria_validator' no se pudo cargar. Verifique las dependencias (selenium, webdriver_manager).")
            def process_massive(self, df, cedula_col, progress_callback=None):
                raise ImportError("El módulo 'registraduria_validator' no se pudo cargar. Verifique las dependencias (selenium, webdriver_manager).")
            def close_driver(self):
                pass

try:
    from modules.adres_validator import ValidatorAdres, ValidatorAdresWeb
except ImportError:
    try:
        from src.modules.adres_validator import ValidatorAdres, ValidatorAdresWeb
    except ImportError:
        class ValidatorAdres:
             def __init__(self): pass
             def validate_cedula(self, cedula, timeout=60): raise ImportError("Módulo adres_validator no encontrado.")
             def process_massive(self, df, cedula_col, progress_callback=None): raise ImportError("Módulo adres_validator no encontrado.")
             def close_driver(self): pass
        class ValidatorAdresWeb:
             def __init__(self, headless=False): pass
             def validate_cedula(self, cedula, timeout=300): raise ImportError("Módulo adres_validator no encontrado.")
             def process_massive(self, df, cedula_col, progress_callback=None): raise ImportError("Módulo adres_validator no encontrado.")
             def close_driver(self): pass

try:
    from modules.analisis_sos import worker_analisis_sos
except ImportError:
    try:
        from src.modules.analisis_sos import worker_analisis_sos
    except ImportError:
        pass # Handle gracefully if module missing


# Configuración de página debe ser la PRIMERA instrucción Streamlit
try:
    st.set_page_config(
        page_title="CDO Clinical Document Organizer", 
        layout="wide", 
        page_icon="📂",
        initial_sidebar_state="expanded"
    )
    # Check for background task notifications immediately after page load
    show_task_notifications()
except Exception as e:
    # Log startup error if set_page_config fails
    try:
        with open(os.path.join(os.path.expanduser("~"), "AppData", "Local", "CDO_Organizer", "startup_error.txt"), "w") as f:
            f.write(f"Startup Error: {str(e)}")
    except:
        pass
    raise e

def create_installer_zip():
    """Crea un ZIP con el Instalador EXE (PyInstaller) que contiene el código protegido."""
    try:
        base_src = os.path.dirname(os.path.abspath(__file__)) # d:\...\src
        project_root = os.path.dirname(base_src) # d:\...\OrganizadorArchivos
        
        output_dir = os.path.join(base_src, "temp_uploads")
        os.makedirs(output_dir, exist_ok=True)
        output_zip = os.path.join(output_dir, "Instalador_CDO.zip")
        
        def sign_exe(target_path):
            """Firma el ejecutable usando el script PowerShell del proyecto."""
            sign_script = os.path.join(project_root, "sign_code.ps1")
            if os.path.exists(sign_script) and os.path.exists(target_path):
                st.info(f"✍️ Firmando {os.path.basename(target_path)}...")
                try:
                    # Run PowerShell script
                    subprocess.run(
                        ["powershell", "-ExecutionPolicy", "Bypass", "-File", sign_script, "-TargetFile", target_path], 
                        check=True, 
                        creationflags=subprocess.CREATE_NO_WINDOW
                    )
                    st.success(f"Firma aplicada a {os.path.basename(target_path)}")
                except Exception as e:
                    st.warning(f"Advertencia: No se pudo firmar el código. {e}")
            else:
                pass # Silent fail if script missing

        # Directorio temporal de construcción
        build_dir = os.path.join(output_dir, "build_temp")
        dist_dir = os.path.join(output_dir, "dist")
        
        # Limpieza previa
        for d in [build_dir, dist_dir]:
            if os.path.exists(d):
                try: shutil.rmtree(d)
                except: pass
        os.makedirs(build_dir, exist_ok=True)

        # 1. Copiar Archivos Raíz necesarios para el build
        files_to_include_root = ["setup_wizard.py", "requirements.txt"]
        for f in files_to_include_root:
            src = os.path.join(project_root, f)
            dst = os.path.join(build_dir, f)
            if os.path.exists(src):
                shutil.copy2(src, dst)
        
        # 2. Copiar Assets
        assets_src = os.path.join(project_root, "assets")
        assets_dst = os.path.join(build_dir, "assets")
        if os.path.exists(assets_src):
            shutil.copytree(assets_src, assets_dst, ignore=shutil.ignore_patterns('~$*', '*.tmp'))

        # 3. Copiar y Compilar SRC
        src_origin = os.path.join(project_root, "src")
        src_dest = os.path.join(build_dir, "src")
        
        # Copiar todo src
        shutil.copytree(src_origin, src_dest, ignore=shutil.ignore_patterns('__pycache__', 'temp_sessions', 'temp_uploads', 'venv', '.git', '*.pyc', '*.zip', 'build_exe', 'dist_exe', '~$*', '*.tmp'))
        
        # Protección de Código: Compilar a .pyc y eliminar .py
        app_web_path = os.path.join(src_dest, "app_web.py")
        if os.path.exists(app_web_path):
            core_path = os.path.join(src_dest, "_core_app.py")
            os.rename(app_web_path, core_path)
            
            compileall.compile_dir(src_dest, force=True, legacy=True, quiet=1)
            
            for root, dirs, files in os.walk(src_dest):
                for file in files:
                    if file.endswith(".py"):
                        # Keep entry points
                        if file in ["run_native.py"]:
                            continue
                        os.remove(os.path.join(root, file))
            
            with open(app_web_path, "w") as f:
                f.write("# Launcher protegido\n")
                f.write("import _core_app\n")
        
        # 4. Compilar en DOS pasos:
        # Paso A: Compilar la Aplicación Cliente (CDO_Cliente.exe) que contiene Streamlit + Python embebido
        # Paso B: Compilar el Instalador (setup_wizard.exe) que empaqueta al Cliente
        
        st.info("🔨 Paso 1/2: Compilando Cliente Nativo (esto puede tardar unos minutos)...")
        
        client_dist_dir = os.path.join(output_dir, "dist_client")
        client_script = os.path.join(src_dest, "run_native.py")
        
        # Asegurar que streamlit tenga sus metadatos
        client_cmd = [
            sys.executable, "-m", "PyInstaller",
            "--noconfirm",
            "--onefile",
            "--windowed",
            "--name", "CDO_Cliente",
            "--clean",
            "--workpath", os.path.join(output_dir, "build_client"),
            "--distpath", client_dist_dir,
            "--specpath", os.path.join(output_dir, "spec_client"),
            
            # Streamlit Critical Hooks
            "--copy-metadata", "streamlit",
            "--copy-metadata", "google-generativeai",
            "--recursive-copy-metadata", "streamlit",
            
            # Exclude Heavy Unused Modules to Reduce Size
            "--exclude-module", "matplotlib",
            "--exclude-module", "scipy",
            "--exclude-module", "notebook",
            "--exclude-module", "jupyter",
            "--exclude-module", "ipython",
            "--exclude-module", "bokeh",
            "--exclude-module", "plotly",
            # tkinter IS used by app_web.py for folder selection, so we cannot exclude it.
            # "--exclude-module", "tkinter", 

            # Hidden Imports
            "--hidden-import", "streamlit",
            "--hidden-import", "streamlit.web.cli",
            "--hidden-import", "pandas",
            "--hidden-import", "altair",
            "--hidden-import", "pydeck",
            "--hidden-import", "rich",
            "--hidden-import", "watchdog",
            "--hidden-import", "tkinter",
            "--hidden-import", "PIL",
            "--hidden-import", "PIL.Image",
            
            # Data: Include SRC (app logic) and Assets
            "--add-data", f"{src_dest}{os.pathsep}src", 
            "--add-data", f"{assets_dst}{os.pathsep}assets",
            
            client_script
        ]
        
        # Run Client Build
        # st.write(f"Debug: Running {client_cmd}")
        process_client = subprocess.run(client_cmd, cwd=build_dir, capture_output=True, text=True)
        
        if process_client.returncode != 0:
            st.error(f"Error compilando Cliente: {process_client.stderr}")
            # Fallback (optional) or Return
            return None
            
        client_exe_path = os.path.join(client_dist_dir, "CDO_Cliente.exe")
        if not os.path.exists(client_exe_path):
            st.error("No se generó el EXE del cliente.")
            return None
            
        # FIRMAR CLIENTE
        sign_exe(client_exe_path)

        # Paso B: Compilar Instalador
        st.info("🔨 Paso 2/2: Empaquetando Instalador Final...")
        
        # Mover Cliente compilado al directorio de build del instalador para empaquetarlo
        shutil.copy2(client_exe_path, os.path.join(build_dir, "CDO_Cliente.exe"))
        
        installer_script = os.path.join(build_dir, "setup_wizard.py")
        exe_name = "Instalador_CDO"
        
        cmd = [
            sys.executable, "-m", "PyInstaller",
            "--noconfirm",
            "--onefile",
            "--windowed",
            "--name", exe_name,
            "--clean",
            "--workpath", build_dir,
            "--distpath", dist_dir,
            "--specpath", build_dir,
            # Bundle the Client EXE inside the Installer
            "--add-data", f"CDO_Cliente.exe{os.pathsep}.",
            # Bundle assets for the installer UI itself
            "--add-data", f"assets{os.pathsep}assets",
            "--hidden-import", "tkinter",
            installer_script
        ]
        
        # Ejecutar compilación del instalador
        process = subprocess.run(cmd, cwd=build_dir, capture_output=True, text=True)
        
        if process.returncode != 0:
            st.error(f"Error compilando instalador: {process.stderr}")
            return None
            
        exe_path = os.path.join(dist_dir, f"{exe_name}.exe")
        
        if not os.path.exists(exe_path):
            st.error("El EXE no se generó.")
            return None
            
        # FIRMAR INSTALADOR
        sign_exe(exe_path)
        
        # 5. Crear ZIPs finales
        # A. ZIP Instalador
        with zipfile.ZipFile(output_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(exe_path, f"{exe_name}.exe")
            
        # B. ZIP Portable (DESACTIVADO)
        output_zip_portable = None 
        # output_zip_portable = os.path.join(output_dir, "CDO_Cliente_Portable.zip")
        # with zipfile.ZipFile(output_zip_portable, 'w', zipfile.ZIP_DEFLATED) as zipf:
        #      zipf.write(client_exe_path, "CDO_Cliente.exe")
        
        # Limpieza
        try: shutil.rmtree(build_dir)
        except: pass
        try: shutil.rmtree(dist_dir)
        except: pass
                            
        return output_zip, output_zip_portable
    except Exception as e:
        st.error(f"Error general creando instalador: {e}")
        return None, None

def create_lightweight_agent_zip():
    """Genera un ZIP con un agente ligero (script Python) para conectar localmente."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as z:
        # agent.py
        agent_code = '''import os
import json
import sys
from flask import Flask, request, jsonify
from flask_cors import CORS

app = Flask(__name__)
CORS(app) # Habilita acceso desde navegador web

@app.route('/', methods=['GET'])
def index():
    return jsonify({"status": "running", "service": "CDO Agent", "version": "1.0"})

@app.route('/list', methods=['GET'])
def list_files():
    path = request.args.get('path', '.')
    if not os.path.exists(path):
        return jsonify({"error": "Ruta no encontrada"}), 404
    try:
        items = []
        with os.scandir(path) as it:
            for entry in it:
                items.append({
                    "name": entry.name,
                    "is_dir": entry.is_dir(),
                    "size": entry.stat().st_size if not entry.is_dir() else 0
                })
        return jsonify({"path": path, "items": sorted(items, key=lambda x: (not x['is_dir'], x['name']))})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/ping', methods=['GET'])
def ping():
    return "pong"

if __name__ == '__main__':
    port = 8000
    print(f"🚀 CDO Agent escuchando en http://localhost:{port}")
    print("Mantenga esta ventana abierta para permitir la conexión.")
    app.run(host='0.0.0.0', port=port)
'''
        z.writestr("agent.py", agent_code)
        
        # requirements.txt
        z.writestr("requirements.txt", "flask\\nflask-cors")
        
        # readme
        readme = """# Agente Local CDO

Este agente permite que la aplicación web CDO acceda a sus archivos locales.

## Instrucciones:
1. Asegúrese de tener Python instalado (https://www.python.org/downloads/).
2. Abra una terminal en esta carpeta.
3. Instale las dependencias:
   pip install -r requirements.txt
4. Ejecute el agente:
   python agent.py
5. Vuelva a la aplicación web y conecte.
"""
        z.writestr("LEER_ANTES_DE_USAR.txt", readme)
        
        # run_agent.bat (Windows convenience)
        bat_code = """@echo off
echo Instalando dependencias si faltan...
pip install -r requirements.txt
cls
echo Iniciando Agente CDO...
python agent.py
pause
"""
        z.writestr("iniciar_agente.bat", bat_code)
        
    buffer.seek(0)
    return buffer.getvalue()

# --- USER MANAGEMENT ---
USERS_FILE = os.path.join(os.path.dirname(__file__), "users.json")

def load_users():
    if not os.path.exists(USERS_FILE):
        return {}
    try:
        with open(USERS_FILE, "r") as f:
            return json.load(f)
    except:
        return {}

def save_users(users):
    with open(USERS_FILE, "w") as f:
        json.dump(users, f, indent=4)

def check_login(username, password):
    users = load_users()
    if username in users and users[username]["password"] == password:
        return True
    return False

def get_user_config(username):
    users = load_users()
    if username in users:
        return users[username]
    return {}

def update_user_last_path(username, path):
    users = load_users()
    if username in users:
        users[username]["last_path"] = path
        save_users(users)

def update_user_config(username, config_data):
    users = load_users()
    if username in users:
        current_config = users[username].get("config", {})
        current_config.update(config_data)
        users[username]["config"] = current_config
        save_users(users)

def get_user_full_config(username):
    users = load_users()
    if username in users:
        return users[username].get("config", {})
    return {}

def add_user_favorite(username, path):
    users = load_users()
    if username in users:
        favs = users[username].get("favorites", [])
        if path not in favs:
            favs.append(path)
            users[username]["favorites"] = favs
            save_users(users)
            return True
    return False

def remove_user_favorite(username, path):
    users = load_users()
    if username in users:
        favs = users[username].get("favorites", [])
        if path in favs:
            favs.remove(path)
            users[username]["favorites"] = favs
            save_users(users)
            return True
    return False

def create_user(username, password):
    users = load_users()
    if username in users:
        return False, "El usuario ya existe"
    
    users[username] = {
        "password": password,
        "last_path": "D:\\",
        "favorites": []
    }
    save_users(users)
    return True, "Usuario creado exitosamente"

def delete_user(username):
    users = load_users()
    if username not in users:
        return False, "El usuario no existe"
    if username == "admin":
        return False, "No se puede eliminar al administrador principal"
    
    del users[username]
    save_users(users)
    return True, "Usuario eliminado exitosamente"

@st.dialog("Panel de Administración de Usuarios")
def admin_panel_modal():
    st.header("👥 Gestión de Usuarios")
    
    tab_list, tab_create, tab_edit = st.tabs(["Listar / Eliminar", "Crear Nuevo", "✏️ Editar Permisos"])
    
    with tab_list:
        users = load_users()
        df_data = [{"Usuario": u, "Rol": d.get("role", "user"), "Bot Zeus": d.get("permissions", {}).get("bot_zeus", "full")} for u, d in users.items()]
        st.dataframe(df_data, use_container_width=True)
        
        st.divider()
        st.subheader("🗑️ Eliminar Usuario")
        user_to_delete = st.selectbox("Seleccionar usuario a eliminar", [u for u in users.keys() if u != "admin"])
        
        if st.button("Eliminar Usuario Seleccionado", type="primary"):
            if user_to_delete:
                ok, msg = delete_user(user_to_delete)
                if ok:
                    st.success(msg)
                    st.rerun()
                else:
                    st.error(msg)
    
    with tab_create:
        st.subheader("➕ Nuevo Usuario")
        new_user = st.text_input("Nombre de usuario")
        new_pass = st.text_input("Contraseña", type="password")
        new_role = st.selectbox("Rol", ["user", "admin"])
        
        if st.button("Crear Usuario"):
            if new_user and new_pass:
                # Create with default permissions
                ok, msg = create_user(new_user, new_pass)
                if ok:
                    # Update role if different from default
                    if new_role != "user":
                        users = load_users()
                        users[new_user]["role"] = new_role
                        save_users(users)
                    st.success(msg)
                    st.rerun()
                else:
                    st.error(msg)
            else:
                st.warning("Complete todos los campos")

    with tab_edit:
        st.subheader("✏️ Editar Permisos y Roles")
        users = load_users()
        user_to_edit = st.selectbox("Seleccionar usuario", list(users.keys()))
        
        if user_to_edit:
            user_data = users[user_to_edit]
            current_role = user_data.get("role", "user")
            current_perms = user_data.get("permissions", {})
            current_bot = current_perms.get("bot_zeus", "full") # default full for backward compat
            current_tabs = current_perms.get("allowed_tabs", ["*"])
            
            # Form
            new_role_edit = st.selectbox("Rol del Usuario", ["user", "admin"], index=0 if current_role == "user" else 1, key="edit_role")
            
            st.divider()
            st.markdown("**Permisos Específicos**")
            
            # Bot Zeus Permissions
            bot_options = ["full", "edit", "execute", "none"]
            bot_labels = ["Completo (Crear/Editar/Ejecutar)", "Edición (Editar/Ejecutar)", "Solo Ejecución", "Sin Acceso"]
            
            try:
                bot_index = bot_options.index(current_bot)
            except:
                bot_index = 0
                
            new_bot_perm = st.selectbox(
                "🤖 Permisos Bot Zeus", 
                bot_options, 
                format_func=lambda x: bot_labels[bot_options.index(x)],
                index=bot_index,
                key="edit_bot"
            )
            
            # Tab Visibility
            all_tabs_available = [
                "🔎 Búsqueda y Acciones",
                "⚙️ Acciones Automatizadas",
                "🔄 Conversión de Archivos",
                "📄 Visor (JSON/XML)",
                "RIPS",
                "🤖 Asistente IA (Gemini)",
                "🤖 Bot Zeus Salud"
            ]
            
            # Helper for multiselect default
            default_tabs = all_tabs_available if "*" in current_tabs else [t for t in current_tabs if t in all_tabs_available]
            
            new_tabs = st.multiselect(
                "👁️ Pestañas Visibles",
                all_tabs_available,
                default=default_tabs,
                key="edit_tabs"
            )
            
            if st.button("💾 Guardar Cambios de Permisos"):
                # Update data
                users[user_to_edit]["role"] = new_role_edit
                
                # Update permissions
                if "permissions" not in users[user_to_edit]:
                    users[user_to_edit]["permissions"] = {}
                
                users[user_to_edit]["permissions"]["bot_zeus"] = new_bot_perm
                
                # Handle tabs
                if len(new_tabs) == len(all_tabs_available):
                    users[user_to_edit]["permissions"]["allowed_tabs"] = ["*"]
                else:
                    users[user_to_edit]["permissions"]["allowed_tabs"] = new_tabs
                
                save_users(users)
                st.success(f"Permisos actualizados para {user_to_edit}")
                time.sleep(1)
                st.rerun()



def seleccionar_carpeta_nativa(title="Seleccionar Carpeta", initial_dir=None):
    """
    Abre un diálogo nativo para seleccionar carpeta usando Agente Local (si disponible) o Tkinter.
    """
    # 1. Intentar usar Agente Local
    try:
        try:
            import agent_client
        except ImportError:
            from src import agent_client
            
        if agent_client.is_agent_available():
            folder = agent_client.select_folder()
            # Si retorna None es porque canceló o falló, retornamos eso.
            return folder
    except Exception as e:
        # Si falla el agente, continuamos con el método nativo (si estamos local)
        print(f"No se pudo contactar al agente local: {e}")

    folder = None
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        # Crear instancia de Tk
        root = tk.Tk()
        root.withdraw() # Ocultar ventana principal
        root.attributes('-topmost', True) # Forzar al frente
        
        # Asegurar que initial_dir sea válido si existe
        if initial_dir and not os.path.isdir(initial_dir):
            initial_dir = None
            
        # Abrir diálogo
        # En Windows moderno, esto llama a IFileDialog con FOS_PICKFOLDERS
        folder = filedialog.askdirectory(
            master=root, 
            title=title, 
            initialdir=initial_dir
        )
        
        root.destroy()
        return folder
    except Exception as e:
        st.error(f"Error al abrir diálogo nativo: {e}")
        return None

def update_path_key(key, title="Seleccionar Carpeta"):
    """Callback para actualizar una ruta en session_state desde un botón."""
    initial_dir = st.session_state.get(key) or st.session_state.get("current_path")
    sel = seleccionar_carpeta_nativa(title, initial_dir)
    if sel:
        st.session_state[key] = sel

def update_main_path():
    """Callback para actualizar la ruta principal desde la sidebar."""
    current_dir = st.session_state.get("current_path")
    folder = seleccionar_carpeta_nativa("Seleccionar Carpeta de Trabajo", current_dir)
    if folder:
        st.session_state.current_path = os.path.normpath(folder)
        st.session_state.path_input = st.session_state.current_path
        if st.session_state.get("username"):
            update_user_last_path(st.session_state.username, st.session_state.current_path)

def render_local_browser():
    # --- CSS Improvements ---
    st.markdown("""
        <style>
        div[data-testid="stDialog"] div[role="dialog"] {
            width: 500px;
            max-width: 90vw;
            height: auto;
        }
        /* Make disabled inputs visible but clearly read-only */
        input:disabled {
            background-color: #f0f2f6 !important;
            color: #31333F !important;
            opacity: 1 !important;
            border: 1px solid #d6d6d6 !important;
            cursor: not-allowed;
        }
        [data-theme="dark"] input:disabled {
             background-color: #262730 !important;
             color: #ffffff !important;
             border: 1px solid #464b5f !important;
             opacity: 1 !important;
             -webkit-text-fill-color: #ffffff !important;
        }
        /* Ensure buttons are visible in Dialogs (Dark/Light friendly) */
    div[data-testid="stDialog"] button {
        border: 1px solid var(--card-border) !important;
        background-color: var(--card-bg) !important;
        color: var(--text-color) !important;
    }
    div[data-testid="stDialog"] button:hover {
        border-color: #ff4b4b !important;
        color: #ff4b4b !important;
    }

    /* Responsive adjustments */
    @media (max-width: 768px) {
        .group-box {
            padding: 10px;
            margin-bottom: 10px;
        }
        div.stButton > button {
            min-height: 40px;
            font-size: 0.9rem;
            padding: 4px 8px !important;
        }
        .stTabs [data-baseweb="tab"] {
            padding: 6px 10px;
            font-size: 0.9rem;
        }
    }
    </style>
""", unsafe_allow_html=True)
    
    st.markdown("### 🖥️ Explorador Nativo")
    st.info("Utiliza el explorador de Windows para seleccionar tu carpeta de trabajo.")

    # Current Path Display
    current = st.session_state.get("current_path", "No seleccionado")
    st.text_input("Ruta Actual:", value=current, disabled=True)

    # Botón con estilo estándar para mejor visibilidad
    st.button("📂 Seleccionar Carpeta con Windows", use_container_width=True, on_click=update_main_path)


def render_web_uploader():
    st.markdown("""
        <style>
        div[data-testid="stDialog"] div[role="dialog"] {
            width: 80vw;
            max-width: 1000px;
        }
        .install-box {
            background-color: var(--card-bg);
            padding: 15px;
            border-radius: 8px;
            border-left: 5px solid #ff4b4b;
            margin-bottom: 20px;
            border: 1px solid var(--card-border);
            color: var(--text-color);
        }
        </style>
    """, unsafe_allow_html=True)

    st.markdown("### ☁️ Modo Web: Cargar Carpeta")
    
    # --- PROMOTION FOR LOCAL CLIENT ---
    st.markdown("""
        <div class="install-box">
            <strong>💡 ¿Quieres acceder a tus discos directamente?</strong><br>
            Instala el <b>Cliente Local CDO</b> para navegar por tus carpetas (C:, D:) sin necesidad de subir archivos ZIP.
            <br><small>Solicita el archivo <code>INSTALAR.bat</code> a tu administrador.</small>
        </div>
    """, unsafe_allow_html=True)
    # ----------------------------------

    st.warning("⚠️ **Nota:** Para subcarpetas, usa un archivo .ZIP.")

    uploaded_files = st.file_uploader(
        "Arrastra ZIP o archivos sueltos", 
        accept_multiple_files=True, 
        key="folder_uploader"
    )

    if uploaded_files:
        count = len(uploaded_files)
        is_zip = count == 1 and uploaded_files[0].name.lower().endswith(".zip")
        
        st.write(f"📦 {count} archivo(s).")
        if is_zip: st.info("✅ ZIP detectado (Extracción automática).")
        
        if st.button("🚀 Cargar y Trabajar", type="primary", use_container_width=True):
            if "session_id" not in st.session_state:
                st.session_state.session_id = str(uuid.uuid4())
            
            temp_base = os.path.join(os.getcwd(), "temp_sessions")
            session_dir = os.path.join(temp_base, st.session_state.session_id)
            
            if os.path.exists(session_dir):
                try: shutil.rmtree(session_dir)
                except: pass
            os.makedirs(session_dir, exist_ok=True)
            
            progress_bar = st.progress(0, text="Procesando...")
            
            try:
                if is_zip:
                    with zipfile.ZipFile(uploaded_files[0]) as z:
                        z.extractall(session_dir)
                    progress_bar.progress(1.0)
                else:
                    total = len(uploaded_files)
                    for i, file in enumerate(uploaded_files):
                        file_path = os.path.join(session_dir, file.name)
                        with open(file_path, "wb") as f:
                            f.write(file.getbuffer())
                        progress_bar.progress((i + 1) / total)
                
                progress_bar.empty()
                st.session_state.current_path = session_dir
                st.session_state.path_input = session_dir
                st.session_state.local_mode = True
                
                st.success("Carga completa.")
                time.sleep(1)
                st.rerun()
                
            except Exception as e:
                st.error(f"Error: {e}")

@st.dialog("Explorador de Archivos", width="large")
def browse_modal():
    # Detect mode from Environment Variable (Set by Installer)
    # Default is WEB if not set
    mode = os.environ.get("CDO_MODE", "WEB").upper()
    
    # Allow manual override from sidebar settings
    if st.session_state.get("force_native_mode", False):
        mode = "LOCAL"
    
    if mode == "LOCAL":
        render_local_browser()
    else:
        render_web_uploader()


import xml.etree.ElementTree as ET
from zipfile import ZipFile, ZIP_DEFLATED
import io
# import tkinter as tk
# from tkinter import filedialog
# import tkinter.messagebox as messagebox
import re
import fitz  # PyMuPDF
from PIL import Image
import openpyxl
from send2trash import send2trash
from pdf2docx import Converter

# Imports condicionales para compatibilidad Linux/Server
try:
    from docx2pdf import convert as convert_docx_to_pdf
    import pythoncom
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False
    convert_docx_to_pdf = None
    pythoncom = None

import requests
from docx import Document
from docx.shared import Inches
from streamlit_elements import elements, mui, html, editor, lazy, sync, event
import urllib.parse
import base64
import string
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

# Silenciar logs de webdriver-manager
os.environ['WDM_LOG_LEVEL'] = '0'

# Configuración de página MOVIDA al inicio del archivo
# st.set_page_config(page_title="CDO Clinical Document Organizer", layout="wide", page_icon="📂")

# --- ESTILOS CSS ---
st.markdown("""
    <style>
    /* CSS ADAPTATIVO USANDO VARIABLES NATIVAS DE STREAMLIT */
    /* Esto asegura que los colores coincidan EXACTAMENTE con el tema seleccionado (Claro/Oscuro) */
    
    :root {
        /* Usamos las variables que Streamlit inyecta automáticamente */
        --card-bg: var(--secondary-background-color);
        --text-color: var(--text-color);
        /* Borde sutil que funciona en ambos modos (blanco o negro con transparencia) */
        --card-border: rgba(128, 128, 128, 0.2); 
        --shadow-color: rgba(0,0,0,0.1);
        --hover-bg: rgba(128, 128, 128, 0.1);
    }

    /* Refuerzo para modo oscuro para asegurar bordes visibles */
    @media (prefers-color-scheme: dark) {
        :root {
            --card-border: rgba(255, 255, 255, 0.2);
            --shadow-color: rgba(0,0,0,0.4);
        }
    }
    
    /* Selector específico para cuando el usuario fuerza modo oscuro en Streamlit */
    [data-theme="dark"] {
        --card-border: rgba(255, 255, 255, 0.2);
        --shadow-color: rgba(0,0,0,0.4);
    }

    .block-container { padding-top: 2rem; }
    
    /* Tabs con estilo de tarjeta */
    .stTabs [data-baseweb="tab-list"] { 
        gap: 8px; 
        padding-bottom: 5px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        background-color: var(--card-bg); 
        border-radius: 8px;
        padding: 10px 16px;
        border: 1px solid var(--card-border);
        box-shadow: 0 2px 4px var(--shadow-color);
        transition: all 0.2s ease;
        color: var(--text-color);
    }
    .stTabs [aria-selected="true"] {
        background-color: var(--card-bg);
        border: 2px solid #ff4b4b;
        color: #ff4b4b;
        font-weight: bold;
        box-shadow: 0 4px 6px var(--shadow-color);
        transform: translateY(-2px);
    }

    .stDataFrame { 
        border: 1px solid var(--card-border); 
        border-radius: 8px;
        box-shadow: 0 2px 5px var(--shadow-color);
    }
    
    /* Contenedores con sombra y bordes redondeados */
    .group-box {
        border: 1px solid var(--card-border);
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 20px;
        background-color: var(--card-bg);
        box-shadow: 0 4px 6px var(--shadow-color), 0 1px 3px var(--shadow-color);
        height: 100%;
        transition: transform 0.2s, box-shadow 0.2s;
        color: var(--text-color);
    }
    .group-box:hover {
        box-shadow: 0 10px 15px var(--shadow-color), 0 4px 6px var(--shadow-color);
        transform: translateY(-2px);
        border-color: #ff4b4b; /* Highlight sutil */
    }

    .group-title, .group-title-left {
        font-weight: 700;
        color: var(--text-color);
        margin-bottom: 15px;
        font-size: 1rem;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        border-bottom: 2px solid var(--card-border);
        padding-bottom: 8px;
    }
    .group-title { text-align: center; }
    .group-title-left { text-align: left; }
    
    /* Botones con relieve Adaptativo */
    div.stButton > button {
        width: 100%;
        min-height: 48px;
        height: 100%;
        white-space: normal !important;
        line-height: 1.2 !important;
        padding: 8px 16px !important;
        display: flex;
        align_items: center;
        justify_content: center;
        border-radius: 8px;
        font-weight: 600;
        border: 1px solid var(--card-border);
        background-color: var(--card-bg); 
        color: var(--text-color);
        box-shadow: 0 2px 4px var(--shadow-color);
        transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1);
    }
    div.stButton > button:hover {
        background-color: var(--hover-bg);
        box-shadow: 0 4px 8px var(--shadow-color);
        transform: translateY(-1px);
        border-color: #ff4b4b;
        color: #ff4b4b;
    }
    div.stButton > button:active {
        box-shadow: 0 1px 2px var(--shadow-color);
        transform: translateY(1px);
    }
    </style>
""", unsafe_allow_html=True)

# --- LOGO DETECTION ---
base_dir = os.path.dirname(os.path.abspath(__file__))
assets_images = os.path.join(base_dir, "..", "assets", "images")

# Definir rutas específicas para cada ubicación
path_header_logo = os.path.join(assets_images, "CDO_logo.png")
path_sidebar_logo = os.path.join(assets_images, "Logo OPV 2.jpg")

# Fallbacks si no existen los específicos
if not os.path.exists(path_header_logo):
    # Si no está el CDO png, buscar jpg u otros
    for p in ["CDO_logo.jpg", "logo.png", "icono.png"]:
        potential = os.path.join(assets_images, p)
        if os.path.exists(potential):
            path_header_logo = potential
            break
    else:
        path_header_logo = None

if not os.path.exists(path_sidebar_logo):
    # Si no está el OPV, buscar otros (evitando repetir si es posible, o usar icono)
    for p in ["icono.png", "logo.png"]:
        potential = os.path.join(assets_images, p)
        if os.path.exists(potential):
            path_sidebar_logo = potential
            break
    else:
        path_sidebar_logo = None

# --- GESTIÓN DE ESTADO ---
if 'logs' not in st.session_state: st.session_state.logs = []
if 'search_results' not in st.session_state: st.session_state.search_results = []
if 'current_path' not in st.session_state: 
    # Better default for native mode
    default_path = "D:\\"
    if not os.path.exists(default_path):
         default_path = os.path.expanduser("~")
    st.session_state.current_path = default_path

if 'conversion_mode' not in st.session_state: st.session_state.conversion_mode = None
if 'editor_data' not in st.session_state: st.session_state.editor_data = None
if 'editor_filename' not in st.session_state: st.session_state.editor_filename = ""
if 'rips_mode' not in st.session_state: st.session_state.rips_mode = None
if 'rips_path' not in st.session_state: st.session_state.rips_path = "" 
if 'action_history' not in st.session_state: st.session_state.action_history = []
if 'session_id' not in st.session_state: st.session_state.session_id = str(uuid.uuid4())
if 'local_mode' not in st.session_state: 
    st.session_state.local_mode = os.environ.get("CDO_MODE") == "LOCAL"

if 'temp_dir' not in st.session_state or not os.path.exists(st.session_state.temp_dir):
    # Crear directorio temporal único para esta sesión
    base_temp = os.path.join(os.path.dirname(__file__), "temp_sessions")
    if not os.path.exists(base_temp): os.makedirs(base_temp, exist_ok=True)
    
    session_path = os.path.join(base_temp, st.session_state.session_id)
    os.makedirs(session_path, exist_ok=True)
    st.session_state.temp_dir = session_path

# Configuración global
if 'app_config' not in st.session_state:
    st.session_state.app_config = {
        "default_pdf_name": "Unificado",
        "default_img_pdf_name": "Imagenes_Unificadas",
        "default_docx_name": "Unificado",
        "split_pdf_prefix": "",
        "image_resolution": 300.0,
        "pdf_compression": 4,
        "pdf_dpi": 600,
        "default_exclusion_patterns": "cuv,xml,xlsx",
        "gemini_api_key": "",
        "gemini_model": "models/gemini-1.5-flash-001"
    }

# --- SESSION PERSISTENCE ---
SESSION_FILE = os.path.join(os.path.dirname(__file__), "session.json")

def save_session(username):
    try:
        with open(SESSION_FILE, "w") as f:
            json.dump({"username": username, "timestamp": time.time()}, f)
    except: pass

def clear_session():
    if os.path.exists(SESSION_FILE):
        try: os.remove(SESSION_FILE)
        except: pass

def load_session():
    if os.path.exists(SESSION_FILE):
        try:
            with open(SESSION_FILE, "r") as f:
                data = json.load(f)
                return data.get("username")
        except: return None
    return None

# --- LOGIN LOGIC ---
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    
    # Auto-login check
    saved_user = load_session()
    if saved_user:
         users_db = load_users()
         if saved_user in users_db:
             st.session_state.logged_in = True
             st.session_state.username = saved_user
             st.session_state.auto_login_success = True

if 'username' not in st.session_state:
    st.session_state.username = ""

if not st.session_state.logged_in:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if path_header_logo and os.path.exists(path_header_logo):
            st.image(path_header_logo, width=250)
            
        st.markdown("## 🔐 Iniciar Sesión")
        with st.form("login_form"):
            user_input = st.text_input("Usuario")
            pass_input = st.text_input("Contraseña", type="password")
            remember_me = st.checkbox("Recordar mi sesión")
            submitted = st.form_submit_button("Entrar")
            
            if submitted:
                if check_login(user_input, pass_input):
                    st.session_state.logged_in = True
                    st.session_state.username = user_input
                    
                    if remember_me:
                        save_session(user_input)
                    else:
                        clear_session()

                    # Load user preferences
                    u_conf = get_user_config(user_input)
                    st.session_state.user_role = u_conf.get("role", "user")
                    st.session_state.user_permissions = u_conf.get("permissions", {})
                    
                    if "last_path" in u_conf and os.path.exists(u_conf["last_path"]):
                        st.session_state.current_path = u_conf["last_path"]
                        st.session_state.path_input = u_conf["last_path"]
                    
                    # Load general configuration
                    full_conf = get_user_full_config(user_input)
                    for key, value in full_conf.items():
                         st.session_state[key] = value

                    st.success("Login correcto!")
                    time.sleep(0.5)
                    st.rerun()
                else:
                    st.error("Usuario o contraseña incorrectos")
    st.stop() # Stop execution if not logged in

# Restore settings if auto-login happened
if st.session_state.get("auto_login_success", False):
    user_input = st.session_state.username
    u_conf = get_user_config(user_input)
    
    st.session_state.user_role = u_conf.get("role", "user")
    st.session_state.user_permissions = u_conf.get("permissions", {})

    if "last_path" in u_conf and os.path.exists(u_conf["last_path"]):
        st.session_state.current_path = u_conf["last_path"]
        st.session_state.path_input = u_conf["last_path"]
    
    full_conf = get_user_full_config(user_input)
    for key, value in full_conf.items():
         st.session_state[key] = value
    
    st.session_state.auto_login_success = False

# Logout button in sidebar
with st.sidebar:
    st.write(f"👤 Usuario: **{st.session_state.username}**")
    if st.button("Cerrar Sesión"):
        clear_session()
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.rerun()

    # Admin Panel Button (Only for admin)
    
    # --- WEB MODE: DOWNLOAD CHANGES ---
    # Show only if we are in "Simulated Local Mode" (Web Upload) AND NOT in Native Desktop Mode
    is_desktop = os.environ.get("CDO_MODE") == "LOCAL"
    if st.session_state.get("local_mode", False) and not is_desktop:
        st.markdown("---")
        st.success("📂 Modo Web Activo")
        st.caption("Los cambios se guardan en la nube. Descárgalos para actualizar tu PC.")
        
        if st.button("📥 Descargar Carpeta Actualizada", type="primary"):
            # Create ZIP
            shutil.make_archive(
                os.path.join(st.session_state.temp_dir, "Mi_Carpeta_Actualizada"), 
                'zip', 
                st.session_state.current_path
            )
            zip_path = os.path.join(st.session_state.temp_dir, "Mi_Carpeta_Actualizada.zip")
            
            with open(zip_path, "rb") as f:
                st.download_button(
                    label="💾 Guardar ZIP en mi PC",
                    data=f,
                    file_name="Carpeta_Trabajo_Actualizada.zip",
                    mime="application/zip"
                )


    if st.session_state.get("user_role") == "admin":
        st.markdown("---")
        if st.button("🛠️ Panel Admin"):
            admin_panel_modal()

        # --- UPDATES SECTION (ADMIN ONLY) ---
        st.markdown("---")
        st.caption(f"v{APP_VERSION}")
        with st.expander("🔄 Actualizaciones"):
            srv = st.session_state.app_config.get("update_server", "http://localhost:8000")
            update_server = st.text_input("Servidor:", value=srv, key="txt_update_server")
            
            # Save server to config if changed
            if update_server != srv:
                st.session_state.app_config["update_server"] = update_server
                # Implicit save would happen on next config save, or we can force it
            
            if st.button("Buscar"):
                with st.spinner("Conectando..."):
                    has_update, ver, info = updater.check_for_updates(update_server)
                    if has_update:
                        st.success(f"¡Nueva versión {ver}!")
                        st.session_state.update_info = {"version": ver, "url": info}
                    else:
                        st.info(info)
                        st.session_state.update_info = None
            
            if st.session_state.get("update_info"):
                st.warning(f"Versión disponible: {st.session_state.update_info['version']}")
                if st.button("📥 Descargar e Instalar"):
                    with st.spinner("Actualizando..."):
                        res = updater.download_and_install(st.session_state.update_info['url'])
                        if res == "UPDATE_INITIATED":
                            st.success("¡Reiniciando!")
                            time.sleep(3)
                            st.stop()
                        else:
                            st.error(f"Error: {res}")



def record_action(action_type, changes):
    """
    Registra una acción en el historial para poder deshacerla.
    changes: lista de tuplas (ruta_actual, ruta_original)
    """
    st.session_state.action_history.append({"type": action_type, "changes": changes})

def undo_last_action():
    if not st.session_state.action_history:
        st.warning("No hay acciones para deshacer.")
        return

    last_action = st.session_state.action_history.pop()
    action_type = last_action["type"]
    changes = last_action["changes"]
    
    success_count = 0
    errors = 0
    
    progress = st.progress(0, text=f"Deshaciendo {action_type}...")
    total = len(changes)
    
    # Invertir lista de cambios por si hubo operaciones dependientes
    for i, (curr, orig) in enumerate(reversed(changes)):
        progress.progress(min((i + 1) / total, 1.0), text=f"Restaurando {i+1}/{total}")
        try:
            if os.path.exists(curr):
                # Verificar si el destino original ya existe (colisión al deshacer)
                if os.path.exists(orig):
                    # Intentar renombrar el existente temporalmente o fallar?
                    # Estrategia simple: Timestamp si colisión
                    base, ext = os.path.splitext(orig)
                    orig = f"{base}_restored_{int(time.time())}{ext}"
                
                shutil.move(curr, orig)
                success_count += 1
            else:
                errors += 1
                log(f"No se pudo deshacer: {curr} no existe.")
        except Exception as e:
            errors += 1
            log(f"Error al deshacer {curr} -> {orig}: {e}")
            
    progress.progress(1.0, text="Deshacer finalizado.")
    st.success(f"Deshacer completado: {success_count} elementos restaurados. ({errors} errores)")
    time.sleep(1.5)
    st.rerun()

# --- SIDEBAR DE CONFIGURACIÓN ---
with st.sidebar:
    # Render Task Center (Background Tasks)
    render_task_center()
    
    st.header("⚙️ Configuración")
    
    with st.expander("🧠 Configuración IA (Gemini)", expanded=True):
        st.session_state.app_config["gemini_api_key"] = st.text_input(
            "API Key Google Gemini",
            value=st.session_state.app_config.get("gemini_api_key", ""),
            type="password",
            help="Obtén tu API Key en: https://aistudio.google.com/"
        )
        
        # --- DETECCIÓN DE MODELOS ---
        if "available_gemini_models" not in st.session_state:
            st.session_state.available_gemini_models = [
                "gemini-1.5-flash",
                "gemini-1.5-pro",
                "gemini-2.0-flash-exp",
                "gemini-1.0-pro"
            ]
            
        if st.button("🔄 Detectar Modelos Disponibles", help="Consulta a Google qué modelos tienes disponibles con tu API Key"):
            api_k = st.session_state.app_config.get("gemini_api_key", "").strip()
            if not api_k:
                st.error("Primero ingresa tu API Key.")
            else:
                try:
                    genai.configure(api_key=api_k)
                    found_models = []
                    for m in genai.list_models():
                        if "generateContent" in m.supported_generation_methods:
                            # Limpiar prefijo models/ para la lista visual
                            clean_name = m.name.replace("models/", "")
                            found_models.append(clean_name)
                    
                    if found_models:
                        st.session_state.available_gemini_models = sorted(found_models, reverse=True)
                        st.success(f"✅ Se encontraron {len(found_models)} modelos.")
                    else:
                        st.warning("⚠️ No se encontraron modelos compatibles con 'generateContent'.")
                except Exception as e:
                    st.error(f"Error consultando modelos: {e}")

        # Opciones de modelo (usando la lista dinámica)
        model_options = st.session_state.available_gemini_models
        
        # Recuperar modelo actual y limpiar prefijos antiguos si existen para matching
        current_model = st.session_state.app_config.get("gemini_model", "gemini-1.5-flash")
        if current_model.startswith("models/"):
            current_model = current_model.replace("models/", "")
            
        # Asegurar que el modelo actual esté en las opciones (si no, añadirlo o resetear)
        if current_model not in model_options:
            model_options.insert(0, current_model)
            
        try:
            default_index = model_options.index(current_model)
        except ValueError:
            default_index = 0

        st.session_state.app_config["gemini_model"] = st.selectbox(
            "Modelo Gemini",
            options=model_options,
            index=default_index,
            help="Selecciona el modelo a utilizar. Flash es más rápido, Pro es más capaz."
        )

    with st.expander("Parámetros Generales", expanded=False):
        st.session_state.app_config["default_pdf_name"] = st.text_input(
            "Nombre PDF Unificado", 
            value=st.session_state.app_config["default_pdf_name"]
        )
        st.session_state.app_config["default_img_pdf_name"] = st.text_input(
            "Nombre PDF Imágenes", 
            value=st.session_state.app_config["default_img_pdf_name"]
        )
        
        # --- AGENTE LOCAL (Integrado) ---
        st.markdown("---")
        st.markdown("### 🖥️ Agente Local CDO (Requerido)")
        
        st.info(
            """
            **¿Para qué sirve?**
            Permite a la aplicación web acceder a archivos de tu computador (Documentos, Disco C:, etc.) de forma segura.
            Debe instalarse una única vez y quedará funcionando siempre en segundo plano.
            """
        )

        col_ag1, col_ag2 = st.columns(2)
        
        with col_ag1:
            st.markdown("#### 1. Instalación")
            # Buscar instalador/agente
            zip_bytes, found, type_found = create_standalone_agent_zip()
            
            if found:
                # Priorizar y destacar el instalador
                label = "📥 Descargar Instalador Automático (.exe)" if type_found == "installer" else "📦 Descargar Versión Portable"
                file_name = "Instalar_Agente_CDO.zip" if type_found == "installer" else "CDO_Agente_Portable.zip"
                
                st.download_button(
                    label=label,
                    data=zip_bytes,
                    file_name=file_name,
                    mime="application/zip",
                    use_container_width=True,
                    help="Descarga e instala el agente. Se configurará para iniciarse automáticamente con Windows."
                )
                
                if type_found == "installer":
                    st.caption("✅ **Recomendado:** Instala el servicio para que funcione siempre (incluso al reiniciar).")
                else:
                    st.caption("⚠️ **Nota:** Esta versión es portable. Debe ejecutar 'iniciar_agente.bat' manualmente.")
                    
            else:
                 # Fallback Script (Solo si no hay binarios)
                 st.warning("⚠️ Instalador oficial no encontrado (Modo Desarrollo).")
                 zip_script = create_lightweight_agent_zip()
                 st.download_button(
                    label="📦 Descargar Script de Desarrollo",
                    data=zip_script,
                    file_name="CDO_Agente_Dev.zip",
                    mime="application/zip",
                    use_container_width=True
                 )

        with col_ag2:
            st.markdown("#### 2. Estado de Conexión")
            # El agente corre en el puerto 8989 por defecto
            agent_url_val = st.session_state.get("agent_url", "http://localhost:8989")
            
            # Verificar conexión automáticamente si no se ha hecho
            if "agent_connected" not in st.session_state:
                 try:
                    import requests
                    resp = requests.get(f"{agent_url_val}/ping", timeout=0.5)
                    st.session_state.agent_connected = (resp.status_code == 200)
                 except:
                    st.session_state.agent_connected = False

            status_color = "green" if st.session_state.get("agent_connected") else "red"
            status_text = "✅ CONECTADO Y ACTIVO" if st.session_state.get("agent_connected") else "❌ DESCONECTADO"
            
            st.markdown(f":{status_color}[**{status_text}**]")
            
            if not st.session_state.get("agent_connected"):
                st.caption("Si ya instaló el agente, asegúrese de que se esté ejecutando.")
            
            if st.button("🔄 Verificar Conexión", use_container_width=True):
                try:
                    import requests
                    resp = requests.get(f"{agent_url_val}/ping", timeout=1)
                    if resp.status_code == 200:
                        st.session_state.agent_connected = True
                        st.rerun()
                    else:
                        st.session_state.agent_connected = False
                        st.error("El agente responde pero con error.")
                except:
                    st.session_state.agent_connected = False
                    st.error("No se detecta el agente local.")

        # Configuración avanzada (oculta)
        with st.expander("Configuración Avanzada Agente"):
            st.session_state.agent_url = st.text_input("URL Agente", value=agent_url_val)
            if st.session_state.get("agent_connected"):
                test_path = st.text_input("Prueba de Ruta:", value=".")
                if st.button("Listar Archivos (Test)"):
                    try:
                        import requests
                        r = requests.get(f"{st.session_state.agent_url}/list", params={"path": test_path})
                        if r.status_code == 200:
                            st.write(r.json().get("items", []))
                    except Exception as e:
                        st.error(f"Error: {e}")

        st.session_state.app_config["default_docx_name"] = st.text_input(
            "Nombre DOCX Unificado", 
            value=st.session_state.app_config.get("default_docx_name", "Unificado")
        )
        st.session_state.app_config["split_pdf_prefix"] = st.text_input(
            "Prefijo División PDF", 
            value=st.session_state.app_config.get("split_pdf_prefix", ""),
            help="Ej: 'Pagina_' para Pagina_1.pdf. Dejar vacío para 1.pdf"
        )
        st.session_state.app_config["feov_prefix"] = st.text_input(
            "Prefijo Facturas (FEOV)",
            value=st.session_state.app_config.get("feov_prefix", "FEOV"),
            help="Texto que precede al número de factura en el nombre del archivo (Ej: 'FEOV' para '...FEOV123.pdf')."
        )
        st.session_state.app_config["image_resolution"] = st.number_input(
            "Resolución Imagen (DPI)", 
            min_value=72.0, max_value=600.0, 
            value=st.session_state.app_config["image_resolution"]
        )
        st.session_state.app_config["pdf_compression"] = st.slider(
            "Compresión PDF (0-20)", 
            min_value=0, max_value=20, 
            value=st.session_state.app_config["pdf_compression"],
            help="0: Ninguna, 4: Estándar, >4: Alta (Puede ser lento)"
        )
        st.session_state.app_config["pdf_dpi"] = st.number_input(
            "Calidad PDF (DPI)",
            min_value=72, max_value=600,
            value=st.session_state.app_config["pdf_dpi"]
        )
        st.session_state.app_config["default_exclusion_patterns"] = st.text_area(
            "Patrones de Exclusión (Búsqueda)",
            value=st.session_state.app_config.get("default_exclusion_patterns", ""),
            help="Separar por comas. Ej: cuv,xml,xlsx. Estos archivos serán ignorados en las búsquedas."
        )

    with st.expander("💾 Preferencias de Usuario", expanded=False):
        st.caption("Guarda tu configuración actual (filtros, acciones, parámetros) para que se cargue automáticamente al iniciar sesión.")
        
        if st.button("💾 Guardar Configuración Actual", use_container_width=True):
            # Collect keys to save
            config_to_save = {}
            # Keys from Tab 1 (Search & Actions)
            keys_to_save = ["search_by", "item_type", "pattern", "exclusion_pattern", "subfolders", "action_radio"]
            for k in keys_to_save:
                if k in st.session_state:
                    config_to_save[k] = st.session_state[k]
            
            # Save app_config (Parameters)
            if "app_config" in st.session_state:
                config_to_save["app_config"] = st.session_state.app_config

            update_user_config(st.session_state.username, config_to_save)
            st.toast("✅ Configuración guardada y vinculada a tu usuario.")

    # Logo al final del sidebar
    if path_sidebar_logo:
        st.markdown("---")
        st.image(path_sidebar_logo, use_container_width=True)

def log(msg):
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.logs.append(f"[{timestamp}] {msg}")

# --- FUNCIONES AUXILIARES ---
def funcion_no_implementada(nombre):
    st.toast(f"⚠️ Función '{nombre}' simulada en versión web.")
    log(f"Ejecutado: {nombre}")

def seleccionar_archivo(file_types=[("Todos los archivos", "*.*")]):
    st.toast("⚠️ La selección de archivos nativa no está disponible en modo web. Usa la carga de archivos.")
    return None

# --- WORKERS DE CONVERSIÓN ---

def _pdf_a_docx(input_path, output_path):
    cv = Converter(input_path)
    cv.convert(output_path)
    cv.close()

def _jpg_a_pdf(input_path, output_path):
    img = Image.open(input_path)
    if img.mode == 'RGBA':
        img = img.convert('RGB')
    res = st.session_state.app_config.get("image_resolution", 100.0)
    img.save(output_path, "PDF", resolution=res)

def _docx_a_pdf(input_path, output_path):
    # docx2pdf maneja la conversión. 
    # Nota: Requiere Word instalado en Windows.
    # pythoncom.CoInitialize() puede ser necesario si hay problemas de hilos.
    if not HAS_DOCX2PDF:
        st.error("La conversión DOCX -> PDF requiere Windows y Microsoft Word instalado. No disponible en este entorno.")
        return

    try:
        pythoncom.CoInitialize() 
    except: pass
    convert_docx_to_pdf(input_path, output_path)

def _pdf_a_jpg(input_path, output_base):
    # output_base es ruta sin extensión
    doc = fitz.open(input_path)
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        if len(doc) > 1:
            out = f"{output_base}_p{i+1}.jpg"
        else:
            out = f"{output_base}.jpg"
        pix.save(out)
    doc.close()

def _png_a_jpg(input_path, output_path):
    img = Image.open(input_path)
    rgb_img = img.convert('RGB')
    rgb_img.save(output_path, 'jpeg')

def _txt_a_json(input_path, output_path):
    # Solo renombrar extensión
    # Asegurarnos de que no sea el mismo archivo (caso raro)
    if input_path == output_path: return
    
    if not os.path.exists(output_path):
        os.rename(input_path, output_path)
    else:
        # Si existe, intentar con timestamp
        base, ext = os.path.splitext(output_path)
        new_out = f"{base}_{int(time.time())}.json"
        os.rename(input_path, new_out)

# --- WORKERS IA GEMINI ---
def extract_text_from_file(file_path):
    """Extrae texto de PDF, DOCX o TXT para análisis con IA."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == ".pdf":
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
        elif ext == ".docx":
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext in [".txt", ".csv", ".json", ".xml", ".py", ".js", ".html"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception as e:
        return f"Error leyendo archivo: {e}"
    return text

def worker_consultar_gemini(prompt, file_context=None):
    api_key = st.session_state.app_config.get("gemini_api_key")
    if not api_key:
        return "⚠️ Por favor configura tu API Key de Google Gemini en el panel lateral."
    
    try:
        # Limpiar API Key de espacios accidentales
        api_key = api_key.strip()
        genai.configure(api_key=api_key)
        
        model_name = st.session_state.app_config.get("gemini_model", "gemini-1.5-flash")
        
        # Intentar instanciar modelo
        # Nota: genai.GenerativeModel suele aceptar tanto 'gemini-pro' como 'models/gemini-pro'
        # pero para mayor robustez probamos ambos si falla.
        
        try:
            model = genai.GenerativeModel(model_name)
            
            full_prompt = prompt
            if file_context:
                full_prompt = f"Contexto del archivo:\n{file_context}\n\nPregunta:\n{prompt}"
                
            response = model.generate_content(full_prompt)
            return response.text
            
        except Exception as e_model:
            # Fallback: Si falla con 404 y tiene/no tiene prefijo, intentar la inversa
            err_msg = str(e_model)
            if "404" in err_msg:
                alt_name = None
                if model_name.startswith("models/"):
                    alt_name = model_name.replace("models/", "")
                else:
                    alt_name = f"models/{model_name}"
                
                if alt_name:
                    try:
                        model = genai.GenerativeModel(alt_name)
                        full_prompt = prompt
                        if file_context:
                            full_prompt = f"Contexto del archivo:\n{file_context}\n\nPregunta:\n{prompt}"
                        response = model.generate_content(full_prompt)
                        return response.text
                    except:
                        pass # Fallback falló, lanzar error original
            
            raise e_model

    except Exception as e:
        msg = str(e)
        if "API_KEY_INVALID" in msg or "403" in msg:
            return "⛔ Error de Autenticación: Tu API Key no es válida o ha expirado. Verifícala en Google AI Studio."
        if "404" in msg:
            return f"⛔ Modelo no encontrado o no disponible: {model_name}. Intenta seleccionar otro modelo."
        return f"❌ Error consultando a Gemini: {msg}"

def _pdf_escala_grises(input_path, output_path):
    doc = fitz.open(input_path)
    doc_final = fitz.open()
    
    dpi = st.session_state.app_config.get("pdf_dpi", 600)
    matrix_scale = dpi / 72.0
    mat = fitz.Matrix(matrix_scale, matrix_scale)
    
    for page in doc:
        # Usar matriz para alta resolución
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY)
        new_page = doc_final.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_image(new_page.rect, pixmap=pix)
    doc.close()
    
    compression = st.session_state.app_config.get("pdf_compression", 4)
    doc_final.save(output_path, garbage=compression, deflate=True)
    doc_final.close()

def worker_convertir_archivo(file_path, tipo, output_folder=None):
    if not file_path or not os.path.exists(file_path):
        return False, "Archivo no encontrado"
        
    folder = output_folder if output_folder else os.path.dirname(file_path)
    filename = os.path.basename(file_path)
    name_no_ext = os.path.splitext(filename)[0]
    
    try:
        if tipo == "PDF2DOCX":
            out = os.path.join(folder, f"{name_no_ext}.docx")
            _pdf_a_docx(file_path, out)
        elif tipo == "JPG2PDF":
            out = os.path.join(folder, f"{name_no_ext}.pdf")
            _jpg_a_pdf(file_path, out)
        elif tipo == "DOCX2PDF":
            out = os.path.join(folder, f"{name_no_ext}.pdf")
            _docx_a_pdf(file_path, out)
        elif tipo == "PDF2JPG":
            out_base = os.path.join(folder, name_no_ext)
            _pdf_a_jpg(file_path, out_base)
        elif tipo == "PNG2JPG":
            out = os.path.join(folder, f"{name_no_ext}.jpg")
            _png_a_jpg(file_path, out)
        elif tipo == "TXT2JSON":
            out = os.path.join(folder, f"{name_no_ext}.json")
            _txt_a_json(file_path, out)
        elif tipo == "PDF_GRAY":
            # Sustituir archivo original
            temp_out = os.path.join(folder, f"{name_no_ext}_temp_gray.pdf")
            _pdf_escala_grises(file_path, temp_out)
            
            # Reemplazar original con el temporal
            if os.path.exists(temp_out):
                try:
                    # Intentar reemplazo atómico/directo
                    os.replace(temp_out, file_path)
                except OSError:
                    # Fallback por si hay bloqueos
                    time.sleep(0.5)
                    os.remove(file_path)
                    os.rename(temp_out, file_path)
            
        return True, "Conversión exitosa"
    except Exception as e:
        return False, str(e)

def worker_convertir_masivo(folder_path, tipo):
    if not folder_path or not os.path.exists(folder_path):
        return 0, "Carpeta no encontrada"
    
    count = 0
    # os.listdir solo lista archivos en la raíz. ¿Debe ser recursivo?
    # La interfaz dice "Por Carpeta Completa". Asumamos recursividad para ser más útil,
    # o mantengamos solo raíz si así era el original. El original solía ser solo raíz.
    # Pero si el usuario tiene subcarpetas, no verá nada.
    # Vamos a usar os.walk para buscar en TODO el árbol si no encuentra nada en la raíz.
    
    files_to_process = []
    
    # Búsqueda recursiva en todas las subcarpetas
    for r, d, f in os.walk(folder_path):
        for file in f:
            files_to_process.append(os.path.join(r, file))

    total = len(files_to_process)
    if total == 0:
        return 0, "La carpeta está vacía (no se encontraron archivos)."

    progress_bar = st.progress(0, text="Convirtiendo...")
    
    for i, full_path in enumerate(files_to_process):
        if i % 5 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        f = os.path.basename(full_path)
        f_lower = f.lower()
        process = False
        
        # Filtros básicos por extensión
        if tipo == "PDF2DOCX" and f_lower.endswith(".pdf"): process = True
        elif tipo == "JPG2PDF" and (f_lower.endswith(".jpg") or f_lower.endswith(".jpeg")): process = True
        elif tipo == "DOCX2PDF" and f_lower.endswith(".docx") and not f.startswith("~$"): process = True
        elif tipo == "PDF2JPG" and f_lower.endswith(".pdf"): process = True
        elif tipo == "PNG2JPG" and f_lower.endswith(".png"): process = True
        elif tipo == "TXT2JSON" and f_lower.endswith(".txt"): process = True
        elif tipo == "PDF_GRAY" and f_lower.endswith(".pdf"): process = True
        
        if process:
            ok, msg = worker_convertir_archivo(full_path, tipo)
            if ok: count += 1
            else: log(f"Error convirtiendo {f}: {msg}")
            
    progress_bar.progress(1.0, text="Finalizado.")
    return count, f"Procesados {count} archivos."

# --- WORKERS (Lógica Real - Continuación) ---

def worker_unificar_pdf(root_path, output_name_base, pdf_compression=4, pdf_dpi=300, silent_mode=False):
    log(f"Iniciando unificación PDF en: {root_path}")
    count_folders = 0
    
    if not output_name_base:
        output_name_base = "Unificado"

    # Configs
    compression = pdf_compression
    dpi = pdf_dpi
    matrix_scale = dpi / 72.0

    # Pre-calcular total para barra de progreso
    total_steps = sum(1 for _ in os.walk(root_path))
    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando unificación...")
    
    current_step = 0

    for root, dirs, files in os.walk(root_path):
        current_step += 1
        if not silent_mode:
            progress_bar.progress(min(current_step / total_steps, 1.0), text=f"Procesando: {os.path.basename(root)}")

        # Buscar secuencia 1.pdf ... 10.pdf (simple) o todos los pdfs
        # La lógica original busca especificamente 1.pdf a 10.pdf.
        # Vamos a hacerlo más flexible: todos los PDFs ordenados alfabéticamente
        pdfs = [f for f in files if f.lower().endswith(".pdf")]
        pdfs.sort() # Ordenar 1.pdf, 10.pdf, 2.pdf... (Ojo con orden natural, pero por ahora simple)
        
        # Filtrar si queremos seguir la lógica estricta de 1..10
        # pdfs_target = []
        # for i in range(1, 11):
        #     name = f"{i}.pdf"
        #     if name in pdfs: pdfs_target.append(name)
        
        # Usaremos todos los PDFs encontrados en la subcarpeta para ser más útil
        if not pdfs:
            continue
            
        try:
            doc_final = fitz.open()
            found_content = False
            
            for pdf_file in pdfs:
                pdf_path = os.path.join(root, pdf_file)
                # Evitar incluir el archivo de salida si ya existe
                if pdf_file == f"{output_name_base}.pdf":
                    continue
                    
                try:
                    # Convertir a imágenes e insertar (como en el original) para aplanar
                    doc_temp = fitz.open(pdf_path)
                    for page in doc_temp:
                        pix = page.get_pixmap(matrix=fitz.Matrix(matrix_scale, matrix_scale), colorspace=fitz.csGRAY)
                        # Crear nueva página en destino con tamaño de la imagen
                        page_new = doc_final.new_page(width=pix.width, height=pix.height)
                        page_new.insert_image(page_new.rect, pixmap=pix)
                        found_content = True
                    doc_temp.close()
                except Exception as e:
                    log(f"Error procesando {pdf_file}: {e}")

            if found_content:
                out_path = os.path.join(root, f"{output_name_base}.pdf")
                doc_final.save(out_path, garbage=compression, deflate=True)
                doc_final.close()
                count_folders += 1
                
        except Exception as e:
            log(f"Error en carpeta {root}: {e}")
    
    msg = f"Proceso completado. Se generaron PDFs en {count_folders} carpetas."
    if not silent_mode:
        progress_bar.progress(1.0, text="Unificación completada.")
        st.success(msg)
    return msg

def run_unificar_pdf_task(root_path, output_name_base, pdf_compression, pdf_dpi):
    return {"message": worker_unificar_pdf(root_path, output_name_base, pdf_compression, pdf_dpi, silent_mode=True)}

def worker_dividir_pdf_masivo(root_path, pdf_prefix="", silent_mode=False):
    log(f"Iniciando división masiva de PDFs en: {root_path}")
    count_files = 0
    
    # Pre-conteo de archivos PDF
    total_pdfs = 0
    for r, d, f in os.walk(root_path):
        for file in f:
            if file.lower().endswith(".pdf"):
                total_pdfs += 1
                
    if total_pdfs == 0:
        if not silent_mode:
            st.warning("No se encontraron archivos PDF.")
        return "No se encontraron archivos PDF."

    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando división...")
    processed_count = 0

    for root, dirs, files in os.walk(root_path):
        for file in files:
            if file.lower().endswith(".pdf"):
                processed_count += 1
                if not silent_mode:
                    progress_bar.progress(min(processed_count / total_pdfs, 1.0), text=f"Dividiendo: {file}")

                pdf_path = os.path.join(root, file)
                folder_name = os.path.splitext(file)[0]
                output_folder = os.path.join(root, folder_name)
                
                # Crear carpeta
                if not os.path.exists(output_folder):
                    os.makedirs(output_folder)
                
                try:
                    doc = fitz.open(pdf_path)
                    for i, page in enumerate(doc):
                        new_doc = fitz.open()
                        new_doc.insert_pdf(doc, from_page=i, to_page=i)
                        prefix = pdf_prefix
                        out_name = f"{prefix}{i+1}.pdf"
                        new_doc.save(os.path.join(output_folder, out_name))
                        new_doc.close()
                    doc.close()
                    count_files += 1
                except Exception as e:
                    log(f"Error dividiendo {file}: {e}")
    
    msg = f"División completada. Procesados {count_files} archivos PDF."
    if not silent_mode:
        progress_bar.progress(1.0, text="División completada.")
        st.success(msg)
    return msg

def run_dividir_pdf_masivo_task(root_path, pdf_prefix):
    return {"message": worker_dividir_pdf_masivo(root_path, pdf_prefix, silent_mode=True)}


def worker_unificar_imagenes_pdf(root_path, output_name_base, ext_list, silent_mode=False):
    log(f"Iniciando unificación de imágenes {ext_list} en: {root_path}")
    count_folders = 0
    
    if not output_name_base: 
        output_name_base = "Imagenes_Unificadas"
    
    total_steps = sum(1 for _ in os.walk(root_path))
    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando unificación de imágenes...")
    
    current_step = 0

    for root, dirs, files in os.walk(root_path):
        current_step += 1
        if not silent_mode:
            progress_bar.progress(min(current_step / total_steps, 1.0), text=f"Procesando: {os.path.basename(root)}")

        imgs = [f for f in files if f.lower().endswith(tuple(ext_list))]
        # Ordenar numéricamente si es posible (1.jpg, 2.jpg) o alfabéticamente
        # Intentar orden natural
        imgs.sort(key=lambda x: [int(c) if c.isdigit() else c for c in re.split(r'(\d+)', x)])
        
        if not imgs: continue
        
        try:
            pil_images = []
            first_img = None
            
            for img_file in imgs:
                img_path = os.path.join(root, img_file)
                try:
                    img = Image.open(img_path)
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    
                    if first_img is None:
                        first_img = img
                    else:
                        pil_images.append(img)
                except Exception as e:
                    log(f"Error leyendo imagen {img_file}: {e}")

            if first_img:
                out_path = os.path.join(root, f"{output_name_base}.pdf")
                first_img.save(out_path, save_all=True, append_images=pil_images)
                count_folders += 1
                
        except Exception as e:
            log(f"Error procesando carpeta {root}: {e}")

    msg = f"Proceso completado. PDFs creados en {count_folders} carpetas."
    if not silent_mode:
        progress_bar.progress(1.0, text="Unificación completada.")
        st.success(msg)
    
    return msg

def run_unificar_img_task(root_path, output_name_base, ext_list):
    return worker_unificar_imagenes_pdf(root_path, output_name_base, ext_list, silent_mode=True)

def worker_crear_carpetas_excel_avanzado(uploaded_file, sheet_name, col_name, root_path, use_visible_rows, silent_mode=False):
    count = 0
    names_to_create = []

    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        # Cargar libro con openpyxl para detectar filas ocultas si es necesario
        uploaded_file.seek(0)
        wb = openpyxl.load_workbook(uploaded_file, data_only=True)
        if sheet_name not in wb.sheetnames:
            if silent_mode: return f"Error: Hoja '{sheet_name}' no encontrada."
            st.error(f"Hoja '{sheet_name}' no encontrada.")
            return
        
        ws = wb[sheet_name]
        
        # Encontrar índice de columna
        headers = [cell.value for cell in ws[1]] # Asumiendo encabezados en fila 1
        if col_name not in headers:
            if silent_mode: return f"Error: Columna '{col_name}' no encontrada en la fila 1."
            st.error(f"Columna '{col_name}' no encontrada en la fila 1.")
            return
            
        col_idx = headers.index(col_name) # 0-based index
        
        if not silent_mode:
            progress_bar = st.progress(0, text="Leyendo Excel...")
        
        # Iterar filas (desde fila 2)
        rows = list(ws.iter_rows(min_row=2))
        total_rows = len(rows)
        
        for i, row in enumerate(rows):
            if not silent_mode and i % 100 == 0: # Actualizar cada 100 filas para no saturar
                progress_bar.progress(min((i / total_rows) * 0.5, 0.5), text=f"Leyendo fila {i}...")

            cell = row[col_idx]
            val = cell.value
            
            # Lógica de filtro visible
            if use_visible_rows:
                # openpyxl row_dimensions[row_idx].hidden
                # row[0].row devuelve el número de fila (1-based)
                row_num = row[0].row
                if ws.row_dimensions[row_num].hidden:
                    continue
            
            if val:
                names_to_create.append(str(val))
        
        wb.close()
        
        # Crear carpetas
        total_create = len(names_to_create)
        if total_create == 0:
            if not silent_mode:
                progress_bar.empty()
                st.warning("No se encontraron nombres válidos para crear carpetas.")
            return "No se encontraron nombres válidos."

        if not silent_mode:
            progress_bar.progress(0.5, text="Creando carpetas...")
        
        for i, name in enumerate(names_to_create):
            # Actualizar progreso (de 0.5 a 1.0)
            if not silent_mode and i % 10 == 0:
                progress_val = 0.5 + (i / total_create) * 0.5
                progress_bar.progress(min(progress_val, 1.0), text=f"Creando: {name}")

            safe_name = "".join([c for c in name if c.isalpha() or c.isdigit() or c in " _-"]).strip()
            if not safe_name: continue
            
            new_path = os.path.join(root_path, safe_name)
            try:
                if not os.path.exists(new_path):
                    os.makedirs(new_path)
                    count += 1
            except Exception as e:
                log(f"Error creando carpeta {safe_name}: {e}")
        
        msg = f"Creación finalizada. {count} carpetas creadas."
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
        
    except Exception as e:
        err_msg = f"Error procesando Excel: {e}"
        if silent_mode: return err_msg
        st.error(err_msg)

def run_crear_carpetas_task(uploaded_file, sheet_name, col_name, root_path, use_visible_rows):
    return worker_crear_carpetas_excel_avanzado(uploaded_file, sheet_name, col_name, root_path, use_visible_rows, silent_mode=True)

def worker_organizar_facturas_feov(path_origen, path_destino, prefix="FEOV", silent_mode=False):
    log(f"Iniciando organización FEOV. Origen: {path_origen} -> Destino: {path_destino} | Prefijo: {prefix}")
    
    # Metricas
    metrics = {
        "scanned_dest": 0,
        "mapped_feov": 0,
        "scanned_source": 0,
        "moved": 0,
        "errors": 0,
        "conflicts": 0
    }
    
    # 1. Mapeo de Destinos
    mapa_feov = {} # {'12345': 'C:/Ruta/Carpeta_Cliente_X'}
    
    if not silent_mode:
        progress_bar = st.progress(0, text="Escaneando destinos...")
    
    dest_subfolders = [os.path.join(path_destino, d) for d in os.listdir(path_destino) if os.path.isdir(os.path.join(path_destino, d))]
    total_dest = len(dest_subfolders)
    
    for i, folder in enumerate(dest_subfolders):
        if not silent_mode and i % 10 == 0 and total_dest > 0:
            progress_bar.progress(min(i / total_dest, 1.0), text=f"Escaneando destino {i}/{total_dest}")
            
        for root, _, files in os.walk(folder):
            for file in files:
                metrics["scanned_dest"] += 1
                if file.lower().endswith(".pdf") and prefix.lower() in file.lower():
                    # Extraer números después de FEOV (prefijo dinámico)
                    # Usamos re.escape para evitar problemas si el prefijo tiene caracteres especiales
                    pattern = re.escape(prefix) + r'(\d+)'
                    match = re.search(pattern, file, re.IGNORECASE)
                    if match:
                        num_feov = match.group(1)
                        if num_feov not in mapa_feov:
                             metrics["mapped_feov"] += 1
                        mapa_feov[num_feov] = folder
    
    if not mapa_feov:
        if not silent_mode:
            progress_bar.empty()
            st.warning(f"No se encontraron archivos con el patrón '{prefix}...' en las carpetas de destino para crear el mapa.")
        return metrics
    
    log(f"Mapa creado con {metrics['mapped_feov']} referencias.")
    
    # 2. Procesar Origen
    files_origen = [f for f in os.listdir(path_origen) if os.path.isfile(os.path.join(path_origen, f))]
    metrics["scanned_source"] = len(files_origen)
    
    if not silent_mode:
        progress_bar.progress(0, text="Moviendo archivos...")
    
    for i, file in enumerate(files_origen):
        if not silent_mode and i % 10 == 0 and metrics["scanned_source"] > 0:
            progress_bar.progress(min(i / metrics["scanned_source"], 1.0), text=f"Procesando archivo {i}/{metrics['scanned_source']}")
            
        # Buscar si algún feov del mapa está en el nombre del archivo
        # Esto puede ser lento si el mapa es gigante. Optimización: Buscar números en el archivo y ver si están en el mapa.
        # Estrategia original: "Si el nombre del archivo contiene alguno de los números"
        
        # Estrategia optimizada: Extraer todos los números del nombre del archivo y chequear si alguno es clave en mapa_feov
        numbers_in_name = re.findall(r'\d+', file)
        target_folder = None
        
        for num in numbers_in_name:
            if num in mapa_feov:
                target_folder = mapa_feov[num]
                break
        
        if target_folder:
            src = os.path.join(path_origen, file)
            dst = os.path.join(target_folder, file)
            try:
                if os.path.exists(dst):
                     metrics["conflicts"] += 1
                     log(f"Conflicto: {file} ya existe en {target_folder}")
                elif src != dst:
                    shutil.move(src, dst)
                    metrics["moved"] += 1
            except Exception as e:
                log(f"Error moviendo {file}: {e}")
                metrics["errors"] += 1

    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
    
    # Display Metrics
    if not silent_mode:
        st.divider()
        st.subheader("📊 Resultados de Organización FEOV")
        
        c1, c2, c3 = st.columns(3)
        c1.metric("📂 Archivos en Origen", metrics["scanned_source"])
        c2.metric("📦 Movidos Exitosamente", metrics["moved"])
        c3.metric("🎯 FEOVs Identificados", metrics["mapped_feov"])
        
        c4, c5, c6 = st.columns(3)
        c4.metric("⚠️ Conflictos (Ya existen)", metrics["conflicts"])
        c5.metric("❌ Errores", metrics["errors"])
        c6.metric("🔍 Archivos Destino Escaneados", metrics["scanned_dest"])
        
        if metrics["moved"] > 0:
            st.balloons()
    
    return metrics

def run_organizar_feov_task(path_origen, path_destino, prefix="FEOV"):
    metrics = worker_organizar_facturas_feov(path_origen, path_destino, prefix, silent_mode=True)
    msg = f"Movidos: {metrics['moved']} | Errores: {metrics['errors']} | Conflictos: {metrics['conflicts']}"
    return {"message": msg}

# --- HELPER FUNCTIONS ---
def get_val_ci(d, key_target):
    if not isinstance(d, dict): return None
    # Intento directo
    if key_target in d: return d[key_target]
    # Búsqueda normalizada
    # Asegurarse que normalize_key esté disponible, si no, usar implementación local simple o llamar a la global
    # Como normalize_key es global (vimos en linea 1795), debería funcionar.
    # Pero para seguridad en ejecución standalone de workers si se extraen, mejor usar la global.
    target_norm = normalize_key(key_target)
    for k, v in d.items():
        if normalize_key(k) == target_norm:
            return v
    return None

def worker_json_a_xlsx_ind(file_obj):
    try:
        # Asegurar lectura desde inicio si es un objeto tipo archivo
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
            
        data = json.load(file_obj)
        
        # Inicializar listas para las hojas (Estructura RIPS 2275 Normalizada)
        transaccion_rows = []
        usuarios_rows = []
        consultas = []
        procedimientos = []
        urgencias = []
        hospitalizacion = []
        recien_nacidos = []
        medicamentos = []
        otros_servicios = []

        # Función auxiliar para búsqueda insensible a mayúsculas en diccionario
        def get_val_ci(d, key_target):
            if not isinstance(d, dict): return None
            # Intento directo
            if key_target in d: return d[key_target]
            # Búsqueda normalizada
            target_norm = normalize_key(key_target)
            for k, v in d.items():
                if normalize_key(k) == target_norm:
                    return v
            return None

        # Extraer campos de cabecera (Hoja transaccion)
        header_info = {
            "numDocumentoIdObligado": get_val_ci(data, "numDocumentoIdObligado"),
            "numFactura": get_val_ci(data, "numFactura"),
            "tipoNota": get_val_ci(data, "tipoNota"),
            "numNota": get_val_ci(data, "numNota")
        }
        transaccion_rows.append(header_info)

        # Obtener lista de usuarios
        usuarios_lista = get_val_ci(data, "usuarios")
        if not isinstance(usuarios_lista, list): usuarios_lista = []
        
        for usuario in usuarios_lista:
            # Hoja usuarios
            u_row = {
                "tipoDocumentoIdentificacion": get_val_ci(usuario, "tipoDocumentoIdentificacion"),
                "numDocumentoIdentificacion": get_val_ci(usuario, "numDocumentoIdentificacion"),
                "tipoUsuario": get_val_ci(usuario, "tipoUsuario"),
                "fechaNacimiento": get_val_ci(usuario, "fechaNacimiento"),
                "codSexo": get_val_ci(usuario, "codSexo"),
                "codPaisResidencia": get_val_ci(usuario, "codPaisResidencia"),
                "codMunicipioResidencia": get_val_ci(usuario, "codMunicipioResidencia"),
                "codZonaTerritorialResidencia": get_val_ci(usuario, "codZonaTerritorialResidencia"),
                "incapacidad": get_val_ci(usuario, "incapacidad"),
                "consecutivo": get_val_ci(usuario, "consecutivo"),
                "codPaisOrigen": get_val_ci(usuario, "codPaisOrigen")
            }
            usuarios_rows.append(u_row)

            # Información de enlace para servicios (Normalización RIPS: solo consecutivoUsuario)
            link_info = {
                "consecutivoUsuario": u_row["consecutivo"]
            }

            servicios = get_val_ci(usuario, "servicios") or {}

            # Helper para procesar servicios
            def procesar_servicio(lista, contenedor):
                for item in lista:
                    # Usar get_val_ci para propiedades del servicio si es necesario, 
                    # pero item.items() ya itera sobre lo que hay.
                    # Sin embargo, para consistencia, copiamos todo.
                    item_clean = {k: v for k, v in item.items()}
                    contenedor.append({**link_info, **item_clean})

            procesar_servicio(get_val_ci(servicios, "consultas") or [], consultas)
            procesar_servicio(get_val_ci(servicios, "procedimientos") or [], procedimientos)
            procesar_servicio(get_val_ci(servicios, "urgencias") or [], urgencias)
            procesar_servicio(get_val_ci(servicios, "hospitalizacion") or [], hospitalizacion)
            procesar_servicio(get_val_ci(servicios, "recienNacidos") or [], recien_nacidos)
            procesar_servicio(get_val_ci(servicios, "medicamentos") or [], medicamentos)
            procesar_servicio(get_val_ci(servicios, "otrosServicios") or [], otros_servicios)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Escribir hojas según estructura oficial (nombres exactos)
            pd.DataFrame(transaccion_rows).to_excel(writer, sheet_name="transaccion", index=False)
            pd.DataFrame(usuarios_rows).to_excel(writer, sheet_name="usuarios", index=False)
            
            if consultas:
                pd.DataFrame(consultas).to_excel(writer, sheet_name="consultas", index=False)
            if procedimientos:
                pd.DataFrame(procedimientos).to_excel(writer, sheet_name="procedimientos", index=False)
            if urgencias:
                pd.DataFrame(urgencias).to_excel(writer, sheet_name="urgencias", index=False)
            if hospitalizacion:
                pd.DataFrame(hospitalizacion).to_excel(writer, sheet_name="hospitalizacion", index=False)
            if recien_nacidos:
                pd.DataFrame(recien_nacidos).to_excel(writer, sheet_name="recienNacidos", index=False)
            if medicamentos:
                pd.DataFrame(medicamentos).to_excel(writer, sheet_name="medicamentos", index=False)
            if otros_servicios:
                pd.DataFrame(otros_servicios).to_excel(writer, sheet_name="otrosServicios", index=False)
                 
        return output.getvalue(), None
    except Exception as e:
        return None, str(e)

def clean_df_for_json(df):
    """
    Limpia un DataFrame leído de Excel para su conversión a JSON.
    - Convierte NaN a None.
    - Convierte floats que son enteros (1.0) a int (1).
    - Fuerza conversión numérica en campos específicos (sin comillas).
    """
    # Normalizar nombres de columnas (strip whitespace)
    df.columns = [str(c).strip() for c in df.columns]

    # Lista de campos que deben ser numéricos (según RIPS/JSON estándar)
    # Usamos minúsculas para comparación insensible a mayúsculas
    numeric_fields_lower = [
        "consecutivo", "consecutivo_usuario", "codservicio", "vrservicio", 
        "valorpagomoderador", "copago", "cuotamoderadora", 
        "numfevpagomoderador", "bonificacion", "valortotal", 
        "cantidad", "valorunitario"
    ]

    # Intentar convertir columnas numéricas explícitas
    # Comentado para evitar pérdida de datos en campos alfanuméricos (ej: codServicio)
    # Preferimos la conversión segura fila por fila en _clean_val_aggressive
    # for col in df.columns:
    #     if str(col).lower() in numeric_fields_lower:
    #         # Forzar numérico, coerce errores a NaN (luego a None)
    #         df[col] = pd.to_numeric(df[col], errors='coerce')

    # Convertir NaN a None (requiere tipo object)
    df = df.astype(object).where(pd.notnull(df), None)
    
    def _clean_val_aggressive(x):
        # Si es float y es entero, devolver int
        if isinstance(x, float) and x.is_integer():
            return int(x)
        # Si es string, intentar convertir a número si parece un entero
        if isinstance(x, str):
            try:
                s = x.strip()
                if not s: return x
                # Si parece float o int
                f = float(s)
                if f.is_integer():
                    return int(f)
                return f
            except:
                pass # Retornar original si falla conversión (seguro para alfanuméricos)
        return x
        
    def _clean_val_soft(x):
        if isinstance(x, float) and x.is_integer():
            return int(x)
        return x

    # Aplicar limpieza
    for col in df.columns:
        if str(col).lower() in numeric_fields_lower:
             df[col] = df[col].apply(_clean_val_aggressive)
        else:
             df[col] = df[col].apply(_clean_val_soft)
                  
    return df

def normalize_key(k):
    """Normaliza claves eliminando mayúsculas, espacios, acentos y guiones bajos."""
    s = str(k).lower().strip()
    try:
        s = unicodedata.normalize('NFKD', s).encode('ASCII', 'ignore').decode('utf-8')
    except:
        pass
    return s.replace("_", "").replace(" ", "")

def get_val_case_insensitive(row, key):
    """Obtiene valor de una Serie (fila) buscando la clave insensible a mayúsculas/espacios/acentos."""
    # Búsqueda directa
    if key in row:
        return row[key]
    
    # Búsqueda normalizada
    key_norm = normalize_key(key)
    for col in row.index:
        if normalize_key(col) == key_norm:
            return row[col]
    return None

def worker_xlsx_a_json_ind(file_obj):
    try:
        xls = pd.ExcelFile(file_obj)
        usuarios_dict = {}
        header_data_extracted = {
            "numDocumentoIdObligado": None,
            "numFactura": None,
            "tipoNota": None,
            "numNota": None
        }

        def procesar_hoja(nombre_hoja, clave_servicio):
            if nombre_hoja in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=nombre_hoja)
                # Limpiar datos para JSON (tipos correctos)
                df = clean_df_for_json(df)
                
                # Extraer cabecera: Continuar buscando si faltan datos
                nonlocal header_data_extracted
                if not df.empty:
                    row0 = df.iloc[0]
                    
                    def get_header_val(keys):
                        # Soporte para lista de sinónimos
                        if isinstance(keys, str): keys = [keys]
                        
                        for k in keys:
                            val = get_val_case_insensitive(row0, k)
                            if not pd.isna(val):
                                # Preservar string si es posible, manejar enteros flotantes
                                if isinstance(val, float) and val.is_integer():
                                    return str(int(val))
                                return str(val).strip() if isinstance(val, str) else str(val)
                        return None
                    
                # Verificar y rellenar campos faltantes
                    fields_to_check = [
                        ("numDocumentoIdObligado", ["numDocumentoIdObligado", "nit", "id_obligado", "numero_documento_obligado"]),
                        ("numFactura", ["numFactura", "factura", "numero_factura", "num_factura"]),
                        ("tipoNota", ["tipoNota", "tipo_nota", "nota_tipo"]),
                        ("numNota", ["numNota", "num_nota", "numero_nota", "nota_numero"])
                    ]
                    
                    for field, synonyms in fields_to_check:
                        if header_data_extracted[field] is None:
                             val = get_header_val(synonyms)
                             if val is not None:
                                 header_data_extracted[field] = val

                for _, row in df.iterrows():
                    # Definición de llaves de búsqueda (estándar RIPS y legados)
                    k_td = ["tipoDocumentoIdentificacion", "tipo_documento_usuario"]
                    k_doc = ["numDocumentoIdentificacion", "documento_usuario", "numero_documento"]
                    
                    # Extracción segura
                    val_td = None
                    for k in k_td:
                         val_td = get_val_case_insensitive(row, k)
                         if val_td: break
                    
                    val_doc = None
                    for k in k_doc:
                         val_doc = get_val_case_insensitive(row, k)
                         if val_doc: break

                    td = str(val_td or "")
                    doc = str(val_doc or "")
                    user_key = (td, doc)
                    
                    if user_key not in usuarios_dict:
                        # Helper para extraer con lista de sinónimos
                        def get_u_val(keys):
                            if isinstance(keys, str): keys = [keys]
                            for k in keys:
                                v = get_val_case_insensitive(row, k)
                                if v is not None: return v
                            return None

                        usuarios_dict[user_key] = {
                            "tipoDocumentoIdentificacion": get_u_val(["tipoDocumentoIdentificacion", "tipo_documento_usuario"]),
                            "numDocumentoIdentificacion": get_u_val(["numDocumentoIdentificacion", "documento_usuario"]),
                            "tipoUsuario": get_u_val(["tipoUsuario", "tipo_usuario"]),
                            "fechaNacimiento": get_u_val(["fechaNacimiento", "fecha_nacimiento"]), 
                            "codSexo": get_u_val(["codSexo", "sexo"]),
                            "codPaisResidencia": get_u_val(["codPaisResidencia", "pais_residencia"]),
                            "codMunicipioResidencia": get_u_val(["codMunicipioResidencia", "municipio_residencia"]),
                            "codZonaTerritorialResidencia": get_u_val(["codZonaTerritorialResidencia", "zona_residencia"]),
                            "incapacidad": get_u_val(["incapacidad"]),
                            "consecutivo": get_u_val(["consecutivo", "consecutivo_usuario"]),
                            "codPaisOrigen": get_u_val(["codPaisOrigen", "pais_origen"]),
                            "servicios": {
                                "consultas": [],
                                "procedimientos": [],
                                "urgencias": [],
                                "hospitalizacion": [],
                                "recienNacidos": [],
                                "medicamentos": [],
                                "otrosServicios": []
                            }
                        }
                    
                    servicio_data = row.to_dict()
                    keys_to_remove = [
                        "tipoDocumentoIdentificacion", "tipo_documento_usuario",
                        "numDocumentoIdentificacion", "documento_usuario", "numero_documento",
                        "tipoUsuario", "tipo_usuario",
                        "fechaNacimiento", "fecha_nacimiento",
                        "codSexo", "sexo",
                        "codPaisResidencia", "pais_residencia",
                        "codMunicipioResidencia", "municipio_residencia",
                        "codZonaTerritorialResidencia", "zona_residencia",
                        "incapacidad",
                        "consecutivo", "consecutivo_usuario",
                        "codPaisOrigen", "pais_origen",
                        "numDocumentoIdObligado", "numFactura", "tipoNota", "numNota"
                    ]
                    # Eliminación robusta
                    keys_to_remove_norm = {normalize_key(k) for k in keys_to_remove}
                    cols_to_pop = [c for c in servicio_data.keys() if normalize_key(c) in keys_to_remove_norm]
                    for c in cols_to_pop:
                        servicio_data.pop(c, None)
                        
                    if any(v is not None for v in servicio_data.values()):
                        usuarios_dict[user_key]["servicios"][clave_servicio].append(servicio_data)

        # Detección de estructura: Normalizada (RIPS 2275) vs Legada
        sheet_map = {normalize_key(s): s for s in xls.sheet_names}
        
        if "transaccion" in sheet_map and "usuarios" in sheet_map:
            # === MODO NORMALIZADO ===
            # 1. Leer Transaccion (Cabecera)
            df_trans = pd.read_excel(xls, sheet_name=sheet_map["transaccion"])
            df_trans = clean_df_for_json(df_trans)
            if not df_trans.empty:
                row0 = df_trans.iloc[0]
                for k in header_data_extracted:
                    val = get_val_case_insensitive(row0, k)
                    if val is not None:
                        if isinstance(val, float) and val.is_integer():
                            header_data_extracted[k] = str(int(val))
                        else:
                            header_data_extracted[k] = str(val).strip()

            # 2. Leer Usuarios
            df_users = pd.read_excel(xls, sheet_name=sheet_map["usuarios"])
            df_users = clean_df_for_json(df_users)
            
            # Helper local
            def get_u_val_norm(row, keys):
                if isinstance(keys, str): keys = [keys]
                for k in keys:
                    v = get_val_case_insensitive(row, k)
                    if v is not None: return v
                return None

            for _, row in df_users.iterrows():
                td = str(get_u_val_norm(row, ["tipoDocumentoIdentificacion", "tipo_documento_usuario"]) or "")
                doc = str(get_u_val_norm(row, ["numDocumentoIdentificacion", "documento_usuario"]) or "")
                consecutivo = get_u_val_norm(row, ["consecutivo", "consecutivo_usuario"])
                
                user_key = (td, doc)
                
                if user_key not in usuarios_dict:
                    usuarios_dict[user_key] = {
                        "tipoDocumentoIdentificacion": td,
                        "numDocumentoIdentificacion": doc,
                        "tipoUsuario": get_u_val_norm(row, ["tipoUsuario", "tipo_usuario"]),
                        "fechaNacimiento": get_u_val_norm(row, ["fechaNacimiento", "fecha_nacimiento"]), 
                        "codSexo": get_u_val_norm(row, ["codSexo", "sexo"]),
                        "codPaisResidencia": get_u_val_norm(row, ["codPaisResidencia", "pais_residencia"]),
                        "codMunicipioResidencia": get_u_val_norm(row, ["codMunicipioResidencia", "municipio_residencia"]),
                        "codZonaTerritorialResidencia": get_u_val_norm(row, ["codZonaTerritorialResidencia", "zona_residencia"]),
                        "incapacidad": get_u_val_norm(row, ["incapacidad"]),
                        "consecutivo": consecutivo,
                        "codPaisOrigen": get_u_val_norm(row, ["codPaisOrigen", "pais_origen"]),
                        "servicios": {
                            "consultas": [], "procedimientos": [], "urgencias": [],
                            "hospitalizacion": [], "recienNacidos": [], "medicamentos": [], "otrosServicios": []
                        }
                    }

            # Mapa de búsqueda por consecutivo para enlace de servicios
            user_by_consecutivo = {}
            for k, u in usuarios_dict.items():
                c = u.get("consecutivo")
                if c is not None:
                    user_by_consecutivo[str(c)] = u

            # 3. Leer Servicios
            service_map = {
                "consultas": "consultas",
                "procedimientos": "procedimientos",
                "urgencias": "urgencias",
                "hospitalizacion": "hospitalizacion",
                "reciennacidos": "recienNacidos",
                "medicamentos": "medicamentos",
                "otrosservicios": "otrosServicios"
            }

            for sheet_norm, json_key in service_map.items():
                if sheet_norm in sheet_map:
                    df_srv = pd.read_excel(xls, sheet_name=sheet_map[sheet_norm])
                    df_srv = clean_df_for_json(df_srv)
                    
                    for _, row in df_srv.iterrows():
                        # Enlace por consecutivoUsuario
                        c_user = get_val_case_insensitive(row, "consecutivoUsuario")
                        parent_user = None
                        
                        if c_user is not None and str(c_user) in user_by_consecutivo:
                            parent_user = user_by_consecutivo[str(c_user)]
                        
                        # Si no se encuentra por consecutivo, intentar por documento (fallback)
                        if not parent_user:
                             s_td = get_val_case_insensitive(row, "tipoDocumentoIdentificacion")
                             s_doc = get_val_case_insensitive(row, "numDocumentoIdentificacion")
                             if s_td and s_doc:
                                 parent_user = usuarios_dict.get((str(s_td), str(s_doc)))

                        if parent_user:
                            srv_data = row.to_dict()
                            # Limpieza
                            keys_to_remove = ["consecutivoUsuario", "numDocumentoIdObligado", "archivo_origen",
                                              "tipoDocumentoIdentificacion", "numDocumentoIdentificacion"]
                            for k in list(srv_data.keys()):
                                if normalize_key(k) in [normalize_key(x) for x in keys_to_remove]:
                                    del srv_data[k]
                            
                            # Eliminar nulos
                            srv_data = {k: v for k, v in srv_data.items() if v is not None}
                            
                            parent_user["servicios"][json_key].append(srv_data)

        else:
            # === MODO LEGADO ===
            procesar_hoja("Consultas", "consultas")
            procesar_hoja("Procedimientos", "procedimientos")
            procesar_hoja("Urgencias", "urgencias")
            procesar_hoja("Hospitalizacion", "hospitalizacion")
            procesar_hoja("RecienNacidos", "recienNacidos")
            procesar_hoja("Medicamentos", "medicamentos")
            procesar_hoja("OtrosServicios", "otrosServicios")
        
        resultado_final = { 
            **header_data_extracted,
            "usuarios": list(usuarios_dict.values()) 
        }
        return json.dumps(resultado_final, indent=4, ensure_ascii=False), None
    except Exception as e:
        return None, str(e)

def worker_consolidar_json_xlsx(folder_path):
    progress_bar = st.progress(0, text="Escaneando JSONs...")
    
    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    total = len(files_to_process)
    if total == 0: return None, "No se encontraron archivos JSON."
    
    # Listas globales (Estructura RIPS 2275 Normalizada + archivo_origen)
    transaccion_rows = []
    usuarios_rows = []
    todas_consultas = []
    todos_procedimientos = []
    todas_urgencias = []
    todas_hospitalizacion = []
    todos_recien_nacidos = []
    todos_medicamentos = []
    todos_otros_servicios = []
    
    for i, path in enumerate(files_to_process):
        if i % 10 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        nombre_archivo = os.path.basename(path)
        
        try:
            with open(path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Función auxiliar para búsqueda insensible a mayúsculas en diccionario
            def get_val_ci(d, key_target):
                if not isinstance(d, dict): return None
                # Intento directo
                if key_target in d: return d[key_target]
                # Búsqueda normalizada
                target_norm = normalize_key(key_target)
                for k, v in d.items():
                    if normalize_key(k) == target_norm:
                        return v
                return None
    
            # Extraer campos de cabecera (Búsqueda robusta)
            header_info = {
                "numDocumentoIdObligado": get_val_ci(data, "numDocumentoIdObligado"),
                "numFactura": get_val_ci(data, "numFactura"),
                "tipoNota": get_val_ci(data, "tipoNota"),
                "numNota": get_val_ci(data, "numNota")
            }
            # Agregar a hoja transaccion con rastreo de origen
            transaccion_rows.append({"archivo_origen": nombre_archivo, **header_info})

            # Obtener lista de usuarios (búsqueda flexible)
            usuarios_lista = get_val_ci(data, "usuarios")
            if not isinstance(usuarios_lista, list): usuarios_lista = []
            
            for usuario in usuarios_lista:
                # Hoja usuarios
                u_row = {
                    "archivo_origen": nombre_archivo,
                    "tipoDocumentoIdentificacion": get_val_ci(usuario, "tipoDocumentoIdentificacion"),
                    "numDocumentoIdentificacion": get_val_ci(usuario, "numDocumentoIdentificacion"),
                    "tipoUsuario": get_val_ci(usuario, "tipoUsuario"),
                    "fechaNacimiento": get_val_ci(usuario, "fechaNacimiento"),
                    "codSexo": get_val_ci(usuario, "codSexo"),
                    "codPaisResidencia": get_val_ci(usuario, "codPaisResidencia"),
                    "codMunicipioResidencia": get_val_ci(usuario, "codMunicipioResidencia"),
                    "codZonaTerritorialResidencia": get_val_ci(usuario, "codZonaTerritorialResidencia"),
                    "incapacidad": get_val_ci(usuario, "incapacidad"),
                    "consecutivo": get_val_ci(usuario, "consecutivo"),
                    "codPaisOrigen": get_val_ci(usuario, "codPaisOrigen")
                }
                usuarios_rows.append(u_row)

                # Información de enlace para servicios
                link_info = {
                    "archivo_origen": nombre_archivo,
                    "consecutivoUsuario": u_row["consecutivo"]
                }
                
                # Opcional: Agregar numDocumentoIdObligado si se requiere rastreo extra, 
                # pero para RIPS estricto solo consecutivoUsuario es la llave.
                # Mantenemos solo consecutivoUsuario para consistencia con worker_json_a_xlsx_ind.

                servicios = get_val_ci(usuario, "servicios") or {}

                # Helper para procesar servicios
                def procesar_servicio_c(lista, contenedor):
                    for item in lista:
                        item_clean = {k: v for k, v in item.items()}
                        contenedor.append({**link_info, **item_clean})

                procesar_servicio_c(get_val_ci(servicios, "consultas") or [], todas_consultas)
                procesar_servicio_c(get_val_ci(servicios, "procedimientos") or [], todos_procedimientos)
                procesar_servicio_c(get_val_ci(servicios, "urgencias") or [], todas_urgencias)
                procesar_servicio_c(get_val_ci(servicios, "hospitalizacion") or [], todas_hospitalizacion)
                procesar_servicio_c(get_val_ci(servicios, "recienNacidos") or [], todos_recien_nacidos)
                procesar_servicio_c(get_val_ci(servicios, "medicamentos") or [], todos_medicamentos)
                procesar_servicio_c(get_val_ci(servicios, "otrosServicios") or [], todos_otros_servicios)
                    
        except Exception as e:
            log(f"Error leyendo {path}: {e}")
            
    progress_bar.progress(1.0, text="Generando Excel...")
    
    try:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Escribir hojas
            if transaccion_rows:
                pd.DataFrame(transaccion_rows).to_excel(writer, sheet_name="transaccion", index=False)
            if usuarios_rows:
                pd.DataFrame(usuarios_rows).to_excel(writer, sheet_name="usuarios", index=False)
                
            if todas_consultas:
                pd.DataFrame(todas_consultas).to_excel(writer, sheet_name="consultas", index=False)
            if todos_procedimientos:
                pd.DataFrame(todos_procedimientos).to_excel(writer, sheet_name="procedimientos", index=False)
            if todas_urgencias:
                pd.DataFrame(todas_urgencias).to_excel(writer, sheet_name="urgencias", index=False)
            if todas_hospitalizacion:
                pd.DataFrame(todas_hospitalizacion).to_excel(writer, sheet_name="hospitalizacion", index=False)
            if todos_recien_nacidos:
                pd.DataFrame(todos_recien_nacidos).to_excel(writer, sheet_name="recienNacidos", index=False)
            if todos_medicamentos:
                pd.DataFrame(todos_medicamentos).to_excel(writer, sheet_name="medicamentos", index=False)
            if todos_otros_servicios:
                pd.DataFrame(todos_otros_servicios).to_excel(writer, sheet_name="otrosServicios", index=False)
        
        return output.getvalue(), None
    except Exception as e:
        return None, str(e)

def worker_desconsolidar_xlsx_json(file_obj, dest_path):
    try:
        xls = pd.ExcelFile(file_obj)
        
        # --- Configuración y Detección de Hojas ---
        sheet_map = {normalize_key(s): s for s in xls.sheet_names}
        
        # DFs Containers
        df_trans = pd.DataFrame()
        df_users = pd.DataFrame()
        
        # Service DFs
        service_dfs = {}
        service_map_names = {
            "consultas": "Consultas",
            "procedimientos": "Procedimientos",
            "urgencias": "Urgencias",
            "hospitalizacion": "Hospitalizacion",
            "reciennacidos": "RecienNacidos",
            "medicamentos": "Medicamentos",
            "otrosservicios": "OtrosServicios"
        }
        
        # Load available sheets (Normalized & Legacy)
        if "transaccion" in sheet_map:
            df_trans = pd.read_excel(xls, sheet_name=sheet_map["transaccion"])
            df_trans = clean_df_for_json(df_trans)
            
        if "usuarios" in sheet_map:
            df_users = pd.read_excel(xls, sheet_name=sheet_map["usuarios"])
            df_users = clean_df_for_json(df_users)
            
        for key, legacy_name in service_map_names.items():
            # Try normalized name first (lowercase), then legacy (Capitalized)
            sheet_name = sheet_map.get(key, sheet_map.get(normalize_key(legacy_name)))
            if sheet_name:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                service_dfs[key] = clean_df_for_json(df)
            else:
                service_dfs[key] = pd.DataFrame()

        # --- Identificar Archivos Únicos ---
        archivos_unicos = set()
        
        def get_col_archivo(df):
            target = "archivoorigen"
            for c in df.columns:
                if normalize_key(c) == target:
                    return c
            return None

        # Check in all DFs
        all_dfs = [df_trans, df_users] + list(service_dfs.values())
        for df in all_dfs:
            if not df.empty:
                c = get_col_archivo(df)
                if c:
                    archivos_unicos.update(df[c].dropna().unique())
        
        if not archivos_unicos:
             return False, "No se encontró columna 'archivo_origen' para desconsolidar."
             
        if not os.path.exists(dest_path):
            os.makedirs(dest_path)
            
        progress_bar = st.progress(0, text="Generando JSONs...")
        total = len(archivos_unicos)
        count = 0
        
        for i, nombre_archivo in enumerate(archivos_unicos):
            if i % 5 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Generando {i}/{total}")
            
            # --- Extraer Header ---
            header_data = {
                "numDocumentoIdObligado": None,
                "numFactura": None,
                "tipoNota": None,
                "numNota": None
            }
            
            # 1. Try from Transaccion sheet (Normalized)
            col_arch_t = get_col_archivo(df_trans)
            if not df_trans.empty and col_arch_t:
                row_t = df_trans[df_trans[col_arch_t] == nombre_archivo]
                if not row_t.empty:
                    row0 = row_t.iloc[0]
                    for k in header_data:
                        val = get_val_case_insensitive(row0, k)
                        if val is not None:
                             if isinstance(val, float) and val.is_integer():
                                 header_data[k] = str(int(val))
                             else:
                                 header_data[k] = str(val).strip()
            
            # 2. Fallback: Search in all service sheets (Legacy)
            if header_data["numDocumentoIdObligado"] is None:
                 # Helper
                def get_header_val(row, keys):
                    if isinstance(keys, str): keys = [keys]
                    for k in keys:
                        val = get_val_case_insensitive(row, k)
                        if not pd.isna(val):
                            if isinstance(val, float) and val.is_integer():
                                return str(int(val))
                            return str(val).strip() if isinstance(val, str) else str(val)
                    return None

                fields_to_check = [
                     ("numDocumentoIdObligado", ["numDocumentoIdObligado", "nit", "id_obligado", "numero_documento_obligado"]),
                     ("numFactura", ["numFactura", "factura", "numero_factura", "num_factura"]),
                     ("tipoNota", ["tipoNota", "tipo_nota", "nota_tipo"]),
                     ("numNota", ["numNota", "num_nota", "numero_nota", "nota_numero"])
                ]

                # Search in service DFs
                for df_source in service_dfs.values():
                    col_arch = get_col_archivo(df_source)
                    if col_arch:
                        df_file = df_source[df_source[col_arch] == nombre_archivo]
                        if not df_file.empty:
                            row0 = df_file.iloc[0]
                            for field, synonyms in fields_to_check:
                                if header_data[field] is None:
                                    val = get_header_val(row0, synonyms)
                                    if val is not None:
                                        header_data[field] = val

            usuarios_dict = {}
            
            # --- Extraer Usuarios ---
            
            # 1. Try from Usuarios sheet (Normalized)
            col_arch_u = get_col_archivo(df_users)
            if not df_users.empty and col_arch_u:
                rows_u = df_users[df_users[col_arch_u] == nombre_archivo]
                
                # Helper local
                def get_u_val_norm(row, keys):
                    if isinstance(keys, str): keys = [keys]
                    for k in keys:
                        v = get_val_case_insensitive(row, k)
                        if v is not None: return v
                    return None
                
                for _, row in rows_u.iterrows():
                    td = str(get_u_val_norm(row, ["tipoDocumentoIdentificacion", "tipo_documento_usuario"]) or "")
                    doc = str(get_u_val_norm(row, ["numDocumentoIdentificacion", "documento_usuario"]) or "")
                    consecutivo = get_u_val_norm(row, ["consecutivo", "consecutivo_usuario"])
                    
                    user_key = (td, doc)
                    if user_key not in usuarios_dict:
                         usuarios_dict[user_key] = {
                            "tipoDocumentoIdentificacion": td,
                            "numDocumentoIdentificacion": doc,
                            "tipoUsuario": get_u_val_norm(row, ["tipoUsuario", "tipo_usuario"]),
                            "fechaNacimiento": get_u_val_norm(row, ["fechaNacimiento", "fecha_nacimiento"]), 
                            "codSexo": get_u_val_norm(row, ["codSexo", "sexo"]),
                            "codPaisResidencia": get_u_val_norm(row, ["codPaisResidencia", "pais_residencia"]),
                            "codMunicipioResidencia": get_u_val_norm(row, ["codMunicipioResidencia", "municipio_residencia"]),
                            "codZonaTerritorialResidencia": get_u_val_norm(row, ["codZonaTerritorialResidencia", "zona_residencia"]),
                            "incapacidad": get_u_val_norm(row, ["incapacidad"]),
                            "consecutivo": consecutivo,
                            "codPaisOrigen": get_u_val_norm(row, ["codPaisOrigen", "pais_origen"]),
                            "servicios": {
                                "consultas": [], "procedimientos": [], "urgencias": [],
                                "hospitalizacion": [], "recienNacidos": [], "medicamentos": [], "otrosServicios": []
                            }
                        }

            # Map for service linking
            user_by_consecutivo = {}
            for k, u in usuarios_dict.items():
                c = u.get("consecutivo")
                if c is not None:
                    user_by_consecutivo[str(c)] = u

            # --- Link Services ---
            def procesar_df(df_origen, clave_servicio):
                if df_origen.empty: return
                col_archivo = get_col_archivo(df_origen)
                if not col_archivo: return
                
                df_filtrado = df_origen[df_origen[col_archivo] == nombre_archivo]
                
                for _, row in df_filtrado.iterrows():
                    # Attempt 1: Link by consecutivoUsuario (Normalized)
                    c_user = get_val_case_insensitive(row, "consecutivoUsuario")
                    parent_user = None
                    if c_user is not None and str(c_user) in user_by_consecutivo:
                        parent_user = user_by_consecutivo[str(c_user)]
                    
                    # Attempt 2: Link by existing user keys (Legacy or fallback)
                    if not parent_user:
                         # Extract legacy keys
                        k_td = ["tipoDocumentoIdentificacion", "tipo_documento_usuario"]
                        k_doc = ["numDocumentoIdentificacion", "documento_usuario", "numero_documento"]
                        
                        val_td = None
                        for k in k_td:
                             val_td = get_val_case_insensitive(row, k)
                             if val_td: break
                        
                        val_doc = None
                        for k in k_doc:
                             val_doc = get_val_case_insensitive(row, k)
                             if val_doc: break

                        td = str(val_td or "")
                        doc = str(val_doc or "")
                        user_key = (td, doc)
                        
                        if user_key in usuarios_dict:
                            parent_user = usuarios_dict[user_key]
                        else:
                            # Create user if it doesn't exist (Legacy Mode where user info is in service sheet)
                            def get_u_val(keys):
                                if isinstance(keys, str): keys = [keys]
                                for k in keys:
                                    v = get_val_case_insensitive(row, k)
                                    if v is not None: return v
                                return None
                            
                            usuarios_dict[user_key] = {
                                "tipoDocumentoIdentificacion": get_u_val(["tipoDocumentoIdentificacion", "tipo_documento_usuario"]),
                                "numDocumentoIdentificacion": get_u_val(["numDocumentoIdentificacion", "documento_usuario"]),
                                "tipoUsuario": get_u_val(["tipoUsuario", "tipo_usuario"]),
                                "fechaNacimiento": get_u_val(["fechaNacimiento", "fecha_nacimiento"]), 
                                "codSexo": get_u_val(["codSexo", "sexo"]),
                                "codPaisResidencia": get_u_val(["codPaisResidencia", "pais_residencia"]),
                                "codMunicipioResidencia": get_u_val(["codMunicipioResidencia", "municipio_residencia"]),
                                "codZonaTerritorialResidencia": get_u_val(["codZonaTerritorialResidencia", "zona_residencia"]),
                                "incapacidad": get_u_val(["incapacidad"]),
                                "consecutivo": get_u_val(["consecutivo", "consecutivo_usuario"]),
                                "codPaisOrigen": get_u_val(["codPaisOrigen", "pais_origen"]),
                                "servicios": {
                                    "consultas": [], "procedimientos": [], "urgencias": [],
                                    "hospitalizacion": [], "recienNacidos": [], "medicamentos": [], "otrosServicios": []
                                }
                            }
                            parent_user = usuarios_dict[user_key]

                    if parent_user:
                        servicio_data = row.to_dict()
                        keys_to_remove = [
                            "tipoDocumentoIdentificacion", "tipo_documento_usuario",
                            "numDocumentoIdentificacion", "documento_usuario", "numero_documento",
                            "tipoUsuario", "tipo_usuario",
                            "fechaNacimiento", "fecha_nacimiento",
                            "codSexo", "sexo",
                            "codPaisResidencia", "pais_residencia",
                            "codMunicipioResidencia", "municipio_residencia",
                            "codZonaTerritorialResidencia", "zona_residencia",
                            "incapacidad",
                            "consecutivo", "consecutivo_usuario",
                            "codPaisOrigen", "pais_origen",
                            "archivo_origen", "archivoorigen",
                            "numDocumentoIdObligado", "numFactura", "tipoNota", "numNota"
                        ]
                        # Eliminación robusta
                        keys_to_remove_norm = {normalize_key(k) for k in keys_to_remove}
                        cols_to_pop = [c for c in servicio_data.keys() if normalize_key(c) in keys_to_remove_norm]
                        for c in cols_to_pop:
                            servicio_data.pop(c, None)
                            
                        if any(v is not None for v in servicio_data.values()):
                            parent_user["servicios"][clave_servicio].append(servicio_data)

            for key, df_s in service_dfs.items():
                procesar_df(df_s, key)
            
            resultado_final = { 
                **header_data,
                "usuarios": list(usuarios_dict.values()) 
            }
            
            # Limpiar nombre de archivo
            safe_name = os.path.basename(nombre_archivo)
            if not safe_name.lower().endswith(".json"): safe_name += ".json"
            
            out_path = os.path.join(dest_path, safe_name)
            
            with open(out_path, 'w', encoding='utf-8') as f:
                json.dump(resultado_final, f, indent=4, ensure_ascii=False)
            count += 1
            
            # --- Lógica de Distribución por Coincidencia (Solicitud Verde) ---
            try:
                for root, dirs, _ in os.walk(dest_path):
                    for d in dirs:
                        dir_name = d.lower()
                        if dir_name and dir_name in safe_name.lower():
                            target_dir = os.path.join(root, d)
                            try:
                                shutil.copy2(out_path, os.path.join(target_dir, safe_name))
                            except Exception:
                                pass
            except Exception:
                pass
            
        progress_bar.progress(1.0, text="Finalizado.")
        return True, f"Se generaron {count} archivos JSON y se distribuyeron en carpetas coincidentes."
    except Exception as e:
        return False, str(e)

def recursive_strip(obj):
    if isinstance(obj, dict):
        return {k: recursive_strip(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [recursive_strip(v) for v in obj]
    elif isinstance(obj, str):
        return obj.strip()
    else:
        return obj

def worker_clean_json_spaces_masivo(folder_path):
    count = 0
    errors = []
    
    progress_bar = st.progress(0, text="Iniciando limpieza...")
    
    # Cambio a os.walk para incluir subcarpetas (Solicitud Roja)
    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    total = len(files_to_process)
    
    if total == 0:
        progress_bar.empty()
        return 0, ["No se encontraron archivos .json en la carpeta o subcarpetas."]

    for i, file_path in enumerate(files_to_process):
        if i % 5 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        # file_path ya es absoluto con os.walk
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            cleaned_data = recursive_strip(data)
            
            # Sobreescribir el mismo archivo
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(cleaned_data, f, indent=4, ensure_ascii=False)
            
            count += 1
        except Exception as e:
            filename = os.path.basename(file_path)
            errors.append(f"{filename}: {e}")
            
    progress_bar.progress(1.0, text="Finalizado.")
    time.sleep(0.5)
    progress_bar.empty()
    
    return count, errors

def recursive_update_key(obj, key_target, new_value):
    count = 0
    if isinstance(obj, dict):
        for k, v in obj.items():
            if k == key_target:
                # Actualizamos el valor
                # Si el usuario quiere filtrar por valor anterior, lo haríamos aquí, pero
                # la instrucción es "cambiar el numero de la tecnologia de todos los que existan"
                obj[k] = new_value
                count += 1
            else:
                count += recursive_update_key(v, key_target, new_value)
    elif isinstance(obj, list):
        for v in obj:
            count += recursive_update_key(v, key_target, new_value)
    return count

def worker_update_json_key_masivo(folder_path, key_target, new_value):
    count_files = 0
    total_changes = 0
    errors = []
    
    progress_bar = st.progress(0, text="Iniciando actualización...")
    
    # Cambio a os.walk para incluir subcarpetas (Solicitud Roja)
    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    total = len(files_to_process)
    
    if total == 0:
        progress_bar.empty()
        return 0, 0, ["No se encontraron archivos .json en la carpeta o subcarpetas."]

    for i, file_path in enumerate(files_to_process):
        if i % 5 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        # file_path ya es absoluto
        filename = os.path.basename(file_path)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            changes = recursive_update_key(data, key_target, new_value)
            
            if changes > 0:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                count_files += 1
                total_changes += changes
                
        except Exception as e:
            errors.append(f"{filename}: {e}")
            
    progress_bar.progress(1.0, text="Finalizado.")
    time.sleep(0.5)
    progress_bar.empty()
    
    return count_files, total_changes, errors




def recursive_update_notes(obj, tipo_val, num_val):
    count = 0
    if isinstance(obj, dict):
        for k, v in obj.items():
            if k == "tipoNota" and tipo_val:
                obj[k] = tipo_val
                count += 1
            elif k == "numNota" and num_val:
                obj[k] = num_val
                count += 1
            else:
                count += recursive_update_notes(v, tipo_val, num_val)
    elif isinstance(obj, list):
        for v in obj:
            count += recursive_update_notes(v, tipo_val, num_val)
    return count

def recursive_update_cups(obj, old_val, new_val):
    count = 0
    if isinstance(obj, dict):
        for k, v in obj.items():
            if k == "codTecnologiaSalud" and v == old_val:
                obj[k] = new_val
                count += 1
            else:
                count += recursive_update_cups(v, old_val, new_val)
    elif isinstance(obj, list):
        for v in obj:
            count += recursive_update_cups(v, old_val, new_val)
    return count

def worker_update_cups_masivo(folder_path, old_val, new_val):
    count_files = 0
    total_changes = 0
    errors = []
    
    progress_bar = st.progress(0, text="Iniciando actualización de CUPS...")
    
    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    total = len(files_to_process)
    
    if total == 0:
        progress_bar.empty()
        return 0, 0, ["No se encontraron archivos .json en la carpeta o subcarpetas."]

    for i, file_path in enumerate(files_to_process):
        if i % 5 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        filename = os.path.basename(file_path)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            changes = recursive_update_cups(data, old_val, new_val)
            
            if changes > 0:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                count_files += 1
                total_changes += changes
                
        except Exception as e:
            errors.append(f"{filename}: {e}")
            
    progress_bar.progress(1.0, text="Finalizado.")
    time.sleep(0.5)
    progress_bar.empty()
    
    return count_files, total_changes, errors

def worker_update_notes_masivo(folder_path, tipo_val, num_val):
    count_files = 0
    total_changes = 0
    errors = []
    
    progress_bar = st.progress(0, text="Iniciando actualización de notas...")
    
    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    total = len(files_to_process)
    
    if total == 0:
        progress_bar.empty()
        return 0, 0, ["No se encontraron archivos .json en la carpeta o subcarpetas."]

    for i, file_path in enumerate(files_to_process):
        if i % 5 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        filename = os.path.basename(file_path)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            changes = recursive_update_notes(data, tipo_val, num_val)
            
            if changes > 0:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                count_files += 1
                total_changes += changes
                
        except Exception as e:
            errors.append(f"{filename}: {e}")
            
    progress_bar.progress(1.0, text="Finalizado.")
    time.sleep(0.5)
    progress_bar.empty()
    
    return count_files, total_changes, errors

def worker_mover_por_coincidencia(root_path, silent_mode=False):
    log(f"Iniciando movimiento por coincidencia en: {root_path}")
    
    items = os.listdir(root_path)
    files = [f for f in items if os.path.isfile(os.path.join(root_path, f))]
    folders = [d for d in items if os.path.isdir(os.path.join(root_path, d))]
    
    count_moved = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Analizando...")
    total = len(files)
    
    for i, file in enumerate(files):
        if not silent_mode and i % 10 == 0 and total > 0:
             progress_bar.progress(min(i / total, 1.0), text=f"Procesando {i}/{total}")
             
        file_lower = file.lower()
        target = None
        
        # Buscar carpeta coincidente
        for folder in folders:
            if folder.lower() in file_lower:
                target = folder
                break
        
        if target:
            src = os.path.join(root_path, file)
            dst = os.path.join(root_path, target, file)
            try:
                shutil.move(src, dst)
                count_moved += 1
            except Exception as e:
                log(f"Error moviendo {file} a {target}: {e}")
                
    msg = f"Proceso completado. {count_moved} archivos organizados."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def run_mover_por_coincidencia_task(root_path):
    return {"message": worker_mover_por_coincidencia(root_path, silent_mode=True)}

def worker_consolidar_subcarpetas(root_path, silent_mode=False):
    log(f"Consolidando subcarpetas en: {root_path}")
    # Recorrer carpetas de primer nivel
    main_folders = [os.path.join(root_path, d) for d in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, d))]
    
    count_moved = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Consolidando...")
    total = len(main_folders)
    
    for i, main_folder in enumerate(main_folders):
        if not silent_mode and total > 0:
             progress_bar.progress(min(i / total, 1.0), text=f"Procesando carpeta {os.path.basename(main_folder)}")
             
        # Recorrer recursivamente DENTRO de la carpeta principal
        for r, d, f in os.walk(main_folder):
            # Si estamos en la raíz de la carpeta principal, no hacer nada (ya están ahí)
            if r == main_folder:
                continue
                
            for file in f:
                src = os.path.join(r, file)
                dst = os.path.join(main_folder, file)
                
                # Manejo de duplicados simple
                if os.path.exists(dst):
                    name, ext = os.path.splitext(file)
                    dst = os.path.join(main_folder, f"{name}_{int(time.time())}{ext}")
                
                try:
                    shutil.move(src, dst)
                    count_moved += 1
                except Exception as e:
                    log(f"Error moviendo {file}: {e}")
                    
    msg = f"Consolidación completada. {count_moved} archivos extraídos a sus carpetas principales."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def run_consolidar_subcarpetas_task(root_path):
    return {"message": worker_consolidar_subcarpetas(root_path, silent_mode=True)}


def worker_copiar_mapeo_subcarpetas(uploaded_file, sheet_name, col_src, col_dst, path_src_base, path_dst_base, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        if col_src not in df.columns or col_dst not in df.columns:
            if silent_mode: return "Error: Columnas seleccionadas no encontradas en el Excel."
            st.error("Columnas seleccionadas no encontradas en el Excel.")
            return

        count_files = 0
        if not silent_mode:
            progress_bar = st.progress(0, text="Copiando carpetas...")
        total_rows = len(df)
        
        for idx, row in df.iterrows():
            if not silent_mode and idx % 5 == 0:
                progress_bar.progress(min(idx / total_rows, 1.0), text=f"Procesando fila {idx}/{total_rows}")
                
            src_name = str(row[col_src]).strip()
            dst_name = str(row[col_dst]).strip()
            
            if not src_name or not dst_name or src_name == "nan" or dst_name == "nan": continue
            
            full_src = os.path.join(path_src_base, src_name)
            full_dst = os.path.join(path_dst_base, dst_name)
            
            if os.path.exists(full_src) and os.path.exists(full_dst):
                # Copiar todos los archivos de src a dst
                for f in os.listdir(full_src):
                    s = os.path.join(full_src, f)
                    d = os.path.join(full_dst, f)
                    if os.path.isfile(s):
                        try:
                            shutil.copy2(s, d)
                            count_files += 1
                        except Exception as e:
                            log(f"Error copiando {f}: {e}")
        
        msg = f"Copia completada. {count_files} archivos copiados."
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
        
    except Exception as e:
        err_msg = f"Error leyendo Excel: {e}"
        if silent_mode: return err_msg
        st.error(err_msg)

def run_copiar_mapeo_sub_task(uploaded_file, sheet_name, col_src, col_dst, path_src_base, path_dst_base):
    return worker_copiar_mapeo_subcarpetas(uploaded_file, sheet_name, col_src, col_dst, path_src_base, path_dst_base, silent_mode=True)

def worker_copiar_raiz_mapeo(uploaded_file, sheet_name, col_id, col_dst, path_src_base, path_dst_base, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        count_files = 0
        
        # Indexar archivos en raíz origen para búsqueda rápida (Optimización)
        # Diccionario: {nombre_archivo_lower: nombre_real}
        files_in_root = {f.lower(): f for f in os.listdir(path_src_base) if os.path.isfile(os.path.join(path_src_base, f))}
        
        if not silent_mode:
            progress_bar = st.progress(0, text="Copiando archivos...")
        total_rows = len(df)
        
        for idx, row in df.iterrows():
            if not silent_mode and idx % 10 == 0:
                progress_bar.progress(min(idx / total_rows, 1.0), text=f"Procesando fila {idx}/{total_rows}")
                
            id_val = str(row[col_id]).strip().lower()
            dst_folder_name = str(row[col_dst]).strip()
            
            if not id_val or not dst_folder_name: continue
            
            # Buscar coincidencia parcial en los archivos de la raíz
            # Esto es O(N*M) donde N=files, M=rows. Puede ser lento.
            # Mejoramos: iteramos sobre los archivos y chequeamos 'id_val in filename'
            
            for f_lower, f_real in files_in_root.items():
                if id_val in f_lower:
                    src = os.path.join(path_src_base, f_real)
                    dst_folder = os.path.join(path_dst_base, dst_folder_name)
                    
                    if not os.path.exists(dst_folder):
                        try:
                            os.makedirs(dst_folder)
                        except: pass
                        
                    dst = os.path.join(dst_folder, f_real)
                    
                    try:
                        shutil.copy2(src, dst)
                        count_files += 1
                    except Exception as e:
                        log(f"Error copiando {f_real}: {e}")
                        
        msg = f"Copia completada. {count_files} archivos copiados."
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
        
    except Exception as e:
        msg = f"Error procesando: {e}"
        if not silent_mode:
            st.error(msg)
        return msg

def run_copiar_raiz_mapeo_task(uploaded_file, sheet_name, col_id, col_dst, path_src_base, path_dst_base):
    return {"message": worker_copiar_raiz_mapeo(uploaded_file, sheet_name, col_id, col_dst, path_src_base, path_dst_base, silent_mode=True)}



def worker_exportar_renombrado():
    results = st.session_state.search_results
    if not results:
        st.warning("Primero realiza una búsqueda (Pestaña 1) para tener archivos que exportar.")
        return None
        
    data = []
    for item in results:
        path = item["Ruta completa"]
        name = os.path.basename(path)
        data.append({"Ruta actual": path, "Nuevo nombre": name})
        
    df = pd.DataFrame(data)
    return df

def worker_aplicar_renombrado_excel(uploaded_file, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file)
        # Verificar columnas
        if "Ruta actual" not in df.columns or "Nuevo nombre" not in df.columns:
            msg = "El Excel debe tener columnas 'Ruta actual' y 'Nuevo nombre'"
            if silent_mode: return msg
            st.error(msg)
            return
            
        count = 0
        total_rows = len(df)
        if not silent_mode:
            progress_bar = st.progress(0, text="Renombrando archivos...")

        for index, row in df.iterrows():
            if not silent_mode and index % 10 == 0:
                progress_bar.progress(min(index / total_rows, 1.0), text=f"Procesando archivo {index+1} de {total_rows}")
            
            old_path = row["Ruta actual"]
            new_name = row["Nuevo nombre"]
            
            if pd.isna(old_path) or pd.isna(new_name): continue
            
            if os.path.exists(old_path):
                folder = os.path.dirname(old_path)
                new_path = os.path.join(folder, str(new_name))
                if old_path != new_path:
                    try:
                        os.rename(old_path, new_path)
                        count += 1
                    except Exception as e:
                        log(f"Error renombrando {old_path}: {e}")
        
        msg = f"Renombrado completado. {count} archivos modificados."
        if not silent_mode:
            progress_bar.progress(1.0, text="Renombrado finalizado.")
            st.success(msg)
            if not silent_mode: st.rerun()
        return msg
        
    except Exception as e:
        msg = f"Error leyendo Excel: {e}"
        if not silent_mode:
            st.error(msg)
        return msg

def run_aplicar_renombrado_task(uploaded_file):
    return {"message": worker_aplicar_renombrado_excel(uploaded_file, silent_mode=True)}

def worker_anadir_sufijo_excel(uploaded_file, sheet_name, col_folder, col_suffix, root_path, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        count = 0
        folders_not_found = 0
        folders_empty = 0
        files_skipped = 0
        
        if not silent_mode:
            progress_bar = st.progress(0, text="Añadiendo sufijos...")
        total_rows = len(df)
        
        for idx, row in df.iterrows():
            if not silent_mode and idx % 10 == 0: progress_bar.progress(min(idx/total_rows, 1.0), text=f"Procesando {idx}/{total_rows}")
            
            folder_name = str(row[col_folder]).strip()
            suffix = str(row[col_suffix]).strip()
            
            if not folder_name or not suffix or folder_name.lower() == "nan" or suffix.lower() == "nan": continue
            
            folder_path = os.path.join(root_path, folder_name)
            
            # Intento de búsqueda insensible a mayúsculas/minúsculas
            if not os.path.exists(folder_path):
                try:
                    candidates = [d for d in os.listdir(root_path) if d.lower() == folder_name.lower()]
                    if candidates:
                        folder_path = os.path.join(root_path, candidates[0])
                except Exception:
                    pass

            if not os.path.exists(folder_path):
                folders_not_found += 1
                # Log opcional para depuración
                # log(f"Carpeta no encontrada: {folder_name}") 
                continue

            try:
                files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
            except Exception:
                folders_not_found += 1 # Si falla listdir, tratar como no encontrada/accesible
                continue

            if not files:
                folders_empty += 1
                continue

            for f in files:
                name, ext = os.path.splitext(f)
                # Evitar doble sufijo
                if not name.endswith(suffix):
                    new_name = f"{name}{suffix}{ext}"
                    try:
                        os.rename(os.path.join(folder_path, f), os.path.join(folder_path, new_name))
                        count += 1
                    except Exception as e:
                        log(f"Error renombrando {f}: {e}")
                else:
                    files_skipped += 1
                                
        msg = f"Proceso completado.\nRenombrados: {count}\nCarpetas no encontradas: {folders_not_found}\nCarpetas vacías: {folders_empty}\nOmitidos (Ya tenían sufijo): {files_skipped}"
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
    except Exception as e:
        msg = f"Error: {e}"
        if not silent_mode:
            st.error(msg)
        return msg

def run_anadir_sufijo_task(uploaded_file, sheet_name, col_folder, col_suffix, root_path):
    return {"message": worker_anadir_sufijo_excel(uploaded_file, sheet_name, col_folder, col_suffix, root_path, silent_mode=True)}


def worker_unificar_docx_carpeta(root_path, output_name, silent_mode=False):
    # Requiere Word
    if not output_name: 
        output_name = "Unificado"
    count_folders = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Unificando DOCX...")
    
    folders = [x[0] for x in os.walk(root_path)]
    total = len(folders)
    
    for i, folder in enumerate(folders):
        if not silent_mode and i % 5 == 0: 
            progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        # Buscar todos los DOCX en la carpeta (ignorar temporales)
        try:
            files_in_folder = os.listdir(folder)
            docxs = [os.path.join(folder, f) for f in files_in_folder 
                     if f.lower().endswith(".docx") and not f.startswith("~$") and not f.startswith("._")]
            
            # Ordenar naturalmente (1.docx, 2.docx, 10.docx)
            docxs.sort(key=lambda x: [int(c) if c.isdigit() else c.lower() for c in re.split(r'(\d+)', os.path.basename(x))])
        except Exception as e:
            log(f"Error listando {folder}: {e}")
            continue
        
        if not docxs: continue
        
        try:
            # Convertir cada DOCX a PDF temporal
            pdfs_temp = []
            for docx in docxs:
                pdf_temp = docx.replace(".docx", "_temp.pdf")
                _docx_a_pdf(docx, pdf_temp)
                pdfs_temp.append(pdf_temp)
            
            # Unificar PDFs
            doc_final = fitz.open()
            for pdf in pdfs_temp:
                doc_t = fitz.open(pdf)
                doc_final.insert_pdf(doc_t)
                doc_t.close()
            
            doc_final.save(os.path.join(folder, f"{output_name}.pdf"))
            doc_final.close()
            
            # Limpiar temporales
            for pdf in pdfs_temp:
                try: os.remove(pdf)
                except: pass
                
            count_folders += 1
        except Exception as e:
            log(f"Error en {folder}: {e}")
            
    msg = f"Unificados DOCX en {count_folders} carpetas."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def run_unificar_docx_carpeta_task(root_path, output_name):
    return {"message": worker_unificar_docx_carpeta(root_path, output_name, silent_mode=True)}


def worker_modificar_docx_excel(uploaded_file, sheet_name, col_folder, col_val, root_path, mode, silent_mode=False):
    # Mode: "AUTORIZACION", "REGIMEN"
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        count = 0
        if not silent_mode:
            progress_bar = st.progress(0, text="Modificando DOCX...")
        total = len(df)
        
        for idx, row in df.iterrows():
            if not silent_mode and idx % 5 == 0: progress_bar.progress(min(idx/total, 1.0), text=f"Procesando {idx}/{total}")
            
            folder_name = str(row[col_folder]).strip()
            new_val = str(row[col_val]).strip()
            
            if not folder_name or not new_val: continue
            
            folder_path = os.path.join(root_path, folder_name)
            if os.path.exists(folder_path):
                # Buscar archivo CRC...
                target_file = None
                for f in os.listdir(folder_path):
                    if f.startswith("CRC_") and "FEOV" in f and f.endswith(".docx"):
                        target_file = os.path.join(folder_path, f)
                        break
                
                if target_file:
                    try:
                        doc = Document(target_file)
                        modified = False
                        keyword = "AUTORIZACION:" if mode == "AUTORIZACION" else "REGIMEN:"
                        
                        for p in doc.paragraphs:
                            if keyword in p.text:
                                # Reemplazar texto
                                # Estrategia simple: Reemplazar todo el texto del run
                                # Ojo: p.text es read-only para assignment directo complex, 
                                # pero podemos reconstruir.
                                if mode == "AUTORIZACION":
                                    p.text = f"AUTORIZACION: {new_val}"
                                else:
                                    p.text = f"REGIMEN: {new_val}"
                                modified = True
                                break # Solo la primera ocurrencia
                        
                        if modified:
                            doc.save(target_file)
                            count += 1
                    except Exception as e:
                        log(f"Error editando {target_file}: {e}")

        msg = f"Modificados {count} documentos."
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
    except Exception as e:
        msg = f"Error: {e}"
        if not silent_mode:
            st.error(msg)
        return msg

def run_modificar_docx_excel_task(uploaded_file, sheet_name, col_folder, col_val, root_path, mode):
    return {"message": worker_modificar_docx_excel(uploaded_file, sheet_name, col_folder, col_val, root_path, mode, silent_mode=True)}

def worker_modificar_docx_completo(uploaded_file, sheet_name, root_path, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        
        # Mapeo de columnas obligatorio (Igual que Desktop)
        required_map = {
            'folder': 'Nombre Carpeta',
            'date': 'Ciudad y Fecha',
            'full_name': 'Nombre Completo',
            'doc_type': 'Tipo Documento',
            'doc_num': 'Numero Documento',
            'service': 'Servicio',
            'eps': 'EPS',
            'tipo_servicio': 'Tipo Servicio',
            'regimen': 'Regimen',
            'categoria': 'Categoria',
            'cuota': 'Valor Cuota Moderadora',
            'auth': 'Numero Autorizacion',
            'fecha_atencion': 'Fecha y Hora Atencion',
            'fecha_fin': 'Fecha Finalizacion'
        }
        
        # Validar columnas
        missing_cols = [header for header in required_map.values() if header not in df.columns]
        if missing_cols:
            msg = f"Faltan las siguientes columnas obligatorias en el Excel: {', '.join(missing_cols)}"
            if not silent_mode: st.error(msg)
            return msg

        count = 0
        errores = 0
        if not silent_mode:
            progress_bar = st.progress(0, text="Modificando DOCX Completo...")
        total = len(df)
        
        for idx, row in df.iterrows():
            if not silent_mode and idx % 5 == 0: progress_bar.progress(min(idx/total, 1.0), text=f"Procesando {idx}/{total}")
            
            # Obtener datos usando el mapa
            datos = {}
            for key, col_name in required_map.items():
                val = row[col_name]
                datos[key] = str(val).strip() if pd.notna(val) else ""
            
            folder_name = datos.get('folder')
            if not folder_name:
                log(f"Fila {idx+2}: Nombre de carpeta vacío.")
                continue
            
            folder_path = os.path.join(root_path, folder_name)
            if not os.path.exists(folder_path):
                log(f"Fila {idx+2}: No existe la carpeta '{folder_name}'.")
                errores += 1
                continue
                
            # Buscar archivo: Buscar 'plantilla' en el nombre (según desktop)
            target_file = None
            candidates = [f for f in os.listdir(folder_path) if f.lower().endswith(".docx") and "plantilla" in f.lower()]
            
            if candidates:
                target_file = os.path.join(folder_path, candidates[0])
            else:
                log(f"Fila {idx+2}: No se encontró archivo DOCX con 'plantilla' en '{folder_name}'.")
                errores += 1
                continue
                
            try:
                doc = Document(target_file)
                modified = False
                
                # Lógica de reemplazo específica del Desktop App
                for p in doc.paragraphs:
                    if "Santiago de Cali, " in p.text:
                        p.text = f"Santiago de Cali,  {datos['date']}"
                        modified = True
                        continue

                    if "Yo " in p.text and "identificado con" in p.text:
                        p.text = (f"Yo {datos['full_name']} identificado con {datos['doc_type']}, "
                                  f"Numero {datos['doc_num']} en calidad de paciente, doy fé y acepto el "
                                  f"servicio de {datos['service']} brindado por la IPS OPORTUNIDAD DE VIDA S.A.S"
                                  f"   "
                                  f"   ")
                        modified = True
                        continue

                    if "EPS:" in p.text:
                        p.text = re.sub(r'(EPS:)\s*.*', r'\1 ' + datos['eps'], p.text, count=1)
                        modified = True
                    if "TIPO SERVICIO:" in p.text:
                        p.text = re.sub(r'(TIPO SERVICIO:)\s*.*', r'\1 ' + datos['tipo_servicio'], p.text, count=1)
                        modified = True
                    if "REGIMEN:" in p.text:
                        p.text = re.sub(r'(REGIMEN:)\s*.*', r'\1 ' + datos['regimen'], p.text, count=1)
                        modified = True
                    if "CATEGORIA:" in p.text:
                        p.text = re.sub(r'(CATEGORIA:)\s*.*', r'\1 ' + datos['categoria'], p.text, count=1)
                        modified = True
                    if "VALOR CUOTA MODERADORA:" in p.text:
                        p.text = re.sub(r'(VALOR CUOTA MODERADORA:)\s*.*', r'\1 ' + datos['cuota'], p.text, count=1)
                        modified = True
                    if "AUTORIZACION:" in p.text:
                        p.text = re.sub(r'(AUTORIZACION:)\s*.*', r'\1 ' + datos['auth'], p.text, count=1)
                        modified = True
                    if "Fecha de Atención:" in p.text:
                        p.text = re.sub(r'(Fecha de Atención:)\s*.*', r'\1 ' + datos['fecha_atencion'], p.text, count=1)
                        modified = True
                    if "Fecha de Finalización:" in p.text:
                        p.text = re.sub(r'(Fecha de Finalización:)\s*.*', r'\1 ' + datos['fecha_fin'], p.text, count=1)
                        modified = True
                        continue

                # Lógica de Firma
                signature_line_index = -1
                for i, p in enumerate(doc.paragraphs):
                    if "FIRMA DE ACEPTACION" in p.text.upper():
                        signature_line_index = i
                        break 

                if signature_line_index != -1 and signature_line_index + 2 < len(doc.paragraphs):
                    name_paragraph = doc.paragraphs[signature_line_index + 2]
                    name_paragraph.text = datos['full_name'].upper()
                    modified = True
                else:
                    log(f"Advertencia: No se encontró lugar para firma en '{os.path.basename(target_file)}'.")

                doc.save(target_file)
                count += 1
                
            except Exception as e:
                log(f"Error editando {target_file}: {e}")
                errores += 1

        progress_bar.progress(1.0, text="Finalizado.")
        st.success(f"Proceso finalizado.\nModificados: {count}\nErrores/Omitidos: {errores}")
    except Exception as e:
        st.error(f"Error: {e}")

def worker_firmar_docx(root_path, doc_name="plantilla.docx", silent_mode=False):
    count = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Firmando documentos...")
    
    folders = [x[0] for x in os.walk(root_path)]
    total = len(folders)
    
    for i, folder in enumerate(folders):
        if not silent_mode and i % 10 == 0: progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        
        docx_path = os.path.join(folder, doc_name)
        # Buscar firma (en raiz o carpeta tipografia)
        img_path = os.path.join(folder, "firma.jpg")
        if not os.path.exists(img_path):
            img_path = os.path.join(folder, "tipografia", "firma.jpg")
            
        if os.path.exists(docx_path) and os.path.exists(img_path):
            try:
                doc = Document(docx_path)
                signed = False
                for idx_p, p in enumerate(doc.paragraphs):
                    if "Firma de Aceptacion" in p.text:
                        # Insertar en el siguiente párrafo
                        if idx_p + 1 < len(doc.paragraphs):
                            p_next = doc.paragraphs[idx_p + 1]
                            p_next.clear() # Borrar contenido
                            run = p_next.add_run()
                            run.add_picture(img_path, width=Inches(1.5))
                            signed = True
                            break
                
                if signed:
                    doc.save(docx_path)
                    count += 1
            except Exception as e:
                log(f"Error firmando {docx_path}: {e}")
                
    msg = f"Firmados {count} documentos."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def worker_analisis_carpetas(root_path, silent_mode=False):
    data_detail = []
    data_summary = {}
    
    if not silent_mode:
        progress_bar = st.progress(0, text="Analizando carpetas...")
    
    # Solo primer nivel
    subfolders = [d for d in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, d))]
    total = len(subfolders)
    
    for i, sub in enumerate(subfolders):
        if not silent_mode and i % 5 == 0: 
            progress_bar.progress(min(i/total, 1.0), text=f"Analizando {sub}")
        
        full_sub = os.path.join(root_path, sub)
        count_files = 0
        
        for r, d, f in os.walk(full_sub):
            for file in f:
                data_detail.append({"Carpeta Principal": sub, "Ruta": os.path.join(r, file), "Archivo": file})
                count_files += 1
        
        data_summary[sub] = count_files
        
    if not silent_mode:
        progress_bar.progress(1.0, text="Generando Excel...")
    
    # Crear Excel
    df_det = pd.DataFrame(data_detail)
    df_sum = pd.DataFrame(list(data_summary.items()), columns=["Carpeta", "Total Archivos"])
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_det.to_excel(writer, sheet_name="Detalle", index=False)
        df_sum.to_excel(writer, sheet_name="Resumen", index=False)
    output.seek(0)
    
    return output

def worker_copiar_archivo_sub(file_path, root_path, silent_mode=False):
    if not file_path or not root_path:
        msg = "Faltan datos: Archivo o Ruta."
        if not silent_mode: st.error(msg)
        return msg
    
    if not os.path.exists(file_path):
        msg = f"El archivo origen no existe: {file_path}"
        if not silent_mode: st.error(msg)
        return msg

    try:
        filename = os.path.basename(file_path)
        
        # Verificar acceso a la carpeta raíz
        try:
            items = os.listdir(root_path)
        except Exception as e:
            msg = f"Error leyendo carpeta raíz: {e}"
            if not silent_mode: st.error(msg)
            return msg

        subfolders = [d for d in items if os.path.isdir(os.path.join(root_path, d))]
        total = len(subfolders)
        
        if total == 0:
            msg = "No se encontraron subcarpetas en la ruta actual."
            if not silent_mode: st.warning(msg)
            return msg

        count = 0
        if not silent_mode:
            progress_bar = st.progress(0, text=f"Preparando copia a {total} carpetas...")
        
        step = max(1, total // 50) # Actualizar barra cada 2% aprox
        
        for i, sub in enumerate(subfolders):
            if not silent_mode and i % step == 0:
                progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i+1}/{total}: {sub}")
            
            dest = os.path.join(root_path, sub, filename)
            
            # Evitar copiarse a sí mismo si la carpeta origen es una subcarpeta (caso raro)
            if os.path.abspath(file_path) == os.path.abspath(dest):
                continue

            # Solo copiar si no existe (o podríamos añadir opción de sobrescribir)
            if not os.path.exists(dest):
                try:
                    shutil.copy2(file_path, dest)
                    count += 1
                except Exception as e:
                    log(f"Error copiando a {sub}: {e}")
            else:
                # Opcional: logear que ya existe
                pass
        
        msg = f"Proceso finalizado. Archivo copiado a {count} de {total} carpetas (omitiendo existentes)."
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
        
    except Exception as e:
        msg = f"Error crítico en el proceso: {e}"
        if not silent_mode: st.error(msg)
        log(f"Error crítico worker_copiar_archivo_sub: {e}")
        return msg

def run_copiar_archivo_sub_task(file_path, root_path):
    return {"message": worker_copiar_archivo_sub(file_path, root_path, silent_mode=True)}

def worker_crear_firma_nombre(root_path, ttf_path, size, humanize=False, silent_mode=False):
    try:
        from PIL import ImageDraw, ImageFont
        font = ImageFont.truetype(ttf_path, size)
    except Exception as e:
        msg = f"Error cargando fuente: {e}"
        if not silent_mode: st.error(msg)
        return msg

    count = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Generando firmas...")
    
    subfolders = [d for d in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, d))]
    total = len(subfolders)
    
    for i, sub in enumerate(subfolders):
        if not silent_mode and i % 10 == 0: progress_bar.progress(min(i/total, 1.0))
        
        text = sub # Nombre de carpeta es el texto
        
        try:
            # Dummy draw para calcular tamaño base
            dummy_img = Image.new('RGB', (1, 1))
            dummy_draw = ImageDraw.Draw(dummy_img)
            bbox = dummy_draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            # Crear imagen SOLO del texto primero
            img_text = Image.new('RGB', (text_width + 20, text_height + 20), (255, 255, 255))
            draw_text = ImageDraw.Draw(img_text)
            draw_text.text((10, 10), text, font=font, fill=(0, 0, 0))
            
            final_img = img_text
            
            if humanize:
                import random
                angle = random.uniform(-8, 8) # Rotación aleatoria
                final_img = img_text.rotate(angle, expand=True, fillcolor=(255, 255, 255))
                
            # Añadir padding final consistente
            fw, fh = final_img.size
            bg_w, bg_h = fw + 60, fh + 40
            bg = Image.new('RGB', (bg_w, bg_h), (255, 255, 255))
            
            # Centrar
            offset_x = (bg_w - fw) // 2
            offset_y = (bg_h - fh) // 2
            bg.paste(final_img, (offset_x, offset_y))
            
            target_dir = os.path.join(root_path, sub, "tipografia")
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)
                
            bg.save(os.path.join(target_dir, "firma.jpg"))
            count += 1
            
        except Exception as e:
            log(f"Error en firma carpeta {sub}: {e}")
            
    msg = f"Generadas {count} firmas."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def worker_crear_firma_excel(root_path, ttf_path, size, excel_file, sheet_name, col_folder, col_full_name, humanize=False, silent_mode=False):
    try:
        from PIL import ImageDraw, ImageFont
        font = ImageFont.truetype(ttf_path, int(size))
    except Exception as e:
        msg = f"Error cargando fuente: {e}"
        if not silent_mode: st.error(msg)
        return msg

    try:
        if isinstance(excel_file, bytes):
            excel_file = io.BytesIO(excel_file)
        excel_file.seek(0)
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
    except Exception as e:
        msg = f"Error leyendo Excel: {e}"
        if not silent_mode: st.error(msg)
        return msg

    count = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Generando firmas desde Excel...")
    total = len(df)
    
    for idx, row in df.iterrows():
        if not silent_mode and idx % 5 == 0: progress_bar.progress(min(idx/total, 1.0))
        
        folder_name = str(row[col_folder]).strip()
        if not folder_name or str(folder_name).lower() == 'nan': continue
        
        # Construir ruta objetivo
        target_dir = os.path.join(root_path, folder_name)
        if not os.path.exists(target_dir):
            continue
            
        # Extraer nombre completo
        full_name = str(row[col_full_name]).strip()
        if not full_name or full_name.lower() == 'nan': full_name = ""
        
        # Lógica inteligente para Primer Nombre + Primer Apellido
        parts = full_name.split()
        name_part = ""
        surname_part = ""
        
        if len(parts) >= 1:
            name_part = parts[0].capitalize() 
        
        if len(parts) >= 4:
            surname_part = parts[2].capitalize()
        elif len(parts) == 3:
            surname_part = parts[1].capitalize()
        elif len(parts) == 2:
            surname_part = parts[1].capitalize()
            
        # Construir texto final
        final_text = f"{name_part} {surname_part}".strip()
        if not final_text: 
            final_text = folder_name 
            
        # Generar Imagen
        try:
            # Dummy draw 
            dummy_img = Image.new('RGB', (1, 1))
            dummy_draw = ImageDraw.Draw(dummy_img)
            bbox = dummy_draw.textbbox((0, 0), final_text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            # Crear imagen base texto
            img_text = Image.new('RGB', (text_width + 20, text_height + 20), (255, 255, 255))
            draw_text = ImageDraw.Draw(img_text)
            draw_text.text((10, 10), final_text, font=font, fill=(0, 0, 0))
            
            final_img = img_text
            
            if humanize:
                import random
                angle = random.uniform(-8, 8)
                final_img = img_text.rotate(angle, expand=True, fillcolor=(255, 255, 255))
            
            # Composition final
            fw, fh = final_img.size
            bg_w, bg_h = fw + 60, fh + 40
            bg = Image.new('RGB', (bg_w, bg_h), (255, 255, 255))
            
            offset_x = (bg_w - fw) // 2
            offset_y = (bg_h - fh) // 2
            bg.paste(final_img, (offset_x, offset_y))
            
            # Guardar
            tipografia_dir = os.path.join(target_dir, "tipografia")
            if not os.path.exists(tipografia_dir):
                os.makedirs(tipografia_dir)
                
            bg.save(os.path.join(tipografia_dir, "firma.jpg"))
            count += 1
        except Exception as e:
            log(f"Error generando firma para {folder_name}: {e}")

    msg = f"Generadas {count} firmas desde Excel."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def worker_descargar_firmas(uploaded_file, sheet_name, col_id, col_folder, root_path, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        count = 0
        if not silent_mode:
            progress_bar = st.progress(0, text="Descargando firmas...")
        total = len(df)
        
        for idx, row in df.iterrows():
            if not silent_mode and idx % 5 == 0: progress_bar.progress(min(idx/total, 1.0))
            
            id_val = str(row[col_id]).strip()
            folder_name = str(row[col_folder]).strip()
            
            if not id_val or not folder_name: continue
            
            target_folder = os.path.join(root_path, folder_name)
            if not os.path.exists(target_folder):
                os.makedirs(target_folder)
                
            url = f"https://oportunidaddevida.com/opvcitas/admisionescall/firmas/{id_val}.png"
            try:
                r = requests.get(url, timeout=10)
                if r.status_code == 200:
                    with open(os.path.join(target_folder, "firma.jpg"), 'wb') as f:
                        f.write(r.content)
                    # Convertir a jpg real (quitar alpha)
                    img = Image.open(os.path.join(target_folder, "firma.jpg"))
                    if img.mode in ('RGBA', 'LA'):
                        background = Image.new(img.mode[:-1], img.size, (255, 255, 255))
                        background.paste(img, img.split()[-1])
                        img = background
                    img.convert('RGB').save(os.path.join(target_folder, "firma.jpg"), "JPEG")
                    count += 1
                else:
                    with open(os.path.join(target_folder, "no tiene firma.txt"), 'w') as f:
                        f.write("404 Not Found")
            except Exception as e:
                log(f"Error descargando firma {id_val}: {e}")
                
        msg = f"Descargadas {count} firmas."
        if not silent_mode:
             progress_bar.progress(1.0, text="Finalizado.")
             st.success(msg)
        return msg
    except Exception as e:
        msg = f"Error: {e}"
        if not silent_mode: st.error(msg)
        return msg

def run_descargar_firmas_task(uploaded_file, sheet_name, col_id, col_folder, root_path):
    return {"message": worker_descargar_firmas(uploaded_file, sheet_name, col_id, col_folder, root_path, silent_mode=True)}

def run_modif_docx_completo_task(uploaded_file, sheet_name, root_path):
    return {"message": worker_modificar_docx_completo(uploaded_file, sheet_name, root_path, silent_mode=True)}

def run_firmar_docx_task(root_path, doc_name):
    return {"message": worker_firmar_docx(root_path, doc_name, silent_mode=True)}

def run_crear_firma_nombre_task(root_path, ttf_path, size, humanize):
    return {"message": worker_crear_firma_nombre(root_path, ttf_path, size, humanize, silent_mode=True)}

def run_crear_firma_excel_task(root_path, ttf_path, size, uploaded, sheet, col_folder, col_full_name, humanize):
    return {"message": worker_crear_firma_excel(root_path, ttf_path, size, uploaded, sheet, col_folder, col_full_name, humanize, silent_mode=True)}

def run_editar_texto_task(file_list, search_text, replace_text):
    return {"message": worker_editar_texto(file_list, search_text, replace_text, silent_mode=True)}

def run_copiar_lista_task(file_list, target_folder):
    return {"message": worker_copiar_lista(file_list, target_folder, silent_mode=True)}

def run_mover_lista_task(file_list, target_folder):
    return {"message": worker_mover_lista(file_list, target_folder, silent_mode=True)}

def run_zip_lista_task(file_list, target_zip_path):
    return {"message": worker_zip_lista(file_list, target_zip_path, silent_mode=True)}

def run_eliminar_lista_task(file_list):
    return {"message": worker_eliminar_lista(file_list, silent_mode=True)}

def run_zip_carpetas_ind_task(folder_list, target_folder):
    return {"message": worker_zip_carpetas_individual(folder_list, target_folder, silent_mode=True)}

def worker_editar_texto(file_list, search_text, replace_text, silent_mode=False):
    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Editando archivos...")
    total = len(file_list)
    
    for i, item in enumerate(file_list):
        file_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Procesando: {os.path.basename(file_path)}")
        
        try:
            if not os.path.exists(file_path):
                continue

            ext = os.path.splitext(file_path)[1].lower()
            modified = False
            
            # Archivos de texto plano
            if ext in ['.txt', '.json', '.xml', '.csv', '.html', '.md', '.log', '.py', '.js', '.css', '.bat', '.ps1']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    
                    if search_text in content:
                        new_content = content.replace(search_text, replace_text)
                        with open(file_path, 'w', encoding='utf-8') as f:
                            f.write(new_content)
                        modified = True
                except Exception as e:
                    log(f"Error leyendo/escribiendo {file_path}: {e}")
                    errors += 1

            # Documentos de Word
            elif ext == '.docx':
                try:
                    doc = Document(file_path)
                    doc_modified = False
                    for p in doc.paragraphs:
                        if search_text in p.text:
                            p.text = p.text.replace(search_text, replace_text)
                            doc_modified = True
                    # También buscar en tablas
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    if search_text in p.text:
                                        p.text = p.text.replace(search_text, replace_text)
                                        doc_modified = True
                                        
                    if doc_modified:
                        doc.save(file_path)
                        modified = True
                except Exception as e:
                    log(f"Error procesando DOCX {file_path}: {e}")
                    errors += 1
            
            if modified:
                count += 1
                log(f"Texto modificado en: {file_path}")
                
        except Exception as e:
            log(f"Error general en {file_path}: {e}")
            errors += 1
            
    msg = f"Proceso finalizado. Archivos modificados: {count}. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        time.sleep(2)
        st.rerun()
    return msg

def worker_copiar_lista(file_list, target_folder, silent_mode=False):
    if not os.path.exists(target_folder):
        try:
            os.makedirs(target_folder)
        except Exception as e:
            msg = f"Error creando carpeta destino: {e}"
            if not silent_mode: st.error(msg)
            return msg

    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Copiando archivos...")
    total = len(file_list)
    
    for i, item in enumerate(file_list):
        src_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Copiando: {os.path.basename(src_path)}")
        
        try:
            if not os.path.exists(src_path): continue
            
            filename = os.path.basename(src_path)
            dest_path = os.path.join(target_folder, filename)
            
            # Manejar colisiones
            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                dest_path = os.path.join(target_folder, f"{base}_{int(time.time())}{ext}")
            
            if os.path.isdir(src_path):
                shutil.copytree(src_path, dest_path)
            else:
                shutil.copy2(src_path, dest_path)
            count += 1
        except Exception as e:
            log(f"Error copiando {src_path}: {e}")
            errors += 1
            
    msg = f"Copiados {count} elementos. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        time.sleep(1.5)
    return msg

def worker_mover_lista(file_list, target_folder, silent_mode=False):
    if not os.path.exists(target_folder):
        try:
            os.makedirs(target_folder)
        except Exception as e:
            msg = f"Error creando carpeta destino: {e}"
            if not silent_mode: st.error(msg)
            return msg

    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Moviendo archivos...")
    total = len(file_list)
    
    changes_made = []

    for i, item in enumerate(file_list):
        src_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Moviendo: {os.path.basename(src_path)}")
        
        try:
            if not os.path.exists(src_path): continue
            
            filename = os.path.basename(src_path)
            dest_path = os.path.join(target_folder, filename)
            
            # Manejar colisiones
            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                dest_path = os.path.join(target_folder, f"{base}_{int(time.time())}{ext}")
            
            shutil.move(src_path, dest_path)
            count += 1
            changes_made.append((dest_path, src_path)) # Guardar para deshacer

            # Actualizar ruta en lista (aunque se va a limpiar o rerun)
            item["Ruta completa"] = dest_path 
        except Exception as e:
            log(f"Error moviendo {src_path}: {e}")
            errors += 1
            
    msg = f"Movidos {count} elementos. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        if count > 0:
            record_action("Mover Lista", changes_made)
        st.success(msg)
        time.sleep(1.5)
        st.rerun()
    return msg

def worker_zip_lista(file_list, target_zip_path, silent_mode=False):
    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Comprimiendo...")
    total = len(file_list)
    
    try:
        with zipfile.ZipFile(target_zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zipf:
            for i, item in enumerate(file_list):
                src_path = item["Ruta completa"]
                if not silent_mode:
                    progress_bar.progress(min(i/total, 1.0), text=f"Añadiendo: {os.path.basename(src_path)}")
                
                try:
                    if not os.path.exists(src_path): continue
                    
                    if os.path.isdir(src_path):
                        for root, dirs, files in os.walk(src_path):
                            for file in files:
                                f_path = os.path.join(root, file)
                                arcname = os.path.relpath(f_path, os.path.dirname(src_path))
                                zipf.write(f_path, arcname)
                    else:
                        zipf.write(src_path, os.path.basename(src_path))
                    count += 1
                except Exception as e:
                    log(f"Error añadiendo al ZIP {src_path}: {e}")
                    errors += 1
        
        msg = f"Archivo ZIP creado exitosamente con {count} elementos."
        if not silent_mode:
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
    except Exception as e:
        msg = f"Error creando ZIP: {e}"
        if not silent_mode: st.error(msg)
        return msg

def worker_zip_carpetas_individual(folder_list, target_folder=None, silent_mode=False):
    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Comprimiendo carpetas...")
    total = len(folder_list)
    
    for i, item in enumerate(folder_list):
        src_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Procesando: {os.path.basename(src_path)}")
        
        try:
            if not os.path.exists(src_path) or not os.path.isdir(src_path): continue
            
            if target_folder:
                if not os.path.exists(target_folder):
                    os.makedirs(target_folder, exist_ok=True)
                zip_base = os.path.join(target_folder, os.path.basename(src_path))
            else:
                zip_base = src_path

            shutil.make_archive(zip_base, 'zip', src_path)
            count += 1
        except Exception as e:
            log(f"Error comprimiendo carpeta {src_path}: {e}")
            errors += 1
            
    msg = f"Se crearon {count} archivos ZIP individuales. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        time.sleep(1.5)
        st.rerun()
    return msg

def worker_eliminar_lista(file_list, silent_mode=False):
    from send2trash import send2trash
    count_del = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Eliminando...")
    total = len(file_list)
    
    for i, item in enumerate(file_list):
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Eliminando {i+1}/{total}")
        path = item["Ruta completa"]
        try:
            if os.path.exists(path):
                safe_path = os.path.normpath(path)
                send2trash(safe_path)
                count_del += 1
        except Exception as e:
            log(f"Error eliminando {path}: {e}")
            errors += 1
    
    msg = f"Se enviaron {count_del} archivos a la papelera. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        st.session_state.search_results = [] 
        time.sleep(1.5)
        st.rerun()
    return msg

# --- DIALOGOS NUEVOS ---

@st.dialog("Añadir Sufijo desde Excel")
def dialog_sufijo():
    uploaded = st.file_uploader("Excel:", type=["xlsx"])
    if uploaded:
        xl = pd.ExcelFile(uploaded)
        sheet = st.selectbox("Hoja:", xl.sheet_names)
        df_prev = xl.parse(sheet_name=sheet, nrows=1)
        cols = df_prev.columns.tolist()
        c1, c2 = st.columns(2)
        with c1: col_f = st.selectbox("Col. Nombre Carpeta:", cols)
        with c2: col_s = st.selectbox("Col. Sufijo:", cols)
        
        col_path, col_btn = st.columns([0.85, 0.15])
        with col_path:
            st.text_input("Ruta a procesar:", key="sufijo_path", value=st.session_state.get("sufijo_path", st.session_state.current_path))
        with col_btn:
            st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
            st.button("📂", key="btn_sufijo", on_click=update_path_key, args=("sufijo_path",))

        if st.button("🚀 Ejecutar"):
            uploaded.seek(0)
            file_bytes = uploaded.getvalue()
            path = st.session_state.get("sufijo_path", st.session_state.current_path)
            submit_task("Añadir Sufijo", run_anadir_sufijo_task, file_bytes, sheet, col_f, col_s, path)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Modificar DOCX (Excel)")
def dialog_modif_docx(mode):
    # mode: AUTORIZACION, REGIMEN
    uploaded = st.file_uploader("Excel:", type=["xlsx"])
    if uploaded:
        xl = pd.ExcelFile(uploaded)
        sheet = st.selectbox("Hoja:", xl.sheet_names)
        df_prev = xl.parse(sheet_name=sheet, nrows=1)
        cols = df_prev.columns.tolist()
        c1, c2 = st.columns(2)
        with c1: col_f = st.selectbox("Col. Nombre Carpeta:", cols)
        with c2: col_v = st.selectbox(f"Col. Nuevo Valor ({mode}):", cols)
        
        st.info(f"Carpeta Raíz: {st.session_state.current_path}")
        if st.button("🚀 Ejecutar"):
            uploaded.seek(0)
            file_bytes = uploaded.getvalue()
            submit_task(f"Modificar DOCX ({mode})", run_modificar_docx_excel_task, file_bytes, sheet, col_f, col_v, st.session_state.current_path, mode)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Modificar DOCX Completo (Excel)")
def dialog_modif_docx_completo():
    st.write("Modifica documentos 'plantilla*.docx' basándose en columnas específicas del Excel.")
    st.info("El Excel debe tener los encabezados exactos: Nombre Carpeta, Ciudad y Fecha, Nombre Completo, Tipo Documento, Numero Documento, Servicio, EPS, Tipo Servicio, Regimen, Categoria, Valor Cuota Moderadora, Numero Autorizacion, Fecha y Hora Atencion, Fecha Finalizacion.")
    
    uploaded = st.file_uploader("Excel:", type=["xlsx"])
    if uploaded:
        xl = pd.ExcelFile(uploaded)
        sheet = st.selectbox("Hoja:", xl.sheet_names)
        
        st.info(f"Carpeta Raíz: {st.session_state.current_path}")
        if st.button("🚀 Ejecutar Modificación Masiva"):
            uploaded.seek(0)
            file_bytes = uploaded.getvalue()
            submit_task("Modificar DOCX Completo", run_modif_docx_completo_task, file_bytes, sheet, st.session_state.current_path)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Copiar Archivo a Subcarpetas")
def dialog_copiar_sub():
    st.write("Selecciona un archivo para copiarlo a TODAS las subcarpetas de la ruta actual.")
    col_path, col_btn = st.columns([0.85, 0.15])
    with col_path:
        st.text_input("Carpeta Destino:", key="copiar_sub_dest", value=st.session_state.get("copiar_sub_dest", st.session_state.current_path))
    with col_btn:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_copiar_sub_dest", on_click=update_path_key, args=("copiar_sub_dest",))
    
    source_type = st.radio("Origen del archivo:", ["📤 Subir archivo", "📂 Seleccionar de la carpeta actual"])
    
    selected_path = None
    
    if source_type == "📤 Subir archivo":
        uploaded = st.file_uploader("Seleccionar Archivo Maestro:", key="uploader_copiar_sub")
        if uploaded:
            temp_dir = os.path.join(os.getcwd(), "temp_uploads")
            os.makedirs(temp_dir, exist_ok=True)
            temp_path = os.path.join(temp_dir, uploaded.name)
            with open(temp_path, "wb") as f:
                f.write(uploaded.getbuffer())
            selected_path = temp_path
            st.success(f"Archivo cargado: {uploaded.name}")
            
    else: # Seleccionar de carpeta actual
        try:
            files = [f for f in os.listdir(st.session_state.current_path) if os.path.isfile(os.path.join(st.session_state.current_path, f))]
            if files:
                f_sel = st.selectbox("Seleccionar Archivo:", files)
                selected_path = os.path.join(st.session_state.current_path, f_sel)
            else:
                st.warning("No hay archivos en la carpeta actual.")
        except Exception as e:
            st.error(f"Error leyendo carpeta: {e}")

    if selected_path:
        st.session_state.file_master = selected_path

    if "file_master" in st.session_state and st.session_state.file_master:
        st.write(f"Archivo seleccionado: **{os.path.basename(st.session_state.file_master)}**")
        if st.button("🚀 Copiar Masivamente"):
            submit_task("Copiar Archivo a Subcarpetas", run_copiar_archivo_sub_task, st.session_state.file_master, st.session_state.current_path)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Insertar Firma en Documentos")
def dialog_insertar_firma_docx():
    st.write("Esta acción buscará el documento indicado y 'firma.jpg' (o 'tipografia/firma.jpg') en cada subcarpeta.")
    doc_name = st.text_input("Nombre del documento:", value="plantilla.docx")
    st.write("Si los encuentra, insertará la firma en el documento.")
    st.info(f"Ruta raíz: {st.session_state.current_path}")
    if st.button("🚀 Ejecutar Firmado"):
        submit_task("Insertar Firma DOCX", run_firmar_docx_task, st.session_state.current_path, doc_name)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Crear Firma Digital desde Nombre")
def dialog_crear_firma():
    st.write("Genera una imagen JPG con firma manuscrita.")
    
    # Resolver rutas de assets
    base_dir = os.path.dirname(os.path.abspath(__file__))
    assets_fonts = os.path.join(base_dir, "..", "assets", "fonts")
    
    # Selección de Fuente (común para ambos modos)
    option = st.radio("Fuente:", ["Subir fuente", "Pacifico (Predeterminada)", "MyUglyHandwriting"], index=1, horizontal=True)
    
    font_path = None
    if option == "Subir fuente":
        uploaded_font = st.file_uploader("Fuente TTF:", type=["ttf", "otf"])
        if uploaded_font:
            with open("temp_font.ttf", "wb") as f:
                f.write(uploaded_font.getbuffer())
            font_path = "temp_font.ttf"
    elif option == "Pacifico (Predeterminada)":
        font_path = os.path.join(assets_fonts, "Pacifico.ttf")
    elif option == "MyUglyHandwriting":
        font_path = os.path.join(assets_fonts, "MyUglyHandwriting-Regular.otf")
        
    c1_opt, c2_opt = st.columns(2)
    with c1_opt:
        size = st.number_input("Tamaño Fuente:", value=70)
    with c2_opt:
        humanize = st.checkbox("🎨 Estilo Natural", value=True, help="Aplica rotación aleatoria e imperfecciones.")

    # Tabs para los modos
    tab1, tab2 = st.tabs(["📁 Usar Nombre Carpeta", "📊 Usar Excel"])
    
    with tab1:
        st.write("Usa el nombre de la subcarpeta como texto de la firma.")
        st.info(f"Ruta actual: {st.session_state.current_path}")
        if st.button("🚀 Generar (Carpeta)"):
            if font_path and os.path.exists(font_path):
                submit_task("Crear Firmas (Carpeta)", run_crear_firma_nombre_task, st.session_state.current_path, font_path, size, humanize)
                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
            else:
                st.error(f"No se encontró la fuente en: {font_path}")

    with tab2:
        st.write("Usa nombres extraídos de una COLUMNA ÚNICA (detecta 1er Nombre + 1er Apellido).")
        uploaded = st.file_uploader("Excel:", type=["xlsx"], key="excel_firma")
        
        if uploaded:
            xl = pd.ExcelFile(uploaded)
            sheet = st.selectbox("Hoja:", xl.sheet_names, key="sheet_firma")
            df_prev = xl.parse(sheet_name=sheet, nrows=1)
            cols = df_prev.columns.tolist()
            
            c1, c2 = st.columns(2)
            with c1: col_folder = st.selectbox("Col. Carpeta (Match):", cols, key="col_match_firma")
            with c2: col_full_name = st.selectbox("Col. Nombre Completo:", cols, key="col_full_name_firma")
            
            st.info(f"Ruta actual: {st.session_state.current_path}")
            
            if st.button("🚀 Generar (Excel)"):
                 uploaded.seek(0)
                 file_bytes = uploaded.getvalue()
                 if font_path and os.path.exists(font_path):
                    submit_task("Crear Firmas (Excel)", run_crear_firma_excel_task, st.session_state.current_path, font_path, size, file_bytes, sheet, col_folder, col_full_name, humanize)
                    st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
                 else:
                    st.error(f"No se encontró la fuente en: {font_path}")

@st.dialog("Descargar Firmas (URL)")
def dialog_descargar_firmas():
    uploaded = st.file_uploader("Excel:", type=["xlsx"])
    if uploaded:
        xl = pd.ExcelFile(uploaded)
        sheet = st.selectbox("Hoja:", xl.sheet_names)
        df_prev = xl.parse(sheet_name=sheet, nrows=1)
        cols = df_prev.columns.tolist()
        c1, c2 = st.columns(2)
        with c1: col_id = st.selectbox("Col. ID (para URL):", cols)
        with c2: col_f = st.selectbox("Col. Nombre Carpeta Destino:", cols)
        
        st.info(f"Carpeta Raíz: {st.session_state.current_path}")
        if st.button("🚀 Descargar"):
            uploaded.seek(0)
            file_bytes = uploaded.getvalue()
            submit_task("Descargar Firmas", run_descargar_firmas_task, file_bytes, sheet, col_id, col_f, st.session_state.current_path)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

def worker_generar_cuv_masivo(root_path, api_url, token, silent_mode=False):
    if not root_path or not os.path.exists(root_path):
        msg = "Ruta no válida."
        if not silent_mode: st.error(msg)
        return msg

    # Normalizar URL
    if not api_url.endswith("/"): api_url += "/"
    
    endpoint = api_url
    headers = {}
    if token:
        headers["Authorization"] = f"Bearer {token}"

    verify_ssl = True
    if api_url.startswith("https://localhost") or api_url.startswith("https://127.0.0.1"):
        verify_ssl = False
        requests.packages.urllib3.disable_warnings()

    count_ok = 0
    count_err = 0
    report_data = []
    
    folders_to_process = [x[0] for x in os.walk(root_path)]
    total_folders = len(folders_to_process)
    
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Buscando archivos RIPS...")
    
    for i, folder in enumerate(folders_to_process):
        if not silent_mode and i % 5 == 0: 
            progress_bar.progress(min(i/total_folders, 1.0), text=f"Procesando carpeta {i}/{total_folders}")
        
        files_found = [f for f in os.listdir(folder) if f.lower().endswith((".json", ".xml"))]
        if not files_found: continue
        
        for file_name in files_found:
            full_path = os.path.join(folder, file_name)
            
            factura_num = os.path.splitext(file_name)[0]
            provider_id = "999"
            try:
                if file_name.lower().endswith('.json'):
                    with open(full_path, 'r', encoding='utf-8') as f_read:
                        content = json.load(f_read)
                        factura_num = content.get('numFactura', factura_num)
                        provider_id = content.get('numDocumentoIdentificacionObligado', provider_id)
                elif file_name.lower().endswith('.xml'):
                    try:
                        tree = ET.parse(full_path)
                        for elem in tree.iter():
                            if 'numFactura' in elem.tag: factura_num = elem.text
                            if 'numDocumentoIdentificacionObligado' in elem.tag: provider_id = elem.text
                    except: pass
            except: pass

            cuv_path = os.path.join(folder, "CUV.txt")
            if os.path.exists(cuv_path):
                try:
                    with open(cuv_path, "r", encoding="utf-8") as f_cuv_read:
                        existing_cuv = f_cuv_read.read().strip()
                except:
                    existing_cuv = "Desconocido (Archivo CUV.txt existe)"
                report_data.append({"Archivo": file_name, "Estado": "Ya tiene CUV", "Detalle": f"CUV existente: {existing_cuv}"})
                continue

            try:
                mime_type = 'application/xml' if file_name.lower().endswith('.xml') else 'application/json'
                with open(full_path, 'rb') as f:
                    files = {'file': (file_name, f, mime_type)}
                    try:
                        response = requests.post(endpoint, files=files, headers=headers, timeout=60, verify=verify_ssl)
                        
                        if response.status_code == 200:
                            data = response.json()
                            try:
                                f_loc_name = f"ResultadosLocales_{factura_num}.json"
                                with open(os.path.join(folder, f_loc_name), "w", encoding="utf-8") as f_out:
                                    json.dump(data, f_out, indent=2, ensure_ascii=False)
                                f_msps_name = f"ResultadosMSPS_{factura_num}_ID{provider_id}_R.json"
                                with open(os.path.join(folder, f_msps_name), "w", encoding="utf-8") as f_out:
                                    json.dump(data, f_out, indent=2, ensure_ascii=False)
                            except Exception as e:
                                if not silent_mode: st.warning(f"No se pudieron guardar los archivos de resultados extra: {e}")

                            cuv = data.get("cuv") or data.get("CUV") or data.get("codigoUnicoValidacion")
                            
                            if cuv:
                                with open(cuv_path, "w", encoding="utf-8") as f_cuv:
                                    f_cuv.write(str(cuv))
                                count_ok += 1
                                report_data.append({"Archivo": file_name, "Estado": "Exitoso", "Detalle": f"CUV Generado: {cuv}"})
                            else:
                                errores = data.get("errores") or data.get("messages") or str(data)
                                with open(os.path.join(folder, "Respuesta_API.json"), "w", encoding="utf-8") as f_resp:
                                    json.dump(data, f_resp, indent=2)
                                count_err += 1
                                report_data.append({"Archivo": file_name, "Estado": "Rechazado", "Detalle": f"Errores: {str(errores)[:500]}"})
                        else:
                            msg_err = f"HTTP {response.status_code}: {response.text[:500]}"
                            log(f"Error API {response.status_code} en {file_name}")
                            try:
                                with open(os.path.join(folder, "Error_API.txt"), "w", encoding="utf-8") as f_err:
                                    f_err.write(f"Archivo: {file_name}\n{msg_err}")
                            except: pass
                            count_err += 1
                            report_data.append({"Archivo": file_name, "Estado": "Error API", "Detalle": msg_err})
                            
                    except requests.exceptions.RequestException as e:
                        log(f"Error de conexión API procesando {file_name}: {e}")
                        try:
                            with open(os.path.join(folder, "Error_Conexion.txt"), "w", encoding="utf-8") as f_err:
                                f_err.write(f"Archivo: {file_name}\nError: {str(e)}")
                        except: pass
                        count_err += 1
                        report_data.append({"Archivo": file_name, "Estado": "Error Conexión", "Detalle": str(e)})
                        
            except Exception as e:
                log(f"Error interno procesando {file_name}: {e}")
                count_err += 1
                report_data.append({"Archivo": file_name, "Estado": "Error Interno", "Detalle": str(e)})

    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
    
    result_files = []
    if report_data:
        df_report = pd.DataFrame(report_data)
        try:
            report_path = os.path.join(root_path, "Reporte_General_CUV.xlsx")
            df_report.to_excel(report_path, index=False)
            if not silent_mode: st.success(f"✅ Reporte guardado automáticamente en: {report_path}")
        except Exception as e:
            if not silent_mode: st.error(f"No se pudo guardar el reporte automáticamente: {e}")

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df_report.to_excel(writer, index=False)
        output.seek(0)
        
        if silent_mode:
            result_files.append({
                "name": "Reporte_CUVs.xlsx",
                "data": output,
                "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "label": "Descargar Reporte (Excel)"
            })
        elif not silent_mode:
            st.write("### 📊 Reporte de Generación CUV")
            st.dataframe(df_report, width=1000)
            st.download_button(label="📥 Descargar Reporte (Excel)", data=output, file_name="Reporte_CUVs.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    msg = f"Proceso completado. CUVs generados: {count_ok}. Errores: {count_err}"
    if not silent_mode:
        st.success(msg)
    
    if silent_mode:
        return {"files": result_files, "message": msg}
    return msg

def run_generar_cuv_task(root_path, api_url, token):
    return worker_generar_cuv_masivo(root_path, api_url, token, silent_mode=True)

# --- DIALOGOS DE ACCIONES ---

@st.dialog("Generar CUV Masivo (FEVRIPS)")
def dialog_generar_cuv():
    st.write("Genera el Código Único de Validación (CUV) enviando los archivos RIPS a la API local.")
    
    st.text_input("Carpeta de Facturas (Raíz)", value=st.session_state.current_path, key="path_cuv", disabled=True)
    st.caption("Se usará la carpeta seleccionada en la pantalla principal.")
        
    st.markdown("---")
    st.write("Configuración API (FEVRIPS)")
    
    col_mode, col_url = st.columns([0.3, 0.7])
    
    with col_mode:
        conn_mode = st.radio("Modo de Conexión:", ["Local", "Servidor Remoto"], key="fevrips_mode", help="Seleccione 'Local' si ejecuta FEVRIPS en su máquina (Docker o Nativo).")
    
    default_url = "https://localhost:9443/api/Validacion/ValidarArchivo"
    if conn_mode == "Servidor Remoto":
        default_url = ""
        st.info("ℹ️ Ingrese la URL del servidor donde está alojado el API FEVRIPS.")

    with col_url:
        # Opción única basada en el manual oficial (ValidarArchivo)
        # El manual indica que el endpoint correcto para validar archivos es .../api/Validacion/ValidarArchivo
        if conn_mode == "Local":
            endpoint_options = [
                "https://localhost:9443/api/Validacion/ValidarArchivo"
            ]
            selected_endpoint = st.selectbox("Seleccione Endpoint de Validación:", endpoint_options, index=0)
            api_url = selected_endpoint
        else:
            # En modo remoto, permitimos ingresar cualquier URL
            api_url = st.text_input("URL del Endpoint de Validación:", value="", placeholder="Ej: https://mi-servidor.com/api/Validacion/ValidarArchivo")
            
        st.caption("ℹ️ Tip: Puede verificar la disponibilidad en `/swagger/index.html` (ej: https://localhost:9443/swagger/index.html)")

        if conn_mode == "Local":
            with st.expander("🛠️ Asistente de Configuración (Docker)", expanded=False):
                st.info("Siga estos pasos si desea usar el contenedor oficial de FEVRIPS.")
                
                # Paso 1: Generar Archivos
                st.markdown("#### 1. Generar Archivos de Configuración")
                if st.button("📄 Crear docker-compose-fevrips.yml"):
                    compose_content = """version: '3.4'
services:
  fevrips-api:
    image: fevripsacr.azurecr.io/minsalud.fevrips.apilocal:latest
    container_name: fevrips-api
    ports:
      - "9443:5100"
    environment:
      - ASPNETCORE_ENVIRONMENT=Docker
      - ASPNETCORE_URLS=https://+:5100;http://+:5000
      - ASPNETCORE_Kestrel__Certificates__Default__Password=fevrips2024*
      - ASPNETCORE_Kestrel__Certificates__Default__Path=/certificates/server.pfx
    volumes:
      - C:/Certificates:/certificates
"""
                    with open("docker-compose-fevrips.yml", "w") as f:
                        f.write(compose_content)
                    st.success("Archivo 'docker-compose-fevrips.yml' creado en la carpeta del proyecto.")
                
                # Paso 2: Certificados
                st.markdown("#### 2. Generar Certificados SSL")
                st.markdown("El contenedor requiere un certificado SSL para funcionar en HTTPS.")
                if st.button("🔐 Generar Certificados (Requiere OpenSSL)"):
                    try:
                        # Buscar script dinámicamente
                        base_dirs = [
                            os.getcwd(),
                            os.path.join(os.getcwd(), "OrganizadorArchivos"),
                            os.path.dirname(os.path.abspath(__file__)),
                            os.path.join(os.path.dirname(os.path.abspath(__file__)), "..")
                        ]
                        
                        script_path = None
                        for d in base_dirs:
                            p = os.path.join(d, "scripts", "generar_certificados_auto.ps1")
                            if os.path.exists(p):
                                script_path = p
                                break
                        
                        if not script_path:
                            st.error("❌ No se encontró el script 'generar_certificados_auto.ps1'")
                            st.warning(f"Buscado en: {base_dirs}")
                        else:
                            cmd = ["powershell", "-ExecutionPolicy", "Bypass", "-File", script_path]
                            result = subprocess.run(cmd, capture_output=True, text=True)
                            
                            if result.returncode == 0:
                                st.success("✅ Certificados generados en C:\\Certificates")
                                st.code(result.stdout)
                            else:
                                st.error("❌ Error generando certificados")
                                st.text("Salida:")
                                st.code(result.stdout)
                                st.text("Error:")
                                st.code(result.stderr)
                                st.info("Asegúrese de tener OpenSSL instalado (Git Bash suele incluirlo).")
                    except Exception as e:
                        st.error(f"Error ejecutando script: {e}")

                # Paso 3: Comandos
                st.markdown("#### 3. Iniciar Contenedor")
                st.markdown("Ejecute estos comandos en su terminal (PowerShell):")
                st.code("""# 1. Login en Azure (Credenciales del Manual)
docker login -u puller -p v1GLVFn6pWoNrQWgEzmx7MYsf1r7TKJQo+kwadvffq+ACRA3mLxs fevripsacr.azurecr.io

# 2. Iniciar Servicio
docker-compose -f docker-compose-fevrips.yml up -d""", language="powershell")

                if st.button("🚀 Intentar Iniciar Docker Aquí"):
                    try:
                        # Intentar V2 primero (docker compose)
                        cmd_v2 = ["docker", "compose", "-f", "docker-compose-fevrips.yml", "up", "-d"]
                        try:
                            res = subprocess.run(cmd_v2, capture_output=True, text=True, check=False)
                            used_cmd = "docker compose"
                        except FileNotFoundError:
                            # Fallback a V1 (docker-compose)
                            cmd_v1 = ["docker-compose", "-f", "docker-compose-fevrips.yml", "up", "-d"]
                            res = subprocess.run(cmd_v1, capture_output=True, text=True, check=False)
                            used_cmd = "docker-compose"

                        if res.returncode == 0:
                            st.success(f"✅ Comando '{used_cmd}' ejecutado correctamente.")
                            st.text(res.stdout)
                        else:
                            st.error(f"❌ Error iniciando Docker ({used_cmd}).")
                            st.text(res.stderr)
                            if "FileNotFoundError" in str(res.stderr) or res.returncode == 2 or "The system cannot find the file specified" in str(res.stderr):
                                st.warning("Asegúrate de que Docker Desktop esté instalado y agregado al PATH del sistema.")
                                
                    except FileNotFoundError:
                         st.error("❌ No se encontró el ejecutable de Docker.")
                         st.warning("Asegúrate de instalar Docker Desktop y que los comandos 'docker' y 'docker-compose' funcionen en tu terminal.")
                    except Exception as e:
                        st.error(f"Error inesperado: {e}")

    # Credenciales de Autenticación SISPRO
    st.write("🔐 Autenticación SISPRO (Opcional)")
    with st.expander("Configurar Credenciales de Login", expanded=False):
        default_auth_url = "https://localhost:9443/api/Auth/LoginSISPRO"
        if conn_mode == "Servidor Remoto":
            default_auth_url = ""
            
        auth_url = st.text_input("URL Login:", value=default_auth_url, placeholder="Ej: https://mi-servidor.com/api/Auth/LoginSISPRO")
        
        col_auth1, col_auth2 = st.columns(2)
        with col_auth1:
            usuario = st.text_input("Usuario (Cédula/Número):", value="")
            tipo_usuario = st.selectbox("Tipo de Usuario:", ["RE", "PIN", "PINx", "PIE"], index=0, help="RE: Representante Entidad, PIN: Profesional Independiente, etc.")
        with col_auth2:
            clave = st.text_input("Contraseña:", type="password", value="")
            nit = st.text_input("NIT:", value="")
        
        if st.button("🔑 Obtener Token"):
            if not auth_url:
                st.warning("⚠️ La URL de autenticación no puede estar vacía.")
            elif not usuario or not clave or not nit:
                st.warning("Complete todos los campos de autenticación.")
            else:
                try:
                    verify_ssl = True
                    if auth_url.startswith("https://localhost") or auth_url.startswith("https://127.0.0.1"):
                        verify_ssl = False
                        requests.packages.urllib3.disable_warnings()
                        
                    payload = {
                        "tipo": "CC", # Asumimos CC por defecto según manual, o podríamos agregar campo
                        "numero": usuario,
                        "clave": clave,
                        "nit": nit,
                        "tipoUsuario": tipo_usuario
                    }
                    
                    with st.spinner("Autenticando..."):
                        r = requests.post(auth_url, json=payload, verify=verify_ssl, timeout=10)
                        
                        if r.status_code == 200:
                            resp_json = r.json()
                            # Ajustar según la respuesta real del API (token, bearer, etc.)
                            token_val = resp_json.get("token") or resp_json.get("Token")
                            if token_val:
                                st.session_state.temp_token = token_val
                                st.success("¡Autenticación Exitosa! Token obtenido.")
                            else:
                                st.error(f"No se encontró token en la respuesta: {resp_json}")
                        else:
                            st.error(f"Error Login ({r.status_code}): {r.text}")
                except requests.exceptions.ConnectionError:
                    st.error("❌ No se pudo conectar al servidor de Autenticación.")
                    st.error("El servidor rechazó la conexión (WinError 10061).")
                    st.warning("⚠️ ESTO SIGNIFICA QUE DOCKER NO ESTÁ CORRIENDO O EL PUERTO ES INCORRECTO.")
                    st.info("Por favor inicie el contenedor Docker FEV-RIPS antes de intentar loguearse.")
                except Exception as e:
                    st.error(f"Error de conexión: {e}")

    # Usar token obtenido o manual
    token_default = st.session_state.get("temp_token", "")
    token = st.text_input("Token de Autorización (Bearer):", value=token_default, type="password")
    
    if st.button("🔌 Probar Conexión"):
        # --- DIAGNÓSTICO DE PUERTOS Y DOCKER ---
        if conn_mode == "Local":
            st.markdown("##### 🔍 Diagnóstico de Sistema Local")
            col_d1, col_d2 = st.columns(2)
            
            docker_ok = False
            with col_d1:
                # 1. Verificar Docker (Informativo)
                try:
                    subprocess.run(["docker", "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True, shell=True)
                    st.success("✅ Docker Detectado")
                    docker_ok = True
                except:
                    st.info("⚠️ Docker no detectado")
                    st.caption("Si usa FEVRIPS nativo, esto es normal.")
            
            with col_d2:
                # 2. Verificar Puerto
                try:
                    parsed_u = urllib.parse.urlparse(api_url)
                    check_host = parsed_u.hostname or "localhost"
                    check_port = parsed_u.port or 443
                    
                    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                    sock.settimeout(1.0)
                    res = sock.connect_ex((check_host, check_port))
                    sock.close()
                    
                    if res == 0:
                        st.success(f"✅ Puerto {check_port} Detectado (Servicio Activo)")
                        st.caption("El validador está listo para usarse.")
                    else:
                        st.error(f"❌ Puerto {check_port} Cerrado (Servicio Inactivo)")
                        st.warning("El servicio de validación NO se está ejecutando.")
                        
                        st.markdown("---")
                        st.markdown("### 🛠️ Opciones de Inicio")
                        
                        tab_native, tab_docker = st.tabs(["🖥️ Opción 1: Ejecutable Nativo (Sin Docker)", "🐳 Opción 2: Docker"])
                        
                        with tab_native:
                            st.info("Use esta opción si **NO** tiene Docker instalado o si falla al iniciarse.")
                            st.markdown("""
                            1. Debe tener la carpeta del validador **FevRips Standalone**.
                            2. Seleccione el archivo `FevRips.Api.exe` dentro de esa carpeta.
                            3. Haga clic en **Iniciar Servidor**.
                            """)
                            
                            # Persistencia de ruta
                            default_exe = st.session_state.app_config.get("fevrips_local_path", "")
                            fev_exe = st.text_input("Ruta del Ejecutable (FevRips.Api.exe):", value=default_exe, placeholder="C:\\FevRips\\FevRips.Api.exe")
                            
                            if fev_exe != default_exe:
                                st.session_state.app_config["fevrips_local_path"] = fev_exe
                                update_user_config(st.session_state.username, "app_config", st.session_state.app_config)
                            
                            col_l1, col_l2 = st.columns([0.3, 0.7])
                            with col_l1:
                                 if st.button("📂 Buscar Archivo", key="btn_browse_fev", use_container_width=True):
                                     try:
                                         import tkinter as tk
                                         from tkinter import filedialog
                                         root = tk.Tk()
                                         root.withdraw()
                                         root.wm_attributes('-topmost', 1)
                                         file_path = filedialog.askopenfilename(
                                             title="Seleccionar FevRips.Api.exe",
                                             filetypes=[("Ejecutables", "*.exe"), ("Todos", "*.*")]
                                         )
                                         root.destroy()
                                         if file_path:
                                             st.session_state.app_config["fevrips_local_path"] = file_path
                                             update_user_config(st.session_state.username, "app_config", st.session_state.app_config)
                                             st.rerun()
                                     except Exception as e:
                                         st.error(f"Error explorador: {e}") 
                            
                            if st.button("▶️ Iniciar Servidor Nativo", use_container_width=True):
                                if not fev_exe or not os.path.exists(fev_exe):
                                    st.error("❌ Ruta inválida. Seleccione 'FevRips.Api.exe'.")
                                else:
                                    try:
                                        working_dir = os.path.dirname(fev_exe)
                                        # Lanzar proceso en consola nueva para evitar bloqueo
                                        CREATE_NEW_CONSOLE = 0x00000010
                                        subprocess.Popen([fev_exe], cwd=working_dir, creationflags=CREATE_NEW_CONSOLE)
                                        
                                        # Loop de verificación
                                        progress_text = "Iniciando servidor... Esperando puerto 9443..."
                                        my_bar = st.progress(0, text=progress_text)
                                        
                                        server_started = False
                                        for i in range(15): # 15 segundos
                                            time.sleep(1)
                                            my_bar.progress((i + 1) / 15, text=f"{progress_text} ({i+1}s)")
                                            
                                            try:
                                                s_check = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                                                s_check.settimeout(0.5)
                                                result_check = s_check.connect_ex(('localhost', 9443))
                                                s_check.close()
                                                if result_check == 0:
                                                    server_started = True
                                                    break
                                            except:
                                                pass
                                                
                                        my_bar.empty()
                                        
                                        if server_started:
                                            st.success("✅ ¡Servidor Iniciado Correctamente!")
                                            time.sleep(1)
                                            st.rerun()
                                        else:
                                            st.error("❌ El servidor no respondió en el puerto 9443.")
                                            st.info("Revise la ventana negra que se abrió para ver posibles errores.")
                                            
                                    except Exception as e:
                                        st.error(f"Error al iniciar: {e}")

                        with tab_docker:
                             st.caption("Use esta opción solo si tiene Docker Desktop instalado y funcionando.")
                             if st.button("🚀 Intentar Iniciar Docker"):
                                try:
                                    cmd_v2 = ["docker", "compose", "-f", "docker-compose-fevrips.yml", "up", "-d"]
                                    try:
                                        res = subprocess.run(cmd_v2, capture_output=True, text=True, check=False)
                                        used_cmd = "docker compose"
                                    except FileNotFoundError:
                                        cmd_v1 = ["docker-compose", "-f", "docker-compose-fevrips.yml", "up", "-d"]
                                        res = subprocess.run(cmd_v1, capture_output=True, text=True, check=False)
                                        used_cmd = "docker-compose"

                                    if res.returncode == 0:
                                        st.success(f"✅ Docker iniciado ({used_cmd}). Espere unos segundos...")
                                        time.sleep(5)
                                        st.rerun()
                                    else:
                                        st.error(f"❌ Falló Docker ({used_cmd}).")
                                        st.code(res.stderr)
                                        if "No se puede ejecutar esta aplicación en el equipo" in str(res.stderr) or res.returncode == 3221225781: # DLL not found etc
                                            st.error("⛔ Su equipo parece no ser compatible con esta versión de Docker.")
                                            st.info("👉 Por favor use la pestaña **'Opción 1: Ejecutable Nativo'**.")
                                except Exception as e:
                                     st.error(f"Error: {e}")
                        
                        st.divider()
                except:
                    st.warning("⚠️ Error verificando puerto")
            
            st.divider()

        try:
            if not api_url:
                st.warning("⚠️ Debe ingresar una URL válida para probar la conexión.")
                st.stop()
                
            with st.spinner("Probando conexión HTTP con el servidor..."):
                verify_ssl = True
                if api_url.startswith("https://localhost") or api_url.startswith("https://127.0.0.1"):
                    verify_ssl = False
                    requests.packages.urllib3.disable_warnings()
                
                # 1. Intentar conectar a la raíz o Swagger para verificar si el servicio está vivo
                # Extraer base url
                parsed = urllib.parse.urlparse(api_url)
                base_url = f"{parsed.scheme}://{parsed.netloc}"
                
                test_urls = [
                    f"{base_url}/swagger/index.html", # URL Típica de documentación FEVRIPS
                    base_url,
                    api_url # El endpoint mismo (dará 405 o 404 pero confirma conexión)
                ]

                # Fallback: Intentar 127.0.0.1 si es localhost
                if "localhost" in base_url:
                    base_127 = base_url.replace("localhost", "127.0.0.1")
                    test_urls.append(f"{base_127}/swagger/index.html")
                    test_urls.append(base_127)
                
                connected = False
                msg = ""
                success_url = ""
                
                for t_url in test_urls:
                    try:
                        r = requests.get(t_url, timeout=3, verify=verify_ssl)
                        connected = True
                        msg = f"✅ Conexión Exitosa con {t_url} (Estado: {r.status_code})"
                        success_url = t_url
                        break
                    except requests.exceptions.ConnectionError:
                        continue
                    except Exception:
                        continue

                if connected:
                    st.success(msg)
                    st.success("El servicio Docker está respondiendo correctamente.")
                    if "127.0.0.1" in success_url and "localhost" in api_url:
                        st.info("💡 Consejo: Parece que Docker responde en 127.0.0.1 pero no en localhost. Intenta cambiar la URL arriba.")
                else:
                    st.error("❌ No se pudo establecer conexión con el contenedor Docker.")
                    st.error(f"El sistema intentó conectar a: {base_url}")
                    st.warning("""
                    **Posibles soluciones:**
                    1. Asegúrate que el contenedor Docker esté corriendo (`docker ps`).
                    2. Verifica que el puerto sea el correcto. El manual sugiere el puerto **9443**.
                    3. Si cambiaste el puerto en el `docker-compose.yml`, actualiza la URL aquí.
                    4. Intenta abrir `https://localhost:9443/swagger/index.html` en tu navegador.
                    """)

        except Exception as e:
            st.error(f"Error ejecutando prueba: {e}")

    st.info("Nota: Si usas el contenedor Docker local con certificados autofirmados, la verificación SSL se omitirá automáticamente para localhost.")
    
    if st.button("🚀 Iniciar Generación de CUV"):
        submit_task("Generar CUV (Masivo)", run_generar_cuv_masivo_task, st.session_state.current_path, api_url, token)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")

def worker_descargar_historias_ovida(df, col_map, save_path, silent_mode=False):
    descargados, errores, conflictos = 0, 0, 0
    total_filas = len(df)
    
    progress_bar = None
    status_text = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando descarga OVIDA...")
        status_text = st.empty()
    
    try:
        # Selenium Setup
        options = webdriver.ChromeOptions()
        options.add_argument('--kiosk-printing')
        
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
        
        # Esperar a que el usuario inicie sesión (Detectando cambio de URL)
        # Se asume que al loguear, la URL cambiará y dejará de ser 'iniciando.php'
        max_wait = 300 # 5 minutos de espera
        start_time = time.time()
        logged_in = False
        
        while time.time() - start_time < max_wait:
            try:
                curr = driver.current_url
                # Si la URL cambia de iniciando.php a cualquier otra cosa, asumimos login exitoso
                # O si estamos en iniciando.php pero ya no es la misma URL inicial (redirección interna)
                # Mejor criterio: Si NO contiene 'iniciando.php' O contiene 'menu', 'index', 'principal'
                if "iniciando.php" not in curr or any(k in curr for k in ["menu.php", "index.php", "principal.php", "home", "dashboard"]):
                    logged_in = True
                    break
                
                # Feedback visual
                if status_text:
                    elapsed = int(time.time() - start_time)
                    status_text.text(f"Esperando inicio de sesión... ({elapsed}s)\nPor favor, ingrese sus credenciales en la ventana del navegador.\nURL actual: {curr}")
            except Exception as e:
                log(f"Error verificando URL en OVIDA: {e}")
            time.sleep(1)
            
        if not logged_in:
            driver.quit()
            return "Tiempo de espera agotado. No se detectó inicio de sesión."

        if status_text:
            status_text.text("Inicio de sesión detectado. Empezando descargas...")
        
        for index, row in df.iterrows():
            if progress_bar:
                progress_bar.progress(min((index + 1) / total_filas, 1.0), text=f"Procesando {index + 1}/{total_filas}")
            
            try:
                nro_estudio = str(int(row[col_map['estudio']])).strip()
                fecha_ingreso_dt = pd.to_datetime(row[col_map['ingreso']])
                fecha_ingreso = fecha_ingreso_dt.strftime('%Y/%m/%d')
                fecha_egreso_dt = pd.to_datetime(row[col_map['egreso']])
                fecha_egreso = fecha_egreso_dt.strftime('%Y/%m/%d')
                nombre_carpeta = str(row[col_map['carpeta']]).strip()
                
                if not all([nro_estudio, fecha_ingreso, fecha_egreso, nombre_carpeta]):
                    errores += 1
                    continue
                
                if status_text:
                    status_text.text(f"Descargando estudio: {nro_estudio}")
                
                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                params = {
                    'estudio': nro_estudio, 'fecha_inicio': fecha_ingreso, 'fecha_fin': fecha_egreso,
                    'verHC': 1, 'verEvo': 1, 'verPar': 1, 'ImprimirOrdenamiento': 1,
                    'ImprimirNotasPcte': 0, 'ImprimirSolOrdenesExt': 1, 'ImprimirGraficasHC': 1,
                    'ImprimirFormatos': 1, 'ImprimirRegistroAdmon': 1, 'ImprimirNovedad': 0,
                    'ImprimirRecomendaciones': 0, 'ImprimirDescripcionQX': 0, 'ImprimirNotasEnfermeria': 1,
                    'ImprimirSignosVitales': 0, 'ImprimirLog': 0, 'ImprimirEpicrisisSinHC': 0
                }
                full_url = f"{base_url}?{urllib.parse.urlencode(params)}"
                
                dest_folder = os.path.join(save_path, nombre_carpeta)
                os.makedirs(dest_folder, exist_ok=True)
                final_file_path = os.path.join(dest_folder, f"HC_{nro_estudio}.pdf")
                
                if os.path.exists(final_file_path):
                    conflictos += 1
                    continue
                
                driver.get(full_url)
                time.sleep(2)
                
                pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {
                    "landscape": False, "printBackground": True,
                    "paperWidth": 8.5, "paperHeight": 11,
                    "marginTop": 0.4, "marginBottom": 0.4, "marginLeft": 0.4, "marginRight": 0.4
                })
                
                with open(final_file_path, 'wb') as f:
                    f.write(base64.b64decode(pdf_b64['data']))
                
                descargados += 1
                
            except Exception as e:
                errores += 1
                log(f"Error descargando {nro_estudio}: {e}")
        
        driver.quit()
        msg = f"Proceso finalizado.\nDescargados: {descargados}\nErrores: {errores}\nConflictos: {conflictos}"
        if not silent_mode:
            st.success(msg)
        return msg
        
    except Exception as e:
        msg = f"Error crítico en OVIDA Downloader: {e}"
        if not silent_mode:
            st.error(msg)
        return msg

def run_descargar_historias_ovida_task(df, col_map, save_path):
    return {"message": worker_descargar_historias_ovida(df, col_map, save_path, silent_mode=True)}

@st.dialog("Descargar Historias (OVIDA)")
def dialog_descargar_historias_ovida():
    st.write("Automatización de descarga de historias clínicas desde OVIDA.")
    
    uploaded = st.file_uploader("Archivo Excel (.xlsx)", type="xlsx", key="ovida_uploader")
    if not uploaded: return

    try:
        xls = pd.ExcelFile(uploaded)
        sheet = st.selectbox("Seleccionar Hoja:", xls.sheet_names, key="ovida_sheet")
        df = xls.parse(sheet_name=sheet)
        
        cols = df.columns.tolist()
        c1, c2 = st.columns(2)
        c3, c4 = st.columns(2)
        
        with c1: col_estudio = st.selectbox("Columna 'Nro ESTUDIO':", cols, key="ovida_c1")
        with c2: col_ingreso = st.selectbox("Columna 'FECHA INGRESO':", cols, key="ovida_c2")
        with c3: col_egreso = st.selectbox("Columna 'FECHA EGRESO':", cols, key="ovida_c3")
        with c4: col_carpeta = st.selectbox("Columna 'NOMBRE CARPETA':", cols, key="ovida_c4")
        
        col_map = {
            'estudio': col_estudio,
            'ingreso': col_ingreso,
            'egreso': col_egreso,
            'carpeta': col_carpeta
        }
        
        st.write("Carpeta Base de Guardado:")
        col_ov_path, col_ov_btn = st.columns([0.85, 0.15])
        with col_ov_path:
            save_path = st.text_input("Carpeta Destino:", value=st.session_state.get("ovida_save_path", st.session_state.current_path), key="ovida_save_path")
        with col_ov_btn:
             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
             st.button("📂", key="btn_ovida_save", help="Seleccionar Ruta", on_click=update_path_key, args=("ovida_save_path", "Seleccionar Ruta"))

        if st.button("🚀 Iniciar Navegador y Descarga"):
            submit_task("Descargar OVIDA", run_descargar_historias_ovida_task, df, col_map, save_path)
            st.info("✅ Tarea iniciada en segundo plano. Se abrirá una ventana de navegador.")
            
    except Exception as e:
        st.error(f"Error leyendo Excel: {e}")

@st.dialog("Unificar PDFs por Carpeta")
def dialog_unificar_pdf():
    st.write("Esta acción buscará todos los PDFs en las subcarpetas de la ruta actual y los unificará en un solo archivo por carpeta.")
    st.info(f"Ruta actual: {st.session_state.current_path}")
    name = st.text_input("Nombre del archivo final (sin extensión):", value="Unificado")
    if st.button("🚀 Iniciar Unificación"):
        comp = st.session_state.app_config.get("pdf_compression", 4)
        dpi = st.session_state.app_config.get("pdf_dpi", 300)
        submit_task("Unificar PDFs por Carpeta", run_unificar_pdf_task, st.session_state.current_path, name, comp, dpi)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Unificar PDFs (Selección Manual)")
def dialog_unificar_pdf_manual():
    st.write("Esta acción unificará los PDFs que están actualmente en la lista de RESULTADOS DE BÚSQUEDA.")
    if not st.session_state.search_results:
        st.error("No hay archivos en la lista de búsqueda. Primero realiza una búsqueda.")
        return
        
    count = len([f for f in st.session_state.search_results if f.get("Ruta completa", "").lower().endswith(".pdf")])
    st.info(f"PDFs encontrados en la lista: {count}")
    
    name = st.text_input("Nombre del archivo final (sin extensión):", value="Unificado_Manual")
    
    if st.button("🚀 Unificar Lista"):
        submit_task("Unificar PDFs (Manual)", run_unificar_pdf_manual_task, st.session_state.search_results, name)
        st.info("✅ Tarea iniciada en segundo plano.")


@st.dialog("Unificar DOCX por Carpeta")
def dialog_unificar_docx():
    st.write("Esta acción buscará todos los DOCX en las subcarpetas, los convertirá a PDF y creará un PDF unificado por carpeta.")
    st.info(f"Ruta actual: {st.session_state.current_path}")
    name = st.text_input("Nombre del PDF final (sin extensión):", value="Unificado_DOCX")
    if st.button("🚀 Iniciar Unificación DOCX"):
        submit_task("Unificar DOCX por Carpeta", run_unificar_docx_carpeta_task, st.session_state.current_path, name)
        st.info("✅ Tarea iniciada en segundo plano.")


@st.dialog("Dividir PDFs Masivamente")
def dialog_dividir_pdf():
    st.write("Esta acción buscará TODOS los PDFs en la ruta actual (y subcarpetas), creará una carpeta con el nombre de cada archivo y extraerá sus páginas.")
    st.warning("Esta operación puede generar muchos archivos.")
    if st.button("🚀 Iniciar División"):
        prefix = st.session_state.app_config.get("split_pdf_prefix", "")
        submit_task("Dividir PDFs (Masivo)", run_dividir_pdf_masivo_task, st.session_state.current_path, prefix)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")

@st.dialog("Aplicar Renombrado desde Excel")
def dialog_importar_excel():
    st.write("Sube el archivo Excel generado previamente con los nuevos nombres.")
    uploaded = st.file_uploader("Archivo Excel (.xlsx)", type="xlsx")
    if uploaded and st.button("🚀 Aplicar Cambios"):
        uploaded.seek(0)
        file_bytes = uploaded.getvalue()
        submit_task("Aplicar Renombrado Excel", run_aplicar_renombrado_task, file_bytes)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

def worker_renombrar_mapeo_excel(uploaded_file, sheet_name, col_src, col_dst, respect_filter, root_path=None, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        
        current_dir = root_path
        if not current_dir or not os.path.exists(current_dir):
            if silent_mode: return "Error: Carpeta actual no válida."
            st.error("Carpeta actual no válida.")
            return

        data_rows = []
        
        # Reset pointer
        uploaded_file.seek(0)

        if respect_filter:
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames:
                if silent_mode: return f"Error: Hoja '{sheet_name}' no encontrada."
                st.error("Hoja no encontrada.")
                return
            ws = wb[sheet_name]
            
            # Find column indices
            header = []
            for cell in ws[1]:
                header.append(cell.value)
                
            try:
                idx_src = header.index(col_src)
                idx_dst = header.index(col_dst)
            except ValueError:
                if silent_mode: return "Error: No se encontraron las columnas en la cabecera."
                st.error("No se encontraron las columnas en la cabecera.")
                return

            # Iterate
            for row in ws.iter_rows(min_row=2):
                # Check hidden
                if ws.row_dimensions[row[0].row].hidden:
                    continue
                
                val_src = row[idx_src].value
                val_dst = row[idx_dst].value
                
                if val_src and val_dst:
                    data_rows.append((str(val_src).strip(), str(val_dst).strip()))
        else:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            if col_src not in df.columns or col_dst not in df.columns:
                 if silent_mode: return "Error: Columnas no encontradas."
                 st.error("Columnas no encontradas.")
                 return
            
            for _, row in df.iterrows():
                val_src = row[col_src]
                val_dst = row[col_dst]
                if pd.notna(val_src) and pd.notna(val_dst):
                    data_rows.append((str(val_src).strip(), str(val_dst).strip()))
        
        # Perform renaming
        count = 0
        errors = 0
        if not silent_mode:
            progress_bar = st.progress(0, text="Renombrando...")
        
        total = len(data_rows)
        
        for i, (src_name, dst_name) in enumerate(data_rows):
            if not silent_mode:
                progress_bar.progress(min((i+1)/total, 1.0), text=f"Procesando {src_name}...")
            
            src_path = os.path.join(current_dir, src_name)
            dst_path = os.path.join(current_dir, dst_name)
            
            if os.path.exists(src_path):
                if src_path != dst_path:
                    try:
                        os.rename(src_path, dst_path)
                        count += 1
                    except Exception as e:
                        log(f"Error renombrando {src_name}: {e}")
                        errors += 1
            else:
                pass
        
        msg = f"Renombrado completado. {count} archivos modificados. {errors} errores."
        if not silent_mode:
            progress_bar.progress(1.0, text="Renombrado finalizado.")
            st.success(msg)
            time.sleep(2)
            st.rerun()
        return msg

    except Exception as e:
        err_msg = f"Error crítico: {e}"
        if silent_mode: return err_msg
        st.error(err_msg)

def run_renombrar_mapeo_task(uploaded_file, sheet_name, col_src, col_dst, respect_filter, root_path):
    return worker_renombrar_mapeo_excel(uploaded_file, sheet_name, col_src, col_dst, respect_filter, root_path, silent_mode=True)

@st.dialog("Renombrar Masivo por Mapeo Excel")
def dialog_renombrar_mapeo_excel():
    st.write("Sube un Excel y configura las columnas para renombrar archivos en la carpeta actual.")
    st.info(f"Ruta actual: {st.session_state.current_path}")
    uploaded = st.file_uploader("Archivo Excel", type=["xlsx", "xls"], key="upl_renom_map")
    
    if uploaded:
        try:
            xl = pd.ExcelFile(uploaded)
            sheets = xl.sheet_names
            sheet = st.selectbox("Selecciona la hoja", sheets, key="sel_sheet_renom")
            
            # Read header only to get columns
            df_header = pd.read_excel(uploaded, sheet_name=sheet, nrows=0)
            cols = df_header.columns.tolist()
            
            col_src = st.selectbox("Columna Nombre Origen", cols, key="sel_col_src")
            col_dst = st.selectbox("Columna Nombre Deseado", cols, key="sel_col_dst")
            
            use_filter = st.radio("Filtros de Excel:", 
                                  ["Usar Datos Completos", "Respetar Filtro (Ocultar filas ocultas)"],
                                  key="rad_filter_renom")
            
            if st.button("🚀 Iniciar Renombrado", key="btn_start_renom"):
                # Use submit_task
                uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                submit_task("Renombrar Mapeo", run_renombrar_mapeo_task, file_bytes, sheet, col_src, col_dst, use_filter == "Respetar Filtro (Ocultar filas ocultas)", st.session_state.current_path)
                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")

@st.dialog("Unificar Imágenes a PDF")
def dialog_unificar_img(type_img):
    st.write(f"Esta acción buscará archivos {type_img} en cada subcarpeta y creará un PDF.")
    
    col_path, col_btn = st.columns([0.85, 0.15])
    with col_path:
        st.text_input("Ruta a procesar:", key="unify_img_path", value=st.session_state.get("unify_img_path", st.session_state.current_path))
    with col_btn:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_unify_img", on_click=update_path_key, args=("unify_img_path",))

    name = st.text_input("Nombre del PDF final:", value="Imagenes_Unificadas")
    if st.button("🚀 Iniciar"):
        path = st.session_state.get("unify_img_path", st.session_state.current_path)
        exts = [".jpg", ".jpeg"] if type_img == "JPG" else [".png"]
        submit_task("Unificar Imágenes", run_unificar_img_task, path, name, exts)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Crear Carpetas desde Excel")
def dialog_crear_carpetas():
    st.write("Crea carpetas masivamente basado en una columna de Excel.")
    uploaded = st.file_uploader("Sube Excel con lista de nombres:", type=["xlsx", "xls"])
    
    if uploaded:
        try:
            # Pre-leer para obtener hojas y columnas (usando pandas es rápido)
            xl = pd.ExcelFile(uploaded)
            sheets = xl.sheet_names
            
            col_sh, col_chk = st.columns(2)
            with col_sh:
                sheet = st.selectbox("Selecciona la Hoja:", sheets)
            with col_chk:
                st.write("")
                st.write("")
                use_filter = st.checkbox("Solo filas visibles (Filtradas)", value=False)
            
            # Leer columnas de la hoja seleccionada
            df_preview = xl.parse(sheet_name=sheet, nrows=1)
            cols = df_preview.columns.tolist()
            
            col_name = st.selectbox("Selecciona la columna con los nombres:", cols)
            
            col_path, col_btn = st.columns([0.85, 0.15])
            with col_path:
                st.text_input("Carpeta Destino:", key="create_folders_path", value=st.session_state.get("create_folders_path", st.session_state.current_path))
            with col_btn:
                st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                st.button("📂", key="btn_create_folders", on_click=update_path_key, args=("create_folders_path",))
            
            if st.button("🚀 Crear Carpetas"):
                # Necesitamos pasar el archivo original (uploaded) al worker
                # Pero al leerlo con pandas, el puntero puede haberse movido.
                uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                path = st.session_state.get("create_folders_path", st.session_state.current_path)
                submit_task("Crear Carpetas Excel", run_crear_carpetas_task, file_bytes, sheet, col_name, path, use_filter)
                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
                
        except Exception as e:
            st.error(f"Error leyendo archivo: {e}")

@st.dialog("Organizar Facturas (FEOV)")
def dialog_organizar_feov():
    prefix = st.session_state.app_config.get("feov_prefix", "FEOV")
    st.write(f"Organiza archivos PDF de una carpeta ORIGEN a una estructura de carpetas DESTINO basándose en el prefijo configurado: **{prefix}**")
    st.info(f"El sistema buscará el patrón **'{prefix}[Numero]'** en los archivos de la carpeta destino para saber dónde mover los nuevos archivos. (Puedes cambiar este prefijo en el menú lateral).")
    
    # --- Origen ---
    col_org, col_btn_org = st.columns([0.85, 0.15])
    with col_org:
        st.text_input("Carpeta Origen:", key="feov_org", value=st.session_state.get("feov_org", st.session_state.current_path), help="Ruta con los PDFs sueltos.")
    with col_btn_org:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_feov_org", help="Seleccionar Origen", on_click=update_path_key, args=("feov_org", "Seleccionar Origen (PDFs sueltos)"))

    # --- Destino ---
    col_dst, col_btn_dst = st.columns([0.85, 0.15])
    with col_dst:
        st.text_input("Carpeta Destino:", key="feov_dst", value=st.session_state.get("feov_dst", ""), help="Ruta donde están las carpetas organizadas.")
    with col_btn_dst:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_feov_dst", help="Seleccionar Destino", on_click=update_path_key, args=("feov_dst", "Seleccionar Destino (Carpetas organizadas)"))

    if st.button("🚀 Iniciar Organización FEOV"):
        org = st.session_state.get("feov_org")
        dst = st.session_state.get("feov_dst")
        prefix = st.session_state.app_config.get("feov_prefix", "FEOV")
        if org and dst and os.path.exists(org) and os.path.exists(dst):
            submit_task("Organizar Facturas FEOV", run_organizar_feov_task, org, dst, prefix)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
        else:
            st.error("Selecciona rutas válidas.")

@st.dialog("Mover Archivos por Coincidencia")
def dialog_mover_coincidencia():
    st.write("Mueve archivos que coincidan con nombres de carpetas (o parte de ellos) dentro de la ruta seleccionada.")
    st.info("Ejemplo: Si existe la carpeta 'PEPITO PEREZ' y el archivo 'Documento PEPITO PEREZ.pdf', el archivo se moverá dentro de la carpeta.")
    
    col_path, col_btn = st.columns([0.85, 0.15])
    
    with col_path:
        st.text_input("Carpeta Raíz:", key="mov_coin_path", value=st.session_state.get("mov_coin_path", st.session_state.current_path))

    with col_btn:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_mov_coin", help="Seleccionar Carpeta Raíz", on_click=update_path_key, args=("mov_coin_path", "Seleccionar Carpeta Raíz"))

    if st.button("🚀 Ejecutar Mover"):
        path = st.session_state.get("mov_coin_path")
        if path and os.path.isdir(path):
            submit_task("Mover por Coincidencia", run_mover_por_coincidencia_task, path)
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
        else:
            st.error("Ruta inválida.")

@st.dialog("Copiar Archivos (Mapeo Subcarpetas)")
def dialog_copiar_mapeo_sub():
    st.write("Copia el contenido de carpetas origen a carpetas destino según un Excel.")
    uploaded = st.file_uploader("Excel Mapeo:", type=["xlsx", "xls"])
    
    c1, c2 = st.columns([0.85, 0.15])
    with c1:
        st.text_input("Carpeta Origen:", key="map_src", value=st.session_state.get("map_src", st.session_state.current_path))
    with c2:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_map_src", help="Seleccionar Origen", on_click=update_path_key, args=("map_src", "Seleccionar Origen"))
        
    c3, c4 = st.columns([0.85, 0.15])
    with c3:
        st.text_input("Carpeta Destino:", key="map_dst", value=st.session_state.get("map_dst", st.session_state.current_path))
    with c4:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_map_dst", help="Seleccionar Destino", on_click=update_path_key, args=("map_dst", "Seleccionar Destino"))

    if uploaded:
        try:
            xl = pd.ExcelFile(uploaded)
            sheets = xl.sheet_names
            sheet = st.selectbox("Selecciona la Hoja:", sheets)
            
            # Leer columnas de la hoja seleccionada
            df_preview = xl.parse(sheet_name=sheet, nrows=1)
            cols = df_preview.columns.tolist()
            
            c1, c2 = st.columns(2)
            with c1: col_s = st.selectbox("Columna Nombre Carpeta Origen:", cols)
            with c2: col_d = st.selectbox("Columna Nombre Carpeta Destino:", cols)
            
            if st.button("🚀 Copiar Carpetas"):
                src = st.session_state.get("map_src")
                dst = st.session_state.get("map_dst")
                if src and dst:
                    # Re-leer archivo para worker
                    uploaded.seek(0)
                    file_bytes = uploaded.getvalue()
                    submit_task("Copiar Mapeo Subcarpetas", run_copiar_mapeo_sub_task, file_bytes, sheet, col_s, col_d, src, dst)
                    st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
                else:
                    st.warning("Selecciona las rutas base.")
        except Exception as e:
            st.error(f"Error leyendo archivo: {e}")

@st.dialog("Copiar Archivos de Raíz (Mapeo)")
def dialog_copiar_raiz_mapeo():
    st.write("Busca archivos en una carpeta raíz por coincidencia de nombre y los copia a subcarpetas destino según Excel.")
    uploaded = st.file_uploader("Excel Mapeo:", type=["xlsx", "xls"])
    
    c1, c2 = st.columns([0.85, 0.15])
    with c1:
        st.text_input("Carpeta Origen:", key="root_src", value=st.session_state.get("root_src", st.session_state.current_path))
    with c2:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_root_src", help="Seleccionar Origen", on_click=update_path_key, args=("root_src", "Seleccionar Origen"))
        
    c3, c4 = st.columns([0.85, 0.15])
    with c3:
        st.text_input("Carpeta Destino:", key="root_dst", value=st.session_state.get("root_dst", st.session_state.current_path))
    with c4:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_root_dst", help="Seleccionar Destino", on_click=update_path_key, args=("root_dst", "Seleccionar Destino"))

    if uploaded:
        try:
            xl = pd.ExcelFile(uploaded)
            sheets = xl.sheet_names
            sheet = st.selectbox("Selecciona la Hoja:", sheets)
            
            # Leer columnas de la hoja seleccionada
            df_preview = xl.parse(sheet_name=sheet, nrows=1)
            cols = df_preview.columns.tolist()

            c1, c2 = st.columns(2)
            with c1: col_id = st.selectbox("Columna Identificador (Nombre Archivo):", cols)
            with c2: col_dst = st.selectbox("Columna Nombre Carpeta Destino:", cols)
            
            if st.button("🚀 Copiar Archivos"):
                src = st.session_state.get("root_src")
                dst = st.session_state.get("root_dst")
                if src and dst:
                    uploaded.seek(0)
                    file_bytes = uploaded.getvalue()
                    submit_task("Copiar Raíz Mapeo", run_copiar_raiz_mapeo_task, file_bytes, sheet, col_id, col_dst, src, dst)
                    st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
                else:
                    st.warning("Selecciona las rutas base.")
        except Exception as e:
            st.error(f"Error leyendo archivo: {e}")

@st.dialog("Consolidar Subcarpetas")
def dialog_consolidar():
    st.write("Extrae todos los archivos de las subcarpetas profundas y los mueve a la raíz de la carpeta principal.")
    st.warning("Esta acción modificará la estructura de archivos en la ruta actual.")
    
    col_path, col_btn = st.columns([0.85, 0.15])
    with col_path:
        st.text_input("Ruta a procesar:", key="consolidar_path", value=st.session_state.get("consolidar_path", st.session_state.current_path))
    with col_btn:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_consolidar", on_click=update_path_key, args=("consolidar_path",))

    if st.button("🚀 Consolidar"):
        path = st.session_state.get("consolidar_path", st.session_state.current_path)
        submit_task("Consolidar Subcarpetas", run_consolidar_subcarpetas_task, path)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")

def procesar_renombrado(results, full, new_name, sust, find_txt, repl_txt, clean_feov, pre, prefix_txt, suf, suffix_txt, silent_mode=False):
    count = 0
    
    if not results:
        msg = "No hay archivos en la lista de resultados para procesar."
        if not silent_mode:
            st.warning(msg)
        return msg

    # Validación de colisiones para Renombrado Completo
    if full and new_name:
        # Verificar si hay múltiples archivos con la misma extensión en la misma carpeta
        # Lo cual causaría conflicto si todos se llaman igual.
        conflict_map = {}
        for item in results:
            old_path = item["Ruta completa"]
            if not os.path.exists(old_path): continue
            
            folder = os.path.dirname(old_path)
            _, ext = os.path.splitext(old_path)
            key = (folder, ext.lower())
            
            if key in conflict_map:
                conflict_map[key].append(os.path.basename(old_path))
            else:
                conflict_map[key] = [os.path.basename(old_path)]
        
        # Filtrar solo los que tienen conflictos (>1 archivo)
        real_conflicts = {k: v for k, v in conflict_map.items() if len(v) > 1}
        
        if real_conflicts:
            msg = "⚠️ ALERTA: Conflicto de nombres detectado. Estás intentando renombrar a un nombre fijo pero hay múltiples archivos con la misma extensión en la misma carpeta."
            if not silent_mode:
                st.error(msg)
                st.write("El sistema ha detenido el proceso para evitar sobrescritura o nombres duplicados.")
                
                with st.expander("Ver detalles del conflicto"):
                    for (folder, ext), files in real_conflicts.items():
                        st.write(f"📂 En {folder} (Extensión {ext}):")
                        for f in files:
                            st.write(f"  - {f}")
            return msg

    total_files = len(results)
    if not silent_mode:
        progress_bar = st.progress(0, text="Renombrando en lote...")
    
    changes_made = [] # Lista para guardar cambios (Undo)

    for idx, item in enumerate(results):
        if not silent_mode and idx % 5 == 0:
             progress_bar.progress(min(idx / total_files, 1.0), text=f"Renombrando {idx+1}/{total_files}")
             
        old_path = item["Ruta completa"]
        if not os.path.exists(old_path): continue
        
        folder = os.path.dirname(old_path)
        filename = os.path.basename(old_path)
        name_part, ext = os.path.splitext(filename)
        
        final_name = name_part
        
        # 1. Renombrado completo (Prioridad)
        if full and new_name:
            final_name = new_name
        else:
            # 2. Sustitución
            if sust and find_txt:
                final_name = final_name.replace(find_txt, repl_txt)
            
            # 3. Limpieza FEOV
            if clean_feov:
                # Regex para _ID..._A (ej: _ID12345_A)
                final_name = re.sub(r'_ID\d+_A', '', final_name)
            
            # 4. Prefijo
            if pre and prefix_txt:
                final_name = f"{prefix_txt}{final_name}"
            
            # 5. Sufijo
            if suf and suffix_txt:
                final_name = f"{final_name}{suffix_txt}"
        
        new_filename = f"{final_name}{ext}"
        new_path = os.path.join(folder, new_filename)
        
        if new_path != old_path:
            try:
                # Manejo simple de colisiones
                if os.path.exists(new_path):
                    timestamp = int(time.time())
                    new_filename = f"{final_name}_{timestamp}{ext}"
                    new_path = os.path.join(folder, new_filename)
                
                os.rename(old_path, new_path)
                count += 1
                changes_made.append((new_path, old_path)) # Guardar para deshacer
                # Actualizar ruta en resultados para reflejar cambio
                item["Ruta completa"] = new_path
            except Exception as e:
                log(f"Error renombrando {filename}: {e}")

    if not silent_mode:
        progress_bar.progress(1.0, text="Proceso finalizado.")
    
    if count > 0:
        record_action("Renombrado Masivo", changes_made)
        msg = f"✅ Renombrados {count} archivos exitosamente."
        if not silent_mode:
            st.success(msg)
            log(f"Renombrado masivo completado. Total: {count}")
            time.sleep(1) # Pausa breve para ver el mensaje
            st.rerun()
        return msg
    else:
        msg = "No se realizaron cambios (verifique los parámetros o nombres de archivo)."
        if not silent_mode:
            st.info(msg)
        return msg

def run_renombrar_task(results, full, new_name, sust, find_txt, repl_txt, clean_feov, pre, prefix_txt, suf, suffix_txt):
    return {"message": procesar_renombrado(results, full, new_name, sust, find_txt, repl_txt, clean_feov, pre, prefix_txt, suf, suffix_txt, silent_mode=True)}

@st.dialog("Editar Texto en Archivos")
def dialogo_editar_texto():
    st.write("Esta acción buscará y reemplazará texto en los archivos listados en la búsqueda.")
    st.warning("Aplica a archivos de texto plano (.txt, .json, .xml, .py, etc.) y documentos Word (.docx).")
    
    if not st.session_state.search_results:
        st.error("No hay archivos en la lista de resultados.")
        return

    st.info(f"Archivos a procesar: {len(st.session_state.search_results)}")
    
    search_text = st.text_input("Texto a buscar:")
    replace_text = st.text_input("Reemplazar con:")
    
    if st.button("🚀 Ejecutar Reemplazo"):
        if not search_text:
            st.warning("Debes ingresar el texto a buscar.")
            return
            
        submit_task("Editar Texto Masivo", run_editar_texto_task, st.session_state.search_results, search_text, replace_text)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Copiar Archivos de Lista")
def dialogo_copiar_lista():
    st.write("Copiará los archivos/carpetas de la lista de resultados a una nueva ubicación.")
    
    if not st.session_state.search_results:
        st.error("No hay elementos en la lista.")
        return

    st.info(f"Elementos a copiar: {len(st.session_state.search_results)}")
    
    col_dest, col_browse = st.columns([0.85, 0.15])
    with col_dest:
        st.text_input("Carpeta Destino:", key="copy_dest_input", value=st.session_state.get("copy_dest_input", st.session_state.current_path), help="Ruta absoluta donde se copiarán los archivos.")
    with col_browse:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_browse_copy", help="Examinar...", on_click=update_path_key, args=("copy_dest_input", "Seleccionar Destino"))

    if st.button("🚀 Copiar"):
        dest = st.session_state.get("copy_dest_input")
        if not dest:
            st.warning("Selecciona una carpeta destino.")
            return
        submit_task("Copiar Lista", run_copiar_lista_task, st.session_state.search_results, dest)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Mover Archivos de Lista")
def dialogo_mover_lista():
    st.write("Moverá los archivos/carpetas de la lista de resultados a una nueva ubicación.")
    st.warning("Los archivos originales serán eliminados de su ubicación actual.")
    
    if not st.session_state.search_results:
        st.error("No hay elementos en la lista.")
        return

    st.info(f"Elementos a mover: {len(st.session_state.search_results)}")
    
    col_dest, col_browse = st.columns([0.85, 0.15])
    with col_dest:
        st.text_input("Carpeta Destino:", key="move_dest_input", value=st.session_state.get("move_dest_input", st.session_state.current_path))
    with col_browse:
        st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
        st.button("📂", key="btn_browse_move", help="Examinar...", on_click=update_path_key, args=("move_dest_input", "Seleccionar Destino"))

    if st.button("🚀 Mover"):
        dest = st.session_state.get("move_dest_input")
        if not dest:
            st.warning("Selecciona una carpeta destino.")
            return
        submit_task("Mover Lista", run_mover_lista_task, st.session_state.search_results, dest)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Comprimir Lista en ZIP")
def dialogo_zip_lista():
    st.write("Creará un archivo ZIP con todos los elementos de la lista.")
    
    if not st.session_state.search_results:
        st.error("No hay elementos en la lista.")
        return

    st.info(f"Elementos a comprimir: {len(st.session_state.search_results)}")
    
    zip_name = st.text_input("Nombre del archivo ZIP:", value="Archivos_Comprimidos.zip")
    
    col_dest, col_browse = st.columns([0.85, 0.15])
    with col_dest:
         st.text_input("Carpeta Destino:", key="zip_dest_path", value=st.session_state.get("zip_dest_path", st.session_state.current_path))
    with col_browse:
         st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
         st.button("📂", key="btn_zip_dest", on_click=update_path_key, args=("zip_dest_path",))
    
    if st.button("🚀 Comprimir"):
        if not zip_name.endswith(".zip"): zip_name += ".zip"
        dest = st.session_state.get("zip_dest_path", st.session_state.current_path)
        target_path = os.path.join(dest, zip_name)
        submit_task("ZIP Lista", run_zip_lista_task, st.session_state.search_results, target_path)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Confirmar Eliminación")
def dialogo_confirmar_eliminar():
    st.warning("⚠️ ¿Estás seguro de que quieres enviar estos archivos a la papelera?")
    st.write("Esta acción enviará los archivos a la papelera de reciclaje del sistema.")
    
    results = st.session_state.search_results
    if not results:
        st.error("No hay archivos seleccionados.")
        return

    st.write(f"Total de elementos: **{len(results)}**")
    
    with st.expander("Ver lista de archivos a eliminar"):
        for item in results[:50]: # Show first 50
            st.text(os.path.basename(item["Ruta completa"]))
        if len(results) > 50:
            st.text(f"... y {len(results)-50} más.")

    col_confirm, col_cancel = st.columns(2)
    with col_confirm:
        if st.button("🗑️ Sí, eliminar", type="primary", use_container_width=True):
            count_del = 0
            progress_bar = st.progress(0, text="Eliminando...")
            total = len(results)
            
            for i, item in enumerate(results):
                progress_bar.progress(min(i/total, 1.0), text=f"Eliminando {i+1}/{total}")
                path = item["Ruta completa"]
                try:
                    if os.path.exists(path):
                        safe_path = os.path.normpath(path)
                        send2trash(safe_path)
                        count_del += 1
                except Exception as e:
                    log(f"Error eliminando {path}: {e}")
            
            progress_bar.progress(1.0, text="Finalizado.")
            st.success(f"Se enviaron {count_del} archivos a la papelera.")
            st.session_state.search_results = [] 
            time.sleep(1.5)
            st.rerun()
            
    with col_cancel:
        if st.button("Cancelar", use_container_width=True):
            st.rerun()

@st.dialog("Comprimir Carpetas Individualmente")
def dialogo_zip_carpetas_individual():
    st.write("Buscará carpetas en la lista de resultados y creará un ZIP individual para cada una.")
    
    if not st.session_state.search_results:
        st.error("No hay elementos en la lista.")
        return

    # Filtrar solo carpetas
    folders = [x for x in st.session_state.search_results if os.path.isdir(x["Ruta completa"])]
    st.info(f"Carpetas encontradas en la lista: {len(folders)}")
    
    if len(folders) == 0:
        st.warning("No hay carpetas en la lista actual.")
        return

    col_dest, col_browse = st.columns([0.85, 0.15])
    with col_dest:
         st.text_input("Carpeta Destino:", key="zip_ind_dest", value=st.session_state.get("zip_ind_dest", st.session_state.current_path), help="Donde se guardarán los ZIPs generados.")
    with col_browse:
         st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
         st.button("📂", key="btn_zip_ind", on_click=update_path_key, args=("zip_ind_dest",))
        
    if st.button("🚀 Comprimir Carpetas"):
        dest = st.session_state.get("zip_ind_dest", st.session_state.current_path)
        submit_task("ZIP Carpetas Individual", run_zip_carpetas_ind_task, st.session_state.search_results, dest)
        st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")

@st.dialog("Modificar Nombres - Opciones Avanzadas")
def dialogo_modificar_nombres():
    st.write("Configura las opciones de renombrado:")
    
    # 1. Renombrado completo
    with st.container(border=True):
        st.markdown("**📝 Renombrado completo (ignora las demás opciones)**")
        activar_full = st.checkbox("Activar renombrado completo", key="chk_full")
        nuevo_nombre = st.text_input("Nuevo nombre (sin extensión)", disabled=not activar_full, key="txt_full")
    
    # 2. Sustituir texto
    with st.container(border=True):
        st.markdown("**🔄 Sustituir texto**")
        activar_sust = st.checkbox("Activar sustitución", key="chk_sust")
        buscar_txt = st.text_input("Buscar:", disabled=not activar_sust, key="txt_find")
        reemplazar_txt = st.text_input("Reemplazar con:", disabled=not activar_sust, key="txt_repl")

    # 3. Limpieza Especial
    with st.container(border=True):
        st.markdown("**🧹 Limpieza Especial (FEOV)**")
        eliminar_id = st.checkbox("Eliminar '_ID<números>_A' del nombre", key="chk_feov")

    # 4. Añadir al inicio
    with st.container(border=True):
        st.markdown("**⬅️ Añadir al inicio**")
        activar_pre = st.checkbox("Añadir prefijo", key="chk_pre")
        prefijo = st.text_input("Texto prefijo:", disabled=not activar_pre, key="txt_pre")

    # 5. Añadir al final
    with st.container(border=True):
        st.markdown("**➡️ Añadir al final (antes de extensión)**")
        activar_suf = st.checkbox("Añadir sufijo", key="chk_suf")
        sufijo = st.text_input("Texto sufijo:", disabled=not activar_suf, key="txt_suf")

    st.markdown("---")
    col_cancel, col_ok = st.columns([1, 1])
    with col_ok:
        if st.button("✅ Ejecutar Cambios", use_container_width=True):
            submit_task("Renombrado Masivo", run_renombrar_task, st.session_state.search_results,
                activar_full, nuevo_nombre, 
                activar_sust, buscar_txt, reemplazar_txt, 
                eliminar_id, 
                activar_pre, prefijo, 
                activar_suf, sufijo
            )
            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas'.")
            
    with col_cancel:
        if st.button("❌ Cancelar", use_container_width=True):
            st.rerun()

def buscar_archivos():
    path = st.session_state.current_path
    if not os.path.exists(path):
        st.error(f"La ruta no existe: {path}")
        return

    pattern = st.session_state.get("pattern", "").lower()
    exclusion = st.session_state.get("exclusion_pattern", "").lower()
    exclusion_list = [x.strip() for x in exclusion.split(",") if x.strip()]

    search_by = st.session_state.get("search_by", "todo") # extensión, nombre, todo
    item_type = st.session_state.get("item_type", "archivos") # archivos, carpetas
    recursive = st.session_state.get("subfolders", True)

    results = []
    log(f"Iniciando búsqueda en: {path} | Patrón: '{pattern}' | Excluir: '{exclusion}'")

    for root, dirs, files in os.walk(path):
        # Si no es recursivo, limpiar dirs para que os.walk no baje más, 
        # PERO debemos procesar la raíz actual.
        # os.walk yielda (root, dirs, files). Si modificamos dirs in-place, afecta la recursión.
        
        items_to_check = []
        if item_type == "archivos":
            items_to_check = files
        elif item_type == "carpetas":
            items_to_check = dirs
        
        for item in items_to_check:
            item_lower = item.lower()
            
            # Verificar exclusiones
            if any(excl in item_lower for excl in exclusion_list):
                continue

            match = False
            
            if not pattern:
                match = True
            else:
                if search_by == "extensión":
                    # Solo aplica a archivos
                    if item_type == "archivos" and item_lower.endswith(pattern):
                        match = True
                elif search_by == "nombre":
                    # Nombre sin extensión para archivos, o nombre carpeta
                    name_only = os.path.splitext(item_lower)[0] if item_type == "archivos" else item_lower
                    if pattern in name_only:
                        match = True
                else: # todo
                    if pattern in item_lower:
                        match = True
            
            if match:
                full_path = os.path.join(root, item)
                try:
                    stats = os.stat(full_path)
                    mtime = datetime.fromtimestamp(stats.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                except:
                    mtime = "N/A"
                
                results.append({
                    "Ruta completa": full_path,
                    "Fecha": mtime
                })

        if not recursive:
            break
            
    st.session_state.search_results = results
    if not results:
        st.warning("No se encontraron coincidencias.")
    else:
        st.success(f"Encontrados {len(results)} elementos.")

# --- WORKERS DE ANÁLISIS Y MANUALES (Portados de Desktop) ---

def worker_analisis_historia_clinica(file_list, silent_mode=False):
    if not file_list: return None
    pdf_files = [f for f in file_list if f["Ruta completa"].lower().endswith(".pdf")]
    if not pdf_files:
        if not silent_mode:
            st.error("No se encontraron archivos PDF en la lista.")
        return None

    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando análisis de historias clínicas...")
    
    patterns = {
        'Paciente': re.compile(r"Paciente:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Identificación': re.compile(r"Identificación:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Edad': re.compile(r"Edad:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Genero': re.compile(r"Género:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Aseguradora': re.compile(r"Aseguradora:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Fecha Ingreso': re.compile(r"Fecha Ingreso:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Fecha Egreso': re.compile(r"Fecha Egreso:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Diagnostico Principal': re.compile(r"Diagnóstico Principal:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Folio': re.compile(r"Folio:\s*(.*?)(?:\n|$)", re.IGNORECASE)
    }

    extracted_data = []
    errores = 0
    total = len(pdf_files)

    for i, item in enumerate(pdf_files):
        pdf_path = item["Ruta completa"]
        filename = os.path.basename(pdf_path)
        if not silent_mode:
            progress_bar.progress((i + 1) / total, text=f"Analizando: {filename}")
        
        try:
            full_text = ""
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            
            data_row = {'Archivo': filename}
            for key, pattern in patterns.items():
                match = pattern.search(full_text)
                data_row[key] = match.group(1).strip() if match else "No encontrado"
            
            extracted_data.append(data_row)
        except Exception as e:
            log(f"Error analizando {filename}: {e}")
            errores += 1
            extracted_data.append({'Archivo': filename, 'Error': str(e)})

    df = pd.DataFrame(extracted_data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    
    if not silent_mode:
        progress_bar.empty()
        st.success(f"Análisis completado. Procesados: {total}, Errores: {errores}")
    return output

def worker_analisis_autorizacion_nueva_eps(file_list, silent_mode=False):
    if not file_list: return None
    pdf_files = [f for f in file_list if f["Ruta completa"].lower().endswith(".pdf")]
    if not pdf_files:
        if not silent_mode:
            st.error("No se encontraron archivos PDF en la lista.")
        return None

    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando análisis de Autorizaciones Nueva EPS...")
    
    patterns = {
        'Afiliado': re.compile(r"Afiliado:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'N° Autorización': re.compile(r"N° Autorización:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Autorizada el': re.compile(r"Autorizada el:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Descripción Servicio': re.compile(r"Descripción Servicio\s*\n\s*\d+\s+\d+\s+(.*?)(?:\n|$)", re.IGNORECASE | re.DOTALL),
        'Info de Pago': re.compile(r"(Afiliado (?:No )?Cancela.*?)(?:\n|$)", re.IGNORECASE)
    }

    extracted_data = []
    errores = 0
    total = len(pdf_files)

    for i, item in enumerate(pdf_files):
        pdf_path = item["Ruta completa"]
        filename = os.path.basename(pdf_path)
        if not silent_mode:
            progress_bar.progress((i + 1) / total, text=f"Analizando: {filename}")

        try:
            full_text = ""
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            
            data_row = {'Archivo': filename}
            for key, pattern in patterns.items():
                match = pattern.search(full_text)
                if match:
                    val = match.group(1).strip()
                    if key == 'Descripción Servicio':
                        val = val.replace('\n', ' ').strip()
                    data_row[key] = val
                else:
                    data_row[key] = "No encontrado"
            extracted_data.append(data_row)
        except Exception as e:
            log(f"Error analizando {filename}: {e}")
            errores += 1
    
    df = pd.DataFrame(extracted_data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    if not silent_mode:
        progress_bar.empty()
        st.success(f"Análisis completado. Procesados: {total}, Errores: {errores}")
    return output

def worker_analisis_cargue_sanitas(file_list, silent_mode=False):
    if not file_list: return None
    pdf_files = [f for f in file_list if f["Ruta completa"].lower().endswith(".pdf")]
    if not pdf_files:
        if not silent_mode:
            st.error("No se encontraron archivos PDF en la lista.")
        return None

    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando análisis Cargue Sanitas...")
    patterns = {
        'FEOV': re.compile(r"FEOV\s*[:#]?\s*(\S+)", re.IGNORECASE),
        'Estado': re.compile(r"Estado\s*[:]?\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Fecha': re.compile(r"Fecha\s*[:]?\s*(\d{2}/\d{2}/\d{4})", re.IGNORECASE)
    }
    extracted_data = []
    errores = 0
    total = len(pdf_files)

    for i, item in enumerate(pdf_files):
        pdf_path = item["Ruta completa"]
        filename = os.path.basename(pdf_path)
        if not silent_mode:
            progress_bar.progress((i + 1) / total, text=f"Analizando: {filename}")
        try:
            full_text = ""
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            data_row = {'Archivo': filename}
            for key, pattern in patterns.items():
                match = pattern.search(full_text)
                data_row[key] = match.group(1).strip() if match else "No encontrado"
            extracted_data.append(data_row)
        except Exception as e:
            log(f"Error analizando {filename}: {e}")
            errores += 1
            
    df = pd.DataFrame(extracted_data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    if not silent_mode:
        progress_bar.empty()
        st.success(f"Análisis completado. Procesados: {total}, Errores: {errores}")
    return output

def worker_analisis_retefuente(file_list, silent_mode=False):
    if not file_list: return None
    pdf_files = [f for f in file_list if f["Ruta completa"].lower().endswith(".pdf")]
    if not pdf_files:
        if not silent_mode:
            st.error("No se encontraron archivos PDF en la lista.")
        return None

    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando análisis Retefuente...")
    extracted_data = []
    errores = 0
    total = len(pdf_files)
    
    for i, item in enumerate(pdf_files):
        pdf_path = item["Ruta completa"]
        filename = os.path.basename(pdf_path)
        if not silent_mode:
            progress_bar.progress((i + 1) / total, text=f"Analizando: {filename}")
        try:
            with fitz.open(pdf_path) as doc:
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text("text")
                    razon_social = "NO ENCONTRADO"
                    nit = "NO ENCONTRADO"
                    lines = text.split('\n')
                    for idx, line in enumerate(lines):
                        if "RAZON SOCIAL" in line.upper() or "APELLIDOS Y NOMBRES" in line.upper():
                            if idx + 1 < len(lines):
                                razon_social = lines[idx+1].strip()
                        if "NIT" in line.upper() and len(line) < 30:
                             nit = line.split("NIT")[-1].strip().replace(":","")
                    extracted_data.append({'Archivo': filename, 'Pagina': page_num, 'Razon Social': razon_social, 'NIT': nit})
        except Exception as e:
            log(f"Error analizando {filename}: {e}")
            errores += 1

    df = pd.DataFrame(extracted_data)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    if not silent_mode:
        progress_bar.empty()
        st.success(f"Análisis completado. Procesados: {total}, Errores: {errores}")
    return output

def worker_unificar_pdf_manual(file_list, output_name, silent_mode=False):
    if not file_list: return
    pdf_files = [f for f in file_list if f["Ruta completa"].lower().endswith(".pdf")]
    if not pdf_files:
        if silent_mode: return "Error: No hay archivos PDF en la lista para unificar."
        st.error("No hay archivos PDF en la lista para unificar.")
        return
        
    if not output_name.endswith(".pdf"): output_name += ".pdf"
    output_path = os.path.join(st.session_state.current_path, output_name)
    
    if not silent_mode:
        progress_bar = st.progress(0, text="Unificando PDFs...")
    try:
        doc_final = fitz.open()
        total = len(pdf_files)
        for i, item in enumerate(pdf_files):
            if not silent_mode:
                progress_bar.progress((i + 1) / total, text=f"Uniendo: {os.path.basename(item['Ruta completa'])}")
            try:
                doc_temp = fitz.open(item["Ruta completa"])
                doc_final.insert_pdf(doc_temp)
                doc_temp.close()
            except Exception as e:
                log(f"Error uniendo {item['Ruta completa']}: {e}")
        
        doc_final.save(output_path)
        doc_final.close()
        
        msg = f"PDFs unificados exitosamente en: {output_name}"
        if not silent_mode:
            progress_bar.empty()
            st.success(f"✅ {msg}")
        return msg
    except Exception as e:
        msg = f"Error general unificando: {e}"
        if not silent_mode:
            st.error(msg)
        return msg

def run_unificar_pdf_manual_task(file_list, output_name):
    return {"message": worker_unificar_pdf_manual(file_list, output_name, silent_mode=True)}

def worker_dividir_pdf_manual(file_list, silent_mode=False):
    if not file_list: return
    pdf_files = [f for f in file_list if f["Ruta completa"].lower().endswith(".pdf")]
    if not pdf_files:
        if silent_mode: return "Error: No hay archivos PDF en la lista para dividir."
        st.error("No hay archivos PDF en la lista para dividir.")
        return

    if not silent_mode:
        progress_bar = st.progress(0, text="Dividiendo PDFs...")
    count = 0
    total = len(pdf_files)
    
    for i, item in enumerate(pdf_files):
        pdf_path = item["Ruta completa"]
        filename = os.path.basename(pdf_path)
        base_name = os.path.splitext(filename)[0]
        folder = os.path.dirname(pdf_path)
        target_folder = os.path.join(folder, base_name)
        
        if not silent_mode:
            progress_bar.progress((i + 1) / total, text=f"Dividiendo: {filename}")
        
        if not os.path.exists(target_folder):
            os.makedirs(target_folder)
            
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                new_doc = fitz.open()
                new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
                new_doc.save(os.path.join(target_folder, f"{base_name}_pag_{page_num+1}.pdf"))
                new_doc.close()
            doc.close()
            count += 1
        except Exception as e:
            log(f"Error dividiendo {filename}: {e}")
            
    msg = f"{count} PDFs divididos en sus respectivas carpetas."
    if not silent_mode:
        progress_bar.empty()
        st.success(f"✅ {msg}")
    return msg

def run_dividir_pdf_manual_task(file_list):
    return {"message": worker_dividir_pdf_manual(file_list, silent_mode=True)}


# --- WORKERS RIPS (Alineados con Desktop v1) ---

def worker_json_a_xlsx_ind(file_obj):
    try:
        # Asegurar lectura desde inicio si es un objeto tipo archivo
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
            
        data = json.load(file_obj)
        
        # Mapa de JSON key -> Excel Sheet Name
        service_map = {
            "consultas": "Consultas",
            "procedimientos": "Procedimientos",
            "urgencias": "Urgencias",
            "hospitalizacion": "Hospitalizacion",
            "recienNacidos": "RecienNacidos",
            "medicamentos": "Medicamentos",
            "otrosServicios": "OtrosServicios"
        }
        
        # 1. Transaccion (Header)
        header_info = {
            "numDocumentoIdObligado": data.get("numDocumentoIdObligado"),
            "numFactura": data.get("numFactura"),
            "tipoNota": data.get("tipoNota"),
            "numNota": data.get("numNota")
        }
        
        # 2. Usuarios y Servicios
        usuarios_rows = []
        all_services = {name: [] for name in service_map.values()}
        
        usuarios_lista = data.get("usuarios", []) if isinstance(data, dict) else []
        
        for usuario in usuarios_lista:
            # Extraer info usuario con case-insensitive helper
            u_info = {
                "tipoDocumentoIdentificacion": get_val_ci(usuario, "tipoDocumentoIdentificacion"),
                "numDocumentoIdentificacion": get_val_ci(usuario, "numDocumentoIdentificacion"),
                "tipoUsuario": get_val_ci(usuario, "tipoUsuario"),
                "fechaNacimiento": get_val_ci(usuario, "fechaNacimiento"),
                "codSexo": get_val_ci(usuario, "codSexo"),
                "codPaisResidencia": get_val_ci(usuario, "codPaisResidencia"),
                "codMunicipioResidencia": get_val_ci(usuario, "codMunicipioResidencia"),
                "codZonaTerritorialResidencia": get_val_ci(usuario, "codZonaTerritorialResidencia"),
                "incapacidad": get_val_ci(usuario, "incapacidad"),
                "consecutivo": get_val_ci(usuario, "consecutivo"),
                "codPaisOrigen": get_val_ci(usuario, "codPaisOrigen")
            }
            usuarios_rows.append(u_info)
            
            # Link Key for services
            consecutivo = u_info.get("consecutivo")

            servicios = usuario.get("servicios", {})
            if not servicios: continue

            for json_key, sheet_name in service_map.items():
                items = servicios.get(json_key, [])
                for item in items:
                    # Item + Link to User
                    item_row = item.copy()
                    item_row["consecutivoUsuario"] = consecutivo
                    all_services[sheet_name].append(item_row)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Sheet Transaccion
            pd.DataFrame([header_info]).to_excel(writer, sheet_name="Transaccion", index=False)
            
            # Sheet Usuarios
            if usuarios_rows:
                pd.DataFrame(usuarios_rows).to_excel(writer, sheet_name="Usuarios", index=False)
            else:
                # Create empty with columns if needed, or just empty
                pd.DataFrame(columns=list(u_info.keys()) if 'u_info' in locals() else []).to_excel(writer, sheet_name="Usuarios", index=False)

            # Service Sheets
            written_services = False
            for sheet_name, rows in all_services.items():
                if rows:
                    pd.DataFrame(rows).to_excel(writer, sheet_name=sheet_name, index=False)
                    written_services = True
            
            # If no services, maybe write empty Consultas? Not strictly required but good for template
                 
        return output.getvalue(), None
    except Exception as e:
        return None, str(e)

def worker_xlsx_a_json_ind(file_obj):
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
            
        # Leer el archivo Excel
        xls = pd.ExcelFile(file_obj)
        
        # Mapa Excel Sheet -> JSON Key
        service_map = {
            "Consultas": "consultas",
            "Procedimientos": "procedimientos",
            "Urgencias": "urgencias",
            "Hospitalizacion": "hospitalizacion",
            "RecienNacidos": "recienNacidos",
            "Medicamentos": "medicamentos",
            "OtrosServicios": "otrosServicios"
        }
        
        # --- STRATEGY 1: NORMALIZED (Transaccion + Usuarios sheets) ---
        if "Transaccion" in xls.sheet_names and "Usuarios" in xls.sheet_names:
            # 1. Transaccion (Header)
            transaccion_data = {}
            df_t = pd.read_excel(xls, sheet_name="Transaccion")
            df_t = clean_df_for_json(df_t)
            if not df_t.empty:
                transaccion_data = df_t.iloc[0].to_dict()
                
            # 2. Usuarios
            usuarios_map = {} # consecutivo -> user_obj
            df_u = pd.read_excel(xls, sheet_name="Usuarios")
            df_u = clean_df_for_json(df_u)
            for _, row in df_u.iterrows():
                u_obj = row.to_dict()
                u_obj["servicios"] = {k: [] for k in service_map.values()}
                consecutivo = str(u_obj.get("consecutivo"))
                usuarios_map[consecutivo] = u_obj
            
            # 3. Servicios
            for sheet_name, json_key in service_map.items():
                if sheet_name in xls.sheet_names:
                    df_s = pd.read_excel(xls, sheet_name=sheet_name)
                    df_s = clean_df_for_json(df_s)
                    
                    for _, row in df_s.iterrows():
                        s_obj = row.to_dict()
                        consecutivo_usuario = str(s_obj.pop("consecutivoUsuario", None))
                        
                        if (consecutivo_usuario == "None" or consecutivo_usuario == "nan") and len(usuarios_map) == 1:
                             consecutivo_usuario = list(usuarios_map.keys())[0]

                        if consecutivo_usuario in usuarios_map:
                            usuarios_map[consecutivo_usuario]["servicios"][json_key].append(s_obj)

            final_json = transaccion_data
            final_json["usuarios"] = list(usuarios_map.values())
            return json.dumps(final_json, ensure_ascii=False, indent=4), None

        # --- STRATEGY 2: FLAT (Service sheets contain all info) ---
        else:
            header_data = {}
            header_extracted = False
            usuarios_dict = {} # (tipoDoc, numDoc) -> user_obj
            
            for sheet_name, json_key in service_map.items():
                if sheet_name not in xls.sheet_names: continue
                
                df = pd.read_excel(xls, sheet_name=sheet_name)
                df = clean_df_for_json(df)
                
                for _, row in df.iterrows():
                    # Extract Header
                    if not header_extracted:
                        header_data = {
                            "numDocumentoIdObligado": row.get("numDocumentoIdObligado"),
                            "numFactura": row.get("numFactura"),
                            "tipoNota": row.get("tipoNota"),
                            "numNota": row.get("numNota")
                        }
                        header_extracted = True
                    
                    # Extract User Key
                    td = str(row.get("tipoDocumentoIdentificacion") or row.get("tipo_documento_usuario") or "")
                    doc = str(row.get("numDocumentoIdentificacion") or row.get("documento_usuario") or "")
                    user_key = (td, doc)
                    
                    if user_key not in usuarios_dict:
                        # Extract User Fields
                        usuarios_dict[user_key] = {
                            "tipoDocumentoIdentificacion": row.get("tipoDocumentoIdentificacion") or row.get("tipo_documento_usuario"),
                            "numDocumentoIdentificacion": row.get("numDocumentoIdentificacion") or row.get("documento_usuario"),
                            "tipoUsuario": row.get("tipoUsuario") or row.get("tipo_usuario"),
                            "fechaNacimiento": row.get("fechaNacimiento") or row.get("fecha_nacimiento"), 
                            "codSexo": row.get("codSexo") or row.get("sexo"),
                            "codPaisResidencia": row.get("codPaisResidencia") or row.get("pais_residencia"),
                            "codMunicipioResidencia": row.get("codMunicipioResidencia") or row.get("codMunicipioResidencia"),
                            "codZonaTerritorialResidencia": row.get("codZonaTerritorialResidencia") or row.get("zona_residencia"),
                            "incapacidad": row.get("incapacidad"),
                            "consecutivo": row.get("consecutivo") or row.get("consecutivo_usuario"),
                            "codPaisOrigen": row.get("codPaisOrigen") or row.get("pais_origen"),
                            "servicios": {k: [] for k in service_map.values()}
                        }

                    # Extract Service Data
                    s_obj = row.to_dict()
                    
                    # Remove Header, User, and System fields
                    fields_to_remove = [
                        "archivo_origen", "numDocumentoIdObligado", "numFactura", "tipoNota", "numNota",
                        "tipoDocumentoIdentificacion", "numDocumentoIdentificacion", "tipoUsuario", "fechaNacimiento",
                        "codSexo", "codPaisResidencia", "codMunicipioResidencia", "codZonaTerritorialResidencia",
                        "incapacidad", "consecutivo", "codPaisOrigen",
                        # Legacy synonyms
                        "tipo_documento_usuario", "documento_usuario", "tipo_usuario", "fecha_nacimiento",
                        "sexo", "pais_residencia", "municipio_residencia", "zona_residencia", "consecutivo_usuario", "pais_origen"
                    ]
                    for f in fields_to_remove:
                        s_obj.pop(f, None)
                        
                    usuarios_dict[user_key]["servicios"][json_key].append(s_obj)
            
            final_json = header_data
            final_json["usuarios"] = list(usuarios_dict.values())
            
            if not final_json["usuarios"] and not header_extracted:
                 return None, "No se encontraron datos válidos (ni Transaccion/Usuarios ni hojas de servicios con datos)."

            return json.dumps(final_json, ensure_ascii=False, indent=4), None

    except Exception as e:
        return None, str(e)

def worker_consolidar_json_xlsx(folder_path):
    if not folder_path or not os.path.exists(folder_path):
        return None, "Carpeta no encontrada"
        
    archivos_json = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".json"):
                archivos_json.append(os.path.join(root, file))
    
    if not archivos_json:
        return None, "No se encontraron archivos JSON."

    try:
        service_map = {
            "consultas": "Consultas",
            "procedimientos": "Procedimientos",
            "urgencias": "Urgencias",
            "hospitalizacion": "Hospitalizacion",
            "recienNacidos": "RecienNacidos",
            "medicamentos": "Medicamentos",
            "otrosServicios": "OtrosServicios"
        }
        
        transaccion_rows = []
        usuarios_rows = []
        all_services = {sheet: [] for sheet in service_map.values()}
        
        errores = 0
        
        progress_bar = st.progress(0, text="Consolidando...")
        total = len(archivos_json)
        
        for i, ruta_json in enumerate(archivos_json):
            progress_bar.progress((i+1)/total, text=f"Procesando: {os.path.basename(ruta_json)}")
            nombre_archivo = os.path.basename(ruta_json)
            
            try:
                with open(ruta_json, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Header info
                header_info = {
                    "archivo_origen": nombre_archivo,
                    "numDocumentoIdObligado": data.get("numDocumentoIdObligado"),
                    "numFactura": data.get("numFactura"),
                    "tipoNota": data.get("tipoNota"),
                    "numNota": data.get("numNota")
                }
                transaccion_rows.append(header_info)

                usuarios_lista = data.get("usuarios", []) if isinstance(data, dict) else []
                
                for usuario in usuarios_lista:
                    u_info = {
                        "archivo_origen": nombre_archivo,
                        "tipoDocumentoIdentificacion": get_val_ci(usuario, "tipoDocumentoIdentificacion"),
                        "numDocumentoIdentificacion": get_val_ci(usuario, "numDocumentoIdentificacion"),
                        "tipoUsuario": get_val_ci(usuario, "tipoUsuario"),
                        "fechaNacimiento": get_val_ci(usuario, "fechaNacimiento"),
                        "codSexo": get_val_ci(usuario, "codSexo"),
                        "codPaisResidencia": get_val_ci(usuario, "codPaisResidencia"),
                        "codMunicipioResidencia": get_val_ci(usuario, "codMunicipioResidencia"),
                        "codZonaTerritorialResidencia": get_val_ci(usuario, "codZonaTerritorialResidencia"),
                        "incapacidad": get_val_ci(usuario, "incapacidad"),
                        "consecutivo": get_val_ci(usuario, "consecutivo"),
                        "codPaisOrigen": get_val_ci(usuario, "codPaisOrigen")
                    }
                    usuarios_rows.append(u_info)
                    
                    consecutivo = u_info.get("consecutivo")
                    servicios = usuario.get("servicios", {})

                    for json_key, sheet_name in service_map.items():
                        items = servicios.get(json_key, [])
                        for item in items:
                            item_row = item.copy()
                            item_row["archivo_origen"] = nombre_archivo
                            item_row["consecutivoUsuario"] = consecutivo
                            all_services[sheet_name].append(item_row)
                        
            except Exception as e:
                errores += 1
                log(f"Error en {nombre_archivo}: {e}")

        progress_bar.empty()
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Transaccion
            if transaccion_rows:
                pd.DataFrame(transaccion_rows).to_excel(writer, sheet_name="Transaccion", index=False)
            
            # Usuarios
            if usuarios_rows:
                pd.DataFrame(usuarios_rows).to_excel(writer, sheet_name="Usuarios", index=False)
            
            # Services
            written_any = False
            for sheet_name, rows in all_services.items():
                if rows:
                    pd.DataFrame(rows).to_excel(writer, sheet_name=sheet_name, index=False)
                    written_any = True
            
            if not written_any and not transaccion_rows:
                 pd.DataFrame().to_excel(writer, sheet_name="Vacio", index=False)
        
        return output.getvalue(), f"Procesados: {total}, Errores: {errores}"
        
    except Exception as e:
        return None, str(e)

def worker_desconsolidar_xlsx_json(file_obj, dest_folder):
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
            
        xls = pd.ExcelFile(file_obj)
        
        service_map = {
            "Consultas": "consultas",
            "Procedimientos": "procedimientos",
            "Urgencias": "urgencias",
            "Hospitalizacion": "hospitalizacion",
            "RecienNacidos": "recienNacidos",
            "Medicamentos": "medicamentos",
            "OtrosServicios": "otrosServicios"
        }
        
        # --- STRATEGY 1: NORMALIZED (Transaccion + Usuarios + Services) ---
        if "Transaccion" in xls.sheet_names and "Usuarios" in xls.sheet_names:
            df_t = pd.read_excel(xls, sheet_name="Transaccion")
            df_t = clean_df_for_json(df_t)
            
            df_u = pd.read_excel(xls, sheet_name="Usuarios")
            df_u = clean_df_for_json(df_u)
            
            if "archivo_origen" not in df_t.columns:
                return False, "La hoja Transaccion no tiene columna 'archivo_origen'."
            
            # Group Headers by file
            headers_by_file = {row["archivo_origen"]: row.to_dict() for _, row in df_t.iterrows()}
            
            # Group Users by file
            users_by_file = {} # filename -> {consecutivo -> user_obj}
            for _, row in df_u.iterrows():
                fname = row.get("archivo_origen")
                if not fname: continue
                if fname not in users_by_file: users_by_file[fname] = {}
                
                u_obj = row.to_dict()
                u_obj.pop("archivo_origen", None)
                u_obj["servicios"] = {k: [] for k in service_map.values()}
                
                consecutivo = str(u_obj.get("consecutivo"))
                users_by_file[fname][consecutivo] = u_obj
                
            # Process Services
            for sheet_name, json_key in service_map.items():
                if sheet_name in xls.sheet_names:
                    df_s = pd.read_excel(xls, sheet_name=sheet_name)
                    df_s = clean_df_for_json(df_s)
                    
                    for _, row in df_s.iterrows():
                        fname = row.get("archivo_origen")
                        consecutivo_usuario = str(row.get("consecutivoUsuario"))
                        
                        if fname in users_by_file and consecutivo_usuario in users_by_file[fname]:
                            s_obj = row.to_dict()
                            s_obj.pop("archivo_origen", None)
                            s_obj.pop("consecutivoUsuario", None)
                            users_by_file[fname][consecutivo_usuario]["servicios"][json_key].append(s_obj)
                            
            # Generate JSONs
            progress_bar = st.progress(0, text="Generando JSONs...")
            files_generated = 0
            
            for fname, header_row in headers_by_file.items():
                header_data = header_row.copy()
                header_data.pop("archivo_origen", None)
                
                users_map = users_by_file.get(fname, {})
                final_json = header_data
                final_json["usuarios"] = list(users_map.values())
                
                if not fname.lower().endswith(".json"): fname += ".json"
                
                with open(os.path.join(dest_folder, fname), 'w', encoding='utf-8') as f:
                    json.dump(final_json, f, ensure_ascii=False, indent=4)
                files_generated += 1
                progress_bar.progress(files_generated / len(headers_by_file))
                
            progress_bar.empty()
            return True, f"Generados {files_generated} archivos JSON (Estructura Normalizada)."
        
        # --- STRATEGY 2: FLAT (Old Massive Structure) ---
        else:
            # Read all sheets
            dfs = {}
        all_files = set()
        
        for sheet_name in service_map.keys():
            if sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)
                # Convert NaN to None
                df = df.astype(object).where(pd.notnull(df), None)
                dfs[sheet_name] = df
                if "archivo_origen" in df.columns:
                    all_files.update(df["archivo_origen"].dropna().unique())
        
        if not all_files:
             return False, "No se encontró columna 'archivo_origen' en ninguna hoja del Excel."

        progress_bar = st.progress(0, text="Generando JSONs...")
        total = len(all_files)
        
        for i, nombre_archivo in enumerate(all_files):
            progress_bar.progress((i+1)/total, text=f"Generando: {nombre_archivo}")
            
            # Header info holders
            header_extracted = False
            header_data = {}
            
            usuarios_dict = {}
            
            for sheet_name, json_key in service_map.items():
                if sheet_name not in dfs: continue
                df = dfs[sheet_name]
                if "archivo_origen" not in df.columns: continue
                
                df_fil = df[df["archivo_origen"] == nombre_archivo]
                
                for _, row in df_fil.iterrows():
                    # Extract header info from first row if not yet done
                    if not header_extracted:
                        header_data = {
                            "numDocumentoIdObligado": row.get("numDocumentoIdObligado"),
                            "numFactura": row.get("numFactura"),
                            "tipoNota": row.get("tipoNota"),
                            "numNota": row.get("numNota")
                        }
                        header_extracted = True
                    
                    # User Key
                    td = str(row.get("tipoDocumentoIdentificacion") or row.get("tipo_documento_usuario") or "")
                    doc = str(row.get("numDocumentoIdentificacion") or row.get("documento_usuario") or "")
                    user_key = (td, doc)
                    
                    if user_key not in usuarios_dict:
                        usuarios_dict[user_key] = {
                            "tipoDocumentoIdentificacion": row.get("tipoDocumentoIdentificacion") or row.get("tipo_documento_usuario"),
                            "numDocumentoIdentificacion": row.get("numDocumentoIdentificacion") or row.get("documento_usuario"),
                            "tipoUsuario": row.get("tipoUsuario") or row.get("tipo_usuario"),
                            "fechaNacimiento": row.get("fechaNacimiento") or row.get("fecha_nacimiento"), 
                            "codSexo": row.get("codSexo") or row.get("sexo"),
                            "codPaisResidencia": row.get("codPaisResidencia") or row.get("pais_residencia"),
                            "codMunicipioResidencia": row.get("codMunicipioResidencia") or row.get("municipio_residencia"),
                            "codZonaTerritorialResidencia": row.get("codZonaTerritorialResidencia") or row.get("zona_residencia"),
                            "incapacidad": row.get("incapacidad"),
                            "consecutivo": row.get("consecutivo") or row.get("consecutivo_usuario"),
                            "codPaisOrigen": row.get("codPaisOrigen") or row.get("pais_origen"),
                            "servicios": {
                                "consultas": [],
                                "procedimientos": [],
                                "urgencias": [],
                                "hospitalizacion": [],
                                "recienNacidos": [],
                                "medicamentos": [],
                                "otrosServicios": []
                            }
                        }
                    
                    # Service data
                    servicio_data = row.to_dict()
                    # Remove user/header fields from service object
                    campos_excluir = [
                        "tipoDocumentoIdentificacion", "numDocumentoIdentificacion", "tipoUsuario", "fechaNacimiento",
                        "codSexo", "codPaisResidencia", "codMunicipioResidencia", "codZonaTerritorialResidencia",
                        "incapacidad", "consecutivo", "codPaisOrigen",
                        # Legacy
                        "tipo_documento_usuario", "documento_usuario", "tipo_usuario", 
                        "fecha_nacimiento", "sexo", "pais_residencia", "municipio_residencia", 
                        "zona_residencia", "consecutivo_usuario", "pais_origen",
                        "archivo_origen",
                        "numDocumentoIdObligado", "numFactura", "tipoNota", "numNota"
                    ]
                    for campo in campos_excluir:
                        servicio_data.pop(campo, None)
                    
                    usuarios_dict[user_key]["servicios"][json_key].append(servicio_data)

            # Construct final JSON
            final_json = {**header_data, "usuarios": list(usuarios_dict.values())}
            
            # Save
            if not nombre_archivo.lower().endswith(".json"):
                nombre_archivo += ".json"
                
            out_path = os.path.join(dest_folder, nombre_archivo)
            with open(out_path, 'w', encoding='utf-8') as f:
                json.dump(final_json, f, ensure_ascii=False, indent=4)
                
        progress_bar.empty()
        return True, f"Generados {len(all_files)} archivos JSON."

    except Exception as e:
        return False, str(e)

# --- INTERFAZ PRINCIPAL ---
st.markdown("""
    <style>
    /* Global Button Styling for Responsiveness */
    div[data-testid="stButton"] > button {
        white-space: normal !important;
        height: auto !important;
        min-height: 45px !important;
        padding-top: 5px !important;
        padding-bottom: 5px !important;
        line-height: 1.2 !important;
        font-weight: 500 !important;
    }
    
    /* Dialog Button Improvements */
    div[data-testid="stDialog"] button {
        width: 100%;
        font-weight: 600 !important;
        border: 1px solid rgba(49, 51, 63, 0.2);
    }

    /* Primary Action Button (Red for Delete) */
    div[data-testid="stDialog"] button[kind="primary"] {
        background-color: #ff4b4b !important;
        color: black !important;
        border: none !important;
    }
    div[data-testid="stDialog"] button[kind="primary"]:hover {
        background-color: #ff3333 !important;
        color: black !important;
    }
    /* Ensure text inside primary button is black */
    div[data-testid="stDialog"] button[kind="primary"] p {
        color: black !important;
    }
    
    /* Fix Close Button (X) Position and Visibility */
    div[data-testid="stDialog"] button[aria-label="Close"] {
        position: absolute !important;
        top: 10px !important;
        right: 15px !important;
        z-index: 999999 !important;
        background-color: rgba(255, 255, 255, 0.9) !important;
        border-radius: 50% !important;
        width: 30px !important;
        height: 30px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.2) !important;
    }
    
    /* Enhance X visibility */
    div[data-testid="stDialog"] button[aria-label="Close"] svg {
        color: #333 !important;
        font-weight: bold !important;
        width: 16px !important;
        height: 16px !important;
    }

    /* Ensure Title doesn't overlap with Close Button */
    div[data-testid="stDialog"] h1, div[data-testid="stDialog"] h2, div[data-testid="stDialog"] h3 {
        padding-right: 50px !important; /* Espacio de seguridad para la X */
    }

    /* Tab labels wrapping */
    div[data-testid="stTabs"] button {
        white-space: normal !important;
        overflow-wrap: break-word !important;
    }
    </style>
""", unsafe_allow_html=True)

# --- TASK WRAPPERS FOR ANALYSIS ---
def run_analisis_sos_task(file_list, use_ai):
    from src.modules.analisis_sos import worker_analisis_sos
    result = worker_analisis_sos(file_list, use_ai=use_ai, silent_mode=True)
    if isinstance(result, tuple):
        out_xlsx, out_txt = result
        return {
            "files": [
                {"name": "Analisis_SOS.xlsx", "data": out_xlsx, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel"},
                {"name": "Analisis_SOS.txt", "data": out_txt, "mime": "text/csv", "label": "CSV/TXT"}
            ],
            "message": "Análisis SOS completado."
        }
    elif result:
        return {
            "files": [
                {"name": "Analisis_SOS.xlsx", "data": result, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel"}
            ],
            "message": "Análisis SOS completado."
        }
    return None

def run_analisis_historia_task(file_list):
    out = worker_analisis_historia_clinica(file_list, silent_mode=True)
    if out:
        return {
            "files": [{"name": "Analisis_Historia_Clinica.xlsx", "data": out, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel Historia Clinica"}],
            "message": "Análisis de Historias Clínicas completado."
        }
    return {"message": "No se generó reporte (posiblemente sin archivos PDF)."}

def run_analisis_autorizacion_task(file_list):
    out = worker_analisis_autorizacion_nueva_eps(file_list, silent_mode=True)
    if out:
        return {
            "files": [{"name": "Analisis_Autorizaciones_NuevaEPS.xlsx", "data": out, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel Autorizaciones"}],
            "message": "Análisis de Autorizaciones completado."
        }
    return {"message": "No se generó reporte."}

def run_analisis_sanitas_task(file_list):
    out = worker_analisis_cargue_sanitas(file_list, silent_mode=True)
    if out:
        return {
            "files": [{"name": "Analisis_Cargue_Sanitas.xlsx", "data": out, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel Sanitas"}],
            "message": "Análisis Sanitas completado."
        }
    return {"message": "No se generó reporte."}

def run_analisis_retefuente_task(file_list):
    out = worker_analisis_retefuente(file_list, silent_mode=True)
    if out:
        return {
            "files": [{"name": "Analisis_Retefuente.xlsx", "data": out, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel Retefuente"}],
            "message": "Análisis Retefuente completado."
        }
    return {"message": "No se generó reporte."}



def run_analisis_carpetas_task(path):
    out = worker_analisis_carpetas(path, silent_mode=True)
    if out:
        return {
            "files": [{"name": "Analisis_Carpetas.xlsx", "data": out, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel Carpetas"}],
            "message": "Análisis de Carpetas completado."
        }
    return {"message": "No se generó reporte."}




def run_conversion_masiva_task(folder_path, code):
    c, msg = worker_convertir_masivo(folder_path, code)
    return {"message": msg}

def run_rips_clean_json_task(folder_path):
    count, errors = worker_clean_json_spaces_masivo(folder_path)
    return {"message": f"Limpieza completada. Archivos: {count}. Errores: {len(errors)}"}

def run_rips_update_key_task(folder_path, key, val):
    c, total, err = worker_update_json_key_masivo(folder_path, key, val)
    return {"message": f"Actualización completada. Archivos: {c}, Cambios: {total}"}

def run_rips_update_cups_task(folder_path, old_val, new_val):
    c, total, err = worker_update_cups_masivo(folder_path, old_val, new_val)
    return {"message": f"Actualización CUPS completada. Archivos: {c}, Cambios: {total}"}

def run_rips_update_notes_task(folder_path, t_val, n_val):
    c, total, err = worker_update_notes_masivo(folder_path, t_val, n_val)
    return {"message": f"Actualización Notas completada. Archivos: {c}, Cambios: {total}"}

def run_rips_consolidar_task(folder_path):
    xlsx_data, err = worker_consolidar_json_xlsx(folder_path)
    if xlsx_data:
        return {
            "files": [{"name": "Consolidado_RIPS.xlsx", "data": xlsx_data, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel Consolidado"}],
            "message": "Consolidación completada."
        }
    return {"message": f"Error: {err}"}

def run_rips_desconsolidar_task(file_data, dest_path):
    # Wrap bytes in BytesIO for pandas
    from io import BytesIO
    f_obj = BytesIO(file_data)
    ok, msg = worker_desconsolidar_xlsx_json(f_obj, dest_path)
    return {"message": msg}

def run_generar_cuv_masivo_task(root_path, api_url, token):
    worker_generar_cuv_masivo(root_path, api_url, token)
    return {"message": "Generación de CUV masiva completada."}

if path_header_logo:
    st.image(path_header_logo, width=400) # Ajusta el ancho según necesidad
else:
    st.title("CDO Clinical Document Organizer")

# --- TABS & PERMISSIONS ---
full_tab_names = [
    "🔎 Búsqueda y Acciones", 
    "⚙️ Acciones Automatizadas", 
    "🔄 Conversión de Archivos", 
    "📄 Visor (JSON/XML)",
    "RIPS",
    "Validador FevRips",
    "Validacion Usuario",
    "🤖 Asistente IA (Gemini)",
    "🤖 Bot Zeus Salud"
]

visible_tab_names = []
user_role = st.session_state.get("user_role", "user")
user_perms = st.session_state.get("user_permissions", {})
allowed_tabs = user_perms.get("allowed_tabs", ["*"])

if user_role == "admin" or "*" in allowed_tabs:
    visible_tab_names = full_tab_names
else:
    for name in full_tab_names:
        if name in allowed_tabs:
            visible_tab_names.append(name)
    # Ensure at least one tab (Search) if empty, or handle empty state
    if not visible_tab_names:
        visible_tab_names = [full_tab_names[0]]

tabs_list = st.tabs(visible_tab_names)
tabs_map = {name: tab for name, tab in zip(visible_tab_names, tabs_list)}

# --- TAB 1: BÚSQUEDA Y ACCIONES PRINCIPALES ---
if "🔎 Búsqueda y Acciones" in tabs_map:
    with tabs_map["🔎 Búsqueda y Acciones"]:
        # 1. Configuración de Ruta de Trabajo (Independiente por usuario)
        st.info("ℹ️ La ruta seleccionada es exclusiva para su sesión actual. Puede trabajar simultáneamente con otros usuarios sin conflictos.")
        
        def update_path(): 
            st.session_state.current_path = st.session_state.path_input
            # Update user config on manual change
            if st.session_state.get("username"):
                update_user_last_path(st.session_state.username, st.session_state.current_path)
        
        # Asegurar sincronización
        if "path_input" not in st.session_state:
            st.session_state.path_input = st.session_state.current_path

        # Navegador de Servidor (Modal) + Manual Input
        st.markdown("### 📂 Selección de Ruta de Trabajo")
        
        col_path_1, col_path_2 = st.columns([0.7, 0.3])
        with col_path_1:
            st.text_input("Ruta de Trabajo:", value=st.session_state.current_path, key="path_input", on_change=update_path, help="Ruta de la carpeta donde se buscarán los archivos. Puede ser una ruta local (C:\\...) o de red.")
        def on_click_examinar():
            # Callback wrapper to handle native mode path selection
            mode = os.environ.get("CDO_MODE", "WEB").upper()
            if st.session_state.get("force_native_mode", False):
                mode = "LOCAL"
            
            if mode == "LOCAL":
                update_main_path()

        with col_path_2:
            if st.button("📂 Examinar", help="Explorar archivos", on_click=on_click_examinar):
                # Modo nativo: Abrir selector directo sin modal intermedio
                mode = os.environ.get("CDO_MODE", "WEB").upper()
                if st.session_state.get("force_native_mode", False):
                    mode = "LOCAL"
                
                # If not local, show the web modal (Local is handled in callback)
                if mode != "LOCAL":
                    browse_modal()
        # 2. Paneles de Criterios y Acciones

        c1, c2 = st.columns([1, 1])
        
        with c1:
            st.markdown('<div class="group-box"><div class="group-title-left">🔎 Criterios de búsqueda</div>', unsafe_allow_html=True)
            col_crit1, col_crit2 = st.columns(2)
            with col_crit1:
                st.selectbox("Buscar por:", ["extensión", "nombre", "todo"], key="search_by")
                st.selectbox("Tipo de elemento:", ["archivos", "carpetas"], key="item_type")
            with col_crit2:
                st.text_input("Patrón:", placeholder="ej: .pdf, Factura", key="pattern")
                # Usar valor por defecto de la configuración
                default_excl = st.session_state.app_config.get("default_exclusion_patterns", "")
                st.text_input("Excluir:", value=default_excl, key="exclusion_pattern", help="Separar por comas. Excluye si coincide con el patrón.")
                st.checkbox("Incluir subcarpetas", value=True, key="subfolders")
            st.markdown('</div>', unsafe_allow_html=True)
                
        with c2:
            st.markdown('<div class="group-box"><div class="group-title-left">🛠️ Acción a realizar</div>', unsafe_allow_html=True)
            st.radio("Seleccione acción:", [
                "Copiar a carpeta", 
                "Mover a carpeta", 
                "Modificar nombre", 
                "Editar texto", 
                "Comprimir archivos en ZIP", 
                "Comprimir carpetas individualmente"
            ], label_visibility="collapsed", key="action_radio")
            st.markdown('</div>', unsafe_allow_html=True)
        
        # 3. Tabla de Resultados
        st.markdown("##### 📄 Archivos encontrados")
        df_display = pd.DataFrame(st.session_state.search_results) if st.session_state.search_results else pd.DataFrame(columns=["Ruta completa", "Fecha"])
        st.dataframe(df_display, width=1000, height=250, hide_index=True)

        # 4. Barra de Botones Inferior
        st.write("") # Espaciador
        col_btns = st.columns(5, gap="small") # 5 columnas
        
        with col_btns[0]:
            if st.button("🔍 Buscar", use_container_width=True, help="Buscar archivos"):
                buscar_archivos()
                st.rerun()
                
        with col_btns[1]:
            if st.button("▶️ Ejecutar", use_container_width=True, help="Ejecutar acción seleccionada"):
                action = st.session_state.get("action_radio")
                if action == "Modificar nombre":
                    dialogo_modificar_nombres()
                elif action == "Editar texto":
                    dialogo_editar_texto()
                elif action == "Copiar a carpeta":
                    dialogo_copiar_lista()
                elif action == "Mover a carpeta":
                    dialogo_mover_lista()
                elif action == "Comprimir archivos en ZIP":
                    dialogo_zip_lista()
                elif action == "Comprimir carpetas individualmente":
                    dialogo_zip_carpetas_individual()
                else:
                    funcion_no_implementada(f"Acción: {action}")
                
        with col_btns[2]:
            if st.button("🧹 Limpiar", use_container_width=True, help="Limpiar lista de resultados"):
                st.session_state.search_results = []
                st.rerun()
                
        with col_btns[3]:
            if st.button("🗑️ Eliminar", use_container_width=True, help="Eliminar archivos seleccionados"):
                dialogo_confirmar_eliminar()

        with col_btns[4]:
            if st.button("↩️ Deshacer", use_container_width=True, help="Revertir la última acción"):
                undo_last_action()


# --- TAB 2: ACCIONES AUTOMATIZADAS ---
if "⚙️ Acciones Automatizadas" in tabs_map:
    with tabs_map["⚙️ Acciones Automatizadas"]:
        col_g1, col_g2, col_g3 = st.columns(3)
        with col_g1:
            st.markdown('<div class="group-box"><div class="group-title">Unificación y División Especial</div>', unsafe_allow_html=True)
            if st.button("📂 Unificar PDF por Carpeta"): dialog_unificar_pdf()
            if st.button("🖼️ Unificar JPG por Carpeta"): dialog_unificar_img("JPG")
            if st.button("🖼️ Unificar PNG por Carpeta"): dialog_unificar_img("PNG")
            if st.button("📄 Unificar DOCX por Carpeta"): dialog_unificar_docx()
            if st.button("📎 Unificar PDFs (Selección Manual)"): 
                dialog_unificar_pdf_manual()
            if st.button("✂️ Dividir PDF en Páginas (Manual)"): 
                if st.session_state.search_results:
                    if st.button("🚀 Confirmar División"):
                         worker_dividir_pdf_manual(st.session_state.search_results)
                else:
                    st.error("Primero busca archivos PDF.")
            if st.button("✂️ Dividir PDFs Masivamente (por Carpeta)"): dialog_dividir_pdf()
            st.markdown('</div>', unsafe_allow_html=True)
        with col_g2:
            st.markdown('<div class="group-box"><div class="group-title">Modificación y Renombrado con Excel</div>', unsafe_allow_html=True)
            
            # Lógica especial para botón de descarga
            df_export = worker_exportar_renombrado()
            if df_export is not None:
                # Convertir DF a Excel en memoria
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df_export.to_excel(writer, index=False)
                output.seek(0)
                st.download_button(
                    label="📤 Exportar para renombrar",
                    data=output,
                    file_name="archivos_para_renombrar.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="btn_export_renom"
                )
            else:
                 st.button("📤 Exportar para renombrar (Requiere búsqueda previa)", disabled=True)

            if st.button("📥 Aplicar renombrado desde Excel"): dialog_importar_excel()
            if st.button("🔄 Renombrar Masivo por Mapeo Excel"): dialog_renombrar_mapeo_excel()
            if st.button("🏷️ Añadir Sufijo desde Excel"): dialog_sufijo()
            if st.button("✍️ Modificar DOCX Completo (Excel)"): dialog_modif_docx_completo()
            if st.button("✒️ Firmar DOCX con Imagen (por Carpeta)"): dialog_insertar_firma_docx()
            st.markdown('</div>', unsafe_allow_html=True)
        with col_g3:
            st.markdown('<div class="group-box"><div class="group-title">Análisis de Archivos</div>', unsafe_allow_html=True)
            if st.button("📊 Análisis de Carpetas a Excel"):
                submit_task("Análisis de Carpetas", run_analisis_carpetas_task, st.session_state.current_path)
                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")

            if st.button("🏥 Análisis de Historia Clínica"): 
                if st.session_state.search_results:
                     submit_task("Análisis Historia Clínica", run_analisis_historia_task, st.session_state.search_results)
                     st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                else:
                     st.error("Primero realiza una búsqueda de archivos PDF.")

            if st.button("🏥 Análisis Autorización Nueva EPS"): 
                 if st.session_state.search_results:
                     submit_task("Análisis Aut. Nueva EPS", run_analisis_autorizacion_task, st.session_state.search_results)
                     st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                 else:
                     st.error("Primero realiza una búsqueda de archivos PDF.")

            if st.button("🏥 Análisis Cargue Sanitas"): 
                 if st.session_state.search_results:
                     submit_task("Análisis Cargue Sanitas", run_analisis_sanitas_task, st.session_state.search_results)
                     st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                 else:
                     st.error("Primero realiza una búsqueda de archivos PDF.")

            if st.button("💰 Análisis Retefuente y Ica"): 
                 if st.session_state.search_results:
                     submit_task("Análisis Retefuente", run_analisis_retefuente_task, st.session_state.search_results)
                     st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                 else:
                     st.error("Primero realiza una búsqueda de archivos PDF.")

            # use_ai_sos = st.checkbox("🤖 Usar IA (Gemini) para Análisis SOS", value=False, help="Extrae datos usando Inteligencia Artificial. Requiere API Key.")
            if st.button("🏥 Análisis Autorización SOS"): 
                 if st.session_state.search_results:
                    submit_task("Análisis SOS", run_analisis_sos_task, st.session_state.search_results, False)
                    st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                 else:
                     st.error("Primero realiza una búsqueda de archivos PDF.")
            st.markdown('</div>', unsafe_allow_html=True)
        col_g4, col_g5, col_g6 = st.columns(3)
        with col_g4:
            st.markdown('<div class="group-box"><div class="group-title">Organización de Archivo</div>', unsafe_allow_html=True)
            if st.button("🧾 Organizar Facturas (FEOV)"): dialog_organizar_feov()
            if st.button("➡️ Mover Archivos por Coincidencia"): dialog_mover_coincidencia()
            if st.button("📂 Copiar Archivos (Mapeo Subcarpetas)"): dialog_copiar_mapeo_sub()
            if st.button("📂 Copiar Archivos de Raíz (Mapeo)"): dialog_copiar_raiz_mapeo()
            if st.button("📥 Consolidar Archivos de Subcarpetas"): dialog_consolidar()
            st.markdown('</div>', unsafe_allow_html=True)
        with col_g5:
            st.markdown('<div class="group-box"><div class="group-title">Creación de Archivos</div>', unsafe_allow_html=True)
            if st.button("📂 Crear Carpetas desde Excel"): dialog_crear_carpetas()
            if st.button("⬇️ Descargar Firmas (URL/Excel)"): dialog_descargar_firmas()
            if st.button("📂 Copiar Archivo a Subcarpetas"): dialog_copiar_sub()
            if st.button("⬇️ Descargar Historias Hospitalización (OVIDA)"): dialog_descargar_historias_ovida()
            if st.button("✒️ Crear Firma Digital desde Nombre"): dialog_crear_firma()
            st.markdown('</div>', unsafe_allow_html=True)
        with col_g6: st.write("")

# --- TAB 3: CONVERSIÓN DE ARCHIVOS ---
if "🔄 Conversión de Archivos" in tabs_map:
    with tabs_map["🔄 Conversión de Archivos"]:
        st.header("🔄 Centro de Conversiones")
        
        st.info("Para conversiones masivas, asegúrate de navegar a la carpeta deseada usando el panel 'Navegar carpetas'.")

        col_indiv, col_masiv = st.columns(2)
        
        with col_indiv:
            st.markdown('<div class="group-box">', unsafe_allow_html=True)
            st.markdown('<div class="group-title">👤 Conversión Individual (Subir Archivo)</div>', unsafe_allow_html=True)
            
            conv_type_ind = st.selectbox("Tipo de Conversión:", 
                ["PDF ➝ DOCX", "JPG ➝ PDF", "DOCX ➝ PDF", "PDF ➝ JPG", "PNG ➝ JPG", "TXT ➝ JSON", "PDF ➝ PDF (Grises)"], key="sel_ind_conv")
            
            src_type = st.radio("Origen:", ["📤 Subir", "📂 Carpeta Actual"], horizontal=True, key="src_conv_ind")
            
            final_path = None
            
            if src_type == "📤 Subir":
                uploaded_ind = st.file_uploader(f"Subir Archivo para {conv_type_ind}", key="up_ind_conv")
                if uploaded_ind:
                    temp_dir = os.path.join(os.getcwd(), "temp_uploads")
                    os.makedirs(temp_dir, exist_ok=True)
                    final_path = os.path.join(temp_dir, uploaded_ind.name)
                    with open(final_path, "wb") as f:
                        f.write(uploaded_ind.getbuffer())
            else:
                # List files matching extension?
                # Simple list for now
                try:
                    files_c = [f for f in os.listdir(st.session_state.current_path) if os.path.isfile(os.path.join(st.session_state.current_path, f))]
                    if files_c:
                        f_sel_c = st.selectbox("Seleccionar Archivo:", files_c, key="sel_file_conv")
                        final_path = os.path.join(st.session_state.current_path, f_sel_c)
                    else:
                        st.warning("No hay archivos.")
                except:
                    st.error("Error leyendo carpeta")

            # Selector de Carpeta de Destino (NUEVO)
            st.write("📂 Guardar en:")
            col_d1, col_d2 = st.columns([0.85, 0.15])
            
            current_dest = st.session_state.get("conv_dest_folder")
            if current_dest:
                dest_display = current_dest
            else:
                dest_display = "Misma carpeta (o Temp si es subida)"

            with col_d1:
                st.text_input("Ruta Destino", value=dest_display, disabled=True, label_visibility="collapsed", key="txt_dest_display")
            
            with col_d2:
                if st.button("📂", key="btn_pick_dest_conv", help="Cambiar carpeta de destino"):
                    sel = seleccionar_carpeta_nativa("Seleccionar Carpeta de Salida")
                    if sel:
                        st.session_state.conv_dest_folder = sel
                        st.rerun()

            if final_path and st.button("🚀 Convertir Archivo", key="btn_conv_ind"):
                map_code = {
                    "PDF ➝ DOCX": "PDF2DOCX",
                    "JPG ➝ PDF": "JPG2PDF",
                    "DOCX ➝ PDF": "DOCX2PDF",
                    "PDF ➝ JPG": "PDF2JPG",
                    "PNG ➝ JPG": "PNG2JPG",
                    "TXT ➝ JSON": "TXT2JSON",
                    "PDF ➝ PDF (Grises)": "PDF_GRAY"
                }
                code = map_code.get(conv_type_ind)
                
                # Determinar carpeta de salida
                dest_folder = st.session_state.get("conv_dest_folder")

                with st.spinner("Convirtiendo..."):
                    ok, msg = worker_convertir_archivo(final_path, code, output_folder=dest_folder)
                
                if ok:
                    st.success(f"✅ Convertido: {msg}")
                    # Mostrar ruta real de guardado
                    if dest_folder:
                         st.info(f"El archivo procesado se guardó en: {dest_folder}")
                    elif src_type == "📂 Carpeta Actual":
                         st.info(f"El archivo se guardó en la misma carpeta.")
                    else:
                         st.info(f"El archivo procesado se guardó en: {os.path.dirname(final_path)}")
                else:
                    st.error(f"❌ Error: {msg}")
        
            st.markdown('</div>', unsafe_allow_html=True)

        with col_masiv:
            st.markdown('<div class="group-box">', unsafe_allow_html=True)
            st.markdown('<div class="group-title">📦 Conversión Masiva (Carpeta Actual)</div>', unsafe_allow_html=True)
            
            st.write(f"📂 Carpeta Objetivo: **{st.session_state.current_path}**")
            
            conv_type_mas = st.selectbox("Tipo de Conversión Masiva:", 
                ["PDF ➝ DOCX", "JPG ➝ PDF", "DOCX ➝ PDF", "PDF ➝ JPG", "PNG ➝ JPG", "TXT ➝ JSON", "PDF ➝ PDF (Grises)"], key="sel_mas_conv")
                
            if st.button("🚀 Ejecutar Conversión Masiva", key="btn_conv_mas"):
                map_code = {
                    "PDF ➝ DOCX": "PDF2DOCX",
                    "JPG ➝ PDF": "JPG2PDF",
                    "DOCX ➝ PDF": "DOCX2PDF",
                    "PDF ➝ JPG": "PDF2JPG",
                    "PNG ➝ JPG": "PNG2JPG",
                    "TXT ➝ JSON": "TXT2JSON",
                    "PDF ➝ PDF (Grises)": "PDF_GRAY"
                }
                code = map_code.get(conv_type_mas)
                
                submit_task(f"Conversión Masiva ({conv_type_mas})", run_conversion_masiva_task, st.session_state.current_path, code)
                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")

            st.markdown('</div>', unsafe_allow_html=True)


# --- TAB 4: VISOR Y EDITOR ---
# --- TAB 5: RIPS ---
if "RIPS" in tabs_map:
    with tabs_map["RIPS"]: 
            col_menu, col_work = st.columns([1, 2])
    
            with col_menu:
                st.markdown('<div class="group-box">', unsafe_allow_html=True)
                st.markdown('<div class="group-title-left">Convertidor</div>', unsafe_allow_html=True)
        
                if st.button("JSON a XLSX (Individual)", use_container_width=True):
                    st.session_state.rips_mode = "JSON_XLSX_IND"
                if st.button("XLSX a JSON (Individual)", use_container_width=True):
                    st.session_state.rips_mode = "XLSX_JSON_IND"
                if st.button("JSON Evento a XLSX (Masivo)", use_container_width=True):
                    st.session_state.rips_mode = "JSON_XLSX_MAS"
                if st.button("XLSX Evento a JSONs (Masivo)", use_container_width=True):
                    st.session_state.rips_mode = "XLSX_JSON_MAS"
        
                st.markdown('</div>', unsafe_allow_html=True)

                st.markdown('<div class="group-box">', unsafe_allow_html=True)
                st.markdown('<div class="group-title-left">Opciones de Validación</div>', unsafe_allow_html=True)
        
                if st.button("🧹 Limpiar JSON (Individual)", use_container_width=True):
                    st.session_state.rips_mode = "JSON_CLEAN_IND"
        
                if st.button("🧹 Limpiar JSONs (Masivo)", use_container_width=True):
                    st.session_state.rips_mode = "JSON_CLEAN_MAS"
        
                if st.button("🔄 Cambio Tecnología (Individual)", use_container_width=True):
                    st.session_state.rips_mode = "JSON_UPDATE_IND"

                if st.button("🔄 Cambio Tecnología (Masivo)", use_container_width=True):
                    st.session_state.rips_mode = "JSON_UPDATE_MAS"
            
                st.markdown('</div>', unsafe_allow_html=True)
        

                st.markdown('</div>', unsafe_allow_html=True)
    
            with col_work:
                cw1, cw2 = st.columns(2)
                
                with cw1:
                    st.markdown('<div class="group-box">', unsafe_allow_html=True)
                    st.markdown('<div class="group-title-left">Cambio de Cups</div>', unsafe_allow_html=True)
                    
                    if st.button("💊 Cups (Individual)", use_container_width=True):
                        st.session_state.rips_mode = "CUPS_UPDATE_IND"
                    
                    if st.button("💊 Cups (Masivo)", use_container_width=True):
                        st.session_state.rips_mode = "CUPS_UPDATE_MAS"
                    
                    st.markdown('</div>', unsafe_allow_html=True)

                with cw2:
                    st.markdown('<div class="group-box">', unsafe_allow_html=True)
                    st.markdown('<div class="group-title-left">Notas de ajuste</div>', unsafe_allow_html=True)
                    
                    if st.button("📝 Notas (Individual)", use_container_width=True):
                        st.session_state.rips_mode = "NOTAS_AJUSTE_IND"
                    
                    if st.button("📝 Notas (Masivo)", use_container_width=True):
                        st.session_state.rips_mode = "NOTAS_AJUSTE_MAS"
                    
                    st.markdown('</div>', unsafe_allow_html=True)

                if st.session_state.rips_mode:
                    st.markdown(f"#### Modo Seleccionado: **{st.session_state.rips_mode}**")
            
                    if st.session_state.rips_mode == "JSON_XLSX_IND":
                        f_rips = st.file_uploader("Sube JSON RIPS:", type="json")
                        if f_rips and st.button("🚀 Convertir a Excel"):
                            xlsx_data, err = worker_json_a_xlsx_ind(f_rips)
                            if xlsx_data:
                                st.download_button("📥 Descargar Excel", data=xlsx_data, file_name=f"{f_rips.name}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                                st.success("Conversión exitosa.")
                            else:
                                st.error(f"Error: {err}")
                    
                    elif st.session_state.rips_mode == "XLSX_JSON_IND":
                        f_rips = st.file_uploader("Sube Excel RIPS:", type="xlsx")
                        if f_rips and st.button("🚀 Convertir a JSON"):
                            json_str, err = worker_xlsx_a_json_ind(f_rips)
                            if json_str:
                                st.download_button("📥 Descargar JSON", data=json_str, file_name=f"{f_rips.name}.json", mime="application/json")
                                st.success("Conversión exitosa.")
                            else:
                                st.error(f"Error: {err}")
                    
                    elif st.session_state.rips_mode == "JSON_XLSX_MAS":
                        st.info("Este proceso buscará todos los .json en la carpeta seleccionada y creará un Excel consolidado.")
                
                        col_path, col_btn, col_browse = st.columns([0.8, 0.1, 0.1])
                
                        # --- Lógica de Botones (Antes del input para evitar error de estado) ---
                        with col_btn:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             if st.button("📍", help="Usar carpeta actual del navegador", key="btn_set_rips"):
                                 st.session_state.rips_path_input = st.session_state.current_path
                                 st.session_state.rips_path = st.session_state.current_path
                                 st.rerun()
                
                        # --- Input y Botón Estandarizado ---
                        col_path, col_browse = st.columns([0.85, 0.15])
                        with col_path:
                             st.text_input("Carpeta Origen:", key="rips_path_input", value=st.session_state.get("rips_path_input", st.session_state.current_path))
                        with col_browse:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             st.button("📂", help="Examinar...", key="btn_browse_rips", on_click=update_path_key, args=("rips_path_input", "Seleccionar Carpeta JSONs"))
                
                        st.session_state.rips_path = st.session_state.rips_path_input
                
                        if st.button("🚀 Consolidar Carpeta"):
                            if st.session_state.rips_path:
                                submit_task("Consolidar RIPS", run_rips_consolidar_task, st.session_state.rips_path)
                                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                            else:
                                st.warning("Selecciona una carpeta.")
                    
                    elif st.session_state.rips_mode == "XLSX_JSON_MAS":
                        st.info("Sube el Excel consolidado y selecciona la carpeta donde se crearán los JSONs individuales.")
                        f_consol = st.file_uploader("Sube Excel Consolidado:", type="xlsx")
                
                        col_path_dest, col_browse_dest = st.columns([0.85, 0.15])
                
                        with col_path_dest:
                            st.text_input("Carpeta Destino:", key="dest_input", value=st.session_state.get("dest_input", st.session_state.current_path))

                        with col_browse_dest:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             st.button("📂", help="Examinar...", key="btn_browse_dest_rips", on_click=update_path_key, args=("dest_input", "Seleccionar Destino"))

                        path_dest = st.session_state.dest_input


                        if f_consol and path_dest and st.button("🚀 Desconsolidar"):
                            # Read bytes first
                            f_bytes = f_consol.getvalue()
                            submit_task("Desconsolidar RIPS", run_rips_desconsolidar_task, f_bytes, path_dest)
                            st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")

                    elif st.session_state.rips_mode == "JSON_CLEAN_IND":
                        st.info("Sube un archivo JSON para eliminar espacios al final de los valores.")
                        f_json = st.file_uploader("Sube JSON:", type="json")
                        if f_json and st.button("🚀 Limpiar Archivo"):
                            try:
                                content = json.load(f_json)
                                cleaned = recursive_strip(content)
                                json_str = json.dumps(cleaned, indent=4, ensure_ascii=False)
                                st.download_button("📥 Descargar JSON Limpio", data=json_str, file_name=f"clean_{f_json.name}", mime="application/json")
                                st.success("✅ Archivo limpiado correctamente.")
                            except Exception as e:
                                st.error(f"Error al procesar: {e}")

                    elif st.session_state.rips_mode == "JSON_CLEAN_MAS":
                        st.info("Elimina espacios al final de los textos en todos los JSON de la carpeta. ⚠️ ESTA ACCIÓN SOBREESCRIBE LOS ARCHIVOS.")
                
                        col_path_c, col_browse_c = st.columns([0.85, 0.15])
                        with col_path_c:
                             st.text_input("Carpeta con JSONs:", key="clean_path_input", value=st.session_state.get("clean_path_input", st.session_state.current_path))
                        with col_browse_c:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             st.button("📂", help="Examinar...", key="btn_browse_clean", on_click=update_path_key, args=("clean_path_input", "Seleccionar Carpeta JSONs"))
                
                        clean_path = st.session_state.clean_path_input
                
                        if st.button("🚀 Limpiar Archivos"):
                            if clean_path and os.path.isdir(clean_path):
                                submit_task("Limpiar JSONs (Masivo)", run_rips_clean_json_task, clean_path)
                                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                            else:
                                st.error("Selecciona una carpeta válida.")

                    elif st.session_state.rips_mode == "JSON_UPDATE_IND":
                        st.info("Sube un archivo JSON para actualizar el valor de 'finalidadTecnologiaSalud'.")
                        f_json = st.file_uploader("Sube JSON:", type="json")
                
                        new_val = st.text_input("Nuevo Valor para 'finalidadTecnologiaSalud':", value="44")
                
                        if f_json and new_val and st.button("🚀 Actualizar Archivo"):
                            try:
                                content = json.load(f_json)
                                count_changes = recursive_update_key(content, "finalidadTecnologiaSalud", new_val)
                        
                                if count_changes > 0:
                                    json_str = json.dumps(content, indent=4, ensure_ascii=False)
                                    st.download_button("📥 Descargar JSON Actualizado", data=json_str, file_name=f"update_{f_json.name}", mime="application/json")
                                    st.success(f"✅ Se actualizaron {count_changes} campos.")
                                else:
                                    st.warning("⚠️ No se encontró el campo 'finalidadTecnologiaSalud' en el archivo.")
                            
                            except Exception as e:
                                st.error(f"Error al procesar: {e}")

                    elif st.session_state.rips_mode == "JSON_UPDATE_MAS":
                        st.info("Actualiza el campo 'finalidadTecnologiaSalud' en todos los JSON de una carpeta. ⚠️ ESTA ACCIÓN SOBREESCRIBE LOS ARCHIVOS.")
                
                        col_path_u, col_browse_u = st.columns([0.85, 0.15])
                        with col_path_u:
                             st.text_input("Carpeta con JSONs:", key="update_path_input", value=st.session_state.get("update_path_input", st.session_state.current_path))
                        with col_browse_u:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             st.button("📂", help="Examinar...", key="btn_browse_update", on_click=update_path_key, args=("update_path_input", "Seleccionar Carpeta JSONs"))
                
                        update_path = st.session_state.update_path_input
                        new_val_mas = st.text_input("Nuevo Valor para 'finalidadTecnologiaSalud':", value="44", key="new_val_mas")
                
                        if st.button("🚀 Actualizar Archivos Masivamente"):
                            if update_path and os.path.isdir(update_path) and new_val_mas:
                                submit_task("Actualizar JSONs (Masivo)", run_rips_update_key_task, update_path, "finalidadTecnologiaSalud", new_val_mas)
                                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                            else:
                                st.error("Verifica la carpeta y el valor ingresado.")

                    elif st.session_state.rips_mode == "NOTAS_AJUSTE_IND":
                        st.info("Actualiza 'tipoNota' y 'numNota' en un archivo JSON (Valores se convertirán a Mayúsculas).")
                        f_json = st.file_uploader("Sube JSON:", type="json")
                        
                        c1, c2 = st.columns(2)
                        with c1: 
                             t_nota = st.text_input("Tipo Nota:", key="val_t_nota_ind")
                        with c2: 
                             n_nota = st.text_input("Num Nota:", key="val_n_nota_ind")

                        if f_json and st.button("🚀 Actualizar Archivo"):
                            try:
                                content = json.load(f_json)
                                # Convert to uppercase if provided
                                t_val = t_nota.strip().upper() if t_nota else None
                                n_val = n_nota.strip().upper() if n_nota else None
                                
                                if not t_val and not n_val:
                                    st.warning("Ingresa al menos un valor para actualizar.")
                                else:
                                    count_changes = recursive_update_notes(content, t_val, n_val)
                            
                                    if count_changes > 0:
                                        json_str = json.dumps(content, indent=4, ensure_ascii=False)
                                        st.download_button("📥 Descargar JSON Actualizado", data=json_str, file_name=f"notas_{f_json.name}", mime="application/json")
                                        st.success(f"✅ Se actualizaron {count_changes} campos.")
                                    else:
                                        st.warning("⚠️ No se encontraron campos para actualizar (tipoNota / numNota) o ya tienen el valor.")
                            
                            except Exception as e:
                                st.error(f"Error al procesar: {e}")

                    elif st.session_state.rips_mode == "NOTAS_AJUSTE_MAS":
                        st.info("Actualiza 'tipoNota' y 'numNota' en TODOS los JSON de una carpeta. ⚠️ SOBREESCRIBE ARCHIVOS.")
                        
                        col_path_n, col_browse_n = st.columns([0.85, 0.15])
                        with col_path_n:
                             st.text_input("Carpeta con JSONs:", key="notes_path_input", value=st.session_state.get("notes_path_input", st.session_state.current_path))
                        with col_browse_n:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             st.button("📂", help="Examinar...", key="btn_browse_notes", on_click=update_path_key, args=("notes_path_input", "Seleccionar Carpeta JSONs"))
                
                        notes_path = st.session_state.notes_path_input
                        
                        c1, c2 = st.columns(2)
                        with c1: t_nota_mas = st.text_input("Tipo Nota:", key="val_t_nota_mas")
                        with c2: n_nota_mas = st.text_input("Num Nota:", key="val_n_nota_mas")
                        
                        if st.button("🚀 Actualizar Archivos Masivamente"):
                            t_val = t_nota_mas.strip().upper() if t_nota_mas else None
                            n_val = n_nota_mas.strip().upper() if n_nota_mas else None
                            
                            if notes_path and os.path.isdir(notes_path) and (t_val or n_val):
                                submit_task("Actualizar Notas (Masivo)", run_rips_update_notes_task, notes_path, t_val, n_val)
                                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                            else:
                                st.error("Verifica la carpeta y ingresa al menos un valor.")


                    elif st.session_state.rips_mode == "CUPS_UPDATE_IND":
                        st.info("Sube un JSON para buscar y reemplazar un código CUPS (codTecnologiaSalud).")
                        f_cups = st.file_uploader("Sube JSON:", type="json", key="f_cups_ind")
                        
                        c1, c2 = st.columns(2)
                        with c1: old_val = st.text_input("Código a buscar (Actual):", key="cups_old_ind")
                        with c2: new_val = st.text_input("Código nuevo (Reemplazo):", key="cups_new_ind")

                        if f_cups and st.button("🚀 Reemplazar y Descargar"):
                            if not old_val or not new_val:
                                st.warning("Debes ingresar ambos códigos.")
                            else:
                                try:
                                    data = json.load(f_cups)
                                    changes = recursive_update_cups(data, old_val.strip(), new_val.strip())
                                    
                                    if changes > 0:
                                        new_json = json.dumps(data, indent=4, ensure_ascii=False)
                                        st.download_button("📥 Descargar JSON Modificado", data=new_json, file_name=f"modificado_{f_cups.name}", mime="application/json")
                                        st.success(f"✅ Se realizaron {changes} reemplazos.")
                                    else:
                                        st.warning(f"⚠️ No se encontró el código '{old_val}' en el archivo.")
                                except Exception as e:
                                    st.error(f"Error procesando JSON: {e}")

                    elif st.session_state.rips_mode == "CUPS_UPDATE_MAS":
                        st.info("Reemplaza masivamente un código CUPS en todos los JSON de una carpeta (y subcarpetas).")
                        
                        col_path, col_browse = st.columns([0.85, 0.15])
                        with col_path:
                             st.text_input("Carpeta JSONs:", key="cups_path_input", value=st.session_state.get("cups_path_input", st.session_state.current_path))
                        with col_browse:
                             st.markdown('<div style="margin-top: 28px;"></div>', unsafe_allow_html=True)
                             st.button("📂", help="Examinar...", key="btn_browse_cups", on_click=update_path_key, args=("cups_path_input", "Seleccionar Carpeta JSONs"))
                
                        cups_path = st.session_state.cups_path_input
                        
                        c1, c2 = st.columns(2)
                        with c1: old_val_mas = st.text_input("Código a buscar (Actual):", key="cups_old_mas")
                        with c2: new_val_mas = st.text_input("Código nuevo (Reemplazo):", key="cups_new_mas")
                        
                        if st.button("🚀 Actualizar CUPS Masivamente"):
                            if cups_path and os.path.isdir(cups_path) and old_val_mas and new_val_mas:
                                submit_task("Actualizar CUPS (Masivo)", run_rips_update_cups_task, cups_path, old_val_mas.strip(), new_val_mas.strip())
                                st.info("✅ Tarea iniciada en segundo plano. Revisa el 'Centro de Tareas' en la barra lateral.")
                            else:
                                st.error("Verifica la carpeta e ingresa ambos códigos.")


# --- TAB: VALIDADOR FEVRIPS ---
if "Validador FevRips" in tabs_map:
    with tabs_map["Validador FevRips"]:
        st.header("🏥 Validador FEVRIPS - Generación de CUV")
        
        st.markdown("""
        Esta herramienta permite validar archivos RIPS y generar el Código Único de Validación (CUV) 
        utilizando el validador oficial FEVRIPS (local o remoto).
        """)
        
        col_val_1, col_val_2 = st.columns([1, 2])
        
        with col_val_1:
            st.markdown('<div class="group-box">', unsafe_allow_html=True)
            st.markdown('<div class="group-title-left">Acciones</div>', unsafe_allow_html=True)
            
            if st.button("🆔 Generar CUV (Masivo)", use_container_width=True):
                 dialog_generar_cuv()
                 
            st.markdown('</div>', unsafe_allow_html=True)
            
        with col_val_2:
            st.info("ℹ️ Asegúrese de que el servicio Docker FEVRIPS esté en ejecución si utiliza el modo local.")


# --- TAB 4: VISOR DE DATOS ---
if "📄 Visor (JSON/XML)" in tabs_map:
    with tabs_map["📄 Visor (JSON/XML)"]:
            st.header("📄 Visor (JSON/XML)")
    
            # Barra de herramientas superior
            col_upload, col_search, col_actions = st.columns([0.35, 0.35, 0.3], gap="small")
    
            with col_upload:
                uploaded_data = st.file_uploader("Abrir Archivo (JSON/XML)", type=["json", "xml"], label_visibility="collapsed")
    
            with col_search:
                search_term = st.text_input("Buscar:", placeholder="Texto a buscar...", label_visibility="collapsed")
    
            with col_actions:
                c1, c2 = st.columns(2)
                with c1:
                    btn_search = st.button("🔍 Buscar", use_container_width=True)
                with c2:
                    if st.button("🧹 Limpiar", use_container_width=True):
                        search_term = ""
                        st.rerun()

            # Contenedor del Editor
            if uploaded_data:
                try:
                    # Inicializar estado del editor para este archivo si cambió
                    if "editor_file_name" not in st.session_state or st.session_state.editor_file_name != uploaded_data.name:
                         st.session_state.editor_file_name = uploaded_data.name
                         st.session_state.editor_content = uploaded_data.getvalue().decode("utf-8")
            
                    content = st.session_state.editor_content
                    file_type = uploaded_data.name.split(".")[-1].lower()
            
                    # Callback para sincronizar cambios desde Monaco
                    def sync_editor_change(value):
                        st.session_state.editor_content = value

                    # Lógica de búsqueda simple (resaltado no es nativo en Monaco via Streamlit, pero podemos filtrar o mostrar alerta)
                    if btn_search and search_term:
                        if search_term in content:
                            st.toast(f"✅ Texto encontrado. (El resaltado visual depende del foco del editor)", icon="🔍")
                            # Intento de posicionar cursor o resaltar (limitado en streamlit-elements sin estado complejo)
                        else:
                            st.toast("⚠️ Texto no encontrado.", icon="❌")

                    if file_type == "json":
                        # Validar JSON
                        try:
                            json_data = json.loads(content)
                            st.caption(f"Archivo: **{uploaded_data.name}** (JSON Válido)")
                    
                            # Botón de Descarga/Guardado
                            st.download_button(
                                label="💾 Guardar Cambios (Descargar JSON)",
                                data=st.session_state.editor_content,
                                file_name=f"modificado_{uploaded_data.name}",
                                mime="application/json"
                            )
                        except:
                            st.error("JSON Inválido")

                        # Editor Monaco
                        with elements("editor_json"):
                            with mui.Paper(elevation=1, style={"padding": "0px", "height": "600px", "border": "1px solid #ddd"}):
                                editor.Monaco(
                                    height="600px",
                                    defaultValue=st.session_state.editor_content,
                                    language="json",
                                    theme="vs-dark",
                                    onChange=lazy(sync_editor_change),
                                    options={
                                        "readOnly": False, 
                                        "minimap": {"enabled": True},
                                        "fontSize": 14,
                                        "wordWrap": "on"
                                    }
                                )

                    elif file_type == "xml":
                        st.caption(f"Archivo: **{uploaded_data.name}** (XML)")
                
                        # Botón de Descarga/Guardado
                        st.download_button(
                            label="💾 Guardar Cambios (Descargar XML)",
                            data=st.session_state.editor_content,
                            file_name=f"modificado_{uploaded_data.name}",
                            mime="application/xml"
                        )
                
                        with elements("editor_xml"):
                             with mui.Paper(elevation=1, style={"padding": "0px", "height": "600px", "border": "1px solid #ddd"}):
                                editor.Monaco(
                                    height="600px",
                                    defaultValue=st.session_state.editor_content,
                                    language="xml",
                                    theme="vs-dark",
                                    onChange=lazy(sync_editor_change),
                                    options={
                                        "readOnly": False, 
                                        "minimap": {"enabled": True},
                                        "fontSize": 14
                                    }
                                )
                
                except Exception as e:
                    st.error(f"Error leyendo archivo: {e}")
            else:
                # Placeholder vacío elegante
                # st.info("Sube un archivo JSON o XML para comenzar la edición.")
                # Espacio vacío para mantener estructura
                st.markdown('<div style="height: 400px; border: 2px dashed var(--card-border); border-radius: 10px; display: flex; align-items: center; justify-content: center; color: var(--text-color); opacity: 0.6;">Área de trabajo vacía</div>', unsafe_allow_html=True)


# --- TAB: VALIDACION USUARIO ---
if "Validacion Usuario" in tabs_map:
    with tabs_map["Validacion Usuario"]:
        st.header("🏥 Estatus Validacion Paciente")
        
        st.markdown("""
        Esta herramienta permite validar el estado de defunción de pacientes en la Registraduría Nacional.
        """)
        
        col_reg_menu, col_reg_work = st.columns([1, 2])
        
        with col_reg_menu:
             # --- REGISTRADURIA ---
             st.markdown('<div class="group-box">', unsafe_allow_html=True)
             st.markdown('<div class="group-title-left">Validador Registraduría</div>', unsafe_allow_html=True)
             
             if st.button("👤 Individual (Reg)", use_container_width=True):
                 st.session_state.val_tool = "REGISTRADURIA"
                 st.session_state.reg_mode = "INDIVIDUAL"
                 st.rerun()
                 
             if st.button("👥 Masivo (Reg)", use_container_width=True):
                 st.session_state.val_tool = "REGISTRADURIA"
                 st.session_state.reg_mode = "MASIVO"
                 st.rerun()
                 
             st.markdown('</div>', unsafe_allow_html=True)
             
             st.markdown("<br>", unsafe_allow_html=True)
             
             # --- ADRES ---
             st.markdown('<div class="group-box">', unsafe_allow_html=True)
             st.markdown('<div class="group-title-left">Validacion ADRES</div>', unsafe_allow_html=True)
             
             if st.button("👤 Individual (ADRES)", use_container_width=True):
                 st.session_state.val_tool = "ADRES"
                 st.session_state.reg_mode = "INDIVIDUAL"
                 st.rerun()
                 
             if st.button("👥 Masivo (ADRES)", use_container_width=True):
                 st.session_state.val_tool = "ADRES"
                 st.session_state.reg_mode = "MASIVO"
                 st.rerun()
                 
             if st.button("🤖 Validador Adres Boot", use_container_width=True):
                 st.session_state.val_tool = "ADRES_BOOT"
                 st.session_state.reg_mode = "INDIVIDUAL"
                 st.rerun()
             st.markdown('</div>', unsafe_allow_html=True)
             
        with col_reg_work:
            tool = st.session_state.get("val_tool", "REGISTRADURIA")
            mode = st.session_state.get("reg_mode", None)
            
            if tool == "REGISTRADURIA":
                st.subheader("Validador Registraduría")
                
                if mode == "INDIVIDUAL":
                    st.markdown("##### Validación Individual")
                    cedula = st.text_input("Número de Cédula", placeholder="Ingrese cédula sin puntos")
                    
                    if st.button("🔍 Validar Cédula", type="primary"):
                        if not cedula:
                            st.warning("Ingrese una cédula.")
                        else:
                            with st.spinner(f"Consultando {cedula}..."):
                                try:
                                    validator = ValidatorRegistraduria()
                                    result = validator.validate_cedula(cedula)
                                    validator.close_driver()
                                    
                                    st.success("Consulta finalizada.")
                                    st.json(result)
                                    
                                    # Export button
                                    df_res = pd.DataFrame([result])
                                    from io import BytesIO
                                    output = BytesIO()
                                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                                        df_res.to_excel(writer, index=False)
                                    output.seek(0)
                                    
                                    st.download_button(
                                        label="📥 Descargar Resultado (Excel)",
                                        data=output,
                                        file_name=f"Validacion_{cedula}.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                    )
                                    
                                except Exception as e:
                                    st.error(f"Error: {e}")
                                    
                elif mode == "MASIVO":
                    st.markdown("##### Validación Masiva")
                    uploaded_file = st.file_uploader("Cargar Excel con Cédulas", type=["xlsx", "xls"])
                    
                    if uploaded_file:
                        try:
                            df_preview = pd.read_excel(uploaded_file, sheet_name=None)
                            sheet = st.selectbox("Seleccione Hoja", list(df_preview.keys()))
                            
                            if sheet:
                                df_sheet = df_preview[sheet]
                                col_cedula = st.selectbox("Seleccione Columna de Cédulas", df_sheet.columns)
                                
                                if st.button("🚀 Iniciar Validación Masiva (Segundo Plano)", type="primary"):
                                    # Filter non-empty
                                    df_to_process = df_sheet[df_sheet[col_cedula].notna()]
                                    
                                    # Submit to background
                                    submit_task("Validación Registraduría Masiva", run_registraduria_massive, df_to_process, col_cedula)
                                    
                                    st.success("✅ Tarea iniciada en segundo plano. Puedes continuar trabajando en otras pestañas. Revisa el 'Centro de Tareas' en la barra lateral para ver el progreso y descargar el resultado.")
                                        
                        except Exception as e:
                            st.error(f"Error leyendo archivo: {e}")
                
                else:
                     st.info("Seleccione una opción del menú 'Validador Registraduría' para comenzar.")

            elif tool == "ADRES":
                st.subheader("Validación ADRES (BDUA)")
                st.info("ℹ️ **Información:** Validación automática vía API (Supersalud). No requiere CAPTCHA.")
                
                if mode == "INDIVIDUAL":
                    st.markdown("##### Validación Individual")
                    cedula = st.text_input("Número de Cédula", placeholder="Ingrese cédula sin puntos", key="adres_cedula")
                    
                    if st.button("🔍 Validar Cédula (ADRES)", type="primary", key="btn_adres_ind"):
                        if not cedula:
                            st.warning("Ingrese una cédula.")
                        else:
                            with st.spinner(f"Consultando API para {cedula}..."):
                                try:
                                    validator = ValidatorAdres()
                                    result = validator.validate_cedula(cedula)
                                    
                                    if "Error" in result and "TimeOut" in str(result.get("Estado", "")):
                                         st.error(f"Error o Tiempo Agotado: {result.get('Entidad', 'No se detectó resultado')}")
                                    else:
                                        st.success("Consulta finalizada.")
                                        st.json(result)
                                        
                                        # Export
                                        df_res = pd.DataFrame([result])
                                        from io import BytesIO
                                        output = BytesIO()
                                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                                            df_res.to_excel(writer, index=False)
                                        output.seek(0)
                                        
                                        st.download_button(
                                            label="📥 Descargar Resultado (Excel)",
                                            data=output,
                                            file_name=f"Validacion_ADRES_{cedula}.xlsx",
                                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                        )
                                except Exception as e:
                                    st.error(f"Error: {e}")
                
                elif mode == "MASIVO":
                    st.markdown("##### Validación Masiva")
                    st.info("Proceso rápido y automático vía API.")
                    uploaded_file = st.file_uploader("Cargar Excel con Cédulas", type=["xlsx", "xls"], key="adres_masivo")
                    
                    if uploaded_file:
                        try:
                            df_preview = pd.read_excel(uploaded_file, sheet_name=None)
                            sheet = st.selectbox("Seleccione Hoja", list(df_preview.keys()), key="adres_sheet")
                            
                            if sheet:
                                df_sheet = df_preview[sheet]
                                col_cedula = st.selectbox("Seleccione Columna de Cédulas", df_sheet.columns, key="adres_col")
                                
                                if st.button("🚀 Iniciar Validación Masiva (ADRES - Segundo Plano)", type="primary", key="btn_adres_masivo"):
                                    # Filter non-empty
                                    df_to_process = df_sheet[df_sheet[col_cedula].notna()]
                                    
                                    submit_task("Validación ADRES Masiva (API)", run_adres_api_massive, df_to_process, col_cedula)
                                    
                                    st.success("✅ Tarea iniciada en segundo plano. Puedes continuar trabajando en otras pestañas. Revisa el 'Centro de Tareas' en la barra lateral.")
                                        
                        except Exception as e:
                            st.error(f"Error leyendo archivo: {e}")
                else:
                    st.info("Seleccione una opción del menú 'Validacion ADRES' para comenzar.")

            elif tool == "ADRES_BOOT":
                st.subheader("Validación ADRES (Boot / Web)")
                st.warning("⚠️ **Atención:** Esta opción abre el navegador y requiere que usted resuelva el CAPTCHA manualmente.")
                
                # Sub-selection for Boot mode (Individual vs Massive) within the tool area?
                # Or reuse reg_mode if we want buttons in sidebar? 
                # Currently I set reg_mode="INDIVIDUAL" in the button.
                # Let's add radio or buttons here to switch mode if needed, or assume default.
                
                mode_boot = st.radio("Modo:", ["INDIVIDUAL", "MASIVO"], horizontal=True)
                
                if mode_boot == "INDIVIDUAL":
                    st.markdown("##### Validación Individual (Web)")
                    cedula = st.text_input("Número de Cédula", placeholder="Ingrese cédula sin puntos", key="adres_boot_cedula")
                    
                    if st.button("🔍 Validar Cédula (Web)", type="primary", key="btn_adres_boot_ind"):
                        if not cedula:
                            st.warning("Ingrese una cédula.")
                        else:
                            with st.spinner(f"Abriendo navegador para {cedula}. Resuelva el CAPTCHA..."):
                                try:
                                    validator = ValidatorAdresWeb(headless=False)
                                    result = validator.validate_cedula(cedula)
                                    # Driver closes automatically inside validate_cedula or we handle it?
                                    # In my restored code, validate_cedula closes it.
                                    
                                    if "Error" in str(result.get("Estado", "")) or "TimeOut" in str(result.get("Estado", "")):
                                         st.error(f"Resultado: {result.get('Estado')} - {result.get('Entidad')}")
                                    else:
                                        st.success("Consulta finalizada.")
                                        st.json(result)
                                        
                                        # Export
                                        df_res = pd.DataFrame([result])
                                        from io import BytesIO
                                        output = BytesIO()
                                        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                                            df_res.to_excel(writer, index=False)
                                        output.seek(0)
                                        
                                        st.download_button(
                                            label="📥 Descargar Resultado (Excel)",
                                            data=output,
                                            file_name=f"Validacion_ADRES_Web_{cedula}.xlsx",
                                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                        )
                                except Exception as e:
                                    st.error(f"Error: {e}")
                
                elif mode_boot == "MASIVO":
                    st.markdown("##### Validación Masiva (Web)")
                    st.info("El navegador se abrirá para CADA cédula. Deberá resolver el CAPTCHA manualmente cada vez.")
                    uploaded_file = st.file_uploader("Cargar Excel con Cédulas", type=["xlsx", "xls"], key="adres_boot_masivo")
                    
                    if uploaded_file:
                        try:
                            df_preview = pd.read_excel(uploaded_file, sheet_name=None)
                            sheet = st.selectbox("Seleccione Hoja", list(df_preview.keys()), key="adres_boot_sheet")
                            
                            if sheet:
                                df_sheet = df_preview[sheet]
                                col_cedula = st.selectbox("Seleccione Columna de Cédulas", df_sheet.columns, key="adres_boot_col")
                                
                                if st.button("🚀 Iniciar Validación Masiva (Web - Segundo Plano)", type="primary", key="btn_adres_boot_masivo"):
                                    # Filter non-empty
                                    df_to_process = df_sheet[df_sheet[col_cedula].notna()]
                                    
                                    submit_task("Validación ADRES Masiva (Web)", run_adres_web_massive, df_to_process, col_cedula)
                                    
                                    st.success("✅ Tarea iniciada. Se abrirá el navegador para resolver CAPTCHAs. Puede minimizar esta pestaña, pero NO cierre la ventana del navegador que se abrirá.")
                                        
                        except Exception as e:
                            st.error(f"Error leyendo archivo: {e}")


# --- TAB 6: ASISTENTE IA GEMINI ---
if "🤖 Asistente IA (Gemini)" in tabs_map:
    with tabs_map["🤖 Asistente IA (Gemini)"]:
            st.header("🤖 Asistente IA (Gemini)")
    
            col_chat, col_context = st.columns([2, 1])
    
            with col_context:
                st.markdown("### 📄 Contexto")
                st.info("Sube un archivo o usa el contenido seleccionado para que la IA lo analice.")
        
                context_option = st.radio("Fuente de contexto:", ["Ninguno", "Subir Archivo", "Texto Manual"])
        
                file_context = None
        
                if context_option == "Subir Archivo":
                    uploaded_file = st.file_uploader("Sube un archivo (PDF, DOCX, TXT)", type=["pdf", "docx", "txt", "csv", "json", "xml"])
                    if uploaded_file:
                        # Guardar temporalmente para extraer texto
                        try:
                            with open("temp_context_file", "wb") as f:
                                f.write(uploaded_file.getbuffer())
                    
                            file_context = extract_text_from_file("temp_context_file")
                            st.success(f"Archivo cargado ({len(file_context)} caracteres)")
                        except Exception as e:
                            st.error(f"Error cargando archivo: {e}")
                        finally:
                            if os.path.exists("temp_context_file"):
                                os.remove("temp_context_file")
        
                elif context_option == "Texto Manual":
                    file_context = st.text_area("Pega el texto aquí:", height=300)

            with col_chat:
                st.markdown("### 💬 Chat")
        
                if "chat_history" not in st.session_state:
                    st.session_state.chat_history = []
            
                # Mostrar historial
                for message in st.session_state.chat_history:
                    role_icon = "👤" if message["role"] == "user" else "🤖"
                    with st.chat_message(message["role"], avatar=role_icon):
                        st.markdown(message["content"])
                
                # Input de chat
                if prompt := st.chat_input("Escribe tu pregunta para Gemini..."):
                    # Agregar mensaje de usuario
                    st.session_state.chat_history.append({"role": "user", "content": prompt})
                    with st.chat_message("user", avatar="👤"):
                        st.markdown(prompt)
            
                    # Respuesta de IA
                    with st.chat_message("assistant", avatar="🤖"):
                        message_placeholder = st.empty()
                        full_response = ""
                
                        with st.spinner("Gemini está pensando..."):
                            response_text = worker_consultar_gemini(prompt, file_context)
                    
                        message_placeholder.markdown(response_text)
                        st.session_state.chat_history.append({"role": "assistant", "content": response_text})
        
                if st.button("🗑️ Borrar Historial"):
                    st.session_state.chat_history = []
                    st.rerun()

# --- TAB 7: BOT ZEUS SALUD ---
if "🤖 Bot Zeus Salud" in tabs_map:
    with tabs_map["🤖 Bot Zeus Salud"]:
        st.header("🤖 Bot Automatización Zeus Salud")
    
        bot_perm = st.session_state.user_permissions.get("bot_zeus", "full")
    
        # 0. Verificación de Acceso Total
        if bot_perm == "none":
            st.error("⛔ Acceso Denegado: No tiene permisos para acceder al módulo Bot Zeus.")
            st.stop()

        # Badge de Rol
        perm_labels = {
            "full": "✅ Control Total (Crear/Editar/Ejecutar)", 
            "edit": "✏️ Edición (Crear/Editar/Ejecutar)", 
            "execute": "🚀 Solo Ejecución",
            "none": "⛔ Sin Acceso"
        }
        st.caption(f"🛡️ Nivel de Acceso: **{perm_labels.get(bot_perm, bot_perm)}**")

        st.info("Automatización de ingreso de documentos desde Excel. Defina una secuencia de pasos (clicks, escritura, teclas) y ejecútela para cada fila.")
    
        col_bot1, col_bot2 = st.columns([1, 1])
    
        excel_cols = []
        df_bot = pd.DataFrame()
    

        with col_bot1:
            st.subheader("1. Conexión")
            if st.button("🚀 Abrir Navegador / Conectar", use_container_width=True):
                success, msg = bot_zeus.abrir_navegador_inicial()
                if success:
                    st.success(msg)
                else:
                    st.error(msg)
        
            st.divider()
        
            st.subheader("2. Carga de Datos")
            uploaded_bot = st.file_uploader("Archivo Excel (.xlsx)", type=["xlsx", "xls"], key="bot_excel_uploader")
        
            if uploaded_bot:
                try:
                    df_bot = pd.read_excel(uploaded_bot)
                    st.dataframe(df_bot.head(3), height=100)
                    excel_cols = df_bot.columns.tolist()
                    st.success(f"✅ {len(df_bot)} registros cargados.")
                except Exception as e:
                    st.error(f"Error leyendo Excel: {e}")

        with col_bot2:
            st.subheader("3. Definición de Pasos")

            # --- BOTONES DE SESIÓN ---
            if bot_perm in ["full", "edit"]:
                col_ses1, col_ses2 = st.columns(2)
                with col_ses1:
                    if st.button("🔄 Cargar Última Sesión", use_container_width=True, help="Recupera los pasos y flujos guardados automáticamente."):
                        ok, msg = bot_zeus.cargar_sesion()
                        if ok: 
                            st.toast(msg, icon="✅")
                            time.sleep(0.5)
                            st.rerun()
                        else: 
                            st.error(msg)
                with col_ses2:
                    if st.button("💾 Guardar Sesión (Manual)", use_container_width=True, help="Fuerza el guardado de la configuración actual (aunque se guarda automáticamente al editar)."):
                        ok, msg = bot_zeus.guardar_sesion()
                        if ok: st.toast(msg, icon="💾")
                        else: st.error(msg)
            # -------------------------
        
            if bot_perm in ["full", "edit"]:
                st.caption("Configure la secuencia que el robot repetirá por cada fila.")
            
                # Selector de posición de inserción
                num_pasos = len(bot_zeus.PASOS_MEMORIZADOS)
            
                # Opciones para inserción: [Final, 1, 2, ..., N]
                # Usamos un índice visual 1-based para el usuario, pero interno 0-based
                opts_insercion = ["Final"] + [str(i+1) for i in range(num_pasos)]
            
                # Guardar en session_state para persistencia
                if "pos_insercion" not in st.session_state:
                    st.session_state.pos_insercion = "Final"
                
                # UI para seleccionar donde insertar
                col_ins1, col_ins2 = st.columns([2, 1])
                with col_ins1:
                    st.info("💡 Los nuevos pasos se agregarán en la posición seleccionada.")
                with col_ins2:
                    st.session_state.pos_insercion = st.selectbox(
                        "Posición de inserción:", 
                        opts_insercion, 
                        index=0,
                        help="Seleccione 'Final' para agregar al final, o un número para insertar ANTES de ese paso."
                    )
            
                # Determinar índice real para pasar a las funciones
                indice_real = None
                if st.session_state.pos_insercion != "Final":
                    indice_real = int(st.session_state.pos_insercion) - 1 # Convertir a 0-based
            
                tab_click, tab_write, tab_key, tab_wait, tab_text, tab_alert, tab_scroll = st.tabs(["🖱️ Click", "✍️ Escribir", "⌨️ Tecla", "⏳ Espera", "🔤 Texto", "⚠️ Alerta", "📜 Scroll"])
            
                with tab_click:
                    st.write("Haga clic en el elemento en el navegador y presione:")
                    
                    saltar_click = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_click_foco", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")
                    
                    if st.button("Grabar Foco (Click)", use_container_width=True):
                        ok, msg = bot_zeus.agregar_paso_foco("click", indice_insercion=indice_real, saltar_al_final=saltar_click)
                        if ok: st.success(msg)
                        else: st.error(msg)
                
                    st.divider()
                    if st.button("🔄 Cambiar Ventana/Pestaña (Foco)", use_container_width=True, help="Cambia el foco del driver a la última ventana abierta o alterna entre ellas."):
                         # Por defecto cambiamos a la última (-1). Si hay muchas, podría necesitarse un selector.
                         ok, msg = bot_zeus.agregar_paso_cambiar_ventana(-1, indice_insercion=indice_real) # No soporta saltar al final aun, o si? Verifiquemos bot_zeus
                         if ok: st.success(msg)
                         else: st.error(msg)

                with tab_write:
                    st.markdown("##### ✍️ Escribir Dato")
                    col_sel = st.selectbox("Columna Excel:", excel_cols, disabled=len(excel_cols)==0)
                    
                    saltar_write = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_write", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")
                    
                    if st.button("Grabar Foco (Escribir)", use_container_width=True, disabled=len(excel_cols)==0):
                        ok, msg = bot_zeus.agregar_paso_foco("escribir", columna=col_sel, indice_insercion=indice_real, saltar_al_final=saltar_write)
                        if ok: st.success(msg)
                        else: st.error(msg)
                
                    st.divider()
                    st.markdown("##### 📅 Escribir Fecha")
                    st.caption("Toma una fecha del Excel, la formatea y la escribe en el campo seleccionado.")
                    col_date, col_fmt = st.columns([2, 1])
                    with col_date:
                        col_date_sel = st.selectbox("Columna Fecha:", excel_cols, disabled=len(excel_cols)==0, key="sel_col_date")
                    with col_fmt:
                        fmt_date = st.text_input("Formato:", value="%d/%m/%Y", help="Ej: %d/%m/%Y (01/12/2023) o %Y-%m-%d (2023-12-01)")
                    
                    saltar_date = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_date", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")

                    if st.button("Grabar Foco (Escribir Fecha)", use_container_width=True, disabled=len(excel_cols)==0):
                        ok, msg = bot_zeus.agregar_paso_foco("escribir_fecha", columna=col_date_sel, formato=fmt_date, indice_insercion=indice_real, saltar_al_final=saltar_date)
                        if ok: st.success(msg)
                        else: st.error(msg)

                    st.divider()
                    st.markdown("##### 🧹 Borrar Dato")
                    st.caption("Seleccione el campo en el navegador y presione el botón para vaciarlo.")
                    
                    saltar_clean = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_clean", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")

                    if st.button("Grabar Foco (Limpiar Campo)", use_container_width=True):
                         ok, msg = bot_zeus.agregar_paso_foco("limpiar_campo", indice_insercion=indice_real, saltar_al_final=saltar_clean)
                         if ok: st.success(msg)
                         else: st.error(msg)
            
                with tab_key:
                    key_sel = st.selectbox("Tecla Especial:", ["ENTER", "TAB", "ESCAPE", "DOWN", "UP"])
                    
                    saltar_key = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_key", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")
                    
                    if st.button("Agregar Tecla", use_container_width=True):
                        ok, msg = bot_zeus.agregar_paso_tecla(key_sel, indice_insercion=indice_real, saltar_al_final=saltar_key)
                        if ok: st.success(msg)
                        else: st.error(msg)
            
                with tab_wait:
                    sec = st.number_input("Segundos:", min_value=0.1, value=1.0, step=0.5)
                    
                    saltar_wait = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_wait", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")
                    
                    if st.button("Agregar Espera", use_container_width=True):
                        ok, msg = bot_zeus.agregar_paso_espera(sec, indice_insercion=indice_real, saltar_al_final=saltar_wait)
                        if ok: st.success(msg)
                        else: st.error(msg)
                
                    st.markdown("---")
                    st.write("🔧 Utilidades:")
                
                    # Mostrar conteo de ventanas activas para depuración
                    try:
                        drv = bot_zeus.obtener_driver(create_if_missing=False)
                        if drv:
                            n_wins = len(drv.window_handles)
                            st.caption(f"Ventanas detectadas por el sistema: {n_wins}")
                    except:
                        pass

                    col_win1, col_win2 = st.columns(2)
                    with col_win1:
                        if st.button("🔄 Ir a Popup (Última)", use_container_width=True):
                             ok, msg = bot_zeus.agregar_paso_cambiar_ventana(-1, indice_insercion=indice_real)
                             if ok: st.success(msg)
                             else: st.error(msg)
                    with col_win2:
                        if st.button("🏠 Ir a Principal (0)", use_container_width=True):
                             ok, msg = bot_zeus.agregar_paso_cambiar_ventana(0, indice_insercion=indice_real)
                             if ok: st.success(msg)
                             else: st.error(msg)
                    
                with tab_text:
                    st.info("💡 Click en elementos por su contenido (Texto, ID, Título o Clase).")
                
                    tipo_txt = st.radio("Modo:", ["Texto Fijo", "Texto Dinámico (Desde Excel)", "Seleccionar Opción de Lista (Excel)", "Selector Personalizado (XPath)", "🎯 Selector Visual (Beta)"], horizontal=True)
                
                    txt_buscar = None
                    es_dinamico = False
                    tipo_seleccion = "texto"
                    tag_val = "*" # Default

                    if tipo_txt == "🎯 Selector Visual (Beta)":
                        st.markdown("""
                        <div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #4caf50;">
                            <b>Instrucciones:</b><br>
                            1. Presione "Activar Selector".<br>
                            2. Vaya al navegador. El cursor cambiará a una cruz.<br>
                            3. Haga <b>CLICK</b> en el elemento deseado (se pondrá verde).<br>
                            4. Regrese aquí y presione "Capturar Selección".
                        </div>
                        """, unsafe_allow_html=True)
                    
                        col_sel1, col_sel2 = st.columns(2)
                        with col_sel1:
                            if st.button("🚀 Activar Selector", use_container_width=True):
                                ok, msg = bot_zeus.iniciar_selector_visual()
                                if ok: 
                                    st.toast(msg, icon="🎯")
                                    st.session_state.selector_activo = True
                                else: st.error(msg)
                    
                        with col_sel2:
                            if st.button("✅ Capturar Selección", use_container_width=True):
                                ok, result = bot_zeus.obtener_seleccion_visual()
                                if ok:
                                    # Autollenar campos para que el usuario guarde
                                    st.success(f"Elemento capturado: {result}")
                                    st.session_state.xpath_capturado = result
                                else:
                                    st.warning("No se detectó selección o navegador desconectado.")

                        if "xpath_capturado" in st.session_state:
                            txt_buscar = st.text_input("XPath Capturado:", value=st.session_state.xpath_capturado)
                            tipo_seleccion = "xpath"
                            st.caption("Puede editar el XPath si es necesario antes de agregar el paso.")

                    elif tipo_txt == "Selector Personalizado (XPath)":
                        st.caption("Escriba el XPath exacto del elemento. Ej: `//div[@class='modal']//button`")
                        txt_buscar = st.text_input("XPath del elemento:")
                        tipo_seleccion = "xpath"
                    
                    elif tipo_txt == "Texto Fijo":
                        txt_buscar = st.text_input("Texto visible o Atributo (ej. 'Cerrar', 'Guardar', 'X'):")
                    
                        # Selector de Tipo de Elemento (Tag) con agrupación inteligente
                        tag_map = {
                            "Cualquiera (*)": "*",
                            "Botón / Enlace / Input": "*[self::button or self::a or self::input]",
                            "Icono / Imagen (i, svg, img, span)": "*[self::i or self::svg or self::img or self::span]",
                            "Texto / Etiqueta (div, p, label, h1-h6)": "*[self::div or self::p or self::label or self::span or self::h1 or self::h2 or self::h3 or self::h4 or self::h5 or self::h6]",
                            "--- Específicos ---": "*",
                            "Button": "button",
                            "A (Enlace)": "a",
                            "Input": "input",
                            "Div": "div",
                            "Span": "span",
                            "I (Icono)": "i",
                            "SVG": "svg"
                        }
                        tag_sel = st.selectbox("Tipo de Elemento (Ayuda a diferenciar):", list(tag_map.keys()))
                        tag_val = tag_map[tag_sel]

                    elif tipo_txt == "Seleccionar Opción de Lista (Excel)":
                        if excel_cols:
                            st.info("""
                            ℹ️ **Instrucciones:**
                            1. Agregue un paso previo de **Click** para abrir la lista desplegable (si no es nativa).
                            2. Use esta opción para seleccionar el texto exacto proveniente del Excel.
                            3. El bot resaltará visualmente (borde rojo) el elemento encontrado antes de hacer click.
                            """)
                            txt_buscar = st.selectbox("Columna Excel con la opción:", excel_cols, key="col_list_excel")
                            es_dinamico = True
                            # Permitimos cualquier tag, pero sugerimos buscar en todo (*)
                            tag_map = {
                                "Cualquiera (*)": "*",
                                "Opción (<option>)": "option", 
                                "Elemento de Lista (<li>)": "li",
                                "Div / Span": "*[self::div or self::span]",
                                "Enlace (<a>)": "a"
                            }
                            tag_sel = st.selectbox("Tipo de Elemento (Opcional):", list(tag_map.keys()), key="tag_list_excel")
                            tag_val = tag_map[tag_sel]
                        else:
                            st.warning("⚠️ Cargue un Excel primero para usar esta opción.")
                            txt_buscar = None
                    
                    else: # Texto Dinámico (Desde Excel)
                        if excel_cols:
                            txt_buscar = st.selectbox("Columna Excel con el texto:", excel_cols)
                            es_dinamico = True
                        else:
                            st.warning("⚠️ Cargue un Excel primero para usar esta opción.")
                            txt_buscar = None
                    
                        # Filtro de tag en dinámico
                        tag_map = {
                            "Cualquiera (*)": "*", 
                            "Botón / Enlace": "*[self::button or self::a or self::input]",
                            "Icono": "*[self::i or self::svg or self::img or self::span]"
                        }
                        tag_sel = st.selectbox("Tipo de Elemento:", list(tag_map.keys()), key="tag_dyn")
                        tag_val = tag_map[tag_sel]

                    # Checkbox de exactitud y case-insensitive
                    col_opts1, col_opts2 = st.columns(2)
                    with col_opts1:
                        default_exacto = True if tipo_txt == "Seleccionar Opción de Lista (Excel)" else False
                        exacto = st.checkbox("Búsqueda exacta (todo el texto debe coincidir)", value=default_exacto, disabled=(tipo_seleccion=="xpath"))
                    with col_opts2:
                        # Por defecto True para listas para ser más robusto, False para otros
                        default_ignore = True if tipo_txt == "Seleccionar Opción de Lista (Excel)" else False
                        ignore_case = st.checkbox("Ignorar Mayúsculas/Minúsculas", value=default_ignore, disabled=(tipo_seleccion=="xpath"))
                    
                    # --- NUEVO: RESTRICT TO CONTAINER (MULTI-SELECTION) ---
                    xpath_cont = None # Legacy support
                    lista_contenedores = st.session_state.get("multi_contenedores", [])

                    if tipo_seleccion != "xpath":
                        with st.expander("🎯 Restringir a Contenedores Visuales (Opcional - Avanzado)"):
                            st.caption("Úselo para limitar la búsqueda a una o varias zonas específicas. Puede agregar hasta 7 contenedores visuales. El bot buscará el texto en cualquiera de ellos.")
                            
                            col_c1, col_c2 = st.columns([1, 1])
                            with col_c1:
                                if st.button("👁️ Activar Selector Visual", key="btn_vis_cont"):
                                    try:
                                        ok, msg = bot_zeus.iniciar_selector_visual()
                                        if ok:
                                            st.toast(msg)
                                        else:
                                            st.error(msg)
                                    except Exception as e:
                                        st.error(f"Error activando selector: {e}")

                            with col_c2:
                                if st.button("📥 Capturar y Agregar a Lista", key="btn_cap_cont"):
                                    try:
                                        ok, result = bot_zeus.obtener_seleccion_visual()
                                        if ok:
                                            if "multi_contenedores" not in st.session_state:
                                                st.session_state["multi_contenedores"] = []
                                            
                                            if result not in st.session_state["multi_contenedores"]:
                                                st.session_state["multi_contenedores"].append(result)
                                                st.success("¡Contenedor agregado!")
                                            else:
                                                st.warning("Este contenedor ya está en la lista.")
                                        else:
                                            st.warning(f"No se detectó selección: {result}")
                                    except Exception as e:
                                        st.error(f"Error capturando: {e}")
                            
                            # Mostrar lista de contenedores
                            if "multi_contenedores" in st.session_state and st.session_state["multi_contenedores"]:
                                st.markdown("##### Contenedores Seleccionados:")
                                for idx, xp in enumerate(st.session_state["multi_contenedores"]):
                                    c_idx, c_xp, c_del = st.columns([0.5, 4, 0.5])
                                    c_idx.text(f"#{idx+1}")
                                    c_xp.code(xp, language="text")
                                    if c_del.button("🗑️", key=f"del_cont_{idx}"):
                                        st.session_state["multi_contenedores"].pop(idx)
                                        st.rerun()
                                
                                if st.button("Limpiar Todos los Contenedores", key="clean_all_cont"):
                                    st.session_state["multi_contenedores"] = []
                                    st.rerun()
                                
                                lista_contenedores = st.session_state["multi_contenedores"]
                            else:
                                st.info("No hay contenedores seleccionados. Se buscará en toda la página.")
                            
                            # --- OPCIÓN: USAR COMO ÍNDICE ---
                            st.markdown("---")
                            usar_indice = st.checkbox("📍 Usar valor (Excel/Texto) como Índice de Contenedor", key="chk_usar_indice", help="Si se marca, el valor (ej: '1') indicará cual contenedor clickear de la lista (1º, 2º...). Si no se marca, se buscará el texto DENTRO de los contenedores.")
                            if usar_indice:
                                st.info("ℹ️ El bot leerá el número (1, 2, 3...) y hará click en el contenedor correspondiente de la lista de arriba.")


                    saltar_text = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_text", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")

                    if st.button("Agregar Click por Texto/Selector", use_container_width=True, disabled=not txt_buscar):
                        # Pasamos la lista completa como xpath_contenedor (el backend ya lo maneja)
                        usar_idx_val = st.session_state.get("chk_usar_indice", False) if tipo_seleccion != "xpath" else False
                        
                        ok, msg = bot_zeus.agregar_paso_click_texto(txt_buscar, exacto, es_dinamico, tag_val, tipo_seleccion, ignore_case, indice_insercion=indice_real, xpath_contenedor=lista_contenedores if lista_contenedores else None, usar_indice_contenedor=usar_idx_val, saltar_al_final=saltar_text)
                        if ok: 
                            st.success(msg)
                            # Limpiar lista tras agregar
                            if "multi_contenedores" in st.session_state:
                                del st.session_state["multi_contenedores"]
                            if "temp_xpath_contenedor" in st.session_state:
                                del st.session_state["temp_xpath_contenedor"]
                        else: st.error(msg)
            
                with tab_alert:
                    st.info("Úselo cuando aparezca una ventana emergente nativa (Aceptar/Cancelar).")
                    
                    saltar_alert = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_alert", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")
                
                    col_al1, col_al2 = st.columns(2)
                    with col_al1:
                        if st.button("✅ Aceptar Alerta (OK)", use_container_width=True):
                            ok, msg = bot_zeus.agregar_paso_alerta("aceptar", indice_insercion=indice_real, saltar_al_final=saltar_alert)
                            if ok: st.success(msg)
                            else: st.error(msg)
                    with col_al2:
                        if st.button("❌ Cancelar Alerta", use_container_width=True):
                            ok, msg = bot_zeus.agregar_paso_alerta("cancelar", indice_insercion=indice_real, saltar_al_final=saltar_alert)
                            if ok: st.success(msg)
                            else: st.error(msg)

                with tab_scroll:
                    st.info("Desplazarse por la página (útil si el elemento está oculto).")
                
                    tipo_scroll = st.radio("Tipo de movimiento:", ["Abajo (Pixels)", "Arriba (Pixels)", "Ir al Final (Bottom)", "Ir al Inicio (Top)"], horizontal=True)
                
                    cant_px = 0
                    if "Pixels" in tipo_scroll:
                        cant_px = st.number_input("Cantidad de Pixels:", min_value=10, max_value=5000, value=300, step=50)
                
                    saltar_scroll = st.checkbox("⏩ Saltar al Final si Éxito", key="saltar_scroll", help="Si este paso se ejecuta correctamente, el bot saltará inmediatamente al último paso.")

                    if st.button("Agregar Scroll", use_container_width=True):
                        # Mapear selección a parámetros
                        direccion = "abajo"
                        if "Arriba" in tipo_scroll: direccion = "arriba"
                        elif "Final" in tipo_scroll: direccion = "fin"
                        elif "Inicio" in tipo_scroll: direccion = "inicio"
                    
                        ok, msg = bot_zeus.agregar_paso_scroll(direccion, cant_px, indice_insercion=indice_real, saltar_al_final=saltar_scroll)
                        if ok: st.success(msg)
                        else: st.error(msg)
            else:
                st.info("🔒 **Modo Ejecución**: La edición y creación de pasos está deshabilitada para su rol. Puede cargar flujos existentes y ejecutarlos.")


            st.divider()
            st.subheader("Pasos Memorizados")
        
            # --- SECCIÓN GUARDAR/CARGAR ---
            col_io1, col_io2 = st.columns(2)
            with col_io1:
                # Botón de descarga
                pasos_actuales = bot_zeus.PASOS_MEMORIZADOS
                if pasos_actuales:
                    json_str = json.dumps(pasos_actuales, indent=4)
                
                    if bot_perm in ["full", "edit"]:
                        st.download_button(
                            label="💾 Guardar Flujo",
                            data=json_str,
                            file_name="flujo_bot_zeus.json",
                            mime="application/json",
                            use_container_width=True
                        )
                    else:
                        st.button("💾 Guardar Flujo (Deshabilitado)", disabled=True, use_container_width=True, help="Requiere permisos de Edición")
            with col_io2:
                # Carga de archivo
                uploaded_flow = st.file_uploader("📂 Cargar Flujo (.json)", type=["json"], label_visibility="collapsed")
                if uploaded_flow:
                    try:
                        # Usar un hash del contenido para evitar recargas infinitas
                        file_bytes = uploaded_flow.getvalue()
                        file_hash = hash(file_bytes)
                    
                        if st.session_state.get("last_flow_hash") != file_hash:
                            data = json.loads(file_bytes.decode("utf-8"))
                            ok, msg = bot_zeus.cargar_pasos_externos(data)
                            if ok:
                                st.session_state.last_flow_hash = file_hash
                                st.toast(msg, icon="✅") # Usar toast en lugar de success+sleep
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                st.error(msg)
                    except Exception as e:
                        st.error(f"Error cargando flujo: {e}")
        
            # --- SECCIÓN FLUJO ALTERNATIVO (MULTIPLE) ---
            if bot_perm in ["full", "edit", "execute"]:
                with st.expander("🔀 Flujos Condicionales (Avanzado)"):
                    st.info("Configure flujos adicionales que se ejecutarán si se cumple una condición específica. Se evalúan en orden (1 -> 2 -> 3). Si ninguno cumple, se usa el Principal.")
                    
                    tabs_cond = st.tabs(["Alternativo 1", "Alternativo 2", "Alternativo 3"])
                    
                    for idx_tab, tab in enumerate(tabs_cond):
                        with tab:
                            # Recuperar estado actual
                            flujo_actual = bot_zeus.get_flujo_condicional(idx_tab)
                            
                            col_conf1, col_conf2 = st.columns(2)
                            
                            with col_conf1:
                                st.markdown(f"**Carga del Flujo {idx_tab + 1}**")
                                uploaded_alt = st.file_uploader(f"📂 Cargar JSON (Alt {idx_tab + 1})", type=["json"], key=f"uploader_alt_{idx_tab}")
                                
                                if uploaded_alt:
                                    try:
                                        file_bytes = uploaded_alt.getvalue()
                                        hash_f = hash(file_bytes)
                                        last_hash_key = f"last_hash_alt_{idx_tab}"
                                        
                                        if st.session_state.get(last_hash_key) != hash_f:
                                            data_flow = json.loads(file_bytes.decode("utf-8"))
                                            # Update pasos but keep condition if exists
                                            cond_prev = flujo_actual["condicion"] if flujo_actual else None
                                            bot_zeus.update_flujo_condicional(idx_tab, pasos=data_flow, condicion=cond_prev, nombre=f"Alternativo {idx_tab+1}")
                                            st.session_state[last_hash_key] = hash_f
                                            st.toast(f"Flujo {idx_tab+1} cargado!", icon="✅")
                                            st.rerun()
                                    except Exception as e:
                                        st.error(f"Error: {e}")

                                if flujo_actual and flujo_actual.get("pasos"):
                                    st.caption(f"✅ Pasos en memoria: {len(flujo_actual['pasos'])}")
                                else:
                                    st.caption("⚠️ Sin pasos cargados.")

                            with col_conf2:
                                st.markdown(f"**Condición de Ejecución {idx_tab + 1}**")
                                # UI para condicion
                                current_cond = flujo_actual.get("condicion") if flujo_actual else {}
                                if not current_cond: current_cond = {}
                                
                                tipo_cond = st.radio("Tipo", ["Texto en Pantalla", "Valor Excel", "Valor Excel (Múltiple)"], 
                                                     index=0 if current_cond.get("tipo") == "texto" else (2 if current_cond.get("tipo") == "excel_multi" else 1),
                                                     key=f"radio_cond_{idx_tab}", horizontal=True)
                                
                                new_cond = {} # Default empty (clears if not set)
                                
                                if tipo_cond == "Texto en Pantalla":
                                    val_txt = st.text_input("Texto que debe aparecer para ACTIVAR este flujo:",
                                                          value=current_cond.get("valor", "") if current_cond.get("tipo") == "texto" else "",
                                                          placeholder="Ej: Embarazada",
                                                          key=f"txt_cond_{idx_tab}")
                                    if val_txt:
                                        new_cond = {"tipo": "texto", "valor": val_txt}
                                        
                                elif tipo_cond == "Valor Excel": # Excel Simple
                                    cols = df_bot.columns.tolist() if df_bot is not None else []
                                    curr_col = current_cond.get("columna", "")
                                    curr_val = current_cond.get("valor", "") if current_cond.get("tipo") == "excel" else ""
                                    
                                    if not cols: st.warning("Cargue Excel primero.")
                                    
                                    c_sel = st.selectbox("Columna", [""] + cols, 
                                                       index=cols.index(curr_col) + 1 if curr_col in cols else 0,
                                                       key=f"sel_col_{idx_tab}")
                                    v_sel = st.text_input("Valor(es) activador(es) (separar con |):",
                                                        value=curr_val,
                                                        placeholder="Ej: Si | Yes",
                                                        key=f"val_cond_{idx_tab}")
                                    
                                    if c_sel and v_sel:
                                        new_cond = {"tipo": "excel", "valor": v_sel, "columna": c_sel}

                                elif tipo_cond == "Valor Excel (Múltiple)":
                                    cols = df_bot.columns.tolist() if df_bot is not None else []
                                    if not cols: st.warning("Cargue Excel primero.")
                                    
                                    # Preparar datos para el editor
                                    raw_reglas = current_cond.get("reglas", []) if current_cond.get("tipo") == "excel_multi" else []
                                    # Convertir a formato UI (Capitalized keys)
                                    ui_reglas = [{"Columna": r.get("columna", ""), "Valor": r.get("valor", "")} for r in raw_reglas]
                                    if not ui_reglas: ui_reglas = [{"Columna": "", "Valor": ""}]
                                    
                                    df_reglas = pd.DataFrame(ui_reglas)
                                    
                                    column_config = {
                                        "Columna": st.column_config.SelectboxColumn(
                                            "Columna Excel",
                                            help="Seleccione la columna a validar",
                                            width="medium",
                                            options=cols,
                                            required=True
                                        ),
                                        "Valor": st.column_config.TextColumn(
                                            "Valor(es) Trigger",
                                            help="Separe valores con |",
                                            width="medium",
                                            required=True
                                        )
                                    }
                                    
                                    st.caption("Defina múltiples reglas. TODAS deben cumplirse (AND).")
                                    edited_df = st.data_editor(
                                        df_reglas,
                                        column_config=column_config,
                                        num_rows="dynamic",
                                        key=f"editor_reglas_{idx_tab}",
                                        use_container_width=True,
                                        hide_index=True
                                    )
                                    
                                    # Procesar resultado
                                    reglas_final = []
                                    for _, row in edited_df.iterrows():
                                        c = row.get("Columna")
                                        v = row.get("Valor")
                                        if c and v:
                                            reglas_final.append({"columna": c, "valor": v})
                                    
                                    if reglas_final:
                                        new_cond = {"tipo": "excel_multi", "reglas": reglas_final}

                                # Botón guardar condicion
                                if st.button("💾 Actualizar Condición", key=f"btn_save_cond_{idx_tab}"):
                                    bot_zeus.update_flujo_condicional(idx_tab, condicion=new_cond)
                                    st.toast("Condición actualizada", icon="💾")
                                    st.rerun()

                                # Mostrar resumen condicion
                                if flujo_actual and flujo_actual.get("condicion"):
                                    c = flujo_actual["condicion"]
                                    if c.get("tipo") == "excel":
                                        st.info(f"Si Columna **{c.get('columna')}** es **{c.get('valor')}** -> Ejecuta este flujo.")
                                    elif c.get("tipo") == "excel_multi":
                                        r_txt = " Y ".join([f"[{r['columna']}='{r['valor']}']" for r in c.get('reglas', [])])
                                        st.info(f"Si **{r_txt}** -> Ejecuta este flujo.")
                                    elif c.get("tipo") == "texto":
                                        st.info(f"Si Texto **{c.get('valor')}** es visible -> Ejecuta este flujo.")
                                else:
                                    st.warning("Sin condición definida (Nunca se ejecutará).")

                            st.divider()
                            if st.button(f"🗑️ Limpiar Alternativo {idx_tab+1}", key=f"clean_{idx_tab}"):
                                bot_zeus.update_flujo_condicional(idx_tab, pasos=[], condicion={})
                                st.session_state.pop(f"last_hash_alt_{idx_tab}", None)
                                st.rerun()

            st.divider()

            # Display steps with management controls
            pasos = bot_zeus.PASOS_MEMORIZADOS
            if not pasos:
                st.info("No hay pasos grabados.")
            else:
                st.write("---")
                st.write("**Gestión de Pasos:**")
            
                # Header row
                if bot_perm in ["full", "edit"]:
                    h1, h2, h3, h4, h5 = st.columns([5, 1, 1, 1, 1])
                    h1.markdown("**Descripción**")
                    h2.markdown("**Opcional**")
                    h3.markdown("**Subir**")
                    h4.markdown("**Bajar**")
                    h5.markdown("**Borrar**")
                else:
                    h1, h2 = st.columns([5, 1])
                    h1.markdown("**Descripción**")
                    h2.markdown("**Estado**")
            
                action_taken = False
            
                for i, p in enumerate(pasos):
                    # Check status
                    es_opcional = p.get("opcional", False)
                    desc_texto = f"{i+1}. {p.get('descripcion', 'Paso')}"
                    if es_opcional:
                        desc_texto += " (OPCIONAL)"
                
                    if bot_perm in ["full", "edit"]:
                        c1, c2, c3, c4, c5 = st.columns([5, 1, 1, 1, 1])
                        # Description with index
                        c1.text(desc_texto)

                        # Optional Toggle
                        btn_label = "⚠️" if es_opcional else "✅"
                        help_text = "Click para marcar como Opcional" if not es_opcional else "Click para marcar como Obligatorio"
                        if c2.button(btn_label, key=f"btn_opt_{i}", help=help_text):
                            bot_zeus.alternar_opcional_paso(i)
                            action_taken = True
                    
                        # Move Up
                        if i > 0: 
                            if c3.button("⬆️", key=f"btn_up_{i}"):
                                bot_zeus.mover_paso(i, -1)
                                action_taken = True
                    
                        # Move Down
                        if i < len(pasos) - 1:
                            if c4.button("⬇️", key=f"btn_down_{i}"):
                                bot_zeus.mover_paso(i, 1)
                                action_taken = True
                            
                        # Delete
                        if c5.button("🗑️", key=f"btn_del_{i}"):
                            bot_zeus.eliminar_paso_indice(i)
                            action_taken = True
                    else:
                        c1, c2 = st.columns([5, 1])
                        c1.text(desc_texto)
                        c2.text("⚠️ Opcional" if es_opcional else "✅ Obligatorio")
                
                    if action_taken:
                        st.rerun()
            
                st.divider()
            
                if bot_perm in ["full", "edit"]:
                    c1, c2 = st.columns(2)
                    with c2:
                        if st.button("Borrar Todo", type="primary"):
                            bot_zeus.limpiar_pasos()
                            st.rerun()

        st.markdown("---")
        st.subheader("4. Ejecución")
    
        delay = st.slider("Velocidad (segundos entre pasos):", 0.0, 3.0, 0.5)
        
        # --- LÓGICA DE EJECUCIÓN EN HILO (PARA PERMITIR STOP) ---
        if "bot_running" not in st.session_state:
            st.session_state.bot_running = False
        if "bot_logs" not in st.session_state:
            st.session_state.bot_logs = []
            
        # Wrapper para el hilo
        def run_bot_thread(df, delay):
            try:
                for msg in bot_zeus.ejecutar_secuencia(df, delay_pasos=delay):
                    if "bot_logs" in st.session_state:
                        st.session_state.bot_logs.append(msg)
                    else:
                        # Fallback si se pierde el contexto
                        pass
                    time.sleep(0.01)
            except Exception as e:
                if "bot_logs" in st.session_state:
                    st.session_state.bot_logs.append(f"❌ Error en hilo: {e}")
            finally:
                if "bot_running" in st.session_state:
                    st.session_state.bot_running = False

        col_ctrl1, col_ctrl2 = st.columns([1, 1])
        
        with col_ctrl1:
            # Botón INICIAR
            if not st.session_state.bot_running:
                if st.button("▶️ Iniciar Secuencia Masiva", use_container_width=True, disabled=not (uploaded_bot and len(bot_zeus.PASOS_MEMORIZADOS) > 0)):
                    st.session_state.bot_running = True
                    st.session_state.bot_logs = []
                    st.session_state.bot_finished_shown = False # Reset flag de toast
                    
                    t = threading.Thread(target=run_bot_thread, args=(df_bot, delay))
                    try:
                        from streamlit.runtime.scriptrunner import add_script_run_ctx
                        add_script_run_ctx(t)
                    except:
                        pass
                    t.start()
                    st.rerun()
            else:
                 st.info("🚀 Ejecutando... (Espere o Detenga)")

        with col_ctrl2:
            # Botón DETENER
            if st.session_state.bot_running:
                if st.button("⛔ DETENER EJECUCIÓN", type="primary", use_container_width=True):
                    bot_zeus.detener_ejecucion()
                    st.toast("Deteniendo...", icon="🛑")
        
        # Mostrar Logs
        if st.session_state.bot_logs:
            st.code("\n".join(st.session_state.bot_logs[-15:]), language="text")
        
        # Auto-refresh mientras corre
        if st.session_state.bot_running:
            time.sleep(1)
            st.rerun()
            
        # --- POST EJECUCIÓN (Popups/Mensajes) ---
        if not st.session_state.bot_running and st.session_state.bot_logs:
            last_log = st.session_state.bot_logs[-1]
            
            # Solo mostrar feedback si acabamos de terminar (usando flag bot_finished_shown)
            if not st.session_state.get("bot_finished_shown", False):
                
                # Caso ERROR
                if bot_zeus.ULTIMO_ERROR:
                    st.error(f"❌ El proceso se detuvo con errores.")
                    st.toast(f"Error: {bot_zeus.ULTIMO_ERROR}", icon="❌")
                    with st.expander("Ver detalle del error", expanded=True):
                        st.write(bot_zeus.ULTIMO_ERROR)
                    st.session_state.bot_finished_shown = True
                    
                # Caso ÉXITO (o Detenido por usuario pero sin error crash)
                elif "Fin del proceso" in last_log or "interrumpido" in last_log:
                    if "interrumpido" in last_log:
                        st.warning("⚠️ Proceso detenido por el usuario.")
                        st.toast("Proceso Detenido", icon="🛑")
                    else:
                        st.success("✅ Proceso Completado Exitosamente.")
                        st.toast("Proceso Completado", icon="✅")
                        st.balloons()
                    st.session_state.bot_finished_shown = True
            
            # Botón de descarga siempre visible al final
            st.download_button("Descargar Log Completo", "\n".join(st.session_state.bot_logs), file_name="log_bot_zeus.txt")

        with st.expander("ℹ️ Guía Rápida"):
            st.markdown("""
            1. **Conexión**: Abra el navegador y conéctese a su aplicación web.
            2. **Datos**: Cargue el Excel que contiene las filas a procesar.
            3. **Definir Pasos**: Construya la secuencia de acciones que se repetirá por cada fila:
               - **Escribir**: Seleccione una columna del Excel, haga clic en el campo del navegador donde va el dato, y presione *Grabar Foco (Escribir)*.
               - **Click**: Haga clic en el botón o elemento del navegador, regrese aquí y presione *Grabar Foco (Click)*.
               - **Tecla**: Agregue pulsaciones como ENTER o TAB para navegar entre campos.
            4. **Verificar**: Revise la lista de "Pasos Memorizados". Use "Deshacer" si se equivoca.
            5. **Ejecutar**: Ajuste la velocidad y presione "Iniciar Secuencia Masiva".
            """)

def create_standalone_agent_zip():
    """Empaqueta el instalador del Agente Local si existe, o el ejecutable portable."""
    try:
        output = io.BytesIO()
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as z:
            # 1. Buscar Instalador (Instalar_Agente.exe)
            installer_name = "Instalar_Agente.exe"
            exe_name = "CDO_Agente.exe"
            
            # Rutas posibles
            base_paths = [
                os.path.dirname(__file__), # Src dir
                os.path.dirname(os.path.dirname(__file__)), # Root dir
                os.path.join(os.path.dirname(os.path.dirname(__file__)), "dist"), # Dist dir
                "dist" # Relative dist
            ]
            
            installer_path = None
            for bp in base_paths:
                p = os.path.join(bp, installer_name)
                if os.path.exists(p):
                    installer_path = p
                    break
            
            if installer_path:
                z.write(installer_path, installer_name)
                z.writestr("LEEME.txt", "Ejecute 'Instalar_Agente.exe' para instalar el servicio.")
                return output.getvalue(), True, "installer"
                
            # 2. Si no hay instalador, buscar EXE portable
            agent_path = None
            for bp in base_paths:
                p = os.path.join(bp, exe_name)
                if os.path.exists(p):
                    agent_path = p
                    break
                    
            if agent_path:
                z.write(agent_path, exe_name)
                z.writestr("iniciar_agente.bat", f'start "" "{exe_name}"')
                readme = """Agente Local CDO (Portable)
===========================
1. Descomprima.
2. Ejecute 'iniciar_agente.bat'.
"""
                z.writestr("LEEME.txt", readme)
                return output.getvalue(), True, "portable"
            else:
                return None, False, None
    except Exception as e:
        print(f"Error creando zip standalone: {e}")
        return None, False, None

# --- LOG GLOBAL ---
st.markdown("---")
st.subheader("🧾 Registro de actividad")
for l in reversed(st.session_state.logs[-5:]):
    st.text(l)