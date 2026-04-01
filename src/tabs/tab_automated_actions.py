import streamlit as st
import os
import time
import shutil
import pandas as pd
import json


@st.cache_data(show_spinner=False, max_entries=5)
def _get_excel_sheet_names(file_bytes):
    import pandas as pd
    import io
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    return xls.sheet_names

@st.cache_data(show_spinner=False, max_entries=10)
def _get_excel_preview(file_bytes, sheet_name, nrows=5):
    import pandas as pd
    import io
    return pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name, nrows=nrows)

def close_auto_dialog():
    # Only clear uploaders starting with "up_" or specific known keys to avoid breaking session
    keys_to_clear = [k for k in st.session_state.keys() if k.startswith("up_") or "uploader" in k or k.endswith("_up")]
    for k in keys_to_clear:
        if k in st.session_state:
            del st.session_state[k]
    
    if "active_auto_dialog" in st.session_state:
        del st.session_state["active_auto_dialog"]
    st.rerun()

def _should_delegate(path_or_list):
    import os
    if not path_or_list: return False
    path = path_or_list
    if isinstance(path_or_list, list):
        path = path_or_list[0]
    if isinstance(path, dict):
        path = path.get("Ruta completa", "")
    return not os.path.exists(path)

import re
import random
import math
from datetime import datetime
import zipfile
import io
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata
try:
    import openpyxl
except ImportError:
    openpyxl = None
import requests
import base64
import urllib.parse
import xml.etree.ElementTree as ET
import sys

# Fix import path for agent_client
agent_client_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if agent_client_path not in sys.path:
    sys.path.append(agent_client_path)

try:
    from agent_client import send_command, wait_for_result
except ImportError:
    try:
        from src.agent_client import send_command, wait_for_result
    except ImportError:
        send_command = None
        wait_for_result = None

# --- CONDITIONAL IMPORTS FOR ANALYSIS WORKERS ---
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None
    
# --- IMPORTS & SETUP ---
try:
    from modules.registraduria_validator import ValidatorRegistraduria
    from modules.adres_validator import ValidatorAdres, ValidatorAdresWeb
except ImportError:
    try:
        from src.modules.registraduria_validator import ValidatorRegistraduria
        from src.modules.adres_validator import ValidatorAdres, ValidatorAdresWeb
    except ImportError:
        pass

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service
    from webdriver_manager.chrome import ChromeDriverManager
except ImportError:
    pass

try:
    from gui_utils import abrir_dialogo_carpeta_nativo, update_path_key, render_path_selector, render_file_selector, render_download_button
except ImportError:
    try:
        from src.gui_utils import abrir_dialogo_carpeta_nativo, update_path_key, render_path_selector, render_file_selector, render_download_button
    except ImportError:
        def abrir_dialogo_carpeta_nativo(title="Seleccionar Carpeta", initial_dir=None):
            st.warning("Selector de carpeta nativo no disponible.")
            return None

        def update_path_key(key, new_path, widget_key=None):
            if new_path:
                st.session_state[key] = new_path
                if widget_key:
                    st.session_state[widget_key] = new_path
        
        def render_path_selector(label, key, default_path=None, help_text=None, omit_checkbox=False):
            st.warning("render_path_selector no disponible")
            return default_path

        def render_file_selector(label, key, default_path=None, help_text=None, file_types=None, omit_checkbox=False):
            st.warning("render_file_selector no disponible")
            return default_path

        def render_download_button(folder_path, key, label="📦 Descargar ZIP"):
            st.warning("Descarga no disponible (Error importando gui_utils)")

try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

try:
    from docx2pdf import convert as convert_docx_to_pdf
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

try:
    import pyperclip
except ImportError:
    pyperclip = None

import google.generativeai as genai

# Helper for callback-based folder selection
# update_path_key imported from gui_utils

# --- HELPERS ---

def find_folder_path(base_path, folder_name):
    """
    Intenta encontrar una carpeta usando los resultados de búsqueda en sesión.
    Si no la encuentra, asume que es una subcarpeta directa de base_path.
    """
    target_name = str(folder_name).strip().lower()
    
    # DEBUG: Log attempt
    print(f"DEBUG find_folder_path: Buscando '{target_name}'...")

    # 1. Buscar en resultados de búsqueda (si existen)
    if "search_results" in st.session_state and st.session_state.search_results:
        print(f"DEBUG: Revisando {len(st.session_state.search_results)} resultados en cache.")
        if len(st.session_state.search_results) > 0:
             print(f"DEBUG: Keys del primer item: {st.session_state.search_results[0].keys()}")
        for item in st.session_state.search_results:
            # Normalizar claves (puede venir como 'name' o 'Nombre')
            i_name = str(item.get("name", item.get("Nombre", ""))).strip().lower()
            i_type = str(item.get("type", item.get("Tipo", ""))).strip().lower()
            i_path = item.get("path", item.get("Ruta completa", ""))
            
            # Check type variants
            is_folder = i_type in ["folder", "carpeta", "directory"]
            
            if is_folder and i_name == target_name:
                print(f"DEBUG: Encontrado en cache: {i_path}")
                # En modo nativo (AWS/Agente), confiamos en la ruta del resultado de búsqueda
                # ya que el servidor no puede verificar la existencia en el cliente.
                return i_path

    # 2. Fallback: Subcarpeta directa
    fallback = os.path.join(base_path, str(folder_name).strip()) if base_path else None
    print(f"DEBUG: No encontrado en cache. Usando fallback: {fallback}")
    return fallback



def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]

def clean_df_for_json(df):
    """Limpia un DataFrame para conversión a JSON, manejando NaN y tipos."""
    df.columns = [str(c).strip() for c in df.columns]
    df = df.where(pd.notnull(df), None)
    numeric_fields = [
        "consecutivo", "consecutivo_usuario", "codservicio", "vrservicio", 
        "valorpagomoderador", "copago", "cuotamoderadora", 
        "numfevpagomoderador", "bonificacion", "valortotal", 
        "cantidad", "valorunitario"
    ]
    for col in df.columns:
        if col.lower() in numeric_fields:
            df[col] = pd.to_numeric(df[col], errors='coerce')
    return df

def get_val_ci(data_dict, key):
    if not isinstance(data_dict, dict): return None
    for k, v in data_dict.items():
        if k.lower() == key.lower():
            return v
    return None

def recursive_clean_json(data):
    if isinstance(data, dict):
        return {k: recursive_clean_json(v) for k, v in data.items() if v is not None}
    elif isinstance(data, list):
        return [recursive_clean_json(item) for item in data]
    else:
        return data

def recursive_strip(data):
    if isinstance(data, dict):
        return {k: recursive_strip(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [recursive_strip(i) for i in data]
    elif isinstance(data, str):
        return data.strip()
    return data

def recursive_update_key(data, key, new_val):
    count = 0
    if isinstance(data, dict):
        for k, v in data.items():
            if k == key:
                data[k] = new_val
                count += 1
            elif isinstance(v, (dict, list)):
                count += recursive_update_key(v, key, new_val)
    elif isinstance(data, list):
        for item in data:
            count += recursive_update_key(item, key, new_val)
    return count

def recursive_update_cups(data, old_val, new_val):
    count = 0
    if isinstance(data, dict):
        for k, v in data.items():
            if k == "codServicio" and str(v).strip() == str(old_val).strip():
                data[k] = new_val
                count += 1
            elif isinstance(v, (dict, list)):
                count += recursive_update_cups(v, old_val, new_val)
    elif isinstance(data, list):
        for item in data:
            count += recursive_update_cups(item, old_val, new_val)
    return count

def recursive_update_notes(data, target_text, new_note):
    count = 0
    if isinstance(data, dict):
        for k, v in data.items():
            if isinstance(v, str) and target_text in v:
                data[k] = new_note
                count += 1
            elif isinstance(v, (dict, list)):
                count += recursive_update_notes(v, target_text, new_note)
    elif isinstance(data, list):
        for item in data:
            count += recursive_update_notes(item, target_text, new_note)
    return count

# --- HELPERS: SIGNATURES ---

def _dibujar_trazo_vocal(draw, x_base, y_centro, ascii_val, colores):
    """Dibuja un trazo característico de vocal"""
    color = random.choice(colores)
    grosor = random.randint(2, 3)
    
    # Crear arco o curva
    altura_arco = 20 + (ascii_val % 15)
    puntos = []
    
    for i in range(20):
        angulo = (i / 19.0) * math.pi
        x = x_base + i * 2
        y = y_centro - math.sin(angulo) * altura_arco + random.randint(-2, 2)
        puntos.append((x, y))
    
    # Dibujar curva
    for i in range(len(puntos) - 1):
        draw.line([puntos[i], puntos[i + 1]], fill=color, width=grosor)

def _dibujar_trazo_consonante_dura(draw, x_base, y_centro, ascii_val, colores):
    """Dibuja un trazo característico de consonante dura"""
    color = random.choice(colores)
    grosor = random.randint(3, 4)
    
    # Línea con ángulos y cambios de dirección
    x = x_base
    y = y_centro + random.randint(-10, 10)
    
    # Línea ascendente
    draw.line([(x, y), (x + 15, y - 20)], fill=color, width=grosor)
    # Línea horizontal
    draw.line([(x + 15, y - 20), (x + 30, y - 15)], fill=color, width=grosor)
    # Línea descendente
    draw.line([(x + 30, y - 15), (x + 40, y + 10)], fill=color, width=grosor)

def _dibujar_trazo_generico(draw, x_base, y_centro, ascii_val, colores, width, height):
    """Dibuja un trazo genérico fluido"""
    color = random.choice(colores)
    grosor = random.randint(2, 3)
    
    # Crear línea ondulada
    puntos = []
    for i in range(30):
        x = x_base + i * 1.5
        # Onda basada en el valor ASCII
        onda = math.sin((x - x_base) * 0.2 + ascii_val * 0.1) * 15
        y = y_centro + onda + random.randint(-3, 3)
        puntos.append((x, y))
    
    # Dibujar línea ondulada
    for i in range(len(puntos) - 1):
        draw.line([puntos[i], puntos[i + 1]], fill=color, width=grosor)
    
    # Añadir algunos puntos extra para dar textura
    for _ in range(random.randint(2, 4)):
        punto_x = random.randint(int(x_base), int(x_base + 40))
        punto_y = y_centro + random.randint(-5, 5)
        draw.ellipse([punto_x - 1, punto_y - 1, punto_x + 1, punto_y + 1], fill=color)
    
    # Añadir algunos detalles decorativos adicionales
    # Pequeñas líneas adicionales
    for _ in range(random.randint(1, 3)):
        x_rand = random.randint(20, width - 20)
        y_rand = random.randint(height // 3, 2 * height // 3)
        longitud = random.randint(10, 25)
        angulo = random.uniform(-0.5, 0.5)
        
        x_fin = x_rand + int(longitud * math.cos(angulo))
        y_fin = y_rand + int(longitud * math.sin(angulo))
        
        draw.line([(x_rand, y_rand), (x_fin, y_fin)], fill=random.choice(colores), width=random.randint(1, 2))
    
    # Añadir un toque final con una línea más prominente
    y_linea_principal = height // 2 + random.randint(-10, 10)
    x_inicio_principal = random.randint(25, 45)
    x_fin_principal = width - random.randint(25, 45)
    
    # Crear la línea principal más larga
    puntos_principales = []
    for x in range(x_inicio_principal, x_fin_principal, 3):
        onda = math.sin((x - x_inicio_principal) * 0.08) * 8
        y = y_linea_principal + onda + random.randint(-2, 2)
        puntos_principales.append((x, y))
    
    # Dibujar línea principal
    if len(puntos_principales) > 1:
        for i in range(len(puntos_principales) - 1):
            draw.line([puntos_principales[i], puntos_principales[i + 1]], fill='black', width=3)

def _crear_firma_estilizada(texto):
    """
    Crea una firma digital estilizada sin usar fuentes tipográficas.
    Convierte cada letra del texto en un trazo manuscrito único.
    """
    # Dimensiones de la imagen
    width = max(400, len(texto) * 40)
    height = 150
    
    # Crear imagen base con fondo blanco
    if Image is None: return None
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    
    # Convertir cada letra del texto en un trazo manuscrito
    colores = ['black', 'gray', 'darkgray']
    
    # Procesar cada letra del nombre
    for i, letra in enumerate(texto):
        if letra.isspace():
            continue
            
        # Posición basada en la letra (más espaciado)
        x_base = 30 + (i * (width - 60) // len(texto))
        y_centro = height // 2
        
        # Crear trazo único basado en la letra
        # Usar el código ASCII de la letra para determinar el estilo
        ascii_val = ord(letra.upper()) if letra.isalpha() else ord('A')
        
        # Determinar la forma del trazo basada en la letra
        if letra.upper() in 'AEIOU':
            # Vocales: líneas curvas y abiertas
            _dibujar_trazo_vocal(draw, x_base, y_centro, ascii_val, colores)
        elif letra.upper() in 'BCDFG':
            # Consonantes duras: líneas fuertes y angulares
            _dibujar_trazo_consonante_dura(draw, x_base, y_centro, ascii_val, colores)
        else:
            # Otras letras: líneas fluidas
            _dibujar_trazo_generico(draw, x_base, y_centro, ascii_val, colores, width, height)
    
    return image

# --- WORKERS: ORGANIZATION ---

def worker_mover_por_coincidencia(root_path, silent_mode=False, return_zip=False):
    if not silent_mode: st.info(f"Iniciando movimiento por coincidencia en: {root_path}")
    
    try:
        items = os.listdir(root_path)
    except Exception as e:
        return {"error": f"Error leyendo directorio: {e}"}

    files = [f for f in items if os.path.isfile(os.path.join(root_path, f))]
    folders = [d for d in items if os.path.isdir(os.path.join(root_path, d))]
    
    count_moved = 0
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Analizando...")
    total = len(files)
    
    for i, file in enumerate(files):
        if not silent_mode and i % 10 == 0 and total > 0:
             progress_bar.progress(min(i / total, 1.0), text=f"Procesando {i}/{total}")
             
        file_lower = file.lower()
        target = None
        for folder in folders:
            if folder.lower() in file_lower:
                target = folder
                break
        
        if target:
            src = os.path.join(root_path, file)
            dst = os.path.join(root_path, target, file)
            try:
                shutil.move(src, dst)
                count_moved += 1
            except Exception as e:
                if not silent_mode: st.error(f"Error moviendo {file} a {target}: {e}")
                
    msg = f"Proceso completado. {count_moved} archivos organizados."
    if not silent_mode:
        if progress_bar: progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)

    if return_zip:
        try:
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(root_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, root_path)
                        zf.write(file_path, arcname)
            
            # Cleanup temp folder if we are returning zip
            try: shutil.rmtree(root_path, ignore_errors=True)
            except: pass
            
            return {
                "files": [{
                    "name": f"Organizados_Coincidencia_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Organizados (ZIP)"
                }],
                "message": msg
            }
        except Exception as e:
            return {"error": f"Error creando ZIP: {e}", "message": msg}
            
    return {"message": msg}

def worker_consolidar_subcarpetas(root_path, silent_mode=False, return_zip=False):
    is_native_mode = False
    if not os.environ.get("CDO_AGENT_MODE") == "1":
        try:
            is_native_mode = st.session_state.get('force_native_mode', True)
        except Exception:
            pass

    if is_native_mode:
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "default")
            
            if not silent_mode:
                st.info("Enviando tarea al Agente Local...")
            
            task_id = send_command(username, "consolidar_subcarpetas", {"root_path": root_path})
            if task_id:
                if not silent_mode:
                    with st.spinner("Agente Local procesando (Consolidando)..."):
                        res = wait_for_result(task_id, timeout=300)
                else:
                    res = wait_for_result(task_id, timeout=300)
                
                if isinstance(res, dict):
                    if "error" in res:
                        return {"error": f"Error del Agente: {res['error']}"}
                    return {"message": res.get("message", "Operación completada por el agente.")}
                return {"error": "Respuesta inesperada del agente."}
            else:
                return {"error": "No se pudo conectar con el Agente Local."}
        except ImportError:
            return {"error": "No se pudo importar cliente del agente."}
        except Exception as e:
            return {"error": f"Error comunicando con el agente: {e}"}

    if not silent_mode: st.info(f"Consolidando subcarpetas en: {root_path}")
    
    try:
        main_folders = [d for d in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, d))]
    except Exception as e:
        return {"error": f"Error leyendo directorio base: {e}"}

    if not main_folders: return {"error": "No se encontraron carpetas para procesar."}
    
    copiados = 0
    conflictos = 0
    errores = 0
    
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Consolidando...")
    total = len(main_folders)
    
    for i, folder_name in enumerate(main_folders):
        if not silent_mode and total > 0:
             progress_bar.progress(min(i / total, 1.0), text=f"Procesando carpeta {folder_name}")
        
        main_folder_path = os.path.join(root_path, folder_name)
        
        for sub_root, _, files in os.walk(main_folder_path):
            if sub_root == main_folder_path:
                continue
                
            for file_name in files:
                source_path = os.path.join(sub_root, file_name)
                dest_path = os.path.join(main_folder_path, file_name)

                try:
                    if os.path.exists(dest_path):
                        conflictos += 1
                        continue
                    
                    shutil.copy2(source_path, dest_path)
                    copiados += 1
                except Exception as e:
                    errores += 1
                    
    msg = f"Consolidación completada. Copiados: {copiados}, Conflictos: {conflictos}, Errores: {errores}."
    if not silent_mode:
        if progress_bar: progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)

    if return_zip:
        try:
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(root_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, root_path)
                        zf.write(file_path, arcname)
            
            try: shutil.rmtree(root_path, ignore_errors=True)
            except: pass
            
            return {
                "files": [{
                    "name": f"Consolidados_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Consolidados (ZIP)"
                }],
                "message": msg
            }
        except Exception as e:
            return {"error": f"Error creando ZIP: {e}", "message": msg}

    return {"message": msg}

def worker_firmar_docx_con_imagen_masivo(base_path, docx_filename, signature_filename, silent_mode=False, return_zip=False):
    is_native_mode = st.session_state.get('force_native_mode', True)
    
    try:
        if is_native_mode:
            if not send_command:
                return {"error": "Error: Modo nativo activado pero cliente agente no disponible."}
            
            username = st.session_state.get("username", "admin")
            if not silent_mode:
                st.info(f"Enviando tarea al agente local para firmar docx...")
                
            task_id = send_command(username, "sign_docx_massive", {
                "base_path": base_path,
                "docx_filename": docx_filename,
                "signature_filename": signature_filename
            })
            
            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=300)
                if not silent_mode: status_placeholder.empty()
                
                if res and "error" not in res:
                    c = res.get("count", 0)
                    errs = res.get("errors", [])
                    return {"message": f"Proceso finalizado. Modificados: {c}, Errores: {len(errs)}"}
                else:
                    return {"error": f"Error del agente: {res.get('error', 'Desconocido') if res else 'Desconocido'}"}
            else:
                return {"error": "No se pudo crear la tarea."}
        
        # Server Execution
        folders_to_process = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        if not folders_to_process: return {"error": "No se encontraron carpetas para procesar."}
        
        procesados = 0
        errores = 0
        
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Modificando documentos...")
            
        for i, folder_name in enumerate(folders_to_process):
            if not silent_mode:
                progress_bar.progress((i + 1) / len(folders_to_process), text=f"Procesando: {folder_name}")
                
            folder_path = os.path.join(base_path, folder_name)
            docx_path = os.path.join(folder_path, docx_filename)
            
            signature_path = None
            possible_sig_paths = [
                os.path.join(folder_path, signature_filename),
                os.path.join(folder_path, "tipografia", signature_filename)
            ]
            for path in possible_sig_paths:
                if os.path.exists(path):
                    signature_path = path
                    break
            
            if not os.path.exists(docx_path) or not signature_path:
                errores += 1
                continue
                
            # Worker logic for single file
            try:
                doc = Document(docx_path)
                anchor_text = "Firma de Aceptacion"
                signature_p_index = -1

                for idx, p in enumerate(doc.paragraphs):
                    if anchor_text.lower() in p.text.lower():
                        target_index = idx + 1
                        if target_index < len(doc.paragraphs):
                            signature_p_index = target_index
                        break
                
                if signature_p_index != -1:
                    signature_p = doc.paragraphs[signature_p_index]
                    p_element = signature_p._p
                    p_element.clear_content()
                    run = signature_p.add_run()
                    run.add_picture(signature_path, width=Inches(1.5))
                    signature_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    doc.save(docx_path)
                    procesados += 1
                else:
                    errores += 1
            except Exception:
                errores += 1
                
        msg = f"Proceso finalizado. Modificados: {procesados}, Errores/Omitidos: {errores}"
        
        if return_zip:
            try:
                mem_zip = io.BytesIO()
                with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, dirs, files in os.walk(base_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, base_path)
                            zf.write(file_path, arcname)
                
                try: shutil.rmtree(base_path, ignore_errors=True)
                except: pass
                
                return {
                    "files": [{
                        "name": f"Firmados_Masivo_{int(time.time())}.zip",
                        "data": mem_zip.getvalue(),
                        "label": "Descargar Firmados (ZIP)"
                    }],
                    "message": msg
                }
            except Exception as e:
                return {"error": f"Error creando ZIP: {e}", "message": msg}

        return {"message": msg}
    except Exception as e:
        return {"error": f"Error general: {e}"}

def worker_txt_a_json_individual(file_list, silent_mode=False, return_zip_from_folder=None):
    count = 0
    errores = 0
    for file_path in file_list:
        try:
            base, _ = os.path.splitext(file_path)
            new_path = base + ".json"
            os.rename(file_path, new_path)
            count += 1
        except Exception as e:
            errores += 1
            if not silent_mode: st.warning(f"Error renombrando {file_path}: {e}")
            
    msg = f"Renombrados {count} archivos. Errores: {errores}"
    
    if return_zip_from_folder and os.path.isdir(return_zip_from_folder):
        try:
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(return_zip_from_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, return_zip_from_folder)
                        zf.write(file_path, arcname)
            
            try: shutil.rmtree(return_zip_from_folder, ignore_errors=True)
            except: pass
            
            return {
                "files": [{
                    "name": f"Renombrados_JSON_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Renombrados (ZIP)"
                }],
                "message": msg
            }
        except Exception as e:
            return {"error": f"Error creando ZIP: {e}", "message": msg}
            
    return {"message": msg}

def worker_organizar_facturas_feov(root_path, target_path, silent_mode=False, return_zip=False):
    if not root_path or not target_path:
        return {"error": "Error: Rutas de origen o destino no válidas."}

    if not silent_mode: st.info("Iniciando organización de facturas FEOV...")
    
    regex = re.compile(r'FEOV(\d+)', re.IGNORECASE)
    destinos_map = {}
    
    # Paso 1: Mapear destinos
    try:
        list_carpetas_destino = [d for d in os.listdir(target_path) if os.path.isdir(os.path.join(target_path, d))]
    except Exception as e:
        return {"error": f"Error leyendo destinos: {e}"}

    for nombre_carpeta_destino in list_carpetas_destino:
        ruta_carpeta_destino = os.path.join(target_path, nombre_carpeta_destino)
        try:
            for archivo in os.listdir(ruta_carpeta_destino):
                if archivo.lower().endswith('.pdf'):
                    match = regex.search(archivo)
                    if match:
                        numero_factura = match.group(1)
                        destinos_map[numero_factura] = ruta_carpeta_destino
                        break
        except Exception as e:
            if not silent_mode: st.warning(f"Error procesando {nombre_carpeta_destino}: {e}")

    if not destinos_map:
        return {"error": "No se encontraron facturas FEOV en las carpetas de destino."}

    # Paso 2: Mover archivos
    movidos, errores, conflictos = 0, 0, 0
    
    progress_bar = None
    if not silent_mode: progress_bar = st.progress(0, text="Organizando...")
    
    files_to_move = []
    for root, _, files in os.walk(root_path):
        for f in files:
            files_to_move.append((root, f))
            
    total = len(files_to_move)
    
    for i, (root, file_to_move) in enumerate(files_to_move):
        if not silent_mode and i % 10 == 0 and total > 0:
             progress_bar.progress(min(i / total, 1.0), text=f"Procesando {i}/{total}")

        moved = False
        for numero_factura, ruta_destino_final in destinos_map.items():
            if numero_factura in file_to_move:
                try:
                    ruta_origen_archivo = os.path.join(root, file_to_move)
                    ruta_final_archivo = os.path.join(ruta_destino_final, file_to_move)

                    if os.path.exists(ruta_final_archivo):
                        conflictos += 1
                    else:
                        shutil.move(ruta_origen_archivo, ruta_destino_final)
                        movidos += 1
                    moved = True
                    break 
                except Exception:
                    errores += 1
                    break
    
    msg = f"Proceso finalizado. Movidos: {movidos}, Conflictos: {conflictos}, Errores: {errores}"
    if not silent_mode: 
        if progress_bar: progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        
    if return_zip:
        try:
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(target_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, target_path)
                        zf.write(file_path, arcname)
            
            try: shutil.rmtree(target_path, ignore_errors=True)
            except: pass
            # We should also clean root_path if it was temp?
            # But root_path might be same as target_path or different.
            # Assuming root_path is handled by caller or is distinct.
            
            return {
                "files": [{
                    "name": f"Organizados_FEOV_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Organizados (ZIP)"
                }],
                "message": msg
            }
        except Exception as e:
            return {"error": f"Error creando ZIP: {e}", "message": msg}
            
    return {"message": msg}
        
    return f"Organización FEOV finalizada. Movidos: {movidos}, Conflictos: {conflictos}, Errores: {errores}"



def worker_crear_carpetas_desde_excel(excel_path, sheet_name, col_idx, target_folder=None, visible_only=False, silent_mode=False):
    if not excel_path or not sheet_name or col_idx is None:
        return {"error": "Faltan parámetros (Excel, Hoja o Columna)."}
        
    try:
        is_temp = False
        is_native_mode = st.session_state.get('force_native_mode', True)

        if not target_folder:
            is_temp = True
            target_folder = os.path.join(os.getcwd(), "temp_downloads", f"carpetas_{int(time.time())}")
            
        # Only create folder on server if NOT native mode or if it IS temp
        if not is_native_mode or is_temp:
            os.makedirs(target_folder, exist_ok=True)

        nombres_carpetas_raw = []
        if visible_only:
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            ws = wb[sheet_name]
            # openpyxl uses 1-based indexing
            col_1based = col_idx + 1 
            for i in range(2, ws.max_row + 1):
                if not ws.row_dimensions[i].hidden:
                    val = ws.cell(row=i, column=col_1based).value
                    if val: nombres_carpetas_raw.append(str(val))
        else:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            nombres_carpetas_raw = df.iloc[:, col_idx].dropna().astype(str).tolist()
            
        if not nombres_carpetas_raw: return {"error": "No se encontraron nombres."}
        
        # --- NATIVE AGENT EXECUTION ---
        if is_native_mode and not is_temp:
             if not send_command:
                 return {"error": "Error: Modo nativo activado pero el cliente del agente no está disponible (send_command)."}
                 
             username = st.session_state.get("username", "admin")
             
             if not silent_mode:
                 st.info(f"Enviando tarea al agente local para crear {len(nombres_carpetas_raw)} carpetas...")
             
             task_id = send_command(username, "create_folders_from_list", {
                 "base_path": target_folder,
                 "names": nombres_carpetas_raw,
                 "unique": True
             })
             if task_id:
                 # Poll
                 status_placeholder = st.empty()
                 if not silent_mode:
                     status_placeholder.text("Esperando agente...")
                 
                 res = wait_for_result(task_id, timeout=30)
                 
                 if not silent_mode: status_placeholder.empty()

                 if res and "error" not in res:
                     # Successful result is the result dict itself
                     count = res.get("count", 0)
                     errors = res.get("errors", [])
                     
                     msg = f"Creadas: {count}"
                     if errors:
                         msg += f", Errores: {len(errors)}"
                     return {"message": f"{msg} en {target_folder} (Agente)"}
                 else:
                     err = res.get("error") if res else "Error desconocido o tiempo de espera agotado"
                     return {"error": f"Error del agente: {err}"}
             else:
                 return {"error": "Error enviando tarea: No se pudo crear la tarea."}

        # --- SERVER SIDE EXECUTION ---
        creadas, errores = 0, 0
        progress_bar = None
        if not silent_mode: progress_bar = st.progress(0, text="Creando carpetas...")
        
        for i, nombre in enumerate(nombres_carpetas_raw):
            if not silent_mode: progress_bar.progress((i + 1) / len(nombres_carpetas_raw))
            nombre_base = "".join(c for c in nombre if c.isalnum() or c in " _-").rstrip()
            if not nombre_base: continue
            
            ruta_final = os.path.join(target_folder, nombre_base)
            if os.path.exists(ruta_final):
                c = 2
                while os.path.exists(os.path.join(target_folder, f"{nombre_base} ({c})")):
                    c += 1
                ruta_final = os.path.join(target_folder, f"{nombre_base} ({c})")
            
            try:
                os.makedirs(ruta_final)
                creadas += 1
            except: errores += 1
            
        msg = f"Creadas: {creadas}, Errores: {errores}"
        
        if is_temp:
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(target_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, target_folder)
                        zf.write(file_path, arcname)
                        
            # Clean up temp folder
            try: shutil.rmtree(target_folder, ignore_errors=True)
            except: pass
            
            return {
                "files": [{
                    "name": f"Carpetas_Creadas_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Carpetas (ZIP)"
                }],
                "message": msg
            }
        else:
            return {"message": f"{msg} en {target_folder}"}
            
    except Exception as e:
        return {"error": f"Error: {e}"}

# --- WORKERS: PDF ---

def worker_unificar_pdfs_list(file_list, output_path=None, sort_method="Nombre", silent_mode=False):
    try:
        if not file_list: return {"error": "No hay archivos para unificar."}
        
        # Sort files
        if sort_method == "Nombre":
            file_list.sort(key=lambda x: x.name if hasattr(x, 'name') else os.path.basename(x))
        
        doc_final = fitz.open()
        for f in file_list:
            try:
                # Handle both file paths and BytesIO/UploadedFile
                if isinstance(f, str):
                    doc = fitz.open(f)
                else:
                    f.seek(0)
                    doc = fitz.open(stream=f.read(), filetype="pdf")
                
                doc_final.insert_pdf(doc)
                doc.close()
            except Exception as e:
                if not silent_mode: st.warning(f"Omitiendo archivo por error: {e}")
        
        if output_path and not output_path.startswith("Unificado"):
            doc_final.save(output_path)
            doc_final.close()
            return {"message": f"PDF Unificado creado en: {output_path}"}
        else:
            # Return bytes for download
            out_buffer = io.BytesIO()
            doc_final.save(out_buffer)
            doc_final.close()
            return {
                "files": [{
                    "name": f"Unificado_{int(time.time())}.pdf",
                    "data": out_buffer.getvalue(),
                    "label": "Descargar PDF Unificado"
                }],
                "message": "PDF Unificado correctamente."
            }
            
    except Exception as e:
        return {"error": f"Error unificando PDFs: {e}"}

def worker_dividir_pdf_paginas(input_pdf, output_folder=None, silent_mode=False):
    try:
        if isinstance(input_pdf, str):
            doc = fitz.open(input_pdf)
            name_base = os.path.splitext(os.path.basename(input_pdf))[0]
        else:
            input_pdf.seek(0)
            doc = fitz.open(stream=input_pdf.read(), filetype="pdf")
            name_base = os.path.splitext(input_pdf.name)[0]
            
        is_temp = False
        if not output_folder:
            is_temp = True
            output_folder = os.path.join(os.getcwd(), "temp_downloads", f"split_{int(time.time())}")
            
        os.makedirs(output_folder, exist_ok=True)
        
        for i in range(len(doc)):
            # out_name = f"{name_base}_pag_{i+1}.pdf"
            # Instead of _pag_, let's format it nicely
            out_name = f"{name_base} - Pagina {i+1}.pdf"
            
            # Create new PDF for single page
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=i, to_page=i)
            new_doc.save(os.path.join(output_folder, out_name))
            new_doc.close()
            
        msg = f"PDF dividido en {len(doc)} páginas."
        
        if True: # Always return zip for manual web downloads to ensure it works
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(output_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, output_folder)
                        zf.write(file_path, arcname)
            
            # If it was temp, clean it up. If it wasn't temp, leave the files on disk but still return the zip.
            if is_temp:
                try: shutil.rmtree(output_folder, ignore_errors=True)
                except: pass
            
            return {
                "files": [{
                    "name": f"PDF_Dividido_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Páginas (ZIP)"
                }],
                "message": msg
            }

    except Exception as e:
        return {"error": f"Error dividiendo PDF: {e}"}

def worker_unificar_imagenes_pdf(folder_path, output_name="Unificado.pdf", silent_mode=False):
    try:
        images = [
            os.path.join(folder_path, f) 
            for f in sorted(os.listdir(folder_path), key=natural_sort_key) 
            if f.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp'))
        ]
        
        if not images:
            return {"error": "No se encontraron imágenes en la carpeta."}
            
        img_list = []
        first_img = None
        
        for img_path in images:
            try:
                img = Image.open(img_path).convert('RGB')
                if first_img is None:
                    first_img = img
                else:
                    img_list.append(img)
            except Exception as e:
                if not silent_mode: st.warning(f"Error cargando imagen {img_path}: {e}")
                
        if first_img:
            output = io.BytesIO()
            first_img.save(output, format="PDF", save_all=True, append_images=img_list)
            pdf_bytes = output.getvalue()
            
            # Save to disk only if it's a persistent folder (Native Mode)
            # We assume folder_path is persistent if it's not in temp_uploads/temp_downloads
            is_temp = "temp_" in folder_path.lower()
            
            if not is_temp:
                pdf_path = os.path.join(folder_path, output_name)
                try:
                    with open(pdf_path, "wb") as f:
                        f.write(pdf_bytes)
                except: pass

            return {
                "files": [{
                    "name": os.path.basename(output_name),
                    "data": pdf_bytes,
                    "label": "Descargar PDF Unificado"
                }],
                "message": f"PDF creado exitosamente ({len(images)} imágenes)"
            }
        else:
            return {"error": "No se pudieron procesar las imágenes."}
    except Exception as e:
        return {"error": f"Error: {e}"}

# --- WORKERS: PDF EXTENDED ---

def worker_unificar_por_carpeta(carpeta_base, nombre_final_base, silent_mode=False):
    if not carpeta_base or not os.path.isdir(carpeta_base): return {"error": "Carpeta base inválida."}
    
    subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
    if not subcarpetas: return {"error": "No se encontraron subcarpetas."}

    log = []
    pdfs_creados = 0
    
    for carpeta in subcarpetas:
        nombre_subcarpeta = os.path.basename(carpeta)
        
        archivos_pdf_a_procesar = []
        for num_pdf in range(1, 11):
            nombre_archivo_buscado = f"{num_pdf}.pdf"
            ruta_archivo_buscado = os.path.join(carpeta, nombre_archivo_buscado)
            if os.path.exists(ruta_archivo_buscado):
                archivos_pdf_a_procesar.append(ruta_archivo_buscado)

        if not archivos_pdf_a_procesar:
            continue
        
        nombre_final = f"{nombre_final_base}.pdf"
        ruta_salida = os.path.join(carpeta, nombre_final)
        
        try:
            doc_final = fitz.open()
            for ruta_pdf in archivos_pdf_a_procesar:
                with fitz.open(ruta_pdf) as doc_origen:
                    for page_origen in doc_origen:
                        pix = page_origen.get_pixmap(dpi=300, colorspace=fitz.csGRAY)
                        pagina_nueva = doc_final.new_page(width=pix.width, height=pix.height)
                        pagina_nueva.insert_image(pagina_nueva.rect, pixmap=pix)
            
            if len(doc_final) > 0:
                doc_final.save(ruta_salida, garbage=4, deflate=True)
                pdfs_creados += 1
            doc_final.close()
        except Exception as e:
            log.append(f"Error en {nombre_subcarpeta}: {e}")

    msg = f"Proceso finalizado. {pdfs_creados} PDFs creados." + (" Errores: " + "; ".join(log) if log else "")
    
    # Always attempt to zip if it looks like a temp folder (or just return the ability to zip)
    # Since this modifies in-place, we can offer a download of the whole structure.
    
    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(carpeta_base):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, carpeta_base)
                zf.write(file_path, arcname)
    
    return {
        "files": [{
            "name": f"Procesados_{int(time.time())}.zip",
            "data": mem_zip.getvalue(),
            "label": "Descargar Resultados (ZIP)"
        }],
        "message": msg
    }

def worker_unificar_imagenes_por_carpeta_rec(carpeta_base, nombre_final_base, tipo_imagen="JPG", silent_mode=False):
    # tipo_imagen: "JPG" or "PNG"
    if not carpeta_base or not os.path.isdir(carpeta_base): return {"error": "Carpeta base inválida."}
    
    ext_map = {
        "JPG": ['.jpg', '.jpeg'],
        "PNG": ['.png']
    }
    exts = ext_map.get(tipo_imagen, ['.jpg'])

    subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
    if not subcarpetas: return {"error": "No se encontraron subcarpetas."}
    
    pdfs_creados = 0
    log = []

    for carpeta in subcarpetas:
        nombre_subcarpeta = os.path.basename(carpeta)
        
        archivos_img_a_procesar = []
        # Buscar 1.jpg, 2.jpg, ...
        for num_img in range(1, 11):
            ruta_encontrada = None
            for ext in exts:
                nombre_archivo = f"{num_img}{ext}"
                ruta_archivo = os.path.join(carpeta, nombre_archivo)
                if os.path.exists(ruta_archivo):
                    ruta_encontrada = ruta_archivo
                    break
            if ruta_encontrada:
                archivos_img_a_procesar.append(ruta_encontrada)
        
        if not archivos_img_a_procesar:
            continue

        try:
            lista_imagenes_procesadas = []
            for ruta_img in archivos_img_a_procesar:
                img = Image.open(ruta_img)
                if tipo_imagen == "PNG":
                     if img.mode in ('RGBA', 'LA'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1]) # Use alpha channel as mask
                        img = background
                     elif img.mode != 'RGB':
                        img = img.convert('RGB')
                else:
                    img = img.convert('L') # Grayscale for JPG as per original code

                lista_imagenes_procesadas.append(img)
            
            if lista_imagenes_procesadas:
                nombre_pdf = f"{nombre_final_base}.pdf"
                ruta_salida = os.path.join(carpeta, nombre_pdf)
                lista_imagenes_procesadas[0].save(
                    ruta_salida, 
                    save_all=True, 
                    append_images=lista_imagenes_procesadas[1:], 
                    resolution=300.0
                )
                pdfs_creados += 1
        except Exception as e:
            log.append(f"Error en {nombre_subcarpeta}: {e}")

    msg = f"Proceso finalizado. {pdfs_creados} PDFs creados." + (" Errores: " + "; ".join(log) if log else "")
    
    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(carpeta_base):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, carpeta_base)
                zf.write(file_path, arcname)

    return {
        "files": [{
            "name": f"Procesados_Img_{int(time.time())}.zip",
            "data": mem_zip.getvalue(),
            "label": "Descargar Resultados (ZIP)"
        }],
        "message": msg
    }

def worker_unificar_docx_por_carpeta(carpeta_base, nombre_final_base, silent_mode=False):
    if not HAS_DOCX2PDF: return {"error": "docx2pdf no está instalado."}
    if not carpeta_base or not os.path.isdir(carpeta_base): return {"error": "Carpeta base inválida."}
    
    subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
    pdfs_creados = 0
    log = []

    for carpeta in subcarpetas:
        nombre_subcarpeta = os.path.basename(carpeta)
        archivos_docx_a_procesar = []
        for num_doc in range(1, 11):
            nombre_archivo = f"{num_doc}.docx"
            ruta_archivo = os.path.join(carpeta, nombre_archivo)
            if os.path.exists(ruta_archivo):
                archivos_docx_a_procesar.append(ruta_archivo)
        
        if not archivos_docx_a_procesar: continue

        pdfs_temporales = []
        try:
            for ruta_docx in archivos_docx_a_procesar:
                nombre_temp_pdf = os.path.splitext(os.path.basename(ruta_docx))[0] + "_temp.pdf"
                ruta_temp_pdf = os.path.join(carpeta, nombre_temp_pdf)
                try:
                    convert_docx_to_pdf(ruta_docx, ruta_temp_pdf)
                    if os.path.exists(ruta_temp_pdf):
                        pdfs_temporales.append(ruta_temp_pdf)
                except: pass
            
            if pdfs_temporales:
                nombre_pdf_final = f"{nombre_final_base}.pdf"
                ruta_salida = os.path.join(carpeta, nombre_pdf_final)
                
                doc_final = fitz.open()
                for pdf_temp in pdfs_temporales:
                    try:
                        with fitz.open(pdf_temp) as doc_temp:
                            doc_final.insert_pdf(doc_temp)
                    except Exception as e:
                        log.append(f"Error uniendo {pdf_temp}: {e}")
                
                doc_final.save(ruta_salida)
                doc_final.close()
                pdfs_creados += 1
                
                for pdf_temp in pdfs_temporales:
                    try: os.remove(pdf_temp)
                    except Exception as e:
                        log.append(f"Error eliminando {pdf_temp}: {e}")
        except Exception as e:
            log.append(f"Error en {nombre_subcarpeta}: {e}")

    msg = f"Proceso finalizado. {pdfs_creados} PDFs creados." + (" Errores: " + "; ".join(log) if log else "")

    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(carpeta_base):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, carpeta_base)
                zf.write(file_path, arcname)

    return {
        "files": [{
            "name": f"Procesados_Docx_{int(time.time())}.zip",
            "data": mem_zip.getvalue(),
            "label": "Descargar Resultados (ZIP)"
        }],
        "message": msg
    }

def worker_dividir_pdfs_masivamente(carpeta_base, silent_mode=False):
    if not carpeta_base or not os.path.isdir(carpeta_base): return {"error": "Carpeta inválida."}
    
    pdfs_a_procesar = []
    for root, _, files in os.walk(carpeta_base):
        for file in files:
            if file.lower().endswith('.pdf'):
                pdfs_a_procesar.append(os.path.join(root, file))
    
    if not pdfs_a_procesar: return {"error": "No se encontraron PDFs."}
    
    count = 0
    for ruta_pdf_original in pdfs_a_procesar:
        try:
            nombre_base_original = os.path.splitext(os.path.basename(ruta_pdf_original))[0]
            directorio_origen = os.path.dirname(ruta_pdf_original)
            ruta_carpeta_salida = os.path.join(directorio_origen, nombre_base_original)
            os.makedirs(ruta_carpeta_salida, exist_ok=True)
            
            with fitz.open(ruta_pdf_original) as doc_origen:
                if len(doc_origen) == 0: continue
                for i in range(len(doc_origen)):
                    doc_nuevo = fitz.open()
                    doc_nuevo.insert_pdf(doc_origen, from_page=i, to_page=i)
                    doc_nuevo.save(os.path.join(ruta_carpeta_salida, f"{i+1}.pdf"))
                    doc_nuevo.close()
            count += 1
        except: pass
        
    msg = f"Divididos {count} PDFs masivamente."
    
    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(carpeta_base):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, carpeta_base)
                zf.write(file_path, arcname)

    return {
        "files": [{
            "name": f"Divididos_Masivo_{int(time.time())}.zip",
            "data": mem_zip.getvalue(),
            "label": "Descargar Resultados (ZIP)"
        }],
        "message": msg
    }



# --- WORKERS: RIPS ---

def worker_json_a_xlsx_ind(file_obj, silent_mode=False):
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
        data = json.load(file_obj)
        
        service_map = {
            "consultas": "Consultas", "procedimientos": "Procedimientos", "urgencias": "Urgencias",
            "hospitalizacion": "Hospitalizacion", "recienNacidos": "RecienNacidos",
            "medicamentos": "Medicamentos", "otrosServicios": "OtrosServicios"
        }
        
        header_info = {
            "numDocumentoIdObligado": data.get("numDocumentoIdObligado"),
            "numFactura": data.get("numFactura"),
            "tipoNota": data.get("tipoNota"),
            "numNota": data.get("numNota")
        }
        
        usuarios_rows = []
        all_services = {name: [] for name in service_map.values()}
        usuarios_lista = data.get("usuarios", []) if isinstance(data, dict) else []
        
        for usuario in usuarios_lista:
            u_info = {
                "tipoDocumentoIdentificacion": get_val_ci(usuario, "tipoDocumentoIdentificacion"),
                "numDocumentoIdentificacion": get_val_ci(usuario, "numDocumentoIdentificacion"),
                "tipoUsuario": get_val_ci(usuario, "tipoUsuario"),
                "fechaNacimiento": get_val_ci(usuario, "fechaNacimiento"),
                "codSexo": get_val_ci(usuario, "codSexo"),
                "codPaisResidencia": get_val_ci(usuario, "codPaisResidencia"),
                "codMunicipioResidencia": get_val_ci(usuario, "codMunicipioResidencia"),
                "codZonaTerritorialResidencia": get_val_ci(usuario, "codZonaTerritorialResidencia"),
                "incapacidad": get_val_ci(usuario, "incapacidad"),
                "consecutivo": get_val_ci(usuario, "consecutivo"),
                "codPaisOrigen": get_val_ci(usuario, "codPaisOrigen"),
            }
            usuarios_rows.append(u_info)
            servicios = usuario.get("servicios", {})
            for json_key, sheet_name in service_map.items():
                items = get_val_ci(servicios, json_key)
                if items and isinstance(items, list):
                    for item in items:
                        item["consecutivoUsuario"] = u_info["consecutivo"]
                        all_services[sheet_name].append(item)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame([header_info]).to_excel(writer, sheet_name="Transaccion", index=False)
            if usuarios_rows:
                pd.DataFrame(usuarios_rows).to_excel(writer, sheet_name="Usuarios", index=False)
            for sheet_name, rows in all_services.items():
                if rows:
                    pd.DataFrame(rows).to_excel(writer, sheet_name=sheet_name, index=False)
        
        return {
            "files": [{
                "name": f"RIPS_Convertido_{int(time.time())}.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Excel (RIPS)"
            }],
            "message": "Conversión JSON a Excel exitosa."
        }
    except Exception as e:
        return {"error": str(e)}

def worker_consolidar_json_xlsx(folder_path, silent_mode=False):
    try:
        json_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.json')]
        if not json_files:
            return {"error": "No hay archivos JSON en la carpeta."}
        
        master_header = []
        master_users = []
        master_services = {
            "Consultas": [], "Procedimientos": [], "Urgencias": [], 
            "Hospitalizacion": [], "RecienNacidos": [], "Medicamentos": [], "OtrosServicios": []
        }
        service_map = {
            "consultas": "Consultas", "procedimientos": "Procedimientos", "urgencias": "Urgencias",
            "hospitalizacion": "Hospitalizacion", "recienNacidos": "RecienNacidos",
            "medicamentos": "Medicamentos", "otrosServicios": "OtrosServicios"
        }
        
        for fname in json_files:
            with open(os.path.join(folder_path, fname), 'r', encoding='utf-8') as f:
                data = json.load(f)
            h_info = {
                "archivo_origen": fname,
                "numDocumentoIdObligado": data.get("numDocumentoIdObligado"),
                "numFactura": data.get("numFactura")
            }
            master_header.append(h_info)
            usuarios = data.get("usuarios", [])
            for u in usuarios:
                u_clean = {k: v for k, v in u.items() if k != "servicios"}
                u_clean["archivo_origen"] = fname
                master_users.append(u_clean)
                servicios = u.get("servicios", {})
                for j_key, s_name in service_map.items():
                    items = get_val_ci(servicios, j_key)
                    if items:
                        for item in items:
                            item["archivo_origen"] = fname
                            item["consecutivoUsuario"] = u.get("consecutivo")
                            master_services[s_name].append(item)
                            
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            pd.DataFrame(master_header).to_excel(writer, sheet_name="Transaccion", index=False)
            pd.DataFrame(master_users).to_excel(writer, sheet_name="Usuarios", index=False)
            for s_name, rows in master_services.items():
                if rows:
                    pd.DataFrame(rows).to_excel(writer, sheet_name=s_name, index=False)
        
        return {
            "files": [{
                "name": f"Consolidado_RIPS_{int(time.time())}.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Consolidado (XLSX)"
            }],
            "message": f"Consolidados {len(json_files)} archivos."
        }
    except Exception as e:
        return {"error": str(e)}

def worker_xlsx_a_json_ind(file_obj, silent_mode=False):
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
        xls = pd.ExcelFile(file_obj)
        service_map = {
            "Consultas": "consultas", "Procedimientos": "procedimientos", "Urgencias": "urgencias",
            "Hospitalizacion": "hospitalizacion", "RecienNacidos": "recienNacidos",
            "Medicamentos": "medicamentos", "OtrosServicios": "otrosServicios"
        }
        if "Transaccion" in xls.sheet_names and "Usuarios" in xls.sheet_names:
            df_t = pd.read_excel(xls, sheet_name="Transaccion")
            df_t = clean_df_for_json(df_t)
            transaccion_data = df_t.iloc[0].to_dict() if not df_t.empty else {}
            
            usuarios_map = {}
            df_u = pd.read_excel(xls, sheet_name="Usuarios")
            df_u = clean_df_for_json(df_u)
            for _, row in df_u.iterrows():
                u_obj = row.to_dict()
                u_obj["servicios"] = {k: [] for k in service_map.values()}
                usuarios_map[str(u_obj.get("consecutivo"))] = u_obj
            
            for sheet_name, json_key in service_map.items():
                if sheet_name in xls.sheet_names:
                    df_s = pd.read_excel(xls, sheet_name=sheet_name)
                    df_s = clean_df_for_json(df_s)
                    for _, row in df_s.iterrows():
                        s_obj = row.to_dict()
                        cons_u = str(s_obj.pop("consecutivoUsuario", None))
                        if cons_u in usuarios_map:
                            usuarios_map[cons_u]["servicios"][json_key].append(s_obj)
            
            final_json = transaccion_data
            final_json["usuarios"] = list(usuarios_map.values())
            json_str = json.dumps(final_json, ensure_ascii=False, indent=4)
            
            return {
                "files": [{
                    "name": f"RIPS_Generado_{int(time.time())}.json",
                    "data": json_str.encode('utf-8'),
                    "label": "Descargar JSON (RIPS)"
                }],
                "message": "Conversión Excel a JSON exitosa."
            }
        return {"error": "Formato Excel inválido (Faltan hojas Transaccion/Usuarios)"}
    except Exception as e:
        return {"error": str(e)}

def worker_rips_excel_to_json_original(file_obj, silent_mode=False):
    """
    Worker para convertir Excel a JSON con estructura RIPS (Original).
    Espera hojas: Consultas, Procedimientos, OtrosServicios.
    """
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
        xls = pd.ExcelFile(file_obj)
        
        usuarios_dict = {}
        
        def procesar_hoja(nombre_hoja, clave_servicio):
            if nombre_hoja in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=nombre_hoja)
                df = df.astype(object).where(pd.notnull(df), None)
                
                for _, row in df.iterrows():
                    td = str(row.get("tipo_documento_usuario", ""))
                    doc = str(row.get("documento_usuario", ""))
                    user_key = (td, doc)
                    
                    if user_key not in usuarios_dict:
                        usuarios_dict[user_key] = {
                            "tipoDocumentoIdentificacion": row.get("tipo_documento_usuario"),
                            "numDocumentoIdentificacion": row.get("documento_usuario"),
                            "tipoUsuario": row.get("tipo_usuario"),
                            "fechaNacimiento": row.get("fecha_nacimiento"), 
                            "codSexo": row.get("sexo"),
                            "codPaisResidencia": row.get("pais_residencia"),
                            "municipio_residencia": row.get("municipio_residencia"),
                            "codZonaTerritorialResidencia": row.get("zona_residencia"),
                            "incapacidad": row.get("incapacidad"),
                            "consecutivo": row.get("consecutivo_usuario"),
                            "codPaisOrigen": row.get("pais_origen"),
                            "servicios": {
                                "consultas": [],
                                "procedimientos": [],
                                "otrosServicios": []
                            }
                        }
                    
                    servicio_data = row.to_dict()
                    keys_to_remove = [
                        "tipo_documento_usuario", "documento_usuario", "tipo_usuario", 
                        "fecha_nacimiento", "sexo", "pais_residencia", "municipio_residencia", 
                        "zona_residencia", "incapacidad", "consecutivo_usuario", "pais_origen"
                    ]
                    for k in keys_to_remove:
                        servicio_data.pop(k, None)
                        
                    if any(v is not None for v in servicio_data.values()):
                        usuarios_dict[user_key]["servicios"][clave_servicio].append(servicio_data)

        procesar_hoja("Consultas", "consultas")
        procesar_hoja("Procedimientos", "procedimientos")
        procesar_hoja("OtrosServicios", "otrosServicios")
        
        resultado_final = {
            "usuarios": list(usuarios_dict.values())
        }
        
        json_str = json.dumps(resultado_final, ensure_ascii=False, indent=4)
        
        return {
            "files": [{
                "name": f"RIPS_Original_{int(time.time())}.json",
                "data": json_str.encode('utf-8'),
                "label": "Descargar JSON (Original)"
            }],
            "message": "Conversión Excel a JSON (Original) exitosa."
        }
        
    except Exception as e:
        return {"error": str(e)}

def worker_rips_json_to_excel_original(file_obj, silent_mode=False):
    """
    Worker para convertir JSON a Excel con estructura RIPS (Original).
    Genera hojas: Consultas, Procedimientos, OtrosServicios.
    """
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
            data = json.load(file_obj)
        else:
            with open(file_obj, "r", encoding="utf-8") as f:
                data = json.load(f)

        consultas = []
        procedimientos = []
        otros_servicios = []

        for usuario in data.get("usuarios", []):
            base_info = {
                "tipo_documento_usuario": usuario.get("tipoDocumentoIdentificacion"),
                "documento_usuario": usuario.get("numDocumentoIdentificacion"),
                "tipo_usuario": usuario.get("tipoUsuario"),
                "fecha_nacimiento": usuario.get("fechaNacimiento"),
                "sexo": usuario.get("codSexo"),
                "pais_residencia": usuario.get("codPaisResidencia"),
                "municipio_residencia": usuario.get("codMunicipioResidencia"),
                "zona_residencia": usuario.get("codZonaTerritorialResidencia"),
                "incapacidad": usuario.get("incapacidad"),
                "consecutivo_usuario": usuario.get("consecutivo"),
                "pais_origen": usuario.get("codPaisOrigen")
            }

            servicios = usuario.get("servicios", {})

            for consulta in servicios.get("consultas", []):
                consultas.append({**base_info, **consulta})

            for procedimiento in servicios.get("procedimientos", []):
                procedimientos.append({**base_info, **procedimiento})

            for otro in servicios.get("otrosServicios", []):
                otros_servicios.append({**base_info, **otro})

        output = io.BytesIO()
        with pd.ExcelWriter(output) as writer:
            if consultas:
                pd.DataFrame(consultas).to_excel(writer, sheet_name="Consultas", index=False)
            if procedimientos:
                pd.DataFrame(procedimientos).to_excel(writer, sheet_name="Procedimientos", index=False)
            if otros_servicios:
                pd.DataFrame(otros_servicios).to_excel(writer, sheet_name="OtrosServicios", index=False)
            
            if not consultas and not procedimientos and not otros_servicios:
                 pd.DataFrame().to_excel(writer, sheet_name="Vacio", index=False)
                 
        return {
            "files": [{
                "name": f"RIPS_Original_{int(time.time())}.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Excel (Original)"
            }],
            "message": "Conversión JSON a Excel (Original) exitosa."
        }
    except Exception as e:
        return {"error": str(e)}

def worker_desconsolidar_xlsx_json(file_obj, dest_folder, silent_mode=False):
    try:
        if hasattr(file_obj, 'seek'):
            file_obj.seek(0)
        xls = pd.ExcelFile(file_obj)
        if "Transaccion" not in xls.sheet_names:
            return {"error": "Falta hoja Transaccion"}
        df_t = pd.read_excel(xls, sheet_name="Transaccion")
        if "archivo_origen" not in df_t.columns:
            return {"error": "Falta columna 'archivo_origen' en Transaccion"}
            
        service_map = {
            "Consultas": "consultas", "Procedimientos": "procedimientos", "Urgencias": "urgencias",
            "Hospitalizacion": "hospitalizacion", "RecienNacidos": "recienNacidos",
            "Medicamentos": "medicamentos", "OtrosServicios": "otrosServicios"
        }
        df_t = clean_df_for_json(df_t)
        headers_by_file = {row["archivo_origen"]: row.to_dict() for _, row in df_t.iterrows()}
        
        users_by_file = {}
        if "Usuarios" in xls.sheet_names:
            df_u = clean_df_for_json(pd.read_excel(xls, sheet_name="Usuarios"))
            for _, row in df_u.iterrows():
                fname = row.get("archivo_origen")
                if fname not in users_by_file: users_by_file[fname] = []
                users_by_file[fname].append(row.to_dict())
                
        services_by_file = {}
        for s_name, j_key in service_map.items():
            if s_name in xls.sheet_names:
                df_s = clean_df_for_json(pd.read_excel(xls, sheet_name=s_name))
                for _, row in df_s.iterrows():
                    fname = row.get("archivo_origen")
                    if fname:
                        if fname not in services_by_file: services_by_file[fname] = {k: [] for k in service_map.values()}
                        services_by_file[fname][j_key].append(row.to_dict())

        # Create ZIP in memory for the results
        mem_zip = io.BytesIO()
        with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
            count = 0
            for fname, header in headers_by_file.items():
                header.pop("archivo_origen", None)
                final = header
                users = []
                for u in users_by_file.get(fname, []):
                    u.pop("archivo_origen", None)
                    u_cons = u.get("consecutivo")
                    u["servicios"] = {k: [] for k in service_map.values()}
                    if fname in services_by_file:
                        for s_key, items in services_by_file[fname].items():
                            for item in items:
                                if item.get("consecutivoUsuario") == u_cons:
                                    i_clean = item.copy()
                                    i_clean.pop("archivo_origen", None)
                                    i_clean.pop("consecutivoUsuario", None)
                                    u["servicios"][s_key].append(i_clean)
                    users.append(u)
                final["usuarios"] = users
                
                # Write to zip
                json_content = json.dumps(final, ensure_ascii=False, indent=4)
                zf.writestr(fname, json_content)
                count += 1
                
        return {
            "files": [{
                "name": f"Desconsolidado_RIPS_{int(time.time())}.zip",
                "data": mem_zip.getvalue(),
                "label": "Descargar Desconsolidados (ZIP)"
            }],
            "message": f"Desconsolidados {count} archivos."
        }
    except Exception as e:
        return {"error": str(e)}

# --- WORKERS: EXCEL / RENAMING ---

def worker_aplicar_renombrado_excel(excel_path, folder_path, silent_mode=False, return_zip=False):
    is_native_mode = st.session_state.get('force_native_mode', True)

    try:
        df = pd.read_excel(excel_path)
        if "Nombre Actual" not in df.columns or "Nombre Nuevo" not in df.columns:
            return "Error: Excel debe tener columnas 'Nombre Actual' y 'Nombre Nuevo'"
        
        records = []
        for _, row in df.iterrows():
            if pd.notna(row["Nombre Actual"]) and pd.notna(row["Nombre Nuevo"]):
                records.append({
                    "current_name": str(row["Nombre Actual"]).strip(),
                    "new_name": str(row["Nombre Nuevo"]).strip()
                })

        if not records:
            return "No se encontraron registros válidos para renombrar."

        if is_native_mode:
            if not send_command:
                return "Error: Modo nativo activado pero cliente agente no disponible."
            
            username = st.session_state.get("username", "admin")
            if not silent_mode:
                st.info(f"Enviando tarea al agente local para renombrar {len(records)} archivos...")

            files_to_rename = []
            for item in records:
                curr = item["current_name"]
                new = item["new_name"]
                if "." not in new:
                    _, ext = os.path.splitext(curr)
                    new += ext
                    
                # Fix: Instead of assuming the file is directly in folder_path,
                # we should allow curr to be a relative path from folder_path if the 
                # export function provided it that way. 
                # However, the user states "LOS RENOMBRE PERO LOS ENVIO A ESA RUTA RECUERDA QUE DEBEN QUEDAR EN LA MISMA CARPETA CONTENEDORA"
                # This implies 'curr' contains subfolder structure, but 'new' only contains the filename.
                
                old_full_path = os.path.join(folder_path, curr)
                # Fix: In Windows, the Excel might contain backslashes but python joins with forward slashes 
                # or vice-versa, causing dirname to fail if the separator isn't recognized.
                # Let's normalize the path first.
                old_full_path = os.path.normpath(old_full_path)
                
                # Ensure the new file stays in the exact same subfolder as the old one
                old_dir = os.path.dirname(old_full_path)
                new_full_path = os.path.join(old_dir, os.path.basename(new))
                
                files_to_rename.append({
                    "old_path": old_full_path,
                    "new_path": new_full_path
                })

            task_id = send_command(username, "rename_files", {
                "files": files_to_rename
            })

            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=300)
                
                if not silent_mode: status_placeholder.empty()

                if res and "error" not in res:
                     count = res.get("count", 0)
                     errors = res.get("errors", [])
                     msg = f"Finalizado correctamente (Renombrados: {count})"
                     if errors:
                         msg += f". Hubo {len(errors)} errores."
                     return msg
                else:
                     return f"Error del agente: {res.get('error', 'Desconocido') if res else 'Desconocido'}"
            else:
                return "No se pudo crear la tarea en el servidor."

        # Server Execution
        count = 0
        for item in records:
            curr = item["current_name"]
            new = item["new_name"]
            
            curr_path = os.path.join(folder_path, curr)
            curr_path = os.path.normpath(curr_path)
            
            if os.path.exists(curr_path):
                if "." not in new:
                    _, ext = os.path.splitext(curr)
                    new += ext
                
                old_dir = os.path.dirname(curr_path)
                new_path = os.path.join(old_dir, os.path.basename(new))
                
                try:
                    os.rename(curr_path, new_path)
                    count += 1
                except: pass
        
        msg = f"Renombrados {count} archivos."
        return msg
    except Exception as e:
        return f"Error: {e}"

def worker_anadir_sufijo_excel(excel_path, sheet_name, col_folder, col_suffix, root_path, use_filter=False, silent_mode=False, return_zip=False):
    """
    Versión sincronizada con app_web.py:
    1. Lee Excel (Carpeta, Sufijo).
    2. Busca la carpeta dentro de root_path.
    3. Si existe, renombra TODOS los archivos dentro de esa carpeta añadiendo el sufijo.
    """
    try:
        if isinstance(excel_path, bytes):
            excel_path = io.BytesIO(excel_path)
            if hasattr(excel_path, 'seek'):
                excel_path.seek(0)

        # Helper para limpiar y normalizar valores
        def clean_val(v):
            if v is None: return ""
            s = str(v).strip()
            if s.endswith(".0"): s = s[:-2]
            # Normalizar unicode (quitar tildes, etc)
            s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
            return s.lower()

        data_rows = []
        # --- Extracción de datos ---
        if use_filter:
            import openpyxl
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada"
            ws = wb[sheet_name]
            header = [cell.value for cell in ws[1]]
            try:
                header_map = {str(h).strip().lower(): i for i, h in enumerate(header) if h is not None}
                idx_folder = header_map.get(str(col_folder).strip().lower())
                idx_suffix = header_map.get(str(col_suffix).strip().lower())
                
                if idx_folder is None or idx_suffix is None:
                    return f"Columnas '{col_folder}' o '{col_suffix}' no encontradas."
            except Exception as e: return f"Error al buscar columnas: {e}"
            
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    v_folder = row[idx_folder].value if idx_folder < len(row) else None
                    v_suffix = row[idx_suffix].value if idx_suffix < len(row) else None
                    
                    val_folder_norm = clean_val(v_folder)
                    val_folder_raw = str(v_folder).strip() if v_folder is not None else ""
                    if val_folder_raw.endswith(".0"): val_folder_raw = val_folder_raw[:-2]

                    val_suffix_clean = str(v_suffix).strip() if v_suffix is not None else ""
                    if val_suffix_clean.endswith(".0"): val_suffix_clean = val_suffix_clean[:-2]
                    
                    if val_folder_norm and val_suffix_clean:
                        data_rows.append((val_folder_norm, val_folder_raw, val_suffix_clean))
        else:
            try:
                df = pd.read_excel(excel_path, sheet_name=sheet_name, dtype=str)
            except Exception as e:
                return f"Error leyendo Excel: {e}"
                
            df.columns = [str(c).strip() for c in df.columns]
            if col_folder not in df.columns or col_suffix not in df.columns:
                return f"Columnas no encontradas."
            
            for _, row in df.iterrows():
                if pd.notna(row[col_folder]) and pd.notna(row[col_suffix]):
                    v_folder = row[col_folder]
                    v_suffix = row[col_suffix]
                    
                    val_folder_norm = clean_val(v_folder)
                    val_folder_raw = str(v_folder).strip()
                    if val_folder_raw.endswith(".0"): val_folder_raw = val_folder_raw[:-2]

                    val_suffix_clean = str(v_suffix).strip()
                    if val_suffix_clean.endswith(".0"): val_suffix_clean = val_suffix_clean[:-2]

                    if val_folder_norm and val_suffix_clean:
                        data_rows.append((val_folder_norm, val_folder_raw, val_suffix_clean))

        if not data_rows:
            return "No se encontraron datos válidos."

        is_native_mode = st.session_state.get('force_native_mode', True)
        
        if is_native_mode:
             if not send_command:
                 return "Error: Modo nativo activado pero el cliente del agente no está disponible."
                 
             username = st.session_state.get("username", "admin")
             
             items = [{"key": row[0], "suffix": row[2]} for row in data_rows]
             
             if not silent_mode:
                 st.info(f"Enviando tarea al agente local para añadir sufijos a {len(items)} carpetas...")
             
             task_id = send_command(username, "bulk_rename", {
                 "path": root_path,
                 "items": items,
                 "separator": "", # No separator, suffix string should contain it if needed or we just append
                 "rename_folders": False, # Solo renombrar los archivos internos de la carpeta encontrada
                 "rename_internal_files": True
             })
             
             if task_id:
                 status_placeholder = st.empty()
                 if not silent_mode:
                     status_placeholder.text("Esperando agente...")
                 
                 res = wait_for_result(task_id, timeout=60)
                 
                 if not silent_mode: status_placeholder.empty()

                 if res and "error" not in res:
                     count = res.get("count", 0)
                     errors = res.get("errors", [])
                     msg = f"Proceso completado por el Agente. {count} elementos renombrados."
                     if errors:
                         msg += f" Hubo {len(errors)} errores."
                     return msg
                 else:
                     err = res.get("error") if res else "Error desconocido o tiempo de espera agotado"
                     return f"Error del agente: {err}"
             else:
                 return "Error enviando tarea al agente."

        count_files = 0
        count_folders = 0
        
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Procesando carpetas...")

        total = len(data_rows)
        
        # Cachear lista de carpetas en root_path para búsqueda rápida normalizada
        try:
            root_subdirs = {clean_val(d): d for d in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, d))}
        except Exception as e:
            return f"Error leyendo carpeta raíz: {e}"

        for i, (folder_norm, folder_raw, suffix) in enumerate(data_rows):
            if not silent_mode and progress_bar:
                progress_bar.progress((i + 1) / total)
            
            # 1. Intentar match exacto (rápido)
            target_path = os.path.join(root_path, folder_raw)
            if not os.path.isdir(target_path):
                # 2. Intentar match normalizado
                matched_real_name = root_subdirs.get(folder_norm)
                if matched_real_name:
                    target_path = os.path.join(root_path, matched_real_name)
                else:
                    continue # No existe la carpeta
            
            # Procesar archivos en la carpeta encontrada
            count_folders += 1
            try:
                for f in os.listdir(target_path):
                    f_full = os.path.join(target_path, f)
                    if os.path.isfile(f_full):
                        name, ext = os.path.splitext(f)
                        # Evitar doble sufijo
                        if not name.endswith(suffix):
                            new_name = f"{name}{suffix}{ext}"
                            try:
                                os.rename(f_full, os.path.join(target_path, new_name))
                                count_files += 1
                            except: pass
            except: pass

        if not silent_mode and progress_bar: progress_bar.empty()
        msg = f"Proceso completado. {count_files} archivos renombrados en {count_folders} carpetas encontradas."
        
        if return_zip:
            try:
                mem_zip = io.BytesIO()
                with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, dirs, files in os.walk(root_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, root_path)
                            zf.write(file_path, arcname)
                
                try: shutil.rmtree(root_path, ignore_errors=True)
                except: pass
                
                return {
                    "files": [{
                        "name": f"Renombrados_Sufijo_{int(time.time())}.zip",
                        "data": mem_zip.getvalue(),
                        "label": "Descargar Renombrados (ZIP)"
                    }],
                    "message": msg
                }
            except Exception as e:
                return {"error": f"Error creando ZIP: {e}", "message": msg}
        
        return {"message": msg}
    except Exception as e:
        return {"error": f"Error crítico: {e}"}

# --- WORKERS: DOCX / FIRMAS ---

def worker_unificar_docx_carpeta(folder_path, output_name="Unificado.docx", silent_mode=False, return_zip=False):
    try:
        files = sorted([f for f in os.listdir(folder_path) if f.lower().endswith('.docx')], key=natural_sort_key)
        if not files: return {"error": "No hay archivos .docx"}
        master = Document(os.path.join(folder_path, files[0]))
        master.add_page_break()
        for f in files[1:]:
            doc = Document(os.path.join(folder_path, f))
            for element in doc.element.body:
                master.element.body.append(element)
            master.add_page_break()
        
        out_path = os.path.join(folder_path, output_name)
        master.save(out_path)
        msg = f"Unificados {len(files)} DOCX."
        
        if return_zip:
            try:
                # If unifying, maybe user wants just the result? 
                # Or the whole folder? Usually just the result if it's "Unify".
                # But let's zip the whole folder to be safe/flexible, or just the file.
                # Given "folder_path" might contain the source files, zipping all is safest backup.
                # OR if it's a temp folder from upload, we definitely want the result.
                # Let's zip just the output file if it's "Unify", but wait, user might want to see sources?
                # Usually "Unificar" -> "Descargar Resultado".
                
                # Let's zip the whole folder for consistency with other workers.
                mem_zip = io.BytesIO()
                with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, dirs, files in os.walk(folder_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, folder_path)
                            zf.write(file_path, arcname)
                
                try: shutil.rmtree(folder_path, ignore_errors=True)
                except: pass
                
                return {
                    "files": [{
                        "name": f"Unificado_DOCX_{int(time.time())}.zip",
                        "data": mem_zip.getvalue(),
                        "label": "Descargar Unificado (ZIP)"
                    }],
                    "message": msg
                }
            except Exception as e:
                return {"error": f"Error creando ZIP: {e}", "message": msg}
                
        return {"message": msg}
    except Exception as e:
        return {"error": f"Error: {e}"}

def worker_crear_firma_nombre(nombre, documento, output_folder=None, silent_mode=False):
    try:
        img = Image.new('RGB', (400, 100), color='white')
        d = ImageDraw.Draw(img)
        try: font = ImageFont.truetype("arial.ttf", 24)
        except: font = ImageFont.load_default()
        d.text((10, 10), f"Firmado por: {nombre}", fill='black', font=font)
        d.text((10, 50), f"Doc: {documento}", fill='black', font=font)
        
        filename = f"Firma_{documento}.png"
        
        if output_folder:
            out_path = os.path.join(output_folder, filename)
            img.save(out_path)
            return {"message": f"Firma guardada en {out_path}", "files": []}
            
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        return {
            "files": [{
                "name": filename,
                "data": img_byte_arr.getvalue(),
                "label": "Descargar Firma (PNG)"
            }],
            "message": "Firma creada exitosamente."
        }
    except Exception as e:
        return {"error": f"Error creando firma: {e}"}

def worker_firmar_docx(docx_path, firma_path, output_path=None, silent_mode=False):
    try:
        doc = Document(docx_path)
        doc.add_picture(firma_path, width=Pt(150))
        
        if output_path:
            doc.save(output_path)
            return {"message": "Documento firmado guardado.", "files": []}
            
        out_buffer = io.BytesIO()
        doc.save(out_buffer)
        
        return {
            "files": [{
                "name": f"Firmado_{int(time.time())}.docx",
                "data": out_buffer.getvalue(),
                "label": "Descargar Documento Firmado"
            }],
            "message": "Documento firmado exitosamente."
        }
    except Exception as e:
        return {"error": f"Error firmando documento: {e}"}

def worker_modificar_docx_excel(uploaded_file, sheet_name, col_folder, col_val, root_path, mode, silent_mode=False, return_zip=False):
    try:
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        modificados = 0
        docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)
        col_folder_idx = ord(col_folder.upper()) - ord('A')
        col_val_idx = ord(col_val.upper()) - ord('A')
        for index, row in df.iterrows():
            try:
                folder_name = str(row.iloc[col_folder_idx]).strip()
                new_val = str(row.iloc[col_val_idx]).strip()
                if not folder_name or not new_val: continue
                target_dir = os.path.join(root_path, folder_name)
                if not os.path.isdir(target_dir): continue
                target_docx = next((os.path.join(target_dir, f) for f in os.listdir(target_dir) if docx_pattern.match(f)), None)
                if target_docx:
                    doc = Document(target_docx)
                    modified = False
                    keyword = "REGIMEN:" if mode == "Regimen" else f"{mode}:"
                    for p in doc.paragraphs:
                        if keyword in p.text.upper():
                            p.text = re.sub(rf'({keyword})\s*.*', r'\1 ' + new_val, p.text, flags=re.IGNORECASE)
                            modified = True
                            break
                    if modified:
                        doc.save(target_docx)
                        modificados += 1
            except: pass
            
        msg = f"Modificados {modificados} documentos."
        
        if return_zip:
            try:
                mem_zip = io.BytesIO()
                with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, dirs, files in os.walk(root_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, root_path)
                            zf.write(file_path, arcname)
                return {
                    "files": [{
                        "name": f"Modificados_{mode}_{int(time.time())}.zip",
                        "data": mem_zip.getvalue(),
                        "label": "Descargar Modificados (ZIP)"
                    }],
                    "message": msg
                }
            except Exception as e:
                return {"error": f"Error creando ZIP: {e}", "message": msg}
        
        return {"message": msg}
    except Exception as e:
        return {"error": f"Error: {e}"}

def _create_column_map_from_headers(df):
    required_map = {
        'folder': 'Nombre Carpeta', 'date': 'Ciudad y Fecha', 'full_name': 'Nombre Completo',
        'doc_type': 'Tipo Documento', 'doc_num': 'Numero Documento', 'service': 'Servicio',
        'eps': 'EPS', 'tipo_servicio': 'Tipo Servicio', 'regimen': 'Regimen',
        'categoria': 'Categoria', 'cuota': 'Valor Cuota Moderadora', 'auth': 'Numero Autorizacion',
        'fecha_atencion': 'Fecha y Hora Atencion', 'fecha_fin': 'Fecha Finalizacion'
    }
    excel_headers = df.columns
    missing_cols = [header for header in required_map.values() if header not in excel_headers]
    if missing_cols: return None, missing_cols
    return required_map, []

def worker_modificar_docx_completo(uploaded_file, sheet_name, root_path, use_filter=False, silent_mode=False):
    try:
        is_native_mode = st.session_state.get('force_native_mode', True)
        should_delegate = is_native_mode and not silent_mode and _should_delegate(root_path)
        
        if isinstance(uploaded_file, bytes): uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)

        df = None
        if use_filter:
            # Re-read file if seekable, otherwise assuming it's fresh or handled
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
                
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada."
            ws = wb[sheet_name]
            
            data = []
            # Read headers from first row
            headers = [cell.value for cell in ws[1]]
            
            # Read visible rows
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    data.append([cell.value for cell in row])
            
            if data:
                df = pd.DataFrame(data, columns=headers)
            else:
                return "No hay datos visibles para procesar."
        else:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)

        column_map, missing = _create_column_map_from_headers(df)
        if not column_map:
            return f"Error: Faltan columnas en el Excel: {', '.join(missing)}"
            
        modificados = 0
        errores = 0
        
        if should_delegate:
            if not send_command:
                return "Error: Modo nativo activado pero cliente agente no disponible."
            
            username = st.session_state.get("username", "admin")
            
            tasks = []
            for index, row in df.iterrows():
                datos = {key: str(row[col_name]).strip() if pd.notna(row[col_name]) else "" for key, col_name in column_map.items()}
                folder_name = datos.get('folder')
                if not folder_name: continue
                
                tasks.append({
                    "rel_path": folder_name,
                    "datos": datos
                })
            
            if not tasks:
                return "No se encontraron tareas para procesar."
                
            if not silent_mode:
                st.info(f"Enviando {len(tasks)} tareas al agente local...")
                
            task_id = send_command(username, "fill_docx_ovida_full", {
                "base_path": root_path,
                "tasks": tasks
            })
            
            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=300)
                if not silent_mode: status_placeholder.empty()
                
                if res and isinstance(res, dict) and "error" not in res:
                    c = res.get("result", {}).get("count", 0) if "result" in res else res.get("count", 0)
                    errs = res.get("result", {}).get("errors", []) if "result" in res else res.get("errors", [])
                    return f"Proceso finalizado. Modificados: {c}, Errores: {len(errs)}"
                else:
                    error_msg = res.get('error', 'Desconocido') if isinstance(res, dict) else 'Respuesta no es un diccionario'
                    return f"Error del agente: {error_msg}"
            else:
                return "No se pudo crear la tarea."
        
        # Server execution
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Modificando DOCX...")
            
        for index, row in df.iterrows():
            if not silent_mode:
                progress_bar.progress((index + 1) / len(df), text=f"Procesando fila {index+1}")

            try:
                datos = {key: str(row[col_name]).strip() if pd.notna(row[col_name]) else "" for key, col_name in column_map.items()}
                folder_name = datos.get('folder')
                if not folder_name: continue
                
                target_dir = os.path.join(root_path, folder_name)
                if not os.path.isdir(target_dir): 
                    errores += 1
                    continue
                    
                target_docx = None
                if os.path.exists(target_dir):
                    for f in os.listdir(target_dir):
                        if f.lower().endswith('.docx') and not f.startswith('~'):
                            target_docx = os.path.join(target_dir, f)
                            break
                        
                if not target_docx:
                    errores += 1
                    continue
                
                doc = Document(target_docx)
                for p in doc.paragraphs:
                    if "Santiago de Cali, " in p.text: 
                        p.text = f"Santiago de Cali,  {datos['date']}"
                    
                    if "Yo " in p.text and "identificado con" in p.text:
                        p.text = f"Yo {datos['full_name']} identificado con {datos['doc_type']}, Numero {datos['doc_num']} en calidad de paciente, doy fé y acepto el servicio de {datos['service']} brindado por la IPS OPORTUNIDAD DE VIDA S.A.S"
                    
                    replacements = {
                        "EPS:": datos['eps'], "TIPO SERVICIO:": datos['tipo_servicio'],
                        "REGIMEN:": datos['regimen'], "CATEGORIA:": datos['categoria'],
                        "VALOR CUOTA MODERADORA:": datos['cuota'], "AUTORIZACION:": datos['auth'],
                        "Fecha de Atención:": datos['fecha_atencion'], "Fecha de Finalización:": datos['fecha_fin']
                    }
                    for key, val in replacements.items():
                        if key in p.text:
                            p.text = re.sub(rf'({key})\s*.*', r'\1 ' + str(val), p.text, count=1)
                
                sig_idx = -1
                for i, p in enumerate(doc.paragraphs):
                    if "FIRMA DE ACEPTACION" in p.text.upper():
                        sig_idx = i
                        break
                if sig_idx != -1 and sig_idx + 2 < len(doc.paragraphs):
                    doc.paragraphs[sig_idx + 2].text = datos['full_name'].upper()
                
                doc.save(target_docx)
                modificados += 1
            except Exception:
                errores += 1
        
        if not silent_mode: progress_bar.empty()
        return f"Proceso finalizado. Modificados: {modificados}, Errores/No encontrados: {errores}"
    except Exception as e:
        return f"Error general: {e}"

def worker_crear_carpetas_excel_avanzado(uploaded_file, sheet_name, col_name, base_path, use_filter=False, silent_mode=False):
    try:
        # --- NATIVE MODE CHECK ---
        is_native_mode = st.session_state.get('force_native_mode', True)

        if not is_native_mode and not os.path.isdir(base_path):
            return "La ruta base seleccionada no es válida."

        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        if hasattr(uploaded_file, 'seek'):
            uploaded_file.seek(0)
            
        nombres_carpetas = []
        if use_filter:
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames:
                return f"La hoja '{sheet_name}' no existe."
            ws = wb[sheet_name]
            
            # Find column index (1-based)
            header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
            col_idx = -1
            if header_row:
                for idx, val in enumerate(header_row):
                    if str(val).strip() == col_name:
                        col_idx = idx + 1
                        break
            
            if col_idx == -1:
                return f"No se encontró la columna '{col_name}' en la primera fila."
                
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    # Ensure row has enough columns
                    if col_idx - 1 < len(row):
                        val = row[col_idx-1].value
                        if val: nombres_carpetas.append(str(val))
        else:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            if col_name not in df.columns:
                return f"No se encontró la columna '{col_name}' en el Excel."
            nombres_carpetas = df[col_name].dropna().astype(str).tolist()

        if not nombres_carpetas:
            return "No se encontraron nombres de carpetas para crear."

        # --- EXECUTION ---
        if is_native_mode:
            if not send_command:
                return "Error: Modo nativo activado pero el cliente del agente no está disponible (send_command)."
            
            username = st.session_state.get("username", "admin")
            
            if not silent_mode:
                st.info(f"Enviando tarea al agente local para crear {len(nombres_carpetas)} carpetas...")
            
            task_id = send_command(username, "create_folders_from_list", {
                "base_path": base_path,
                "names": nombres_carpetas,
                "unique": True
            })
            
            if task_id:
                status_placeholder = st.empty()
                if not silent_mode:
                    status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=30)
                
                if not silent_mode: status_placeholder.empty()

                if res and "error" not in res:
                    count = res.get("count", 0)
                    errors = res.get("errors", [])
                    
                    msg = f"Proceso finalizado (Agente). Carpetas creadas: {count}"
                    if errors:
                        msg += f", Errores: {len(errors)}"
                        if not silent_mode:
                            st.error(f"Errores del agente: {'; '.join(errors[:5])}...")
                    return msg
                else:
                    err = res.get("error") if res else "Error desconocido o tiempo de espera agotado"
                    return f"Error del agente: {err}"
            else:
                return "Error enviando tarea: No se pudo crear la tarea."

        else:
            # Server-side execution
            creadas = 0
            errores = 0
            
            progress_bar = None
            if not silent_mode:
                progress_bar = st.progress(0, text="Creando carpetas...")
                
            for i, nombre in enumerate(nombres_carpetas):
                if not silent_mode:
                    progress_bar.progress((i + 1) / len(nombres_carpetas), text=f"Procesando: {nombre}")
                    
                nombre_base = "".join(c for c in nombre if c.isalnum() or c in " _-").rstrip()
                if not nombre_base: continue
                
                ruta_final = os.path.join(base_path, nombre_base)
                
                if os.path.exists(ruta_final):
                    contador = 2
                    nombre_consecutivo = f"{nombre_base} ({contador})"
                    ruta_final = os.path.join(base_path, nombre_consecutivo)
                    while os.path.exists(ruta_final):
                        contador += 1
                        nombre_consecutivo = f"{nombre_base} ({contador})"
                        ruta_final = os.path.join(base_path, nombre_consecutivo)
                
                try:
                    os.makedirs(ruta_final, exist_ok=True)
                    creadas += 1
                except Exception:
                    errores += 1
                    
            if not silent_mode: progress_bar.empty()
            return f"Proceso finalizado. Carpetas creadas: {creadas}, Errores: {errores}"
        
    except Exception as e:
        return f"Error crítico: {e}"

def worker_mover_archivos_por_coincidencia(base_path, silent_mode=False):
    if not base_path or not os.path.isdir(base_path):
        return "Ruta base inválida."
        
    try:
        elementos = os.listdir(base_path)
        archivos = [os.path.join(base_path, e) for e in elementos if os.path.isfile(os.path.join(base_path, e))]
        carpetas = [os.path.join(base_path, e) for e in elementos if os.path.isdir(os.path.join(base_path, e))]
    except Exception as e:
        return f"Error leyendo directorio: {e}"
        
    if not archivos or not carpetas:
        return "No hay archivos o carpetas suficientes para procesar."
        
    movidos, errores = 0, 0
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Moviendo archivos...")
        
    for i, ruta_archivo in enumerate(archivos):
        nombre_archivo = os.path.basename(ruta_archivo)
        if not silent_mode:
            progress_bar.progress((i + 1) / len(archivos), text=f"Verificando: {nombre_archivo}")
            
        for ruta_carpeta in carpetas:
            nombre_carpeta = os.path.basename(ruta_carpeta)
            if nombre_carpeta.lower() in nombre_archivo.lower():
                try:
                    shutil.move(ruta_archivo, ruta_carpeta)
                    movidos += 1
                    break
                except Exception:
                    errores += 1
                    break
                    
    if not silent_mode: progress_bar.empty()
    return f"Proceso finalizado. Movidos: {movidos}, Errores: {errores}"

def worker_anadir_sufijo_desde_excel(uploaded_file, sheet_name, col_folder, col_suffix, base_path, use_filter=False, silent_mode=False, item_type="both"):
    try:
        data_rows = []
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        if hasattr(uploaded_file, 'seek'):
            uploaded_file.seek(0)

        # --- NATIVE MODE CHECK ---
        is_native_mode = st.session_state.get('force_native_mode', True)

        if use_filter:
            # Re-read file if seekable, otherwise assuming it's fresh or handled
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
                
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada"
            ws = wb[sheet_name]
            
            header = [cell.value for cell in ws[1]]
            try:
                idx_folder = header.index(col_folder)
                idx_suffix = header.index(col_suffix)
            except ValueError: return "Columnas no encontradas en encabezado"
            
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    val_folder = row[idx_folder].value
                    val_suffix = row[idx_suffix].value
                    if val_folder and val_suffix:
                        data_rows.append({"folder": str(val_folder).strip(), "suffix": str(val_suffix).strip()})
        else:
            # Re-read file if seekable
            if hasattr(uploaded_file, 'seek'):
                uploaded_file.seek(0)
                
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            if col_folder not in df.columns or col_suffix not in df.columns:
                return f"Columnas '{col_folder}' o '{col_suffix}' no encontradas."
            for _, row in df.iterrows():
                if pd.notna(row[col_folder]) and pd.notna(row[col_suffix]):
                    data_rows.append({"folder": str(row[col_folder]).strip(), "suffix": str(row[col_suffix]).strip()})

        if not data_rows:
            return "No se encontraron filas válidas para procesar."

        # --- AGENT EXECUTION ---
        if is_native_mode:
            if not send_command:
                return "Error: Modo nativo activado pero cliente agente no disponible."
            
            username = st.session_state.get("username", "admin")
            if not silent_mode:
                st.info(f"Enviando tarea al agente local para procesar {len(data_rows)} items...")

            # Map to bulk_rename format
            items_for_agent = [{"key": r["folder"], "suffix": r["suffix"]} for r in data_rows]
            
            task_id = send_command(username, "bulk_rename", {
                "path": base_path,
                "items": items_for_agent,
                "separator": "", # Suffix in excel usually includes separator or logic handles it? Logic below uses direct concatenation.
                # Actually server logic below does: new_name = f"{name}{suffix}{ext}" -> No separator added!
                # But bulk_rename usually expects separator.
                # Let's use empty separator here to match server logic, OR assume user provides suffix with separator.
                "item_type": item_type
            })

            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                # Timeout razonable para operaciones de renombrado masivo
                res = wait_for_result(task_id, timeout=300)
                
                if not silent_mode: status_placeholder.empty()

                if res and "result" in res:
                    r = res["result"]
                    count = r.get("count", 0)
                    errs = r.get("errors", [])
                    return f"Finalizado. Renombrados: {count}, Errores: {len(errs)}"
                elif res and "error" in res:
                     return f"Error del agente: {res.get('error')}"
                else:
                     return "Respuesta inválida del agente."
            else:
                return "No se pudo crear la tarea en el servidor."

        # --- SERVER EXECUTION ---
        carpetas_procesadas, archivos_renombrados, errores = 0, 0, 0
        
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Añadiendo sufijos...")
            
        total = len(data_rows)
        
        scope_folders = str(item_type).lower() in ["todo", "both", "carpetas", "folders", "directory"]
        scope_files = str(item_type).lower() in ["todo", "both", "archivos", "files", "file"]

        # Helper to resolve path using search results
        def _resolve_path(name, base):
            if "search_results" in st.session_state and st.session_state.search_results:
                name_norm = str(name).strip().lower()
                for res in st.session_state.search_results:
                    r_name = str(res.get("name", res.get("Nombre", ""))).strip().lower()
                    r_path = res.get("path", res.get("Ruta completa", ""))
                    if r_name == name_norm and r_path:
                        return r_path
            return os.path.join(base, name)

        for index, item in enumerate(data_rows):
            folder_name = item["folder"]
            suffix = item["suffix"]
            
            if not silent_mode:
                progress_bar.progress((index + 1) / total)
            
            target_path = _resolve_path(folder_name, base_path)
            
            # 1. RENAME FOLDER ITSELF (If scope_folders)
            if scope_folders and os.path.isdir(target_path):
                if not folder_name.endswith(suffix): # Simple check, suffix might contain separator
                     new_folder_name = f"{folder_name}{suffix}"
                     # Use dirname of target_path to reconstruct new path
                     parent_dir = os.path.dirname(target_path)
                     new_folder_path = os.path.join(parent_dir, new_folder_name)
                     try:
                         os.rename(target_path, new_folder_path)
                         carpetas_procesadas += 1
                         target_path = new_folder_path # Update for file processing
                     except Exception:
                         errores += 1
            
            # 2. RENAME FILES INSIDE (If scope_files)
            if scope_files:
                if os.path.isdir(target_path):
                    # It's a folder, rename files inside
                    for fname in os.listdir(target_path):
                        fpath = os.path.join(target_path, fname)
                        if os.path.isfile(fpath):
                            name, ext = os.path.splitext(fname)
                            new_name = f"{name}{suffix}{ext}"
                            new_path = os.path.join(target_path, new_name)
                            
                            if fpath == new_path: continue
                            if os.path.exists(new_path):
                                errores += 1
                                continue
                                
                            try:
                                os.rename(fpath, new_path)
                                archivos_renombrados += 1
                            except Exception:
                                errores += 1
                elif os.path.isfile(target_path):
                     # It's a specific file (found via search or direct path)
                     # Rename the file itself
                     parent_dir = os.path.dirname(target_path)
                     fname = os.path.basename(target_path)
                     name, ext = os.path.splitext(fname)
                     
                     if not name.endswith(suffix):
                         new_name = f"{name}{suffix}{ext}"
                         new_path = os.path.join(parent_dir, new_name)
                         try:
                             os.rename(target_path, new_path)
                             archivos_renombrados += 1
                         except Exception:
                             errores += 1
            elif not os.path.isdir(target_path) and not (scope_files and os.path.isfile(target_path)):
                 # Only count error if it wasn't handled above
                 if not os.path.exists(target_path):
                    errores += 1 # Path not found

                        
        if not silent_mode: progress_bar.empty()
        return f"Finalizado. Carpetas Renombradas: {carpetas_procesadas}, Archivos Renombrados: {archivos_renombrados}, Errores: {errores}"
    except Exception as e:
        return f"Error crítico: {e}"

# --- WORKERS: MISSING FILE OPS ---

def worker_crear_carpetas_desde_excel(excel_path, sheet_name, col_name, output_folder, filter_hidden=False, silent_mode=False):
    try:
        if filter_hidden:
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            ws = wb[sheet_name]
            df_col_idx = -1
            for idx, cell in enumerate(ws[1]):
                if cell.value == col_name:
                    df_col_idx = idx
                    break
            if df_col_idx == -1: return "Columna no encontrada."
            names = []
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    val = row[df_col_idx].value
                    if val: names.append(str(val))
        else:
            df = pd.read_excel(excel_path, sheet_name=sheet_name)
            if col_name not in df.columns: return "Columna no encontrada."
            names = df[col_name].dropna().astype(str).tolist()
            
        count = 0
        errores = 0
        for name in names:
            safe_name = "".join(c for c in name if c.isalnum() or c in " _-").strip()
            if not safe_name: continue
            target = os.path.join(output_folder, safe_name)
            if os.path.exists(target):
                i = 2
                while os.path.exists(f"{target} ({i})"): i += 1
                target = f"{target} ({i})"
            try:
                os.makedirs(target)
                count += 1
            except: errores += 1
        return f"Creadas {count} carpetas. Errores: {errores}"
    except Exception as e:
        return f"Error: {e}"

def worker_copiar_archivos_desde_mapeo(excel_path, sheet_name, col_src, col_dst, base_src, base_dst, silent_mode=False):
    try:
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        if col_src not in df.columns or col_dst not in df.columns: return "Columnas no encontradas."
        
        copiados = 0
        conflictos = 0
        errores = 0
        
        for _, row in df.iterrows():
            src_folder = str(row[col_src]).strip()
            dst_folder = str(row[col_dst]).strip()
            if not src_folder or not dst_folder: continue
            
            full_src = os.path.join(base_src, src_folder)
            full_dst = os.path.join(base_dst, dst_folder)
            
            if not os.path.isdir(full_src) or not os.path.isdir(full_dst): continue
            
            for f in os.listdir(full_src):
                f_src = os.path.join(full_src, f)
                if os.path.isfile(f_src):
                    f_dst = os.path.join(full_dst, f)
                    if os.path.exists(f_dst):
                        conflictos += 1
                        continue
                    try:
                        shutil.copy2(f_src, f_dst)
                        copiados += 1
                    except: errores += 1
        return f"Copiados {copiados}. Conflictos: {conflictos}. Errores: {errores}."
    except Exception as e:
        return f"Error: {e}"





def worker_txt_a_json_masivo(folder_path, silent_mode=False):
    try:
        count = 0
        for f in os.listdir(folder_path):
            if f.lower().endswith('.txt'):
                try:
                    base = os.path.splitext(f)[0]
                    src = os.path.join(folder_path, f)
                    dst = os.path.join(folder_path, base + ".json")
                    if not os.path.exists(dst):
                        os.rename(src, dst)
                        count += 1
                except: pass
        return f"Renombrados {count} TXT a JSON."
    except Exception as e:
        return f"Error: {e}"


def worker_analisis_carpetas(root_path, silent_mode=False):
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(root_path):
        if not silent_mode: st.info(f"Delegando análisis de carpetas al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_carpetas", {"path": root_path})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    if not silent_mode: st.info(f"Analizando: {root_path}")
    data_summary = []
    data_details = []
    
    for root, dirs, files in os.walk(root_path):
        folder_name = os.path.basename(root)
        count = len(files)
        
        # Calculate size safely handling inaccessible/missing files
        size = 0
        for f in files:
            try:
                full_path = os.path.join(root, f)
                if os.path.exists(full_path) and os.path.isfile(full_path):
                    size += os.path.getsize(full_path)
            except Exception:
                pass
                
        data_summary.append({"Carpeta": folder_name, "Archivos": count, "Peso (KB)": round(size/1024, 2), "Ruta": root})
        
        for f in files:
            file_size = 0
            try:
                full_path = os.path.join(root, f)
                if os.path.exists(full_path) and os.path.isfile(full_path):
                    file_size = os.path.getsize(full_path)
            except Exception:
                pass
                
            data_details.append({
                "Carpeta Principal": folder_name,
                "Ruta": os.path.join(root, f),
                "Archivo": f,
                "Peso (KB)": round(file_size/1024, 2)
            })
    
    if not data_summary:
        if not silent_mode: st.warning("Carpeta vacía o sin acceso.")
        return None

    df_summary = pd.DataFrame(data_summary)
    df_details = pd.DataFrame(data_details)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_details.to_excel(writer, index=False, sheet_name='Detalle')
        df_summary.to_excel(writer, index=False, sheet_name='Resumen')
    
    return {
        "files": [{
            "name": f"Reporte_Carpetas_{int(time.time())}.xlsx",
            "data": output.getvalue(),
            "label": "Descargar Reporte"
        }],
        "message": f"Reporte generado con {len(data_summary)} carpetas analizadas."
    }

# --- WORKERS: CONVERSION & AI ---

def _pdf_a_docx(input_path, output_path):
    try:
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()
    except Exception as e:
        print(f"Error pdf2docx: {e}")

def _jpg_a_pdf(input_path, output_path):
    img = Image.open(input_path)
    if img.mode == 'RGBA':
        img = img.convert('RGB')
    res = st.session_state.app_config.get("image_resolution", 100.0) if 'app_config' in st.session_state else 100.0
    img.save(output_path, "PDF", resolution=res)

def _docx_a_pdf(input_path, output_path):
    """
    Convierte DOCX a PDF usando automatización COM directa (win32com) si está disponible.
    Fallback a docx2pdf.
    """
    success = False
    if HAS_WIN32COM:
        try:
            import pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            in_file_abs = os.path.abspath(input_path)
            out_file_abs = os.path.abspath(output_path)
            
            doc = word.Documents.Open(in_file_abs)
            doc.SaveAs(out_file_abs, FileFormat=17) # wdFormatPDF = 17
            doc.Close(False)
            word.Quit()
            success = True
        except Exception as e:
            print(f"Error win32com DOCX->PDF: {e}. Intentando fallback...")
    
    if not success and HAS_DOCX2PDF:
        try:
            # docx2pdf sometimes needs pythoncom init too if running in thread
            try:
                import pythoncom
                pythoncom.CoInitialize() 
            except: pass
            convert_docx_to_pdf(input_path, output_path)
        except Exception as e:
            print(f"Error docx2pdf: {e}")

def _pdf_a_jpg(input_path, output_base):
    doc = fitz.open(input_path)
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        out = f"{output_base}_p{i+1}.jpg" if len(doc) > 1 else f"{output_base}.jpg"
        pix.save(out)
    doc.close()

def _png_a_jpg(input_path, output_path):
    img = Image.open(input_path)
    rgb_img = img.convert('RGB')
    rgb_img.save(output_path, 'jpeg')

def _txt_a_json(input_path, output_path):
    if input_path == output_path: return
    if not os.path.exists(output_path):
        os.rename(input_path, output_path)
    else:
        base, ext = os.path.splitext(output_path)
        new_out = f"{base}_{int(time.time())}.json"
        os.rename(input_path, new_out)

def _pdf_escala_grises(input_path, output_path):
    doc = fitz.open(input_path)
    doc_final = fitz.open()
    dpi = st.session_state.app_config.get("pdf_dpi", 600) if 'app_config' in st.session_state else 600
    matrix_scale = dpi / 72.0
    mat = fitz.Matrix(matrix_scale, matrix_scale)
    for page in doc:
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY)
        new_page = doc_final.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_image(new_page.rect, pixmap=pix)
    doc.close()
    compression = st.session_state.app_config.get("pdf_compression", 4) if 'app_config' in st.session_state else 4
    doc_final.save(output_path, garbage=compression, deflate=True)
    doc_final.close()

def worker_convertir_archivo(file_path, tipo, output_folder=None, silent_mode=False):
    if not file_path or not os.path.exists(file_path):
        return False, "Archivo no encontrado"
    folder = output_folder if output_folder else os.path.dirname(file_path)
    filename = os.path.basename(file_path)
    name_no_ext = os.path.splitext(filename)[0]
    try:
        if tipo == "PDF2DOCX":
            out = os.path.join(folder, f"{name_no_ext}.docx")
            _pdf_a_docx(file_path, out)
        elif tipo == "JPG2PDF":
            out = os.path.join(folder, f"{name_no_ext}.pdf")
            _jpg_a_pdf(file_path, out)
        elif tipo == "DOCX2PDF":
            out = os.path.join(folder, f"{name_no_ext}.pdf")
            _docx_a_pdf(file_path, out)
        elif tipo == "PDF2JPG":
            out_base = os.path.join(folder, name_no_ext)
            _pdf_a_jpg(file_path, out_base)
        elif tipo == "PNG2JPG":
            out = os.path.join(folder, f"{name_no_ext}.jpg")
            _png_a_jpg(file_path, out)
        elif tipo == "TXT2JSON":
            out = os.path.join(folder, f"{name_no_ext}.json")
            _txt_a_json(file_path, out)
        elif tipo == "PDF_GRAY":
            temp_out = os.path.join(folder, f"{name_no_ext}_temp_gray.pdf")
            _pdf_escala_grises(file_path, temp_out)
            if os.path.exists(temp_out):
                try:
                    os.replace(temp_out, file_path)
                except OSError:
                    time.sleep(0.5)
                    os.remove(file_path)
                    os.rename(temp_out, file_path)
        return True, "Conversión exitosa"
    except Exception as e:
        return False, str(e)

def worker_convertir_masivo(folder_path, tipo, silent_mode=False):
    if not folder_path or not os.path.exists(folder_path):
        return 0, "Carpeta no encontrada"
    count = 0
    files_to_process = []
    for r, d, f in os.walk(folder_path):
        for file in f:
            files_to_process.append(os.path.join(r, file))
    total = len(files_to_process)
    if total == 0:
        return 0, "Carpeta vacía"
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Convirtiendo...")
    for i, full_path in enumerate(files_to_process):
        if not silent_mode and i % 5 == 0: 
            progress_bar.progress(min(i/total, 1.0), text=f"Procesando {i}/{total}")
        f = os.path.basename(full_path)
        f_lower = f.lower()
        process = False
        if tipo == "PDF2DOCX" and f_lower.endswith(".pdf"): process = True
        elif tipo == "JPG2PDF" and (f_lower.endswith(".jpg") or f_lower.endswith(".jpeg")): process = True
        elif tipo == "DOCX2PDF" and f_lower.endswith(".docx") and not f.startswith("~$"): process = True
        elif tipo == "PDF2JPG" and f_lower.endswith(".pdf"): process = True
        elif tipo == "PNG2JPG" and f_lower.endswith(".png"): process = True
        elif tipo == "TXT2JSON" and f_lower.endswith(".txt"): process = True
        elif tipo == "PDF_GRAY" and f_lower.endswith(".pdf"): process = True
        if process:
            ok, msg = worker_convertir_archivo(full_path, tipo, silent_mode=True)
            if ok: count += 1
            else: 
                if not silent_mode: print(f"Error convirtiendo {f}: {msg}")
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
    return count, f"Procesados {count} archivos."

def worker_consultar_gemini(prompt, file_context=None, silent_mode=False):
    api_key = st.session_state.app_config.get("gemini_api_key") if 'app_config' in st.session_state else None
    if not api_key: return "⚠️ Configura tu API Key de Google Gemini."
    try:
        genai.configure(api_key=api_key.strip())
        model_name = st.session_state.app_config.get("gemini_model", "gemini-1.5-flash") if 'app_config' in st.session_state else "gemini-1.5-flash"
        model = genai.GenerativeModel(model_name)
        full_prompt = f"Contexto:\n{file_context}\n\n{prompt}" if file_context else prompt
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        return f"Error Gemini: {e}"

# --- WORKERS: MAPEO / OTROS ---

def worker_copiar_raiz_mapeo(uploaded_file, sheet_name, col_id, col_dst, path_src_base, path_dst_base, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        count_files = 0
        files_in_root = {f.lower(): f for f in os.listdir(path_src_base) if os.path.isfile(os.path.join(path_src_base, f))}
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Copiando...")
        total_rows = len(df)
        for idx, row in df.iterrows():
            if not silent_mode and idx % 10 == 0 and total_rows > 0:
                progress_bar.progress(min(idx / total_rows, 1.0), text=f"Procesando {idx}/{total_rows}")
            id_val = str(row[col_id]).strip().lower()
            dst_folder_name = str(row[col_dst]).strip()
            if not id_val or not dst_folder_name: continue
            for f_lower, f_real in files_in_root.items():
                if id_val in f_lower:
                    src = os.path.join(path_src_base, f_real)
                    dst_folder = os.path.join(path_dst_base, dst_folder_name)
                    if not os.path.exists(dst_folder):
                        try: os.makedirs(dst_folder)
                        except: pass
                    try:
                        shutil.copy2(src, os.path.join(dst_folder, f_real))
                        count_files += 1
                    except Exception: pass
        msg = f"Copia completada. {count_files} archivos copiados."
        if not silent_mode:
            if progress_bar: progress_bar.progress(1.0, text="Finalizado.")
            st.success(msg)
        return msg
    except Exception as e:
        return f"Error: {e}"

def worker_exportar_renombrado(search_results, silent_mode=False):
    if not search_results: return None
    data = []
    for item in search_results:
        path = item.get("Ruta completa", "")
        if path:
            data.append({"Ruta actual": path, "Nuevo nombre": os.path.basename(path)})
    return pd.DataFrame(data)

def worker_renombrar_mapeo_excel(uploaded_file, sheet_name, col_src, col_dst, use_filter, root_path=None, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        if hasattr(uploaded_file, 'seek'):
            uploaded_file.seek(0)
        data_rows = []
        if use_filter:
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada."
            ws = wb[sheet_name]
            header = [cell.value for cell in ws[1]]
            try:
                idx_src = header.index(col_src)
                idx_dst = header.index(col_dst)
            except: return "Columnas no encontradas."
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    val_src = row[idx_src].value
                    val_dst = row[idx_dst].value
                    if val_src and val_dst:
                        data_rows.append((str(val_src).strip(), str(val_dst).strip()))
        else:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
            if col_src not in df.columns or col_dst not in df.columns: return "Columnas no encontradas."
            for _, row in df.iterrows():
                if pd.notna(row[col_src]) and pd.notna(row[col_dst]):
                    data_rows.append((str(row[col_src]).strip(), str(row[col_dst]).strip()))
        is_native_mode = st.session_state.get('force_native_mode', True)
        if is_native_mode:
            if not send_command:
                return "Error: Modo nativo activado pero cliente agente no disponible."
            
            username = st.session_state.get("username", "admin")
            
            files_to_rename = []
            for src_name, dst_name in data_rows:
                src_path = os.path.join(root_path, src_name)
                dst_path = os.path.join(root_path, dst_name)
                files_to_rename.append({"old_path": src_path, "new_path": dst_path})
                
            if not files_to_rename:
                return "No hay archivos para renombrar."
                
            task_id = send_command(username, "rename_files", {"files": files_to_rename})
            
            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=300)
                if not silent_mode: status_placeholder.empty()
                
                if res and "error" not in res:
                    c = res.get("count", 0)
                    errs = res.get("errors", [])
                    return f"Renombrados {c} archivos. Errores: {len(errs)}"
                else:
                    return f"Error del agente: {res.get('error', 'Desconocido') if res else 'Desconocido'}"
            else:
                return "No se pudo crear la tarea."

        # Server Execution
        count = 0
        progress_bar = None
        if not silent_mode: progress_bar = st.progress(0, text="Renombrando...")
        for i, (src_name, dst_name) in enumerate(data_rows):
            if not silent_mode: progress_bar.progress(min((i+1)/len(data_rows), 1.0))
            src_path = os.path.join(root_path, src_name)
            dst_path = os.path.join(root_path, dst_name)
            if os.path.exists(src_path) and src_path != dst_path:
                try:
                    os.rename(src_path, dst_path)
                    count += 1
                except: pass
        if not silent_mode: st.success(f"Renombrados {count} archivos.")
        return f"Renombrados {count} archivos."
    except Exception as e: return f"Error: {e}"



def worker_copiar_mapeo_subcarpetas(uploaded_file, sheet_name, col_src, col_dst, path_src_base, path_dst_base, use_filter=False, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes): uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        
        df = None
        if use_filter:
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada."
            ws = wb[sheet_name]
            
            data = []
            headers = [cell.value for cell in ws[1]]
            
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    data.append([cell.value for cell in row])
            
            if data:
                df = pd.DataFrame(data, columns=headers)
            else:
                return "No hay datos visibles."
        else:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)

        if col_src not in df.columns or col_dst not in df.columns:
            return f"Columnas no encontradas: {col_src}, {col_dst}"

        copiados = 0
        for _, row in df.iterrows():
            src_n = str(row[col_src]).strip()
            dst_n = str(row[col_dst]).strip()
            if not src_n or not dst_n: continue
            src_full = os.path.join(path_src_base, src_n)
            dst_full = os.path.join(path_dst_base, dst_n)
            
            if os.path.isdir(src_full):
                os.makedirs(dst_full, exist_ok=True)
                for f in os.listdir(src_full):
                    src_f = os.path.join(src_full, f)
                    dst_f = os.path.join(dst_full, f)
                    if os.path.isfile(src_f) and not os.path.exists(dst_f):
                        try:
                            shutil.copy2(src_f, dst_f)
                            copiados += 1
                        except: pass
        return f"Copiados: {copiados}"
    except Exception as e: return f"Error: {e}"

# --- WORKERS: FIRMA DIGITAL & CONSOLIDACION ---

def _crear_firma_estilizada(texto):
    """
    Crea una firma digital estilizada sin usar fuentes tipográficas.
    Convierte cada letra del texto en un trazo manuscrito único.
    """
    width = max(400, len(texto) * 40)
    height = 150
    image = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(image)
    colores = ['black', 'gray', 'darkgray']

    def _dibujar_trazo_vocal(draw, x_base, y_centro, ascii_val, colores):
        color = random.choice(colores)
        grosor = random.randint(2, 3)
        altura_arco = 20 + (ascii_val % 15)
        puntos = []
        for i in range(20):
            angulo = (i / 19.0) * math.pi
            x = x_base + i * 2
            y = y_centro - math.sin(angulo) * altura_arco + random.randint(-2, 2)
            puntos.append((x, y))
        for i in range(len(puntos) - 1):
            draw.line([puntos[i], puntos[i + 1]], fill=color, width=grosor)

    def _dibujar_trazo_consonante_dura(draw, x_base, y_centro, ascii_val, colores):
        color = random.choice(colores)
        grosor = random.randint(3, 4)
        x = x_base
        y = y_centro + random.randint(-10, 10)
        draw.line([(x, y), (x + 15, y - 20)], fill=color, width=grosor)
        draw.line([(x + 15, y - 20), (x + 30, y - 15)], fill=color, width=grosor)
        draw.line([(x + 30, y - 15), (x + 40, y + 10)], fill=color, width=grosor)

    def _dibujar_trazo_generico(draw, x_base, y_centro, ascii_val, colores):
        color = random.choice(colores)
        grosor = random.randint(2, 3)
        puntos = []
        for i in range(30):
            x = x_base + i * 1.5
            onda = math.sin((x - x_base) * 0.2 + ascii_val * 0.1) * 15
            y = y_centro + onda + random.randint(-3, 3)
            puntos.append((x, y))
        for i in range(len(puntos) - 1):
            draw.line([puntos[i], puntos[i + 1]], fill=color, width=grosor)
        for _ in range(random.randint(2, 4)):
            punto_x = random.randint(int(x_base), int(x_base + 40))
            punto_y = int(y_centro + random.randint(-5, 5))
            draw.ellipse([punto_x - 1, punto_y - 1, punto_x + 1, punto_y + 1], fill=color)

    for i, letra in enumerate(texto):
        if letra.isspace(): continue
        x_base = 30 + (i * (width - 60) // len(texto))
        y_centro = height // 2
        ascii_val = ord(letra.upper()) if letra.isalpha() else ord('A')
        if letra.upper() in 'AEIOU':
            _dibujar_trazo_vocal(draw, x_base, y_centro, ascii_val, colores)
        elif letra.upper() in 'BCDFG':
            _dibujar_trazo_consonante_dura(draw, x_base, y_centro, ascii_val, colores)
        else:
            _dibujar_trazo_generico(draw, x_base, y_centro, ascii_val, colores)
            
    return image

def worker_crear_firma_digital(base_path, font_path, font_size, silent_mode=False):
    if not base_path or not font_path: return "Error: Rutas inválidas."
    try:
        font = ImageFont.truetype(font_path, font_size)
    except Exception as e:
        return f"Error cargando fuente: {e}"
        
    try:
        folders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
    except Exception as e:
        return f"Error leyendo carpetas: {e}"
        
    if not folders: return "No hay carpetas."
    
    creadas, errores = 0, 0
    progress_bar = None
    if not silent_mode: progress_bar = st.progress(0, text="Creando firmas...")
    
    for i, folder_name in enumerate(folders):
        if not silent_mode: progress_bar.progress((i + 1) / len(folders))
        try:
            text_to_draw = folder_name
            temp_img = Image.new('RGB', (1, 1))
            draw_temp = ImageDraw.Draw(temp_img)
            bbox = draw_temp.textbbox((0, 0), text_to_draw, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            padding = 20
            img_width = text_width + (2 * padding)
            img_height = text_height + (2 * padding)
            
            image = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(image)
            draw.text((padding, padding), text_to_draw, font=font, fill='black')
            
            tipografia_folder = os.path.join(base_path, folder_name, "tipografia")
            os.makedirs(tipografia_folder, exist_ok=True)
            image.save(os.path.join(tipografia_folder, "firma.jpg"), 'JPEG', quality=95)
            creadas += 1
        except:
            errores += 1
            
    if not silent_mode: progress_bar.empty()
    return f"Firmas creadas: {creadas}. Errores: {errores}."

def worker_consolidar_archivos_subcarpetas(base_path, silent_mode=False):
    if not base_path or not os.path.isdir(base_path): return "Ruta inválida."
    
    try:
        main_folders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
    except Exception as e: return f"Error leyendo base: {e}"
    
    if not main_folders: return "No hay carpetas."
    
    copiados, conflictos, errores = 0, 0, 0
    progress_bar = None
    if not silent_mode: progress_bar = st.progress(0, text="Consolidando...")
    
    for i, folder_name in enumerate(main_folders):
        if not silent_mode: progress_bar.progress((i + 1) / len(main_folders))
        main_folder_path = os.path.join(base_path, folder_name)
        
        for sub_root, _, files in os.walk(main_folder_path):
            if sub_root == main_folder_path: continue
            for file_name in files:
                src = os.path.join(sub_root, file_name)
                dst = os.path.join(main_folder_path, file_name)
                try:
                    if os.path.exists(dst):
                        conflictos += 1
                        continue
                    shutil.copy2(src, dst)
                    copiados += 1
                except: errores += 1
                
    if not silent_mode: progress_bar.empty()
    return f"Copiados: {copiados}. Conflictos: {conflictos}. Errores: {errores}."

def worker_copiar_archivos_desde_raiz_mapeo(uploaded_file, sheet_name, col_id, col_folder, root_src, root_dst, use_filter=False, silent_mode=False):
    try:
        if isinstance(uploaded_file, bytes): uploaded_file = io.BytesIO(uploaded_file)
        uploaded_file.seek(0)
        
        df = None
        if use_filter:
            import openpyxl
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada."
            ws = wb[sheet_name]
            
            data = []
            headers = [cell.value for cell in ws[1]]
            
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    data.append([cell.value for cell in row])
            
            if data:
                df = pd.DataFrame(data, columns=headers)
            else:
                return "No hay datos visibles."
        else:
            df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        
        archivos_origen = [f for f in os.listdir(root_src) if os.path.isfile(os.path.join(root_src, f))]
        copiados, no_encontrados, conflictos, errores = 0, 0, 0, 0
        
        progress_bar = None
        if not silent_mode: progress_bar = st.progress(0, text="Copiando...")
        
        for i, row in df.iterrows():
            if not silent_mode: progress_bar.progress((i + 1) / len(df))
            
            id_val = str(row[col_id]).strip()
            folder_val = str(row[col_folder]).strip()
            if not id_val or not folder_val: continue
            
            found_file = None
            for f in archivos_origen:
                if id_val.lower() in f.lower():
                    found_file = f
                    break
            
            if not found_file:
                no_encontrados += 1
                continue
                
            dst_dir = os.path.join(root_dst, folder_val)
            os.makedirs(dst_dir, exist_ok=True)
            dst_file = os.path.join(dst_dir, found_file)
            
            if os.path.exists(dst_file):
                conflictos += 1
                continue
                
            try:
                shutil.copy2(os.path.join(root_src, found_file), dst_file)
                copiados += 1
            except: errores += 1
            
        if not silent_mode: progress_bar.empty()
        return f"Copiados: {copiados}. No encontrados: {no_encontrados}. Conflictos: {conflictos}. Errores: {errores}."
    except Exception as e: return f"Error: {e}"

def worker_copiar_archivo_a_subcarpetas(file_path, dest_base_path, silent_mode=False):
    if not file_path or not dest_base_path: return "Rutas inválidas."
    
    try:
        subcarpetas = [os.path.join(dest_base_path, d) for d in os.listdir(dest_base_path) if os.path.isdir(os.path.join(dest_base_path, d))]
    except Exception as e: return f"Error leyendo destinos: {e}"
    
    if not subcarpetas: return "No hay subcarpetas."
    
    copiados, conflictos, errores = 0, 0, 0
    fname = os.path.basename(file_path)
    
    progress_bar = None
    if not silent_mode: progress_bar = st.progress(0, text=f"Copiando {fname}...")
    
    for i, sub in enumerate(subcarpetas):
        if not silent_mode: progress_bar.progress((i + 1) / len(subcarpetas))
        dst = os.path.join(sub, fname)
        if os.path.exists(dst):
            conflictos += 1
            continue
        try:
            shutil.copy2(file_path, dst)
            copiados += 1
        except: errores += 1
        
    if not silent_mode: progress_bar.empty()
    return f"Copiados: {copiados}. Conflictos: {conflictos}. Errores: {errores}."

def worker_descargar_historias_hospitalizacion_ovida(uploaded_file, sheet_name, col_map, save_path=None, silent_mode=False):
    # Requires Selenium
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from webdriver_manager.chrome import ChromeDriverManager
        import base64
        import zipfile
    except ImportError: return {"error": "Falta Selenium/WebDriverManager."}

    try:
        if isinstance(uploaded_file, bytes): uploaded_file = io.BytesIO(uploaded_file)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
    except Exception as e: return {"error": f"Error Excel: {e}"}
    
    driver = None
    
    # --- NATIVE MODE CHECK ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    
    try:
        # Determine output directory
        is_temp = False
        if not save_path:
            is_temp = True
            save_path = os.path.join(os.getcwd(), "temp_downloads", f"ovida_{int(time.time())}")
            # En server mode creamos la carpeta. En native mode, el agente lo hará (o usaremos ruta temporal allá si se implementara)
            if not is_native_mode:
                os.makedirs(save_path, exist_ok=True)
        elif not os.path.exists(save_path) and not is_native_mode:
             os.makedirs(save_path, exist_ok=True)

        # --- AGENT EXECUTION ---
        if is_native_mode:
            try:
                # Prepare records for agent
                records = []
                for _, row in df.iterrows():
                    try:
                        # Map columns
                        estudio = str(int(row[col_map['estudio']])).strip() if col_map['estudio'] in row else ""
                        
                        ingreso = row[col_map['ingreso']] if col_map['ingreso'] in row else ""
                        if isinstance(ingreso, pd.Timestamp):
                            ingreso = ingreso.strftime('%Y/%m/%d')
                        else:
                             ingreso = str(ingreso)
                             
                        egreso = row[col_map['egreso']] if col_map['egreso'] in row else ""
                        if isinstance(egreso, pd.Timestamp):
                            egreso = egreso.strftime('%Y/%m/%d')
                        else:
                             egreso = str(egreso)
                        
                        carpeta = str(row[col_map['carpeta']]).strip() if col_map['carpeta'] in row else ""

                        records.append({
                            "nro_estudio": estudio,
                            "fecha_ingreso": ingreso,
                            "fecha_salida": egreso,
                            "rel_path": carpeta
                        })
                    except Exception as row_err:
                        print(f"Error procesando fila hospitalización para agente: {row_err}")
                        continue

                if not records:
                    return {"error": "No se encontraron registros válidos."}

                if not send_command:
                     return {"error": "Error: Modo nativo activado pero cliente agente no disponible."}
                
                username = st.session_state.get("username", "admin")
                
                if not silent_mode:
                    st.info(f"Enviando tarea al agente local (Hospitalización) para {len(records)} historias...")
                    
                # Reutilizamos el comando 'download_ovida' del agente que ya soporta esta estructura
                # Nota: El agente usa una URL generica que puede servir para hospitalización si los parámetros coinciden.
                # Revisando el agente, usa 'reporte_historia_general.php' con verHC=1, verEvo=1, etc.
                # Esto parece ser lo estándar para 'Historia Completa'.
                
                task_id = send_command(username, "download_ovida", {
                    "base_path": save_path,
                    "records": records
                })
                
                if task_id:
                    status_placeholder = st.empty()
                    if not silent_mode: status_placeholder.text("Esperando agente...")
                    
                    res = wait_for_result(task_id, timeout=600)
                    
                    if not silent_mode: status_placeholder.empty()

                    if res and "status" in res:
                        if res["status"] == "success":
                             msg = res.get("message", "Finalizado correctamente")
                             return {"message": f"{msg} (Agente)"}
                        else:
                             return {"error": f"Error del agente: {res.get('message', 'Desconocido')}"}
                    else:
                         return {"error": f"Error del agente: {res.get('error', 'Desconocido')}"}
                else:
                    return {"error": "No se pudo crear la tarea."}

            except Exception as e:
                return {"error": f"Excepción agente: {str(e)}"}

        # --- SERVER EXECUTION ---
        options = webdriver.ChromeOptions()
        # Headless mode for server environment if possible, but OVIDA might require GUI for login
        # If running on server without display, we need headless.
        # But the code waits for manual login...
        # For now, let's keep it as is but be aware of headless limitations.
        # If we are in Web Mode (AWS), we probably can't open a browser window for the user to see.
        # This function might fail in AWS unless we use headless and auto-login (which we don't have credentials for).
        # Assuming for now this is used with "Agente Local" or user accepts it won't work on pure headless server without tunneling.
        
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
        
        timeout = 300
        start = time.time()
        logged_in = False
        while time.time() - start < timeout:
            try:
                if "App/Vistas" in driver.current_url:
                    logged_in = True
                    break
            except: pass
            time.sleep(2)
            
        if not logged_in:
            driver.quit()
            return {"error": "No se detectó inicio de sesión."}
            
        descargados, errores, conflictos = 0, 0, 0
        progress_bar = None
        if not silent_mode: progress_bar = st.progress(0, text="Descargando hospitalización...")
        
        for i, row in df.iterrows():
            if not silent_mode: progress_bar.progress((i + 1) / len(df))
            try:
                estudio = str(int(row[col_map['estudio']])).strip()
                ingreso = pd.to_datetime(row[col_map['ingreso']]).strftime('%Y/%m/%d')
                egreso = pd.to_datetime(row[col_map['egreso']]).strftime('%Y/%m/%d')
                carpeta = str(row[col_map['carpeta']]).strip()
                
                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                params = {
                    'estudio': estudio, 'fecha_inicio': ingreso, 'fecha_fin': egreso,
                    'verHC': 1, 'verEvo': 1, 'verPar': 1, 'ImprimirOrdenamiento': 1,
                    'ImprimirSolOrdenesExt': 1, 'ImprimirGraficasHC': 1,
                    'ImprimirFormatos': 1, 'ImprimirRegistroAdmon': 1,
                    'ImprimirNotasEnfermeria': 1
                }
                full_url = f"{base_url}?{urllib.parse.urlencode(params)}"
                
                dest_dir = os.path.join(save_path, carpeta)
                os.makedirs(dest_dir, exist_ok=True)
                dest_file = os.path.join(dest_dir, f"HC_{estudio}.pdf")
                
                if os.path.exists(dest_file):
                    conflictos += 1
                    continue
                    
                driver.get(full_url)
                time.sleep(2)
                pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {"landscape": False, "printBackground": True})
                with open(dest_file, 'wb') as f:
                    f.write(base64.b64decode(pdf_b64['data']))
                descargados += 1
            except: errores += 1
            
        if not silent_mode: progress_bar.empty()
        
        result_msg = f"Descargados: {descargados}. Errores: {errores}. Conflictos: {conflictos}."
        
        if is_temp:
            # Zip results
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(save_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, save_path)
                        zf.write(file_path, arcname)
            
            # Cleanup
            try:
                shutil.rmtree(save_path, ignore_errors=True)
            except: pass
            
            return {
                "files": [{
                    "name": f"Historias_OVIDA_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Historias (ZIP)"
                }],
                "message": result_msg
            }
        else:
            return {"message": result_msg}

    except Exception as e:
        return {"error": f"Error crítico: {e}"}
    finally:
        if driver: driver.quit()

# --- WORKERS: CDO VALIDATORS ---

def worker_registraduria_masiva(df, col_cedula, headless=True, update_progress=None, silent_mode=False):
    try:
        validator = ValidatorRegistraduria(headless=headless)
    except Exception as e:
        if not silent_mode: st.error(f"Error initializing ValidatorRegistraduria: {e}")
        return {"error": f"Error: {e}"}
    cb = update_progress if update_progress else lambda c, t, **kwargs: None
    try:
        df_results = validator.process_massive(df, col_cedula, progress_callback=cb)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_results.to_excel(writer, index=False)
        
        return {
            "files": [{
                "name": f"Resultados_Registraduria_{int(time.time())}.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Resultados"
            }],
            "message": f"Procesados {len(df_results)} registros."
        }
    except Exception as e:
        return {"error": f"Error processing massive Registraduria: {e}"}

def worker_adres_api_masiva(df, col_cedula, col_tipo_doc=None, default_tipo_doc="CC", update_progress=None, silent_mode=False):
    try:
        validator = ValidatorAdres()
    except Exception as e:
        return {"error": f"Error initializing ValidatorAdres: {e}"}
    cb = update_progress if update_progress else lambda c, t, **kwargs: None
    try:
        df_results = validator.process_massive(df, col_cedula, tipo_doc_col=col_tipo_doc, default_tipo_doc=default_tipo_doc, progress_callback=cb)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_results.to_excel(writer, index=False)
            
        return {
            "files": [{
                "name": f"Resultados_ADRES_API_{int(time.time())}.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Resultados"
            }],
            "message": f"Procesados {len(df_results)} registros."
        }
    except Exception as e:
        return {"error": f"Error processing massive ADRES API: {e}"}

def worker_adres_web_massive(df, col_cedula, col_tipo_doc=None, default_tipo_doc="CC", update_progress=None, silent_mode=False):
    try:
        validator = ValidatorAdresWeb(headless=False)
    except Exception as e:
        return {"error": f"Error initializing ValidatorAdresWeb: {e}"}
    cb = update_progress if update_progress else lambda c, t, **kwargs: None
    try:
        df_results = validator.process_massive(df, col_cedula, tipo_doc_col=col_tipo_doc, default_tipo_doc=default_tipo_doc, progress_callback=cb)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_results.to_excel(writer, index=False)
            
        return {
            "files": [{
                "name": f"Resultados_ADRES_WEB_{int(time.time())}.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Resultados"
            }],
            "message": f"Procesados {len(df_results)} registros."
        }
    except Exception as e:
        return {"error": f"Error processing massive ADRES Web: {e}"}

# --- DIALOGS (Proxies to Workers with UI) ---

@st.dialog("Importar Excel para Renombrado")
def dialog_importar_excel():
    st.write("### Renombrar archivos usando Excel")
    uploaded = st.file_uploader("Subir Excel", type=["xlsx", "xls"], key="up_historia")
    if uploaded:
        default_path = st.session_state.get("current_path", os.getcwd())
        target_path = render_path_selector("Carpeta donde aplicar cambios", "ren_excel_folder", default_path=default_path)
        folder = target_path
        
        if st.button("Aplicar Renombrado"):
            try:
                with st.spinner("Renombrando..."):
                    uploaded.seek(0)
                    result = worker_aplicar_renombrado_excel(uploaded, folder)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(folder, "dl_ren_excel", "📦 Descargar Carpeta Modificada (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_importar_excel"):
        close_auto_dialog()

@st.dialog("Añadir Sufijo desde Excel")
def dialog_sufijo():
    st.write("### Añadir Sufijo a Archivos (Por Carpeta)")
    st.info("Esta opción busca las carpetas listadas en el Excel y añade el sufijo a TODOS los archivos dentro de ellas.")
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")
        
    uploaded = st.file_uploader("Subir Excel", type=["xlsx", "xls"], key="up_analisis_auth")
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Seleccione la Hoja", sheet_names, key="suf_sheet")
            df_preview = _get_excel_preview(file_bytes, sheet, nrows=1)
            cols = df_preview.columns.tolist()
            
            col_folder = st.selectbox("Columna Nombre Carpeta", cols, index=0, key="suf_col_folder")
            col_suffix = st.selectbox("Columna Sufijo", cols, index=min(1, len(cols)-1), key="suf_col_suf")
            
            use_filter = st.checkbox("Usar filtro de Excel (Filas visibles)", value=False, key="suf_filter")
            
            default_path = st.session_state.get("current_path", os.getcwd())
            target_path = render_path_selector("Carpeta Raíz (donde están las carpetas)", "suf_folder", default_path=default_path)
            folder = target_path
            
            if st.button("Aplicar Sufijos"):
                try:
                    with st.spinner("Procesando sufijos..."):
                        # Reset pointer for worker
                        uploaded.seek(0)
                        # Run synchronously
                        result = worker_anadir_sufijo_excel(uploaded, sheet, col_folder, col_suffix, folder, use_filter)
                        st.success(result)
                        close_auto_dialog()
#                         render_download_button(folder, "dl_sufijo", "📦 Descargar Resultados (ZIP)")
                        # time.sleep(2)
                        # st.rerun()
                except Exception as e:
                    st.error(f"Error al procesar: {e}")
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")

    if st.button("Cerrar", key="btn_close_sufijo"):
        close_auto_dialog()

@st.dialog("Renombrar por Mapeo Excel")
def dialog_renombrar_mapeo_excel():
    st.write("### Renombrar Archivos (Mapeo Columna A -> Columna B)")
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")
        
    uploaded = st.file_uploader("Subir Excel", type=["xlsx", "xls"], key="up_analisis_sanitas")
    
    sheet = None
    col_src = None
    col_dst = None
    use_filter = True
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Nombre Hoja", sheet_names, key="ren_map_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                c1, c2 = st.columns(2)
                col_src = c1.selectbox("Columna Nombre Actual", df_preview.columns, key="ren_map_src")
                col_dst = c2.selectbox("Columna Nombre Nuevo", df_preview.columns, key="ren_map_dst")
                use_filter = st.checkbox("Usar filtro de Excel (solo visibles)", value=True, key="ren_map_filter")
        except Exception as e:
            st.error(f"Error: {e}")

    st.write("Carpeta Objetivo:")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    target_path = render_path_selector("Ruta", "ren_map_folder", default_path=default_path)
    folder = target_path
    
    if st.button("Renombrar"):
        if uploaded and sheet and col_src and col_dst and folder:
            try:
                with st.spinner("Renombrando..."):
                    if hasattr(uploaded, 'seek'):
                        uploaded.seek(0)
                    result = worker_renombrar_mapeo_excel(uploaded, sheet, col_src, col_dst, use_filter, folder)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(folder, "dl_ren_map", "📦 Descargar Carpeta Modificada (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_ren_map"):
        close_auto_dialog()

@st.dialog("Modificar DOCX Completo")
def dialog_modif_docx_completo():
    st.write("### Modificación Masiva de DOCX (Plantillas)")
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")
        
    uploaded = st.file_uploader("Subir Excel de Datos", type=["xlsx"], key="up_mod_docx_full")
    
    sheet = None
    use_filter = False
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Nombre Hoja", sheet_names, key="mod_full_sheet")
            use_filter = st.checkbox("Usar filtros de Excel (solo filas visibles)", value=False, key="mod_full_filter")
        except Exception as e:
            st.error(f"Error: {e}")

    st.write("Carpeta Objetivo:")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    target_path = render_path_selector("Ruta", "mod_full_folder", default_path=default_path)
    folder = target_path
    
    if st.button("Ejecutar Modificación"):
        if uploaded and sheet and folder:
            try:
                with st.spinner("Modificando documentos..."):
                    uploaded.seek(0)
                    result = worker_modificar_docx_completo(uploaded, sheet, folder, use_filter)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(folder, "dl_mod_docx_full", "📦 Descargar Carpeta Modificada (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_mod_full"):
        close_auto_dialog()

@st.dialog("Insertar Firma en DOCX (Masivo)")
def dialog_insertar_firma_docx():
    st.write("### Insertar Firma (Imagen) en DOCX")
    st.write("Inserta una imagen de firma en documentos DOCX dentro de subcarpetas.")
    st.info("Busca la imagen de firma y la inserta en el DOCX donde diga 'Firma de Aceptacion'.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")

    st.write("Carpeta Base:")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    target_path = render_path_selector("Ruta", "firma_docx_base", default_path=default_path)
    
    base_path = target_path
    docx_name = st.text_input("Nombre del DOCX", value="Consentimiento.docx")
    sig_name = st.text_input("Nombre de la Firma (Imagen)", value="firma.jpg")
    
    if st.button("Iniciar Inserción de Firmas"):
        if base_path and docx_name and sig_name:
            try:
                with st.spinner("Insertando firmas..."):
                    result = worker_firmar_docx_con_imagen_masivo(base_path, docx_name, sig_name)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(base_path, "dl_sign_docx", "📦 Descargar Destino (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.error("Complete todos los campos.")

    if st.button("Cerrar", key="btn_close_ins_firma"):
        close_auto_dialog()

@st.dialog("Generar CUV (FEVRIPS)")
def dialog_generar_cuv():
    st.write("### Generar CUV masivo")
    st.info("Funcionalidad pendiente de integración completa.")

@st.dialog("RIPS: Limpieza JSON")
def dialog_rips_limpieza_json():
    st.write("### Limpieza de espacios en JSON")
    st.info("Esta herramienta elimina espacios extra en claves y valores de archivos JSON.")
    # Implementation placeholder

@st.dialog("RIPS: Actualizar Clave")
def dialog_rips_update_key():
    st.write("### Actualizar Clave en JSON")
    uploaded_files = st.file_uploader("Seleccionar archivos JSON", type=["json"], accept_multiple_files=True, key="up_json_analysis")
    key_to_update = st.text_input("Clave a buscar")
    new_value = st.text_input("Nuevo valor")
    if uploaded_files and key_to_update and st.button("Actualizar Clave"):
        # Placeholder
        pass

# --- WORKERS: ANALYSIS & EXTRACTION ---

def worker_analisis_historia_clinica(file_list, silent_mode=False):
    """
    Analiza masivamente los archivos PDF para extraer datos de historias clínicas.
    Retorna bytes de Excel.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(file_list):
        if not silent_mode: st.info(f"Delegando análisis HC al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_hc", {"file_list": file_list})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    if not file_list:
        if not silent_mode: st.error("No hay archivos para analizar.")
        return None

    # Filter PDFs
    archivos_pdf = [f for f in file_list if f.lower().endswith('.pdf')]
    if not archivos_pdf:
        if not silent_mode: st.warning("No se encontraron archivos PDF.")
        return None

    patterns = {
        'Identificación': re.compile(r"Identificaci[oó]n:?\s*(.*?)(?=\s+Paciente:|\n|$)", re.IGNORECASE),
        # 'Paciente': re.compile(r"Paciente:?\s*(.*?)(?=\s+Fecha|Ingreso:|\n|$)", re.IGNORECASE),
        'Fecha Ingreso': re.compile(r"Fecha Ingreso:?\s*(.*?)(?=\s+Hora|\n|$)", re.IGNORECASE),
        'Hora Ing': re.compile(r"Hora Ing:?\s*(.*?)(?=\s+Ingreso:|\n|$)", re.IGNORECASE),
        'Ingreso': re.compile(r"(?<!Fecha\s)Ingreso:?\s*(\d+)", re.IGNORECASE),
        'Fecha Atencion': re.compile(r"Fecha Atenci[oó]n:?\s*(.*?)(?=\s+Fecha|\n|$)", re.IGNORECASE),
        'Fecha Cierre HC': re.compile(r"Fecha Cierre HC:?\s*(.*?)(?=\s+Fecha|\n|$)", re.IGNORECASE),
        'Fecha Naci': re.compile(r"Fecha Naci:?\s*(.*?)(?=\s+Edad:|\n|$)", re.IGNORECASE),
        'Edad': re.compile(r"Edad:?\s*(.*?)(?=\s+Sexo:|\n|$)", re.IGNORECASE),
        'Sexo': re.compile(r"Sexo:?\s*(.*?)(?=\n|$)", re.IGNORECASE),
        'Nro.Historia': re.compile(r"Nro\.Historia:?\s*(.*?)(?=\s+Tipo Usuario:|\n|$)", re.IGNORECASE),
        'Tipo Usuario': re.compile(r"Tipo Usuario:?\s*(.*?)(?=\n|$)", re.IGNORECASE),
        'Telefono': re.compile(r"Tel[ée]fono:?\s*([\d\-\s]+)(?=\s+Estrato:|\n|$)", re.IGNORECASE),
        'Estrato': re.compile(r"Estrato:?\s*(.*?)(?=\s+Municipio:|\n|$)", re.IGNORECASE),
        'Municipio': re.compile(r"Municipio:?\s*(.*?)(?=\n|$)", re.IGNORECASE),
        'Dirección': re.compile(r"Direcci[oó]n:?\s*(.*?)(?=\s+Estado Civil:|\n|$)", re.IGNORECASE),
        'Estado Civil': re.compile(r"Estado Civil:?\s*(.*?)(?=\n|$)", re.IGNORECASE),
        'Empresa': re.compile(r"Empresa:?\s*(.*?)(?=\s+Contrato:|\n|$)", re.IGNORECASE),
        'Contrato': re.compile(r"Contrato:?\s*(.*?)(?=\n|$)", re.IGNORECASE),
        'Acompañante': re.compile(r"Acompañante:?\s*(.*?)(?=\s+Tel\. Acompañante:|\n|$)", re.IGNORECASE),
        # 'Tel. Acompañante': re.compile(r"Tel\. Acompañante:?\s*(.*?)(?=\n|$)", re.IGNORECASE),
        'Discapacidad': re.compile(r"DISCAPACIDAD\s*(.*?)(?=\s+DESCRIPCION|\n\n|$)", re.IGNORECASE | re.DOTALL),
        'Motivo Consulta': re.compile(r"MOTIVO DE CONSULTA\s*(.*?)(?=\s+ENFERMEDAD ACTUAL|\s+ANTECEDENTES|\s+REVISION POR SISTEMAS|$)", re.IGNORECASE | re.DOTALL)
    }

    extracted_data = []
    
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0)
        status_text = st.empty()

    for i, pdf_path in enumerate(archivos_pdf):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(archivos_pdf))
            status_text.text(f"Procesando: {os.path.basename(pdf_path)}")

        try:
            full_text = ""
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            
            record = {'Archivo': os.path.basename(pdf_path)}
            for key, pattern in patterns.items():
                match = pattern.search(full_text)
                if match:
                    val = match.group(1).strip()
                    # Clean up common issues
                    val = val.replace('\n', ' ').strip()
                    record[key] = val
                else:
                    record[key] = ""

            extracted_data.append(record)

        except Exception as e:
            if not silent_mode: st.error(f"Error procesando {os.path.basename(pdf_path)}: {e}")

    if not extracted_data:
        if not silent_mode: st.warning("No se extrajeron datos.")
        return None

    try:
        column_order = [
            'Archivo', 'Identificación', 'Fecha Ingreso', 'Hora Ing', 'Ingreso', 
            'Fecha Atencion', 'Fecha Cierre HC', 'Fecha Naci', 'Edad', 'Sexo', 
            'Nro.Historia', 'Tipo Usuario', 'Telefono', 'Estrato', 'Municipio', 
            'Dirección', 'Estado Civil', 'Empresa', 'Contrato', 
            'Acompañante', 'Discapacidad', 'Motivo Consulta'
        ]
        df = pd.DataFrame(extracted_data)
        # Reorder if columns exist
        cols = [c for c in column_order if c in df.columns] + [c for c in df.columns if c not in column_order]
        df = df[cols]
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False)
        return {
            "files": [{
                "name": "Analisis_Historia_Clinica.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Análisis HC"
            }],
            "message": f"Procesados: {len(extracted_data)} registros."
        }
    except Exception as e:
        if not silent_mode: st.error(f"Error generando Excel: {e}")
        return None

def worker_leer_pdf_retefuente(file_list, silent_mode=False):
    """
    Lee archivos PDF de Retefuente y extrae RAZON SOCIAL y NIT.
    Retorna bytes de Excel.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(file_list):
        if not silent_mode: st.info(f"Delegando análisis Retefuente al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_rete", {"file_list": file_list})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------
    
    if not file_list: return None
    
    archivos_pdf = [f for f in file_list if f.lower().endswith('.pdf')]
    if not archivos_pdf: return None

    resultados_datos = []
    
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0)
        status_text = st.empty()

    for i, ruta_pdf in enumerate(archivos_pdf):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(archivos_pdf))
            status_text.text(f"Procesando: {os.path.basename(ruta_pdf)}")
            
        try:
            with fitz.open(ruta_pdf) as doc:
                for num_pagina, page in enumerate(doc, start=1):
                    blocks = page.get_text("blocks")
                    blocks.sort(key=lambda b: b[1]) # Sort vertically
                    
                    label_block = None
                    nit_label_block = None
                    nombre_encontrado = "NO ENCONTRADO"
                    nit_encontrado = "NO ENCONTRADO"
                    
                    # 1. Find key labels
                    for b in blocks:
                        text_clean = " ".join(b[4].split()).upper()
                        if "PRACTICO LA RETENCION" in text_clean:
                            label_block = b
                            break
                    
                    # Find NIT label
                    if label_block:
                        lx0, ly0, lx1, ly1 = label_block[:4]
                        for b in blocks:
                            bx0, by0 = b[:2]
                            text_clean = " ".join(b[4].split()).upper()
                            if bx0 > lx0 and abs(by0 - ly0) < 30:
                                if "NIT" in text_clean or "C.C." in text_clean:
                                    nit_label_block = b
                                    break
                    
                    if not nit_label_block:
                        for b in blocks:
                            text_clean = " ".join(b[4].split()).upper()
                            if "NIT." in text_clean and "C.C." in text_clean:
                                nit_label_block = b
                                break

                    # 2. Extract NAME
                    if label_block:
                        lx0, ly0, lx1, ly1 = label_block[:4]
                        candidates = []
                        for b in blocks:
                            if b == label_block: continue
                            bx0, by0 = b[:2]
                            if by0 > ly0 and abs(bx0 - lx0) < 100: 
                                candidates.append(b)
                        candidates.sort(key=lambda b: b[1])
                        
                        for cand in candidates:
                            text_cand = cand[4].strip()
                            upper_cand = text_cand.upper()
                            if not text_cand: continue
                            if "DIRECCION" in upper_cand: break
                            if "NIT" in upper_cand or "C.C." in upper_cand: continue
                            nombre_encontrado = " ".join(text_cand.split())
                            break
                            
                    # 3. Extract NIT
                    if nit_label_block:
                        nx0, ny0 = nit_label_block[:2]
                        nit_candidates = []
                        for b in blocks:
                            if b == nit_label_block: continue
                            bx0, by0 = b[:2]
                            if by0 > ny0 and abs(bx0 - nx0) < 80:
                                nit_candidates.append(b)
                        nit_candidates.sort(key=lambda b: b[1])
                        
                        for cand in nit_candidates:
                            text_cand = cand[4].strip()
                            upper_cand = text_cand.upper()
                            if not text_cand: continue
                            if "CIUDAD" in upper_cand: break
                            nit_encontrado = " ".join(text_cand.split())
                            break
                    
                    # Fallback for Name
                    if nombre_encontrado == "NO ENCONTRADO":
                        full_text = page.get_text("text")
                        lines = [l.strip() for l in full_text.split('\n') if l.strip()]
                        for idx, line in enumerate(lines):
                            if "PRACTICO LA RETENCION" in line.upper():
                                if idx + 1 < len(lines):
                                    potential = lines[idx+1]
                                    if "NIT" not in potential.upper() and "DIRECCION" not in potential.upper():
                                         nombre_encontrado = potential
                                break
                    
                    # Cleanup and separation
                    if nombre_encontrado != "NO ENCONTRADO":
                        match_mix = re.search(r'^(.*?)(\d{6,}[\d\s]*)$', nombre_encontrado)
                        if match_mix:
                            nombre_limpio = match_mix.group(1).strip()
                            nit_extraido = match_mix.group(2).replace(" ", "").strip()
                            nombre_encontrado = nombre_limpio
                            if nit_encontrado == "NO ENCONTRADO" or not any(c.isdigit() for c in nit_encontrado):
                                nit_encontrado = nit_extraido
                            elif any(c.isalpha() for c in nit_encontrado):
                                 nit_encontrado = nit_extraido

                    if nombre_encontrado != "NO ENCONTRADO" and not nombre_encontrado.lower().endswith('.pdf'):
                        nombre_encontrado += ".pdf"

                    resultados_datos.append({
                        "Archivo": os.path.basename(ruta_pdf),
                        "Página": num_pagina,
                        "RAZON SOCIAL / NOMBRE": nombre_encontrado,
                        "NIT / C.C.": nit_encontrado
                    })
                    
        except Exception as e:
            resultados_datos.append({
                "Archivo": os.path.basename(ruta_pdf),
                "Página": "Error",
                "RAZON SOCIAL / NOMBRE": f"ERROR: {str(e)}"
            })

    if resultados_datos:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            pd.DataFrame(resultados_datos).to_excel(writer, index=False)
        return {
            "files": [{
                "name": "Analisis_Retefuente.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Retefuente"
            }],
            "message": f"Procesados: {len(resultados_datos)} registros."
        }
    return None

def worker_analisis_emssanar(file_list, silent_mode=False):
    """
    Analiza archivos PDF de Emssanar (Autorizaciones).
    Extrae información del encabezado, paciente, servicios y pagos.
    Resalta columnas en amarillo.
    Retorna bytes de Excel.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(file_list):
        if not silent_mode: st.info(f"Delegando análisis Emssanar al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_emssanar", {"file_list": file_list})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    if not file_list: return None
    
    archivos_pdf = [f for f in file_list if f.lower().endswith('.pdf')]
    if not archivos_pdf: return None

    extracted_data = []
    
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0)
        status_text = st.empty()

    for i, pdf_path in enumerate(archivos_pdf):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(archivos_pdf))
            status_text.text(f"Procesando: {os.path.basename(pdf_path)}")

        try:
            record = {'Archivo': os.path.basename(pdf_path)}
            
            with fitz.open(pdf_path) as doc:
                page = doc[0] # Asumimos primera página
                text = page.get_text("text")
                words = page.get_text("words") # (x0, y0, x1, y1, "word", block_no, line_no, word_no)
                
                clean_text = re.sub(r'\s+', ' ', text)
                
                # --- 1. Número de Autorización ---
                # Usamos una búsqueda más robusta:
                # 1. Regex directo si el texto está limpio.
                # 2. Búsqueda por coordenadas si el texto está fragmentado.
                m_auth = re.search(r'N[UÚ]MERO DE AUTORIZACI[OÓ]N[:\.]?\s*(\d+)', clean_text, re.IGNORECASE)
                if m_auth:
                    record['Numero Autorizacion'] = m_auth.group(1)
                else:
                    # Búsqueda espacial
                    label_rects = page.search_for("NÚMERO DE AUTORIZACIÓN")
                    if label_rects:
                        r = label_rects[0]
                        # Buscar a la derecha
                        val_rect = fitz.Rect(r.x1, r.y0 - 2, page.rect.width, r.y1 + 2)
                        val = page.get_text("text", clip=val_rect).strip()
                        # Extraer dígitos largos
                        nums = re.findall(r'\d{5,}', val)
                        record['Numero Autorizacion'] = nums[0] if nums else ""
                    else:
                        # Intento fallback: Buscar cualquier número largo al inicio del documento
                        m_auth_fallback = re.search(r'AUTORIZACI[OÓ]N\s*:?\s*(\d{5,})', clean_text, re.IGNORECASE)
                        record['Numero Autorizacion'] = m_auth_fallback.group(1) if m_auth_fallback else ""

                # --- 2. Número Documento ---
                m_doc = re.search(r'(\d{6,15})\s*N[uú]mero documento de identificaci[oó]n', clean_text, re.IGNORECASE)
                if not m_doc:
                     m_doc = re.search(r'N[uú]mero documento de identificaci[oó]n\s*(\d{6,15})', clean_text, re.IGNORECASE)
                
                if not m_doc:
                      label_rect = page.search_for("Número documento de identificación")
                      if label_rect:
                           l_rect = label_rect[0]
                           search_rect = fitz.Rect(l_rect.x0 - 50, l_rect.y0 - 30, l_rect.x1 + 50, l_rect.y0)
                           nearby_text = page.get_text("text", clip=search_rect).strip()
                           nums = re.findall(r'\d{6,15}', nearby_text)
                           if nums: record['Documento Paciente'] = nums[0]
                           else: record['Documento Paciente'] = ""
                      else:
                           record['Documento Paciente'] = ""
                else:
                     record['Documento Paciente'] = m_doc.group(1)

                # --- 3. Tipo Documento (Detectar X) ---
                tipos_doc = [
                    "Registro civil", "Tarjeta de identidad", "Cédula de ciudadanía", "Cédula de extranjería",
                    "Pasaporte", "Adulto sin identificación", "Menor sin identificación", "Permiso especial de permanencia"
                ]
                found_type = ""
                xs = [w for w in words if w[4].strip().upper() == 'X']
                for x_word in xs:
                    x_rect = fitz.Rect(x_word[:4])
                    search_area = fitz.Rect(x_rect.x1, x_rect.y0 - 5, page.rect.width, x_rect.y1 + 5)
                    text_right = page.get_text("text", clip=search_area).strip()
                    text_right = re.sub(r'\s+', ' ', text_right)
                    for t in tipos_doc:
                        if t.lower() in text_right.lower():
                            found_type = t
                            break
                    if found_type: break
                record['Tipo Documento'] = found_type

                # --- 4. Nombre Paciente (Mejorado) ---
                # Buscar texto entre "DATOS DEL PACIENTE" y la siguiente etiqueta estructural
                label_datos = page.search_for("DATOS DEL PACIENTE")
                label_tipo = page.search_for("Tipo Documento") # Suele estar debajo
                
                nombre_encontrado = ""
                if label_datos:
                    y_top = label_datos[0].y1
                    # Definir limite inferior
                    y_bottom = label_tipo[0].y0 if label_tipo else y_top + 100 
                    
                    # Buscar texto en esa franja
                    name_rect = fitz.Rect(0, y_top, page.rect.width, y_bottom)
                    name_text = page.get_text("text", clip=name_rect).strip()
                    # Limpieza
                    name_text = re.sub(r'\s+', ' ', name_text)
                    # Eliminar etiquetas comunes que podrían aparecer
                    name_text = re.sub(r'(1er Apellido|2do Apellido|1er Nombre|2do Nombre)', '', name_text, flags=re.IGNORECASE)
                    nombre_encontrado = name_text.strip()
                
                if not nombre_encontrado:
                    # Fallback regex original
                    m_nombres = re.search(r'DATOS DEL PACIENTE\s+([A-Z\s]+?)\s+(?:Tipo Documento|1er Apellido)', clean_text)
                    if m_nombres:
                        nombre_encontrado = m_nombres.group(1).strip()
                
                record['Nombre Paciente'] = nombre_encontrado

                # --- 5. Valor a Pagar por el Usuario ---
                label_valor = page.search_for("Valor a pagar por el usuario")
                val_pagar = "0"
                if label_valor:
                    r = label_valor[0]
                    # Buscar debajo de la etiqueta
                    # Asumimos que el valor está en una caja debajo, digamos 30px de alto
                    search_rect = fitz.Rect(r.x0 - 20, r.y1, r.x1 + 20, r.y1 + 40)
                    val_text = page.get_text("text", clip=search_rect).strip()
                    # Extraer números/moneda
                    matches = re.findall(r'[\d\.,]+', val_text)
                    if matches:
                        val_pagar = matches[0]
                record['Valor a Pagar'] = val_pagar

                # --- 6. Concepto (Cuota moderadora, etc) ---
                conceptos_posibles = ["Cuota moderadora", "Copago", "Cuota de recuperación", "Pagos compartidos"]
                found_concepto = "exento de pago"
                
                # Reutilizamos las Xs encontradas
                for x_word in xs:
                    x_rect = fitz.Rect(x_word[:4])
                    search_area = fitz.Rect(x_rect.x1, x_rect.y0 - 5, page.rect.width, x_rect.y1 + 5)
                    text_right = page.get_text("text", clip=search_area).strip()
                    text_right = re.sub(r'\s+', ' ', text_right).lower()
                    
                    for c in conceptos_posibles:
                        if c.lower() in text_right:
                            found_concepto = c
                            break
                    if found_concepto != "exento de pago":
                        break
                record['Concepto'] = found_concepto


                # --- Otros Campos ---
                m_nit = re.search(r'NIT/CC:?\s*([\d\.-]+)', clean_text, re.IGNORECASE)
                record['NIT Entidad'] = m_nit.group(1) if m_nit else ""

                m_regimen = re.search(r'R[ée]gimen afiliaci[oó]n:?\s*(\w+)', clean_text, re.IGNORECASE)
                record['Regimen'] = m_regimen.group(1) if m_regimen else ""

                m_depto = re.search(r'Departamento:?\s*(.*?)(?=\s+Municipio:)', clean_text, re.IGNORECASE)
                record['Departamento'] = m_depto.group(1).strip() if m_depto else ""
                
                found_codes = re.findall(r'(?<!\d)(\d{6})(?!\d)', clean_text)
                valid_codes = []
                for c in found_codes:
                    if c not in [record.get('Documento Paciente', ''), '900438792']: 
                         valid_codes.append(c)
                record['Codigos Servicios'] = " | ".join(set(valid_codes))

                m_desc = re.search(r'SE AUTORIZA\s*\((\d+)\)\s*-\s*([^-\.]+)', clean_text, re.IGNORECASE)
                if m_desc:
                     record['Descripcion Servicio'] = m_desc.group(2).strip()
                else:
                     m_desc2 = re.search(r'(CONSULTA DE [A-Z\s]+)', clean_text)
                     record['Descripcion Servicio'] = m_desc2.group(1).strip() if m_desc2 else ""

                m_cant = re.search(r'\s(\d+)\s+Consulta', clean_text, re.IGNORECASE)
                record['Cantidad'] = m_cant.group(1) if m_cant else "1"

            extracted_data.append(record)

        except Exception as e:
            if not silent_mode: st.error(f"Error procesando {os.path.basename(pdf_path)}: {e}")

    if extracted_data:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df = pd.DataFrame(extracted_data)
            
            # Columnas prioritarias y a resaltar
            # Reemplazamos Porcentaje por Valor a Pagar y Concepto
            cols_highlight = [
                'Numero Autorizacion', 'Documento Paciente', 'Tipo Documento', 'NIT Entidad',
                'Nombre Paciente', 'Regimen', 'Departamento', 'Codigos Servicios',
                'Descripcion Servicio', 'Cantidad', 'Valor a Pagar', 'Concepto'
            ]
            
            # Reordenar
            cols_final = ['Archivo'] + [c for c in cols_highlight if c in df.columns] + [c for c in df.columns if c not in cols_highlight and c != 'Archivo']
            df = df[cols_final]
            
            sheet_name = 'Resultados'
            df.to_excel(writer, index=False, sheet_name=sheet_name)
            
            # Formato Amarillo
            workbook = writer.book
            worksheet = writer.sheets[sheet_name]
            yellow_format = workbook.add_format({'bg_color': '#FFFF00', 'border': 1})
            header_format = workbook.add_format({'bg_color': '#FFFF00', 'bold': True, 'border': 1})

            # Aplicar formato
            for col_num, value in enumerate(df.columns.values):
                if value in cols_highlight:
                    worksheet.write(0, col_num, value, header_format)
                    worksheet.conditional_format(1, col_num, len(df), col_num, {
                        'type': 'no_errors',
                        'format': yellow_format
                    })
                else:
                    worksheet.set_column(col_num, col_num, 20)

            
        return {
            "files": [{
                "name": "Analisis_Autorizaciones_Emssanar_Completo.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Emssanar Completo"
            }],
            "message": f"Procesados: {len(extracted_data)} registros."
        }
    return None

def worker_analisis_fomag(file_list, silent_mode=False):
    """
    Analiza archivos PDF de Autorizaciones FOMAG.
    Extrae información del encabezado, paciente, prestador y servicios.
    Retorna bytes de Excel.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode:
        if not silent_mode: st.info(f"Delegando análisis Fomag al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_fomag", {"file_list": file_list})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    if not file_list: return None
    
    archivos_pdf = [f for f in file_list if f.lower().endswith('.pdf')]
    if not archivos_pdf: return None

    extracted_data = []
    
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0)
        status_text = st.empty()

    def get_value_below(page, label_text, h=30, w_add=0, match_idx=0):
        """Busca una etiqueta y devuelve el texto inmediatamente debajo."""
        labels = [label_text] if isinstance(label_text, str) else label_text
        all_rects = []
        for l in labels:
            all_rects.extend(page.search_for(l))
        
        all_rects.sort(key=lambda r: r.y0)
        
        if not all_rects or len(all_rects) <= match_idx: return ""
        r = all_rects[match_idx]
        
        # Rectángulo de búsqueda debajo de la etiqueta
        search_rect = fitz.Rect(r.x0, r.y1, r.x1 + w_add, r.y1 + h)
        val = page.get_text("text", clip=search_rect).strip()
        return re.sub(r'\s+', ' ', val)
    
    def get_field_dynamic(page, labels, next_labels=None, h=30, min_width=100, match_idx=0, single_line=False, stop_words=None, x_offset=0):
        """
        Busca 'labels' y trata de determinar el ancho hasta el siguiente campo ('next_labels').
        Si no encuentra next_labels, usa min_width.
        """
        if isinstance(labels, str): labels = [labels]
        if isinstance(next_labels, str): next_labels = [next_labels]
        
        # 1. Encontrar etiqueta principal
        candidates = []
        for l in labels:
            candidates.extend(page.search_for(l))
        candidates.sort(key=lambda r: (r.y0, r.x0)) # Sort by Y then X
        
        if not candidates or len(candidates) <= match_idx: return ""
        curr_rect = candidates[match_idx]
        
        # 2. Encontrar etiqueta siguiente (para limitar ancho)
        limit_x = curr_rect.x1 + min_width # Default limit
        
        if next_labels:
            next_candidates = []
            for nl in next_labels:
                next_candidates.extend(page.search_for(nl))
            
            # Filtrar solo aquellos que están en la misma "línea" (y similar) y a la derecha
            valid_next = [nr for nr in next_candidates 
                          if abs(nr.y0 - curr_rect.y0) < 10 and nr.x0 > curr_rect.x0]
            
            if valid_next:
                valid_next.sort(key=lambda r: r.x0)
                limit_x = valid_next[0].x0 # El inicio del siguiente campo es el límite
            
        # 3. Extraer texto
        # Aplicamos x_offset al inicio de la búsqueda (útil si el texto empieza antes de la etiqueta o para ajustar márgenes)
        search_rect = fitz.Rect(curr_rect.x0 + x_offset, curr_rect.y1, limit_x, curr_rect.y1 + h)
        val = page.get_text("text", clip=search_rect).strip()
        
        if single_line and '\n' in val:
            val = val.split('\n')[0].strip()
            
        if stop_words:
            for sw in stop_words:
                # Case insensitive stop word check
                if sw.lower() in val.lower():
                    # Find the actual index to split
                    idx = val.lower().find(sw.lower())
                    if idx != -1:
                        val = val[:idx].strip()
        
        return re.sub(r'\s+', ' ', val)

    for i, pdf_path in enumerate(archivos_pdf):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(archivos_pdf))
            status_text.text(f"Procesando: {os.path.basename(pdf_path)}")

        try:
            record = {'Archivo': os.path.basename(pdf_path)}
            
            with fitz.open(pdf_path) as doc:
                page = doc[0]
                text = page.get_text("text")
                clean_text = re.sub(r'\s+', ' ', text)

                # --- 1. Encabezado ---
                m_fecha = re.search(r'Fecha de Gestion de Red:?\s*([\d-]+)', clean_text, re.IGNORECASE)
                record['Fecha Gestion'] = m_fecha.group(1) if m_fecha else ""

                m_orden = re.search(r'N[uú]mero de Orden:?\s*(\d+)', clean_text, re.IGNORECASE)
                record['Numero Orden'] = m_orden.group(1) if m_orden else ""

                # --- 2. Paciente ---
                # Usamos get_field_dynamic con x_offset negativo para capturar nombres largos alineados a la izquierda
                # single_line=False para evitar cortar nombres en dos líneas, luego limpiamos
                record['Nombre Paciente'] = get_field_dynamic(page, "Nombre Paciente", next_labels=["Sexo", "Identificación", "Identificacion"], min_width=250, h=40, single_line=False, stop_words=["Telefono", "Teléfono", "Direccion"], x_offset=-40)
                if record['Nombre Paciente']:
                     # Limpieza básica: quitar saltos de línea y posibles letras basura al inicio (ej: "E FLOREZ")
                     record['Nombre Paciente'] = re.sub(r'\s+', ' ', record['Nombre Paciente'])
                     # Si empieza con una letra sola y espacio, y luego sigue texto largo, asumimos basura
                     if re.match(r'^[A-Z]\s+[A-Z]{3,}', record['Nombre Paciente']):
                         # Pero cuidado con iniciales. Solo quitamos si es "E" o similar conocido
                         # Por ahora dejamos todo, mejor tener "E FLOREZ" que "OREZ"
                         pass

                record['Sexo'] = get_field_dynamic(page, "Sexo", next_labels=["Identificación", "Identificacion", "Edad"], min_width=50, h=20, single_line=True)
                
                # Identificacion: Aumentamos ancho y altura, y añadimos variaciones de etiqueta
                # single_line=False porque a veces aparece basura "E D I" en la primera línea y el ID en la segunda
                record['Identificacion'] = get_field_dynamic(page, ["Identificación", "Identificacion", "Identificación del Paciente", "Identificacion del Paciente"], next_labels=["Edad", "Nacimiento"], min_width=200, h=40, single_line=False, stop_words=["Correo", "Edad"])
                # Limpieza extra de Identificacion
                if record['Identificacion']:
                     # Eliminar prefijos comunes de basura o etiquetas repetidas
                     val_id = record['Identificacion']
                     val_id = re.sub(r'^(E\s?D\s?I\s?|Trans\s?)', '', val_id, flags=re.IGNORECASE).strip()
                     # Buscar patrón de documento si existe (CC, TI, etc seguido de numeros)
                     m_id = re.search(r'((?:CC|TI|RC|CE|PA|CD|SC|PE)\s*[-:]?\s*\d+)', val_id, re.IGNORECASE)
                     if m_id:
                         val_id = m_id.group(1)
                     record['Identificacion'] = val_id

                record['Edad'] = get_field_dynamic(page, "Edad", next_labels=["Nacimiento"], min_width=50, h=20, single_line=True, stop_words=["Munic", "Municipio"])
                # record['Nacimiento'] = get_field_dynamic(page, "Nacimiento", next_labels=["Direccion", "Dirección"], min_width=100, h=15, single_line=True)
                
                # Direccion aparece 2 veces (Paciente y Prestador)
                # record['Direccion Paciente'] = get_field_dynamic(page, ["Direccion", "Dirección"], next_labels=["Telefono", "Teléfono"], min_width=200, match_idx=0, h=15, single_line=True)
                # record['Telefono Paciente'] = get_field_dynamic(page, ["Telefono", "Teléfono"], next_labels=["Correo"], min_width=100, match_idx=0, h=15, single_line=True)
                # record['Correo'] = get_field_dynamic(page, ["Correo"], next_labels=["Municipio"], min_width=150, h=15, single_line=True)
                # record['Municipio Paciente'] = get_field_dynamic(page, "Municipio", next_labels=["Nombre Prestador"], min_width=150, match_idx=0, h=15, single_line=True)

                # --- 3. Prestador ---
                # record['Nombre Prestador'] = get_value_below(page, "Nombre Prestador", w_add=200)
                # record['Direccion Prestador'] = get_value_below(page, ["Direccion", "Dirección"], w_add=200, match_idx=1)
                
                # NIT Prestador (el segundo NIT en el documento, el primero es del encabezado)
                # record['NIT Prestador'] = get_value_below(page, "NIT", match_idx=1, w_add=50)
                
                # record['Telefono Prestador'] = get_value_below(page, ["Telefono", "Teléfono"], w_add=50, match_idx=1)
                
                # Cod Habilitacion: Aumentamos w_add y variaciones (incluyendo acento grave `ò`)
                record['Cod Habilitacion'] = get_value_below(page, ["Cod Habilitación", "Cod Habilitacion", "Cod Habilitaciòn", "Código de Habilitación", "Codigo de Habilitacion", "Cód. Habilitación"], w_add=150, h=35)
                # Limpieza Cod Habilitacion (quitar basura como "O M")
                if record['Cod Habilitacion']:
                    m_cod = re.search(r'(\d{8,12})', record['Cod Habilitacion'])
                    if m_cod:
                        record['Cod Habilitacion'] = m_cod.group(1)
                # record['Municipio Prestador'] = get_value_below(page, "Municipio", w_add=100, match_idx=1)
                # record['Diagnostico DX'] = get_value_below(page, "Diagnostico DX", w_add=50)

                # --- 4. Servicio ---
                # Estrategia por columnas basada en encabezados para mayor precisión
                
                # 1. Buscar encabezado de referencia (Consecutivo)
                header_consec = None
                for l in ["N°.Consecutivo", "N°. Consecutivo", "Consecutivo"]:
                    res = page.search_for(l)
                    if res:
                        header_consec = res[0]
                        break
                
                # 2. Buscar encabezado de Codigo (Cerca del consecutivo)
                header_cod = None
                for l in ["Código", "Codigo", "Cód"]:
                    res = page.search_for(l)
                    # Debe estar cerca en Y del consecutivo
                    if res and header_consec and abs(res[0].y0 - header_consec.y0) < 15:
                        header_cod = res[0]
                        break
                        
                # 3. Buscar encabezado de Descripcion/Nombre
                header_desc = None
                for l in ["Descripción", "Descripcion", "Nombre"]:
                    res = page.search_for(l)
                    if res and header_consec and abs(res[0].y0 - header_consec.y0) < 15:
                        header_desc = res[0]
                        break

                # 4. Buscar encabezado Cantidad
                header_cant = None
                for l in ["Cant", "Cantidad", "Cant."]:
                    res = page.search_for(l)
                    if res and header_consec and abs(res[0].y0 - header_consec.y0) < 15:
                        header_cant = res[0]
                        break
                
                if header_consec:
                    y_top = header_consec.y1 + 2
                    # y_bot_single: Para campos que deben ser de una sola línea (Consecutivo, Codigo, Cantidad)
                    # Reducimos drásticamente la altura para evitar leer la siguiente línea (Observación, etc.)
                    y_bot_single = y_top + 12 
                    
                    # y_bot_desc: Para la descripción, permitimos más altura porque puede ser multilinea
                    y_bot_desc = y_top + 45
                    
                    # Definir limites X
                    x_consec = header_consec.x0
                    x_cod = header_cod.x0 if header_cod else x_consec + 80
                    x_desc = header_desc.x0 if header_desc else x_cod + 80
                    x_cant = header_cant.x0 if header_cant else page.rect.width - 50
                    
                    # Extraer Consecutivo (Single Line Strict)
                    r_consec = fitz.Rect(x_consec, y_top, x_cod, y_bot_single)
                    val_consec = page.get_text("text", clip=r_consec).strip()
                    if '\n' in val_consec: val_consec = val_consec.split('\n')[0].strip()
                    # Limpieza extra: Si aparece "Observación" o digitos raros pegados
                    if "Observ" in val_consec: val_consec = val_consec.split("Observ")[0].strip()
                    record['Consecutivo'] = val_consec
                    
                    # Extraer Codigo (Single Line Strict)
                    r_cod = fitz.Rect(x_cod, y_top, x_desc, y_bot_single)
                    val_cod = page.get_text("text", clip=r_cod).strip()
                    if '\n' in val_cod: val_cod = val_cod.split('\n')[0].strip()
                    record['Codigo Servicio'] = val_cod
                    
                    # Extraer Nombre/Descripcion (Multi Line)
                    r_desc = fitz.Rect(x_desc, y_top, x_cant, y_bot_desc)
                    nom_serv = page.get_text("text", clip=r_desc).strip().replace('\n', ' ')
                    
                    # Extraer Cantidad (Single Line Strict)
                    r_cant = fitz.Rect(x_cant, y_top, page.rect.width, y_bot_single)
                    val_cant = page.get_text("text", clip=r_cant).strip()
                    if '\n' in val_cant: val_cant = val_cant.split('\n')[0].strip()
                    # Limpieza de emails o horas que aparecen abajo
                    for garbage in ["@", ":", "fomag"]: 
                        if garbage in val_cant: val_cant = val_cant.split(garbage)[0].strip()
                        # Si cortamos por : (hora), a veces queda el número antes. Ej "1 6:25" -> "1 6".
                        # Mejor: tomar solo el primer token numérico si hay espacios
                        if ' ' in val_cant:
                             tokens = val_cant.split()
                             # Tomar el primer token que sea numero o digito
                             if tokens: val_cant = tokens[0]
                    record['Cantidad'] = val_cant

                    # BUSCAR OBSERVACIÓN (Aparece debajo, usualmente alineada a la izquierda)
                    # Definimos un area amplia debajo de la fila principal
                    # Buscamos DESDE y_bot_single hacia abajo
                    r_obs_search = fitz.Rect(0, y_bot_single, page.rect.width, y_bot_desc + 20)
                    text_obs_area = page.get_text("text", clip=r_obs_search)
                    
                    # Buscamos la línea que empiece con Observación:
                    m_obs = re.search(r'(Observaci[óo]n:.*?)(?:\n|$)', text_obs_area, re.IGNORECASE)
                    if m_obs:
                        obs_text = m_obs.group(1).strip()
                        # Solo agregar si no está ya en el nombre (a veces se lee doble)
                        if obs_text not in nom_serv:
                            nom_serv += f" | {obs_text}"
                    
                    record['Nombre Servicio'] = nom_serv

                else:
                    # Fallback si no encuentra encabezados
                    pass

            extracted_data.append(record)

        except Exception as e:
            if not silent_mode: st.error(f"Error procesando {os.path.basename(pdf_path)}: {e}")

    if extracted_data:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df = pd.DataFrame(extracted_data)
            df.to_excel(writer, index=False, sheet_name='Resultados')
            
            # Ajuste de columnas
            workbook = writer.book
            worksheet = writer.sheets['Resultados']
            for i, col in enumerate(df.columns):
                worksheet.set_column(i, i, 20)

        return {
            "files": [{
                "name": "Analisis_Autorizaciones_FOMAG.xlsx",
                "data": output.getvalue(),
                "label": "Descargar FOMAG"
            }],
            "message": f"Procesados: {len(extracted_data)} registros."
        }
    return None


def worker_analisis_sos(file_list, silent_mode=False, use_ai=False, api_key=None):
    """
    Analiza archivos PDF de SOS (Autorizaciones).
    Soporta modo 'Studio' (Reglas/PDFPlumber) y opcionalmente IA.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(file_list):
        if not silent_mode: st.info(f"Delegando análisis SOS al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            payload = {
                "file_list": file_list,
                "use_ai": use_ai,
                "api_key": api_key
            }
            task_id = send_command(username, "analisis_sos", payload)
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=600) # Longer timeout for AI
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    # ... Rest of local implementation ...
    if not file_list: return None
    
    archivos_pdf = [f for f in file_list if f.lower().endswith('.pdf')]
    if not archivos_pdf: return None
    
    extracted_data = []
    
    # Helper for SOS extraction (Studio Logic)
    def extract_sos_studio(pdf_path):
        if not pdfplumber: return {}
        data_res = {"valid_extraction": False}
        try:
            with pdfplumber.open(pdf_path) as pdf:
                if not pdf.pages: return {}
                pagina = pdf.pages[0]
                texto = pagina.extract_text() or ""
                
                # Regex Extraction
                patrones = {
                    "Fecha Consulta": r"Fecha Consulta[:\s]+(\d{2}/\d{2}/\d{4})",
                    "Identificación": r"Identificación[:\s]+(\d+)",
                    "Afiliado": r"Afiliado[:\s]+(.+?)(?=\s+Identificación|Plan|\n|$)",
                    "Plan": r"Plan[:\s]+(.+?)(?=\s+Rango|\n|$)",
                    "Derecho": r"Derecho[:\s]+(.+?)(?=\s+Ambito|\s+IPS Primaria|\n|$)",
                    "IPS Primaria": r"IPS Primaria[:\s]+(.+?)(?=\s+IPS Solicitante|\n|$)",
                    "IPS Solicitante": r"IPS Solicitante[:\s]+(.+?)(?=\n|$)",
                    "Ambito": r"Ambito[:\s]+([A-Z\s]+)"
                }
                for k, p in patrones.items():
                    m = re.search(p, texto, re.IGNORECASE)
                    if m: data_res[k] = m.group(1).strip()

                # Table Extraction
                tabla = pagina.extract_table()
                if tabla:
                    codigos, nombres, cantidades, respuestas, autorizaciones = [], [], [], [], []
                    for fila in tabla:
                        if not fila: continue
                        row_str = "".join([str(c) for c in fila if c]).lower()
                        if "código" in row_str and "prestación" in row_str: continue
                        if "autorizador" in row_str and "linea" in row_str: continue
                        if "(cid:" in row_str: return {"valid_extraction": False} # Garbage check

                        val_codigo, val_nombre, val_cant, val_resp, val_auth = "", "", "", "", ""

                        # 4-col format check
                        if len(fila) == 4 and (str(fila[0]).isdigit() or str(fila[3]).isdigit()):
                            val_codigo = str(fila[0]).strip()
                            val_resp = str(fila[1]).strip()
                            val_auth = str(fila[3]).strip()
                            val_nombre = f"Ver P-Autorización {val_codigo}"
                            val_cant = "1"
                        else:
                            val_codigo = str(fila[0]).strip() if len(fila) > 0 and fila[0] else ""
                            val_nombre = str(fila[1]).strip().replace("\n", " ") if len(fila) > 1 and fila[1] else ""
                            val_cant = str(fila[2]).strip() if len(fila) > 2 and fila[2] else ""
                            val_resp = str(fila[3]).strip() if len(fila) > 3 and fila[3] else ""
                            if len(fila) > 7 and fila[7]: val_auth = str(fila[7]).strip()
                        
                        if val_codigo and not val_codigo.replace(".","").isdigit(): continue

                        if val_codigo or val_nombre:
                            if val_codigo: codigos.append(val_codigo)
                            if val_nombre: nombres.append(val_nombre)
                            if val_cant: cantidades.append(val_cant)
                            if val_resp: respuestas.append(val_resp)
                            if val_auth: autorizaciones.append(val_auth)
                            data_res["valid_extraction"] = True
                    
                    if codigos: data_res["Código Prestación"] = " | ".join(codigos)
                    if nombres: data_res["Nombre Prestación"] = " | ".join(nombres)
                    if cantidades: data_res["Cantidad"] = " | ".join(cantidades)
                    if respuestas: data_res["Respuesta EPS"] = " | ".join(respuestas)
                    if autorizaciones: data_res["No. Autorización"] = " | ".join(autorizaciones)
            
            return data_res
        except Exception:
            return {}

    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0)
        status_text = st.empty()

    for i, pdf_path in enumerate(archivos_pdf):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(archivos_pdf))
            status_text.text(f"Procesando: {os.path.basename(pdf_path)}")
            
        record = {
            "Archivo": os.path.basename(pdf_path),
            "Fecha Consulta": "", "Afiliado": "", "Identificación": "", "Plan": "",
            "IPS Primaria": "", "Código Prestación": "", "Nombre Prestación": "",
            "Cantidad": "", "Respuesta EPS": "", "No. Autorización": "",
            "Ambito": "", "Derecho": "", "IPS Solicitante": ""
        }
        
        # Try Studio Extraction first
        studio_data = extract_sos_studio(pdf_path)
        if studio_data.get("valid_extraction"):
            record.update(studio_data)
            record["_DEBUG_STRATEGY"] = "Studio_Rules"
        elif use_ai and api_key and genai:
             # Placeholder for AI logic if needed, but keeping it simple for now
             # Since the module had complex AI logic, we can defer or implement later if requested.
             # For now, we focus on the rule-based part which is robust.
             record["_DEBUG_STRATEGY"] = "AI_Not_Implemented_Yet"
        else:
            record["_DEBUG_STRATEGY"] = "Failed"

        extracted_data.append(record)

    if extracted_data:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            pd.DataFrame(extracted_data).to_excel(writer, index=False)
        return {
            "files": [{
                "name": "Analisis_SOS.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Análisis SOS"
            }],
            "message": f"Procesados: {len(extracted_data)} registros."
        }
    return None

def worker_analisis_autorizacion_nueva_eps(file_list, silent_mode=False):
    """
    Analiza archivos PDF de Autorizaciones Nueva EPS usando PyMuPDF (fitz).
    Retorna bytes de Excel.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(file_list):
        if not silent_mode: st.info(f"Delegando análisis Nueva EPS al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_neps", {"file_list": file_list})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    if not fitz:
        if not silent_mode: st.error("Librería 'fitz' (PyMuPDF) no instalada.")
        return None

    data_res = []
    
    # Regex patterns (from OrganizadorArchivos_v1.py)
    patterns = {
        'Afiliado': re.compile(r"Afiliado:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'N° Autorización': re.compile(r"N° Autorización:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Autorizada el': re.compile(r"Autorizada el:\s*(.*?)(?:\n|$)", re.IGNORECASE),
        'Descripción Servicio': re.compile(r"Descripción Servicio\s*\n\s*\d+\s+\d+\s+(.*?)(?:\n|$)", re.IGNORECASE | re.DOTALL),
        'Info de Pago': re.compile(r"(Afiliado (?:No )?Cancela.*?)(?:\n|$)", re.IGNORECASE)
    }

    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Analizando Autorizaciones...")

    for i, file_path in enumerate(file_list):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(file_list), text=f"Procesando: {os.path.basename(file_path)}")

        try:
            full_text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            
            row = {'Archivo': os.path.basename(file_path)}
            for key, pattern in patterns.items():
                match = pattern.search(full_text)
                if match:
                    val = match.group(1).strip()
                    # Clean up 'Descripción Servicio' which might capture too much
                    if key == 'Descripción Servicio':
                        val = val.split('\n')[0].strip()
                    row[key] = val
                else:
                    row[key] = ""
            data_res.append(row)
        except Exception as e:
            if not silent_mode: st.warning(f"Error en {os.path.basename(file_path)}: {e}")
            data_res.append({'Archivo': os.path.basename(file_path), 'Error': str(e)})

    if data_res:
        # Define column order as in original
        column_order = ['Archivo', 'Afiliado', 'N° Autorización', 'Autorizada el', 'Descripción Servicio', 'Info de Pago']
        # Ensure all columns exist
        for col in column_order:
            if col not in data_res[0]: # Check first row structure mostly
                 pass # Pandas handles missing cols but better to be safe? 
                      # Actually constructing DataFrame from dict list handles it fine.
        
        df = pd.DataFrame(data_res)
        # Reorder if columns present
        existing_cols = [c for c in column_order if c in df.columns]
        other_cols = [c for c in df.columns if c not in column_order]
        df = df[existing_cols + other_cols]

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False)
        return {
            "files": [{
                "name": "Analisis_Autorizaciones_NuevaEPS.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Autorizaciones"
            }],
            "message": f"Procesados: {len(data_res)} registros."
        }
    return None

def worker_analisis_cargue_sanitas(file_list, silent_mode=False):
    """
    Analiza archivos PDF de Cargue Sanitas (FEOV) usando PyMuPDF (fitz).
    Retorna bytes de Excel.
    """
    # --- Agent Delegation ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    if is_native_mode and not silent_mode and _should_delegate(file_list):
        if not silent_mode: st.info(f"Delegando análisis Sanitas al Agente Local...")
        try:
            from src.agent_client import send_command, wait_for_result
            username = st.session_state.get("username", "admin")
            task_id = send_command(username, "analisis_sanitas", {"file_list": file_list})
            if not task_id: return {"error": "No se pudo enviar la tarea al agente."}
            
            res = wait_for_result(task_id, timeout=300)
            if res and "error" not in res:
                return res
            else:
                return {"error": f"Error en agente: {res.get('error') if res else 'Sin respuesta'}"}
        except Exception as e:
            return {"error": f"Fallo en delegación a agente: {e}"}
    # ------------------------

    if not fitz:
        if not silent_mode: st.error("Librería 'fitz' (PyMuPDF) no instalada.")
        return None

    data_res = []
    
    patterns = {
        'Factura (FEOV)': re.compile(r"FEOV(\d+)", re.IGNORECASE),
        'Fecha y hora de cargue': re.compile(r"(\d{1,2}\s+\w+\s+\d{4}\s*-\s*\d{1,2}:\d{2})", re.IGNORECASE)
    }

    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Analizando Cargue Sanitas...")

    for i, file_path in enumerate(file_list):
        if not silent_mode and progress_bar:
            progress_bar.progress((i + 1) / len(file_list), text=f"Procesando: {os.path.basename(file_path)}")

        try:
            full_text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    full_text += page.get_text("text") + "\n"
            
            row = {'Archivo': os.path.basename(file_path)}
            for key, pattern in patterns.items():
                match = pattern.search(full_text)
                row[key] = match.group(1).strip() if match else ""
            
            # Extract numeric FEOV
            feov_match = patterns['Factura (FEOV)'].search(full_text)
            if feov_match:
                row['Factura (FEOV)'] = feov_match.group(1)

            data_res.append(row)
        except Exception as e:
            if not silent_mode: st.warning(f"Error en {os.path.basename(file_path)}: {e}")
            data_res.append({'Archivo': os.path.basename(file_path), 'Error': str(e)})

    if data_res:
        column_order = ['Archivo', 'Factura (FEOV)', 'Fecha y hora de cargue']
        df = pd.DataFrame(data_res)
        
        # Reorder
        existing_cols = [c for c in column_order if c in df.columns]
        other_cols = [c for c in df.columns if c not in column_order]
        df = df[existing_cols + other_cols]

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False)
        return {
            "files": [{
                "name": "Analisis_Cargue_Sanitas.xlsx",
                "data": output.getvalue(),
                "label": "Descargar Sanitas"
            }],
            "message": f"Procesados: {len(data_res)} registros."
        }
    return None

# --- WORKERS: WEB SCRAPING & DOWNLOADS ---

def worker_descargar_firmas(uploaded_file, sheet_name, col_id, col_folder, root_path=None, silent_mode=False):
    """
    Descarga firmas desde una URL base usando un Excel para mapear IDs a carpetas.
    Soporta modo nativo (Agente) y modo web.
    """
    if not requests or not Image:
        if not silent_mode: st.error("Faltan librerías: requests o Pillow.")
        return {"error": "Librerías faltantes."}

    try:
        is_native = st.session_state.get("force_native_mode", True)

        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
        base_url = "https://oportunidaddevida.com/opvcitas/admisionescall/firmas/"
        
        # Determine output directory
        is_temp = False
        if not root_path:
            is_temp = True
            root_path = os.path.join(os.getcwd(), "temp_downloads", f"firmas_{int(time.time())}")
        
        # Collect tasks
        tasks = []
        preview_dest = ""
        
        total = len(df)
        progress_bar = None
        status_text = None

        if not silent_mode:
            progress_bar = st.progress(0)
            status_text = st.empty()

        for i, row in df.iterrows():
            if not silent_mode and progress_bar:
                progress_bar.progress((i + 1) / total)

            id_firma = str(row[col_id]).strip()
            nombre_carpeta = str(row[col_folder]).strip()
            
            if not id_firma or not nombre_carpeta or pd.isna(row[col_id]) or pd.isna(row[col_folder]):
                continue

            # Use helper to find folder path (supports search results)
            # In Native Mode, this returns the constructed path string (G:/...) even if not exists locally
            target_dir = find_folder_path(root_path, nombre_carpeta)
            dest_path = os.path.join(target_dir, "tipografia", "firma.jpg")
            url_completa = f"{base_url}{id_firma}.png"
            
            tasks.append({
                "url": url_completa,
                "dest_path": dest_path
            })
            
            if not preview_dest: preview_dest = dest_path

        if not tasks:
            return "No se encontraron registros válidos."

        if is_native:
            # Agent Mode
            import agent_client
            if not agent_client:
                 return "Error: Módulo agent_client no disponible."
            
            username = st.session_state.get("username", "default")
            
            print(f"DEBUG: Enviando {len(tasks)} descargas al agente. Primer destino: {preview_dest}")
            
            task_id = agent_client.send_command(username, "download_files", {
                "tasks": tasks
            })
            
            if not silent_mode: 
                progress_bar.empty()
                status_text.empty()
                
            if task_id:
                return f"Tarea enviada al agente (ID: {task_id}). Archivos a descargar: {len(tasks)}. (Ejemplo destino: {preview_dest})"
            else:
                return "Error: No se pudo enviar la tarea al agente."

        else:
            # Web Mode (Server Side)
            descargados = 0
            errores = 0
            
            os.makedirs(root_path, exist_ok=True)
            
            for task in tasks:
                url = task.get("url")
                dest = task.get("dest_path")
                target_dir = os.path.dirname(dest)
                os.makedirs(target_dir, exist_ok=True)
                
                try:
                    response = requests.get(url, stream=True, timeout=10)
                    if response.status_code == 200:
                        if not response.content:
                            raise ValueError("Contenido vacío")
                        img = Image.open(io.BytesIO(response.content)).convert('RGB')
                        img.save(dest, "JPEG")
                        descargados += 1
                    else:
                        with open(os.path.join(target_dir, "no tiene firma.txt"), 'w') as f:
                            f.write(f"No firma: {url} - {response.status_code}")
                        errores += 1
                except Exception:
                    errores += 1

            if not silent_mode: 
                progress_bar.empty()
                status_text.empty()

            result_msg = f"Proceso finalizado (Servidor). Descargados: {descargados}. Errores/No encontrados: {errores}."
            
            if is_temp:
                 # Zip logic (kept for web mode fallback, though user wants it removed, keeping it for web mode is safer)
                 # ... (omitted for brevity, user mainly uses native mode)
                 return {"message": result_msg} # Just return message to avoid zip button in native context
            
            return {"message": result_msg}

    except Exception as e:
        return {"error": f"Error crítico: {e}"}

def worker_descargar_historias_ovida(uploaded_file, sheet_name, col_estudio, col_ingreso, col_egreso, col_carpeta, download_path=None, silent_mode=False):
    """
    Descarga historias clínicas de OVIDA usando Selenium (Chrome Headless/GUI).
    """
    try:
        from selenium import webdriver
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.chrome.service import Service
        from webdriver_manager.chrome import ChromeDriverManager
    except ImportError:
        return {"error": "Error: Selenium/WebDriverManager no instalado."}

    # Determine output directory
    is_temp = False
    if not download_path:
        is_temp = True
        download_path = os.path.join(os.getcwd(), "temp_downloads", f"ovida_{int(time.time())}")
    
    # --- NATIVE MODE CHECK ---
    is_native_mode = st.session_state.get('force_native_mode', True)
    
    if is_native_mode:
        # En modo nativo, download_path debe ser una ruta válida en la máquina del cliente
        # Si no se proveyó una ruta (is_temp=True), no podemos continuar en modo nativo 
        # a menos que el agente maneje descargas temporales (que no es el caso usual).
        # Asumiremos que si está en modo nativo, el usuario seleccionó una carpeta o se usará una por defecto.
        if is_temp:
             # Si no hay ruta seleccionada, intentar usar una por defecto en el cliente? 
             # Mejor pedir que seleccionen carpeta.
             pass 
    else:
        os.makedirs(download_path, exist_ok=True)

    try:
        if isinstance(uploaded_file, bytes):
            uploaded_file = io.BytesIO(uploaded_file)
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
    except Exception as e:
        return {"error": f"Error leyendo Excel: {e}"}

    # --- AGENT EXECUTION ---
    if is_native_mode:
        try:
            # Prepare records for agent
            records = []
            for _, row in df.iterrows():
                try:
                    # Extract values based on column mappings
                    estudio = str(row[col_estudio]).strip() if col_estudio in row else ""
                    ingreso = row[col_ingreso] if col_ingreso in row else ""
                    egreso = row[col_egreso] if col_egreso in row else ""
                    carpeta = str(row[col_carpeta]).strip() if col_carpeta in row else ""
                    
                    # Convert dates to string format expected by agent (YYYY/MM/DD) if they are datetime objects
                    if isinstance(ingreso, pd.Timestamp):
                        ingreso = ingreso.strftime('%Y/%m/%d')
                    else:
                         ingreso = str(ingreso)
                         
                    if isinstance(egreso, pd.Timestamp):
                        egreso = egreso.strftime('%Y/%m/%d')
                    else:
                        egreso = str(egreso)

                    records.append({
                        "nro_estudio": estudio,
                        "fecha_ingreso": ingreso,
                        "fecha_salida": egreso,
                        "rel_path": carpeta
                    })
                except Exception as row_err:
                    print(f"Error procesando fila para agente: {row_err}")
                    continue

            if not records:
                return {"error": "No se encontraron registros válidos para procesar."}

            if not send_command:
                 return {"error": "Error: Modo nativo activado pero el cliente del agente no está disponible."}
            
            username = st.session_state.get("username", "admin")
            
            if not silent_mode:
                st.info(f"Enviando tarea al agente local para descargar {len(records)} historias...")
                
            task_id = send_command(username, "download_ovida", {
                "base_path": download_path,
                "records": records
            })
            
            if task_id:
                # Poll
                status_placeholder = st.empty()
                if not silent_mode:
                    status_placeholder.text("Esperando agente (esto puede tomar varios minutos)...")
                
                # Timeout largo para descargas masivas
                res = wait_for_result(task_id, timeout=600) 
                
                if not silent_mode: status_placeholder.empty()

                if res and "status" in res:
                    if res["status"] == "success":
                         msg = res.get("message", "Finalizado correctamente")
                         stats = res.get("stats", {})
                         return {"message": f"{msg} (Agente)", "stats": stats}
                    else:
                         return {"error": f"Error del agente: {res.get('message', 'Desconocido')}"}
                elif res and "error" in res: # Fallback for other error formats
                     return {"error": f"Error del agente: {res.get('error')}"}
                else:
                    return {"error": "Tiempo de espera agotado o respuesta inválida del agente."}
            else:
                return {"error": "No se pudo crear la tarea en el servidor."}

        except Exception as e:
            return {"error": f"Excepción preparando tarea agente: {str(e)}"}

    # --- SERVER SIDE EXECUTION (Legacy) ---
    driver = None
    try:
        options = webdriver.ChromeOptions()
        prefs = {
            "download.default_directory": download_path,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True
        }
        options.add_experimental_option("prefs", prefs)
        
        # Open visible browser for login
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
        
        if not silent_mode:
            st.warning("⚠️ Se abrió una ventana de Chrome. INICIE SESIÓN en OVIDA manualmente. El proceso continuará automáticamente cuando detecte el ingreso.")
        
        # Wait for login (detect change to main page or timeout)
        # Timeout 5 minutes
        timeout = 300 
        start_time = time.time()
        while time.time() - start_time < timeout:
             try:
                 # Check if URL changed from login page
                 if "iniciando.php" not in driver.current_url and "login" not in driver.current_url.lower():
                     break
             except:
                 pass
             time.sleep(1)
        
        if time.time() - start_time >= timeout:
            driver.quit()
            return {"error": "Error: Tiempo de espera de inicio de sesión agotado."}
        
        # Additional wait for session to stabilize
        time.sleep(2)
        
        # Verify login success
        logged_in = False
        start_check = time.time()
        while time.time() - start_check < 30:
            try:
                if "App/Vistas" in driver.current_url:
                    logged_in = True
                    break
            except: pass
            time.sleep(1)
            
        if not logged_in:
             # Just a warning, maybe URL is different but session is active
             pass

        if not silent_mode: st.info("Inicio de sesión detectado. Comenzando descargas...")

        descargados = 0
        errores = 0
        conflictos = 0
        
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0)
            status_text = st.empty()
            
        total = len(df)
        
        for i, row in df.iterrows():
            if not silent_mode and progress_bar:
                progress_bar.progress((i + 1) / total)
            
            try:
                estudio = str(int(row[col_estudio])).strip()
                f_ing = pd.to_datetime(row[col_ingreso]).strftime('%Y/%m/%d')
                f_egr = pd.to_datetime(row[col_egreso]).strftime('%Y/%m/%d')
                carpeta = str(row[col_carpeta]).strip()
                
                if not all([estudio, f_ing, f_egr, carpeta]):
                    errores += 1
                    continue

                dest_dir = os.path.join(download_path, carpeta)
                os.makedirs(dest_dir, exist_ok=True)
                final_path = os.path.join(dest_dir, f"HC_{estudio}.pdf")
                
                if os.path.exists(final_path):
                    conflictos += 1
                    continue
                    
                if not silent_mode: status_text.text(f"Descargando Estudio: {estudio}")

                # URL construction
                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                params = {
                    'estudio': estudio, 'fecha_inicio': f_ing, 'fecha_fin': f_egr,
                    'verHC': 1, 'verEvo': 1, 'verPar': 1, 'ImprimirOrdenamiento': 1,
                    'ImprimirNotasPcte': 0, 'ImprimirSolOrdenesExt': 1, 'ImprimirGraficasHC': 1,
                    'ImprimirFormatos': 1, 'ImprimirRegistroAdmon': 1, 'ImprimirNovedad': 0,
                    'ImprimirRecomendaciones': 0, 'ImprimirDescripcionQX': 0, 'ImprimirNotasEnfermeria': 1,
                    'ImprimirSignosVitales': 0, 'ImprimirLog': 0, 'ImprimirEpicrisisSinHC': 0
                }
                full_url = f"{base_url}?{urllib.parse.urlencode(params)}"
                
                driver.get(full_url)
                time.sleep(2) # Wait for render
                
                pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {
                    "landscape": False, "printBackground": True,
                    "paperWidth": 8.5, "paperHeight": 11,
                    "marginTop": 0.4, "marginBottom": 0.4, "marginLeft": 0.4, "marginRight": 0.4
                })
                
                pdf_data = base64.b64decode(pdf_b64['data'])
                with open(final_path, 'wb') as f:
                    f.write(pdf_data)
                
                descargados += 1
                
            except Exception as e:
                errores += 1
                if not silent_mode: st.warning(f"Error en estudio {estudio}: {e}")
        
        result_msg = f"Finalizado. Descargados: {descargados}, Errores: {errores}, Conflictos: {conflictos}."
        
        if is_temp:
            # Zip results
            mem_zip = io.BytesIO()
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(download_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, download_path)
                        zf.write(file_path, arcname)
            
            # Cleanup
            try:
                shutil.rmtree(download_path, ignore_errors=True)
            except: pass
            
            return {
                "files": [{
                    "name": f"Historias_OVIDA_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "Descargar Historias (ZIP)"
                }],
                "message": result_msg
            }
        else:
            return {"message": result_msg}

    except Exception as e:
        return {"error": f"Error crítico: {e}"}
    finally:
        if driver: driver.quit()

# --- DIALOGS FOR DOWNLOADS ---

@st.dialog("Descargar Firmas (Excel)")
def dialog_descargar_firmas():
    st.write("Cargue un Excel con IDs de firma y nombres de carpeta.")
    
    uploaded = st.file_uploader("Archivo Excel", type=["xlsx", "xls"], key="firmas_up")
    
    sheet_name = "Hoja1"
    cols = []
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet_name = st.selectbox("Seleccione la Hoja", sheet_names, key="firmas_sheet_sel")
            df_preview = _get_excel_preview(file_bytes, sheet_name, nrows=1)
            cols = df_preview.columns.tolist()
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")

    if cols:
        col_id = st.selectbox("Columna ID Firma", cols, index=cols.index("id_firma") if "id_firma" in cols else 0, key="firmas_col_id_sel")
        col_folder = st.selectbox("Columna Nombre Carpeta", cols, index=cols.index("nombre_carpeta") if "nombre_carpeta" in cols else 0, key="firmas_col_folder_sel")
    else:
        col_id = st.text_input("Columna ID Firma", value="id_firma", key="firmas_col_id")
        col_folder = st.text_input("Columna Nombre Carpeta", value="nombre_carpeta", key="firmas_col_folder")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    root_path = render_path_selector(
        key="firmas_path",
        label="Ruta Raíz Descarga",
        default_path=default_path
    )
    
    if st.button("Iniciar Descarga"):
        if uploaded and sheet_name and col_id and col_folder and root_path:
            try:
                with st.spinner("Descargando firmas..."):
                    uploaded.seek(0)
                    result = worker_descargar_firmas(uploaded, sheet_name, col_id, col_folder, root_path)
                    st.success(result)
                    close_auto_dialog()
                    # render_download_button(root_path, "dl_sigs_root", "📦 Descargar Firmas (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.error("Complete todos los campos.")

    if st.button("Cerrar", key="btn_close_desc_firmas"):
        close_auto_dialog()

@st.dialog("Descargar Historias OVIDA")
def dialog_descargar_historias_ovida():
    st.write("Automatización de descargas desde OVIDA (Requiere Credenciales).")
    st.warning("Se abrirá un navegador Chrome. Debe iniciar sesión manualmente cuando se indique.")
    
    uploaded = st.file_uploader("Archivo Excel (Pacientes)", type=["xlsx", "xls"], key="ovida_up")
    
    sheet_name = "Hoja1"
    cols = []
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet_name = st.selectbox("Seleccione la Hoja", sheet_names, key="ovida_sheet_sel")
            df_preview = _get_excel_preview(file_bytes, sheet_name, nrows=1)
            cols = df_preview.columns.tolist()
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")

    c1, c2 = st.columns(2)
    with c1:
        if cols:
            col_estudio = st.selectbox("Columna Estudio", cols, index=cols.index("estudio") if "estudio" in cols else 0, key="ovida_est_sel")
            col_ingreso = st.selectbox("Columna Fecha Ingreso", cols, index=cols.index("f_ingreso") if "f_ingreso" in cols else 0, key="ovida_ing_sel")
        else:
            col_estudio = st.text_input("Columna Estudio", value="estudio", key="ovida_est")
            col_ingreso = st.text_input("Columna Fecha Ingreso", value="f_ingreso", key="ovida_ing")
            
    with c2:
        if cols:
            col_egreso = st.selectbox("Columna Fecha Egreso", cols, index=cols.index("f_egreso") if "f_egreso" in cols else 0, key="ovida_egr_sel")
            col_carpeta = st.selectbox("Columna Carpeta Destino", cols, index=cols.index("carpeta") if "carpeta" in cols else 0, key="ovida_carp_sel")
        else:
            col_egreso = st.text_input("Columna Fecha Egreso", value="f_egreso", key="ovida_egr")
            col_carpeta = st.text_input("Columna Carpeta Destino", value="carpeta", key="ovida_carp")
        
    default_path = st.session_state.get("current_path", os.getcwd())
    download_path = render_path_selector(
        key="ovida_path",
        label="Ruta Descarga Base",
        default_path=default_path
    )
    
    if st.button("Iniciar Descarga Masiva"):
        if uploaded and sheet_name and col_estudio and col_ingreso and col_egreso and col_carpeta and download_path:
            # Re-read file to ensure pointer is at start or handled by worker
            uploaded.seek(0)
            try:
                with st.spinner("Descargando historias de OVIDA..."):
                    result = worker_descargar_historias_ovida(uploaded, sheet_name, col_estudio, col_ingreso, col_egreso, col_carpeta, download_path)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(download_path, "dl_ovida_root", "📦 Descargar Historias (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.error("Complete todos los campos.")

    if st.button("Cerrar", key="btn_close_desc_ovida"):
        close_auto_dialog()

def worker_organizar_facturas_por_pdf_avanzado(carpeta_destinos, carpeta_origen, silent_mode=False):
    try:
        regex = re.compile(r'FEOV(\d+)', re.IGNORECASE)
        destinos_map = {}
        
        # 1. Map destinations
        list_carpetas_destino = [d for d in os.listdir(carpeta_destinos) if os.path.isdir(os.path.join(carpeta_destinos, d))]
        
        for nombre_carpeta_destino in list_carpetas_destino:
            ruta_carpeta_destino = os.path.join(carpeta_destinos, nombre_carpeta_destino)
            for archivo in os.listdir(ruta_carpeta_destino):
                if archivo.lower().endswith('.pdf'):
                    match = regex.search(archivo)
                    if match:
                        numero_factura = match.group(1)
                        destinos_map[numero_factura] = ruta_carpeta_destino
                        break
        
        if not destinos_map:
            return "No se encontraron PDFs con patrón FEOV en las carpetas de destino."
            
        # 2. Move files
        movidos, conflictos, errores = 0, 0, 0
        
        files_to_move = []
        for root, _, files in os.walk(carpeta_origen):
            for f in files:
                files_to_move.append((root, f))
                
        if not files_to_move:
            return "No hay archivos en la carpeta de origen."
            
        if not silent_mode:
            progress_bar = st.progress(0, text="Organizando...")
            
        for i, (root, file_to_move) in enumerate(files_to_move):
            if not silent_mode and i % 10 == 0:
                progress_bar.progress((i + 1) / len(files_to_move), text=f"Procesando: {file_to_move}")
                
            moved_this = False
            for numero_factura, ruta_destino_final in destinos_map.items():
                if numero_factura in file_to_move:
                    try:
                        ruta_origen_archivo = os.path.join(root, file_to_move)
                        ruta_final_archivo = os.path.join(ruta_destino_final, file_to_move)

                        if os.path.exists(ruta_final_archivo):
                            conflictos += 1
                        else:
                            shutil.move(ruta_origen_archivo, ruta_destino_final)
                            movidos += 1
                        moved_this = True
                        break
                    except Exception:
                        errores += 1
                        break
            
        if not silent_mode:
            progress_bar.empty()
            
        return f"Movidos: {movidos}, Conflictos: {conflictos}, Errores: {errores}"
    except Exception as e:
        return f"Error: {e}"

def worker_json_evento_a_xlsx_masivo(carpeta_origen, archivo_salida, silent_mode=False):
    if not carpeta_origen or not archivo_salida: return "Rutas inválidas."
    
    archivos_json = []
    for root, dirs, files in os.walk(carpeta_origen):
        for file in files:
            if file.lower().endswith(".json"):
                archivos_json.append(os.path.join(root, file))
    
    if not archivos_json:
        return "No se encontraron archivos JSON en la carpeta seleccionada."
        
    progress_bar = None
    if not silent_mode:
        progress_bar = st.progress(0, text="Consolidando JSONs...")
    
    try:
        todas_consultas = []
        todos_procedimientos = []
        todos_otros_servicios = []
        errores = 0
        
        for i, ruta_json in enumerate(archivos_json):
            if not silent_mode:
                progress_bar.progress((i + 1) / len(archivos_json), text=f"Procesando: {os.path.basename(ruta_json)}")
            
            try:
                with open(ruta_json, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                usuarios_lista = data.get("usuarios", []) if isinstance(data, dict) else []
                
                for usuario in usuarios_lista:
                    base_info = {
                        "archivo_origen": os.path.basename(ruta_json),
                        "tipo_documento_usuario": usuario.get("tipoDocumentoIdentificacion"),
                        "documento_usuario": usuario.get("numDocumentoIdentificacion"),
                        "tipo_usuario": usuario.get("tipoUsuario"),
                        "fecha_nacimiento": usuario.get("fechaNacimiento"),
                        "sexo": usuario.get("codSexo"),
                        "pais_residencia": usuario.get("codPaisResidencia"),
                        "municipio_residencia": usuario.get("codMunicipioResidencia"),
                        "zona_residencia": usuario.get("codZonaTerritorialResidencia"),
                        "incapacidad": usuario.get("incapacidad"),
                        "consecutivo_usuario": usuario.get("consecutivo"),
                        "pais_origen": usuario.get("codPaisOrigen")
                    }

                    servicios = usuario.get("servicios", {})

                    for consulta in servicios.get("consultas", []):
                        todas_consultas.append({**base_info, **consulta})

                    for procedimiento in servicios.get("procedimientos", []):
                        todos_procedimientos.append({**base_info, **procedimiento})

                    for otro in servicios.get("otrosServicios", []):
                        todos_otros_servicios.append({**base_info, **otro})
                    
            except Exception:
                errores += 1
        
        if todas_consultas or todos_procedimientos or todos_otros_servicios:
            with pd.ExcelWriter(archivo_salida, engine="openpyxl") as writer:
                if todas_consultas:
                    pd.DataFrame(todas_consultas).to_excel(writer, sheet_name="Consultas", index=False)
                if todos_procedimientos:
                    pd.DataFrame(todos_procedimientos).to_excel(writer, sheet_name="Procedimientos", index=False)
                if todos_otros_servicios:
                    pd.DataFrame(todos_otros_servicios).to_excel(writer, sheet_name="OtrosServicios", index=False)
                
                if not (todas_consultas or todos_procedimientos or todos_otros_servicios):
                     pd.DataFrame().to_excel(writer, sheet_name="Vacio", index=False)

            if not silent_mode: progress_bar.empty()
            total_reg = len(todas_consultas) + len(todos_procedimientos) + len(todos_otros_servicios)
            return f"Consolidación completada. Registros: {total_reg}, Errores: {errores}"
        else:
            if not silent_mode: progress_bar.empty()
            return "No se encontraron datos RIPS válidos para exportar."
            
    except Exception as e:
        return f"Error general: {e}"

def worker_xlsx_evento_a_json_masivo(archivo_excel, carpeta_destino, silent_mode=False):
    if not archivo_excel or not carpeta_destino: return "Rutas inválidas."
    
    try:
        xls = pd.ExcelFile(archivo_excel)
        df_consultas = pd.DataFrame()
        df_procedimientos = pd.DataFrame()
        df_otros = pd.DataFrame()
        
        if "Consultas" in xls.sheet_names:
            df_consultas = pd.read_excel(xls, sheet_name="Consultas")
        if "Procedimientos" in xls.sheet_names:
            df_procedimientos = pd.read_excel(xls, sheet_name="Procedimientos")
        if "OtrosServicios" in xls.sheet_names:
            df_otros = pd.read_excel(xls, sheet_name="OtrosServicios")
            
        df_consultas = df_consultas.astype(object).where(pd.notnull(df_consultas), None)
        df_procedimientos = df_procedimientos.astype(object).where(pd.notnull(df_procedimientos), None)
        df_otros = df_otros.astype(object).where(pd.notnull(df_otros), None)

        archivos_unicos = set()
        if "archivo_origen" in df_consultas.columns:
            archivos_unicos.update(df_consultas["archivo_origen"].dropna().unique())
        if "archivo_origen" in df_procedimientos.columns:
            archivos_unicos.update(df_procedimientos["archivo_origen"].dropna().unique())
        if "archivo_origen" in df_otros.columns:
            archivos_unicos.update(df_otros["archivo_origen"].dropna().unique())
        
        if not archivos_unicos:
            return "No se encontró la columna 'archivo_origen' o está vacía."

        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Generando JSONs...")
        
        errores = 0
        generados = 0
        
        for i, nombre_archivo in enumerate(archivos_unicos):
            if not silent_mode:
                progress_bar.progress((i + 1) / len(archivos_unicos), text=f"Generando: {nombre_archivo}")
            
            try:
                usuarios_dict = {}
                
                def procesar_df(df_origen, clave_servicio):
                    if df_origen.empty or "archivo_origen" not in df_origen.columns: return
                    df_filtrado = df_origen[df_origen["archivo_origen"] == nombre_archivo]
                    
                    for _, row in df_filtrado.iterrows():
                        td = str(row.get("tipo_documento_usuario", ""))
                        doc = str(row.get("documento_usuario", ""))
                        user_key = (td, doc)
                        
                        if user_key not in usuarios_dict:
                            usuarios_dict[user_key] = {
                                "tipoDocumentoIdentificacion": row.get("tipo_documento_usuario"),
                                "numDocumentoIdentificacion": row.get("documento_usuario"),
                                "tipoUsuario": row.get("tipo_usuario"),
                                "fechaNacimiento": row.get("fecha_nacimiento"), 
                                "codSexo": row.get("sexo"),
                                "codPaisResidencia": row.get("pais_residencia"),
                                "codMunicipioResidencia": row.get("municipio_residencia"),
                                "codZonaTerritorialResidencia": row.get("zona_residencia"),
                                "incapacidad": row.get("incapacidad"),
                                "consecutivo": row.get("consecutivo_usuario"),
                                "codPaisOrigen": row.get("pais_origen"),
                                "servicios": { "consultas": [], "procedimientos": [], "otrosServicios": [] }
                            }
                        
                        servicio_data = row.to_dict()
                        keys_to_remove = [
                            "tipo_documento_usuario", "documento_usuario", "tipo_usuario", 
                            "fecha_nacimiento", "sexo", "pais_residencia", "municipio_residencia", 
                            "zona_residencia", "incapacidad", "consecutivo_usuario", "pais_origen",
                            "archivo_origen"
                        ]
                        for k in keys_to_remove: servicio_data.pop(k, None)
                        if any(v is not None for v in servicio_data.values()):
                            usuarios_dict[user_key]["servicios"][clave_servicio].append(servicio_data)

                procesar_df(df_consultas, "consultas")
                procesar_df(df_procedimientos, "procedimientos")
                procesar_df(df_otros, "otrosServicios")
                
                resultado_final = { "usuarios": list(usuarios_dict.values()) }
                ruta_salida = os.path.join(carpeta_destino, nombre_archivo)
                if not ruta_salida.lower().endswith(".json"): ruta_salida += ".json"

                with open(ruta_salida, 'w', encoding='utf-8') as f:
                    json.dump(resultado_final, f, ensure_ascii=False, indent=4)
                generados += 1
                
            except Exception:
                errores += 1
        
        if not silent_mode: progress_bar.empty()
        return f"Se generaron {generados} archivos JSON. Errores: {errores}"
        
    except Exception as e:
        return f"Error leyendo Excel: {e}"

def worker_autorizacion_docx_desde_excel(carpeta_origen, archivo_excel, sheet_name, col_carpeta, col_auth, use_filter=False, silent_mode=False):
    if not carpeta_origen or not archivo_excel: return "Rutas inválidas."
    
    is_native_mode = st.session_state.get('force_native_mode', True)
    
    try:
        if isinstance(archivo_excel, bytes): archivo_excel = io.BytesIO(archivo_excel)
        archivo_excel.seek(0)
        
        df = None
        if use_filter:
            import openpyxl
            wb = openpyxl.load_workbook(archivo_excel, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada."
            ws = wb[sheet_name]
            
            data = []
            headers = [cell.value for cell in ws[1]]
            
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    data.append([cell.value for cell in row])
            
            if data:
                df = pd.DataFrame(data, columns=headers)
            else:
                return "No hay datos visibles."
        else:
            df = pd.read_excel(archivo_excel, sheet_name=sheet_name)
        
        if col_carpeta not in df.columns or col_auth not in df.columns:
            return f"Columnas no encontradas: {col_carpeta}, {col_auth}"

        if is_native_mode:
            if not send_command:
                return "Error: Modo nativo activado pero cliente agente no disponible."
            
            username = st.session_state.get("username", "admin")
            tasks = []
            for index, fila in df.iterrows():
                nombre_carpeta = fila[col_carpeta]
                nueva_autorizacion = fila[col_auth]
                
                if pd.isna(nombre_carpeta) or pd.isna(nueva_autorizacion): continue
                
                nombre_carpeta = str(nombre_carpeta).strip()
                nueva_autorizacion = str(int(nueva_autorizacion)) if isinstance(nueva_autorizacion, (float, int)) else str(nueva_autorizacion).strip()
                
                if not nombre_carpeta or not nueva_autorizacion: continue
                
                tasks.append({
                    "rel_path": nombre_carpeta,
                    "file_pattern": r'CRC_.*_FEOV.*\.docx$',
                    "regex_replacements": [
                        (r'(AUTORIZACION:)\s*.*', r'\g<1> ' + nueva_autorizacion)
                    ]
                })
                
            if not tasks:
                return "No se encontraron tareas válidas para procesar."
                
            if not silent_mode:
                st.info(f"Enviando {len(tasks)} tareas al agente local...")
                
            task_id = send_command(username, "fill_docx", {
                "base_path": carpeta_origen,
                "tasks": tasks
            })
            
            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=300)
                
                if not silent_mode: status_placeholder.empty()
                
                if res and "error" not in res:
                    count = res.get("count", 0)
                    errors = res.get("errors", [])
                    msg = f"Finalizado correctamente (Modificados: {count})"
                    if errors:
                        msg += f". Hubo {len(errors)} errores."
                    return msg
                else:
                    return f"Error del agente: {res.get('error', 'Desconocido') if res else 'Desconocido'}"
            else:
                return "No se pudo crear la tarea en el servidor."

        # SERVER MODE
        modificados, errores_carpeta, errores_docx, errores_proceso = 0, 0, 0, 0
        docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)
        
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Modificando DOCX...")
            
        for index, fila in df.iterrows():
            if not silent_mode:
                progress_bar.progress((index + 1) / len(df), text=f"Procesando fila {index+1}")
                
            nombre_carpeta = fila[col_carpeta]
            nueva_autorizacion = fila[col_auth]
            
            if pd.isna(nombre_carpeta) or pd.isna(nueva_autorizacion): continue

            nombre_carpeta = str(nombre_carpeta).strip()
            nueva_autorizacion = str(int(nueva_autorizacion)) if isinstance(nueva_autorizacion, (float, int)) else str(nueva_autorizacion).strip()
            
            if not nombre_carpeta or not nueva_autorizacion: continue
            
            ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
            if not os.path.isdir(ruta_carpeta_especifica):
                errores_carpeta += 1
                continue
            
            ruta_docx_encontrada = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if docx_pattern.match(f)), None)

            if not ruta_docx_encontrada:
                errores_docx += 1
                continue

            try:
                doc = Document(ruta_docx_encontrada)
                fue_modificado = False
                for p in doc.paragraphs:
                    if "AUTORIZACION:" in p.text.upper():
                        p.text = re.sub(r'(AUTORIZACION:)\s*.*', r'\1 ' + str(nueva_autorizacion), p.text, flags=re.IGNORECASE)
                        fue_modificado = True
                        break
                
                if fue_modificado:
                    doc.save(ruta_docx_encontrada)
                    modificados += 1
                else:
                    errores_proceso += 1
            except Exception:
                errores_proceso += 1
        
        if not silent_mode: progress_bar.empty()
        return f"Modificados: {modificados}, Carpetas no encontradas: {errores_carpeta}, DOCX no encontrados: {errores_docx}, Errores proceso: {errores_proceso}"
    except Exception as e:
        return f"Error general: {e}"

def worker_regimen_docx_desde_excel(carpeta_origen, archivo_excel, sheet_name, col_carpeta, col_regimen, use_filter=False, silent_mode=False):
    if not carpeta_origen or not archivo_excel: return "Rutas inválidas."
    
    is_native_mode = st.session_state.get('force_native_mode', True)
    
    try:
        if isinstance(archivo_excel, bytes): archivo_excel = io.BytesIO(archivo_excel)
        archivo_excel.seek(0)
        
        df = None
        if use_filter:
            import openpyxl
            wb = openpyxl.load_workbook(archivo_excel, data_only=True)
            if sheet_name not in wb.sheetnames: return "Hoja no encontrada."
            ws = wb[sheet_name]
            
            data = []
            headers = [cell.value for cell in ws[1]]
            
            for row in ws.iter_rows(min_row=2):
                if not ws.row_dimensions[row[0].row].hidden:
                    data.append([cell.value for cell in row])
            
            if data:
                df = pd.DataFrame(data, columns=headers)
            else:
                return "No hay datos visibles."
        else:
            df = pd.read_excel(archivo_excel, sheet_name=sheet_name)
        
        if col_carpeta not in df.columns or col_regimen not in df.columns:
            return f"Columnas no encontradas: {col_carpeta}, {col_regimen}"

        if is_native_mode:
            if not send_command:
                return "Error: Modo nativo activado pero cliente agente no disponible."
            
            username = st.session_state.get("username", "admin")
            tasks = []
            for index, fila in df.iterrows():
                nombre_carpeta = fila[col_carpeta]
                nuevo_regimen = fila[col_regimen]
                
                if pd.isna(nombre_carpeta) or pd.isna(nuevo_regimen): continue
                
                nombre_carpeta = str(nombre_carpeta).strip()
                nuevo_regimen = str(nuevo_regimen).strip()
                
                if not nombre_carpeta or not nuevo_regimen: continue
                
                tasks.append({
                    "rel_path": nombre_carpeta,
                    "file_pattern": r'CRC_.*_FEOV.*\.docx$',
                    "regex_replacements": [
                        (r'(REGIMEN:)\s*.*', r'\g<1> ' + nuevo_regimen)
                    ]
                })
                
            if not tasks:
                return "No se encontraron tareas válidas para procesar."
                
            if not silent_mode:
                st.info(f"Enviando {len(tasks)} tareas al agente local...")
                
            task_id = send_command(username, "fill_docx", {
                "base_path": carpeta_origen,
                "tasks": tasks
            })
            
            if task_id:
                status_placeholder = st.empty()
                if not silent_mode: status_placeholder.text("Esperando agente...")
                
                res = wait_for_result(task_id, timeout=300)
                
                if not silent_mode: status_placeholder.empty()
                
                if res and "error" not in res:
                    count = res.get("count", 0)
                    errors = res.get("errors", [])
                    msg = f"Finalizado correctamente (Modificados: {count})"
                    if errors:
                        msg += f". Hubo {len(errors)} errores."
                    return msg
                else:
                    return f"Error del agente: {res.get('error', 'Desconocido') if res else 'Desconocido'}"
            else:
                return "No se pudo crear la tarea en el servidor."

        # SERVER MODE
        modificados, errores_carpeta, errores_docx, errores_proceso = 0, 0, 0, 0
        docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)
        
        progress_bar = None
        if not silent_mode:
            progress_bar = st.progress(0, text="Modificando Régimen...")
            
        for index, fila in df.iterrows():
            if not silent_mode:
                progress_bar.progress((index + 1) / len(df), text=f"Procesando fila {index+1}")
                
            nombre_carpeta = fila[col_carpeta]
            nuevo_regimen = fila[col_regimen]
            
            if pd.isna(nombre_carpeta) or pd.isna(nuevo_regimen): continue

            nombre_carpeta = str(nombre_carpeta).strip()
            nuevo_regimen = str(nuevo_regimen).strip()
            
            if not nombre_carpeta or not nuevo_regimen: continue
            
            ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
            if not os.path.isdir(ruta_carpeta_especifica):
                errores_carpeta += 1
                continue
            
            ruta_docx_encontrada = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if docx_pattern.match(f)), None)

            if not ruta_docx_encontrada:
                errores_docx += 1
                continue

            try:
                doc = Document(ruta_docx_encontrada)
                fue_modificado = False
                for p in doc.paragraphs:
                    if "REGIMEN:" in p.text.upper():
                        p.text = re.sub(r'(REGIMEN:)\s*.*', r'\1 ' + str(nuevo_regimen), p.text, flags=re.IGNORECASE)
                        fue_modificado = True
                        break
                
                if fue_modificado:
                    doc.save(ruta_docx_encontrada)
                    modificados += 1
                else:
                    errores_proceso += 1
            except Exception:
                errores_proceso += 1
        
        if not silent_mode: progress_bar.empty()
        return f"Modificados: {modificados}, Carpetas no encontradas: {errores_carpeta}, DOCX no encontrados: {errores_docx}, Errores proceso: {errores_proceso}"
    except Exception as e:
        return f"Error general: {e}"

# --- DIALOGS FOR NEW WORKERS ---

# --- WORKERS: FIRMA DIGITAL ---

def worker_crear_firma_nombre(root_path, ttf_path, size, humanize=False, silent_mode=False):
    try:
        from PIL import ImageDraw, ImageFont
        font = ImageFont.truetype(ttf_path, size)
    except Exception as e:
        msg = f"Error cargando fuente: {e}"
        if not silent_mode: st.error(msg)
        return msg

    count = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Generando firmas...")
    
    try:
        subfolders = [d for d in os.listdir(root_path) if os.path.isdir(os.path.join(root_path, d))]
    except Exception as e:
        return f"Error leyendo carpetas: {e}"
        
    total = len(subfolders)
    
    for i, sub in enumerate(subfolders):
        if not silent_mode and i % 10 == 0: progress_bar.progress(min(i/total, 1.0))
        
        text = sub # Nombre de carpeta es el texto
        
        try:
            # Dummy draw para calcular tamaño base
            dummy_img = Image.new('RGB', (1, 1))
            dummy_draw = ImageDraw.Draw(dummy_img)
            bbox = dummy_draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            # Crear imagen SOLO del texto primero
            img_text = Image.new('RGB', (text_width + 20, text_height + 20), (255, 255, 255))
            draw_text = ImageDraw.Draw(img_text)
            draw_text.text((10, 10), text, font=font, fill=(0, 0, 0))
            
            final_img = img_text
            
            if humanize:
                import random
                angle = random.uniform(-8, 8) # Rotación aleatoria
                final_img = img_text.rotate(angle, expand=True, fillcolor=(255, 255, 255))
                
            # Añadir padding final consistente
            fw, fh = final_img.size
            bg_w, bg_h = fw + 60, fh + 40
            bg = Image.new('RGB', (bg_w, bg_h), (255, 255, 255))
            
            # Centrar
            offset_x = (bg_w - fw) // 2
            offset_y = (bg_h - fh) // 2
            bg.paste(final_img, (offset_x, offset_y))
            
            target_dir = os.path.join(root_path, sub, "tipografia")
            if not os.path.exists(target_dir):
                os.makedirs(target_dir)
                
            bg.save(os.path.join(target_dir, "firma.jpg"))
            count += 1
            
        except Exception as e:
            if not silent_mode: st.warning(f"Error en firma carpeta {sub}: {e}")
            
    msg = f"Generadas {count} firmas."
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

def worker_crear_firma_excel(root_path, ttf_path, size, excel_file, sheet_name, col_folder, col_full_name, humanize=False, silent_mode=False):
    is_native = st.session_state.get("force_native_mode", True)
    
    try:
        from PIL import ImageDraw, ImageFont
        font = ImageFont.truetype(ttf_path, int(size))
    except Exception as e:
        msg = f"Error cargando fuente: {e}"
        if not silent_mode: st.error(msg)
        return msg

    try:
        if isinstance(excel_file, bytes):
            excel_file = io.BytesIO(excel_file)
        if hasattr(excel_file, 'seek'):
            excel_file.seek(0)
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
    except Exception as e:
        msg = f"Error leyendo Excel: {e}"
        if not silent_mode: st.error(msg)
        return msg

    count = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Generando firmas desde Excel...")
    total = len(df)
    
    # Store files for agent (Native Mode)
    agent_files = []
    
    for idx, row in df.iterrows():
        if not silent_mode and idx % 5 == 0: progress_bar.progress(min(idx/total, 1.0))
        
        folder_name = str(row[col_folder]).strip()
        if not folder_name or str(folder_name).lower() == 'nan': continue
        
        # Construir ruta objetivo
        target_dir = find_folder_path(root_path, folder_name)
        
        # Validar existencia solo si NO estamos en modo nativo (servidor local)
        if not is_native and not os.path.exists(target_dir):
            continue
            
        # Extraer nombre completo
        full_name = str(row[col_full_name]).strip()
        if not full_name or full_name.lower() == 'nan': full_name = ""
        
        # Lógica inteligente para Primer Nombre + Primer Apellido
        parts = full_name.split()
        name_part = ""
        surname_part = ""
        
        if len(parts) >= 1:
            name_part = parts[0].capitalize() 
        
        if len(parts) >= 4:
            surname_part = parts[2].capitalize()
        elif len(parts) == 3:
            surname_part = parts[1].capitalize()
        elif len(parts) == 2:
            surname_part = parts[1].capitalize()
            
        # Construir texto final
        final_text = f"{name_part} {surname_part}".strip()
        if not final_text: 
            final_text = folder_name 
            
        # Generar Imagen
        try:
            # Dummy draw 
            dummy_img = Image.new('RGB', (1, 1))
            dummy_draw = ImageDraw.Draw(dummy_img)
            bbox = dummy_draw.textbbox((0, 0), final_text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            # Crear imagen base texto
            img_text = Image.new('RGB', (text_width + 20, text_height + 20), (255, 255, 255))
            draw_text = ImageDraw.Draw(img_text)
            draw_text.text((10, 10), final_text, font=font, fill=(0, 0, 0))
            
            final_img = img_text
            
            if humanize:
                import random
                angle = random.uniform(-8, 8)
                final_img = img_text.rotate(angle, expand=True, fillcolor=(255, 255, 255))
            
            # Composition final
            fw, fh = final_img.size
            bg_w, bg_h = fw + 60, fh + 40
            bg = Image.new('RGB', (bg_w, bg_h), (255, 255, 255))
            
            offset_x = (bg_w - fw) // 2
            offset_y = (bg_h - fh) // 2
            bg.paste(final_img, (offset_x, offset_y))
            
            # Guardar
            if is_native:
                # Convertir a base64 y agregar a lista para agente
                buf = io.BytesIO()
                bg.save(buf, format="JPEG")
                img_b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
                
                # Rutas destino (tipografía y raíz)
                path_tipografia = os.path.join(target_dir, "tipografia", "firma.jpg")
                path_raiz = os.path.join(target_dir, "firma.jpg")
                
                agent_files.append({"path": path_tipografia, "content_b64": img_b64})
                agent_files.append({"path": path_raiz, "content_b64": img_b64})
                count += 1
            else:
                # Guardar localmente (Servidor)
                tipografia_dir = os.path.join(target_dir, "tipografia")
                if not os.path.exists(tipografia_dir):
                    os.makedirs(tipografia_dir)
                    
                bg.save(os.path.join(tipografia_dir, "firma.jpg"))
                
                # Guardar también en la carpeta raíz (Requerimiento Usuario)
                bg.save(os.path.join(target_dir, "firma.jpg"))
                count += 1

        except Exception as e:
            if not silent_mode: st.warning(f"Error generando firma para {folder_name}: {e}")

    # Enviar al agente si es necesario
    if is_native and agent_files:
        import agent_client
        if not agent_client:
             return "Error: Módulo agent_client no disponible."
        
        username = st.session_state.get("username", "default")
        print(f"DEBUG: Enviando {len(agent_files)} archivos de firma al agente.")
        
        task_id = agent_client.send_command(username, "write_files", {
            "files": agent_files
        })
        
        if task_id:
            msg = f"Tarea de creación de firmas enviada al agente (ID: {task_id}). Archivos a crear: {len(agent_files)}."
        else:
            msg = "Error enviando tarea al agente."
    else:
        msg = f"Generadas {count} firmas desde Excel."

    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
    return msg

# --- DIALOGS FOR NEW WORKERS ---

@st.dialog("Crear Firma Digital desde Nombre")
def dialog_crear_firma():
    st.write("Genera una imagen JPG con firma manuscrita.")
    
    # Resolver rutas de assets
    base_dir = os.path.dirname(os.path.abspath(__file__))
    # Ajustar ruta para subir dos niveles desde src/tabs hasta la raíz, y luego a assets
    assets_fonts = os.path.join(base_dir, "..", "..", "assets", "fonts")
    
    # Selección de Fuente (común para ambos modos)
    option = st.radio("Fuente:", ["Subir fuente", "Pacifico (Predeterminada)", "MyUglyHandwriting"], index=1, horizontal=True)
    
    font_path = None
    if option == "Subir fuente":
        uploaded_font = st.file_uploader("Fuente TTF:", type=["ttf", "otf"], key="up_font_firma")
        if uploaded_font:
            with open("temp_font.ttf", "wb") as f:
                f.write(uploaded_font.getbuffer())
            font_path = "temp_font.ttf"
    elif option == "Pacifico (Predeterminada)":
        font_path = os.path.join(assets_fonts, "Pacifico.ttf")
    elif option == "MyUglyHandwriting":
        font_path = os.path.join(assets_fonts, "MyUglyHandwriting-Regular.otf")
        
    c1_opt, c2_opt = st.columns(2)
    with c1_opt:
        size = st.number_input("Tamaño Fuente:", value=70)
    with c2_opt:
        humanize = st.checkbox("🎨 Estilo Natural", value=True, help="Aplica rotación aleatoria e imperfecciones.")

    # Folder Selector
    st.write("Carpeta Base de Trabajo:")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    target_path = render_path_selector(
        key="crear_firma_path",
        label="Ruta",
        default_path=default_path
    )
    current_path = target_path

    # Tabs para los modos
    tab1, tab2 = st.tabs(["📁 Usar Nombre Carpeta", "📊 Usar Excel"])
    
    with tab1:
        st.write("Usa el nombre de la subcarpeta como texto de la firma.")
        st.info(f"Ruta actual: {current_path}")
        if st.button("🚀 Generar (Carpeta)"):
            if font_path and os.path.exists(font_path):
                try:
                    with st.spinner("Generando firmas..."):
                        result = worker_crear_firma_nombre(current_path, font_path, size, humanize)
                        st.success(result)
                        close_auto_dialog()
                        # render_download_button(current_path, "dl_firma_folder", "📦 Descargar Firmas (ZIP)")
                except Exception as e:
                    st.error(f"Error: {e}")
            else:
                st.error(f"No se encontró la fuente en: {font_path}")

    with tab2:
        st.write("Usa nombres extraídos de una COLUMNA ÚNICA (detecta 1er Nombre + 1er Apellido).")
        uploaded = st.file_uploader("Excel:", type=["xlsx", "xls"], key="excel_firma")
        
        if uploaded:
            try:
                if hasattr(uploaded, 'seek'): uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                sheet_names = _get_excel_sheet_names(file_bytes)
                sheet = st.selectbox("Hoja:", sheet_names, key="sheet_firma")
                df_prev = _get_excel_preview(file_bytes, sheet, nrows=1)
                cols = df_prev.columns.tolist()
                
                c1, c2 = st.columns(2)
                with c1: col_folder = st.selectbox("Col. Carpeta (Match):", cols, key="col_match_firma")
                with c2: col_full_name = st.selectbox("Col. Nombre Completo:", cols, key="col_full_name_firma")
                
                st.info(f"Ruta actual: {current_path}")
                
                if st.button("🚀 Generar (Excel)"):
                     uploaded.seek(0)
                     file_bytes = uploaded.getvalue()
                     if font_path and os.path.exists(font_path):
                        try:
                            with st.spinner("Generando firmas..."):
                                result = worker_crear_firma_excel(current_path, font_path, size, file_bytes, sheet, col_folder, col_full_name, humanize)
                                st.success(result)
                                close_auto_dialog()
                                # render_download_button(current_path, "dl_firma_excel", "📦 Descargar Resultados (ZIP)")
                            # st.rerun()
                        except Exception as e:
                            st.error(f"Error: {e}")
                     else:
                        st.error(f"No se encontró la fuente en: {font_path}")
            except Exception as e:
                st.error(f"Error leyendo Excel: {e}")

    if st.button("Cerrar", key="btn_close_crear_firma"):
        close_auto_dialog()

@st.dialog("Organización FEOV Avanzada")
def dialog_organizar_feov_avanzado():
    st.write("Mueve archivos de 'Origen' a subcarpetas en 'Destino' basándose en el número de factura FEOV del PDF destino.")
    
    st.write("1. Carpeta Destino (contiene subcarpetas con PDFs FEOV...)")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    
    path_dest = render_path_selector(
        key="feov_adv_dest",
        label="Ruta Destino",
        default_path=default_path
    )
    
    st.write("2. Carpeta Origen (archivos a mover)")
    
    path_orig = render_path_selector(
        key="feov_adv_orig",
        label="Ruta Origen",
        default_path=default_path
    )
    
    if st.button("Iniciar Organización Avanzada"):
        if path_dest and path_orig:
            try:
                with st.spinner("Organizando facturas..."):
                    result = worker_organizar_facturas_por_pdf_avanzado(path_dest, path_orig)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(path_dest, "dl_feov_adv", "📦 Descargar Destino (ZIP)")
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.error("Seleccione ambas carpetas.")

    if st.button("Cerrar", key="btn_close_feov_adv"):
        close_auto_dialog()

@st.dialog("Autorización DOCX desde Excel")
def dialog_autorizacion_docx():
    st.write("Modifica el campo AUTORIZACION en DOCX masivamente.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")
        
    uploaded = st.file_uploader("Excel", type=["xlsx"], key="auth_up")
    sheet = None
    col_folder = None
    col_auth = None
    use_filter = False
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Hoja", sheet_names, key="auth_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                c1, c2 = st.columns(2)
                col_folder = c1.selectbox("Columna Carpeta", df_preview.columns, key="auth_col_folder")
                col_auth = c2.selectbox("Columna Autorización", df_preview.columns, key="auth_col_val")
                use_filter = st.checkbox("Usar filtros de Excel (solo filas visibles)", value=False, key="auth_filter")
        except Exception as e:
            st.error(f"Error: {e}")

    default_path = st.session_state.get("current_path", os.getcwd())
    base_path = render_path_selector(
        key="auth_base",
        label="Carpeta Base",
        default_path=default_path
    )
    
    if st.button("Iniciar Modificación"):
        if uploaded and base_path and col_folder and col_auth:
            try:
                uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                with st.spinner("Modificando DOCX..."):
                    result = worker_autorizacion_docx_desde_excel(base_path, file_bytes, sheet, col_folder, col_auth, use_filter)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(base_path, "dl_auth_docx", "📦 Descargar DOCX Modificados (ZIP)")
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_auth_docx"):
        close_auto_dialog()

@st.dialog("Régimen DOCX desde Excel")
def dialog_regimen_docx():
    st.write("Modifica el campo REGIMEN en DOCX masivamente.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")
        
    uploaded = st.file_uploader("Excel", type=["xlsx"], key="reg_up")
    sheet = None
    col_folder = None
    col_reg = None
    use_filter = False
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Hoja", sheet_names, key="reg_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                c1, c2 = st.columns(2)
                col_folder = c1.selectbox("Columna Carpeta", df_preview.columns, key="reg_col_folder")
                col_reg = c2.selectbox("Columna Régimen", df_preview.columns, key="reg_col_val")
                use_filter = st.checkbox("Usar filtros de Excel (solo filas visibles)", value=False, key="reg_filter")
        except Exception as e:
            st.error(f"Error: {e}")

    default_path = st.session_state.get("current_path", os.getcwd())
    base_path = render_path_selector(
        key="reg_base",
        label="Carpeta Base",
        default_path=default_path
    )
    
    if st.button("Iniciar Modificación"):
        if uploaded and base_path and col_folder and col_reg:
            try:
                uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                with st.spinner("Modificando Régimen..."):
                    result = worker_regimen_docx_desde_excel(base_path, file_bytes, sheet, col_folder, col_reg, use_filter)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(base_path, "dl_reg_docx", "📦 Descargar DOCX Modificados (ZIP)")
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_reg_docx"):
        close_auto_dialog()

def worker_distribuir_base_archivo(file_source, is_upload_bytes, excel_bytes, sheet_name, col_folder, base_path):
    try:
        # Check Native Mode
        is_native = st.session_state.get("force_native_mode", True)
        
        # 1. Read Excel
        if isinstance(excel_bytes, bytes):
            df = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=sheet_name)
        else:
            df = pd.read_excel(excel_bytes, sheet_name=sheet_name)
        
        # Determine source file name and content
        src_filename = "Archivo_Distribuido.dat"
        src_content = None
        
        if is_upload_bytes:
            # file_source is UploadedFile object or bytes
            if hasattr(file_source, "name"): src_filename = file_source.name
            
            if hasattr(file_source, "getvalue"): src_content = file_source.getvalue()
            else: src_content = file_source # assume bytes
        else:
            # file_source is path string (server side)
            if os.path.exists(file_source):
                 src_filename = os.path.basename(file_source)
                 with open(file_source, "rb") as f:
                     src_content = f.read()
            else:
                 return f"Error: Archivo origen no existe en servidor: {file_source}"

        # Collect destinations
        dest_paths = []
        preview_dest = ""
        
        for index, row in df.iterrows():
            folder_val = row[col_folder]
            if pd.isna(folder_val): continue
            folder_name = str(folder_val).strip()
            
            # Sanitize folder name
            folder_name = "".join([c for c in folder_name if c.isalnum() or c in (' ', '-', '_', '.')]).strip()
            
            if not folder_name: continue

            # Use helper to find folder path (supports search results)
            # In Native Mode on AWS, this returns the constructed path string (G:/...) even if not exists locally
            dest_dir = find_folder_path(base_path, folder_name)
            
            # Clean up potential double separators if base_path ends with /
            # os.path.join handles it usually, but let's be safe
            dest_file = os.path.join(dest_dir, src_filename)
            dest_paths.append(dest_file)
            
            if not preview_dest: preview_dest = dest_file

        if not dest_paths:
            return "No se encontraron carpetas destino válidas en el Excel."

        # Execute Distribution
        if is_native:
            # Agent Mode
            import agent_client
            if not agent_client:
                 return "Error: Módulo agent_client no disponible."
            
            try:
                username = st.session_state.get("username", "default")
                
                # Encode content
                content_b64 = base64.b64encode(src_content).decode('utf-8')
                
                print(f"DEBUG: Enviando {len(dest_paths)} archivos al agente. Primer destino: {preview_dest}")
                
                task_id = agent_client.send_command(username, "distribute_files", {
                    "content_b64": content_b64,
                    "paths": dest_paths
                })
                
                if task_id:
                    return f"Tarea enviada al agente (ID: {task_id}). Archivos a copiar: {len(dest_paths)}. (Ejemplo destino: {preview_dest})"
                else:
                    return "Error: No se pudo enviar la tarea al agente."
            except Exception as e:
                return f"Error enviando al agente: {e}"
        else:
            # Server Local Mode
            count = 0
            errors = 0
            
            for dest_file in dest_paths:
                try:
                    dest_dir = os.path.dirname(dest_file)
                    os.makedirs(dest_dir, exist_ok=True)
                    
                    with open(dest_file, "wb") as f:
                        f.write(src_content)
                    count += 1
                except Exception as e:
                    errors += 1
                    print(f"Error copying to {dest_file}: {e}")
                    
            msg = f"Proceso completado (Servidor). Archivos copiados: {count}. Errores: {errors}."
            if count > 0:
                 msg += f" (Ejemplo: {preview_dest})"
            return msg

    except Exception as e:
        return f"Error crítico: {e}"

@st.dialog("Distribuir Archivo (Base Excel)")
def dialog_distribuir_base():
    st.write("Copia un archivo a múltiples carpetas definidas en un Excel.")
    
    # 1. Excel Base
    uploaded_excel = st.file_uploader("Cargar Base (Excel)", type=["xlsx", "xls"], key="dist_base_excel")
    
    sheet = None
    col_folder = None
    
    if uploaded_excel:
        try:
            file_bytes = uploaded_excel.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Hoja", sheet_names, key="dist_base_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                col_folder = st.selectbox("Columna Nombre Carpeta", df_preview.columns, key="dist_base_col")
        except Exception as e:
            st.error(f"Error Excel: {e}")

    # 2. Archivo a Distribuir
    st.divider()
    st.write("Archivo a Distribuir:")
    
    # Siempre usar file_uploader para permitir cargar cualquier archivo desde el PC
    file_to_distribute = st.file_uploader("Cargar Archivo a Distribuir", key="dist_base_file_up")
    is_upload_bytes = True if file_to_distribute else False

    # 3. Carpeta Destino Base
    st.divider()
    default_path = st.session_state.get("current_path", os.getcwd())
    base_dest_path = render_path_selector(
        key="dist_base_dest",
        label="Carpeta Destino (Raíz)",
        help_text="Donde se encuentran o crearán las carpetas del Excel.",
        default_path=default_path
    )
    
    # 4. Action
    if st.button("🚀 Distribuir"):
        if uploaded_excel and sheet and col_folder and file_to_distribute and base_dest_path:
            try:
                if hasattr(uploaded_excel, 'seek'):
                    uploaded_excel.seek(0)
                excel_bytes = uploaded_excel.getvalue()
                
                with st.spinner("Distribuyendo archivo..."):
                    result = worker_distribuir_base_archivo(
                        file_to_distribute, 
                        is_upload_bytes, 
                        excel_bytes, 
                        sheet, 
                        col_folder, 
                        base_dest_path
                    )
                    st.success(result)
                    close_auto_dialog()
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.warning("Complete todos los campos.")

    if st.button("Cerrar", key="btn_close_dist_base"):
        close_auto_dialog()

@st.dialog("Crear Carpetas desde Excel")
def dialog_crear_carpetas_excel():
    st.write("Crea estructura de carpetas basada en una columna de Excel.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")

    uploaded = st.file_uploader("Excel", type=["xlsx", "xls"], key="create_fold_up")
    
    sheet = None
    col_name = None
    use_filter = False
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Hoja", sheet_names, key="create_fold_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                col_name = st.selectbox("Nombre Columna Carpetas", df_preview.columns, key="create_fold_col")
                use_filter = st.checkbox("Usar filtros de Excel (solo filas visibles)", value=False, key="create_fold_filter")
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")
            
    default_path = st.session_state.get("current_path", os.getcwd())
    base_path = render_path_selector(
        key="create_fold_base",
        label="Carpeta Base",
        default_path=default_path
    )
    
    if st.button("Crear Carpetas"):
        if uploaded and base_path and col_name:
            try:
                if hasattr(uploaded, 'seek'):
                    uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                with st.spinner("Creando carpetas..."):
                    result = worker_crear_carpetas_excel_avanzado(file_bytes, sheet, col_name, base_path, use_filter)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(base_path, "dl_create_fold_excel", "📦 Descargar Estructura (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_cr_fold"):
        close_auto_dialog()

def worker_copiar_archivo_a_subcarpetas(archivo_a_copiar, carpeta_destino_base, silent_mode=False):
    # Agent integration
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and _should_delegate(carpeta_destino_base) and agent_client:
        try:
            username = st.session_state.get("username", "default")
            task_id = agent_client.send_command(username, "copiar_archivo_a_subcarpetas", {
                "archivo": archivo_a_copiar,
                "carpeta_base": carpeta_destino_base
            })
            if task_id:
                if not silent_mode:
                    with st.spinner("Copiando archivo vía Agente..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    res = agent_client.wait_for_result(task_id, timeout=600)
                if res and isinstance(res, dict) and "message" in res:
                    return res["message"]
                elif res and isinstance(res, dict) and "error" in res:
                    return f"Error: {res['error']}"
        except Exception as e:
            pass

    try:
        subcarpetas = [os.path.join(carpeta_destino_base, d) for d in os.listdir(carpeta_destino_base) if os.path.isdir(os.path.join(carpeta_destino_base, d))]
    except Exception as e:
        return f"Error al leer la carpeta destino: {e}"

    if not subcarpetas:
        return "No se encontraron subcarpetas en la ruta destino."

    copiados = 0
    conflictos = 0
    errores = 0
    nombre_archivo = os.path.basename(archivo_a_copiar)

    if not silent_mode:
        progress_bar = st.progress(0, text="Iniciando copia...")
    
    total = len(subcarpetas)
    for i, subcarpeta in enumerate(subcarpetas):
        if not silent_mode:
            progress_bar.progress(min((i + 1) / total, 1.0), text=f"Copiando a {os.path.basename(subcarpeta)}...")
        
        destino_final = os.path.join(subcarpeta, nombre_archivo)
        if os.path.exists(destino_final):
            conflictos += 1
            continue
        try:
            shutil.copy2(archivo_a_copiar, destino_final)
            copiados += 1
        except Exception as e:
            errores += 1

    if not silent_mode:
        progress_bar.progress(1.0, text="Copia completada.")
        
    return f"Copia completada. Copiados: {copiados}. Conflictos: {conflictos}. Errores: {errores}."

@st.dialog("Copiar Archivo a Subcarpetas")
def dialog_copiar_archivo_a_subcarpetas():
    st.write("Copia un único archivo a todas las subcarpetas de una carpeta destino.")
    
    # Modo Validación
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: Las rutas deben estar accesibles por el servidor.")

    default_path = st.session_state.get("current_path", os.getcwd())
    
    file_path = render_path_selector(
        key="copy_sub_file",
        label="Archivo a Copiar",
        default_path=default_path
    )
    
    target_path = render_path_selector(
        key="copy_sub_target",
        label="Carpeta Destino (que contiene subcarpetas)",
        default_path=default_path
    )
    
    if st.button("🚀 Copiar", key="btn_init_copy_sub"):
        if file_path and target_path:
            if not os.path.isfile(file_path):
                st.warning("El origen debe ser un archivo.")
            elif not os.path.isdir(target_path):
                st.warning("El destino debe ser una carpeta.")
            else:
                with st.spinner("Copiando archivo..."):
                    result = worker_copiar_archivo_a_subcarpetas(file_path, target_path)
                    st.success(result)
                    close_auto_dialog()

    if st.button("❌ Cerrar", key="btn_close_copy_sub"):
        close_auto_dialog()

@st.dialog("Copiar Mapeo Subcarpetas")
def dialog_copiar_mapeo():
    st.write("Copia archivos entre carpetas basándose en un mapeo Excel.")
    uploaded = st.file_uploader("Excel", type=["xlsx", "xls"], key="copy_map_up")
    
    sheet = None
    col_src = None
    col_dst = None
    use_filter = False
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Hoja", sheet_names, key="copy_map_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                c1, c2 = st.columns(2)
                col_src = c1.selectbox("Columna Origen", df_preview.columns, key="copy_map_src")
                col_dst = c2.selectbox("Columna Destino", df_preview.columns, key="copy_map_dst")
                use_filter = st.checkbox("Usar filtros de Excel (solo filas visibles)", value=False, key="copy_map_filter")
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")
    
    st.write("Rutas Base:")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    
    src_base = render_path_selector(
        key="copy_map_src_base",
        label="Carpeta Origen Base",
        default_path=default_path
    )

    dst_base = render_path_selector(
        key="copy_map_dst_base",
        label="Carpeta Destino Base",
        default_path=default_path
    )
    
    if st.button("Iniciar Copia", key="btn_init_copy_map_sub"):
        if uploaded and src_base and dst_base and col_src and col_dst:
            try:
                uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                with st.spinner("Copiando archivos..."):
                    result = worker_copiar_mapeo_subcarpetas(file_bytes, sheet, col_src, col_dst, src_base, dst_base, use_filter)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(dst_base, "dl_copy_map_sub", "📦 Descargar Destino (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_cop_map"):
        close_auto_dialog()

@st.dialog("Copiar desde Raíz (Mapeo)")
def dialog_copiar_raiz():
    st.write("Copia archivos desde una raíz única a carpetas destino según Excel.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")

    uploaded = st.file_uploader("Excel", type=["xlsx", "xls"], key="copy_root_up")
    
    sheet = None
    col_id = None
    col_folder = None
    use_filter = False
    
    if uploaded:
        try:
            file_bytes = uploaded.getvalue()
            sheet_names = _get_excel_sheet_names(file_bytes)
            sheet = st.selectbox("Hoja", sheet_names, key="copy_root_sheet")
            if sheet:
                df_preview = _get_excel_preview(file_bytes, sheet, nrows=5)
                c1, c2 = st.columns(2)
                col_id = c1.selectbox("Columna ID/Nombre Archivo", df_preview.columns, key="copy_root_id")
                col_folder = c2.selectbox("Columna Carpeta Destino", df_preview.columns, key="copy_root_folder")
                use_filter = st.checkbox("Usar filtros de Excel (solo filas visibles)", value=False, key="copy_root_filter")
        except Exception as e:
            st.error(f"Error leyendo Excel: {e}")

    st.write("Rutas:")
    
    default_path = st.session_state.get("current_path", os.getcwd())
    
    root_src = render_path_selector(
        key="copy_root_src_base",
        label="Carpeta Origen (Archivos)",
        default_path=default_path
    )

    root_dst = render_path_selector(
        key="copy_root_dst_base",
        label="Carpeta Destino (Carpetas)",
        default_path=default_path
    )
    
    if st.button("Iniciar Copia", key="btn_init_copy_root_map"):
        if uploaded and root_src and root_dst and col_id and col_folder:
            try:
                uploaded.seek(0)
                file_bytes = uploaded.getvalue()
                with st.spinner("Copiando archivos..."):
                    result = worker_copiar_archivos_desde_raiz_mapeo(file_bytes, sheet, col_id, col_folder, root_src, root_dst, use_filter)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(root_dst, "dl_copy_root_map", "📦 Descargar Destino (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar", key="btn_close_cop_raiz"):
        close_auto_dialog()

@st.dialog("RIPS Eventos Masivos")
def dialog_rips_masivos():
    st.write("Conversión masiva entre JSON (Eventos) y Excel.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")

    mode = st.radio("Modo", ["JSON -> Excel", "Excel -> JSON"])
    
    if mode == "JSON -> Excel":
        default_path = st.session_state.get("current_path", os.getcwd())
        folder_src = render_path_selector(
            key="rips_json_src",
            label="Carpeta JSONs",
            default_path=default_path
        )

        file_dst = st.text_input("Nombre Archivo Salida (.xlsx)", "Consolidado.xlsx")
        if st.button("Convertir JSONs a Excel"):
            if folder_src:
                try:
                    with st.spinner("Convirtiendo JSONs..."):
                        excel_path = os.path.join(folder_src, file_dst)
                        result = worker_json_evento_a_xlsx_masivo(folder_src, excel_path)
                        st.success(result)
#                         render_download_button(excel_path, f"dl_rips_json_excel_{int(time.time())}", f"📥 Descargar {file_dst}")
                        # if os.path.exists(excel_path):
                        #     with open(excel_path, "rb") as f:
                        #         st.download_button("📥 Descargar Excel", f, file_name=file_dst)
                        # time.sleep(2)
                        # st.rerun()
                except Exception as e:
                    st.error(f"Error: {e}")
    else:
        file_src = st.file_uploader("Excel Eventos", type=["xlsx"], key="up_excel_eventos")
        
        default_path = st.session_state.get("current_path", os.getcwd())
        folder_dst = render_path_selector(
            key="rips_excel_dst",
            label="Carpeta Destino JSONs",
            default_path=default_path
        )

        if st.button("Convertir Excel a JSONs"):
            if file_src and folder_dst:
                try:
                    # Need to save temp excel first? worker takes path
                    t_path = os.path.join(folder_dst, "temp_eventos.xlsx")
                    with open(t_path, "wb") as f: f.write(file_src.getbuffer())
                    
                    with st.spinner("Generando JSONs..."):
                        result = worker_xlsx_evento_a_json_masivo(t_path, folder_dst)
                        st.success(result)
#                         render_download_button(folder_dst, "dl_rips_excel_json", "📦 Descargar JSONs Generados (ZIP)")
                        # time.sleep(2)
                        # st.rerun()
                except Exception as e:
                    st.error(f"Error: {e}")

# --- DIALOGS: RENAMING ---

@st.dialog("Exportar para Renombrar")
def dialog_exportar_renombrado():
    st.write("Genera un Excel con los archivos de una carpeta para renombrarlos masivamente.")
    
    c1, c2 = st.columns([0.8, 0.2])
    
    default_path = st.session_state.get("current_path", os.getcwd())
    folder = render_path_selector(
        key="renombrar_export_src",
        label="Carpeta a Analizar",
        default_path=default_path
    )
    
    dest_path = render_path_selector(
        key="renombrar_export_dst",
        label="Carpeta Destino (donde se guardará el Excel)",
        default_path=default_path
    )
    
    if st.button("Generar Excel"):
        if folder and dest_path:
            try:
                is_native_mode = st.session_state.get('force_native_mode', True)
                
                data = []
                
                if is_native_mode:
                    if not send_command:
                        st.error("Error: Modo nativo activado pero cliente agente no disponible.")
                        return
                    
                    username = st.session_state.get("username", "admin")
                    with st.spinner("Buscando archivos con el agente local..."):
                        task_id = send_command(username, "search_files", {
                            "path": folder,
                            "patterns": ["*"],
                            "recursive": False,
                            "item_type": "file"
                        })
                        if task_id:
                            res = wait_for_result(task_id, timeout=60)
                            if res and "status" in res and res["status"] == "success":
                                items = res.get("result", {}).get("items", [])
                                for item in items:
                                    if item.get("type") == "file":
                                        fname = item.get("name")
                                        data.append({"Nombre Actual": fname, "Nombre Nuevo": fname})
                            else:
                                st.error(f"Error del agente: {res.get('error', 'Desconocido')}")
                                return
                        else:
                            st.error("No se pudo crear la tarea de búsqueda.")
                            return
                else:
                    for f in os.listdir(folder):
                        if os.path.isfile(os.path.join(folder, f)):
                            data.append({"Nombre Actual": f, "Nombre Nuevo": f})
                
                if data:
                    df = pd.DataFrame(data)
                    out_path = os.path.join(dest_path, "Renombrar_Archivos.xlsx")
                    
                    if is_native_mode:
                        import io
                        import base64
                        mem_excel = io.BytesIO()
                        df.to_excel(mem_excel, index=False)
                        b64_content = base64.b64encode(mem_excel.getvalue()).decode('utf-8')
                        
                        with st.spinner("Guardando Excel mediante el agente..."):
                            write_task = send_command(username, "write_files", {
                                "files": [{
                                    "path": out_path,
                                    "content": b64_content,
                                    "encoding": "base64"
                                }]
                            })
                            if write_task:
                                write_res = wait_for_result(write_task)
                                if write_res and "error" not in write_res and not write_res.get("errors"):
                                    st.success(f"Excel generado exitosamente en: {out_path}")
                                    close_auto_dialog()
                                else:
                                    st.error("Error guardando el archivo.")
                    else:
                        df.to_excel(out_path, index=False)
                        st.success(f"Excel generado en: {out_path}")
                        close_auto_dialog()
                else:
                    st.warning("No se encontraron archivos en la carpeta.")
            except Exception as e:
                st.error(f"Error: {e}")

    if st.button("Cerrar"):
        close_auto_dialog()

@st.dialog("Aplicar Renombrado (Excel)")
def dialog_aplicar_renombrado():
    st.write("Renombra archivos basándose en un Excel (NombreActual -> NuevoNombre).")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")
        
    excel_file = st.file_uploader("Archivo Excel", type=["xlsx"], key="up_excel_unificar")
    
    c1, c2 = st.columns([0.8, 0.2])
    
    default_path = st.session_state.get("current_path", os.getcwd())
    folder = render_path_selector(
        key="renombrar_apply_src",
        label="Carpeta",
        default_path=default_path
    )
    
    if st.button("Aplicar Cambios"):
        if excel_file and folder:
            try:
                # En lugar de guardar temporalmente en 'folder' (que puede ser ruta cliente en modo nativo),
                # leemos el Excel directamente desde memoria.
                with st.spinner("Renombrando archivos..."):
                    result = worker_aplicar_renombrado_excel(excel_file, folder)
                    st.success(result)
                    close_auto_dialog()
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.error("Seleccione archivo Excel y carpeta.")

    if st.button("Cerrar", key="btn_close_app_ren"):
        close_auto_dialog()

@st.dialog("Copiar Archivo a Subcarpetas")
def dialog_copiar_archivo_a_subcarpetas():
    st.write("Copia un archivo seleccionado a todas las subcarpetas del destino.")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible.")

    file_to_copy = st.file_uploader("Archivo a Copiar", key="copy_sub_file")
    
    c1, c2 = st.columns([0.8, 0.2])
    
    default_path = st.session_state.get("current_path", os.getcwd())
    dest_base_path = render_path_selector(
        key="copy_sub_dest",
        label="Carpeta Destino Base",
        default_path=default_path
    )
    
    if st.button("Iniciar Copia a Subcarpetas"):
        if file_to_copy and dest_base_path:
            try:
                # Save temp file
                t_path = os.path.join(dest_base_path, file_to_copy.name)
                with open(t_path, "wb") as f:
                    f.write(file_to_copy.getbuffer())
                
                with st.spinner("Copiando archivos..."):
                    result = worker_copiar_archivo_a_subcarpetas(t_path, dest_base_path)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(dest_base_path, "dl_copy_sub", "📦 Descargar Destino (ZIP)")
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.error("Seleccione archivo y carpeta destino.")

    if st.button("Cerrar", key="btn_close_cop_sub"):
        close_auto_dialog()

@st.dialog("Organizar Facturas (FEOV)")
def dialog_organizar_feov():
    st.write("Organiza facturas PDF moviéndolas a subcarpetas según su número FEOV.")
    st.info("1. Selecciona la carpeta DESTINO (donde están las carpetas numeradas).\n2. Selecciona la carpeta ORIGEN (donde están los archivos desordenados).")
    
    # Validación de Modo
    if not st.session_state.get("force_native_mode", True):
        st.warning("⚠️ Modo Web: La selección de carpetas nativa no está disponible. Use las rutas manuales.")

    default_path = st.session_state.get("current_path", os.getcwd())
    
    target_path = render_path_selector(
        key="feov_target",
        label="Carpeta DESTINO",
        default_path=default_path
    )

    source_path = render_path_selector(
        key="feov_source",
        label="Carpeta ORIGEN",
        default_path=default_path
    )
    
    if st.button("Organizar Facturas"):
        if target_path and source_path:
            try:
                with st.spinner("Organizando facturas..."):
                    result = worker_organizar_facturas_feov(source_path, target_path)
                    st.success(result)
                    close_auto_dialog()
#                     render_download_button(target_path, "dl_feov", "📦 Descargar Facturas Organizadas (ZIP)")
                    # time.sleep(2)
                    # st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")
        else:
            st.warning("Seleccione ambas carpetas.")

    if st.button("Cerrar", key="btn_close_org_feov"):
        close_auto_dialog()

# --- RENDER MAIN TAB ---

def render():
    st.markdown("## ⚙️ Acciones Automatizadas")
    
    # Obtener ruta global por defecto (la buscada al inicio)
    default_path = st.session_state.get("current_path", os.path.expanduser("~"))

    # Crear pestañas principales
    tab_unif, tab_org, tab_modif, tab_an, tab_create = st.tabs([
        "Unificación y División", 
        "Organización", 
        "Modificación y Renombrado", 
        "Análisis", 
        "Creación y Otros"
    ])

    # --- TAB 1: Unificación y División ---
    with tab_unif:
        st.caption("Operaciones de unión y división de archivos PDF, imágenes y DOCX.")
        
        col_u1, col_u2 = st.columns(2)
        
        with col_u1:
            st.subheader("Operaciones por Carpeta")
                # Selector de ruta estandarizado
            path_unif = render_path_selector(
                label="Carpeta de Trabajo",
                key="tab_unif_folder",
                default_path=default_path
            )
            
            if st.button("🗂️ Unificar PDF por Carpeta", key="btn_unif_pdf"):
                is_native = st.session_state.get('force_native_mode', True)
                if not path_unif:
                    st.error("Carpeta base inválida.")
                elif not is_native and not os.path.isdir(path_unif):
                    st.error("Carpeta base inválida.")
                else:
                    try:
                        with st.spinner("Unificando PDFs..."):
                            if is_native:
                                username = st.session_state.get("username", "admin")
                                task_id = send_command(username, "unify_pdf_folder", {"base_path": path_unif, "output_name": "Unificado.pdf"})
                                if task_id:
                                    res = wait_for_result(task_id, timeout=120)
                                    if res and "error" not in res:
                                        st.success(res.get("message", "PDFs unificados correctamente por el agente."))
                                    else:
                                        st.error(f"Error del agente: {res.get('error', 'Tiempo de espera agotado')}")
                                else:
                                    st.error("Error enviando tarea al agente.")
                            else:
                                result = worker_unificar_por_carpeta(path_unif, "Unificado")
                                if isinstance(result, dict) and "error" in result:
                                    st.error(result["error"])
                                else:
                                    if isinstance(result, dict) and "message" in result:
                                        st.success(result["message"])
                                    else:
                                        st.success(result)
                                    out_file = os.path.join(path_unif, "Unificado.pdf")
                                    # Try to find if any subfolder has a unificado.pdf since it's a batch operation
                                    render_download_button(path_unif, f"dl_unif_pdf_{int(time.time())}", "📦 Descargar Resultados (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")
            
            if st.button("🖼️ Unificar JPG por Carpeta", key="btn_unif_jpg"):
                is_native = st.session_state.get('force_native_mode', True)
                if not path_unif:
                    st.error("Carpeta base inválida.")
                elif not is_native and not os.path.isdir(path_unif):
                    st.error("Carpeta base inválida.")
                else:
                    try:
                        with st.spinner("Unificando JPGs..."):
                            if is_native:
                                username = st.session_state.get("username", "admin")
                                task_id = send_command(username, "unify_img_folder", {"base_path": path_unif, "output_name": "Unificado.pdf", "img_type": "JPG"})
                                if task_id:
                                    res = wait_for_result(task_id, timeout=120)
                                    if res and "error" not in res:
                                        st.success(res.get("message", "JPGs unificados correctamente por el agente."))
                                    else:
                                        st.error(f"Error del agente: {res.get('error', 'Tiempo de espera agotado')}")
                                else:
                                    st.error("Error enviando tarea al agente.")
                            else:
                                result = worker_unificar_imagenes_por_carpeta_rec(path_unif, "Unificado.pdf", "JPG")
                                if isinstance(result, dict) and "error" in result:
                                    st.error(result["error"])
                                else:
                                    if isinstance(result, dict) and "message" in result:
                                        st.success(result["message"])
                                    else:
                                        st.success(result)
                                    out_file = os.path.join(path_unif, "Unificado.pdf")
                                    render_download_button(path_unif, f"dl_unif_jpg_{int(time.time())}", "📦 Descargar Resultados (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")
                
            if st.button("🖼️ Unificar PNG por Carpeta", key="btn_unif_png"):
                is_native = st.session_state.get('force_native_mode', True)
                if not path_unif:
                    st.error("Carpeta base inválida.")
                elif not is_native and not os.path.isdir(path_unif):
                    st.error("Carpeta base inválida.")
                else:
                    try:
                        with st.spinner("Unificando PNGs..."):
                            if is_native:
                                username = st.session_state.get("username", "admin")
                                task_id = send_command(username, "unify_img_folder", {"base_path": path_unif, "output_name": "Unificado.pdf", "img_type": "PNG"})
                                if task_id:
                                    res = wait_for_result(task_id, timeout=120)
                                    if res and "error" not in res:
                                        st.success(res.get("message", "PNGs unificados correctamente por el agente."))
                                    else:
                                        st.error(f"Error del agente: {res.get('error', 'Tiempo de espera agotado')}")
                                else:
                                    st.error("Error enviando tarea al agente.")
                            else:
                                result = worker_unificar_imagenes_por_carpeta_rec(path_unif, "Unificado.pdf", "PNG")
                                if isinstance(result, dict) and "error" in result:
                                    st.error(result["error"])
                                else:
                                    if isinstance(result, dict) and "message" in result:
                                        st.success(result["message"])
                                    else:
                                        st.success(result)
                                    out_file = os.path.join(path_unif, "Unificado.pdf")
                                    render_download_button(path_unif, f"dl_unif_png_{int(time.time())}", "📦 Descargar Resultados (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")
                
            if st.button("📄 Unificar DOCX por Carpeta", key="btn_unif_docx"):
                is_native = st.session_state.get('force_native_mode', True)
                if not path_unif:
                    st.error("Carpeta base inválida.")
                elif not is_native and not os.path.isdir(path_unif):
                    st.error("Carpeta base inválida.")
                else:
                    try:
                        with st.spinner("Unificando DOCX..."):
                            # Use native agent delegation if active
                            if is_native:
                                username = st.session_state.get("username", "admin")
                                task_id = send_command(username, "unify_docx_folder", {"base_path": path_unif, "output_name": "Unificado.pdf"})
                                if task_id:
                                    res = wait_for_result(task_id, timeout=120)
                                    if res and "error" not in res:
                                        st.success(res.get("message", "DOCX unificados correctamente por el agente."))
                                    else:
                                        err = res.get("error") if res else "Tiempo de espera agotado"
                                        st.error(f"Error del agente: {err}")
                                else:
                                    st.error("Error enviando tarea al agente.")
                            else:
                                result = worker_unificar_docx_por_carpeta(path_unif, "Unificado.pdf")
                                if isinstance(result, dict) and "error" in result:
                                    st.error(result["error"])
                                else:
                                    if isinstance(result, dict) and "message" in result:
                                        st.success(result["message"])
                                    else:
                                        st.success(result)
                                    out_file = os.path.join(path_unif, "Unificado.pdf")
                                    render_download_button(path_unif, f"dl_unif_docx_{int(time.time())}", "📦 Descargar Resultados (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")

            st.divider()
            if st.button("✂️ Dividir PDFs Masivamente", key="btn_split_mass"):
                is_native = st.session_state.get('force_native_mode', True)
                if not path_unif:
                    st.error("Carpeta base inválida.")
                elif not is_native and not os.path.isdir(path_unif):
                    st.error("Carpeta base inválida.")
                else:
                    try:
                        with st.spinner("Dividiendo PDFs..."):
                            if is_native:
                                username = st.session_state.get("username", "admin")
                                task_id = send_command(username, "split_pdf_massive", {"base_path": path_unif})
                                if task_id:
                                    res = wait_for_result(task_id, timeout=120)
                                    if res and "error" not in res:
                                        st.success(res.get("message", "PDFs divididos correctamente por el agente."))
                                    else:
                                        st.error(f"Error del agente: {res.get('error', 'Tiempo de espera agotado')}")
                                else:
                                    st.error("Error enviando tarea al agente.")
                            else:
                                result = worker_dividir_pdfs_masivamente(path_unif)
                                if isinstance(result, dict) and "error" in result:
                                    st.error(result["error"])
                                else:
                                    if isinstance(result, dict) and "message" in result:
                                        st.success(result["message"])
                                    else:
                                        st.success(result)
                                    div_folder = os.path.join(path_unif, "Dividido")
                                    if os.path.exists(div_folder):
                                        render_download_button(div_folder, "dl_split_mass", "📦 Descargar Carpeta Dividido (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")

        with col_u2:
            st.subheader("Operaciones Manuales")
            
            # Manual PDF Unify
            uploaded_pdfs = st.file_uploader("Unificar PDFs (Manual)", type=['pdf'], accept_multiple_files=True, key="col1_pdf_man")
            if st.button("🧷 Unificar Seleccionados", key="btn_unif_sel"):
                 if uploaded_pdfs:
                     try:
                         with st.spinner("Unificando PDFs seleccionados..."):
                             out_path = "Unificado_Manual.pdf" # Pass generic name to force stream download
                             result = worker_unificar_pdfs_list(uploaded_pdfs, out_path)
                             if isinstance(result, dict) and "error" in result:
                                 st.error(result["error"])
                             else:
                                 if isinstance(result, dict) and "message" in result:
                                     st.success(result["message"])
                                     if "files" in result:
                                         for f in result["files"]:
                                             st.download_button(
                                                 label=f.get("label", f"⬇️ Descargar {f['name']}"),
                                                 data=f["data"],
                                                 file_name=f["name"],
                                                 key=f"dl_unif_man_{f['name']}_{int(time.time())}"
                                             )
                                     else:
                                         render_download_button(out_path, f"dl_unif_man_{int(time.time())}", "⬇️ Descargar Unificado Manual")
                                 else:
                                     st.success(result)
                                     render_download_button(out_path, f"dl_unif_man_{int(time.time())}", "⬇️ Descargar Unificado Manual")
                     except Exception as e:
                         st.error(f"Error: {e}")

            st.divider()

            # Manual Split
            uploaded_split = st.file_uploader("Dividir PDF (Manual)", type=['pdf'], key="col1_split_man")
            if st.button("✂️ Dividir en Páginas", key="btn_split_man"):
                if uploaded_split:
                    try:
                        with st.spinner("Dividiendo PDF..."):
                            out_folder = os.path.join(st.session_state.get('current_path', '.'), "Dividido")
                            result = worker_dividir_pdf_paginas(uploaded_split, out_folder)
                            if isinstance(result, dict) and "error" in result:
                                st.error(result["error"])
                            else:
                                if isinstance(result, dict) and "message" in result:
                                    st.success(result["message"])
                                    if "files" in result:
                                        for f in result["files"]:
                                            st.download_button(
                                                label=f.get("label", f"📦 Descargar {f['name']}"),
                                                data=f["data"],
                                                file_name=f["name"],
                                                key=f"dl_split_man_{f['name']}_{int(time.time())}"
                                            )
                                    else:
                                        render_download_button(out_folder, "dl_split_man", "📦 Descargar Páginas Divididas (ZIP)")
                                else:
                                    st.success(result)
                                    render_download_button(out_folder, "dl_split_man", "📦 Descargar Páginas Divididas (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")

    # --- TAB 2: Organización ---
    with tab_org:
        st.caption("Organización de facturas, movimiento por coincidencia y consolidación.")
        
        # Selector de ruta estandarizado
        path_org = render_path_selector(
            label="Carpeta de Trabajo",
            key="tab_org_folder",
            default_path=default_path
        )
        
        col_o1, col_o2 = st.columns(2)
        with col_o1:
            if st.button("📥 Organizar Facturas (FEOV)", key="btn_org_feov_new"):
                st.session_state.active_auto_dialog = "organizar_feov"
                st.rerun()
                
            if st.button("📂➡️📁 Mover por Coincidencia", key="btn_org_move"):
                if path_org:
                    try:
                        with st.spinner("Moviendo archivos por coincidencia..."):
                            result = worker_mover_por_coincidencia(path_org)
                            st.success(result)
#                             render_download_button(path_org, "dl_org_move", "📦 Descargar Resultado (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")
                


        with col_o2:
            if st.button("📤 Copiar Archivo a Subcarpetas", key="btn_org_copy_sub_new"):
                st.session_state.active_auto_dialog = "copiar_mapeo_sub"
                st.rerun()
                
            if st.button("🗺️ Copiar Archivos (Mapeo Sub)", key="btn_org_map_sub_new"):
                st.session_state.active_auto_dialog = "copiar_mapeo_subcarpetas"
                st.rerun()
                
            if st.button("📜 Copiar Archivos Raíz (Mapeo)", key="btn_org_map_root_new"):
                st.session_state.active_auto_dialog = "copiar_mapeo_raiz"
                st.rerun()
                
            if st.button("📤 Consolidar Subcarpetas", key="btn_org_consol"):
                if path_org:
                    try:
                        with st.spinner("Consolidando subcarpetas..."):
                            result = worker_consolidar_subcarpetas(path_org)
                            st.success(result)
#                             render_download_button(path_org, "dl_org_consol", "📦 Descargar Consolidado (ZIP)")
                    except Exception as e:
                        st.error(f"Error: {e}")

    # --- TAB 3: Modificación ---
    with tab_modif:
        st.caption("Renombrado masivo con Excel y modificación de documentos DOCX.")
        
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            if st.button("📤 Exportar para renombrar", key="btn_mod_exp_new"):
                st.session_state.active_auto_dialog = "exportar_nombres"
                st.rerun()
                
            if st.button("📥 Aplicar renombrado Excel", key="btn_mod_app_new"):
                st.session_state.active_auto_dialog = "aplicar_nombres"
                st.rerun()
                
            if st.button("🏷️ Añadir Sufijo desde Excel", key="btn_mod_suf_new"):
                st.session_state.active_auto_dialog = "sufijo_archivos"
                st.rerun()

            if st.button("📝 Renombrar Masivo por Mapeo Excel", key="btn_mod_map_new"):
                st.session_state.active_auto_dialog = "renombrar_excel"
                st.rerun()
                
        with col_m2:
            if st.button("✍️ Modif. DOCX Completo", key="btn_mod_full_new"):
                st.session_state.active_auto_dialog = "modif_docx"
                st.rerun()
                
            if st.button("🖋️ Firmar DOCX con Imagen", key="btn_mod_sign_new"):
                st.session_state.active_auto_dialog = "insertar_firma_docx"
                st.rerun()

    # --- TAB 4: Análisis ---
    with tab_an:
        st.caption("Análisis y extracción de datos de historias clínicas y otros documentos.")
        
        # Selector de ruta estandarizado
        path_an = render_path_selector(
            label="Carpeta de Análisis",
            key="tab_an_folder",
            default_path=default_path
        )
        
        is_native_mode = st.session_state.get('force_native_mode', True)
        
        # Obtener lista de PDFs para análisis (Recursivo)
        files_pdf = []
        path_is_valid = False
        
        if path_an:
            if is_native_mode:
                # En modo nativo, asumimos que la ruta es válida y pasamos la ruta al agente para que la expanda
                path_is_valid = True
                files_pdf = [path_an]
            elif os.path.exists(path_an):
                path_is_valid = True
                for root, dirs, files in os.walk(path_an):
                    for f in files:
                        if f.lower().endswith('.pdf'):
                            files_pdf.append(os.path.join(root, f))
        
        if path_is_valid and not is_native_mode:
            st.caption(f"📂 Se encontraron {len(files_pdf)} archivos PDF en {path_an} (búsqueda recursiva).")
        elif path_is_valid and is_native_mode:
            st.caption(f"📂 El Agente Local buscará archivos PDF en: {path_an}")
        
        col_a1, col_a2 = st.columns(2)

        def run_analysis_sync(func, args, key_prefix):
            try:
                # Prevent multiple immediate clicks from firing multiple identical requests
                is_running_key = f"running_{key_prefix}"
                if st.session_state.get(is_running_key, False):
                    st.warning("El análisis ya se está ejecutando. Por favor espere.")
                    return
                
                st.session_state[is_running_key] = True
                
                with st.spinner("Procesando..."):
                    result = func(*args)
                
                if result and isinstance(result, dict):
                    if "error" in result:
                        st.error(f"Error en análisis: {result['error']}")
                    elif "files" in result:
                        st.session_state[f"analysis_result_{key_prefix}"] = result
                        st.session_state["last_run"] = int(time.time())
                    else:
                         st.warning("El resultado no tiene el formato esperado para descarga directa.")
                elif result:
                     st.warning("El resultado no tiene el formato esperado para descarga directa.")
                else:
                    st.warning("No se generaron resultados.")
            except Exception as e:
                st.error(f"Error: {e}")
            finally:
                st.session_state[is_running_key] = False

        # Renderizar resultados guardados en session_state
        def render_analysis_results(key_prefix):
            state_key = f"analysis_result_{key_prefix}"
            if state_key in st.session_state:
                result = st.session_state[state_key]
                st.success(result.get("message", "Análisis completado."))
                for i, f in enumerate(result["files"]):
                    data = f["data"]
                    if hasattr(data, "getvalue"): 
                        data = data.getvalue()
                    elif isinstance(data, str):
                        try:
                            import base64
                            import binascii
                            data = base64.b64decode(data, validate=True)
                        except Exception:
                            # Not base64, assume raw text
                            data = data.encode('utf-8')
                    
                    st.download_button(
                        label=f["label"] if "label" in f else f"📥 Descargar {f['name']}",
                        data=data,
                        file_name=f["name"],
                        mime=f.get("mime", "application/octet-stream"),
                        key=f"{key_prefix}_dl_{i}_{st.session_state.get('last_run', 0)}"
                    )

        with col_a1:
            if st.button("📊 Análisis Carpetas (Excel)", key="btn_an_folders"):
                if path_an:
                     # Clear previous result to avoid confusion
                     if "analysis_result_an_folders" in st.session_state:
                         del st.session_state["analysis_result_an_folders"]
                     run_analysis_sync(worker_analisis_carpetas, [path_an], "an_folders")
                else:
                    st.warning("Seleccione una carpeta válida.")
            render_analysis_results("an_folders")
            
            if st.button("📊 Análisis SOS", key="btn_an_sos"):
                 if files_pdf: 
                     run_analysis_sync(worker_analisis_sos, [files_pdf], "an_sos")
                 else: 
                     st.warning("No se encontraron PDFs.")
            render_analysis_results("an_sos")

        with col_a2:
            if st.button("📊 Análisis Historia Clínica", key="btn_an_hc"):
                if files_pdf:
                    run_analysis_sync(worker_analisis_historia_clinica, [files_pdf], "an_hc")
                else:
                    st.warning("No se encontraron PDFs.")
            render_analysis_results("an_hc")
            
            if st.button("📊 Análisis Autoriz. Nueva EPS", key="btn_an_neps"):
                if files_pdf:
                    run_analysis_sync(worker_analisis_autorizacion_nueva_eps, [files_pdf], "an_neps")
                else:
                    st.warning("No se encontraron PDFs.")
            render_analysis_results("an_neps")

            if st.button("📊 Análisis Cargue Sanitas", key="btn_an_sanitas"):
                 if files_pdf:
                    run_analysis_sync(worker_analisis_cargue_sanitas, [files_pdf], "an_sanitas")
                 else:
                    st.warning("No se encontraron PDFs.")
            render_analysis_results("an_sanitas")

            if st.button("📊 Análisis Retefuente/ICA", key="btn_an_rete"):
                 if files_pdf:
                    run_analysis_sync(worker_leer_pdf_retefuente, [files_pdf], "an_rete")
                 else:
                    st.warning("No se encontraron PDFs.")
            render_analysis_results("an_rete")

            if st.button("📊 Análisis Aut. Emssanar", key="btn_an_emssanar"):
                 if files_pdf:
                    run_analysis_sync(worker_analisis_emssanar, [files_pdf], "an_emssanar")
                 else:
                    st.warning("No se encontraron PDFs.")
            render_analysis_results("an_emssanar")

            if st.button("📊 Análisis Aut. FOMAG", key="btn_an_fomag"):
                 if files_pdf:
                    run_analysis_sync(worker_analisis_fomag, [files_pdf], "an_fomag")
                 else:
                    st.warning("No se encontraron PDFs.")
            render_analysis_results("an_fomag")

    # --- TAB 5: Creación y Otros ---
    with tab_create:
        st.caption("Creación de carpetas, firmas digitales y distribución de archivos.")
        
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            st.subheader("Creación")
            if st.button("📂 Crear Carpetas (Excel)", key="btn_cr_folders"):
                st.session_state.active_auto_dialog = "crear_carpetas"
                st.rerun()
                
            if st.button("⬇️ Descargar Firmas", key="btn_cr_sigs"):
                st.session_state.active_auto_dialog = "descargar_firmas"
                st.rerun()
                
            if st.button("⬇️ Descargar Hist. OVIDA", key="btn_cr_ovida"):
                st.session_state.active_auto_dialog = "descargar_ovida"
                st.rerun()
                
            if st.button("✒️ Crear Firma Digital", key="btn_cr_dig_sig"):
                st.session_state.active_auto_dialog = "crear_firma"
                st.rerun()

        with col_c2:
            st.subheader("Distribución / Otros")
            
            if st.button("📂 Distribuir Base", key="btn_dist_base"):
                st.session_state.active_auto_dialog = "distribuir_base"
                st.rerun()

    # Move dialog triggers to the root scope to avoid "Only one dialog" exception
    active_auto_dialog = st.session_state.get("active_auto_dialog")
    if active_auto_dialog:
        if active_auto_dialog == "crear_carpetas":
            dialog_crear_carpetas_excel()
        elif active_auto_dialog == "descargar_firmas":
            dialog_descargar_firmas()
        elif active_auto_dialog == "descargar_ovida":
            dialog_descargar_historias_ovida()
        elif active_auto_dialog == "crear_firma":
            dialog_crear_firma()
        elif active_auto_dialog == "distribuir_base":
            dialog_distribuir_base()
        elif active_auto_dialog == "organizar_feov":
            dialog_organizar_feov()
        elif active_auto_dialog == "mover_coincidencia":
            dialog_organizar_feov_avanzado()
        elif active_auto_dialog == "copiar_mapeo_sub":
            dialog_copiar_archivo_a_subcarpetas()
        elif active_auto_dialog == "copiar_mapeo_subcarpetas":
            dialog_copiar_mapeo()
        elif active_auto_dialog == "copiar_mapeo_raiz":
            dialog_copiar_raiz()
        elif active_auto_dialog == "autorizacion_docx":
            dialog_autorizacion_docx()
        elif active_auto_dialog == "regimen_docx":
            dialog_regimen_docx()
        elif active_auto_dialog == "sufijo_archivos":
            dialog_sufijo()
        elif active_auto_dialog == "renombrar_excel":
            dialog_renombrar_mapeo_excel()
        elif active_auto_dialog == "exportar_nombres":
            dialog_exportar_renombrado()
        elif active_auto_dialog == "aplicar_nombres":
            dialog_aplicar_renombrado()
        elif active_auto_dialog == "modif_docx":
            dialog_modif_docx_completo()
        elif active_auto_dialog == "insertar_firma_docx":
            dialog_insertar_firma_docx()
        else:
            if "active_auto_dialog" in st.session_state:
                del st.session_state["active_auto_dialog"]


