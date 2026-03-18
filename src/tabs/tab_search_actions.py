import streamlit as st
import os
import time
from datetime import datetime
import shutil
import re
import pandas as pd
import zipfile
import tempfile
import io

try:
    import agent_client
except ImportError:
    try:
        from src import agent_client
    except ImportError:
        agent_client = None

try:
    from send2trash import send2trash
except ImportError:
    send2trash = None

try:
    from docx import Document
except ImportError:
    Document = None

# Imports from project


try:
    from gui_utils import abrir_dialogo_carpeta_nativo, render_path_selector, update_path_key
except ImportError:
    try:
        from src.gui_utils import abrir_dialogo_carpeta_nativo, render_path_selector, update_path_key
    except ImportError:
        abrir_dialogo_carpeta_nativo = None
        
        def render_path_selector(label, key, default_path=None, help_text=None, omit_checkbox=False):
            st.warning("render_path_selector no disponible")
            return default_path
            
        def update_path_key(key, new_path, widget_key=None):
             pass

        def render_download_button(folder_path, key, label="📦 Descargar ZIP"):
            pass

# --- HELPER FUNCTIONS ---

def log(msg):
    timestamp = datetime.now().strftime("%H:%M:%S")
    if "logs" not in st.session_state:
        st.session_state.logs = []
    st.session_state.logs.append(f"[{timestamp}] {msg}")
    print(msg)

def record_action(action_type, changes):
    """
    Registra una acción en el historial para poder deshacerla.
    changes: lista de tuplas (ruta_actual, ruta_original)
    """
    if "action_history" not in st.session_state:
        st.session_state.action_history = []
    st.session_state.action_history.append({"type": action_type, "changes": changes})

def undo_last_action():
    if "action_history" not in st.session_state or not st.session_state.action_history:
        st.warning("No hay acciones para deshacer.")
        return

    last_action = st.session_state.action_history.pop()
    action_type = last_action["type"]
    changes = last_action["changes"]
    
    success_count = 0
    errors = 0
    
    progress = st.progress(0, text=f"Deshaciendo {action_type}...")
    total = len(changes)
    
    # Invertir lista de cambios por si hubo operaciones dependientes
    for i, (curr, orig) in enumerate(reversed(changes)):
        progress.progress(min((i + 1) / total, 1.0), text=f"Restaurando {i+1}/{total}")
        try:
            if os.path.exists(curr):
                # Verificar si el destino original ya existe (colisión al deshacer)
                if os.path.exists(orig):
                    # Intentar renombrar el existente temporalmente o fallar?
                    # Estrategia simple: Timestamp si colisión
                    base, ext = os.path.splitext(orig)
                    orig = f"{base}_restored_{int(time.time())}{ext}"
                
                shutil.move(curr, orig)
                success_count += 1
            else:
                errors += 1
                log(f"No se pudo deshacer: {curr} no existe.")
        except Exception as e:
            errors += 1
            log(f"Error al deshacer {curr} -> {orig}: {e}")
            
    progress.progress(1.0, text="Deshacer finalizado.")
    st.success(f"Deshacer completado: {success_count} elementos restaurados. ({errors} errores)")
    time.sleep(1.5)
    st.rerun()



def funcion_no_implementada(nombre):
    st.toast(f"⚠️ Función '{nombre}' simulada en versión web.")
    log(f"Ejecutado: {nombre}")

def buscar_archivos_wrapper():
    """Wrapper para llamar a buscar_archivos desde el botón y forzar rerun"""
    buscar_archivos()

def buscar_archivos():
    path = st.session_state.get("current_path", "")
    
    # Check for empty path but allow skipping existence check if in native mode (Agent will check)
    is_native = st.session_state.get("force_native_mode", True)
    
    if not path:
        st.warning("⚠️ Por favor, seleccione una carpeta válida para analizar.")
        return

    # In WEB mode, we must check existence locally. In NATIVE mode, path is on user's PC, so server can't check it.
    if not is_native and not os.path.exists(path):
         st.warning(f"⚠️ La ruta no existe en el servidor: {path}")
         return

    pattern_input = st.session_state.get("pattern", "")
    # Split by newlines first, then replace commas with newlines to split by them too, or just regex split
    patterns = [p.strip().lower() for p in pattern_input.replace(',', '\n').split('\n') if p.strip()]
    
    exclusion = st.session_state.get("exclusion_pattern", "").lower()
    exclusion_list = [x.strip() for x in exclusion.split(",") if x.strip()]

    search_by = st.session_state.get("search_by", "todo") # extensión, nombre, todo
    item_type = st.session_state.get("item_type", "archivos") # archivos, carpetas
    recursive = st.session_state.get("subfolders", True)
    search_empty_folders = st.session_state.get("search_empty_folders", False)

    results = []
    log(f"Iniciando búsqueda en: {path} | Patrones: '{patterns}' | Excluir: '{exclusion}' | Carpetas vacías: {search_empty_folders}")

    is_native = st.session_state.get("force_native_mode", True)
    
    if is_native:
        if agent_client:
            try:
                 username = st.session_state.get("username", "admin")
                 task_id = agent_client.send_command(username, "search_files", {
                     "path": path,
                     "patterns": patterns,
                     "exclusion_list": exclusion_list,
                     "search_by": search_by,
                     "item_type": item_type,
                     "recursive": recursive,
                     "search_empty_folders": search_empty_folders
                 })
                 
                 if task_id:
                                  with st.spinner("Buscando en equipo local (vía Agente)..."):
                                      res = agent_client.wait_for_result(task_id, timeout=60)
                                      
                                      # DEBUG: Inspect raw result
                                      log(f"DEBUG RAW RESULT TYPE: {type(res)}")
                                      log(f"DEBUG RAW RESULT CONTENT: {str(res)[:500]}...") # Log first 500 chars
                                      
                                      # Debug UI
                                      with st.expander("🕵️ Debug Agente (Raw Response)", expanded=False):
                                           st.write(f"Tipo: {type(res)}")
                                           st.json(res)

                                      if res and isinstance(res, list):
                                           # Safety check: if list of strings, convert to dicts
                                           if res and isinstance(res[0], str):
                                                safe_results = []
                                                for r in res:
                                                     try:
                                                         is_file = os.path.isfile(r)
                                                     except:
                                                         is_file = True # Assume file if check fails
                                                     safe_results.append({
                                                         "Ruta completa": r, 
                                                         "Nombre": os.path.basename(r), 
                                                         "Tipo": "Archivo" if is_file else "Carpeta",
                                                         "Fecha": "N/A"
                                                     })
                                                results.extend(safe_results)
                                           else:
                                                # Assume list of dicts, normalize them
                                                norm_results = []
                                                for item in res:
                                                     if isinstance(item, dict):
                                                          path = item.get("path", item.get("Ruta completa", ""))
                                                          name = item.get("name", item.get("Nombre", os.path.basename(path)))
                                                          itype = item.get("type", item.get("Tipo", "Archivo"))
                                                          date = item.get("mtime", item.get("Fecha", "N/A"))
                                                          norm_results.append({
                                                               "Ruta completa": path,
                                                               "Nombre": name,
                                                               "Tipo": itype,
                                                               "Fecha": date
                                                          })
                                                results.extend(norm_results)
                                      elif res and isinstance(res, dict):
                                           items = res.get("items", [])
                                           if not items and "result" in res:
                                                # Fallback if wrapped differently
                                                if isinstance(res["result"], list):
                                                     items = res["result"]
                                           
                                           if items:
                                                norm_results = []
                                                for item in items:
                                                     if isinstance(item, dict):
                                                          path = item.get("path", item.get("Ruta completa", ""))
                                                          name = item.get("name", item.get("Nombre", os.path.basename(path)))
                                                          itype = item.get("type", item.get("Tipo", "Archivo"))
                                                          date = item.get("mtime", item.get("Fecha", "N/A"))
                                                          norm_results.append({
                                                               "Ruta completa": path,
                                                               "Nombre": name,
                                                               "Tipo": itype,
                                                               "Fecha": date
                                                          })
                                                     elif isinstance(item, str):
                                                          norm_results.append({
                                                               "Ruta completa": item,
                                                               "Nombre": os.path.basename(item),
                                                               "Tipo": "Archivo",
                                                               "Fecha": "N/A"
                                                          })
                                                results.extend(norm_results)
                                           
                                           if "errors" in res and res["errors"]:
                                                st.warning(f"Errores reportados por agente: {res['errors']}")
                                      else:
                                           error_msg = res.get('error') if isinstance(res, dict) else 'Respuesta inválida o vacía'
                                           if not res:
                                                st.warning("El agente devolvió una respuesta vacía.")
                                           else:
                                                st.error(f"Error en la búsqueda del agente: {error_msg}")
                                           # st.code(f"Raw Response: {res}") # Show raw response for debugging
                 else:
                     st.error("No se pudo conectar con el agente para iniciar la búsqueda.")
            except Exception as e:
                st.error(f"Error al comunicar con el agente: {e}")
        else:
             st.error("Módulo de agente no disponible.")

    else:
        for root, dirs, files in os.walk(path):
            # Si no es recursivo, limpiar dirs para que os.walk no baje más, 
            # PERO debemos procesar la raíz actual.
            # os.walk yielda (root, dirs, files). Si modificamos dirs in-place, afecta la recursión.
            
            items_to_check = []
            if search_empty_folders:
                 # Si buscamos carpetas vacías, iteramos sobre dirs
                 items_to_check = dirs
            elif item_type == "archivos":
                items_to_check = files
            elif item_type == "carpetas":
                items_to_check = dirs
            
            for item in items_to_check:
                item_lower = item.lower()
                full_path = os.path.join(root, item)
                
                # Verificar exclusiones
                if any(excl in item_lower for excl in exclusion_list):
                    continue

                match = False
                
                if search_empty_folders:
                    # Lógica específica para carpetas vacías
                    try:
                        if os.path.isdir(full_path) and not os.listdir(full_path):
                             match = True
                    except PermissionError:
                        continue
                    except Exception as e:
                        log(f"Error accediendo a {full_path}: {e}")
                        continue
                else:
                    # Lógica normal de patrones
                    if not patterns:
                        match = True
                    else:
                        # Verificar si coincide con ALGUNO de los patrones
                        for pat in patterns:
                            if search_by == "extensión":
                                # Solo aplica a archivos
                                if item_type == "archivos" and item_lower.endswith(pat):
                                    match = True
                                    break
                            elif search_by == "nombre":
                                # Nombre sin extensión para archivos, o nombre carpeta
                                name_only = os.path.splitext(item_lower)[0] if item_type == "archivos" else item_lower
                                if pat in name_only:
                                    match = True
                                    break
                            else: # todo
                                if pat in item_lower:
                                    match = True
                                    break
                
                if match:
                    try:
                        stats = os.stat(full_path)
                        mtime = datetime.fromtimestamp(stats.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                    except:
                        mtime = "N/A"
                    
                    results.append({
                        "Ruta completa": full_path,
                        "Fecha": mtime
                    })

            if not recursive:
                break
            
    st.session_state.search_results = results
    if not results:
        st.warning("No se encontraron coincidencias.")
    else:
        st.success(f"Encontrados {len(results)} elementos.")
        st.info(f"Total de archivos encontrados: {len(results)}")

# --- WORKERS ---

def procesar_renombrado(results, full, new_name, sust, find_txt, repl_txt, clean_feov, pre, prefix_txt, suf, suffix_txt, activar_num=False, inicio_num=1, silent_mode=False):
    count = 0
    is_native = st.session_state.get("force_native_mode", True)
    
    if not results:
        msg = "No hay archivos en la lista de resultados para procesar."
        if not silent_mode:
            st.warning(msg)
        return msg

    # Validación de colisiones para Renombrado Completo (Solo Local)
    if not is_native and full and new_name and not activar_num:
        # Verificar si hay múltiples archivos con la misma extensión en la misma carpeta
        conflict_map = {}
        for item in results:
            old_path = item["Ruta completa"]
            if not os.path.exists(old_path): continue
            
            folder = os.path.dirname(old_path)
            _, ext = os.path.splitext(old_path)
            key = (folder, ext.lower())
            
            if key in conflict_map:
                conflict_map[key].append(os.path.basename(old_path))
            else:
                conflict_map[key] = [os.path.basename(old_path)]
        
        # Filtrar solo los que tienen conflictos (>1 archivo)
        real_conflicts = {k: v for k, v in conflict_map.items() if len(v) > 1}
        
        if real_conflicts:
            msg = "⚠️ ALERTA: Conflicto de nombres detectado. Estás intentando renombrar a un nombre fijo pero hay múltiples archivos con la misma extensión en la misma carpeta. Activa la numeración consecutiva para evitar esto."
            if not silent_mode:
                st.error(msg)
                st.write("El sistema ha detenido el proceso para evitar sobrescritura o nombres duplicados.")
                
                with st.expander("Ver detalles del conflicto"):
                    for (folder, ext), files in real_conflicts.items():
                        st.write(f"📂 En {folder} (Extensión {ext}):")
                        for f in files:
                            st.write(f"  - {f}")
            return msg

    total_files = len(results)
    if not silent_mode:
        progress_bar = st.progress(0, text="Renombrando en lote...")
    
    changes_made = [] # Lista para guardar cambios (Undo)
    batch_renames = [] # Para modo nativo (Agente)
    
    current_num = inicio_num

    for idx, item in enumerate(results):
        if not silent_mode and idx % 5 == 0:
             progress_bar.progress(min(idx / total_files, 1.0), text=f"Renombrando {idx+1}/{total_files}")
             
        old_path = item["Ruta completa"]
        if not is_native and not os.path.exists(old_path): continue
        if is_native and not old_path: continue
        
        folder = os.path.dirname(old_path)
        filename = os.path.basename(old_path)
        name_part, ext = os.path.splitext(filename)
        
        final_name = name_part
        
        # 1. Renombrado completo (Prioridad)
        if full and new_name:
            final_name = new_name
        else:
            # 2. Sustitución
            if sust and find_txt:
                final_name = final_name.replace(find_txt, repl_txt)
            
            # 3. Limpieza FEOV
            if clean_feov:
                # Regex para _ID..._A (ej: _ID12345_A)
                final_name = re.sub(r'_ID\d+_A', '', final_name)
            
            # 4. Prefijo
            if pre and prefix_txt:
                final_name = f"{prefix_txt}{final_name}"
            
            # 5. Sufijo
            if suf and suffix_txt:
                final_name = f"{final_name}{suffix_txt}"

        # 6. Numeración Consecutiva
        if activar_num:
            # Formato simple: _1, _2, etc. O podría ser con padding _001?
            # Usaremos formato simple por ahora o padding si son muchos
            padding = len(str(total_files + inicio_num))
            num_str = str(current_num).zfill(padding)
            final_name = f"{final_name}_{num_str}"
            current_num += 1
        
        new_filename = f"{final_name}{ext}"
        new_path = os.path.join(folder, new_filename)
        
        if new_path != old_path:
            if is_native:
                batch_renames.append({
                    "old_path": old_path,
                    "new_path": new_path,
                    "item": item
                })
            else:
                try:
                    # Manejo simple de colisiones
                    if os.path.exists(new_path):
                        timestamp = int(time.time())
                        new_filename = f"{final_name}_{timestamp}{ext}"
                        new_path = os.path.join(folder, new_filename)
                    
                    os.rename(old_path, new_path)
                    count += 1
                    changes_made.append((new_path, old_path)) # Guardar para deshacer
                    # Actualizar ruta en resultados para reflejar cambio
                    item["Ruta completa"] = new_path
                except Exception as e:
                    pass
                    # log(f"Error renombrando {filename}: {e}")

    # Procesar Lote Nativo (Agente)
    if is_native:
        if not agent_client:
            if not silent_mode: st.error("Error: Módulo agent_client no cargado.")
            return "Error: Agente no disponible"
            
        if not batch_renames:
            if not silent_mode: st.warning("No se generaron cambios para procesar (nombres idénticos o rutas vacías).")
            # DEBUG
            if not silent_mode and results:
                 st.write(f"DEBUG: Primer item old: {results[0].get('Ruta completa')}")
            return "No hay cambios pendientes."

        try:
            username = st.session_state.get("username", "default")
            
            if not silent_mode:
                st.info(f"Enviando {len(batch_renames)} archivos al agente para renombrar...")
            
            task_id = agent_client.send_command(username, "rename_files", {
                "files": [{"old_path": r["old_path"], "new_path": r["new_path"]} for r in batch_renames]
            })
            
            if task_id:
                if not silent_mode:
                     progress_bar.progress(0.9, text="Enviando tarea al agente...")
                
                # Esperar resultado
                res = agent_client.wait_for_result(task_id, timeout=600)
                
                if res and isinstance(res, dict):
                    count = res.get("count", 0)
                    errors = res.get("errors", [])
                    
                    # Actualizar items si éxito total (optimista)
                    # Si hay errores, no sabemos cuáles fallaron, así que no actualizamos items para no confundir
                    if count == len(batch_renames) and not errors:
                        for r in batch_renames:
                            r["item"]["Ruta completa"] = r["new_path"]
                    
                    if errors and not silent_mode:
                        st.error(f"Errores reportados por el agente: {len(errors)}")
                        with st.expander("Ver errores"):
                            for e in errors: st.write(e)
                else:
                    if not silent_mode: st.error("Respuesta inválida del agente")
            else:
                 if not silent_mode: st.error("No se pudo conectar con el agente")
        except Exception as e:
             if not silent_mode: st.error(f"Error enviando tarea al agente: {e}")

    if not silent_mode:
        progress_bar.progress(1.0, text="Proceso finalizado.")
    
    if count > 0:
        if not is_native:
            record_action("Renombrado Masivo", changes_made)
            
        msg = f"✅ Renombrados {count} archivos exitosamente."
        if not silent_mode:
            st.success(msg)
            
            # Offer download of the root folder (only relevant for local/web mode, not native)
            if not is_native:
                root_path = st.session_state.get("current_path")
                if root_path and os.path.exists(root_path):
#                      render_download_button(root_path, "dl_rename_mass", "📦 Descargar Carpeta Completa (ZIP)")
                    pass
                 
            # log(f"Renombrado masivo completado. Total: {count}")
            # time.sleep(1) # Pausa breve para ver el mensaje - Removed to allow interaction with download button
            # st.rerun() # Rerun prevents clicking the download button
        return msg
    else:
        msg = "No se realizaron cambios (verifique los parámetros o nombres de archivo)."
        if not silent_mode:
            st.info(msg)
        return msg

def run_renombrar_task(results, full, new_name, sust, find_txt, repl_txt, clean_feov, pre, prefix_txt, suf, suffix_txt, activar_num, inicio_num):
    return {"message": procesar_renombrado(results, full, new_name, sust, find_txt, repl_txt, clean_feov, pre, prefix_txt, suf, suffix_txt, activar_num, inicio_num, silent_mode=True)}

# --- ADDITIONAL WORKERS (Migrated from app_web.py) ---

def worker_editar_texto(file_list, search_text, replace_text, silent_mode=False):
    # Native Mode Agent Integration
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and agent_client:
        try:
            username = st.session_state.get("username", "default")
            if not silent_mode:
                st.info("Enviando tarea de edición de texto al Agente Local...")
            
            # Extract paths if file_list contains dicts
            items = []
            for f in file_list:
                if isinstance(f, dict):
                    items.append(f.get("Ruta completa"))
                else:
                    items.append(f)
            
            task_id = agent_client.send_command(username, "edit_text", {
                "items": items,
                "find": search_text,
                "replace": replace_text
            })
            
            if task_id:
                if not silent_mode:
                    with st.spinner("Procesando con Agente Local..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    # In silent mode we might block or not? 
                    # Usually workers are called from UI, so blocking is fine.
                    res = agent_client.wait_for_result(task_id, timeout=600)

                # Fix: wait_for_result returns the result payload directly
                if res and isinstance(res, dict) and "count" in res:
                    r_data = res
                    count = r_data.get("count", 0)
                    errors = r_data.get("errors", [])
                    msg = f"Agente: Texto modificado en {count} archivos."
                    if errors:
                        msg += f" Errores: {len(errors)}"
                    
                    if not silent_mode:
                        st.success(msg)
                        if errors:
                            with st.expander("Ver errores"):
                                for e in errors: st.write(e)
                    return msg
                elif res and isinstance(res, dict) and "error" in res:
                     err_msg = res.get("error")
                     if not silent_mode: st.error(f"Error del Agente: {err_msg}")
                     return f"Error Agente: {err_msg}"
                else:
                    err_msg = "Respuesta inesperada del Agente"
                    if not silent_mode: st.error(f"Error del Agente: {err_msg}")
                    return f"Error Agente: {err_msg}"
            else:
                if not silent_mode: st.error("No se pudo conectar con el Agente.")
        except Exception as e:
            if not silent_mode: st.error(f"Excepción Agente: {e}")

    # Web Mode Logic
    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Editando archivos...")
    total = len(file_list)
    
    for i, item in enumerate(file_list):
        file_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Procesando: {os.path.basename(file_path)}")
        
        try:
            if not os.path.exists(file_path):
                continue

            ext = os.path.splitext(file_path)[1].lower()
            modified = False
            
            # Archivos de texto plano
            if ext in ['.txt', '.json', '.xml', '.csv', '.html', '.md', '.log', '.py', '.js', '.css', '.bat', '.ps1']:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    
                    if search_text in content:
                        new_content = content.replace(search_text, replace_text)
                        with open(file_path, 'w', encoding='utf-8') as f:
                            f.write(new_content)
                        modified = True
                except Exception as e:
                    log(f"Error leyendo/escribiendo {file_path}: {e}")
                    errors += 1

            # Documentos de Word
            elif ext == '.docx':
                try:
                    if Document is not None:
                        doc = Document(file_path)
                        doc_modified = False
                        for p in doc.paragraphs:
                            if search_text in p.text:
                                p.text = p.text.replace(search_text, replace_text)
                                doc_modified = True
                        # También buscar en tablas
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        if search_text in p.text:
                                            p.text = p.text.replace(search_text, replace_text)
                                            doc_modified = True
                                            
                        if doc_modified:
                            doc.save(file_path)
                            modified = True
                    else:
                        log(f"No se puede procesar DOCX (python-docx no disponible): {file_path}")
                        errors += 1
                        if not silent_mode:
                            st.warning(f"Omitido {os.path.basename(file_path)}: Falta 'python-docx'")
                except Exception as e:
                    log(f"Error procesando DOCX {file_path}: {e}")
                    errors += 1
                    if not silent_mode:
                        st.error(f"Error en {os.path.basename(file_path)}: {e}")
            
            if modified:
                count += 1
                log(f"Texto modificado en: {file_path}")
                
        except Exception as e:
            log(f"Error general en {file_path}: {e}")
            errors += 1
            if not silent_mode:
                st.error(f"Error en {os.path.basename(file_path)}: {e}")
            
    msg = f"Proceso finalizado. Archivos modificados: {count}. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        
        # Offer download of the root folder
        root_path = st.session_state.get("current_path")
        if root_path and os.path.exists(root_path):
#              render_download_button(root_path, "dl_text_edit", "📦 Descargar Carpeta Completa (ZIP)")
             pass
             
        # time.sleep(2)
        # st.rerun()
    return msg

def run_editar_texto_task(file_list, search_text, replace_text):
    return {"message": worker_editar_texto(file_list, search_text, replace_text, silent_mode=True)}

def worker_copiar_lista(file_list, target_folder, silent_mode=False):
    # Native Mode Agent Integration
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and agent_client:
        try:
            username = st.session_state.get("username", "default")
            
            # Extract paths if file_list contains dicts
            items = []
            for f in file_list:
                if isinstance(f, dict):
                    items.append(f.get("Ruta completa"))
                else:
                    items.append(f)
            
            task_id = agent_client.send_command(username, "copy_files", {
                "items": items,
                "target_folder": target_folder
            })
            
            if task_id:
                if not silent_mode:
                    with st.spinner("Copiando archivos vía Agente..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    res = agent_client.wait_for_result(task_id, timeout=600)

                if res and isinstance(res, dict) and "count" in res:
                    r_data = res
                    count = r_data.get("count", 0)
                    errors = r_data.get("errors", [])
                    msg = f"Agente: Copiados {count} archivos."
                    if errors: msg += f" Errores: {len(errors)}"
                    
                    if not silent_mode:
                        st.success(msg)
                        if errors:
                             with st.expander("Errores de copia"):
                                 for e in errors: st.write(e)
                    return msg
                elif res and isinstance(res, dict) and "error" in res:
                    err_msg = res.get("error")
                    if not silent_mode: st.error(f"Fallo Agente: {err_msg}")
                    return f"Error: {err_msg}"
                else:
                    err = str(res) if res else "Error desconocido"
                    if not silent_mode: st.error(f"Fallo Agente: {err}")
                    return f"Error: {err}"
        except Exception as e:
            if not silent_mode: st.error(f"Error conexión Agente: {e}")

    if not os.path.exists(target_folder):
        try:
            os.makedirs(target_folder)
        except Exception as e:
            msg = f"Error creando carpeta destino: {e}"
            if not silent_mode: st.error(msg)
            return msg

    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Copiando archivos...")
    total = len(file_list)
    
    for i, item in enumerate(file_list):
        src_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Copiando: {os.path.basename(src_path)}")
        
        try:
            if not os.path.exists(src_path): continue
            
            filename = os.path.basename(src_path)
            dest_path = os.path.join(target_folder, filename)
            
            # Manejar colisiones
            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                dest_path = os.path.join(target_folder, f"{base}_{int(time.time())}{ext}")
            
            if os.path.isdir(src_path):
                shutil.copytree(src_path, dest_path)
            else:
                shutil.copy2(src_path, dest_path)
            count += 1
        except Exception as e:
            log(f"Error copiando {src_path}: {e}")
            errors += 1
            
    msg = f"Copiados {count} elementos. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        
#         render_download_button(target_folder, "dl_copy_list", "📦 Descargar Archivos Copiados (ZIP)")
        
        # time.sleep(1.5)
        # st.rerun()
    return msg

def run_copiar_lista_task(file_list, target_folder):
    return {"message": worker_copiar_lista(file_list, target_folder, silent_mode=True)}

def worker_mover_lista(file_list, target_folder, silent_mode=False):
    # Native Mode Agent Integration
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and agent_client:
        try:
            username = st.session_state.get("username", "default")
            
            # Extract paths if file_list contains dicts
            items = []
            for f in file_list:
                if isinstance(f, dict):
                    items.append(f.get("Ruta completa"))
                else:
                    items.append(f)
            
            task_id = agent_client.send_command(username, "move_files", {
                "items": items,
                "target_folder": target_folder
            })
            
            if task_id:
                if not silent_mode:
                    with st.spinner("Moviendo archivos vía Agente..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    res = agent_client.wait_for_result(task_id, timeout=600)

                if res and isinstance(res, dict) and "count" in res:
                    r_data = res
                    count = r_data.get("count", 0)
                    errors = r_data.get("errors", [])
                    msg = f"Agente: Movidos {count} archivos."
                    if errors: msg += f" Errores: {len(errors)}"
                    
                    if not silent_mode:
                        st.success(msg)
                        # Clear results as files are moved
                        st.session_state.search_results = []
                        if errors:
                             with st.expander("Errores de movimiento"):
                                 for e in errors: st.write(e)
                        time.sleep(1.5)
                        st.rerun()
                    return msg
                elif res and isinstance(res, dict) and "error" in res:
                    err_msg = res.get("error")
                    if not silent_mode: st.error(f"Fallo Agente: {err_msg}")
                    return f"Error: {err_msg}"
                else:
                    err = str(res) if res else "Error desconocido"
                    if not silent_mode: st.error(f"Fallo Agente: {err}")
                    return f"Error: {err}"
        except Exception as e:
            if not silent_mode: st.error(f"Error conexión Agente: {e}")

    if not os.path.exists(target_folder):
        try:
            os.makedirs(target_folder)
        except Exception as e:
            msg = f"Error creando carpeta destino: {e}"
            if not silent_mode: st.error(msg)
            return msg

    count = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Moviendo archivos...")
    total = len(file_list)
    
    changes_made = []

    for i, item in enumerate(file_list):
        src_path = item["Ruta completa"]
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Moviendo: {os.path.basename(src_path)}")
        
        try:
            if not os.path.exists(src_path): continue
            
            filename = os.path.basename(src_path)
            dest_path = os.path.join(target_folder, filename)
            
            # Manejar colisiones
            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                dest_path = os.path.join(target_folder, f"{base}_{int(time.time())}{ext}")
            
            shutil.move(src_path, dest_path)
            count += 1
            changes_made.append((dest_path, src_path)) # Guardar para deshacer

            # Actualizar ruta en lista (aunque se va a limpiar o rerun)
            item["Ruta completa"] = dest_path
        except Exception as e:
            log(f"Error moviendo {src_path}: {e}")
            errors += 1
            
    msg = f"Movidos {count} elementos. Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        if count > 0:
            record_action("Mover Lista", changes_made)
        st.success(msg)
        
#         render_download_button(target_folder, "dl_move_list", "📦 Descargar Archivos Movidos (ZIP)")
        
        # time.sleep(1.5)
        # st.rerun()
    return msg

def run_mover_lista_task(file_list, target_folder):
    return {"message": worker_mover_lista(file_list, target_folder, silent_mode=True)}


def worker_eliminar_lista(file_list, force_delete=False, silent_mode=False):
    # Native Mode Agent Integration
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and agent_client:
        try:
            username = st.session_state.get("username", "default")
            # Extract paths if file_list contains dicts
            items = []
            for f in file_list:
                if isinstance(f, dict):
                    items.append(f.get("Ruta completa"))
                else:
                    items.append(f)
            
            task_id = agent_client.send_command(username, "delete_files", {
                "items": items,
                "force": force_delete
            })
            
            if task_id:
                if not silent_mode:
                    with st.spinner("Eliminando archivos vía Agente..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    res = agent_client.wait_for_result(task_id, timeout=600)

                if res and isinstance(res, dict) and "count" in res:
                    r_data = res
                    count = r_data.get("count", 0)
                    errors = r_data.get("errors", [])
                    msg = f"Agente: Eliminados {count} archivos."
                    if errors: msg += f" Errores: {len(errors)}"
                    
                    if not silent_mode:
                        st.success(msg)
                        st.session_state.search_results = []
                        if errors:
                             with st.expander("Errores de eliminación"):
                                 for e in errors: st.write(e)
                        time.sleep(1.5)
                        st.rerun()
                    return msg
                elif res and isinstance(res, dict) and "error" in res:
                    err_msg = res.get("error")
                    if not silent_mode: st.error(f"Fallo Agente: {err_msg}")
                    return f"Error: {err_msg}"
                else:
                    err = str(res) if res else "Error desconocido"
                    if not silent_mode: st.error(f"Fallo Agente: {err}")
                    return f"Error: {err}"
        except Exception as e:
            if not silent_mode: st.error(f"Error conexión Agente: {e}")

    from send2trash import send2trash
    count_del = 0
    errors = 0
    if not silent_mode:
        progress_bar = st.progress(0, text="Eliminando...")
    total = len(file_list)
    
    for i, item in enumerate(file_list):
        if not silent_mode:
            progress_bar.progress(min(i/total, 1.0), text=f"Eliminando {i+1}/{total}")
        
        path = item["Ruta completa"] if isinstance(item, dict) else item
        
        try:
            if os.path.exists(path):
                safe_path = os.path.normpath(path)
                
                if force_delete:
                    if os.path.isdir(safe_path):
                        shutil.rmtree(safe_path)
                    else:
                        os.remove(safe_path)
                    count_del += 1
                else:
                    send2trash(safe_path)
                    count_del += 1
        except Exception as e:
            log(f"Error eliminando {path}: {e}")
            errors += 1
    
    msg = f"Se enviaron {count_del} archivos a la papelera (Forzado={force_delete}). Errores: {errors}"
    if not silent_mode:
        progress_bar.progress(1.0, text="Finalizado.")
        st.success(msg)
        st.session_state.search_results = [] 
        time.sleep(1.5)
        st.rerun()
    return msg

def run_eliminar_lista_task(file_list):
    return {"message": worker_eliminar_lista(file_list, force_delete=False, silent_mode=True)}

def worker_comprimir_zip(file_list, output_path, silent_mode=False):
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and agent_client:
        try:
            username = st.session_state.get("username", "default")
            
            items = []
            for f in file_list:
                if isinstance(f, dict):
                    items.append(f.get("Ruta completa"))
                else:
                    items.append(f)
            
            task_id = agent_client.send_command(username, "compress_zip", {
                "items": items,
                "output_path": output_path
            })
            
            if task_id:
                if not silent_mode:
                    with st.spinner("Comprimiendo archivos vía Agente..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    res = agent_client.wait_for_result(task_id, timeout=600)

                if res and res.get("status") == "SUCCESS":
                    r_data = res.get("result", {})
                    count = r_data.get("count", 0)
                    errors = r_data.get("errors", [])
                    msg = f"Agente: Comprimidos {count} elementos en {os.path.basename(output_path)}."
                    if errors: msg += f" Errores: {len(errors)}"
                    
                    if not silent_mode:
                        st.success(msg)
                        if errors:
                             with st.expander("Errores de compresión"):
                                 for e in errors: st.write(e)
                    return msg
                else:
                    err = res.get("result") if res else "Error desconocido"
                    if not silent_mode: st.error(f"Fallo Agente: {err}")
                    return f"Error: {err}"
        except Exception as e:
            if not silent_mode: st.error(f"Error conexión Agente: {e}")
            return f"Error: {e}"
    return "Modo nativo no activo o agente no conectado."

def worker_comprimir_individual(file_list, silent_mode=False):
    is_native = st.session_state.get("force_native_mode", True)
    if is_native and agent_client:
        try:
            username = st.session_state.get("username", "default")
            
            items = []
            for f in file_list:
                if isinstance(f, dict):
                    items.append(f.get("Ruta completa"))
                else:
                    items.append(f)
            
            task_id = agent_client.send_command(username, "compress_individual", {
                "items": items
            })
            
            if task_id:
                if not silent_mode:
                    with st.spinner("Comprimiendo individualmente vía Agente..."):
                        res = agent_client.wait_for_result(task_id, timeout=600)
                else:
                    res = agent_client.wait_for_result(task_id, timeout=600)

                if res and res.get("status") == "SUCCESS":
                    r_data = res.get("result", {})
                    count = r_data.get("count", 0)
                    errors = r_data.get("errors", [])
                    msg = f"Agente: Comprimidos {count} elementos individualmente."
                    if errors: msg += f" Errores: {len(errors)}"
                    
                    if not silent_mode:
                        st.success(msg)
                        if errors:
                             with st.expander("Errores de compresión"):
                                 for e in errors: st.write(e)
                    return msg
                else:
                    err = res.get("result") if res else "Error desconocido"
                    if not silent_mode: st.error(f"Fallo Agente: {err}")
                    return f"Error: {err}"
        except Exception as e:
            if not silent_mode: st.error(f"Error conexión Agente: {e}")
            return f"Error: {e}"
    return "Modo nativo no activo o agente no conectado."

# --- DIALOGS ---

@st.dialog("Modificar Nombres - Opciones Avanzadas")
def dialogo_modificar_nombres():
    st.write("Configura las opciones de renombrado:")
    
    # 1. Renombrado completo
    with st.container(border=True):
        st.markdown("**📝 Renombrado completo (ignora las demás opciones)**")
        activar_full = st.checkbox("Activar renombrado completo", key="chk_full")
        nuevo_nombre = st.text_input("Nuevo nombre (sin extensión)", disabled=not activar_full, key="txt_full")
    
    # 2. Sustituir texto
    with st.container(border=True):
        st.markdown("**🔄 Sustituir texto**")
        activar_sust = st.checkbox("Activar sustitución", key="chk_sust")
        buscar_txt = st.text_input("Buscar:", disabled=not activar_sust, key="txt_find")
        reemplazar_txt = st.text_input("Reemplazar con:", disabled=not activar_sust, key="txt_repl")

    # 3. Limpieza Especial
    with st.container(border=True):
        st.markdown("**🧹 Limpieza Especial (FEOV)**")
        eliminar_id = st.checkbox("Eliminar '_ID<números>_A' del nombre", key="chk_feov")

    # 4. Añadir al inicio
    with st.container(border=True):
        st.markdown("**⬅️ Añadir al inicio**")
        activar_pre = st.checkbox("Añadir prefijo", key="chk_pre")
        prefijo = st.text_input("Texto prefijo:", disabled=not activar_pre, key="txt_pre")

    # 5. Añadir al final
    with st.container(border=True):
        st.markdown("**➡️ Añadir al final (antes de extensión)**")
        activar_suf = st.checkbox("Añadir sufijo", key="chk_suf")
        sufijo = st.text_input("Texto sufijo:", disabled=not activar_suf, key="txt_suf")

    # 6. Numeración Consecutiva
    with st.container(border=True):
        st.markdown("**🔢 Numeración Consecutiva**")
        activar_num = st.checkbox("Añadir numeración consecutiva al final", key="chk_num")
        inicio_num = st.number_input("Iniciar desde:", min_value=1, value=1, disabled=not activar_num, key="num_inicio")

    st.markdown("---")
    col_cancel, col_ok = st.columns([1, 1])
    with col_ok:
        if st.button("✅ Ejecutar Cambios", use_container_width=True):
            procesar_renombrado(st.session_state.get("search_results", []),
                activar_full, nuevo_nombre, 
                activar_sust, buscar_txt, reemplazar_txt, 
                eliminar_id, 
                activar_pre, prefijo, 
                activar_suf, sufijo,
                activar_num, inicio_num,
                silent_mode=False
            )
            
    with col_cancel:
        if st.button("❌ Cancelar", use_container_width=True):
            st.session_state.active_action_dialog = None
            st.rerun()

@st.dialog("Editar Texto en Archivos")
def dialogo_editar_texto():
    st.write("Esta acción buscará y reemplazará texto en los archivos listados en la búsqueda.")
    st.warning("Aplica a archivos de texto plano (.txt, .json, .xml, .py, etc.) y documentos Word (.docx).")
    
    if Document is None:
        st.warning("⚠️ La librería 'python-docx' no está instalada. No se podrán editar archivos .docx.")
    
    if not st.session_state.get("search_results", []):
        st.error("No hay archivos en la lista de resultados.")
        return

    st.info(f"Archivos a procesar: {len(st.session_state.get('search_results', []))}")
    
    search_text = st.text_input("Texto a buscar:")
    replace_text = st.text_input("Reemplazar con:")
    
    col_ok, col_cancel = st.columns([1, 1])
    with col_ok:
        if st.button("🚀 Ejecutar Reemplazo", use_container_width=True):
            if not search_text:
                st.warning("Debes ingresar el texto a buscar.")
            else:
                worker_editar_texto(st.session_state.get("search_results", []), search_text, replace_text, silent_mode=False)
            
    with col_cancel:
        if st.button("❌ Cerrar", use_container_width=True):
            st.session_state.active_action_dialog = None
            st.rerun()

@st.dialog("Copiar Archivos de Lista")
def dialogo_copiar_lista():
    st.write("Copiará los archivos/carpetas de la lista de resultados a una nueva ubicación.")
    
    if not st.session_state.get("search_results", []):
        st.error("No hay elementos en la lista.")
        return

    st.info(f"Elementos a copiar: {len(st.session_state.get('search_results', []))}")
    
    current_global_path = st.session_state.get("current_path", os.getcwd())
    if not current_global_path: current_global_path = os.getcwd()

    st.text_input("Carpeta Origen (Búsqueda inicial):", value=current_global_path, disabled=True)

    # Use omit_checkbox=True to force the use of the custom path behavior
    target_copy_path = render_path_selector(
        label="Carpeta Destino:",
        key="copy_dest_input",
        default_path=current_global_path,
        omit_checkbox=True
    )
    
    col_ok, col_cancel = st.columns([1, 1])
    with col_ok:
        if st.button("🚀 Copiar", use_container_width=True):
            dest = st.session_state.get("copy_dest_input")
            if not dest:
                st.warning("Selecciona una carpeta destino.")
            else:
                worker_copiar_lista(st.session_state.get("search_results", []), dest, silent_mode=False)
            
    with col_cancel:
        if st.button("❌ Cerrar", use_container_width=True, key="btn_close_copy"):
            st.session_state.active_action_dialog = None
            st.rerun()

@st.dialog("Mover Archivos de Lista")
def dialogo_mover_lista():
    st.write("Moverá los archivos/carpetas de la lista de resultados a una nueva ubicación.")
    st.warning("Los archivos originales serán eliminados de su ubicación actual.")
    
    if not st.session_state.get("search_results", []):
        st.error("No hay elementos en la lista.")
        return

    st.info(f"Elementos a mover: {len(st.session_state.get('search_results', []))}")
    
    current_global_path = st.session_state.get("current_path", os.getcwd())
    if not current_global_path: current_global_path = os.getcwd()

    st.text_input("Carpeta Origen (Búsqueda inicial):", value=current_global_path, disabled=True)

    # Selector de Ruta Estandarizado
    # Use omit_checkbox=True to force the use of the custom path behavior
    target_move_path = render_path_selector(
        label="Carpeta Destino:",
        key="move_dest_input",
        default_path=current_global_path,
        help_text="Donde se moverán los archivos.",
        omit_checkbox=True
    )

    col_ok, col_cancel = st.columns([1, 1])
    with col_ok:
        if st.button("🚀 Mover", use_container_width=True):
            if not target_move_path:
                st.warning("Selecciona una carpeta destino.")
            else:
                worker_mover_lista(st.session_state.get("search_results", []), target_move_path, silent_mode=False)

    with col_cancel:
        if st.button("❌ Cerrar", use_container_width=True, key="btn_close_move"):
            st.session_state.active_action_dialog = None
            st.rerun()


@st.dialog("Confirmar Eliminación")
def dialogo_confirmar_eliminar():
    st.warning("⚠️ ¿Estás seguro de que quieres enviar estos archivos a la papelera?")
    st.write("Esta acción enviará los archivos a la papelera de reciclaje del sistema.")
    
    results = st.session_state.search_results
    if not results:
        st.error("No hay archivos seleccionados.")
        return

    st.write(f"Total de elementos: **{len(results)}**")
    
    with st.expander("Ver lista de archivos a eliminar"):
        for item in results[:50]: # Show first 50
            st.text(os.path.basename(item["Ruta completa"]))
        if len(results) > 50:
            st.text(f"... y {len(results)-50} más.")

    # Opción de forzar borrado permanente si falla la papelera
    # En modo Web (no nativo), por defecto sugerimos borrado permanente ya que la papelera del servidor no es accesible
    is_web_mode = not st.session_state.get("force_native_mode", True)
    default_force = is_web_mode
    
    force_delete = st.checkbox("Forzar borrado permanente (Irreversible)", value=default_force, help="Elimina los archivos definitivamente sin pasar por la papelera. Recomendado para modo Web/Servidor.")

    col_confirm, col_cancel = st.columns(2)
    with col_confirm:
        if st.button("🗑️ Sí, eliminar", type="primary", use_container_width=True):
            worker_eliminar_lista(results, force_delete=force_delete, silent_mode=False)
            
    with col_cancel:
        if st.button("Cancelar", use_container_width=True):
            st.session_state.active_action_dialog = None
            st.rerun()

@st.dialog("Comprimir en ZIP")
def dialogo_comprimir_zip():
    st.write("Comprimirá todos los elementos seleccionados en un único archivo ZIP.")
    
    if not st.session_state.get("search_results", []):
        st.error("No hay elementos seleccionados.")
        return
        
    st.info(f"Elementos a comprimir: {len(st.session_state.get('search_results', []))}")
    
    # Default name and location
    current_path = st.session_state.get("current_path", os.getcwd())
    default_name = f"archivo_{int(time.time())}.zip"
    
    zip_name = st.text_input("Nombre del archivo ZIP:", value=default_name)
    
    # Folder selector for destination
    target_path = render_path_selector(
        label="Carpeta Destino:",
        key="zip_dest_input",
        default_path=current_path
    )
    
    if st.button("🚀 Comprimir"):
        if not zip_name:
            st.warning("Ingresa un nombre para el archivo.")
            return
        if not target_path:
            st.warning("Selecciona una carpeta destino.")
            return
            
        full_zip_path = os.path.join(target_path, zip_name)
        if not full_zip_path.lower().endswith(".zip"):
            full_zip_path += ".zip"
            
        worker_comprimir_zip(st.session_state.get("search_results", []), full_zip_path, silent_mode=False)
        
    if st.button("❌ Cerrar", key="btn_close_zip"):
        st.session_state.active_action_dialog = None
        st.rerun()

@st.dialog("Comprimir Individualmente")
def dialogo_comprimir_individual():
    st.write("Cada carpeta/archivo seleccionado se comprimirá en su propio archivo ZIP en la misma ubicación.")
    st.warning("Esta acción creará múltiples archivos ZIP.")
    
    if not st.session_state.get("search_results", []):
        st.error("No hay elementos seleccionados.")
        return
        
    st.info(f"Elementos a procesar: {len(st.session_state.get('search_results', []))}")
    
    col_ok, col_cancel = st.columns([1, 1])
    with col_ok:
        if st.button("🚀 Comprimir Individualmente"):
             worker_comprimir_individual(st.session_state.get("search_results", []), silent_mode=False)
    
    with col_cancel:
        if st.button("❌ Cerrar", key="btn_close_ind"):
            st.session_state.active_action_dialog = None
            st.rerun()

# --- RENDER FUNCTION ---



def handle_zip_upload():
    """Maneja la carga de archivos ZIP para el modo web."""
    uploaded_file = st.file_uploader("📂 Cargar Entorno de Trabajo (ZIP)", type="zip", key="zip_uploader", help="Sube un archivo ZIP con las carpetas y archivos que deseas modificar.")
    
    if uploaded_file is not None:
        try:
            from gui_utils import extract_uploaded_zip
            final_path = extract_uploaded_zip(uploaded_file)
            
            if final_path:
                st.session_state.current_path = final_path
                st.session_state.path_input = final_path
                # st.session_state["path_input_widget"] = final_path # Evitar conflicto con el widget
                
                st.success(f"✅ Entorno cargado: {os.path.basename(final_path)}")
                st.rerun()
        except ImportError:
            st.error("Error importando utilidad de ZIP (gui_utils).")
        except Exception as e:
            st.error(f"Error al procesar ZIP: {e}")

def handle_zip_download(current_path):
    """Muestra botón de descarga para la carpeta actual."""
    if not os.path.exists(current_path):
        st.warning("Ruta actual no válida para descarga.")
        return
    
    # Check if directory is empty
    if not os.listdir(current_path):
        st.info("La carpeta está vacía.")
        return

    # Button to trigger compression
    # We use a unique key based on path modification time or just path to avoid constant re-zipping?
    # Actually, st.download_button needs the data upfront or a callback.
    # For large folders, generating ZIP on every render is bad.
    # We can use a callback or just generate it when clicked? No, download_button needs data.
    # Alternative: "Preparar Descarga" button -> Generates ZIP -> Shows Download Button.
    
    if st.button("📦 Preparar Descarga (ZIP)", key="btn_prep_download"):
        with st.spinner("Comprimiendo carpeta..."):
            try:
                mem_zip = io.BytesIO()
                with zipfile.ZipFile(mem_zip, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
                    # Calculate root len to make paths relative
                    # If we are in /tmp/work/FolderA, we want zip to contain FolderA/... or just contents?
                    # User said: "descargar la carpeta concervando como tenemos el modelo"
                    # Usually means preserving structure relative to current path.
                    
                    root_len = len(current_path) + 1
                    
                    for root, dirs, files in os.walk(current_path):
                        for file in files:
                            file_path = os.path.join(root, file)
                            archive_name = file_path[root_len:]
                            zf.write(file_path, archive_name)
                
                st.session_state["ready_to_download_zip"] = mem_zip.getvalue()
                st.session_state["ready_to_download_name"] = f"trabajo_procesado_{int(time.time())}.zip"
                st.success("✅ Archivo ZIP listo para descargar.")
            except Exception as e:
                st.error(f"Error al comprimir: {e}")

    if "ready_to_download_zip" in st.session_state:
        st.download_button(
            label="⬇️ Descargar ZIP",
            data=st.session_state["ready_to_download_zip"],
            file_name=st.session_state.get("ready_to_download_name", "trabajo.zip"),
            mime="application/zip",
            key="btn_download_zip_final"
        )

def render(container):
    """
    Renderiza el panel de búsqueda y gestión de archivos.
    """
    if "current_path" not in st.session_state:
         st.session_state.current_path = "" # Inicialmente vacío
    
    # Asegurar inicialización de otras variables de estado
    if "search_results" not in st.session_state:
        st.session_state.search_results = []
    if "app_config" not in st.session_state:
        st.session_state.app_config = {}

    # Layout Principal
    with container:
        st.markdown("## 🔍 Búsqueda y Gestión de Archivos")
        
        # --- SELECCIÓN DE RUTA ---
        # Usamos el selector estandarizado que maneja la sincronización y diálogo nativo
        current_path = render_path_selector(
            "Ruta a analizar:", 
            key="current_path", 
            default_path=st.session_state.get("current_path", os.getcwd()),
            help_text="Selecciona la carpeta base para realizar búsquedas y operaciones.",
            omit_checkbox=True
        )

        # Información de ayuda general
        # c_info, c_help = st.columns([0.8, 0.2]) # Ya integrado arriba
        
        
        # 2. Paneles de Criterios y Acciones

        c1, c2 = st.columns([1, 1])
        
        with c1:
            st.markdown('<div class="group-box"><div class="group-title-left">🔎 Criterios de búsqueda</div>', unsafe_allow_html=True)
            col_crit1, col_crit2 = st.columns(2)
            with col_crit1:
                st.selectbox("Buscar por:", ["extensión", "nombre", "todo"], key="search_by")
                st.selectbox("Tipo de elemento:", ["archivos", "carpetas"], key="item_type")
            with col_crit2:
                st.text_area("Patrones (uno por línea o separados por comas):", key="pattern", height=100, help="Ingresa múltiples términos para búsqueda masiva.")
                # Usar valor por defecto de la configuración
                default_excl = st.session_state.app_config.get("default_exclusion_patterns", "")
                st.text_input("Excluir:", value=default_excl, key="exclusion_pattern", help="Separar por comas. Excluye si coincide con el patrón.")
                st.checkbox("Incluir subcarpetas", value=True, key="subfolders")
                st.checkbox("Buscar solo carpetas vacías", value=False, key="search_empty_folders", help="Si se marca, ignora los patrones y busca carpetas sin contenido.")
            st.markdown('</div>', unsafe_allow_html=True)
                
        with c2:
            st.markdown('<div class="group-box"><div class="group-title-left">🛠️ Acción a realizar</div>', unsafe_allow_html=True)
            st.radio("Seleccione acción:", [
                "Copiar a carpeta", 
                "Mover a carpeta", 
                "Modificar nombre", 
                "Editar texto",
                "Comprimir en ZIP",
                "Comprimir individualmente"
            ], label_visibility="collapsed", key="action_radio")
            st.markdown('</div>', unsafe_allow_html=True)
        
        # 3. Tabla de Resultados
        search_results = st.session_state.get("search_results", [])
        num_found = len(search_results)
        st.markdown(f"##### 📄 Archivos encontrados ({num_found})")
        df_display = pd.DataFrame(search_results) if search_results else pd.DataFrame(columns=["Ruta completa", "Fecha"])
        st.dataframe(df_display, width=1000, height=250, hide_index=True)

        # 4. Barra de Botones Inferior
        st.write("") # Espaciador
        col_btns = st.columns(5, gap="small") # 5 columnas
        
        with col_btns[0]:
            if st.button("🔍 Buscar", use_container_width=True, help="Buscar archivos"):
                buscar_archivos_wrapper()
                st.rerun()
                
        with col_btns[1]:
            if st.button("▶️ Ejecutar", use_container_width=True, help="Ejecutar acción seleccionada"):
                action = st.session_state.get("action_radio")
                st.session_state.active_action_dialog = action
                st.rerun()

            # Manejo de diálogos persistentes
            active_dialog = st.session_state.get("active_action_dialog")
            
            if active_dialog == "Modificar nombre":
                dialogo_modificar_nombres()
            elif active_dialog == "Editar texto":
                dialogo_editar_texto()
            elif active_dialog == "Copiar a carpeta":
                dialogo_copiar_lista()
            elif active_dialog == "Mover a carpeta":
                dialogo_mover_lista()
            elif active_dialog == "Comprimir en ZIP":
                dialogo_comprimir_zip()
            elif active_dialog == "Comprimir individualmente":
                dialogo_comprimir_individual()
            elif active_dialog:
                # Caso fallback o error
                funcion_no_implementada(f"Acción: {active_dialog}")
                st.session_state.active_action_dialog = None
                
        with col_btns[2]:
            if st.button("🧹 Limpiar", use_container_width=True, help="Limpiar lista de resultados"):
                st.session_state.search_results = []
                st.rerun()
                
        with col_btns[3]:
            if st.button("🗑️ Eliminar", use_container_width=True, help="Eliminar archivos seleccionados"):
                dialogo_confirmar_eliminar()

        with col_btns[4]:
            if st.button("↩️ Deshacer", use_container_width=True, help="Revertir la última acción"):
                undo_last_action()
