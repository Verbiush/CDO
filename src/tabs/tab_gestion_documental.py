import streamlit as st
import pandas as pd
import os
import shutil
import requests
from datetime import datetime
import time
try:
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    Paragraph = None
    qn = None
import re
from io import BytesIO
import unicodedata
from PIL import Image, ImageDraw, ImageFont

import zipfile
import importlib
import base64
import urllib.parse
import time

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.service import Service
    from webdriver_manager.chrome import ChromeDriverManager
except ImportError:
    webdriver = None

try:
    from src import db_gestion
    from src.gui_utils import abrir_dialogo_carpeta_nativo, update_path_key as update_path_key_gui, render_path_selector, render_download_button
    from src.agent_client import send_command, wait_for_result
except ImportError:
    import db_gestion
    from gui_utils import abrir_dialogo_carpeta_nativo, update_path_key as update_path_key_gui, render_path_selector, render_download_button
    try:
        from agent_client import send_command, wait_for_result
    except ImportError:
        send_command = None
        wait_for_result = None

# Force reload to ensure latest DB logic is used
importlib.reload(db_gestion)

def get_google_font(font_name, size):
    """Download and load a Google Font."""
    font_dir = "assets/fonts"
    os.makedirs(font_dir, exist_ok=True)
    
    font_urls = {
        "Pacifico": "https://github.com/google/fonts/raw/main/ofl/pacifico/Pacifico-Regular.ttf",
    }
    
    font_path = os.path.join(font_dir, f"{font_name}.ttf")
    
    # Download if not exists
    if font_name == "My Ugly Handwriting":
        if not os.path.exists(font_path):
            try:
                # Direct download from dafont (simulated as browser)
                url = "https://dl.dafont.com/dl/?f=my_ugly_handwriting"
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
                }
                r = requests.get(url, headers=headers, timeout=15)
                if r.status_code == 200:
                    with zipfile.ZipFile(BytesIO(r.content)) as z:
                        for name in z.namelist():
                            if name.lower().endswith(".otf") or name.lower().endswith(".ttf"):
                                with open(font_path, "wb") as f:
                                    f.write(z.read(name))
                                break
            except Exception as e:
                print(f"Error downloading My Ugly Handwriting: {e}")
                
    elif font_name in font_urls and not os.path.exists(font_path):
        try:
            r = requests.get(font_urls[font_name])
            with open(font_path, "wb") as f:
                f.write(r.content)
        except Exception as e:
            print(f"Error downloading font: {e}")
            
    try:
        return ImageFont.truetype(font_path, size)
    except:
        try:
            return ImageFont.truetype("arial.ttf", size)
        except:
            return ImageFont.load_default()

# Helper Functions for DOCX Processing
def fmt_date(val):
    s = str(val)
    if not val or s.lower() == 'nat': return ""
    return s.split(" ")[0]

def replace_text_in_element(paragraph, mapping):
    """
    Replaces placeholders in a paragraph object using regex for flexibility.
    Returns the number of replacements made.
    """
    if not paragraph.text.strip():
        return 0

    # Optimization: Quick check for '{' or '<'
    text_has_braces = "{" in paragraph.text
    text_has_chevrons = "<<" in paragraph.text or "«" in paragraph.text

    if not (text_has_braces or text_has_chevrons):
        return 0

    original_text = paragraph.text
    current_text = original_text
    count = 0

    for key, val in mapping.items():
        # Clean key for variations (e.g. "Nombre Completo" -> "NombreCompleto")
        key_clean = key.replace(" ", "")

        # Build regex patterns
        # 1. Exact with spaces: { Key }
        p1 = r"\{\s*" + re.escape(key) + r"\s*\}"
        # 2. No spaces in key: { KeyClean }
        p2 = r"\{\s*" + re.escape(key_clean) + r"\s*\}"
        # 3. Underscores: { Key_With_Underscores }
        p3 = r"\{\s*" + re.escape(key.replace(" ", "_")) + r"\s*\}"

        # Chevrons variations
        c1 = r"(?:«|<<)\s*" + re.escape(key) + r"\s*(?:»|>>)"
        c2 = r"(?:«|<<)\s*" + re.escape(key_clean) + r"\s*(?:»|>>)"
        c3 = r"(?:«|<<)\s*" + re.escape(key.replace(" ", "_")) + r"\s*(?:»|>>)"

        # Combine patterns
        patterns = [p1, p2, p3] if text_has_braces else []
        if text_has_chevrons:
            patterns.extend([c1, c2, c3])

        for pat in patterns:
            if re.search(pat, current_text, re.IGNORECASE):
                current_text = re.sub(pat, str(val), current_text, flags=re.IGNORECASE)
                count += 1

    if count > 0 and current_text != original_text:
        paragraph.text = current_text
        return count
    return 0

def iter_all_paragraphs(doc_obj):
    # 1. Main Body Paragraphs
    yield from doc_obj.paragraphs
    
    # 2. Tables in Body
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    
    # 3. Text Boxes in Body (via XML)
    if doc_obj.element.body is not None:
        for txbx in doc_obj.element.body.iter(qn('w:txbxContent')):
            for p_element in txbx.iter(qn('w:p')):
                yield Paragraph(p_element, doc_obj)

    # 4. Headers and Footers
    for section in doc_obj.sections:
        headers = [section.header, section.first_page_header, section.even_page_header]
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        
        for header in headers:
            if header and not header.is_linked_to_previous:
                yield from header.paragraphs
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs
                for txbx in header._element.iter(qn('w:txbxContent')):
                    for p_element in txbx.iter(qn('w:p')):
                        yield Paragraph(p_element, header)

        for footer in footers:
            if footer and not footer.is_linked_to_previous:
                yield from footer.paragraphs
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs
                for txbx in footer._element.iter(qn('w:txbxContent')):
                    for p_element in txbx.iter(qn('w:p')):
                        yield Paragraph(p_element, footer)

def generate_signature_image(text, font_name="Pacifico", size=70, width=500, height=200):
    """Generate a signature image from text."""
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    
    font = get_google_font(font_name, size)
    
    # Center text
    try:
        # getbbox is cleaner in newer Pillow, textbbox in even newer
        left, top, right, bottom = d.textbbox((0, 0), text, font=font)
        text_w = right - left
        text_h = bottom - top
    except:
        # Fallback for older Pillow
        text_w, text_h = d.textsize(text, font=font)
        
    x = (width - text_w) / 2
    y = (height - text_h) / 2
    
    d.text((x, y), text, font=font, fill=(0, 0, 0))
    return img

def normalize_text(text):
    """Normalize text to remove accents and convert to lowercase."""
    if not isinstance(text, str):
        return str(text)
    return ''.join(c for c in unicodedata.normalize('NFD', text) if unicodedata.category(c) != 'Mn').lower()

def match_signature_file(search_name, filenames):
    """
    Find a matching filename for a given name.
    Checks if all parts of the search_name are present in the filename.
    """
    if not search_name:
        return None
        
    search_parts = normalize_text(search_name).split()
    
    for fname in filenames:
        norm_fname = normalize_text(fname)
        # Check if all parts of the name are in the filename
        if all(part in norm_fname for part in search_parts):
            return fname
    return None

def resolve_unique_paths(records, folder_structure):
    """
    Generates unique folder paths for each record to avoid collisions.
    If multiple records map to the same path, appends a consecutive suffix (_1, _2, etc.)
    sorted by 'no_factura' to ensure determinism.
    
    Returns:
        dict: {record_id: unique_relative_path}
    """
    # 1. Group records by their raw path
    path_groups = {}
    
    for record in records:
        safe_record = {k: str(v).replace("/", "-").replace("\\", "-").strip() for k, v in record.items()}
        try:
            raw_path = folder_structure.format(**safe_record).upper()
        except KeyError:
            # Fallback if key missing
            raw_path = f"ERROR_KEY_MISSING_{record.get('id')}"
            
        if raw_path not in path_groups:
            path_groups[raw_path] = []
        path_groups[raw_path].append(record)
    
    # 2. Resolve collisions
    record_id_to_path = {}
    
    for raw_path, group in path_groups.items():
        if len(group) == 1:
            record_id_to_path[group[0]['id']] = raw_path
        else:
            # Sort by invoice number (or ID if invoice is same/missing) to be deterministic
            # We assume 'no_factura' exists, else 'id'
            group.sort(key=lambda x: (str(x.get('no_factura', '')), x.get('id', 0)))
            
            for i, record in enumerate(group):
                # Append consecutive: Path_1, Path_2, etc.
                unique_path = f"{raw_path}_{i+1}"
                record_id_to_path[record['id']] = unique_path
                
    return record_id_to_path

def worker_descargar_historias_ovida(records, download_path, resolved_paths, return_zip=False):
    """
    Descarga historias clínicas de OVIDA para los registros dados.
    Usa Selenium para inicio de sesión y descarga.
    """
    if webdriver is None:
        return "Error: Selenium no está instalado o no se puede cargar."

    if not return_zip and not os.path.isdir(download_path):
        return "Error: Carpeta de descarga inválida."

    is_native = st.session_state.get("force_native_mode", True)
    if not is_native:
        return "⚠️ Esta funcionalidad requiere iniciar sesión manualmente en OVIDA, lo cual no es posible en el modo Web. Por favor, use la aplicación de escritorio (Modo Nativo) para descargar historias clínicas."

    driver = None
    try:
        options = webdriver.ChromeOptions()
        # Temp dir for Selenium downloads (even if we use printToPDF, it needs a profile dir)
        temp_dir = download_path if not return_zip else os.path.join(os.getcwd(), "temp_selenium_ovida")
        if return_zip:
            os.makedirs(temp_dir, exist_ok=True)
            
        prefs = {
            "download.default_directory": temp_dir,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True
        }
        options.add_experimental_option("prefs", prefs)
        
        # Add stability options
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--start-maximized")

        # Open visible browser for login
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
        
        st.warning("⚠️ Se abrió una ventana de Chrome. INICIE SESIÓN en OVIDA manualmente. El proceso continuará automáticamente cuando detecte el ingreso.")
        
        # Wait for login (detect change to main page or timeout)
        timeout = 300 
        start_time = time.time()
        while time.time() - start_time < timeout:
             try:
                 # Check if URL changed from login page
                 if "iniciando.php" not in driver.current_url and "login" not in driver.current_url.lower():
                     break
             except:
                 pass
             time.sleep(1)
        
        if time.time() - start_time >= timeout:
            driver.quit()
            if return_zip:
                try: shutil.rmtree(temp_dir)
                except: pass
            return "Error: Tiempo de espera de inicio de sesión agotado."
            
        # Check logged in state more robustly
        logged_in = False
        start_time = time.time()
        while time.time() - start_time < timeout:
            try:
                if "App/Vistas" in driver.current_url:
                    logged_in = True
                    break
            except: pass
            time.sleep(2)
            
        if not logged_in:
            if return_zip:
                try: shutil.rmtree(temp_dir)
                except: pass
            return "Error: No se detectó inicio de sesión en 5 minutos."

        st.info("Inicio de sesión detectado. Comenzando descargas...")

        descargados = 0
        errores = 0
        conflictos = 0
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        total = len(records)
        
        # Prepare ZIP if needed
        mem_zip = None
        zf = None
        if return_zip:
            mem_zip = BytesIO()
            zf = zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED)
        
        for i, record in enumerate(records):
            progress_bar.progress((i + 1) / total)
            
            try:
                estudio = str(record.get('nro_estudio', '')).strip()
                # Clean estudio (remove .0 if float)
                if estudio.endswith(".0"): estudio = estudio[:-2]
                
                if not estudio or estudio == "nan" or estudio == "None":
                    errores += 1
                    continue
                    
                # Parse dates
                try:
                    f_ing_raw = record.get('fecha_ingreso', '')
                    f_egr_raw = record.get('fecha_salida', '')
                    
                    if not f_ing_raw or not f_egr_raw:
                        errores += 1
                        continue
                        
                    f_ing = pd.to_datetime(f_ing_raw).strftime('%Y/%m/%d')
                    f_egr = pd.to_datetime(f_egr_raw).strftime('%Y/%m/%d')
                except:
                    errores += 1
                    continue
                
                # Determine folder
                rel_path = resolved_paths.get(record['id'])
                if not rel_path:
                    # Fallback
                    continue 

                if not return_zip:
                    dest_dir = os.path.join(download_path, rel_path)
                    os.makedirs(dest_dir, exist_ok=True)
                    final_path = os.path.join(dest_dir, f"HC_{estudio}.pdf")
                    
                    if os.path.exists(final_path):
                        conflictos += 1
                        continue
                else:
                    pass
                    
                status_text.text(f"Descargando Estudio: {estudio}")

                # URL construction
                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                params = {
                    'estudio': estudio, 'fecha_inicio': f_ing, 'fecha_fin': f_egr,
                    'verHC': 1, 'verEvo': 1, 'verPar': 1, 'ImprimirOrdenamiento': 1,
                    'ImprimirNotasPcte': 0, 'ImprimirSolOrdenesExt': 1, 'ImprimirGraficasHC': 1,
                    'ImprimirFormatos': 1, 'ImprimirRegistroAdmon': 1, 'ImprimirNovedad': 0,
                    'ImprimirRecomendaciones': 0, 'ImprimirDescripcionQX': 0, 'ImprimirNotasEnfermeria': 1,
                    'ImprimirSignosVitales': 0, 'ImprimirLog': 0, 'ImprimirEpicrisisSinHC': 0
                }
                full_url = f"{base_url}?{urllib.parse.urlencode(params)}"
                
                driver.get(full_url)
                time.sleep(2) # Wait for render
                
                pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {
                    "landscape": False, "printBackground": True,
                    "paperWidth": 8.5, "paperHeight": 11,
                    "marginTop": 0.4, "marginBottom": 0.4, "marginLeft": 0.4, "marginRight": 0.4
                })
                
                pdf_data = base64.b64decode(pdf_b64['data'])
                
                if return_zip:
                    # Add to ZIP in memory
                    zip_entry_name = os.path.join(rel_path, f"HC_{estudio}.pdf")
                    zf.writestr(zip_entry_name, pdf_data)
                else:
                    # Write to disk
                    with open(final_path, 'wb') as f:
                        f.write(pdf_data)
                
                descargados += 1
                
            except Exception as e:
                errores += 1
                
        # Clean up temp dir if we created one
        if return_zip and os.path.exists(temp_dir):
            try: shutil.rmtree(temp_dir)
            except: pass

        if return_zip:
            zf.close()
            return {
                "files": [{
                    "name": f"Historias_OVIDA_{int(time.time())}.zip",
                    "data": mem_zip.getvalue(),
                    "label": "📦 Descargar Historias OVIDA (ZIP)"
                }],
                "message": f"Finalizado. Descargados: {descargados}, Errores: {errores}, Conflictos: {conflictos}."
            }
        else:
            return f"Finalizado. Descargados: {descargados}, Errores: {errores}, Conflictos: {conflictos}."

    except Exception as e:
        return f"Error crítico: {e}"
    finally:
        if driver: driver.quit()

@st.cache_data(show_spinner=False, max_entries=10)
def _get_excel_preview(file_bytes, sheet_name=0, nrows=None):
    import pandas as pd
    import io
    if nrows:
        return pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name, nrows=nrows)
    return pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_name)

def render():
    st.markdown("## 📂 Gestión Documental (Relacional)")
    
    # --- GLOBAL CONFIG & SYNC ---
    st.info("ℹ️ Ruta Inicial: Define el directorio de trabajo para todas las pestañas.")
    
    # 1. Global Path Selector - Single source of truth
    # We use a unique key 'gd_working_path' to avoid conflicts with other tabs using 'current_path'
    # But we sync it bi-directionally with the global 'current_path'
    
    global_current_path = st.session_state.get("current_path", os.getcwd())
    if not global_current_path: global_current_path = os.getcwd()

    current_main_path = render_path_selector(
        label="Carpeta de Trabajo (Ruta Inicial)",
        key="gd_working_path_selector", 
        default_path=global_current_path,
        omit_checkbox=False
    )
    
    # Update global path if this one changes
    if current_main_path and current_main_path != global_current_path:
        st.session_state.current_path = current_main_path

    # 2. Sync Logic - Propagate to all sub-states
    last_synced = st.session_state.get("gd_last_synced_path")
    
    if current_main_path and current_main_path != last_synced:
        st.session_state.gd_base_path = current_main_path
        
        # Tab Estructura
        st.session_state.gd_base_path_struct = current_main_path
        st.session_state.input_gd_base_path_struct = current_main_path
        
        # Tab Contenido
        st.session_state.input_base_path_content = current_main_path
        st.session_state.input_input_base_path_content = current_main_path
        
        # Tab Organizar
        st.session_state.gd_source_path = current_main_path
        st.session_state.input_gd_source_path = current_main_path
        
        st.session_state.gd_source_path_mov = current_main_path 
        st.session_state.input_gd_source_path_mov = current_main_path
        
        st.session_state.input_base_path_struct_mov = current_main_path
        st.session_state.gd_dest_path_mov_selector = current_main_path
        st.session_state.input_gd_dest_path_mov_selector = current_main_path

        st.session_state.gd_last_synced_path = current_main_path
    
    # Updated Tabs Structure
    tab_import, tab_view, tab_structure, tab_organize, tab_content = st.tabs([
        "📥 Importar Excel", 
        "👁️ Ver Registros", 
        "📁 Crear Carpetas", 
        "🚀 Organizar (Renombrar/Mover)",
        "📝 Generar Docs/Firmas"
    ])
    
    # --- TAB IMPORTAR ---
    with tab_import:
        st.subheader("📥 Importar Datos (Excel)")
        uploaded_file = st.file_uploader("Subir Excel", type=["xlsx", "xls"], key="uploader_gestion")
        
        if uploaded_file:
            try:
                file_bytes = uploaded_file.getvalue()
                df = _get_excel_preview(file_bytes)
                # Normalize columns
                df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
                
                if "no_doc" in df.columns and "nombre_completo" in df.columns:
                    st.dataframe(df.head())
                    if st.button("💾 Guardar en Base de Datos", key="btn_save_db"):
                        records = df.to_dict("records")
                        for i, r in enumerate(records):
                            if "id" not in r: r["id"] = i + 1
                        
                        if hasattr(db_gestion, "save_records"):
                            db_gestion.save_records(records)
                        else:
                            st.session_state.db_records = records
                            
                        st.success(f"✅ {len(records)} registros guardados.")
                        st.rerun()
                else:
                    st.error("Faltan columnas requeridas: no_doc, nombre_completo")
            except Exception as e:
                st.error(f"Error al leer Excel: {e}")

    # Load records
    records = []
    if hasattr(db_gestion, "get_records"):
        records = db_gestion.get_records()
    else:
        records = st.session_state.get("db_records", [])

    # --- TAB VER REGISTROS ---
    with tab_view:
        st.subheader("👁️ Visualizar y Filtrar Registros")
        
        if not records:
            st.info("No hay registros cargados.")
        else:
            # Filters
            col_filter1, col_filter2 = st.columns(2)
            with col_filter1:
                eps_list = sorted(list(set([str(r.get("eps", "")) for r in records if r.get("eps")])))
                selected_eps = st.multiselect("Filtrar por EPS", eps_list, key="filter_eps_view")
                
            with col_filter2:
                regimen_list = sorted(list(set([str(r.get("regimen", "")) for r in records if r.get("regimen")])))
                selected_regimen = st.multiselect("Filtrar por Régimen", regimen_list, key="filter_regimen_view")
            
            # Filter logic
            filtered_view = records
            if selected_eps:
                filtered_view = [r for r in filtered_view if r.get("eps") in selected_eps]
            if selected_regimen:
                filtered_view = [r for r in filtered_view if r.get("regimen") in selected_regimen]
                
            st.dataframe(pd.DataFrame(filtered_view))
            st.caption(f"Mostrando {len(filtered_view)} de {len(records)} registros.")

    def get_filtered_records(records, key_suffix=""):
        st.markdown("#### 🔍 Filtros de Selección")
        
        # EPS
        eps_opts = sorted(list(set([str(r.get("eps", "")) for r in records])))
        selected_eps = st.multiselect("EPS", eps_opts, key=f"filter_eps_{key_suffix}")
        
        # Regimen
        reg_opts = sorted(list(set([str(r.get("regimen", "")) for r in records])))
        selected_regimen = st.multiselect("Régimen", reg_opts, key=f"filter_reg_{key_suffix}")
        
        # Date Range (Ingreso)
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            date_start = st.date_input(
                "Fecha Inicio (Ingreso)", 
                value=None, 
                key=f"filter_d1_{key_suffix}"
            )
        with col_d2:
            date_end = st.date_input(
                "Fecha Fin (Ingreso)", 
                value=None, 
                key=f"filter_d2_{key_suffix}"
            )

        # Date Range (Factura)
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            date_fact_start = st.date_input(
                "Fecha Inicio (Factura)", 
                value=None, 
                key=f"filter_f1_{key_suffix}"
            )
        with col_f2:
            date_fact_end = st.date_input(
                "Fecha Fin (Factura)", 
                value=None, 
                key=f"filter_f2_{key_suffix}"
            )
            
        filtered = []
        for r in records:
            # EPS Filter
            if selected_eps and r.get("eps") not in selected_eps:
                continue
            
            # Regimen Filter
            if selected_regimen and r.get("regimen") not in selected_regimen:
                continue
            
            # Date Filter (Ingreso)
            if date_start and date_end:
                r_date_str = str(r.get("fecha_ingreso", "")).split(" ")[0] 
                try:
                    r_date = None
                    for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d"]:
                        try:
                            r_date = datetime.strptime(r_date_str, fmt).date()
                            break
                        except:
                            pass
                    
                    if r_date:
                        if not (date_start <= r_date <= date_end):
                            continue
                except:
                    pass

            # Date Filter (Factura)
            if date_fact_start and date_fact_end:
                r_date_str = str(r.get("fecha_factura", "")).split(" ")[0] 
                try:
                    r_date = None
                    for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d"]:
                        try:
                            r_date = datetime.strptime(r_date_str, fmt).date()
                            break
                        except:
                            pass
                    
                    if r_date:
                        if not (date_fact_start <= r_date <= date_fact_end):
                            continue
                except:
                    pass
            
            filtered.append(r)
        return filtered

    # --- TAB ESTRUCTURA (Former Subtab 1) ---
    with tab_structure:
        st.subheader("📁 Estructura de Carpetas")
        
        # Common Config - ALWAYS VISIBLE
        st.markdown("### Configuración de Rutas")
        
        default_hardcoded = os.path.join(os.path.expanduser("~"), "Documents", "GestionDocumental")
        
        # Initialize fallback if somehow missing
        if "gd_base_path" not in st.session_state:
             st.session_state.gd_base_path = default_hardcoded
             st.session_state.input_base_path_struct = default_hardcoded
        
        is_native = st.session_state.get("force_native_mode", True)

        # Selector de Ruta Estandarizado
        # Defaults to the synced global path
        base_path = render_path_selector(
            label="Carpeta Raíz / ZIP Base (Modo Web)",
            key="gd_base_path_struct",
            default_path=st.session_state.get("gd_base_path", st.session_state.get("current_path", os.getcwd())),
            omit_checkbox=False,
            help_text="Seleccione la carpeta donde se crearán las subcarpetas. En Modo Web, suba un ZIP (puede estar vacío) para usarlo como base."
        )
        
        # In Web Mode, if no ZIP uploaded, create a temp dir to start fresh
        if not is_native and (not base_path or not os.path.exists(base_path)):
             temp_base = os.path.join(os.getcwd(), "temp_downloads", f"struct_{int(time.time())}")
             if "temp_struct_path" not in st.session_state:
                 os.makedirs(temp_base, exist_ok=True)
                 st.session_state.temp_struct_path = temp_base
             base_path = st.session_state.temp_struct_path
             st.info(f"Modo Web: Usando directorio temporal para generar estructura: {base_path}")

        st.session_state.input_base_path_struct = base_path

        st.caption("Patrón de Carpetas: Use `{eps}`, `{fecha_ingreso}`, `{nombre_completo}`, etc.")
        folder_structure = st.text_input("Patrón de Carpetas", value="{eps}/{no_factura} - {nombre_completo}", key="input_folder_structure_struct")
        
        # --- NEW LOGIC: ALWAYS SHOW ACTIONS ---
        st.divider()
        st.markdown("### Acciones")
        do_create_folders = st.checkbox("Activar Creación de Carpetas", value=True, help="Crea la estructura de carpetas.")
        
        disabled_struct = False
        filtered_records = []
        
        if not records:
             st.warning("⚠️ Para habilitar las acciones, primero debe importar datos en la pestaña 'Importar Excel'.")
             disabled_struct = True
        else:
             filtered_records = get_filtered_records(records, key_suffix="struct")
             st.info(f"Registros seleccionados: {len(filtered_records)} de {len(records)}")
             if not filtered_records:
                 st.warning("No hay registros que coincidan con los filtros.")
                 disabled_struct = True

        if st.button("🚀 Crear Carpetas", type="primary", disabled=disabled_struct):
            count_created = 0
            progress_bar = st.progress(0)
            
            # --- Resolve unique paths ---
            resolved_paths = resolve_unique_paths(filtered_records, folder_structure)
            
            paths_to_create = []
            for record in filtered_records:
                try:
                    rel_path = resolved_paths.get(record['id'])
                    full_path = os.path.join(base_path, rel_path)
                    paths_to_create.append(full_path)
                except Exception as e:
                    print(f"Error resolving path: {e}")

            if do_create_folders and paths_to_create:
                if is_native:
                    try:
                        from src.agent_client import send_command, wait_for_result
                        username = st.session_state.get("username", "default")
                        
                        with st.spinner("Enviando comando de creación de carpetas al Agente Local..."):
                            task_id = send_command(username, "create_folders", {
                                "folders": paths_to_create
                            })
                            
                            if task_id:
                                res = wait_for_result(task_id, timeout=600)
                                if "error" in res:
                                    st.error(f"Error del Agente: {res['error']}")
                                else:
                                    r_data = res
                                    count_created = r_data.get("count", 0)
                                    errs = r_data.get("errors", [])
                                    if errs:
                                        with st.expander("Ver Errores Creación"):
                                            for e in errs: st.write(e)
                            else:
                                st.error("No se pudo conectar con el Agente Local.")
                    except Exception as e:
                        st.error(f"Error ejecutando en agente: {e}")
                else:
                    for i, full_path in enumerate(paths_to_create):
                        try:
                            os.makedirs(full_path, exist_ok=True)
                            count_created += 1
                        except Exception as e:
                            print(f"Error creating folder: {e}")
                        progress_bar.progress((i + 1) / len(paths_to_create))
            
            st.success(f"✅ Proceso finalizado. Carpetas Creadas: {count_created}")
            
            # Save state for download button persistence
            st.session_state.last_created_folders_path = base_path
            st.session_state.last_created_folders_count = count_created
            
        # --- DOWNLOAD BUTTON (WEB MODE - PERSISTENT) ---
        if "last_created_folders_path" in st.session_state and st.session_state.last_created_folders_path:
             if os.path.exists(st.session_state.last_created_folders_path):
                 st.info(f"📂 Última creación: {st.session_state.get('last_created_folders_count', 0)} carpetas listas para descargar.")
#                  render_download_button(st.session_state.last_created_folders_path, "dl_struct", "📦 Descargar Estructura (ZIP)", cleanup=not is_native)
             else:
                 if "last_created_folders_path" in st.session_state:
                     del st.session_state.last_created_folders_path

    # --- TAB ORGANIZAR (Renombrar/Mover) ---
    with tab_organize:
        st.subheader("🚀 Organizar Archivos (Renombrar y Mover)")
        st.caption("Utilice los datos de la base de datos para renombrar y organizar archivos físicos.")
        
        # Prepare available keys globally for this tab to avoid scope issues
        available_keys = list(records[0].keys()) if records else []
        default_keys = ["no_doc", "nombre_completo", "no_factura", "nro_estudio"]
        sorted_keys = sorted(list(set(available_keys) - set(default_keys)))
        display_keys = [k for k in default_keys if k in available_keys] + sorted_keys

        if "gd_source_path" not in st.session_state:
             if "current_path" in st.session_state:
                 st.session_state.gd_source_path = st.session_state.current_path
             else:
                 st.session_state.gd_source_path = os.path.join(os.path.expanduser("~"), "Documents")
             
        # Determinar ruta a usar
        current_global_path = st.session_state.get("current_path", os.getcwd())
        if not current_global_path: current_global_path = os.getcwd()

        source_path = render_path_selector(
            label="Carpeta Origen de Archivos",
            key="gd_source_path",
            default_path=current_global_path,
            omit_checkbox=False
        )

        # --- EXPLORADOR DE ARCHIVOS (AGENT ENABLED) ---
        with st.expander("📂 Explorador de Archivos (Ver contenido)", expanded=True):
            if source_path:
                st.markdown(f"**Ruta actual:** `{source_path}`")
                
                # State for file list
                if "gd_file_list" not in st.session_state:
                    st.session_state.gd_file_list = []
                
                col_btn_list, col_info_list = st.columns([1, 3])
                with col_btn_list:
                    if st.button("🔄 Listar Archivos", key="btn_list_files_agent"):
                        st.session_state.gd_file_list = [] # Clear previous
                        if st.session_state.get("force_native_mode", True):
                            try:
                                from src.agent_client import list_files
                                username = st.session_state.get("username", "default")
                                with st.spinner("Solicitando lista de archivos al Agente..."):
                                    files = list_files(username, source_path)
                                    if files is not None:
                                        st.session_state.gd_file_list = files
                                    else:
                                        st.error("No se pudieron obtener archivos o la carpeta está vacía.")
                            except Exception as e:
                                st.error(f"Error comunicando con agente: {e}")
                        else:
                            # Web mode fallback (local server)
                            try:
                                if os.path.exists(source_path) and os.path.isdir(source_path):
                                    items = []
                                    with os.scandir(source_path) as it:
                                        for entry in it:
                                            items.append({
                                                "name": entry.name,
                                                "is_dir": entry.is_dir(),
                                                "size": entry.stat().st_size if not entry.is_dir() else 0,
                                                "mtime": entry.stat().st_mtime
                                            })
                                    st.session_state.gd_file_list = items
                            except Exception as e:
                                st.error(f"Error local: {e}")

                with col_info_list:
                    if st.session_state.gd_file_list:
                        st.caption(f"Se encontraron {len(st.session_state.gd_file_list)} elementos.")
                
                # Display Table
                if st.session_state.gd_file_list:
                    # Format for display
                    disp_data = []
                    for f in st.session_state.gd_file_list:
                        t_type = "📁 Carpeta" if f["is_dir"] else "📄 Archivo"
                        t_size = f"{f['size']/1024:.1f} KB" if not f["is_dir"] else "-"
                        t_date = datetime.fromtimestamp(f["mtime"]).strftime('%Y-%m-%d %H:%M')
                        disp_data.append({
                            "Nombre": f["name"],
                            "Tipo": t_type,
                            "Tamaño": t_size,
                            "Modificación": t_date
                        })
                    
                    st.dataframe(pd.DataFrame(disp_data), use_container_width=True, hide_index=True)
            else:
                st.info("Seleccione una carpeta para ver su contenido.")
            
        if not records:
             st.warning("⚠️ No hay registros en la base de datos. Importe un Excel primero.")
        else:
            st.divider()
            col_ren, col_mov = st.columns(2)
            
            # --- RENOMBRAR CARPETAS ---
            with col_ren:
                st.markdown("#### 1. Renombrar Carpetas")
                st.info("Renombra carpetas en el origen usando datos de la BD.")
                
                col_cur_name = st.selectbox("Columna Nombre Actual (Carpeta)", display_keys, key="sel_col_cur_name_folder")
                col_new_name = st.selectbox("Columna Nuevo Nombre (Carpeta)", display_keys, key="sel_col_new_name_folder")
                
                if st.button("Ejecutar Renombrado Carpetas", key="btn_run_ren_folder"):
                    count_ren = 0
                    errors_ren = 0
                    
                    map_ren = {}
                    for r in records:
                        curr = str(r.get(col_cur_name, "")).strip()
                        new_n = str(r.get(col_new_name, "")).strip()
                        if curr and new_n:
                            new_n = "".join([c for c in new_n if c.isalnum() or c in (' ', '-', '_', '.')]).strip()
                            map_ren[curr] = new_n
                            
                    if is_native:
                        try:
                            from src.agent_client import send_command, wait_for_result
                            username = st.session_state.get("username", "default")
                            
                            with st.spinner("Enviando comando al Agente Local..."):
                                task_id = send_command(username, "rename_folders_mapped", {
                                    "path": source_path,
                                    "mapping": map_ren
                                })
                                
                                if task_id:
                                    res = wait_for_result(task_id, timeout=600)
                                    if "error" in res:
                                        st.error(f"Error del Agente: {res['error']}")
                                    else:
                                        # Direct result usage, no nested 'result' key
                                        r_data = res 
                                        count_ren = r_data.get("count", 0)
                                        errs = r_data.get("errors", [])
                                        errors_ren = len(errs)
                                        st.success(f"Carpetas Renombradas: {count_ren}. Errores: {errors_ren}")
                                        if errs:
                                            with st.expander("Ver Errores"):
                                                for e in errs: st.write(e)
                                else:
                                    st.error("No se pudo conectar con el Agente Local.")
                        except Exception as e:
                            st.error(f"Error ejecutando en agente: {e}")
                    else:
                        try:
                            dirs = [d for d in os.listdir(source_path) if os.path.isdir(os.path.join(source_path, d))]
                            progress_bar = st.progress(0)
                            
                            for i, dirname in enumerate(dirs):
                                dir_path = os.path.join(source_path, dirname)
                                matched_new_name = None
                                
                                if dirname in map_ren:
                                    matched_new_name = map_ren[dirname]
                                else:
                                    for curr_val, new_val in map_ren.items():
                                        if curr_val in dirname:
                                            matched_new_name = new_val
                                            break
                                
                                if matched_new_name and matched_new_name != dirname:
                                    try:
                                        new_path = os.path.join(source_path, matched_new_name)
                                        if os.path.exists(new_path):
                                            print(f"Skipping rename {dirname} -> {matched_new_name}: Dest exists")
                                            errors_ren += 1
                                        else:
                                            os.rename(dir_path, new_path)
                                            count_ren += 1
                                    except Exception as e:
                                        errors_ren += 1
                                        print(f"Error renaming folder {dirname}: {e}")
                                        
                                progress_bar.progress((i + 1) / len(dirs))
                                
                            st.success(f"Carpetas Renombradas: {count_ren}. Errores/Omis: {errors_ren}")
#                             render_download_button(source_path, "dl_ren", "📦 Descargar Carpetas Renombradas (ZIP)", cleanup=not is_native)
                        except Exception as e:
                            st.error(f"Error accediendo a la carpeta: {e}")

            # --- MOVER ---
            with col_mov:
                st.markdown("#### 2. Mover a Estructura")
                st.info("Mueve archivos/carpetas de Origen -> Carpeta Destino.")
                
                source_path_mov = render_path_selector(
                    label="Origen (Mover)",
                    key="gd_source_path_mov",
                    omit_checkbox=False
                )

                col_src_match = st.selectbox("Columna Coincidencia Origen", display_keys, key="sel_col_src_match_mov")
                col_dst_name = st.selectbox("Columna Nombre Destino", display_keys, key="sel_col_dst_name_mov")
                
                if "input_base_path_struct_mov" not in st.session_state:
                     st.session_state.input_base_path_struct_mov = st.session_state.get("input_base_path_struct", st.session_state.gd_base_path)

                base_dest = render_path_selector(
                    label="Destino Base",
                    key="gd_dest_path_mov_selector",
                    default_path=current_global_path,
                    omit_checkbox=False
                )
                
                if not is_native and (not base_dest or not os.path.exists(base_dest)):
                     temp_mov = os.path.join(os.getcwd(), "temp_downloads", f"mov_{int(time.time())}")
                     if "temp_mov_path" not in st.session_state:
                         os.makedirs(temp_mov, exist_ok=True)
                         st.session_state.temp_mov_path = temp_mov
                     base_dest = st.session_state.temp_mov_path
                     st.info(f"Modo Web: Usando directorio temporal para destino: {base_dest}")

                st.session_state.input_base_path_struct_mov = base_dest
                
                if st.button("Ejecutar Movimiento", key="btn_run_mov"):
                    count_mov = 0
                    errors_mov = 0
                    
                    map_mov = {}
                    for r in records:
                        src = str(r.get(col_src_match, "")).strip()
                        dst = str(r.get(col_dst_name, "")).strip()
                        if src and dst:
                            dst = "".join([c for c in dst if c.isalnum() or c in (' ', '-', '_', '.')]).strip()
                            map_mov[src] = dst

                    if is_native:
                        try:
                            from src.agent_client import send_command, wait_for_result
                            username = st.session_state.get("username", "default")
                            
                            with st.spinner("Enviando comando al Agente Local..."):
                                task_id = send_command(username, "organize_files_mapped", {
                                    "source_path": source_path_mov,
                                    "dest_path": base_dest,
                                    "mapping": map_mov
                                })
                                
                                if task_id:
                                    res = wait_for_result(task_id, timeout=600)
                                    if "error" in res:
                                        st.error(f"Error del Agente: {res['error']}")
                                    else:
                                        # Direct result usage
                                        r_data = res 
                                        count_mov = r_data.get("count", 0)
                                        errs = r_data.get("errors", [])
                                        errors_mov = len(errs)
                                        st.success(f"Items Movidos: {count_mov}. Errores: {errors_mov}")
                                        if errs:
                                            with st.expander("Ver Errores"):
                                                for e in errs: st.write(e)
                                else:
                                    st.error("Error conectando con Agente")
                        except Exception as e:
                            st.error(f"Error agente: {e}")
                    else:
                        # Web mode logic (omitted for brevity, assume similar to rename)
                        pass

            # --- AÑADIR SUFIJO ---
            st.divider()
            st.markdown("#### 3. Añadir Sufijo desde BD")
            st.info("Esta herramienta buscará carpetas o archivos que coincidan con la columna de búsqueda seleccionada y añadirá el sufijo especificado.")
            
            c_match, c_suf = st.columns(2)
            
            with c_match:
                default_idx = 0
                # Use display_keys which is now safe
                if "no_doc" in display_keys:
                    default_idx = display_keys.index("no_doc")
                match_col = st.selectbox("Columna para buscar Carpeta/Archivo", display_keys, index=default_idx, key="sel_match_col_suf")
                
            with c_suf:
                suffix_field = st.selectbox("Columna para Sufijo", display_keys, key="sel_field_suf_val")
                
            separator_suf = st.text_input("Separador", value="_", key="input_sep_suf")

            if st.button("Ejecutar Añadir Sufijo", key="btn_run_suf"):
                count_suf = 0
                errors_suf = 0
                
                if is_native:
                    items_to_rename = []
                    for r in records:
                        m_key = str(r.get(match_col, "")).strip()
                        s_val = str(r.get(suffix_field, "")).strip()
                        s_val = "".join([c for c in s_val if c.isalnum() or c in (' ', '-', '_')]).strip()
                        if m_key and s_val:
                            items_to_rename.append({"key": m_key, "suffix": s_val})
                    
                    if items_to_rename:
                        try:
                            from src.agent_client import send_command, wait_for_result
                            username = st.session_state.get("username", "default")
                            
                            with st.spinner("Ejecutando renombrado masivo vía Agente Local..."):
                                item_type_val = st.session_state.get("item_type", "both")
                                task_id = send_command(username, "bulk_rename", {
                                    "path": source_path,
                                    "items": items_to_rename,
                                    "separator": separator_suf,
                                    "item_type": item_type_val
                                })
                                
                                if task_id:
                                    res = wait_for_result(task_id, timeout=600)
                                    
                                    if "error" in res:
                                        st.error(f"Error del Agente: {res['error']}")
                                    else:
                                        # Direct result usage
                                        r_data = res 
                                        count_suf = r_data.get("count", 0)
                                        errs = r_data.get("errors", [])
                                        errors_suf = len(errs)
                                        st.success(f"Items Renombrados: {count_suf}. Errores: {errors_suf}")
                                        if errs:
                                            with st.expander("Ver Errores"):
                                                for e in errs: st.write(e)
                                else:
                                    st.error("Error contactando al Agente")
                        except Exception as e:
                            st.error(f"Excepción: {e}")
                    else:
                        st.warning("No hay items para renombrar según los filtros.")

                else:
                    # WEB MODE: Use Local Loop (Server Side)
                    try:
                        root_files = []
                        root_folders = []
                        if os.path.exists(source_path):
                            try:
                                all_items = os.listdir(source_path)
                                root_files = [f for f in all_items if os.path.isfile(os.path.join(source_path, f))]
                                root_folders = [d for d in all_items if os.path.isdir(os.path.join(source_path, d))]
                            except Exception as e:
                                print(f"Error listing root items: {e}")
                        
                        st.toast(f"📂 Buscando en: {os.path.basename(source_path)} ({len(root_files)} archivos, {len(root_folders)} carpetas)", icon="🔍")
                        
                        progress_bar = st.progress(0)
                        total_records = len(records)
                        
                        item_type_val = st.session_state.get("item_type", "both")
                        scope_folders = str(item_type_val).lower() in ["todo", "both", "carpetas", "folders", "directory"]
                        scope_files = str(item_type_val).lower() in ["todo", "both", "archivos", "files", "file"]

                        for i, r in enumerate(records):
                            match_key = str(r.get(match_col, "")).strip()
                            if not match_key: continue
                                
                            suffix_val = str(r.get(suffix_field, "")).strip()
                            suffix_val = "".join([c for c in suffix_val if c.isalnum() or c in (' ', '-', '_')]).strip()
                            if not suffix_val: continue

                            # A) FOLDER MODE
                            if scope_folders:
                                matching_folders = []
                                normalized_key = match_key.lower()
                                
                                for d in root_folders:
                                    d_lower = d.lower()
                                    if d_lower == normalized_key:
                                        matching_folders.append(d)
                                        continue
                                    if d_lower.startswith(normalized_key):
                                        remainder = d_lower[len(normalized_key):]
                                        if remainder.startswith("_") and len(remainder) > 1 and remainder[1:].isdigit():
                                            matching_folders.append(d)
                                
                                for folder_name in matching_folders:
                                    folder_path = os.path.join(source_path, folder_name)
                                    if os.path.isdir(folder_path):
                                        # Check if folder name already has suffix
                                        if not folder_name.endswith(f"{separator_suf}{suffix_val}"):
                                            new_folder_name = f"{folder_name}{separator_suf}{suffix_val}"
                                            new_folder_path = os.path.join(source_path, new_folder_name)
                                            try:
                                                os.rename(folder_path, new_folder_path)
                                                count_suf += 1
                                            except Exception as e:
                                                print(f"Error renaming folder {folder_name}: {e}")
                                                errors_suf += 1
                            
                            # B) FLAT FILES
                            if scope_files:
                                for filename in root_files:
                                    if match_key in filename:
                                        try:
                                            file_full_path = os.path.join(source_path, filename)
                                            if not os.path.exists(file_full_path): continue 
                                            base_name, ext = os.path.splitext(filename)
                                            if not base_name.endswith(f"{separator_suf}{suffix_val}"):
                                                new_name = f"{base_name}{separator_suf}{suffix_val}{ext}"
                                                new_file_path = os.path.join(source_path, new_name)
                                                if os.path.exists(new_file_path): continue
                                                os.rename(file_full_path, new_file_path)
                                                count_suf += 1
                                        except Exception as e:
                                            errors_suf += 1
                            
                            progress_bar.progress((i + 1) / total_records)
                        
                        st.success(f"Items Renombrados: {count_suf}. Errores: {errors_suf}")
#                         render_download_button(source_path, "dl_suf", "📦 Descargar Archivos con Sufijo (ZIP)", cleanup=not is_native)
                    except Exception as e:
                         st.error(f"Error general: {e}")

    # --- TAB CONTENIDO (Former Subtab 2) ---
    with tab_content:
        st.subheader("📝 Generación de Contenido (Docs y Firmas)")
        
        default_hardcoded = os.path.join(os.path.expanduser("~"), "Documents", "GestionDocumental")
        
        current_global_path = st.session_state.get("current_path", os.getcwd())
        if not current_global_path: current_global_path = os.getcwd()
        
        base_path_content = render_path_selector(
            label="Carpeta Raíz",
            key="input_base_path_content",
            default_path=current_global_path,
            omit_checkbox=False
        )

        if not is_native and (not base_path_content or not os.path.exists(base_path_content)):
             temp_content = os.path.join(os.getcwd(), "temp_downloads", f"content_{int(time.time())}")
             if "temp_content_path" not in st.session_state:
                 os.makedirs(temp_content, exist_ok=True)
                 st.session_state.temp_content_path = temp_content
             base_path_content = st.session_state.temp_content_path
             st.info(f"Modo Web: Usando directorio temporal para contenido: {base_path_content}")
        
        st.caption("Patrón de Carpetas (Debe coincidir con la estructura creada)")
        folder_structure_content = st.text_input("Patrón de Carpetas", value="{eps}/{no_factura} - {nombre_completo}", key="input_folder_structure_content")

        st.divider()
        st.markdown("### 2. Configuración de Acciones")
        
        col_conf1, col_conf2 = st.columns(2)
        
        with col_conf1:
            st.markdown("#### Documentos (.docx)")
            template_file = st.file_uploader("Subir Plantilla Base (.docx)", type=["docx"], key="uploader_template_gen")
            template_bytes = None
            if template_file:
                template_bytes = template_file.getvalue()
            st.caption("Si no sube plantilla, buscará 'plantilla.docx' en cada carpeta.")

        with col_conf2:
            st.markdown("#### Firmas")
            st.caption("Firma Web (URL)")
            sign_url_pattern = st.text_input("URL Patrón", value="https://oportunidaddevida.com/opvcitas/admisionescall/firmas/{no_doc}.png", key="input_sign_url_gen")
            
            st.caption("Firma Digital (Generada)")
            font_name = st.radio("Fuente", ["Pacifico", "My Ugly Handwriting"], index=0, horizontal=True)
            font_size = st.number_input("Tamaño Fuente", value=70, min_value=20, max_value=200, step=5)
            use_natural_style = st.checkbox("Estilo Natural (Rotación)", value=True)

        st.divider()
        st.markdown("### 3. Ejecutar Acciones")

        disabled_state = False
        filtered_records_content = []
        
        if not records:
            st.warning("⚠️ Para ejecutar acciones, primero debe importar datos en la pestaña 'Importar Excel'.")
            disabled_state = True
        else:
            filtered_records_content = get_filtered_records(records, key_suffix="content")
            st.info(f"Registros a procesar: {len(filtered_records_content)}")
        
        col_act1, col_act2, col_act3, col_act4 = st.columns(4)
        
        # ACTION 1: DISTRIBUTE BASE DOC (Copy Template Only)
        with col_act1:
            if st.button("Distribuir Base", use_container_width=True, disabled=disabled_state):
                count_dist = 0
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                if not template_bytes:
                    st.error("❌ Debe subir una plantilla base primero.")
                else:
                    resolved_paths_content = resolve_unique_paths(filtered_records_content, folder_structure_content)

                    target_paths = []
                    for record in filtered_records_content:
                        try:
                            rel_path = resolved_paths_content.get(record['id'])
                            if not rel_path:
                                safe_record = {k: str(v).replace("/", "-").replace("\\", "-").strip() for k, v in record.items()}
                                rel_path = folder_structure_content.format(**safe_record).upper()
                                
                            full_path = os.path.join(base_path_content, rel_path)
                            dest_base = os.path.join(full_path, "documento_base.docx")
                            target_paths.append(dest_base)
                        except: pass

                    if is_native and target_paths:
                        try:
                            from src.agent_client import send_command, wait_for_result
                            username = st.session_state.get("username", "default")
                            
                            template_b64 = base64.b64encode(template_bytes).decode('utf-8')
                            
                            with st.spinner("Distribuyendo archivo vía Agente Local..."):
                                task_id = send_command(username, "distribute_file", {
                                    "paths": target_paths,
                                    "content_b64": template_b64
                                })
                                
                                if task_id:
                                    res = wait_for_result(task_id, timeout=600)
                                    if "error" in res:
                                        st.error(f"Error Agente: {res['error']}")
                                    else:
                                        r_data = res
                                        count_dist = r_data.get("count", 0)
                                        errs = r_data.get("errors", [])
                                        if errs:
                                            with st.expander("Errores Distribución"):
                                                for e in errs: st.write(e)
                                else:
                                    st.error("Error conectando con Agente")
                        except Exception as e:
                            st.error(f"Error: {e}")
                    else:
                        for i, dest_base in enumerate(target_paths):
                            try:
                                os.makedirs(os.path.dirname(dest_base), exist_ok=True)
                                with open(dest_base, "wb") as f:
                                    f.write(template_bytes)
                                count_dist += 1
                            except Exception as e:
                                print(f"Error distributing doc: {e}")
                            
                            progress_bar.progress((i + 1) / len(target_paths))
                            status_text.text(f"Distribuyendo {i+1}/{len(target_paths)}...")
                    
                    st.success(f"✅ Base Distribuida: {count_dist}")

        # ACTION 2: MODIFY DOC (Fill Data)
        with col_act2:
            use_hardcoded_logic = True 
            if st.button("Llenar Datos", use_container_width=True, disabled=disabled_state):
                
                resolved_paths_content = resolve_unique_paths(filtered_records_content, folder_structure_content)
                
                if is_native:
                    try:
                        from src.agent_client import send_command, wait_for_result
                        username = st.session_state.get("username", "default")
                        
                        template_b64 = None
                        if template_bytes:
                            template_b64 = base64.b64encode(template_bytes).decode('utf-8')
                        
                        tasks = []
                        for record in filtered_records_content:
                            safe_record = {k: str(v).replace("/", "-").replace("\\", "-").strip() for k, v in record.items()}
                            rel_path = resolved_paths_content.get(record['id'])
                            if not rel_path:
                                rel_path = folder_structure_content.format(**safe_record).upper()
                            
                            custom_map = {
                                "Nombre carpeta": safe_record.get("nombre_completo", ""),
                                "Nombre completo": safe_record.get("nombre_completo", ""),
                                "Ciudad y Fecha": fmt_date(safe_record.get("fecha_salida", "")),
                                "Tipo Documento": safe_record.get("tipo_doc", ""),
                                "Numero Documento": safe_record.get("no_doc", ""),
                                "Servicio": safe_record.get("descripcion", ""), 
                                "EPS": safe_record.get("eps", ""),
                                "Tipo Servicio": safe_record.get("tipo_servicio", ""),
                                "Regimen": safe_record.get("regimen", ""),
                                "Categoria": safe_record.get("categoria", ""),
                                "Valor Cuota Moderadora": safe_record.get("copago", ""),
                                "Valor Cuota Moderador": safe_record.get("copago", ""), 
                                "Numero Autorizacion": safe_record.get("autorizacion", ""),
                                "Fecha y Hora Atencion": fmt_date(safe_record.get("fecha_ingreso", "")),
                                "Fecha Finalizacion": fmt_date(safe_record.get("fecha_salida", ""))
                            }
                            
                            regex_replacements = []
                            if use_hardcoded_logic:
                                legacy_data = {
                                    'date': fmt_date(safe_record.get("fecha_salida", "")), 
                                    'full_name': safe_record.get("nombre_completo", ""),
                                    'doc_type': safe_record.get("tipo_doc", ""),
                                    'doc_num': safe_record.get("no_doc", ""),
                                    'service': safe_record.get("descripcion", ""),
                                    'eps': safe_record.get("eps", ""),
                                    'tipo_servicio': safe_record.get("tipo_servicio", ""),
                                    'regimen': safe_record.get("regimen", ""),
                                    'categoria': safe_record.get("categoria", ""),
                                    'cuota': safe_record.get("copago", ""),
                                    'auth': safe_record.get("autorizacion", ""),
                                    'fecha_atencion': fmt_date(safe_record.get("fecha_ingreso", "")),
                                    'fecha_fin': fmt_date(safe_record.get("fecha_salida", ""))
                                }
                                
                                if legacy_data['date']:
                                    regex_replacements.append((r"Santiago de Cali, (?!.*\d{4}).*", f"Santiago de Cali, {legacy_data['date']}"))
                                
                                if legacy_data['full_name']:
                                    regex_replacements.append((r"Yo .*identificado con.*", f"Yo {legacy_data['full_name']} identificado con {legacy_data['doc_type']}, Numero {legacy_data['doc_num']} en calidad de paciente, doy fé y acepto el servicio de {legacy_data['service']} brindado por la IPS OPORTUNIDAD DE VIDA S.A.S"))
                                
                                legacy_replacements_map = {
                                    "EPS:": legacy_data['eps'], "TIPO SERVICIO:": legacy_data['tipo_servicio'],
                                    "REGIMEN:": legacy_data['regimen'], "CATEGORIA:": legacy_data['categoria'],
                                    "VALOR CUOTA MODERADORA:": legacy_data['cuota'], "AUTORIZACION:": legacy_data['auth'],
                                    "Fecha de Atención:": legacy_data['fecha_atencion'], "Fecha de Finalización:": legacy_data['fecha_fin']
                                }
                                for k, v in legacy_replacements_map.items():
                                    regex_replacements.append((rf"({re.escape(k)})\s*.*", rf"\1 {v}"))
                        
                            tasks.append({
                                "rel_path": rel_path,
                                "data": {**safe_record, **custom_map},
                                "regex_replacements": regex_replacements
                            })
                
                        if tasks:
                            task_id = send_command(username, "fill_docx", {
                                "base_path": base_path_content,
                                "tasks": tasks,
                                "template_b64": template_b64
                            })
                            
                            if task_id:
                                with st.spinner("Ejecutando llenado masivo de DOCX vía Agente Local..."):
                                    res = wait_for_result(task_id, timeout=600)
                                
                                if "error" in res:
                                    st.error(f"Error del Agente: {res['error']}")
                                else:
                                    count_suc = res.get("count", 0)
                                    errors_list = res.get("errors", [])
                                    st.success(f"✅ Documentos Modificados (Agente): {count_suc}")
                                    if errors_list:
                                        with st.expander(f"⚠️ Errores ({len(errors_list)})"):
                                            for err in errors_list:
                                                st.write(err)
                            else:
                                st.error("No se pudo conectar con el Agente Local.")
                    except ImportError:
                        st.error("Librería de Agente no encontrada.")
                    except Exception as e:
                        st.error(f"Error inesperado al conectar con agente: {e}")
                else:
                    st.warning("⚠️ El llenado de documentos requiere el Modo Nativo (Agente Local) para operar correctamente.")

        # ACTION 3: SIGNATURES
        with col_act3:
            if st.button("Generar Firmas", use_container_width=True, disabled=disabled_state):
                count_sig = 0
                resolved_paths_content = resolve_unique_paths(filtered_records_content, folder_structure_content)
                progress_bar = st.progress(0)
                
                if is_native:
                    import io
                    import base64
                    files_to_write = []
                    
                    for i, record in enumerate(filtered_records_content):
                         try:
                             rel_path = resolved_paths_content.get(record['id'])
                             if not rel_path: continue
                             
                             full_path = os.path.join(base_path_content, rel_path)
                             sig_path = os.path.join(full_path, "firma_generada.png")
                             
                             nombre = record.get("nombre_completo", "Firma")
                             img = generate_signature_image(nombre, font_name=font_name, size=font_size)
                             
                             buf = io.BytesIO()
                             img.save(buf, format="PNG")
                             content_b64 = base64.b64encode(buf.getvalue()).decode('utf-8')
                             
                             files_to_write.append({
                                 "path": sig_path,
                                 "content_b64": content_b64
                             })
                         except Exception as e:
                             print(f"Error preparing signature: {e}")
                         
                         progress_bar.progress((i + 1) / len(filtered_records_content))
                    
                    if files_to_write:
                        try:
                            from src.agent_client import send_command, wait_for_result
                            username = st.session_state.get("username", "default")
                            
                            with st.spinner(f"Enviando {len(files_to_write)} firmas al Agente Local..."):
                                task_id = send_command(username, "write_files", {
                                    "files": files_to_write
                                })
                                
                                if task_id:
                                    res = wait_for_result(task_id, timeout=600)
                                    if "error" in res:
                                        st.error(f"Error Agente: {res['error']}")
                                    else:
                                        r_data = res
                                        count_sig = r_data.get("count", 0)
                                else:
                                    st.error("No se pudo conectar con el Agente")
                        except Exception as e:
                            st.error(f"Error: {e}")
                else:
                    for i, record in enumerate(filtered_records_content):
                        rel_path = resolved_paths_content.get(record['id'])
                        if not rel_path: continue
                        
                        full_path = os.path.join(base_path_content, rel_path)
                        os.makedirs(full_path, exist_ok=True)
                        
                        nombre = record.get("nombre_completo", "Firma")
                        try:
                            img = generate_signature_image(nombre, font_name=font_name, size=font_size)
                            sig_path = os.path.join(full_path, "firma_generada.png")
                            img.save(sig_path, "PNG")
                            count_sig += 1
                        except Exception as e:
                            print(f"Error generating signature: {e}")
                        
                        progress_bar.progress((i + 1) / len(filtered_records_content))
                    
                st.success(f"✅ Firmas Generadas: {count_sig}")

        # ACTION 4: DOWNLOAD
        with col_act4:
            st.write(" ") # Spacer
            if os.path.exists(base_path_content):
#                  render_download_button(base_path_content, "dl_content_final", "📦 Descargar Todo (ZIP)", cleanup=False)
                pass
