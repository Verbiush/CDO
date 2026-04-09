import traceback
import time
from datetime import datetime
import requests
import zipfile
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.service import Service
    from webdriver_manager.chrome import ChromeDriverManager
except ImportError:
    webdriver = None
try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import json
import os
import shutil
import platform
import sys
import threading
import queue
import tkinter as tk
from tkinter import simpledialog, messagebox, scrolledtext, ttk
from getpass import getpass
import re
import base64
from io import BytesIO
import pandas as pd
import io
import logging

# MOCK STREAMLIT FOR LOCAL AGENT TO AVOID IMPORT ERRORS FROM TABS
import sys
try:
    import streamlit as st
except ImportError:
    class MockSt:
        def progress(self, *args, **kwargs):
            class _Prog:
                def progress(self, *a, **k): pass
                def empty(self): pass
            return _Prog()
        def __getattr__(self, name):
            def _dummy(*args, **kwargs): return None
            return _dummy
        @property
        def session_state(self):
            return {}
    sys.modules['streamlit'] = MockSt()

# Configure logging
logging.basicConfig(
    filename='agent_debug.log',
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    force=True
)

# Try to import docx, but don't fail if not present
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
    Paragraph = None
    qn = None

# Add parent directory to path to allow imports from src (if running from source)
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

# Find config file location
def find_config_file():
    # 1. Check CWD
    if os.path.exists("agent_config.json"):
        return "agent_config.json"
    
    # 2. Check script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    path_in_script_dir = os.path.join(script_dir, "agent_config.json")
    if os.path.exists(path_in_script_dir):
        return path_in_script_dir
        
    # 3. Check LOCALAPPDATA/CDO_Organizer
    local_appdata = os.getenv('LOCALAPPDATA', os.path.expanduser("~"))
    path_in_appdata = os.path.join(local_appdata, "CDO_Organizer", "agent_config.json")
    if os.path.exists(path_in_appdata):
        return path_in_appdata
        
    # Default to CWD if not found
    return "agent_config.json"

CONFIG_FILE = find_config_file()

# --- FILE PROCESSING LOGIC ---

def _encode_bytes(data):
    if hasattr(data, 'getvalue'):
        data = data.getvalue()
    if isinstance(data, bytes):
        return base64.b64encode(data).decode('utf-8')
    return data

def _serialize_analysis_result(res):
    if isinstance(res, dict) and "files" in res:
        for f in res["files"]:
            if "data" in f:
                f["data"] = _encode_bytes(f["data"])
    return res

def _expand_file_list(files):
    """
    Expande una lista de archivos/directorios.
    Si un elemento es un directorio, busca recursivamente archivos .pdf.
    """
    expanded = []
    for item in files:
        if os.path.isdir(item):
            for root, _, filenames in os.walk(item):
                for f in filenames:
                    if f.lower().endswith('.pdf'):
                        expanded.append(os.path.join(root, f))
        elif os.path.isfile(item):
            expanded.append(item)
    return expanded

def recursive_update_cups(data, old_val, new_val):
    count = 0
    if isinstance(data, dict):
        for k, v in data.items():
            if k == "codServicio" and str(v).strip() == str(old_val).strip():
                data[k] = new_val
                count += 1
            elif isinstance(v, (dict, list)):
                count += recursive_update_cups(v, old_val, new_val)
    elif isinstance(data, list):
        for item in data:
            count += recursive_update_cups(item, old_val, new_val)
    return count

def process_update_cups(folder_path, old_val, new_val):
    count_files = 0
    total_changes = 0
    errors = []
    
    if not os.path.isdir(folder_path):
        return {"status": "error", "message": "Carpeta no válida"}

    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    for file_path in files_to_process:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        
            changes = recursive_update_cups(data, old_val, new_val)
        
            if changes > 0:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                count_files += 1
                total_changes += changes
        except Exception as e:
            errors.append(f"{os.path.basename(file_path)}: {str(e)}")
            
    return {
        "count_files": count_files,
        "total_changes": total_changes,
        "errors": errors
    }

def recursive_update_key(data, key, new_val):
    count = 0
    if isinstance(data, dict):
        for k, v in data.items():
            if k == key:
                data[k] = new_val
                count += 1
            elif isinstance(v, (dict, list)):
                count += recursive_update_key(v, key, new_val)
    elif isinstance(data, list):
        for item in data:
            count += recursive_update_key(item, key, new_val)
    return count

def process_update_key(folder_path, key_target, new_value):
    count_files = 0
    total_changes = 0
    errors = []
    
    if not os.path.isdir(folder_path):
        return {"count_files": 0, "total_changes": 0, "errors": ["Carpeta no válida"]}

    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    for file_path in files_to_process:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        
            changes = recursive_update_key(data, key_target, new_value)
        
            if changes > 0:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                count_files += 1
                total_changes += changes
        except Exception as e:
            errors.append(f"{os.path.basename(file_path)}: {str(e)}")
            
    return {
        "count_files": count_files,
        "total_changes": total_changes,
        "errors": errors
    }

def recursive_update_notes(data, target_text, new_note):
    count = 0
    if isinstance(data, dict):
        for k, v in data.items():
            if isinstance(v, str) and target_text in v:
                data[k] = new_note
                count += 1
            elif isinstance(v, (dict, list)):
                count += recursive_update_notes(v, target_text, new_note)
    elif isinstance(data, list):
        for item in data:
            count += recursive_update_notes(item, target_text, new_note)
    return count

def process_update_notes(folder_path, target_text, new_note):
    count_files = 0
    total_changes = 0
    errors = []
    
    if not os.path.isdir(folder_path):
        return {"count_files": 0, "total_changes": 0, "errors": ["Carpeta no válida"]}

    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    for file_path in files_to_process:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        
            changes = recursive_update_notes(data, target_text, new_note)
        
            if changes > 0:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                count_files += 1
                total_changes += changes
        except Exception as e:
            errors.append(f"{os.path.basename(file_path)}: {str(e)}")
            
    return {
        "count_files": count_files,
        "total_changes": total_changes,
        "errors": errors
    }

def recursive_strip(data):
    if isinstance(data, dict):
        return {k.strip(): recursive_strip(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [recursive_strip(i) for i in data]
    elif isinstance(data, str):
        return data.strip()
    return data

def process_clean_json(folder_path):
    count_files = 0
    errors = []
    
    if not os.path.isdir(folder_path):
        return {"count_files": 0, "errors": ["Carpeta no válida"]}

    files_to_process = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.json'):
                files_to_process.append(os.path.join(root, file))
    
    for file_path in files_to_process:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        
            cleaned_data = recursive_strip(data)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(cleaned_data, f, indent=4, ensure_ascii=False)
            
            count_files += 1
        except Exception as e:
            errors.append(f"{os.path.basename(file_path)}: {str(e)}")
            
    return {
        "count_files": count_files,
        "errors": errors
    }

def replace_text_in_element(paragraph, mapping):
    if not paragraph.text.strip():
        return 0
    
    text_has_braces = "{" in paragraph.text
    text_has_chevrons = "<<" in paragraph.text or "«" in paragraph.text
    
    if not (text_has_braces or text_has_chevrons):
        return 0

    original_text = paragraph.text
    current_text = original_text
    count = 0
    
    for key, val in mapping.items():
        key_clean = key.replace(" ", "")
        p1 = r"\{\s*" + re.escape(key) + r"\s*\}"
        p2 = r"\{\s*" + re.escape(key_clean) + r"\s*\}"
        p3 = r"\{\s*" + re.escape(key.replace(" ", "_")) + r"\s*\}"
        
        c1 = r"(?:«|<<)\s*" + re.escape(key) + r"\s*(?:»|>>)"
        c2 = r"(?:«|<<)\s*" + re.escape(key_clean) + r"\s*(?:»|>>)"
        c3 = r"(?:«|<<)\s*" + re.escape(key.replace(" ", "_")) + r"\s*(?:»|>>)"
        
        patterns = [p1, p2, p3] if text_has_braces else []
        if text_has_chevrons:
            patterns.extend([c1, c2, c3])
            
        for pat in patterns:
            if re.search(pat, current_text, re.IGNORECASE):
                current_text = re.sub(pat, str(val), current_text, flags=re.IGNORECASE)
                count += 1
    
    if count > 0 and current_text != original_text:
        paragraph.text = current_text
        return count
    return 0

def iter_all_paragraphs(doc_obj):
    yield from doc_obj.paragraphs
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    
    if doc_obj.element.body is not None:
        for txbx in doc_obj.element.body.iter(qn('w:txbxContent')):
            for p_element in txbx.iter(qn('w:p')):
                yield Paragraph(p_element, doc_obj)

    for section in doc_obj.sections:
        headers = [section.header, section.first_page_header, section.even_page_header]
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        
        for header in headers:
            if header and not header.is_linked_to_previous:
                yield from header.paragraphs
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs
                for txbx in header._element.iter(qn('w:txbxContent')):
                    for p_element in txbx.iter(qn('w:p')):
                        yield Paragraph(p_element, header)

        for footer in footers:
            if footer and not footer.is_linked_to_previous:
                yield from footer.paragraphs
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs
                for txbx in footer._element.iter(qn('w:txbxContent')):
                    for p_element in txbx.iter(qn('w:p')):
                        yield Paragraph(p_element, footer)

def process_fill_docx_ovida_full(base_path, tasks):
    count_success = 0
    errors = []
    
    if Document is None:
        return {"count": 0, "errors": ["Librería python-docx no instalada en Agente"]}

    for i, task in enumerate(tasks):
        rel_path = task.get("rel_path")
        datos = task.get("datos", {})
        
        if not rel_path:
            errors.append(f"Task {i}: Falta ruta relativa")
            continue
            
        full_path = os.path.join(base_path, rel_path)
        
        try:
            target_docx = None
            if os.path.exists(full_path):
                for f in os.listdir(full_path):
                    if f.lower().endswith('.docx') and not f.startswith('~'):
                        target_docx = os.path.join(full_path, f)
                        break
                    
            if not target_docx:
                errors.append(f"No se encontró archivo docx en: {full_path}")
                continue
            
            doc = Document(target_docx)
            for p in doc.paragraphs:
                if "Santiago de Cali, " in p.text: 
                    p.text = f"Santiago de Cali,  {datos.get('date', '')}"
                
                if "Yo " in p.text and "identificado con" in p.text:
                    p.text = f"Yo {datos.get('full_name', '')} identificado con {datos.get('doc_type', '')}, Numero {datos.get('doc_num', '')} en calidad de paciente, doy fé y acepto el servicio de {datos.get('service', '')} brindado por la IPS OPORTUNIDAD DE VIDA S.A.S"
                
                replacements = {
                    "EPS:": datos.get('eps', ''), "TIPO SERVICIO:": datos.get('tipo_servicio', ''),
                    "REGIMEN:": datos.get('regimen', ''), "CATEGORIA:": datos.get('categoria', ''),
                    "VALOR CUOTA MODERADORA:": datos.get('cuota', ''), "AUTORIZACION:": datos.get('auth', ''),
                    "Fecha de Atención:": datos.get('fecha_atencion', ''), "Fecha de Finalización:": datos.get('fecha_fin', '')
                }
                for key, val in replacements.items():
                    if key in p.text:
                        p.text = re.sub(rf'({key})\s*.*', r'\1 ' + str(val), p.text, count=1)
            
            sig_idx = -1
            for idx, p in enumerate(doc.paragraphs):
                if "FIRMA DE ACEPTACION" in p.text.upper():
                    sig_idx = idx
                    break
            if sig_idx != -1 and sig_idx + 2 < len(doc.paragraphs):
                doc.paragraphs[sig_idx + 2].text = datos.get('full_name', '').upper()
            
            doc.save(target_docx)
            count_success += 1
                
        except Exception as e:
            errors.append(f"Error procesando {rel_path}: {str(e)}")
            
    return {"count": count_success, "errors": errors}

def process_fill_docx(base_path, tasks, template_b64=None):
    count_success = 0
    errors = []
    
    if Document is None:
        return {"count": 0, "errors": ["Librería python-docx no instalada en Agente"]}

    template_bytes = None
    if template_b64:
        try:
            template_bytes = base64.b64decode(template_b64)
        except Exception as e:
            return {"count": 0, "errors": [f"Error decodificando plantilla: {e}"]}

    for i, task in enumerate(tasks):
        rel_path = task.get("rel_path")
        data = task.get("data", {})
        
        if not rel_path:
            errors.append(f"Task {i}: Falta ruta relativa")
            continue
            
        full_path = os.path.join(base_path, rel_path)
        
        try:
            os.makedirs(full_path, exist_ok=True)
            
            doc_to_process = None
            doc_path = None
            
            # 1. Try distributed base doc
            base_doc_path = os.path.join(full_path, "documento_base.docx")
            if os.path.exists(base_doc_path):
                doc_to_process = Document(base_doc_path)
                doc_path = base_doc_path
            # 2. Try template bytes
            elif template_bytes:
                doc_to_process = Document(BytesIO(template_bytes))
            # 3. Fallback to local files
            else:
                local_candidates = [f for f in os.listdir(full_path) if f.lower().endswith(".docx") and not f.startswith("~")]
                file_pattern = task.get("file_pattern")
                if file_pattern:
                    import re
                    pattern = re.compile(file_pattern, re.IGNORECASE)
                    local_candidates = [f for f in local_candidates if pattern.match(f)]
                
                if local_candidates:
                    if "plantilla.docx" in local_candidates:
                        doc_path = os.path.join(full_path, "plantilla.docx")
                    elif "documento_generado.docx" in local_candidates:
                        doc_path = os.path.join(full_path, "documento_generado.docx")
                    else:
                        doc_path = os.path.join(full_path, local_candidates[0])
                    
                    if doc_path:
                        doc_to_process = Document(doc_path)
            
            if doc_to_process:
                # Perform replacements
                replacements = 0
                regex_replacements = task.get("regex_replacements", [])
                
                for p in iter_all_paragraphs(doc_to_process):
                    # 1. Map-based replacements
                    replacements += replace_text_in_element(p, data)
                    
                    # 2. Regex-based replacements
                    for pat, repl in regex_replacements:
                        if p.text.strip() and re.search(pat, p.text):
                            try:
                                new_text = re.sub(pat, repl, p.text)
                                if new_text != p.text:
                                    p.text = new_text
                                    replacements += 1
                            except Exception as e:
                                pass # Ignore regex errors

                # Save
                dest_doc = doc_path if doc_path else os.path.join(full_path, "documento_base.docx")
                doc_to_process.save(dest_doc)
                count_success += 1
            else:
                errors.append(f"No se encontró plantilla para: {rel_path}")
                
        except Exception as e:
            errors.append(f"Error procesando {rel_path}: {str(e)}")
            
    return {"count": count_success, "errors": errors}

def process_sign_docx_massive(base_path, docx_filename, signature_filename):
    if Document is None:
        return {"count": 0, "errors": ["Librería python-docx no instalada en el agente local."]}

    folders_to_process = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
    if not folders_to_process: 
        return {"count": 0, "errors": ["No se encontraron carpetas para procesar."]}
    
    procesados = 0
    errores = []
    
    for folder_name in folders_to_process:
        folder_path = os.path.join(base_path, folder_name)
        docx_path = os.path.join(folder_path, docx_filename)
        
        signature_path = None
        possible_sig_paths = [
            os.path.join(folder_path, signature_filename),
            os.path.join(folder_path, "tipografia", signature_filename)
        ]
        for path in possible_sig_paths:
            if os.path.exists(path):
                signature_path = path
                break
        
        if not os.path.exists(docx_path) or not signature_path:
            errores.append(f"Archivos no encontrados en {folder_name}")
            continue
            
        try:
            doc = Document(docx_path)
            anchor_text = "Firma de Aceptacion"
            signature_p_index = -1

            for idx, p in enumerate(doc.paragraphs):
                if anchor_text.lower() in p.text.lower():
                    target_index = idx + 1
                    if target_index < len(doc.paragraphs):
                        signature_p_index = target_index
                    break
            
            if signature_p_index != -1:
                signature_p = doc.paragraphs[signature_p_index]
                p_element = signature_p._p
                p_element.clear_content()
                run = signature_p.add_run()
                run.add_picture(signature_path, width=Inches(1.5))
                signature_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.save(docx_path)
                procesados += 1
            else:
                errores.append(f"No se encontró 'Firma de Aceptacion' en {folder_name}")
        except Exception as e:
            errores.append(f"Error procesando {folder_name}: {str(e)}")
            
    return {"count": procesados, "errors": errores}

def process_rename_folders_mapped(path, mapping):
    count_renamed = 0
    errors = []
    
    if not os.path.isdir(path):
        return {"count": 0, "errors": ["Ruta inválida"]}
        
    try:
        dirs = [d for d in os.listdir(path) if os.path.isdir(os.path.join(path, d))]
        
        for dirname in dirs:
            dir_path = os.path.join(path, dirname)
            matched_new_name = None
            
            if dirname in mapping:
                matched_new_name = mapping[dirname]
            else:
                for curr_val, new_val in mapping.items():
                    if curr_val in dirname:
                        matched_new_name = new_val
                        break
            
            if matched_new_name and matched_new_name != dirname:
                try:
                    new_path = os.path.join(path, matched_new_name)
                    if os.path.exists(new_path):
                        errors.append(f"Omitido {dirname}: Destino ya existe")
                    else:
                        os.rename(dir_path, new_path)
                        count_renamed += 1
                except Exception as e:
                    errors.append(f"Error {dirname}: {str(e)}")
                    
    except Exception as e:
        errors.append(f"Error listando directorios: {str(e)}")
        
    return {"count": count_renamed, "errors": errors}

def process_organize_files_mapped(source_path, dest_base_path, mapping):
    count_moved = 0
    errors = []
    
    if not os.path.isdir(source_path):
        return {"count": 0, "errors": ["Ruta Origen inválida"]}
    
    if not os.path.exists(dest_base_path):
        try:
            os.makedirs(dest_base_path)
        except Exception as e:
            return {"count": 0, "errors": [f"No se pudo crear destino base: {str(e)}"]}
            
    try:
        files = [f for f in os.listdir(source_path) if os.path.isfile(os.path.join(source_path, f))]
        
        for filename in files:
            file_path = os.path.join(source_path, filename)
            matched_folder = None
            
            # Find destination folder based on mapping keys in filename
            for key, folder_name in mapping.items():
                if key in filename:
                    matched_folder = folder_name
                    break
            
            if matched_folder:
                try:
                    target_dir = os.path.join(dest_base_path, matched_folder)
                    if not os.path.exists(target_dir):
                        os.makedirs(target_dir)
                        
                    target_path = os.path.join(target_dir, filename)
                    
                    if os.path.exists(target_path):
                        base, ext = os.path.splitext(filename)
                        target_path = os.path.join(target_dir, f"{base}_{int(time.time())}{ext}")
                        
                    shutil.move(file_path, target_path)
                    count_moved += 1
                except Exception as e:
                    errors.append(f"Error moviendo {filename}: {str(e)}")
                    
    except Exception as e:
        errors.append(f"Error procesando archivos: {str(e)}")
        
    return {"count": count_moved, "errors": errors}

def process_create_folders(folders):
    count_created = 0
    errors = []
    
    for folder in folders:
        try:
            if not os.path.exists(folder):
                os.makedirs(folder)
                count_created += 1
        except Exception as e:
            errors.append(f"Error creando {folder}: {str(e)}")
            
    return {"count": count_created, "errors": errors}

def process_create_folders_from_list(base_path, names, unique=False):
    count_created = 0
    errors = []
    
    # En Windows, os.path.exists puede fallar para unidades de red si no están mapeadas igual, 
    # pero asumimos que el Agente corre en el contexto del usuario.
    if not os.path.exists(base_path):
        # Intentar crear la base si no existe? No, mejor reportar error.
        return {"count": 0, "errors": [f"Ruta base no existe: {base_path}"]}
        
    for name in names:
        try:
            # Sanitizar nombre (misma lógica que en el servidor)
            nombre_base = "".join(c for c in name if c.isalnum() or c in " _-").rstrip()
            if not nombre_base: continue
            
            ruta_final = os.path.join(base_path, nombre_base)
            
            if unique and os.path.exists(ruta_final):
                contador = 2
                nombre_consecutivo = f"{nombre_base} ({contador})"
                ruta_final = os.path.join(base_path, nombre_consecutivo)
                while os.path.exists(ruta_final):
                    contador += 1
                    nombre_consecutivo = f"{nombre_base} ({contador})"
                    ruta_final = os.path.join(base_path, nombre_consecutivo)
            
            os.makedirs(ruta_final, exist_ok=True)
            count_created += 1
        except Exception as e:
            errors.append(f"Error creando {name}: {str(e)}")
            
    return {"count": count_created, "errors": errors}

def process_search_files(path, patterns, exclusion_list=None, search_by="name", item_type="both", recursive=True, search_empty_folders=False):
    print(f"DEBUG: Iniciando búsqueda en: '{path}'")
    logging.info(f"Iniciando búsqueda en: '{path}'")
    print(f"DEBUG: Patrones: {patterns}, Exclusiones: {exclusion_list}")
    logging.info(f"Patrones: {patterns}, Exclusiones: {exclusion_list}")
    
    # Normalize item_type to handle UI variations
    item_type_map = {
        "archivos": "file",
        "carpetas": "folder",
        "ambos": "both",
        "files": "file",
        "folders": "folder"
    }
    item_type = item_type_map.get(str(item_type).lower(), item_type)
    
    found_items = []
    errors = []
    
    if not os.path.isdir(path):
        print(f"ERROR: La ruta no existe o no es un directorio: '{path}'")
        logging.error(f"La ruta no existe o no es un directorio: '{path}'")
        return {"items": [], "errors": [f"Ruta inválida o no accesible: {path}"]}
        
    exclusion_list = exclusion_list or []
    
    try:
        # Helper to check exclusion
        def is_excluded(name):
            return any(excl.lower() in name.lower() for excl in exclusion_list)
            
        # Helper to check match
        def is_match(name):
            if not patterns:
                return True
            
            name_lower = name.lower()
            for pat in patterns:
                pat_lower = pat.lower()
                
                # Handle simple glob patterns
                if pat_lower == "*": return True
                if pat_lower.startswith("*."): pat_lower = pat_lower[1:] # Remove * from *.ext
                
                if search_by == "extensión":
                    if name_lower.endswith(pat_lower):
                        return True
                    # Handle case where user inputs "pdf" without dot
                    if not pat_lower.startswith('.') and name_lower.endswith('.' + pat_lower):
                        return True
                else:
                    if pat_lower in name_lower:
                        return True
            return False
            
        if recursive:
            iterator = os.walk(path)
        else:
            # Fake walk for non-recursive
            try:
                iterator = [(path, [d for d in os.listdir(path) if os.path.isdir(os.path.join(path, d))], 
                                   [f for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))])]
            except Exception as e:
                 print(f"ERROR al listar directorio no recursivo: {e}")
                 logging.error(f"ERROR al listar directorio no recursivo: {e}")
                 iterator = []
                               
        scanned_count = 0
        matched_count = 0
        
        for root, dirs, files in iterator:
            # Filter directories
            if item_type in ["folder", "both"]:
                for d in dirs:
                    if is_excluded(d): continue
                    if is_match(d):
                        full_path = os.path.join(root, d)
                        try:
                            is_empty = not os.listdir(full_path)
                            if search_empty_folders and not is_empty: continue
                            
                            found_items.append({
                                "Ruta completa": full_path,
                                "Nombre": d,
                                "Tipo": "Carpeta",
                                "Tamaño": 0,
                                "Fecha": datetime.fromtimestamp(os.path.getmtime(full_path)).strftime('%Y-%m-%d %H:%M:%S'),
                                # Compatibilidad con versión anterior (English keys)
                                "path": full_path,
                                "name": d,
                                "type": "folder",
                                "size": 0
                            })
                            matched_count += 1
                        except Exception as e:
                            logging.warning(f"Error procesando carpeta {full_path}: {e}")
                        
            # Filter files
            if item_type in ["file", "both"] and not search_empty_folders:
                for f in files:
                    scanned_count += 1
                    if is_excluded(f): continue
                    if is_match(f):
                        full_path = os.path.join(root, f)
                        try:
                            size = os.path.getsize(full_path)
                            mtime = os.path.getmtime(full_path)
                            found_items.append({
                                "Ruta completa": full_path,
                                "Nombre": f,
                                "Tipo": "Archivo",
                                "Tamaño": size,
                                "Fecha": datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M:%S'),
                                # Compatibilidad con versión anterior (English keys)
                                "path": full_path,
                                "name": f,
                                "type": "file",
                                "size": size
                            })
                            matched_count += 1
                        except Exception as e:
                             logging.warning(f"Error procesando archivo {full_path}: {e}")
        
        print(f"DEBUG: Búsqueda finalizada. Escaneados: {scanned_count}, Encontrados: {matched_count}, Total items: {len(found_items)}")
        logging.info(f"Búsqueda finalizada. Escaneados: {scanned_count}, Encontrados: {matched_count}, Total items: {len(found_items)}")
                        
    except Exception as e:
        print(f"ERROR CRITICO en process_search_files: {e}")
        logging.critical(f"ERROR CRITICO en process_search_files: {e}", exc_info=True)
        errors.append(str(e))
        
    # Compatibilidad con versión anterior del servidor: devolver lista directa
    return {"items": found_items, "errors": errors}
    # return found_items

def process_distribute_file(paths, content_b64):
    count_distributed = 0
    errors = []
    
    try:
        content_bytes = base64.b64decode(content_b64)
    except Exception as e:
        return {"count": 0, "errors": [f"Error decodificando contenido base64: {str(e)}"]}
        
    for path in paths:
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(path), exist_ok=True)
            
            with open(path, "wb") as f:
                f.write(content_bytes)
            count_distributed += 1
        except Exception as e:
            errors.append(f"Error escribiendo en {path}: {str(e)}")
            
    return {"count": count_distributed, "errors": errors}

def process_write_files(files):
    count_written = 0
    errors = []
    
    for item in files:
        path = item.get("path")
        content_b64 = item.get("content_b64")
        
        if not path or not content_b64:
            errors.append(f"Item incompleto: {item}")
            continue
            
        try:
            content_bytes = base64.b64decode(content_b64)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(path), exist_ok=True)
            
            with open(path, "wb") as f:
                f.write(content_bytes)
            count_written += 1
        except Exception as e:
            errors.append(f"Error escribiendo {path}: {str(e)}")
            
    return {"count": count_written, "errors": errors}

def process_rename_files(files):
    count_renamed = 0
    errors = []
    
    print(f"DEBUG: Iniciando renombrado de {len(files)} archivos")
    
    for item in files:
        old_path = item.get("old_path")
        new_path = item.get("new_path")
        
        if not old_path or not new_path:
            continue
            
        try:
            if os.path.exists(old_path):
                # Ensure destination folder exists
                dest_dir = os.path.dirname(new_path)
                if not os.path.exists(dest_dir):
                    os.makedirs(dest_dir, exist_ok=True)
                
                # Check if destination exists (Windows file in use or exists error)
                if os.path.exists(new_path) and new_path != old_path:
                     try:
                         # Try to remove if it exists (overwrite)
                         if os.path.isfile(new_path):
                             os.remove(new_path)
                     except Exception as e:
                         print(f"WARN: Could not remove existing destination {new_path}: {e}")

                print(f"DEBUG: Renaming '{old_path}' -> '{new_path}'")
                os.rename(old_path, new_path)
                count_renamed += 1
            else:
                msg = f"Archivo no encontrado: {old_path}"
                print(f"ERROR: {msg}")
                errors.append(msg)
        except Exception as e:
            msg = f"Error renombrando {os.path.basename(old_path)}: {str(e)}"
            print(f"ERROR: {msg}")
            errors.append(msg)
            
    print(f"DEBUG: Renombrado finalizado. Éxitos: {count_renamed}, Errores: {len(errors)}")
    return {"count": count_renamed, "errors": errors}

def process_distribute_files(content_b64, paths):
    count_distributed = 0
    errors = []
    
    try:
        content_bytes = base64.b64decode(content_b64)
    except Exception as e:
        return {"count": 0, "errors": [f"Error decodificando base64: {str(e)}"]}
        
    print(f"DEBUG: Distribuyendo archivo a {len(paths)} destinos")
        
    for path in paths:
        try:
            dest_dir = os.path.dirname(path)
            if not os.path.exists(dest_dir):
                os.makedirs(dest_dir, exist_ok=True)
                
            with open(path, "wb") as f:
                f.write(content_bytes)
            count_distributed += 1
        except Exception as e:
            msg = f"Error escribiendo en {path}: {str(e)}"
            print(f"ERROR: {msg}")
            errors.append(msg)
            
    print(f"DEBUG: Distribución finalizada. Éxitos: {count_distributed}, Errores: {len(errors)}")
    return {"count": count_distributed, "errors": errors}

def process_copiar_archivo_a_subcarpetas(archivo, carpeta_base):
    try:
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
    except Exception as e:
        return {"error": f"Error al leer la carpeta destino: {str(e)}"}

    if not subcarpetas:
        return {"message": "No se encontraron subcarpetas en la ruta destino."}

    copiados = 0
    conflictos = 0
    errores = 0
    nombre_archivo = os.path.basename(archivo)
    
    for subcarpeta in subcarpetas:
        destino_final = os.path.join(subcarpeta, nombre_archivo)
        if os.path.exists(destino_final):
            conflictos += 1
            continue
        try:
            import shutil
            shutil.copy2(archivo, destino_final)
            copiados += 1
        except Exception:
            errores += 1
            
    return {"message": f"Copia completada. Copiados: {copiados}. Conflictos: {conflictos}. Errores: {errores}."}

def process_download_files(tasks):
    count_downloaded = 0
    errors = []
    
    print(f"DEBUG: Descargando {len(tasks)} archivos")
    
    for task in tasks:
        url = task.get("url")
        dest_path = task.get("dest_path")
        
        if not url or not dest_path:
            continue
            
        try:
            # Create dest directory
            dest_dir = os.path.dirname(dest_path)
            if not os.path.exists(dest_dir):
                os.makedirs(dest_dir, exist_ok=True)
            
            # Download
            response = requests.get(url, stream=True, timeout=15)
            if response.status_code == 200:
                with open(dest_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192): 
                        if chunk: 
                            f.write(chunk)
                count_downloaded += 1
            else:
                msg = f"Error descargando {url}: Status {response.status_code}"
                print(f"ERROR: {msg}")
                # Create a placeholder file?
                try:
                    with open(dest_path + ".error.txt", "w") as f:
                        f.write(msg)
                except: pass
                errors.append(msg)
                
        except Exception as e:
            msg = f"Error procesando {url}: {str(e)}"
            print(f"ERROR: {msg}")
            errors.append(msg)
            
    print(f"DEBUG: Descarga finalizada. Éxitos: {count_downloaded}, Errores: {len(errors)}")
    return {"count": count_downloaded, "errors": errors}

def process_flat_to_excel(path):
    # Stub implementation to prevent crashes
    return {"status": "error", "message": "Función flat_to_excel no implementada en el Agente Local aún."}

def get_val_ci(data_dict, key):
    if not isinstance(data_dict, dict): return None
    for k, v in data_dict.items():
        if k.lower() == key.lower():
            return v
    return None

def process_consolidate_json(path):
    try:
        import os
        import json
        import pandas as pd
        
        if not path or not os.path.isdir(path):
            return {"status": "error", "error": "La ruta proporcionada no es una carpeta válida."}

        json_files = [f for f in os.listdir(path) if f.lower().endswith('.json')]
        
        if not json_files:
            return {"status": "error", "error": "No hay archivos JSON en la carpeta."}
        
        master_header = []
        master_users = []
        master_services = {
            "Consultas": [], "Procedimientos": [], "Urgencias": [], 
            "Hospitalizacion": [], "RecienNacidos": [], "Medicamentos": [], "OtrosServicios": []
        }
        service_map = {
            "consultas": "Consultas", "procedimientos": "Procedimientos", "urgencias": "Urgencias",
            "hospitalizacion": "Hospitalizacion", "recienNacidos": "RecienNacidos",
            "medicamentos": "Medicamentos", "otrosServicios": "OtrosServicios"
        }
        
        for fname in json_files:
            with open(os.path.join(path, fname), 'r', encoding='utf-8') as f:
                data = json.load(f)
            h_info = {
                "archivo_origen": fname,
                "numDocumentoIdObligado": data.get("numDocumentoIdObligado"),
                "numFactura": data.get("numFactura")
            }
            master_header.append(h_info)
            usuarios = data.get("usuarios", [])
            for u in usuarios:
                u_clean = {k: v for k, v in u.items() if k != "servicios"}
                u_clean["archivo_origen"] = fname
                master_users.append(u_clean)
                servicios = u.get("servicios", {})
                for j_key, s_name in service_map.items():
                    items = get_val_ci(servicios, j_key)
                    if items:
                        for item in items:
                            item_copy = item.copy()
                            item_copy["archivo_origen"] = fname
                            item_copy["consecutivoUsuario"] = u.get("consecutivo")
                            master_services[s_name].append(item_copy)
        
        out_file = os.path.join(path, "RIPS_Consolidado_Agente.xlsx")
        with pd.ExcelWriter(out_file, engine='openpyxl') as writer:
            pd.DataFrame(master_header).to_excel(writer, sheet_name="Transaccion", index=False)
            pd.DataFrame(master_users).to_excel(writer, sheet_name="Usuarios", index=False)
            for s_name, rows in master_services.items():
                if rows:
                    pd.DataFrame(rows).to_excel(writer, sheet_name=s_name, index=False)
                    
        return {"status": "success", "file_path": out_file, "message": f"Consolidados {len(json_files)} archivos."}
    except Exception as e:
        return {"status": "error", "error": str(e)}

def process_desconsolidate_json(file_path, dest_folder):
    # Stub implementation to prevent crashes
    return {"status": "error", "message": "Función desconsolidate_json no implementada en el Agente Local aún."}

def process_bulk_rename(source_path, items, separator="_", item_type="both", rename_folders=True, rename_internal_files=True):
    count_renamed = 0
    errors = []
    
    if not os.path.isdir(source_path):
        return {"count": 0, "errors": ["Carpeta no válida"]}

    try:
        root_items = os.listdir(source_path)
    except Exception as e:
        return {"count": 0, "errors": [f"Error listando directorio: {str(e)}"]}
        
    root_folders = [d for d in root_items if os.path.isdir(os.path.join(source_path, d))]
    root_files = [f for f in root_items if os.path.isfile(os.path.join(source_path, f))]
    
    scope_folders = str(item_type).lower() in ["todo", "both", "carpetas", "folders", "directory"]
    scope_files = str(item_type).lower() in ["todo", "both", "archivos", "files", "file"]
    
    for item in items:
        match_key = str(item.get("key", "")).strip()
        suffix_val = str(item.get("suffix", "")).strip()
        
        if not match_key or not suffix_val:
            continue
            
        normalized_key = match_key.lower()
        
        # 1. Matching Folders
        if scope_folders:
            matching_folders = []
            for d in root_folders:
                if d.lower() == normalized_key or \
                   (d.lower().startswith(normalized_key + "_") and d[len(normalized_key)+1:].isdigit()) or \
                   (d.lower().startswith(normalized_key + " (") and d.endswith(")")):
                   matching_folders.append(d)
            
            for folder_name in matching_folders:
                folder_path = os.path.join(source_path, folder_name)
                
                new_folder_path = folder_path
                # Check if folder name already has suffix
                if rename_folders and not folder_name.endswith(f"{separator}{suffix_val}"):
                    new_folder_name = f"{folder_name}{separator}{suffix_val}"
                    new_folder_path = os.path.join(source_path, new_folder_name)
                    
                    try:
                        os.rename(folder_path, new_folder_path)
                        count_renamed += 1
                    except Exception as e:
                        errors.append(f"Error renombrando carpeta {folder_name}: {str(e)}")
                
                # Rename internal files
                if rename_internal_files:
                    try:
                        if os.path.exists(folder_path) and not rename_folders:
                            target_path = folder_path 
                        elif os.path.exists(new_folder_path):
                            target_path = new_folder_path
                        else:
                            continue
                        
                        for filename in os.listdir(target_path):
                            file_full_path = os.path.join(target_path, filename)
                            if os.path.isfile(file_full_path):
                                base_name, ext = os.path.splitext(filename)
                                if not base_name.endswith(f"{separator}{suffix_val}"):
                                    new_name = f"{base_name}{separator}{suffix_val}{ext}"
                                    try:
                                        os.rename(file_full_path, os.path.join(target_path, new_name))
                                        count_renamed += 1
                                    except: pass
                    except Exception as e:
                        errors.append(f"Error procesando archivos internos de {folder_name}: {str(e)}")

        # 2. Matching Files in Root
        if scope_files:
            for filename in root_files:
                if match_key in filename:
                    file_full_path = os.path.join(source_path, filename)
                    base_name, ext = os.path.splitext(filename)
                    
                    if not base_name.endswith(f"{separator}{suffix_val}"):
                        new_name = f"{base_name}{separator}{suffix_val}{ext}"
                        try:
                            os.rename(file_full_path, os.path.join(source_path, new_name))
                            count_renamed += 1
                        except Exception as e:
                            errors.append(f"Error renombrando archivo {filename}: {str(e)}")

    return {"count": count_renamed, "errors": errors}

def process_validate_rips(base_path, api_url, token=None, verify_ssl=True):
    results = []
    generated_files = []
    
    if not os.path.isdir(base_path):
        return {"status": "error", "message": "Ruta inválida"}

    try:
        files = [f for f in os.listdir(base_path) if f.lower().endswith('.json') 
                 and not f.startswith("Resultados") and not f.startswith("Resp_")]
    except Exception as e:
        return {"status": "error", "message": f"Error listando archivos: {str(e)}"}
        
    if not files:
        return {"status": "warning", "message": "No se encontraron archivos JSON"}

    headers = {"Content-Type": "application/json"}
    if token:
        headers["Authorization"] = f"Bearer {token}"
        
    if not verify_ssl:
        requests.packages.urllib3.disable_warnings()

    count_processed = 0
    
    for fname in files:
        full_path = os.path.join(base_path, fname)
        res_row = {"Archivo": fname, "Estado": "Pendiente", "CUV": "", "Mensaje": ""}
        
        try:
            with open(full_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                
            try:
                r = requests.post(api_url, json=data, headers=headers, verify=verify_ssl, timeout=60)
                res_row["Estado"] = str(r.status_code)
                
                if r.status_code == 200:
                    r_json = r.json()
                    res_row["CUV"] = r_json.get("cuv") or r_json.get("CUV") or ""
                    res_row["Mensaje"] = "Validado Correctamente"
                    
                    # Save results locally
                    factura_num = data.get('numFactura', os.path.splitext(fname)[0])
                    provider_id = data.get('numDocumentoIdentificacionObligado', '999')
                    
                    # 1. ResultadosLocales
                    f_loc_name = f"ResultadosLocales_{factura_num}.json"
                    f_loc_path = os.path.join(base_path, f_loc_name)
                    with open(f_loc_path, "w", encoding="utf-8") as f_out:
                        json.dump(r_json, f_out, indent=2, ensure_ascii=False)
                    generated_files.append(f_loc_name)
                    
                    # 2. ResultadosMSPS
                    f_msps_name = f"ResultadosMSPS_{factura_num}_ID{provider_id}_R.json"
                    f_msps_path = os.path.join(base_path, f_msps_name)
                    with open(f_msps_path, "w", encoding="utf-8") as f_out:
                        json.dump(r_json, f_out, indent=2, ensure_ascii=False)
                    generated_files.append(f_msps_name)
                else:
                    res_row["Mensaje"] = r.text[:200]
            except Exception as e:
                 res_row["Estado"] = "Error Conexión"
                 res_row["Mensaje"] = str(e)
                 
        except Exception as e:
            res_row["Estado"] = "Error Lectura"
            res_row["Mensaje"] = str(e)
            
        results.append(res_row)
        count_processed += 1
        
    return {
        "status": "success",
        "processed": count_processed,
        "results": results,
        "generated_files": generated_files
    }

def process_generate_signature(text, font_name="Pacifico", size=70, width=500, height=200):
    if Image is None:
        return {"error": "Librería Pillow no instalada en Agente"}
        
    try:
        img = Image.new('RGB', (width, height), color=(255, 255, 255))
        d = ImageDraw.Draw(img)
        
        # Simple font loading logic
        font_path = f"assets/fonts/{font_name}.ttf"
        try:
            if os.path.exists(font_path):
                font = ImageFont.truetype(font_path, size)
            else:
                font = ImageFont.truetype("arial.ttf", size)
        except:
            font = ImageFont.load_default()
        
        # Center text
        try:
            left, top, right, bottom = d.textbbox((0, 0), text, font=font)
            text_w = right - left
            text_h = bottom - top
        except:
            text_w, text_h = d.textsize(text, font=font)
            
        x = (width - text_w) / 2
        y = (height - text_h) / 2
        
        d.text((x, y), text, font=font, fill=(0, 0, 0))
        
        buffered = BytesIO()
        img.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
        
        return {"image_b64": img_str}
    except Exception as e:
        return {"error": str(e)}

def process_download_ovida(base_path, records):
    if webdriver is None:
        return {"status": "error", "message": "Selenium no instalado en Agente"}
        
    driver = None
    try:
        options = webdriver.ChromeOptions()
        prefs = {
            "download.default_directory": base_path,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True,
            "safebrowsing.enabled": False,
            "safebrowsing.disable_download_protection": True,
            "profile.default_content_settings.popups": 0,
            "profile.default_content_setting_values.automatic_downloads": 1,
            "profile.content_settings.exceptions.automatic_downloads.*.setting": 1
        }
        options.add_experimental_option("prefs", prefs)
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--start-maximized")
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("--disable-popup-blocking")
        options.add_argument("--allow-running-insecure-content")
        options.add_argument("--disable-web-security")
        options.add_argument("--safebrowsing-disable-download-protection")
        options.add_argument("--safebrowsing-disable-extension-blacklist")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
        
        # Wait for login
        timeout = 300 
        start_time = time.time()
        logged_in = False
        
        while time.time() - start_time < timeout:
             try:
                 if "App/Vistas" in driver.current_url:
                     logged_in = True
                     break
             except: pass
             time.sleep(1)
            
        if not logged_in:
            driver.quit()
            return {"status": "error", "message": "Tiempo de espera agotado (Login)"}

        descargados = 0
        errores = 0
        conflictos = 0
        log_errores = []
        
        for record in records:
            try:
                estudio = str(record.get('nro_estudio', '')).strip()
                if estudio.endswith(".0"): estudio = estudio[:-2]
                
                rel_path = record.get('rel_path', 'Sin_Carpeta')
                
                if not estudio or estudio == "nan" or estudio.lower() == "none":
                    errores += 1
                    log_errores.append(f"Carpeta '{rel_path}' -> No tiene estudio (celda vacía)")
                    continue
                    
                if not rel_path or rel_path == 'Sin_Carpeta':
                    continue 

                dest_dir = os.path.abspath(os.path.join(base_path, rel_path))
                os.makedirs(dest_dir, exist_ok=True)
                final_path = os.path.join(dest_dir, f"HC_{estudio}.pdf")
                
                if os.path.exists(final_path):
                    conflictos += 1
                    continue
                
                # Format dates
                f_ing = pd.to_datetime(record.get('fecha_ingreso', '')).strftime('%Y/%m/%d')
                f_egr = pd.to_datetime(record.get('fecha_salida', '')).strftime('%Y/%m/%d')

                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                # Simple param construction
                params = f"estudio={estudio}&fecha_inicio={f_ing}&fecha_fin={f_egr}&verHC=1&verEvo=1&verPar=1&ImprimirOrdenamiento=1&ImprimirNotasPcte=0&ImprimirSolOrdenesExt=1&ImprimirGraficasHC=1&ImprimirFormatos=1&ImprimirRegistroAdmon=1&ImprimirNovedad=0&ImprimirRecomendaciones=0&ImprimirDescripcionQX=0&ImprimirNotasEnfermeria=1&ImprimirSignosVitales=0&ImprimirLog=0&ImprimirEpicrisisSinHC=0"
                full_url = f"{base_url}?{params}"
                
                driver.get(full_url)
                time.sleep(2)
                
                pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {
                    "landscape": False, "printBackground": True,
                    "paperWidth": 8.5, "paperHeight": 11,
                    "marginTop": 0.4, "marginBottom": 0.4, "marginLeft": 0.4, "marginRight": 0.4
                })
                
                pdf_data = base64.b64decode(pdf_b64['data'])
                with open(final_path, 'wb') as f:
                    f.write(pdf_data)
                
                descargados += 1
            except Exception as e:
                errores += 1
                
        driver.quit()
        return {
            "status": "success", 
            "message": f"Finalizado. Descargados: {descargados}, Errores: {errores}, Conflictos: {conflictos}",
            "stats": {"descargados": descargados, "errores": errores, "conflictos": conflictos}
        }

    except Exception as e:
        if driver: driver.quit()
        return {"status": "error", "message": str(e)}

def process_download_zeus_adjuntos(base_path, records):
    if webdriver is None:
        return {"status": "error", "message": "Selenium no instalado en Agente"}
        
    driver = None
    try:
        options = webdriver.ChromeOptions()
        prefs = {
            "download.default_directory": base_path,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True,
            "safebrowsing.enabled": False,
            "safebrowsing.disable_download_protection": True,
            "profile.default_content_settings.popups": 0,
            "profile.default_content_setting_values.automatic_downloads": 1,
            "profile.content_settings.exceptions.automatic_downloads.*.setting": 1
        }
        options.add_experimental_option("prefs", prefs)
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--start-maximized")
        options.add_argument("--ignore-certificate-errors")
        options.add_argument("--disable-popup-blocking")
        options.add_argument("--allow-running-insecure-content")
        options.add_argument("--disable-web-security")
        options.add_argument("--safebrowsing-disable-download-protection")
        options.add_argument("--safebrowsing-disable-extension-blacklist")

        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
        
        # Wait for login
        timeout = 300 
        start_time = time.time()
        logged_in = False
        
        while time.time() - start_time < timeout:
             try:
                 if "App/Vistas" in driver.current_url:
                     logged_in = True
                     break
             except: pass
             time.sleep(1)
            
        if not logged_in:
            driver.quit()
            return {"status": "error", "message": "Tiempo de espera agotado (Login)"}

        descargados = 0
        errores = 0
        conflictos = 0
        
        for record in records:
            try:
                estudio = str(record.get('nro_estudio', '')).strip()
                if estudio.endswith(".0"): estudio = estudio[:-2]
                
                if not estudio or estudio == "nan":
                    errores += 1
                    continue
                    
                rel_path = record.get('rel_path')
                if not rel_path:
                    continue 

                dest_dir = os.path.abspath(os.path.join(base_path, rel_path))
                os.makedirs(dest_dir, exist_ok=True)
                
                # Configurar el directorio de descarga para este registro usando CDP
                driver.execute_cdp_cmd("Page.setDownloadBehavior", {
                    "behavior": "allow",
                    "downloadPath": dest_dir
                })

                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/ips/App/Vistas/Hc/Adjuntos.php"
                full_url = f"{base_url}?estudio={estudio}"
                
                driver.get(full_url)
                time.sleep(3)
                
                # Clickear enlaces o imágenes que parezcan adjuntos
                elements_to_click = driver.find_elements(By.XPATH, "//a | //img")
                clicked_something = False
                for el in elements_to_click:
                    try:
                        href = el.get_attribute("href") or ""
                        src = el.get_attribute("src") or ""
                        text = el.text.lower() if hasattr(el, "text") and el.text else ""
                        
                        # Buscamos descargar el archivo
                        if "adjuntos.php" not in href and ("pdf" in src.lower() or "archivo" in text or "descargar" in text or "javascript" in href.lower() or "archivo" in src.lower()):
                            handles_before = driver.window_handles
                            
                            # Si es imagen con on_click o a href javascript, o un href a archivo
                            if el.tag_name == "img":
                                try:
                                    el.find_element(By.XPATH, "..").click()
                                except:
                                    el.click()
                            else:
                                el.click()
                                
                            time.sleep(2)
                            clicked_something = True
                            
                            handles_after = driver.window_handles
                            if len(handles_after) > len(handles_before):
                                new_window = [h for h in handles_after if h not in handles_before][0]
                                driver.switch_to.window(new_window)
                                
                                # Set CDP behavior for the new window
                                driver.execute_cdp_cmd("Page.setDownloadBehavior", {
                                    "behavior": "allow",
                                    "downloadPath": dest_dir
                                })
                                time.sleep(1)
                                
                            # Esperar primero a que el archivo comience a descargarse
                            start_dl = time.time()
                            while time.time() - start_dl < 5:
                                if any(f.endswith(".crdownload") or f.endswith(".tmp") for f in os.listdir(dest_dir)):
                                    break
                                time.sleep(0.5)

                            # Luego, esperar hasta 20s a que termine de descargarse
                            timeout_dl = 20
                            start_dl = time.time()
                            while time.time() - start_dl < timeout_dl:
                                if any(f.endswith(".crdownload") or f.endswith(".tmp") for f in os.listdir(dest_dir)):
                                    time.sleep(1)
                                else:
                                    break
                            
                            # Cierra la pestaña si se abrió una nueva, dando un tiempo extra
                            time.sleep(1)
                            if len(driver.window_handles) > len(handles_before):
                                driver.close()
                                driver.switch_to.window(handles_before[0])
                                
                    except: pass
                
                if clicked_something:
                    descargados += 1
                else:
                    errores += 1
                    log_errores.append(f"Estudio '{estudio}' (Carpeta '{rel_path}') -> no tiene estudio / no se pudo descargar el documento")
                    
            except Exception as e:
                errores += 1
                log_errores.append(f"Estudio '{estudio}' (Carpeta '{rel_path}') -> Error interno: {str(e)}")
                
        if log_errores:
            try:
                log_path = os.path.join(base_path, "errores_descarga.txt")
                with open(log_path, "w", encoding="utf-8") as f:
                    f.write("REPORTE DE DOCUMENTOS NO DESCARGADOS\n")
                    f.write("====================================\n\n")
                    for linea in log_errores:
                        f.write(linea + "\n")
            except: pass
            
        driver.quit()
        return {
            "status": "success", 
            "message": f"Finalizado. Descargados: {descargados}, Errores: {errores}, Conflictos: {conflictos}",
            "stats": {"descargados": descargados, "errores": errores, "conflictos": conflictos}
        }

    except Exception as e:
        if driver: driver.quit()
        return {"status": "error", "message": str(e)}

def process_edit_text(items, find_text, replace_text):
    count_modified = 0
    errors = []
    
    if not items:
        return {"count": 0, "errors": ["No se proporcionaron archivos para editar"]}
        
    print(f"DEBUG: Editando texto en {len(items)} archivos. Buscar: '{find_text}', Reemplazar: '{replace_text}'")
    
    text_extensions = ['.txt', '.json', '.xml', '.csv', '.html', '.md', '.log', '.py', '.js', '.css', '.bat', '.ps1', '.sql', '.ini', '.yaml', '.yml']
    
    for file_path in items:
        if not os.path.exists(file_path):
            errors.append(f"Archivo no encontrado: {file_path}")
            continue
            
        if os.path.isdir(file_path):
            continue
            
        try:
            ext = os.path.splitext(file_path)[1].lower()
            modified = False
            
            # Archivos de texto plano
            if ext in text_extensions:
                try:
                    # Try reading with utf-8 first
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    except UnicodeDecodeError:
                        # Fallback to latin-1 or system default
                        with open(file_path, 'r', encoding='latin-1') as f:
                            content = f.read()
                    
                    if find_text in content:
                        new_content = content.replace(find_text, replace_text)
                        with open(file_path, 'w', encoding='utf-8') as f:
                            f.write(new_content)
                        modified = True
                except Exception as e:
                    errors.append(f"Error leyendo/escribiendo {os.path.basename(file_path)}: {e}")

            # Documentos de Word
            elif ext == '.docx':
                try:
                    if Document is not None:
                        doc = Document(file_path)
                        doc_modified = False
                        
                        # Paragraphs
                        for p in doc.paragraphs:
                            if find_text in p.text:
                                p.text = p.text.replace(find_text, replace_text)
                                doc_modified = True
                        
                        # Tables
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        if find_text in p.text:
                                            p.text = p.text.replace(find_text, replace_text)
                                            doc_modified = True
                                            
                        if doc_modified:
                            doc.save(file_path)
                            modified = True
                    else:
                        errors.append(f"Omitido {os.path.basename(file_path)}: Librería python-docx no disponible")
                except Exception as e:
                    errors.append(f"Error procesando DOCX {os.path.basename(file_path)}: {e}")
            
            if modified:
                count_modified += 1
                
        except Exception as e:
            errors.append(f"Error general en {os.path.basename(file_path)}: {e}")
            
    return {"count": count_modified, "errors": errors}

def process_copy_files(items, target_folder):
    count = 0
    errors = []
    
    if not items:
        return {"count": 0, "errors": ["No se proporcionaron archivos para copiar"]}
        
    try:
        if not os.path.exists(target_folder):
            os.makedirs(target_folder)
    except Exception as e:
        return {"count": 0, "errors": [f"Error creando carpeta destino: {e}"]}

    for item in items:
        if not os.path.exists(item):
            errors.append(f"No encontrado: {item}")
            continue
            
        try:
            filename = os.path.basename(item)
            dest_path = os.path.join(target_folder, filename)
            
            # Manejar colisiones
            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                dest_path = os.path.join(target_folder, f"{base}_{int(time.time())}{ext}")
            
            if os.path.isdir(item):
                shutil.copytree(item, dest_path)
            else:
                shutil.copy2(item, dest_path)
            count += 1
        except Exception as e:
            errors.append(f"Error copiando {os.path.basename(item)}: {e}")
            
    return {"count": count, "errors": errors}

def process_move_files(items, target_folder):
    count = 0
    errors = []
    
    if not items:
        return {"count": 0, "errors": ["No se proporcionaron archivos para mover"]}
        
    try:
        if not os.path.exists(target_folder):
            os.makedirs(target_folder)
    except Exception as e:
        return {"count": 0, "errors": [f"Error creando carpeta destino: {e}"]}

    for item in items:
        if not os.path.exists(item):
            errors.append(f"No encontrado: {item}")
            continue
            
        try:
            filename = os.path.basename(item)
            dest_path = os.path.join(target_folder, filename)
            
            # Manejar colisiones
            if os.path.exists(dest_path):
                base, ext = os.path.splitext(filename)
                dest_path = os.path.join(target_folder, f"{base}_{int(time.time())}{ext}")
            
            shutil.move(item, dest_path)
            count += 1
        except Exception as e:
            errors.append(f"Error moviendo {os.path.basename(item)}: {e}")
            
    return {"count": count, "errors": errors}

def process_delete_files(items, force=False):
    count = 0
    errors = []
    
    if not items:
        return {"count": 0, "errors": ["No se proporcionaron archivos para eliminar"]}
        
    # Try importing send2trash if not force delete
    send2trash_func = None
    if not force:
        try:
            from send2trash import send2trash as s2t
            send2trash_func = s2t
        except ImportError:
            return {"count": 0, "errors": ["Librería send2trash no instalada en el Agente. Use borrado forzado."]}

    for item in items:
        if not os.path.exists(item):
            errors.append(f"No encontrado: {item}")
            continue
            
        try:
            if force:
                if os.path.isdir(item):
                    shutil.rmtree(item)
                else:
                    os.remove(item)
            else:
                # Send to trash
                send2trash_func(os.path.normpath(item))
            count += 1
        except Exception as e:
            errors.append(f"Error eliminando {os.path.basename(item)}: {e}")
            
    return {"count": count, "errors": errors}

def process_compress_zip(items, output_path):
    count = 0
    errors = []
    
    if not items:
        return {"count": 0, "errors": ["No se proporcionaron archivos para comprimir"]}
        
    try:
        # Create directory for zip if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for item in items:
                if not os.path.exists(item):
                    errors.append(f"No encontrado: {item}")
                    continue
                    
                if os.path.isfile(item):
                    zipf.write(item, os.path.basename(item))
                    count += 1
                elif os.path.isdir(item):
                    base_folder = os.path.basename(item)
                    for root, dirs, files in os.walk(item):
                        for file in files:
                            file_path = os.path.join(root, file)
                            # Create relative path starting with folder name
                            rel_path = os.path.relpath(file_path, os.path.dirname(item))
                            zipf.write(file_path, rel_path)
                    count += 1
                    
    except Exception as e:
        return {"count": count, "errors": [str(e)]}
        
    return {"count": count, "errors": errors}

def process_compress_individual(items):
    count = 0
    errors = []
    
    if not items:
        return {"count": 0, "errors": ["No se proporcionaron elementos para comprimir"]}
        
    for item in items:
        if not os.path.exists(item):
            errors.append(f"No encontrado: {item}")
            continue
            
        try:
            # Create zip file with same name as item
            zip_name = f"{item}.zip"
            
            with zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED) as zipf:
                if os.path.isdir(item):
                    for root, dirs, files in os.walk(item):
                        for file in files:
                            file_path = os.path.join(root, file)
                            rel_path = os.path.relpath(file_path, item)
                            zipf.write(file_path, rel_path)
                else:
                    zipf.write(item, os.path.basename(item))
            count += 1
            
        except Exception as e:
            errors.append(f"Error comprimiendo {os.path.basename(item)}: {e}")
            
    return {"count": count, "errors": errors}

class AgentWorker:
    def __init__(self, username, task_url, result_url, password=None):
        self.username = username
        self.task_url = task_url
        self.result_url = result_url
        self.password = password
        self.running = False
        self.thread = None
        self.log_callback = None
        self.gui_invoker = None

    def log(self, message):
        print(message)
        logging.info(message)
        if self.log_callback:
            self.log_callback(message)

    def start(self):
        self.running = True
        self.thread = threading.Thread(target=self.run_loop)
        self.thread.daemon = True
        self.thread.start()
        self.log("Worker iniciado.")

    def stop(self):
        self.running = False
        if self.thread:
            self.thread.join(timeout=2)
        self.log("Worker detenido.")

    def run_loop(self):
        while self.running:
            try:
                # Poll for tasks with increased timeout
                auth = (self.username, self.password) if self.password else None
                resp = requests.get(f"{self.task_url}?username={self.username}", auth=auth, timeout=15)
                if resp.status_code == 200:
                    data = resp.json()
                    if data.get("tasks"):
                        for task in data["tasks"]:
                            self.process_task(task)
            except Exception as e:
                self.log(f"Error polling tasks: {e}")
                # pass
            
            time.sleep(2)

    def process_task(self, task):
        task_id = task.get("id")
        command = task.get("command")
        params = task.get("params", {})
        
        self.log(f"Procesando tarea {task_id}: {command}")
        
        result = {"status": "COMPLETED", "result": None}
        
        try:
            if command == "ping":
                result["result"] = "pong"
                
            elif command == "update_cups":
                folder = params.get("path")
                old_val = params.get("old_val")
                new_val = params.get("new_val")
                
                if folder and old_val is not None and new_val is not None:
                    res = process_update_cups(folder, old_val, new_val)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "update_key":
                folder = params.get("path")
                key_target = params.get("key")
                new_value = params.get("value", params.get("val"))
                
                if folder and key_target and new_value is not None:
                    res = process_update_key(folder, key_target, new_value)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "update_notes":
                path = params.get("path")
                target = params.get("target")
                note = params.get("note")
                if path:
                    self.log(f"Actualizando Notas en: {path}")
                    result["result"] = process_update_notes(path, target, note)
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta path"}
                    
            elif command == "clean_json":
                path = params.get("path")
                if path:
                    self.log(f"Limpiando JSONs en: {path}")
                    result["result"] = process_clean_json(path)
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta path"}
                    
            elif command == "consolidate_json":
                path = params.get("path")
                if path:
                    self.log(f"Consolidando JSONs en: {path}")
                    result["result"] = process_consolidate_json(path)
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta path"}

            elif command == "desconsolidate_json":
                file_path = params.get("file_path")
                dest_folder = params.get("dest_folder")
                if file_path and dest_folder:
                    self.log(f"Desconsolidando {file_path} a {dest_folder}")
                    result["result"] = process_desconsolidate_json(file_path, dest_folder)
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}
            
            elif command == "flat_to_excel":
                path = params.get("path")
                if path:
                    self.log(f"Convirtiendo planos en: {path}")
                    result["result"] = process_flat_to_excel(path)
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta path"}
            
            elif command == "bulk_rename":
                path = params.get("path")
                items = params.get("items", [])
                separator = params.get("separator", "_")
                item_type = params.get("item_type", "both")
                rename_folders = params.get("rename_folders", True)
                rename_internal_files = params.get("rename_internal_files", True)
                
                if path:
                    res = process_bulk_rename(path, items, separator, item_type, rename_folders, rename_internal_files)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta path"}

            elif command == "rename_folders_mapped":
                path = params.get("path")
                mapping = params.get("mapping", {})
                
                if path and mapping:
                    res = process_rename_folders_mapped(path, mapping)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "organize_files_mapped":
                source_path = params.get("source_path")
                dest_path = params.get("dest_path")
                mapping = params.get("mapping", {})
                
                if source_path and dest_path and mapping:
                    res = process_organize_files_mapped(source_path, dest_path, mapping)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "create_folders":
                folders = params.get("folders", [])
                
                if folders:
                    res = process_create_folders(folders)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "search_files":
                path = params.get("path")
                patterns = params.get("patterns", [])
                exclusion_list = params.get("exclusion_list", [])
                search_by = params.get("search_by", "name")
                item_type = params.get("item_type", "both")
                recursive = params.get("recursive", True)
                search_empty_folders = params.get("search_empty_folders", False)
                
                self.log(f"Recibida tarea de búsqueda: Path='{path}', Patterns='{patterns}', ItemType='{item_type}'")
                
                if path:
                    res = process_search_files(path, patterns, exclusion_list, search_by, item_type, recursive, search_empty_folders)
                    # res is already a dict {"items": [...], "errors": [...]}
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta el parámetro 'path'"}

            elif command == "fill_docx":
                base_path = params.get("base_path")
                tasks = params.get("tasks", [])
                template_b64 = params.get("template_b64")
                
                if not base_path or not os.path.exists(base_path):
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Ruta base inválida o no encontrada"}
                else:
                    self.log(f"Iniciando llenado de DOCX en: {base_path} ({len(tasks)} tareas)")
                    res = process_fill_docx(base_path, tasks, template_b64)
                    result["result"] = res

            elif command == "fill_docx_ovida_full":
                base_path = params.get("base_path")
                tasks = params.get("tasks", [])
                
                if not base_path or not os.path.exists(base_path):
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Ruta base inválida o no encontrada"}
                else:
                    self.log(f"Iniciando llenado de DOCX OVIDA en: {base_path} ({len(tasks)} tareas)")
                    res = process_fill_docx_ovida_full(base_path, tasks)
                    result["result"] = res

            elif command == "sign_docx_massive":
                base_path = params.get("base_path")
                docx_filename = params.get("docx_filename")
                signature_filename = params.get("signature_filename")
                
                if not base_path or not os.path.exists(base_path):
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Ruta base inválida o no encontrada"}
                else:
                    self.log(f"Iniciando firma de DOCX masiva en: {base_path}")
                    res = process_sign_docx_massive(base_path, docx_filename, signature_filename)
                    result["result"] = res

            elif command == "distribute_file":
                paths = params.get("paths", [])
                content_b64 = params.get("content_b64")
                
                if paths and content_b64:
                    self.log(f"Distribuyendo archivo a {len(paths)} destinos")
                    res = process_distribute_file(paths, content_b64)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "write_files":
                files = params.get("files", [])
                
                if files:
                    self.log(f"Escribiendo {len(files)} archivos")
                    res = process_write_files(files)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "rename_files":
                files = params.get("files", [])
                
                if files:
                    self.log(f"Renombrando {len(files)} archivos")
                    res = process_rename_files(files)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros"}

            elif command == "distribute_files":
                content_b64 = params.get("content_b64")
                paths = params.get("paths", [])
                
                if content_b64 and paths:
                    self.log(f"Distribuyendo archivo a {len(paths)} destinos")
                    res = process_distribute_files(content_b64, paths)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (content_b64, paths)"}

            elif command == "copiar_archivo_a_subcarpetas":
                archivo = params.get("archivo")
                carpeta_base = params.get("carpeta_base")
                
                if archivo and carpeta_base:
                    self.log(f"Copiando archivo a subcarpetas: {archivo}")
                    res = process_copiar_archivo_a_subcarpetas(archivo, carpeta_base)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (archivo, carpeta_base)"}

            elif command == "download_files":
                tasks = params.get("tasks", [])
                
                if tasks:
                    self.log(f"Descargando {len(tasks)} archivos")
                    res = process_download_files(tasks)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (tasks)"}

            elif command == "validate_rips":
                base_path = params.get("base_path")
                api_url = params.get("api_url")
                token = params.get("token")
                verify_ssl = params.get("verify_ssl", True)
                
                if not base_path or not api_url:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (path o api_url)"}
                else:
                    self.log(f"Validando RIPS en: {base_path}")
                    res = process_validate_rips(base_path, api_url, token, verify_ssl)
                    result["result"] = res

            elif command == "list_files":
                path = params.get("path")
                if path:
                    self.log(f"Listando archivos en: {path}")
                    try:
                        if os.path.exists(path) and os.path.isdir(path):
                            items = []
                            # Limit to 500 items to avoid payload issues
                            count = 0
                            with os.scandir(path) as it:
                                for entry in it:
                                    if count > 500: break
                                    items.append({
                                        "name": entry.name,
                                        "is_dir": entry.is_dir(),
                                        "size": entry.stat().st_size if not entry.is_dir() else 0,
                                        "mtime": entry.stat().st_mtime
                                    })
                                    count += 1
                            result["result"] = {"files": items, "count": count}
                        else:
                            result["status"] = "ERROR"
                            result["result"] = {"error": "Ruta no encontrada o no es carpeta"}
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta path"}

            elif command == "browse_folder":
                title = params.get("title", "Seleccionar Carpeta")
                self.log(f"Abriendo selector de carpeta: {title}")
                
                try:
                    import tkinter.filedialog as fd
                    
                    path = None
                    if self.gui_invoker:
                        # Use the GUI thread invoker if available
                        path = self.gui_invoker(fd.askdirectory, title=title)
                    else:
                        # Fallback (unsafe)
                        path = fd.askdirectory(title=title)
                        
                    if path:
                        result["result"] = {"path": path}
                    else:
                        result["result"] = {"cancelled": True}
                        
                except Exception as e:
                    result["status"] = "ERROR"
                    result["result"] = {"error": str(e)}

            elif command == "browse_file":
                title = params.get("title", "Seleccionar Archivo")
                file_types = params.get("file_types", [])
                self.log(f"Abriendo selector de archivo: {title}")
                
                try:
                    import tkinter.filedialog as fd
                    # Convert file_types from list of lists to list of tuples if needed
                    # Streamlit sends: [[label, pattern], ...]
                    # Tkinter expects: [(label, pattern), ...]
                    ft = []
                    if file_types:
                        for item in file_types:
                            if isinstance(item, list) and len(item) >= 2:
                                ft.append((item[0], item[1]))
                            elif isinstance(item, tuple):
                                ft.append(item)
                    
                    if not ft:
                        ft = [("Todos los archivos", "*.*")]
                        
                    path = None
                    if self.gui_invoker:
                        path = self.gui_invoker(fd.askopenfilename, title=title, filetypes=ft)
                    else:
                        path = fd.askopenfilename(title=title, filetypes=ft)

                    if path:
                        result["result"] = {"path": path}
                    else:
                        result["result"] = {"cancelled": True}
                except Exception as e:
                    result["status"] = "ERROR"
                    result["result"] = {"error": str(e)}

            elif command == "generate_signature":
                text = params.get("text")
                font_name = params.get("font_name", "Pacifico")
                size = params.get("size", 70)
                width = params.get("width", 500)
                height = params.get("height", 200)
                
                if text:
                    self.log(f"Generando firma para: {text}")
                    res = process_generate_signature(text, font_name, size, width, height)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta texto para la firma"}

            elif command == "download_ovida":
                base_path = params.get("base_path")
                records = params.get("records", [])
                
                if base_path and records:
                    self.log(f"Descargando historias de OVIDA ({len(records)} registros)")
                    res = process_download_ovida(base_path, records)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (path o records)"}

            elif command == "download_zeus_adjuntos":
                base_path = params.get("base_path")
                records = params.get("records", [])
                
                if base_path and records:
                    self.log(f"Descargando adjuntos de ZeusSalud ({len(records)} registros)")
                    res = process_download_zeus_adjuntos(base_path, records)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (path o records)"}

            elif command == "create_folders_from_list":
                base_path = params.get("base_path")
                names = params.get("names", [])
                unique = params.get("unique", False)
                
                if base_path and names:
                    self.log(f"Creando {len(names)} carpetas en: {base_path}")
                    res = process_create_folders_from_list(base_path, names, unique)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (base_path o names)"}

            elif command == "edit_text":
                items = params.get("items", [])
                find_text = params.get("find", "")
                replace_text = params.get("replace", "")
                
                if items and find_text:
                    self.log(f"Editando texto en {len(items)} archivos")
                    res = process_edit_text(items, find_text, replace_text)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (items, find)"}

            elif command == "delete_files":
                items = params.get("items", [])
                force = params.get("force", False)
                
                if items:
                    self.log(f"Eliminando {len(items)} archivos (Forzado: {force})")
                    res = process_delete_files(items, force)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (items)"}

            elif command == "compress_zip":
                items = params.get("items", [])
                output_path = params.get("output_path")
                
                if items and output_path:
                    self.log(f"Comprimiendo {len(items)} elementos en: {output_path}")
                    res = process_compress_zip(items, output_path)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (items, output_path)"}

            elif command == "compress_individual":
                items = params.get("items", [])
                
                if items:
                    self.log(f"Comprimiendo individualmente {len(items)} elementos")
                    res = process_compress_individual(items)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (items)"}

            elif command == "copy_files":
                items = params.get("items", [])
                target_folder = params.get("target_folder")
                
                if items and target_folder:
                    self.log(f"Copiando {len(items)} archivos a {target_folder}")
                    res = process_copy_files(items, target_folder)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (items, target_folder)"}

            elif command == "move_files":
                items = params.get("items", [])
                target_folder = params.get("target_folder")
                
                if items and target_folder:
                    self.log(f"Moviendo {len(items)} archivos a {target_folder}")
                    res = process_move_files(items, target_folder)
                    result["result"] = res
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (items, target_folder)"}

            elif command == "consolidar_subcarpetas":
                root_path = params.get("root_path")
                if root_path:
                    # Ensure src can be imported
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_consolidar_subcarpetas
                        # Force local execution mode within agent
                        os.environ["CDO_AGENT_MODE"] = "1"
                        res = worker_consolidar_subcarpetas(root_path, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (root_path)"}

            elif command == "unify_pdf_folder":
                base_path = params.get("base_path")
                output_name = params.get("output_name", "Unificado")
                if base_path:
                    # Ensure src can be imported
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_unificar_por_carpeta
                        res = worker_unificar_por_carpeta(base_path, output_name, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (base_path)"}

            elif command == "unify_img_folder":
                base_path = params.get("base_path")
                output_name = params.get("output_name", "Unificado.pdf")
                img_type = params.get("img_type", "JPG")
                if base_path:
                    # Ensure src can be imported
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_unificar_imagenes_por_carpeta_rec
                        res = worker_unificar_imagenes_por_carpeta_rec(base_path, output_name, img_type, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (base_path)"}

            elif command == "unify_docx_folder":
                base_path = params.get("base_path")
                output_name = params.get("output_name", "Unificado")
                if base_path:
                    # Ensure src can be imported
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_unificar_docx_por_carpeta
                        res = worker_unificar_docx_por_carpeta(base_path, output_name, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (base_path)"}

            elif command == "split_pdf_massive":
                base_path = params.get("base_path")
                if base_path:
                    # Ensure src can be imported
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_dividir_pdfs_masivamente
                        res = worker_dividir_pdfs_masivamente(base_path, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (base_path)"}

            elif command == "convert_file":
                file_path = params.get("file_path")
                ctype = params.get("type")
                out_folder = params.get("output_folder")
                sep = params.get("sep", ",")
                if file_path and ctype:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        os.environ["CDO_AGENT_MODE"] = "1"
                        from src.tabs.tab_conversion import worker_convertir_archivo
                        ok, msg = worker_convertir_archivo(file_path, ctype, output_folder=out_folder, sep=sep)
                        result["result"] = {"ok": ok, "message": msg}
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (file_path, type)"}

            elif command == "convert_bulk":
                folder_path = params.get("folder_path")
                ctype = params.get("type")
                out_folder = params.get("output_folder")
                sep = params.get("sep", ",")
                save_in_same_dir = params.get("save_in_same_dir", False)
                if folder_path and ctype:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        os.environ["CDO_AGENT_MODE"] = "1"
                        from src.tabs.tab_conversion import worker_convertir_masivo
                        res = worker_convertir_masivo(folder_path, ctype, output_folder=out_folder, sep=sep, return_zip=False, save_in_same_dir=save_in_same_dir)
                        if isinstance(res, tuple):
                            count, msg = res
                            result["result"] = {"count": count, "message": msg, "errors": []}
                        elif isinstance(res, dict):
                            # Already in dict form
                            result["result"] = {k: v for k, v in res.items() if k in ("count", "message", "errors")}
                        else:
                            result["result"] = {"count": 0, "message": f"Respuesta inesperada: {res}", "errors": []}
                    except Exception as e:
                        result["status"] = "ERROR"
                        result["result"] = {"error": str(e)}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Faltan parámetros (folder_path, type)"}

            elif command == "analisis_carpetas":
                path = params.get("path")
                if path:
                    # Ensure src can be imported
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    
                    try:
                        from src.tabs.tab_automated_actions import worker_analisis_carpetas
                        res = worker_analisis_carpetas(path, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (path)"}

            elif command == "analisis_sos":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                use_ai = params.get("use_ai", False)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.modules.analisis_sos import worker_analisis_sos
                        res = worker_analisis_sos(files, use_ai=use_ai, silent_mode=True)
                        if isinstance(res, tuple):
                            out_xlsx, out_txt = res
                            result["result"] = {
                                "files": [
                                    {"name": "Analisis_SOS.xlsx", "data": _encode_bytes(out_xlsx), "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel"},
                                    {"name": "Analisis_SOS.txt", "data": _encode_bytes(out_txt), "mime": "text/csv", "label": "CSV/TXT"}
                                ],
                                "message": "Análisis SOS completado."
                            }
                        elif res:
                            result["result"] = {
                                "files": [{"name": "Analisis_SOS.xlsx", "data": _encode_bytes(res), "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "label": "Excel"}],
                                "message": "Análisis SOS completado."
                            }
                    except Exception as e:
                        self.log(f"Error importando worker SOS: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}

            elif command == "analisis_hc":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_analisis_historia_clinica
                        res = worker_analisis_historia_clinica(files, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker HC: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}

            elif command == "analisis_neps":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_analisis_autorizacion_nueva_eps
                        res = worker_analisis_autorizacion_nueva_eps(files, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker NEPS: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}

            elif command == "analisis_sanitas":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_analisis_cargue_sanitas
                        res = worker_analisis_cargue_sanitas(files, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker Sanitas: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}

            elif command == "analisis_rete":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_leer_pdf_retefuente
                        res = worker_leer_pdf_retefuente(files, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker Retefuente: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}

            elif command == "analisis_emssanar":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_analisis_emssanar
                        res = worker_analisis_emssanar(files, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker Emssanar: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}

            elif command == "analisis_fomag":
                files_param = params.get("file_list", params.get("files", []))
                files = _expand_file_list(files_param)
                if files:
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    try:
                        from src.tabs.tab_automated_actions import worker_analisis_fomag
                        res = worker_analisis_fomag(files, silent_mode=True)
                        result["result"] = _serialize_analysis_result(res)
                    except Exception as e:
                        self.log(f"Error importando worker FOMAG: {e}\n{traceback.format_exc()}")
                        result["status"] = "ERROR"
                        result["result"] = {"error": f"Error interno: {e}"}
                else:
                    result["status"] = "ERROR"
                    result["result"] = {"error": "Falta parámetro (files)"}
                    
            elif command == "launch_browser":
                url = params.get("url", "https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")
                try:
                    # Ensure bot_zeus can be imported and launch asynchronously
                    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
                    
                    def _open_browser():
                        try:
                            from src import bot_zeus
                            bot_zeus.abrir_navegador_inicial()
                        except Exception as e_inner:
                            self.log(f"Error abriendo navegador (hilo): {e_inner}\n{traceback.format_exc()}")
                    
                    threading.Thread(target=_open_browser, daemon=True).start()
                    result["status"] = "success"
                    result["result"] = {"message": "Navegador lanzado en el Agente Local."}
                except Exception as e:
                    self.log(f"Error abriendo navegador: {e}\n{traceback.format_exc()}")
                    result["status"] = "ERROR"
                    result["result"] = {"error": str(e)}

            else:
                result["status"] = "ERROR"
                result["result"] = {"error": f"Comando desconocido: {command}"}

        except Exception as e:
            result["status"] = "ERROR"
            result["result"] = {"error": str(e)}
            self.log(f"Error procesando tarea: {e}")

        # Send result
        try:
            # FIX: Use RESTful endpoint for AWS server connection
            # Server expects POST /tasks/{task_id}/result
            base_url = self.result_url.rstrip("/")
            post_url = f"{base_url}/{task_id}/result"
            auth = (self.username, self.password) if self.password else None
            
            resp = requests.post(post_url, json={
                "status": result["status"],
                "result": result["result"]
            }, auth=auth)
            
            if resp.status_code == 200:
                self.log(f"Resultado enviado para {task_id}")
            else:
                self.log(f"Error enviando resultado {task_id}: {resp.status_code} - {resp.text}")
        except Exception as e:
            self.log(f"Error enviando resultado: {e}")

# --- GUI ---
def load_config():
    # 1. Check current directory (priority for portable/dev)
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r') as f:
                return json.load(f)
        except:
            pass
            
    # 2. Check LOCALAPPDATA (installed location)
    try:
        app_data = os.getenv('LOCALAPPDATA', os.path.expanduser("~"))
        config_path = os.path.join(app_data, 'CDO_Organizer', 'agent_config.json')
        if os.path.exists(config_path):
            with open(config_path, 'r') as f:
                return json.load(f)
    except:
        pass
        
    return {}

def save_config(config):
    try:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f)
    except Exception as e:
        print(f"Advertencia: No se pudo guardar la configuración en {CONFIG_FILE}: {e}")

class AgentGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Agente Local - Organizador Archivos")
        self.root.geometry("500x400")
        
        self.config = load_config()
        self.worker = None
        
        # UI Elements
        frame = ttk.Frame(root, padding="10")
        frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frame, text="Usuario:").grid(row=0, column=0, sticky=tk.W)
        self.user_var = tk.StringVar(value=self.config.get("username", ""))
        ttk.Entry(frame, textvariable=self.user_var).grid(row=0, column=1, sticky=tk.EW)
        
        ttk.Label(frame, text="Contraseña:").grid(row=1, column=0, sticky=tk.W)
        self.pass_var = tk.StringVar(value=self.config.get("password", ""))
        ttk.Entry(frame, textvariable=self.pass_var, show="*").grid(row=1, column=1, sticky=tk.EW)

        ttk.Label(frame, text="URL Tareas:").grid(row=2, column=0, sticky=tk.W)
        # Fix: Use correct AWS endpoint (port 8000, /tasks/poll)
        default_task_url = "http://3.138.135.181:8000/tasks/poll"
        # Heuristic: if loaded config has "localhost" or "8501" or "/api/", replace with correct default
        loaded_task_url = self.config.get("task_url", default_task_url)
        if "localhost" in loaded_task_url or "8501" in loaded_task_url or "/api/" in loaded_task_url or "3.142.164.128" in loaded_task_url or "3.15.237.186" in loaded_task_url or "18.118.37.215" in loaded_task_url:
             loaded_task_url = default_task_url
             
        self.url_task_var = tk.StringVar(value=loaded_task_url)
        ttk.Entry(frame, textvariable=self.url_task_var).grid(row=2, column=1, sticky=tk.EW)
        
        ttk.Label(frame, text="URL Resultados:").grid(row=3, column=0, sticky=tk.W)
        default_res_url = "http://3.138.135.181:8000/tasks"
        loaded_res_url = self.config.get("result_url", default_res_url)
        if "localhost" in loaded_res_url or "8501" in loaded_res_url or "/api/" in loaded_res_url or "3.142.164.128" in loaded_res_url or "3.15.237.186" in loaded_res_url or "18.118.37.215" in loaded_res_url:
             loaded_res_url = default_res_url
             
        self.url_res_var = tk.StringVar(value=loaded_res_url)
        ttk.Entry(frame, textvariable=self.url_res_var).grid(row=3, column=1, sticky=tk.EW)
        
        self.btn_start = ttk.Button(frame, text="Iniciar Agente", command=self.toggle_agent)
        self.btn_start.grid(row=4, column=0, columnspan=2, pady=10)
        
        self.log_area = scrolledtext.ScrolledText(frame, height=15)
        self.log_area.grid(row=5, column=0, columnspan=2, sticky=tk.NSEW)
        
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(5, weight=1)

    def log(self, msg):
        def _log():
            self.log_area.insert(tk.END, f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n")
            self.log_area.see(tk.END)
        self.root.after(0, _log)

    def invoke_on_gui(self, func, *args, **kwargs):
        """
        Executes a function on the main GUI thread and returns the result.
        This is thread-safe and blocking for the caller.
        """
        result_container = {}
        event = threading.Event()
        
        def wrapper():
            try:
                result_container["result"] = func(*args, **kwargs)
            except Exception as e:
                result_container["error"] = e
            finally:
                event.set()
        
        self.root.after(0, wrapper)
        event.wait()
        
        if "error" in result_container:
            raise result_container["error"]
        return result_container.get("result")

    def toggle_agent(self):
        if self.worker and self.worker.running:
            self.worker.stop()
            self.btn_start.config(text="Iniciar Agente")
            self.log("Agente detenido manualmente.")
        else:
            user = self.user_var.get()
            password = self.pass_var.get()
            t_url = self.url_task_var.get()
            r_url = self.url_res_var.get()
            
            if not user or not t_url:
                messagebox.showerror("Error", "Configure usuario y URLs")
                return
                
            self.config["username"] = user
            self.config["password"] = password
            self.config["task_url"] = t_url
            self.config["result_url"] = r_url
            save_config(self.config)
            
            self.worker = AgentWorker(user, t_url, r_url, password)
            self.worker.log_callback = self.log
            self.worker.gui_invoker = self.invoke_on_gui
            self.worker.start()
            self.btn_start.config(text="Detener Agente")

if __name__ == "__main__":
    root = tk.Tk()
    app = AgentGUI(root)
    root.mainloop()
