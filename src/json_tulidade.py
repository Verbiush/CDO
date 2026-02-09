import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import os
import shutil
import re
import pandas as pd # type: ignore
import threading
from docx import Document # type: ignore
from docx.shared import Inches # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
import json
from send2trash import send2trash # type: ignore
import fitz  # type: ignore # PyMuPDF
from PIL import Image, ImageDraw, ImageFont # type: ignore
import io
from zipfile import ZipFile, ZIP_DEFLATED
import string
from pdf2docx import Converter # type: ignore # type: ignore
from docx2pdf import convert # type: ignore
import fnmatch
import openpyxl  # type: ignore
import xml.etree.ElementTree as ET
import requests # type: ignore # type: ignore
import subprocess
import win32com.client # type: ignore
# EL 'import pyperclip' SE MUEVE DENTRO DE LA CLASE PARA MANEJAR EL ERROR

# === INICIO NUEVA FUNCIONALIDAD OVIDA ===
import time
import urllib.parse
import base64
# Las importaciones de Selenium se manejan dentro de un try/except para no romper la app si no están instaladas
# === FIN NUEVA FUNCIONALIDAD OVIDA ===

# Importar la clase del módulo RIPS
try:
    from rips_module import OrganizadorArchivosApp as RipsApp # type: ignore
except ImportError:
    RipsApp = None # Manejar el caso si el archivo no está presente
    print("Advertencia: No se encontró el módulo 'rips_module.py'. La pestaña RIPS no estará disponible.")


class MainApp: # Renombramos la clase principal para evitar conflicto
    def __init__(self, root):
        self.root = root
        
        # ### CORRECCIÓN: La importación ahora está dentro del try/except ###
        try:
            import pyperclip # type: ignore # type: ignore
        except ImportError:
            messagebox.showwarning("Dependencia Faltante", 
                                   "La librería 'pyperclip' no está instalada.\n\n"
                                   "La función de 'Copiar al portapapeles' en el editor de datos no estará disponible.\n\n"
                                   "Puedes instalarla con: pip install pyperclip",
                                   parent=self.root)
            self.pyperclip = None
            self.pyperclip_available = False
        else:
            self.pyperclip = pyperclip
            self.pyperclip_available = True

        # === INICIO NUEVA FUNCIONALIDAD OVIDA ===
        # Comprobar dependencias de Selenium
        try:
            from selenium import webdriver # type: ignore
            from selenium.webdriver.common.by import By # type: ignore
            from selenium.webdriver.support.ui import WebDriverWait, Select # type: ignore
            from selenium.webdriver.support import expected_conditions as EC # type: ignore
            from selenium.webdriver.chrome.service import Service # type: ignore
            from selenium.webdriver.common.keys import Keys # type: ignore
            from webdriver_manager.chrome import ChromeDriverManager # type: ignore
            self.selenium_available = True
            self.selenium_deps = {
                "webdriver": webdriver, "By": By, "WebDriverWait": WebDriverWait,
                "Select": Select, "EC": EC, "Service": Service, "Keys": Keys,
                "ChromeDriverManager": ChromeDriverManager
            }
        except ImportError:
            self.selenium_available = False
        # === FIN NUEVA FUNCIONALIDAD OVIDA ===

        style = ttk.Style(self.root)
        style.theme_use("vista")

        self.root.title("📂 Organizador de Archivos v5.7 (Mapeo Parcial)")
        self.root.geometry("1100x850")
       
        # Variables de la Tab 1 (Búsqueda)
        self.source_path = tk.StringVar()
        self.search_pattern = tk.StringVar()
        self.filter_type = tk.StringVar(value="extensión")
        self.include_folders = tk.BooleanVar(value=True)
        self.element_type = tk.StringVar(value="archivos")
        self.accion = tk.StringVar(value="copiar a carpeta")
        self.resultados = []

        # --- Variables para el editor de datos (Tab 4) ---
        self.data_editor_path = tk.StringVar()
        self.data_tree = None
        self.save_button = None
        self.parsed_data = None
        self.current_file_path = None
        self.current_data_type = None
        self.is_data_modified = False
        self.embedded_xml_docs = {}
        self.tooltip_window = None 
        self.xml_element_map = {}
        
        # --- Variables para la búsqueda en el Treeview ---
        self.search_data_var = tk.StringVar()
        self.search_results_iids = []

        self.data_context_menu = None
        
        self.create_widgets()
        self._on_filter_type_change()

    def _create_progress_window(self, title, max_value):
        """Crea y devuelve una ventana emergente con una barra de progreso."""
        progress_win = tk.Toplevel(self.root)
        progress_win.title(title)
        progress_win.geometry("400x120")
        progress_win.resizable(False, False)
        progress_win.transient(self.root)
        progress_win.grab_set()

        win_x = self.root.winfo_x() + (self.root.winfo_width() // 2) - 200
        win_y = self.root.winfo_y() + (self.root.winfo_height() // 2) - 60
        progress_win.geometry(f'+{win_x}+{win_y}')

        progress_label = ttk.Label(progress_win, text="Iniciando...", padding=(10, 10))
        progress_label.pack()

        progress_bar = ttk.Progressbar(
            progress_win,
            orient='horizontal',
            length=360,
            mode='determinate',
            maximum=max_value,
            value=0
        )
        progress_bar.pack(pady=10, padx=20)
        
        return progress_win, progress_bar, progress_label
    
    def _on_filter_type_change(self, event=None):
        """Actualiza la etiqueta del patrón según el tipo de filtro seleccionado."""
        filtro = self.filter_type.get()
        if filtro == "Patrón con comodín (*, ?)":
            self.pattern_label.config(text="Patrón (ej: prueba.docx):")
        else:
            self.pattern_label.config(text="Patrón:")

    def create_widgets(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(pady=10, padx=10, fill="both", expand=True)
        
        # ================== TAB 1: Búsqueda y Acciones ==================
        self._setup_search_tab(notebook)
        
        # ================== TAB 2: Acciones Automatizadas ==================
        self._setup_automated_actions_tab(notebook)
        
        # ================== TAB 3: Conversión Individual ==================
        self._setup_individual_conversion_tab(notebook)
            
        # ================== TAB 4: Visor y Editor de Datos ==================
        self._setup_data_editor_tab(notebook)

        # ================== TAB RIPS (Nueva Pestaña) ==================
        if RipsApp: # Si el módulo RIPS se importó correctamente
            self._setup_rips_tab(notebook)
        else:
            # Opcional: Crear una pestaña de aviso si RipsApp no está disponible
            tab_rips_placeholder = ttk.Frame(notebook)
            notebook.add(tab_rips_placeholder, text="RIPS")
            ttk.Label(tab_rips_placeholder, text="El módulo RIPS no se encontró o no se pudo importar.").pack(padx=20, pady=20)
            ttk.Label(tab_rips_placeholder, text="Por favor, asegúrate de que 'rips_module.py' está en el mismo directorio y tiene las dependencias necesarias.").pack(padx=20)

        # ================== Log de Actividad (al final) ==================
        log_frame = ttk.LabelFrame(self.root, text="🧾 Registro de actividad", padding=10)
        log_frame.pack(side="bottom", fill="x", padx=10, pady=(0,10))
        self.log_text = tk.Text(log_frame, height=8, wrap="word", state="disabled")
        self.log_text.pack(fill="both", expand=True)
        log_frame.columnconfigure(0, weight=1)
        
    # Métodos para configurar cada pestaña
    def _setup_search_tab(self, notebook):
        tab1 = ttk.Frame(notebook)
        notebook.add(tab1, text="🔎 Búsqueda y Acciones Principales")
        tab1.columnconfigure(0, weight=1)
        tab1.columnconfigure(1, weight=1)
        tab1.rowconfigure(2, weight=1)

        origen_frame = ttk.LabelFrame(tab1, text="📁 Carpeta de origen", padding=10)
        origen_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=5, padx=5)
        origen_frame.columnconfigure(0, weight=1)
        ttk.Entry(origen_frame, textvariable=self.source_path).grid(row=0, column=0, sticky="ew", padx=(0, 5))
        ttk.Button(origen_frame, text="Examinar", command=self.seleccionar_directorio).grid(row=0, column=1)

        filtro_frame = ttk.LabelFrame(tab1, text="🔎 Criterios de búsqueda", padding=10)
        filtro_frame.grid(row=1, column=0, sticky="nsew", pady=5, padx=5)
        ttk.Label(filtro_frame, text="Buscar por:").grid(row=0, column=0, sticky="w")
        
        filtro_combobox = ttk.Combobox(filtro_frame, textvariable=self.filter_type, 
                                     values=["extensión", "nombre", "Patrón con comodín (*, ?)", "expresión regular"], 
                                     state="readonly")
        filtro_combobox.grid(row=0, column=1, sticky="ew", padx=5)
        filtro_combobox.bind("<<ComboboxSelected>>", self._on_filter_type_change)

        self.pattern_label = ttk.Label(filtro_frame, text="Patrón:")
        self.pattern_label.grid(row=1, column=0, sticky="w", pady=(5, 0))
        ttk.Entry(filtro_frame, textvariable=self.search_pattern).grid(row=1, column=1, sticky="ew", padx=5, pady=(5, 0))
        
        ttk.Checkbutton(filtro_frame, text="Incluir subcarpetas", variable=self.include_folders).grid(row=2, column=0, columnspan=2, sticky="w", pady=(5, 0))
        ttk.Label(filtro_frame, text="Tipo de elemento:").grid(row=3, column=0, sticky="w", pady=(5, 0))
        ttk.Combobox(filtro_frame, textvariable=self.element_type, values=["archivos", "carpetas", "todos"], state="readonly").grid(row=3, column=1, sticky="ew", padx=5, pady=(5, 0))

        accion_frame = ttk.LabelFrame(tab1, text="🛠️ Acción a realizar", padding=10)
        accion_frame.grid(row=1, column=1, sticky="nsew", pady=5, padx=5)
        acciones_principales = ["Copiar a carpeta", "Mover a carpeta", "Modificar nombre", "Editar texto", "Comprimir archivos en ZIP", "Comprimir carpetas individualmente"]
        for i, texto in enumerate(acciones_principales):
            ttk.Radiobutton(accion_frame, text=texto, variable=self.accion, value=texto.lower()).grid(row=i, column=0, sticky="w")

        resultado_frame = ttk.LabelFrame(tab1, text="📄 Archivos encontrados", padding=10)
        resultado_frame.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=5, padx=5)
        resultado_frame.columnconfigure(0, weight=1)
        resultado_frame.rowconfigure(0, weight=1)
        self.tree = ttk.Treeview(resultado_frame, columns=("ruta", "fecha"), show="headings")
        self.tree.heading("ruta", text="Ruta completa")
        self.tree.heading("fecha", text="Fecha")
        self.tree.column("ruta", width=600)
        self.tree.column("fecha", width=150)
        self.tree.grid(row=0, column=0, sticky="nsew")

        acciones_inf_frame = ttk.Frame(tab1, padding=5)
        acciones_inf_frame.grid(row=3, column=0, columnspan=2, sticky="ew", padx=5)
        ttk.Button(acciones_inf_frame, text="Buscar archivos", command=self.buscar_archivos).pack(side="left", padx=(0, 10))
        ttk.Button(acciones_inf_frame, text="Ejecutar acción", command=self.ejecutar_accion).pack(side="left", padx=(0, 10))
        ttk.Button(acciones_inf_frame, text="Limpiar", command=self.limpiar).pack(side="left")
        ttk.Button(acciones_inf_frame, text="🗑️ Eliminar resultados", command=self.eliminar_resultados).pack(side="left", padx=(10, 0))
        ttk.Button(acciones_inf_frame, text="🚪 Salir", command=self.root.quit).pack(side="right")

    def _setup_automated_actions_tab(self, notebook):
        tab2 = ttk.Frame(notebook)
        notebook.add(tab2, text="⚙️ Acciones Automatizadas y por Lotes")

        auto_col1 = ttk.Frame(tab2, padding=10)
        auto_col1.pack(side="left", fill="y", anchor="n", padx=5)
        auto_col2 = ttk.Frame(tab2, padding=10)
        auto_col2.pack(side="left", fill="y", anchor="n", padx=5)
        auto_col3 = ttk.Frame(tab2, padding=10)
        auto_col3.pack(side="left", fill="y", anchor="n", padx=5)
        
        accion_personalizada_frame = ttk.LabelFrame(auto_col1, text="Unificación y División Especial", padding=10)
        accion_personalizada_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(accion_personalizada_frame, text="🗂️ Unificar PDF por Carpeta", command=self.accion_unificar_por_carpeta).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="🖼️ Unificar JPG por Carpeta", command=self.accion_unificar_jpg_por_carpeta).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="🧷 Unificar PDFs (Selección Manual)", command=self.accion_unificar_pdfs).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="✂️ Dividir PDF en Páginas (Manual)", command=self.accion_dividir_pdf_en_paginas).pack(anchor="w", pady=2, fill="x")
        ttk.Button(accion_personalizada_frame, text="✂️ Dividir PDFs Masivamente (por Carpeta)", command=self.accion_dividir_pdfs_masivamente).pack(anchor="w", pady=2, fill="x")

        organizacion_frame = ttk.LabelFrame(auto_col1, text="Organización de Archivo", padding=10)
        organizacion_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(organizacion_frame, text="📥 Organizar Facturas (FEOV)", command=self.accion_organizar_facturas_por_pdf).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📂➡️📁 Mover Archivos por Coincidencia de Nombre", command=self.accion_mover_archivos_por_coincidencia_nombre).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="🗺️ Copiar Archivos (Mapeo Subcarpetas)", command=self.accion_copiar_archivos_desde_mapeo_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📜 Copiar Archivos de Raíz (Mapeo Parcial)", command=self.accion_copiar_archivos_desde_raiz_mapeo_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📊 Análisis de Carpetas a Excel", command=self.exportar_lista).pack(anchor="w", pady=2, fill="x")
        ttk.Button(organizacion_frame, text="📊 Análisis de Historia Clínica", command=self.accion_analisis_historia_clinica).pack(anchor="w", pady=2, fill="x")
        ### NUEVO ###
        ttk.Button(organizacion_frame, text="📊 Análisis Autorización Nueva EPS", command=self.accion_analisis_autorizacion_nueva_eps).pack(anchor="w", pady=2, fill="x")
        ### FIN NUEVO ###
        ### NUEVO ANÁLISIS SANITAS ###
        ttk.Button(organizacion_frame, text="📊 Análisis Cargue Sanitas", command=self.accion_analisis_cargue_sanitas).pack(anchor="w", pady=2, fill="x")
        ### FIN NUEVO ANÁLISIS SANITAS ###
        ttk.Button(organizacion_frame, text="📤 Consolidar Archivos de Subcarpetas", command=self.accion_consolidar_archivos_subcarpetas).pack(anchor="w", pady=2, fill="x")
        
        renombrado_excel_frame = ttk.LabelFrame(auto_col2, text="Modificación y Renombrado con Excel", padding=10)
        renombrado_excel_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(renombrado_excel_frame, text="📤 Exportar para renombrar", command=self.exportar_para_renombrar).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="📥 Aplicar renombrado desde Excel", command=self.importar_y_aplicar_renombrado).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="🏷️ Añadir Sufijo desde Excel", command=self.accion_anadir_sufijo_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✍️ Modificar Autorización DOCX (Excel)", command=self.accion_autorizacion_docx_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✍️ Modificar Régimen DOCX (Excel)", command=self.accion_regimen_docx_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✍️ Modificar DOCX Completo (Excel)", command=self.accion_modificar_docx_completo_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="🖋️ Firmar DOCX con Imagen (por Carpeta)", command=self.accion_firmar_docx_con_imagen).pack(anchor="w", pady=2, fill="x")

        masivo_frame = ttk.LabelFrame(auto_col2, text="Conversión Masiva (desde resultados)", padding=10)
        masivo_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(masivo_frame, text="JSON (FEOV) → XLSX", command=self.accion_json_feov_a_xlsx_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="PDF → DOCX", command=self.accion_pdf_a_docx_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="JPG → PDF", command=self.accion_jpg_a_pdf_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="DOCX → PDF", command=self.accion_docx_a_pdf_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="PDF → JPG", command=self.accion_pdf_a_jpg_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="PNG → JPG", command=self.accion_png_a_jpg_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="TXT → JSON (Renombrar)", command=self.accion_txt_a_json_masivo).pack(anchor="w", pady=2, fill="x")
        ttk.Button(masivo_frame, text="📄 PDF → PDF (Escala de Grises)", command=self.accion_pdf_a_escala_grises_masivo_desde_resultados).pack(anchor="w", pady=2, fill="x")

        renombrado_excel_frame = ttk.LabelFrame(auto_col3, text="Creación de Archivos", padding=10)
        renombrado_excel_frame.pack(fill="x", pady=5, anchor="n")
        ttk.Button(renombrado_excel_frame, text="📂 Crear Carpetas desde Excel", command=self.accion_crear_carpetas_desde_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="⬇️ Descargar Firmas (URL/Excel)", command=self.accion_descargar_firmas_url_excel).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="📤 Copiar Archivo a Subcarpetas", command=self.accion_copiar_archivo_a_subcarpetas).pack(anchor="w", pady=2, fill="x")
          # === INICIO NUEVA FUNCIONALIDAD OVIDA ===
        ttk.Button(renombrado_excel_frame, text="⬇️ Descargar Historias Hospitalización (OVIDA)", command=self.accion_descargar_historias_hospitalizacion_ovida).pack(anchor="w", pady=2, fill="x")
        ttk.Button(renombrado_excel_frame, text="✒️ Crear Firma Digital desde Nombre de Carpeta", command=self.accion_crear_firma_digital).pack(anchor="w", pady=2, fill="x")

    def _setup_individual_conversion_tab(self, notebook):
        tab3 = ttk.Frame(notebook)
        notebook.add(tab3, text="🔄 Conversión de Archivos Individuales")
        
        conversion_frame = ttk.LabelFrame(tab3, text="Seleccionar archivos y convertir", padding=20)
        conversion_frame.pack(pady=20, padx=20)
        ttk.Button(conversion_frame, text="JSON (FEOV) → XLSX", command=self.accion_json_feov_a_xlsx).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="PDF → DOCX", command=self.accion_pdf_a_docx).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="JPG → PDF", command=self.accion_jpg_a_pdf).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="DOCX → PDF", command=self.accion_docx_a_pdf).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="PDF → JPG", command=self.accion_pdf_a_jpg).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="PNG → JPG", command=self.accion_png_a_jpg).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="TXT → JSON (Renombrar)", command=self.accion_txt_a_json).pack(anchor="w", pady=5, fill="x")
        ttk.Button(conversion_frame, text="📄 PDF → PDF (Escala de Grises)", command=self.accion_pdf_a_escala_grises_individual).pack(anchor="w", pady=5, fill="x")
            
    def _setup_data_editor_tab(self, notebook):
        tab4 = ttk.Frame(notebook)
        notebook.add(tab4, text="📄 Visor y Editor de Datos (JSON/XML)")
        self._setup_data_editor_tab(tab4) 

    # ====== Nueva Pestaña para RIPS ======
    def _setup_rips_tab(self, notebook):
        """Configura la pestaña RIPS para cargar la interfaz de OrganizadorArchivosApp."""
        tab_rips = ttk.Frame(notebook)
        notebook.add(tab_rips, text="RIPS")

        # Instanciar la aplicación RIPS dentro de esta pestaña
        # Aseguramos que RipsApp es una clase, no None
        if RipsApp:
            # Creamos un frame para contener la aplicación RIPS
            rips_frame = ttk.Frame(tab_rips, padding=10)
            rips_frame.pack(fill="both", expand=True)
            
            # Instanciamos RipsApp dentro de este frame
            # Pasamos el frame como root, no el root principal de la aplicación
            self.rips_instance = RipsApp(rips_frame) 
        else:
            # Si RipsApp es None, mostramos un mensaje
            ttk.Label(tab_rips, text="El módulo RIPS no está disponible. Asegúrate de que 'rips_module.py' se ha importado correctamente.", padding=20).pack()

    # --- Métodos de la clase original (ahora parte de MainApp) ---
    # ... (Todos los métodos de la clase OrganizadorArchivosApp irían aquí,
    # como self.log, self.seleccionar_directorio, self.buscar_archivos, etc.)
    # Los métodos de la clase original se heredan o se acceden a través de self.

    def log(self, mensaje):
        self.log_text.config(state="normal")
        self.log_text.insert(tk.END, f"{mensaje}\n")
        self.log_text.config(state="disabled")
        self.log_text.see(tk.END)

    def seleccionar_directorio(self):
        ruta = filedialog.askdirectory(parent=self.root)
        if ruta:
            self.source_path.set(ruta)

    def buscar_archivos(self):
        self.tree.delete(*self.tree.get_children())
        carpeta = self.source_path.get()
        patron = self.search_pattern.get()
        filtro = self.filter_type.get()
        incluir_sub = self.include_folders.get()
        tipo = self.element_type.get()
        self.resultados = []

        if not carpeta or not os.path.isdir(carpeta):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida.", parent=self.root)
            return

        for root, dirs, files in os.walk(carpeta):
            if not incluir_sub and root != carpeta:
                continue

            elementos = []
            if tipo == 'carpetas':
                elementos = dirs
            elif tipo == 'archivos':
                elementos = files
            else:
                elementos = dirs + files

            for nombre in elementos:
                ruta = os.path.join(root, nombre)
                es_carpeta = os.path.isdir(ruta)
                match = False

                if not patron:
                    match = True
                elif filtro == "extensión" and not es_carpeta:
                    match = nombre.lower().endswith(patron.lower())
                elif filtro == "nombre":
                    match = patron.lower() in nombre.lower()
                
                elif filtro == "Patrón con comodín (*, ?)":
                    try:
                        user_pattern = patron.lower()
                        effective_pattern = ''
                        if '*' in user_pattern or '?' in user_pattern:
                            effective_pattern = user_pattern
                        else:
                            if '.' in user_pattern:
                                name_part, ext_part = user_pattern.rsplit('.', 1)
                                effective_pattern = f"*{name_part}*.{ext_part}"
                            else:
                                effective_pattern = f"*{user_pattern}*"
                        match = fnmatch.fnmatch(nombre.lower(), effective_pattern)
                    except Exception as e:
                        self.log(f"Error en patrón con comodín: {e}")
                        match = False

                elif filtro == "expresión regular":
                    try:
                        match = bool(re.search(patron, nombre, re.IGNORECASE))
                    except re.error as e:
                        self.log(f"Error en expresión regular: {e}")
                        match = False
                
                if match:
                    try:
                        fecha_mod = os.path.getmtime(ruta)
                        fecha = pd.to_datetime(fecha_mod, unit='s').strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        fecha = "N/A"
                    self.tree.insert('', 'end', values=(ruta, fecha))
                    self.resultados.append(ruta)

        self.log(f"{len(self.resultados)} resultados encontrados.")

    def mostrar_dialogo_modificar_nombre(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("🏷️ Modificar Nombres - Opciones Avanzadas")
        dialog.geometry("520x680") 
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        opciones = {
            'sustituir_activo': tk.BooleanVar(), 'sustituir_buscar': tk.StringVar(), 'sustituir_reemplazar': tk.StringVar(),
            'prefijo_activo': tk.BooleanVar(), 'prefijo_texto': tk.StringVar(),
            'sufijo_activo': tk.BooleanVar(), 'sufijo_texto': tk.StringVar(),
            'renombrar_activo': tk.BooleanVar(), 'renombrar_texto': tk.StringVar(),
            'limpiar_id_activo': tk.BooleanVar(),
            'resultado': None
        }

        main_frame = ttk.Frame(dialog, padding=20)
        main_frame.pack(fill='both', expand=True)
        
        renombrar_frame = ttk.LabelFrame(main_frame, text="📝 Renombrado completo (ignora las demás opciones)", padding=10)
        renombrar_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(renombrar_frame, text="Activar renombrado completo", variable=opciones['renombrar_activo']).pack(anchor='w')
        ttk.Label(renombrar_frame, text="Nuevo nombre (sin extensión):").pack(anchor='w', pady=(5, 0))
        ttk.Entry(renombrar_frame, textvariable=opciones['renombrar_texto']).pack(fill='x', pady=2)

        sustituir_frame = ttk.LabelFrame(main_frame, text="🔄 Sustituir texto", padding=10)
        sustituir_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(sustituir_frame, text="Activar sustitución", variable=opciones['sustituir_activo']).pack(anchor='w')
        ttk.Label(sustituir_frame, text="Buscar:").pack(anchor='w', pady=(5, 0))
        ttk.Entry(sustituir_frame, textvariable=opciones['sustituir_buscar']).pack(fill='x', pady=2)
        ttk.Label(sustituir_frame, text="Reemplazar con:").pack(anchor='w')
        ttk.Entry(sustituir_frame, textvariable=opciones['sustituir_reemplazar']).pack(fill='x', pady=2)

        limpieza_frame = ttk.LabelFrame(main_frame, text="🧹 Limpieza Especial (FEOV)", padding=10)
        limpieza_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(limpieza_frame, text="Eliminar '_ID<números>_A' del nombre", variable=opciones['limpiar_id_activo']).pack(anchor='w', pady=5)

        prefijo_frame = ttk.LabelFrame(main_frame, text="⬅️ Añadir al inicio", padding=10)
        prefijo_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(prefijo_frame, text="Añadir prefijo", variable=opciones['prefijo_activo']).pack(anchor='w')
        ttk.Entry(prefijo_frame, textvariable=opciones['prefijo_texto']).pack(fill='x', pady=2)

        sufijo_frame = ttk.LabelFrame(main_frame, text="➡️ Añadir al final (antes de extensión)", padding=10)
        sufijo_frame.pack(fill='x', pady=5)
        ttk.Checkbutton(sufijo_frame, text="Añadir sufijo", variable=opciones['sufijo_activo']).pack(anchor='w')
        ttk.Entry(sufijo_frame, textvariable=opciones['sufijo_texto']).pack(fill='x', pady=2)

        def aplicar():
            if not (opciones['sustituir_activo'].get() or opciones['prefijo_activo'].get() or opciones['sufijo_activo'].get() or opciones['renombrar_activo'].get() or opciones['limpiar_id_activo'].get()):
                messagebox.showwarning("Advertencia", "Debes activar y rellenar al menos una opción.", parent=dialog)
                return
            opciones['resultado'] = opciones
            dialog.destroy()

        botones_frame = ttk.Frame(main_frame)
        botones_frame.pack(fill='x', pady=20)
        ttk.Button(botones_frame, text="✅ Ejecutar Cambios", command=aplicar).pack(side="right", padx=5)
        ttk.Button(botones_frame, text="❌ Cancelar", command=dialog.destroy).pack(side="right")

        self.root.wait_window(dialog)
        return opciones['resultado']

    def ejecutar_accion(self):
        accion = self.accion.get()
        total_resultados = len(self.resultados)
        if total_resultados == 0:
            self.log("No hay archivos seleccionados para ejecutar la acción.")
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar una acción sobre ellos.", parent=self.root)
            return

        progress_win, progress_bar, progress_label = None, None, None

        try:
            if accion in ["copiar a carpeta", "mover a carpeta"]:
                action_word = "copiar" if accion == "copiar a carpeta" else "mover"
                destino = filedialog.askdirectory(parent=self.root, title=f'Seleccionar carpeta destino para {action_word}')
                if not destino: return
                progress_win, progress_bar, progress_label = self._create_progress_window(f"{action_word.capitalize()}ndo archivos...", total_resultados)
                for i, ruta in enumerate(self.resultados):
                    nombre_archivo = os.path.basename(ruta)
                    progress_label.config(text=f"{action_word.capitalize()}ndo: {nombre_archivo}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    try:
                        if accion == "copiar a carpeta":
                            shutil.copy2(ruta, destino)
                        else:
                            shutil.move(ruta, destino)
                        self.log(f'{action_word.capitalize()}do: {nombre_archivo} a {destino}')
                    except Exception as e:
                        self.log(f'Error al {action_word} {ruta}: {e}')

            elif accion == "modificar nombre":
                opciones = self.mostrar_dialogo_modificar_nombre()
                if not opciones:
                    self.log("Renombrado cancelado por el usuario.")
                    return

                procesados, errores = 0, 0
                progress_win, progress_bar, progress_label = self._create_progress_window("Renombrando archivos...", total_resultados)
                for i, archivo in enumerate(self.resultados):
                    nombre_completo = os.path.basename(archivo)
                    progress_label.config(text=f"Renombrando: {nombre_completo}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    try:
                        carpeta, _ = os.path.split(archivo)
                        nombre_base, extension = os.path.splitext(nombre_completo)
                        nuevo_nombre = nombre_base

                        if opciones['renombrar_activo'].get() and opciones['renombrar_texto'].get():
                            nuevo_nombre = opciones['renombrar_texto'].get()
                        else:
                            if opciones['sustituir_activo'].get() and opciones['sustituir_buscar'].get():
                                nuevo_nombre = nuevo_nombre.replace(opciones['sustituir_buscar'].get(), opciones['sustituir_reemplazar'].get())
                            
                            if opciones['limpiar_id_activo'].get():
                                nuevo_nombre = re.sub(r'_ID\d+_A', '', nuevo_nombre, flags=re.IGNORECASE)
                                                        
                            if opciones['prefijo_activo'].get() and opciones['prefijo_texto'].get():
                                nuevo_nombre = opciones['prefijo_texto'].get() + nuevo_nombre
                            if opciones['sufijo_activo'].get() and opciones['sufijo_texto'].get():
                                nuevo_nombre = nuevo_nombre + opciones['sufijo_texto'].get()

                        if nuevo_nombre != nombre_base:
                            nueva_ruta = os.path.join(carpeta, nuevo_nombre + extension)
                            if os.path.exists(nueva_ruta):
                                self.log(f"⚠️ Error: ya existe un archivo con el nombre '{os.path.basename(nueva_ruta)}'")
                                errores += 1
                                continue
                            os.rename(archivo, nueva_ruta)
                            self.log(f"✅ Renombrado: {nombre_completo} → {os.path.basename(os.path.basename(nueva_ruta))}")
                            procesados += 1
                    except Exception as e:
                        self.log(f"❌ Error al renombrar {archivo}: {e}")
                        errores += 1
                self.log(f"--- Renombrado finalizado. Procesados: {procesados}, Errores: {errores} ---")
                if procesados > 0: self.buscar_archivos()


            elif accion == "editar texto":
                buscar = simpledialog.askstring('Buscar en contenido', 'Texto a buscar:', parent=self.root)
                if buscar is None: return
                reemplazar = simpledialog.askstring('Reemplazar con', 'Nuevo texto:', parent=self.root)
                if reemplazar is None: return
                
                progress_win, progress_bar, progress_label = self._create_progress_window("Editando contenido...", total_resultados)
                for i, archivo in enumerate(self.resultados):
                    nombre_archivo = os.path.basename(archivo)
                    progress_label.config(text=f"Analizando: {nombre_archivo}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    try:
                        ext = os.path.splitext(archivo)[1].lower()
                        if ext in ['.txt', '.json', '.xml', '.csv', '.html', '.md', '.log']:
                            with open(archivo, 'r', encoding='utf-8', errors='ignore') as f:
                                contenido = f.read()
                            if buscar in contenido:
                                contenido_nuevo = contenido.replace(buscar, reemplazar)
                                with open(archivo, 'w', encoding='utf-8') as f:
                                    f.write(contenido_nuevo)
                                self.log(f'Texto modificado en: {archivo}')
                        elif ext == '.docx':
                            doc = Document(archivo)
                            modificado = False
                            for p in doc.paragraphs:
                                if buscar in p.text:
                                    p.text = p.text.replace(buscar, reemplazar)
                                    modificado = True
                            if modificado:
                                doc.save(archivo)
                                self.log(f'Texto modificado en DOCX: {archivo}')
                    except Exception as e:
                        self.log(f'❌ Error al editar {archivo}: {e}')

            elif accion == "comprimir archivos en zip":
                nombre = simpledialog.askstring("Nombre ZIP", "Nombre del archivo ZIP (sin .zip):", parent=self.root)
                if not nombre: return
                carpeta_destino = filedialog.askdirectory(parent=self.root, title="Selecciona carpeta donde guardar el ZIP")
                if not carpeta_destino: return
                ruta_zip = os.path.join(carpeta_destino, nombre + ".zip")
                progress_win, progress_bar, progress_label = self._create_progress_window("Comprimiendo en ZIP...", total_resultados)
                with ZipFile(ruta_zip, "w", ZIP_DEFLATED) as zipf:
                    for i, archivo in enumerate(self.resultados):
                        nombre_archivo = os.path.basename(archivo)
                        progress_label.config(text=f"Añadiendo: {nombre_archivo}")
                        progress_bar['value'] = i + 1
                        self.root.update_idletasks()
                        if os.path.isdir(archivo):
                            for root_dir, _, files_in_dir in os.walk(archivo):
                                for f in files_in_dir:
                                    abs_path = os.path.join(root_dir, f)
                                    arcname = os.path.relpath(abs_path, os.path.dirname(archivo))
                                    zipf.write(abs_path, arcname)
                        else:
                            zipf.write(archivo, os.path.basename(archivo))
                self.log(f"✅ ZIP creado en: {ruta_zip}")

            elif accion == "comprimir carpetas individualmente":
                carpetas_a_comprimir = [r for r in self.resultados if os.path.isdir(r)]
                if not carpetas_a_comprimir:
                    messagebox.showinfo("Información", "La selección no contiene carpetas para comprimir.", parent=self.root)
                    return
                
                carpeta_destino = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta destino para los ZIPs")
                if not carpeta_destino: return
                progress_win, progress_bar, progress_label = self._create_progress_window("Comprimiendo carpetas...", len(carpetas_a_comprimir))
                for i, ruta in enumerate(carpetas_a_comprimir):
                    nombre_zip = os.path.basename(ruta.rstrip(os.sep))
                    progress_label.config(text=f"Comprimiendo: {nombre_zip}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    ruta_zip = os.path.join(carpeta_destino, nombre_zip + ".zip")
                    try:
                        with ZipFile(ruta_zip, "w", ZIP_DEFLATED) as zipf:
                            for root_dir, _, files_in_dir in os.walk(ruta):
                                for f in files_in_dir:
                                    abs_path = os.path.join(root_dir, f)
                                    arcname = os.path.relpath(abs_path, start=ruta)
                                    zipf.write(abs_path, arcname)
                        self.log(f"✅ Carpeta comprimida: {nombre_zip}.zip")
                    except Exception as e:
                        self.log(f"❌ Error al comprimir {ruta}: {e}")
        
        finally:
            if progress_win:
                progress_win.destroy()

    def limpiar(self):
        self.tree.delete(*self.tree.get_children())
        self.log_text.config(state="normal")
        self.log_text.delete("1.0", tk.END)
        self.log_text.config(state="disabled")
        self.source_path.set("")
        self.search_pattern.set("")
        self.resultados = []
        self.log("Campos y resultados limpiados.")

    def eliminar_resultados(self):
        if not self.resultados:
            self.log("No hay nada que eliminar.")
            return
        confirm = messagebox.askyesno("Confirmar eliminación", f"¿Seguro que quieres enviar {len(self.resultados)} elementos a la papelera?", parent=self.root)
        if not confirm: return

        eliminados, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Eliminando...", len(self.resultados))
        try:
            for i, ruta in enumerate(self.resultados):
                nombre_archivo = os.path.basename(ruta)
                progress_label.config(text=f"Eliminando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                try:
                    ruta_normalizada = os.path.normpath(ruta)
                    send2trash(ruta_normalizada)
                    eliminados += 1
                except Exception as e:
                    self.log(f"❌ Error al eliminar {ruta}: {e}")
                    errores += 1
        finally:
            progress_win.destroy()
        
        self.log(f"🗑️ Enviados a la papelera: {eliminados}. Errores: {errores}.")
        if eliminados > 0: self.buscar_archivos()

    def _convertir_masivo(self, extension_filtro, output_ext, conversion_function, title="Convirtiendo archivos..."):
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Busca archivos primero para usar la conversión masiva.", parent=self.root)
            return
        
        archivos_filtrados = [ruta for ruta in self.resultados if any(ruta.lower().endswith(ext) for ext in (extension_filtro if isinstance(extension_filtro, (tuple, list)) else (extension_filtro,)))]
        if not archivos_filtrados:
             messagebox.showinfo("Información", f"No se encontraron archivos con la extensión '{extension_filtro}' en los resultados.", parent=self.root)
             return
        
        progress_win, progress_bar, progress_label = self._create_progress_window(title, len(archivos_filtrados))
        try:
            for i, ruta in enumerate(archivos_filtrados):
                nombre_base = os.path.splitext(os.path.basename(ruta))[0]
                salida = os.path.join(os.path.dirname(ruta), f"{nombre_base}.{output_ext}")

                progress_label.config(text=f"Convirtiendo: {os.path.basename(ruta)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                conversion_function(ruta, salida)
        finally:
            progress_win.destroy()
            self.log(f"--- Conversión masiva finalizada. ---")

    def importar_y_aplicar_renombrado(self):
        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar Excel editado", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            df = pd.read_excel(archivo_excel)
            if not all(col in df.columns for col in ["Ruta actual", "Nuevo nombre"]):
                self.log("❌ El Excel debe contener las columnas 'Ruta actual' y 'Nuevo nombre'.")
                return
            
            cambios, errores = 0, 0
            progress_win, progress_bar, progress_label = self._create_progress_window("Renombrando desde Excel...", len(df))
            try:
                for index, fila in df.iterrows():
                    ruta_actual = fila["Ruta actual"]
                    nuevo_nombre = fila["Nuevo nombre"]

                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()

                    if pd.isna(nuevo_nombre) or pd.isna(ruta_actual) or os.path.basename(ruta_actual) == nuevo_nombre: continue
                    
                    progress_label.config(text=f"Procesando: {os.path.basename(str(ruta_actual))}")
                    nueva_ruta = os.path.join(os.path.dirname(ruta_actual), str(nuevo_nombre))
                    try:
                        if os.path.exists(ruta_actual):
                            if os.path.exists(nueva_ruta):
                                self.log(f"⚠️ Error: ya existe un archivo con el nombre '{nuevo_nombre}'")
                                errores += 1
                                continue
                            os.rename(ruta_actual, nueva_ruta)
                            self.log(f"✅ Renombrado: {os.path.basename(ruta_actual)} → {nuevo_nombre}")
                            cambios += 1
                        else:
                             self.log(f"⚠️ No se encontró: {ruta_actual}")
                    except Exception as e:
                        self.log(f"❌ Error al renombrar {ruta_actual}: {e}")
                        errores += 1
            finally:
                progress_win.destroy()

            self.log(f"Renombrado desde Excel finalizado. Cambios: {cambios}, Errores: {errores}")
            if cambios > 0: self.buscar_archivos()
        except Exception as e:
            self.log(f"Error al procesar el Excel: {e}")

    def _convertir_archivos(self, title, filetypes, output_ext, conversion_function):
        archivos = filedialog.askopenfilenames(parent=self.root, title=title, filetypes=filetypes)
        if not archivos: return
        
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Convirtiendo a .{output_ext}", len(archivos))
        try:
            for i, archivo in enumerate(archivos):
                nombre_base = os.path.splitext(os.path.basename(archivo))[0]
                salida = os.path.join(os.path.dirname(archivo), f"{nombre_base}.{output_ext}")
                
                progress_label.config(text=f"Convirtiendo: {os.path.basename(archivo)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                conversion_function(archivo, salida)
        finally:
            progress_win.destroy()
            self.log("--- Conversión finalizada. ---")
    
    def accion_txt_a_json_masivo(self): self._convertir_masivo(".txt", "json", self._txt_a_json_worker, "Renombrando TXT a JSON...")
    def accion_json_feov_a_xlsx_masivo(self): self._convertir_masivo(".json", "xlsx", self._json_feov_a_xlsx_worker, "Convirtiendo JSON (FEOV) a XLSX...")
    def accion_pdf_a_docx_masivo(self): self._convertir_masivo(".pdf", "docx", self._pdf_a_docx_worker, "Convirtiendo PDF a DOCX...")
    def accion_jpg_a_pdf_masivo(self): self._convertir_masivo((".jpg",".jpeg"), "pdf", self._jpg_a_pdf_worker, "Convirtiendo JPG a PDF...")
    def accion_docx_a_pdf_masivo(self): self._convertir_masivo(".docx", "pdf", self._docx_a_pdf_worker, "Convirtiendo DOCX a PDF...")
    def accion_pdf_a_jpg_masivo(self): self._convertir_masivo(".pdf", "jpg", self._pdf_a_jpg_worker, "Convirtiendo PDF a JPG...")
    def accion_png_a_jpg_masivo(self): self._convertir_masivo(".png", "jpg", self._png_a_jpg_worker, "Convirtiendo PNG a JPG...")

    def accion_json_feov_a_xlsx(self): self._convertir_archivos("Seleccionar JSON (FEOV)", [("JSON", "*.json")], "xlsx", self._json_feov_a_xlsx_worker)
    def accion_pdf_a_docx(self): self._convertir_archivos("Seleccionar PDF(s)", [("PDF", "*.pdf")], "docx", self._pdf_a_docx_worker)
    def accion_jpg_a_pdf(self): self._convertir_archivos("Seleccionar JPG(s)", [("JPG", "*.jpg;*.jpeg")], "pdf", self._jpg_a_pdf_worker)
    def accion_docx_a_pdf(self): self._convertir_archivos("Seleccionar DOCX", [("Word", "*.docx")], "pdf", self._docx_a_pdf_worker)
    def accion_pdf_a_jpg(self): self._convertir_archivos("Seleccionar PDF(s)", [("PDF", "*.pdf")], "jpg", self._pdf_a_jpg_worker)
    def accion_png_a_jpg(self): self._convertir_archivos("Seleccionar PNG(s)", [("PNG", "*.png")], "jpg", self._png_a_jpg_worker)

    def _txt_a_json_worker(self, entrada, salida):
        """Worker para renombrar un archivo de .txt a .json."""
        try:
            if os.path.exists(salida):
                self.log(f"⚠️ Error: El archivo '{os.path.basename(salida)}' ya existe. Se omite el renombrado.")
                return
            os.rename(entrada, salida)
            self.log(f"✅ Renombrado: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error renombrando {os.path.basename(entrada)}: {e}")

    def _json_feov_a_xlsx_worker(self, entrada, salida):
        try:
            with open(entrada, 'r', encoding='utf-8') as f:
                data = json.load(f)
            df = pd.json_normalize(data)
            df.to_excel(salida, index=False)
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")

    def _pdf_a_docx_worker(self, entrada, salida):
        try:
            cv = Converter(entrada)
            cv.convert(salida, start=0, end=None)
            cv.close()
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")

    def _jpg_a_pdf_worker(self, entrada, salida):
        try:
            img = Image.open(entrada)
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            img.save(salida, "PDF", resolution=100.0)
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")
            
    def _docx_a_pdf_worker(self, entrada, salida):
        """
        Convierte DOCX a PDF usando automatización COM directa.
        Esta es la versión robusta que funciona correctamente en el .exe compilado.
        """
        word = None
        doc = None
        try:
            in_file_abs = os.path.abspath(entrada)
            out_file_abs = os.path.abspath(salida)

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(in_file_abs)
            
            wdFormatPDF = 17
            doc.SaveAs(out_file_abs, FileFormat=wdFormatPDF)
            
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")

        except Exception as e:
            error_msg = str(e)
            if "com_error" in error_msg:
                 error_msg += ("\n (Asegúrese de que Microsoft Word esté instalado y no tenga diálogos abiertos que bloqueen la automatización).")
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {error_msg}")
        
        finally:
            if doc:
                doc.Close(False)
            if word:
                word.Quit()
            
    def _pdf_a_jpg_worker(self, entrada, salida):
        try:
            doc = fitz.open(entrada)
            for i, page in enumerate(doc):
                pix = page.get_pixmap()
                if len(doc) == 1:
                    nombre_salida = salida
                else:
                    base, ext = os.path.splitext(salida)
                    nombre_salida = f"{base}_p{i+1}{ext}"
                pix.save(nombre_salida)
            doc.close()
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)} (y páginas adicionales si aplica)")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")
        
    def _png_a_jpg_worker(self, entrada, salida):
        try:
            img = Image.open(entrada).convert("RGB")
            img.save(salida, "jpeg")
            self.log(f"✅ Convertido: {os.path.basename(entrada)} → {os.path.basename(salida)}")
        except Exception as e:
            self.log(f"❌ Error convirtiendo {os.path.basename(entrada)}: {e}")
        
    def accion_txt_a_json(self):
        archivos = filedialog.askopenfilenames(parent=self.root, title="Seleccionar archivos TXT para renombrar a JSON", filetypes=[("TXT", "*.txt")])
        if not archivos: return
        for archivo in archivos:
            try:
                base, _ = os.path.splitext(archivo)
                os.rename(archivo, base + ".json")
                self.log(f"Renombrado: {os.path.basename(archivo)} → {os.path.basename(base)}.json")
            except Exception as e:
                self.log(f"Error renombrando {os.path.basename(archivo)}: {e}")

    def accion_unificar_por_carpeta(self):
        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta principal con subcarpetas")
        if not carpeta_base: return
        
        nombre_final_base = simpledialog.askstring(
            "Nombre del archivo final", 
            "Nombre para los PDF combinados (sin .pdf).\nEste nombre se usará para el archivo resultante en CADA subcarpeta.", 
            parent=self.root
        )
        if not nombre_final_base: return
        
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
        if not subcarpetas:
            self.log("No se encontraron subcarpetas.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Unificando PDFs por carpeta...", len(subcarpetas))
        try:
            for i, carpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(carpeta)
                progress_label.config(text=f"Procesando: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                archivos_pdf_a_procesar = []
                for num_pdf in range(1, 11):
                    nombre_archivo_buscado = f"{num_pdf}.pdf"
                    ruta_archivo_buscado = os.path.join(carpeta, nombre_archivo_buscado)
                    if os.path.exists(ruta_archivo_buscado):
                        archivos_pdf_a_procesar.append(ruta_archivo_buscado)

                if not archivos_pdf_a_procesar:
                    self.log(f"No se encontraron PDFs numerados (1-10) en '{nombre_subcarpeta}', saltando.")
                    continue
                
                nombre_final = f"{nombre_final_base}.pdf"
                ruta_salida = os.path.join(carpeta, nombre_final)
                
                doc_final = fitz.open()
                
                for ruta_pdf in archivos_pdf_a_procesar:
                    try:
                        with fitz.open(ruta_pdf) as doc_origen:
                            for page_origen in doc_origen:
                                pix = page_origen.get_pixmap(dpi=300, colorspace=fitz.csGRAY)
                                pagina_nueva = doc_final.new_page(width=pix.width, height=pix.height)
                                pagina_nueva.insert_image(pagina_nueva.rect, pixmap=pix)
                    except Exception as e:
                        self.log(f"Error procesando '{os.path.basename(ruta_pdf)}' en '{nombre_subcarpeta}': {e}")
                
                if len(doc_final) > 0:
                    doc_final.save(ruta_salida, garbage=4, deflate=True)
                    self.log(f"✅ PDF unificado '{nombre_final}' creado en la carpeta '{nombre_subcarpeta}'.")
                doc_final.close()
        finally:
            if progress_win:
                progress_win.destroy()
    
    def accion_unificar_jpg_por_carpeta(self):
        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta principal con subcarpetas de JPGs")
        if not carpeta_base: return
        
        nombre_final_base = simpledialog.askstring(
            "Nombre del archivo final",
            "Nombre para el PDF combinado (sin .pdf).\nEste nombre se usará para el archivo resultante en CADA subcarpeta.",
            parent=self.root
        )
        if not nombre_final_base: return
        
        subcarpetas = [os.path.join(carpeta_base, d) for d in os.listdir(carpeta_base) if os.path.isdir(os.path.join(carpeta_base, d))]
        if not subcarpetas:
            self.log("No se encontraron subcarpetas.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Unificando JPGs por carpeta...", len(subcarpetas))
        try:
            for i, carpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(carpeta)
                progress_label.config(text=f"Procesando: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                archivos_img_a_procesar = []
                for num_img in range(1, 11):
                    ruta_encontrada = None
                    for ext in ['.jpg', '.jpeg']:
                        nombre_archivo = f"{num_img}{ext}"
                        ruta_archivo = os.path.join(carpeta, nombre_archivo)
                        if os.path.exists(ruta_archivo):
                            ruta_encontrada = ruta_archivo
                            break
                    if ruta_encontrada:
                        archivos_img_a_procesar.append(ruta_encontrada)

                if not archivos_img_a_procesar:
                    self.log(f"No se encontraron JPGs/JPEGs numerados (1-10) en '{nombre_subcarpeta}', saltando.")
                    continue

                lista_imagenes_procesadas = []
                for ruta_img in archivos_img_a_procesar:
                    try:
                        img = Image.open(ruta_img)
                        img_gray = img.convert('L') 
                        lista_imagenes_procesadas.append(img_gray)
                    except Exception as e:
                        self.log(f"Error abriendo imagen '{os.path.basename(ruta_img)}': {e}")
                
                if lista_imagenes_procesadas:
                    nombre_pdf = f"{nombre_final_base}.pdf"
                    ruta_salida = os.path.join(carpeta, nombre_pdf)
                    
                    lista_imagenes_procesadas[0].save(
                        ruta_salida, 
                        save_all=True, 
                        append_images=lista_imagenes_procesadas[1:], 
                        resolution=300.0
                    )
                    self.log(f"✅ PDF '{nombre_pdf}' creado a partir de JPGs en '{nombre_subcarpeta}'.")
        finally:
            if progress_win:
                progress_win.destroy()
    
    def accion_unificar_pdfs(self):
        archivos = filedialog.askopenfilenames(
            parent=self.root,
            title="Seleccionar PDFs a unificar",
            filetypes=[("Archivos PDF", "*.pdf")]
        )
        if len(archivos) < 2:
            self.log("Selecciona al menos dos archivos para unificar.")
            return

        archivo_map = {os.path.basename(p): p for p in archivos}
        orden_dialog = tk.Toplevel(self.root)
        orden_dialog.title("📚 Ordenar PDFs antes de unificar")
        orden_dialog.geometry("600x500")
        orden_dialog.transient(self.root)
        orden_dialog.grab_set()

        list_frame = ttk.Frame(orden_dialog)
        list_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical")
        orden_listbox = tk.Listbox(list_frame, selectmode=tk.SINGLE, yscrollcommand=scrollbar.set)
        scrollbar.config(command=orden_listbox.yview)

        scrollbar.pack(side="right", fill="y")
        orden_listbox.pack(side="left", fill="both", expand=True)

        try:
            import natsort # type: ignore
            sorted_archivos = natsort.natsorted(archivos)
        except ImportError:
            sorted_archivos = sorted(archivos)

        for archivo in sorted_archivos:
            orden_listbox.insert(tk.END, os.path.basename(archivo))
        
        botones_frame = ttk.Frame(orden_dialog)
        botones_frame.pack(pady=10)

        def mover_arriba():
            seleccion = orden_listbox.curselection()
            if not seleccion or seleccion[0] == 0: return
            i = seleccion[0]
            texto = orden_listbox.get(i)
            orden_listbox.delete(i)
            orden_listbox.insert(i - 1, texto)
            orden_listbox.select_set(i - 1)
            orden_listbox.see(i - 1)

        def mover_abajo():
            seleccion = orden_listbox.curselection()
            if not seleccion or seleccion[0] == orden_listbox.size() - 1: return
            i = seleccion[0]
            texto = orden_listbox.get(i)
            orden_listbox.delete(i)
            orden_listbox.insert(i + 1, texto)
            orden_listbox.select_set(i + 1)
            orden_listbox.see(i + 1)

        def unificar():
            ordered_basenames = orden_listbox.get(0, tk.END)
            if not ordered_basenames:
                messagebox.showwarning("Advertencia", "No hay archivos para unificar.", parent=orden_dialog)
                return

            ruta_guardado = filedialog.asksaveasfilename(
                parent=self.root,
                defaultextension=".pdf", 
                filetypes=[("Archivo PDF", "*.pdf")], 
                title="Guardar PDF unificado como"
            )
            if not ruta_guardado: return
            orden_dialog.destroy()
            self.log("Iniciando unificación manual de PDFs...")
            doc_final = fitz.open()

            progress_win, progress_bar, progress_label = self._create_progress_window("Unificando PDFs...", len(ordered_basenames))
            try:
                for i, basename in enumerate(ordered_basenames):
                    progress_label.config(text=f"Procesando: {basename}")
                    progress_bar['value'] = i + 1
                    self.root.update_idletasks()
                    
                    ruta = archivo_map.get(basename)
                    if not ruta:
                        self.log(f"⚠️ No se encontró la ruta para {basename}, saltando.")
                        continue
                    
                    try:
                        with fitz.open(ruta) as doc_origen:
                           doc_final.insert_pdf(doc_origen)
                        self.log(f"📄 Procesado: {os.path.basename(ruta)}")
                    except Exception as e:
                        self.log(f"❌ Error al procesar {ruta}: {e}")
                
                if len(doc_final) > 0:
                    try:
                        doc_final.save(ruta_guardado, garbage=4, deflate=True)
                        self.log(f"✅ PDF combinado guardado en: {ruta_guardado}")
                    except Exception as e:
                        self.log(f"❌ Error al guardar el PDF final: {e}")
                else:
                    self.log("❌ No se generaron páginas. El PDF final no fue guardado.")
                doc_final.close()
            finally:
                if progress_win:
                    progress_win.destroy()

        ttk.Button(botones_frame, text="⬆️ Arriba", command=mover_arriba).pack(side="left", padx=5)
        ttk.Button(botones_frame, text="⬇️ Abajo", command=mover_abajo).pack(side="left", padx=5)
        ttk.Button(botones_frame, text="✅ Unificar y guardar", command=unificar).pack(side="left", padx=15)
        ttk.Button(botones_frame, text="❌ Cancelar", command=orden_dialog.destroy).pack(side="right", padx=5)
        
        self.root.wait_window(orden_dialog)
    
    def accion_dividir_pdf_en_paginas(self):
        archivos_pdf = filedialog.askopenfilenames(
            parent=self.root,
            title="Seleccionar PDF(s) para dividir en páginas",
            filetypes=[("Archivos PDF", "*.pdf")]
        )
        if not archivos_pdf:
            self.log("Operación de división de PDF cancelada.")
            return

        self.log(f"--- Iniciando división de {len(archivos_pdf)} PDF(s) en páginas individuales ---")
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Dividiendo PDFs...", len(archivos_pdf))
        try:
            for i, ruta_pdf_original in enumerate(archivos_pdf):
                nombre_base_original = os.path.splitext(os.path.basename(ruta_pdf_original))[0]
                directorio_origen = os.path.dirname(ruta_pdf_original)

                progress_label.config(text=f"Procesando: {nombre_base_original}.pdf")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                ruta_carpeta_salida = os.path.join(directorio_origen, nombre_base_original)
                os.makedirs(ruta_carpeta_salida, exist_ok=True)

                try:
                    with fitz.open(ruta_pdf_original) as doc_origen:
                        num_paginas = len(doc_origen)
                        if num_paginas == 0:
                            self.log(f"⚠️ El archivo '{nombre_base_original}.pdf' no tiene páginas y será omitido.")
                            continue

                        for num_pagina in range(num_paginas):
                            doc_nuevo = fitz.open()
                            doc_nuevo.insert_pdf(doc_origen, from_page=num_pagina, to_page=num_pagina)
                            
                            nombre_pagina_nueva = f"{num_pagina + 1}.pdf"
                            ruta_pagina_nueva = os.path.join(ruta_carpeta_salida, nombre_pagina_nueva)
                            
                            doc_nuevo.save(ruta_pagina_nueva)
                            doc_nuevo.close()
                        
                        self.log(f"✅ Dividido '{nombre_base_original}.pdf' en {num_paginas} páginas en la carpeta '{nombre_base_original}'")
                
                except Exception as e:
                    self.log(f"❌ Error al procesar '{nombre_base_original}.pdf': {e}")
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log("--- Proceso de división finalizado. ---")
        messagebox.showinfo("Proceso Completado", "La división de los archivos PDF ha finalizado. Revise el registro de actividad para más detalles.", parent=self.root)

    def accion_dividir_pdfs_masivamente(self):
        carpeta_base = filedialog.askdirectory(
            parent=self.root,
            title="Seleccionar carpeta principal para buscar y dividir PDFs"
        )
        if not carpeta_base:
            self.log("Operación de división masiva cancelada.")
            return

        self.log(f"--- Iniciando búsqueda y división masiva de PDFs en: {carpeta_base} ---")

        pdfs_a_procesar = []
        for root, _, files in os.walk(carpeta_base):
            for file in files:
                if file.lower().endswith('.pdf'):
                    pdfs_a_procesar.append(os.path.join(root, file))

        if not pdfs_a_procesar:
            self.log("No se encontraron archivos PDF en la carpeta seleccionada y sus subdirectorios.")
            messagebox.showinfo("Sin Archivos", "No se encontraron archivos PDF para procesar.", parent=self.root)
            return
            
        self.log(f"Se encontraron {len(pdfs_a_procesar)} PDFs para dividir.")
        
        progress_win, progress_bar, progress_label = self._create_progress_window(
            "Dividiendo PDFs masivamente...", len(pdfs_a_procesar)
        )
        
        try:
            for i, ruta_pdf_original in enumerate(pdfs_a_procesar):
                nombre_base_original = os.path.splitext(os.path.basename(ruta_pdf_original))[0]
                directorio_origen = os.path.dirname(ruta_pdf_original)

                progress_label.config(text=f"Procesando: {os.path.basename(ruta_pdf_original)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                ruta_carpeta_salida = os.path.join(directorio_origen, nombre_base_original)
                os.makedirs(ruta_carpeta_salida, exist_ok=True)

                try:
                    with fitz.open(ruta_pdf_original) as doc_origen:
                        num_paginas = len(doc_origen)
                        if num_paginas == 0:
                            self.log(f"⚠️ El archivo '{os.path.basename(ruta_pdf_original)}' no tiene páginas, omitido.")
                            continue

                        for num_pagina in range(num_paginas):
                            doc_nuevo = fitz.open()
                            doc_nuevo.insert_pdf(doc_origen, from_page=num_pagina, to_page=num_pagina)
                            
                            nombre_pagina_nueva = f"{num_pagina + 1}.pdf"
                            ruta_pagina_nueva = os.path.join(ruta_carpeta_salida, nombre_pagina_nueva)
                            
                            doc_nuevo.save(ruta_pagina_nueva)
                            doc_nuevo.close()
                        
                        self.log(f"✅ Dividido '{os.path.basename(ruta_pdf_original)}' en {num_paginas} páginas en la carpeta '{os.path.basename(ruta_carpeta_salida)}'")
                
                except Exception as e:
                    self.log(f"❌ Error al procesar '{os.path.basename(ruta_pdf_original)}': {e}")
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log("--- Proceso de división masiva finalizado. ---")
        messagebox.showinfo("Proceso Completado", "La división masiva de PDFs ha finalizado. Revise el registro de actividad para más detalles.", parent=self.root)

    def _pdf_a_escala_grises_worker(self, entrada, salida):
        es_sobrescritura = (os.path.normpath(entrada) == os.path.normpath(salida))
        ruta_temporal = entrada + "._tmp_grayscale" if es_sobrescritura else None

        try:
            with fitz.open(entrada) as doc_origen:
                if doc_origen.is_encrypted:
                    self.log(f"⚠️ El archivo '{os.path.basename(entrada)}' está encriptado y no se puede procesar.")
                    return
                
                doc_final = fitz.open()
                for page in doc_origen:
                    pix = page.get_pixmap(dpi=300, colorspace=fitz.csGRAY)
                    pagina_nueva = doc_final.new_page(width=pix.width, height=pix.height)
                    pagina_nueva.insert_image(pagina_nueva.rect, pixmap=pix)
                
                ruta_guardado = ruta_temporal if es_sobrescritura else salida
                doc_final.save(ruta_guardado, garbage=4, deflate=True)
                doc_final.close()
            
            if es_sobrescritura:
                os.replace(ruta_temporal, entrada)
                self.log(f"✅ Reemplazado (escala de grises): {os.path.basename(entrada)}")
            else:
                self.log(f"✅ Convertido a escala de grises: {os.path.basename(entrada)} → {os.path.basename(salida)}")

        except Exception as e:
            self.log(f"❌ Error convirtiendo a escala de grises '{os.path.basename(entrada)}': {e}")
        finally:
            if ruta_temporal and os.path.exists(ruta_temporal):
                os.remove(ruta_temporal)

    def accion_pdf_a_escala_grises_individual(self):
        archivos = filedialog.askopenfilenames(
            parent=self.root,
            title="Seleccionar PDF(s) para convertir a escala de grises (REEMPLAZARÁ ORIGINALES)",
            filetypes=[("Archivos PDF", "*.pdf")]
        )
        if not archivos:
            self.log("Operación de conversión a escala de grises cancelada.")
            return
            
        confirm = messagebox.askyesno(
            "Confirmar Sobrescritura",
            f"¿Estás seguro de que quieres convertir {len(archivos)} archivo(s) a escala de grises?\n\n"
            "¡ADVERTENCIA! Esta acción REEMPLAZARÁ los archivos originales y no se puede deshacer.",
            parent=self.root
        )
        if not confirm:
            self.log("Sobrescritura de archivos cancelada por el usuario.")
            return

        total_files = len(archivos)
        self.log(f"--- Iniciando conversión (con reemplazo) a escala de grises para {total_files} archivo(s) ---")
        
        progress_win, progress_bar, progress_label = self._create_progress_window(
            "Convirtiendo y reemplazando...", total_files
        )
        try:
            for i, entrada in enumerate(archivos):
                progress_label.config(text=f"Convirtiendo: {os.path.basename(entrada)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                self._pdf_a_escala_grises_worker(entrada, entrada)
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log("--- Proceso de conversión a escala de grises finalizado. ---")
        messagebox.showinfo("Proceso Completado", "Conversión a escala de grises (con reemplazo) finalizada. Revise el registro.", parent=self.root)
    
    def accion_pdf_a_escala_grises_masivo_desde_resultados(self):
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Busca archivos primero para usar esta conversión masiva.", parent=self.root)
            return

        archivos_filtrados = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        
        if not archivos_filtrados:
            messagebox.showinfo("Información", "No se encontraron archivos PDF en los resultados para convertir.", parent=self.root)
            return

        confirm = messagebox.askyesno(
            "Confirmar Sobrescritura desde Resultados",
            f"¿Estás seguro de que quieres convertir los {len(archivos_filtrados)} PDFs encontrados a escala de grises?\n\n"
            "¡ADVERTENCIA! Esta acción REEMPLAZARÁ los archivos originales y no se puede deshacer.",
            parent=self.root
        )
        if not confirm:
            self.log("Sobrescritura desde resultados cancelada por el usuario.")
            return

        self.log(f"--- Iniciando conversión (con reemplazo) a escala de grises para {len(archivos_filtrados)} archivo(s) desde los resultados ---")
        progress_win, progress_bar, progress_label = self._create_progress_window(
            "Convirtiendo y reemplazando...", len(archivos_filtrados)
        )
        try:
            for i, entrada in enumerate(archivos_filtrados):
                progress_label.config(text=f"Convirtiendo: {os.path.basename(entrada)}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                self._pdf_a_escala_grises_worker(entrada, entrada)
        finally:
            if progress_win:
                progress_win.destroy()

        self.log("--- Proceso de conversión desde resultados finalizado. ---")
        messagebox.showinfo("Proceso Completado", "Conversión a escala de grises (con reemplazo) desde resultados finalizada.", parent=self.root)

    def accion_organizar_facturas_por_pdf(self):
        messagebox.showinfo(
            "Información - Organizar Facturas FEOV",
            "Esta función organiza archivos de facturas en dos pasos:\n\n"
            "1. Primero, selecciona la carpeta 'DESTINO' que contiene las subcarpetas finales (ej: una carpeta llamada '12345' que a su vez contiene un PDF '...FEOV12345.pdf').\n\n"
            "2. Luego, selecciona la carpeta 'ORIGEN' que contiene los archivos desorganizados que quieres mover.\n\n"
            "El programa moverá los archivos del 'ORIGEN' a la carpeta 'DESTINO' correcta si su nombre contiene el número de la factura FEOV.",
            parent=self.root
        )
        carpeta_destinos = filedialog.askdirectory(parent=self.root, title="PASO 1: Selecciona la carpeta principal de DESTINOS")
        if not carpeta_destinos:
            self.log("Operación cancelada. No se seleccionó carpeta de destinos.")
            return

        carpeta_origen_archivos = filedialog.askdirectory(parent=self.root, title="PASO 2: Selecciona la carpeta de ORIGEN con los archivos a mover")
        if not carpeta_origen_archivos:
            self.log("Operación cancelada. No se seleccionó carpeta de origen de archivos.")
            return

        self.log("Iniciando organización de facturas FEOV...")
        regex = re.compile(r'FEOV(\d+)', re.IGNORECASE)
        
        destinos_map = {}
        self.log("--- Analizando carpetas de destino para encontrar números de factura ---")
        
        try:
            list_carpetas_destino = [d for d in os.listdir(carpeta_destinos) if os.path.isdir(os.path.join(carpeta_destinos, d))]
        except FileNotFoundError:
            self.log(f"❌ Error: La carpeta de destinos '{carpeta_destinos}' no fue encontrada.")
            return

        for nombre_carpeta_destino in list_carpetas_destino:
            ruta_carpeta_destino = os.path.join(carpeta_destinos, nombre_carpeta_destino)
            for archivo in os.listdir(ruta_carpeta_destino):
                if archivo.lower().endswith('.pdf'):
                    match = regex.search(archivo)
                    if match:
                        numero_factura = match.group(1)
                        destinos_map[numero_factura] = ruta_carpeta_destino
                        self.log(f"  Encontrado destino para factura '{numero_factura}' en carpeta: '{nombre_carpeta_destino}'")
                        break 

        if not destinos_map:
            self.log("⚠️ No se encontraron PDFs con el patrón 'FEOV<numero>' en las carpetas de destino. Proceso detenido.")
            messagebox.showwarning("Sin Destinos", "No se encontró ningún PDF con el formato 'FEOV...' en las subcarpetas de la carpeta de destino seleccionada.", parent=self.root)
            return

        self.log(f"--- Buscando archivos para mover en '{carpeta_origen_archivos}' ---")
        movidos, errores, conflictos = 0, 0, 0

        total_files_to_scan = sum(len(files) for _, _, files in os.walk(carpeta_origen_archivos))
        if total_files_to_scan == 0:
            self.log("No se encontraron archivos en la carpeta de origen para procesar.")
            return

        progress_win, progress_bar, progress_label = self._create_progress_window("Organizando facturas...", total_files_to_scan)
        file_count = 0
        try:
            for root, _, files in os.walk(carpeta_origen_archivos):
                for file_to_move in files:
                    file_count += 1
                    progress_label.config(text=f"Escaneando: {file_to_move}")
                    progress_bar['value'] = file_count
                    self.root.update_idletasks()

                    for numero_factura, ruta_destino_final in destinos_map.items():
                        if numero_factura in file_to_move:
                            try:
                                ruta_origen_archivo = os.path.join(root, file_to_move)
                                ruta_final_archivo = os.path.join(ruta_destino_final, file_to_move)

                                if os.path.exists(ruta_final_archivo):
                                    self.log(f"  ⚠️ Conflicto: El archivo '{file_to_move}' ya existe en el destino. Se omite.")
                                    conflictos += 1
                                else:
                                    shutil.move(ruta_origen_archivo, ruta_destino_final)
                                    self.log(f"  ✅ Movido '{file_to_move}' a la carpeta '{os.path.basename(ruta_destino_final)}'")
                                    movidos += 1
                                break 
                            except Exception as e:
                                self.log(f"  ❌ Error moviendo '{file_to_move}': {e}")
                                errores += 1
                                break 
        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (f"Proceso de organización finalizado.\n\n"
                       f"- Archivos movidos: {movidos}\n"
                       f"- Conflictos (ya existían): {conflictos}\n"
                       f"- Errores: {errores}")
        self.log(f"--- Organización finalizada. Movidos: {movidos}, Conflictos: {conflictos}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def accion_crear_carpetas_desde_excel(self):
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel", filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
        except Exception as e:
            messagebox.showerror("Error de Archivo", f"No se pudo leer el archivo Excel:\n{e}", parent=self.root)
            return

        hoja_seleccionada = simpledialog.askstring(
            "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
            parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
        )
        if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return
        
        respuesta_filtro = messagebox.askyesnocancel(
            "Usar Filtros de Excel",
            "¿Desea crear carpetas solo para las filas VISIBLES (filtradas)?\n\n"
            "- 'Sí': Solo usará las filas que no están ocultas por un filtro.\n"
            "- 'No': Usará TODAS las filas de la hoja, ignorando cualquier filtro.\n"
            "- 'Cancelar': Abortará la operación.",
            parent=self.root
        )

        if respuesta_filtro is None:
            self.log("Operación cancelada por el usuario.")
            return

        usar_filas_visibles = respuesta_filtro

        nombres_carpetas_raw = []
        indice_columna = -1

        try:
            df_temp = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df_temp.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"
            
            letra_columna = simpledialog.askstring(
                "Seleccionar Columna", f"Ingrese la LETRA de la columna que contiene los nombres para las carpetas.\n(Opciones: {letras_rango})",
                parent=self.root
            )
            if not letra_columna or letra_columna.upper() not in letras_disponibles: return
            indice_columna = ord(letra_columna.upper()) - ord('A')
            
            if usar_filas_visibles:
                self.log("Leyendo solo filas visibles (filtradas) desde el Excel...")
                wb = openpyxl.load_workbook(archivo_excel, data_only=True)
                ws = wb[hoja_seleccionada]
                
                for i in range(2, ws.max_row + 1):
                    if not ws.row_dimensions[i].hidden:
                        valor_celda = ws.cell(row=i, column=indice_columna + 1).value
                        if valor_celda:
                            nombres_carpetas_raw.append(str(valor_celda))
            else:
                self.log("Leyendo todas las filas desde el Excel...")
                nombres_carpetas_raw = df_temp.iloc[:, indice_columna].dropna().astype(str).tolist()

        except Exception as e:
            messagebox.showerror("Error de Lectura", f"No se pudo leer la hoja '{hoja_seleccionada}':\n{e}", parent=self.root)
            return

        if not nombres_carpetas_raw:
            messagebox.showinfo("Sin Datos", "No se encontraron nombres de carpetas para crear con la configuración seleccionada.", parent=self.root)
            self.log("No se encontraron nombres en la columna y hoja especificadas.")
            return

        carpeta_base = filedialog.askdirectory(parent=self.root, title="Seleccionar carpeta donde crear las nuevas carpetas")
        if not carpeta_base: return
        
        creadas, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Creando carpetas...", len(nombres_carpetas_raw))
        try:
            for i, nombre in enumerate(nombres_carpetas_raw):
                nombre_base = "".join(c for c in nombre if c.isalnum() or c in " _-").rstrip()
                progress_label.config(text=f"Procesando: {nombre_base}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                if not nombre_base: continue

                ruta_final = os.path.join(carpeta_base, nombre_base)
                
                if os.path.exists(ruta_final):
                    contador = 2
                    nombre_consecutivo = f"{nombre_base} ({contador})"
                    ruta_final = os.path.join(carpeta_base, nombre_consecutivo)
                    
                    while os.path.exists(ruta_final):
                        contador += 1
                        nombre_consecutivo = f"{nombre_base} ({contador})"
                        ruta_final = os.path.join(carpeta_base, nombre_consecutivo)
                    
                    self.log(f"⚠️ El nombre '{nombre_base}' ya existía. Se creará como '{os.path.basename(ruta_final)}'")
                
                try:
                    os.makedirs(ruta_final)
                    self.log(f"✅ Creada: {os.path.basename(ruta_final)}")
                    creadas += 1
                except Exception as e:
                    self.log(f"❌ Error creando '{os.path.basename(ruta_final)}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()
        
        self.log(f"--- Proceso finalizado. Carpetas creadas: {creadas}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", f"Creación de carpetas finalizada.\n\n- Carpetas creadas: {creadas}\n- Errores: {errores}", parent=self.root)

    def accion_mover_archivos_por_coincidencia_nombre(self):
        carpeta_base = self.source_path.get()
        if not carpeta_base or not os.path.isdir(carpeta_base):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida que contenga los archivos y las carpetas de destino.", parent=self.root)
            return

        self.log("--- Iniciando el movimiento de archivos a carpetas por coincidencia de nombre ---")

        try:
            elementos = os.listdir(carpeta_base)
            archivos = [os.path.join(carpeta_base, e) for e in elementos if os.path.isfile(os.path.join(carpeta_base, e))]
            carpetas = [os.path.join(carpeta_base, e) for e in elementos if os.path.isdir(os.path.join(carpeta_base, e))]
        except Exception as e:
            self.log(f"❌ Error al leer el contenido de la carpeta: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta:\n{e}", parent=self.root)
            return

        if not archivos or not carpetas:
            self.log("No se encontraron suficientes archivos o carpetas para procesar en la ruta de origen.")
            return

        movidos, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Moviendo archivos...", len(archivos))
        try:
            for i, ruta_archivo in enumerate(archivos):
                nombre_archivo = os.path.basename(ruta_archivo)
                progress_label.config(text=f"Verificando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                for ruta_carpeta in carpetas:
                    nombre_carpeta = os.path.basename(ruta_carpeta)

                    if nombre_carpeta.lower() in nombre_archivo.lower():
                        try:
                            shutil.move(ruta_archivo, ruta_carpeta)
                            self.log(f"✅ Movido: '{nombre_archivo}' → a la carpeta '{nombre_carpeta}'")
                            movidos += 1
                            break
                        except Exception as e:
                            self.log(f"❌ Error al mover '{nombre_archivo}': {e}")
                            errores += 1
                            break
        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (f"Proceso finalizado.\n\n"
                       f"- Archivos movidos: {movidos}\n"
                       f"- Errores: {errores}")
        self.log(f"--- Proceso Finalizado. Movidos: {movidos}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def _prompt_for_excel_mapping_details(self):
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel con el mapeo", filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return None, None, None

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return None, None, None

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_origen = simpledialog.askstring("Columna de Origen", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETAS DE ORIGEN.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_origen or letra_col_origen.upper() not in letras_disponibles: return None, None, None

            letra_col_destino = simpledialog.askstring("Columna de Destino", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETAS DE DESTINO.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_destino or letra_col_destino.upper() not in letras_disponibles: return None, None, None

            col_origen_idx = ord(letra_col_origen.upper()) - ord('A')
            col_destino_idx = ord(letra_col_destino.upper()) - ord('A')

            return df, df.columns[col_origen_idx], df.columns[col_destino_idx]

        except Exception as e:
            self.log(f"❌ Error al procesar el archivo Excel: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al leer el Excel: {e}", parent=self.root)
            return None, None, None
            
    def accion_copiar_archivos_desde_mapeo_excel(self):
        self.log("--- Iniciando copia de archivos basada en mapeo de Excel (por subcarpetas) ---")
        
        df, col_origen, col_destino = self._prompt_for_excel_mapping_details()
        if df is None:
            self.log("Operación cancelada durante la selección del Excel.")
            return

        ruta_origen_base = filedialog.askdirectory(parent=self.root, title="PASO 1: Seleccione la CARPETA BASE que contiene las carpetas de Origen")
        if not ruta_origen_base:
            self.log("Operación cancelada. No se seleccionó carpeta base de origen.")
            return
        
        ruta_destino_base = filedialog.askdirectory(parent=self.root, title="PASO 2: Seleccione la CARPETA BASE que contiene las carpetas de Destino")
        if not ruta_destino_base:
            self.log("Operación cancelada. No se seleccionó carpeta base de destino.")
            return

        copiados, errores, conflictos, carpetas_no_encontradas = 0, 0, 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Copiando archivos según mapeo...", len(df))

        try:
            for index, fila in df.iterrows():
                progress_bar['value'] = index + 1
                
                if pd.isna(fila[col_origen]) or pd.isna(fila[col_destino]):
                    continue
                nombre_carpeta_origen = str(fila[col_origen]).strip()
                nombre_carpeta_destino = str(fila[col_destino]).strip()

                progress_label.config(text=f"Mapeando: {nombre_carpeta_origen} -> {nombre_carpeta_destino}")
                self.root.update_idletasks()
                
                ruta_origen_completa = os.path.join(ruta_origen_base, nombre_carpeta_origen)
                ruta_destino_completa = os.path.join(ruta_destino_base, nombre_carpeta_destino)

                if not os.path.isdir(ruta_origen_completa):
                    self.log(f"⚠️  Origen no encontrado: '{ruta_origen_completa}'. Fila {index+2} omitida.")
                    carpetas_no_encontradas += 1
                    continue
                if not os.path.isdir(ruta_destino_completa):
                    self.log(f"⚠️  Destino no encontrado: '{ruta_destino_completa}'. Fila {index+2} omitida.")
                    carpetas_no_encontradas += 1
                    continue
                
                for nombre_archivo in os.listdir(ruta_origen_completa):
                    ruta_archivo_origen = os.path.join(ruta_origen_completa, nombre_archivo)
                    if os.path.isfile(ruta_archivo_origen):
                        ruta_archivo_destino = os.path.join(ruta_destino_completa, nombre_archivo)
                        
                        if os.path.exists(ruta_archivo_destino):
                            self.log(f"  -> Conflicto: '{nombre_archivo}' ya existe en el destino. Omitido.")
                            conflictos += 1
                            continue
                        
                        try:
                            shutil.copy2(ruta_archivo_origen, ruta_archivo_destino)
                            self.log(f"  -> Copiado: '{nombre_archivo}' a la carpeta '{nombre_carpeta_destino}'")
                            copiados += 1
                        except Exception as e:
                            self.log(f"  -> ❌ Error al copiar '{nombre_archivo}': {e}")
                            errores += 1
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen_msg = (
            f"Proceso de copia por mapeo finalizado.\n\n"
            f"- Archivos copiados exitosamente: {copiados}\n"
            f"- Conflictos (archivos ya existentes): {conflictos}\n"
            f"- Carpetas no encontradas: {carpetas_no_encontradas}\n"
            f"- Errores de copia: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def _prompt_for_excel_file_mapping_details(self):
        """Pide al usuario los detalles para el mapeo Excel de IDENTIFICADOR a CARPETA."""
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel con el mapeo (Identificador -> Carpeta)",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return None, None, None

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return None, None, None

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_archivo = simpledialog.askstring("Columna de Identificadores", f"Ingrese la LETRA de la columna con los IDENTIFICADORES DE ARCHIVO (ej: FEOV123).\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_archivo or letra_col_archivo.upper() not in letras_disponibles: return None, None, None

            letra_col_carpeta = simpledialog.askstring("Columna de Carpetas Destino", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETA DE DESTINO.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return None, None, None

            col_archivo_idx = ord(letra_col_archivo.upper()) - ord('A')
            col_carpeta_idx = ord(letra_col_carpeta.upper()) - ord('A')

            return df, df.columns[col_archivo_idx], df.columns[col_carpeta_idx]

        except Exception as e:
            self.log(f"❌ Error al procesar el archivo Excel: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al leer el Excel: {e}", parent=self.root)
            return None, None, None

    def accion_copiar_archivos_desde_raiz_mapeo_excel(self):
        """Copia archivos desde la RAÍZ de una carpeta a subcarpetas, usando coincidencia parcial de nombres desde un Excel."""
        self.log("--- Iniciando copia de archivos desde la RAÍZ (Mapeo Parcial) ---")
        
        df, col_identificador, col_carpeta = self._prompt_for_excel_file_mapping_details()
        if df is None:
            self.log("Operación cancelada durante la selección del Excel.")
            return

        ruta_origen_raiz = filedialog.askdirectory(parent=self.root, title="PASO 1: Seleccione la CARPETA RAÍZ que contiene los archivos de Origen")
        if not ruta_origen_raiz:
            self.log("Operación cancelada. No se seleccionó carpeta raíz de origen.")
            return
        
        ruta_destino_base = filedialog.askdirectory(parent=self.root, title="PASO 2: Seleccione la CARPETA BASE donde se encuentran las carpetas de Destino")
        if not ruta_destino_base:
            self.log("Operación cancelada. No se seleccionó carpeta base de destino.")
            return

        copiados, errores, conflictos, no_encontrados = 0, 0, 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Copiando por mapeo parcial...", len(df))

        try:
            for index, fila in df.iterrows():
                progress_bar['value'] = index + 1
                
                if pd.isna(fila[col_identificador]) or pd.isna(fila[col_carpeta]):
                    continue
                
                identificador_archivo = str(fila[col_identificador]).strip()
                nombre_carpeta_destino = str(fila[col_carpeta]).strip()

                if not identificador_archivo or not nombre_carpeta_destino:
                    continue

                progress_label.config(text=f"Buscando por: {identificador_archivo}")
                self.root.update_idletasks()
                
                ruta_archivo_encontrado = None
                nombre_archivo_real = ""
                
                # Buscamos archivos en la raíz del origen
                for f_name in os.listdir(ruta_origen_raiz):
                    if os.path.isfile(os.path.join(ruta_origen_raiz, f_name)):
                        # <<< INICIO DE LA MODIFICACIÓN: Búsqueda case-insensitive >>>
                        if identificador_archivo.lower() in f_name.lower():
                        # <<< FIN DE LA MODIFICACIÓN >>>
                            nombre_archivo_real = f_name
                            ruta_archivo_encontrado = os.path.join(ruta_origen_raiz, f_name)
                            break 

                if not ruta_archivo_encontrado:
                    self.log(f"⚠️ No se encontró archivo que contenga '{identificador_archivo}' en la raíz. Fila {index+2} omitida.")
                    no_encontrados += 1
                    continue
                
                self.log(f"  -> Coincidencia para '{identificador_archivo}': Se usará el archivo '{nombre_archivo_real}'")
                
                ruta_carpeta_destino_completa = os.path.join(ruta_destino_base, nombre_carpeta_destino)
                
                try:
                    os.makedirs(ruta_carpeta_destino_completa, exist_ok=True)
                except Exception as e:
                    self.log(f"❌ Error creando carpeta destino '{nombre_carpeta_destino}': {e}. Fila {index+2} omitida.")
                    errores += 1
                    continue

                ruta_archivo_destino = os.path.join(ruta_carpeta_destino_completa, nombre_archivo_real)

                if os.path.exists(ruta_archivo_destino):
                    self.log(f"  -> Conflicto: '{nombre_archivo_real}' ya existe en '{nombre_carpeta_destino}'. Omitido.")
                    conflictos += 1
                    continue
                
                try:
                    shutil.copy2(ruta_archivo_encontrado, ruta_archivo_destino)
                    self.log(f"  -> ✅ Copiado: '{nombre_archivo_real}' a la carpeta '{nombre_carpeta_destino}'")
                    copiados += 1
                except Exception as e:
                    self.log(f"  -> ❌ Error al copiar '{nombre_archivo_real}': {e}")
                    errores += 1
        
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen_msg = (
            f"Proceso de copia por mapeo parcial finalizado.\n\n"
            f"- Archivos copiados exitosamente: {copiados}\n"
            f"- Identificadores no encontrados: {no_encontrados}\n"
            f"- Conflictos (archivos ya existentes): {conflictos}\n"
            f"- Errores de copia/creación de carpetas: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
    
    def accion_copiar_archivo_a_subcarpetas(self):
        self.log("--- Iniciando copia de un archivo a múltiples subcarpetas ---")

        archivo_a_copiar = filedialog.askopenfilename(
            parent=self.root,
            title="PASO 1: Selecciona el ARCHIVO que quieres copiar"
        )
        if not archivo_a_copiar:
            self.log("Operación cancelada. No se seleccionó ningún archivo.")
            return

        carpeta_destino_base = filedialog.askdirectory(
            parent=self.root,
            title="PASO 2: Selecciona la CARPETA que contiene las subcarpetas de destino"
        )
        if not carpeta_destino_base:
            self.log("Operación cancelada. No se seleccionó la carpeta de destino.")
            return

        try:
            subcarpetas = [os.path.join(carpeta_destino_base, d) for d in os.listdir(carpeta_destino_base) if os.path.isdir(os.path.join(carpeta_destino_base, d))]
        except Exception as e:
            self.log(f"❌ Error al leer las subcarpetas de '{carpeta_destino_base}': {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta de destino:\n{e}", parent=self.root)
            return
        
        if not subcarpetas:
            self.log("No se encontraron subcarpetas en la ruta seleccionada.")
            messagebox.showinfo("Sin Carpetas", "No se encontraron subcarpetas en la ruta de destino para realizar la copia.", parent=self.root)
            return

        self.log(f"Se copiará el archivo '{os.path.basename(archivo_a_copiar)}' en {len(subcarpetas)} carpetas.")

        copiados, conflictos, errores = 0, 0, 0
        nombre_archivo_origen = os.path.basename(archivo_a_copiar)
        progress_win, progress_bar, progress_label = self._create_progress_window(f"Copiando '{nombre_archivo_origen}'...", len(subcarpetas))

        try:
            for i, ruta_subcarpeta in enumerate(subcarpetas):
                nombre_subcarpeta = os.path.basename(ruta_subcarpeta)
                progress_label.config(text=f"Copiando a: {nombre_subcarpeta}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                ruta_destino_final = os.path.join(ruta_subcarpeta, nombre_archivo_origen)

                if os.path.exists(ruta_destino_final):
                    self.log(f"  ⚠️ Conflicto: '{nombre_archivo_origen}' ya existe en '{nombre_subcarpeta}'. Omitido.")
                    conflictos += 1
                    continue

                try:
                    shutil.copy2(archivo_a_copiar, ruta_destino_final)
                    self.log(f"  ✅ Copiado a '{nombre_subcarpeta}'")
                    copiados += 1
                except Exception as e:
                    self.log(f"  ❌ Error al copiar a '{nombre_subcarpeta}': {e}")
                    errores += 1

        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (
            f"Proceso de copia finalizado.\n\n"
            f"- Archivo copiado exitosamente en: {copiados} carpetas\n"
            f"- Conflictos (archivos ya existentes): {conflictos}\n"
            f"- Errores de copia: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    def _prompt_for_url_download_details(self):
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar archivo Excel con los datos de descarga", filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return None, None, None

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return None, None, None

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_id = simpledialog.askstring("Columna de ID", f"Ingrese la LETRA de la columna con los NÚMEROS para la URL.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_id or letra_col_id.upper() not in letras_disponibles: return None, None, None

            letra_col_carpeta = simpledialog.askstring("Columna de Carpeta", f"Ingrese la LETRA de la columna con los NOMBRES DE CARPETA de destino.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return None, None, None

            col_id_idx = ord(letra_col_id.upper()) - ord('A')
            col_carpeta_idx = ord(letra_col_carpeta.upper()) - ord('A')

            return df, df.columns[col_id_idx], df.columns[col_carpeta_idx]

        except Exception as e:
            self.log(f"❌ Error al procesar el archivo Excel: {e}")
            messagebox.showerror("Error", f"Ocurrió un error al leer el Excel: {e}", parent=self.root)
            return None, None, None

    def accion_descargar_firmas_url_excel(self):
        self.log("--- Iniciando descarga de firmas desde URL por Excel ---")
        
        df, col_id, col_carpeta = self._prompt_for_url_download_details()
        if df is None:
            self.log("Operación cancelada durante la selección del Excel.")
            return

        base_path = filedialog.askdirectory(parent=self.root, title="Seleccione la CARPETA BASE donde se crearán las carpetas de destino")
        if not base_path:
            self.log("Operación cancelada. No se seleccionó carpeta base de destino.")
            return
        
        base_url = "https://oportunidaddevida.com/opvcitas/admisionescall/firmas/"
        descargados, no_encontrados, errores_red = 0, 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Descargando firmas...", len(df))

        try:
            for index, row in df.iterrows():
                progress_bar['value'] = index + 1
                
                if pd.isna(row[col_id]) or pd.isna(row[col_carpeta]):
                    self.log(f"⚠️ Fila {index+2} omitida: contiene datos vacíos.")
                    continue
                
                id_firma = str(row[col_id]).strip()
                nombre_carpeta = str(row[col_carpeta]).strip()
                
                progress_label.config(text=f"Procesando ID: {id_firma}")
                self.root.update_idletasks()

                url_completa = f"{base_url}{id_firma}.png"
                ruta_carpeta_destino = os.path.join(base_path, nombre_carpeta)
                os.makedirs(ruta_carpeta_destino, exist_ok=True)
                
                try:
                    response = requests.get(url_completa, stream=True, timeout=15)
                    
                    if response.status_code == 200:
                        ruta_salida_jpg = os.path.join(ruta_carpeta_destino, "firma.jpg")
                        
                        img = Image.open(io.BytesIO(response.content))
                        img_rgb = img.convert('RGB')
                        img_rgb.save(ruta_salida_jpg, 'JPEG')
                        
                        self.log(f"✅ Firma descargada para ID {id_firma} en carpeta '{nombre_carpeta}'")
                        descargados += 1
                    
                    else:
                        ruta_txt_error = os.path.join(ruta_carpeta_destino, "no tiene firma.txt")
                        with open(ruta_txt_error, 'w', encoding='utf-8') as f:
                            f.write(f"No se encontró firma en la URL: {url_completa}\n")
                            f.write(f"Código de estado: {response.status_code}\n")
                        
                        self.log(f"⚠️ No se encontró firma en URL para ID {id_firma} (Estado: {response.status_code}). Se creó TXT.")
                        no_encontrados += 1

                except requests.exceptions.RequestException as e:
                    self.log(f"❌ Error de red para ID {id_firma}: {e}")
                    errores_red += 1

        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (
            f"Proceso de descarga de firmas finalizado.\n\n"
            f"- Firmas descargadas y guardadas: {descargados}\n"
            f"- Firmas no encontradas (404, etc.): {no_encontrados}\n"
            f"- Errores de red (timeout, etc.): {errores_red}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

    # === INICIO NUEVA FUNCIONALIDAD OVIDA ===


    def accion_descargar_historias_hospitalizacion_ovida(self):
        """Orquesta la descarga de historias de hospitalización desde OVIDA."""
        if not self.selenium_available:
            messagebox.showerror(
                "Dependencias Faltantes",
                "Esta función requiere 'selenium' y 'webdriver-manager'.\n\n"
                "Instálalos con: pip install selenium webdriver-manager",
                parent=self.root
            )
            self.log("❌ Faltan dependencias de Selenium para la descarga de hospitalización.")
            return

        self.log("--- Iniciando descarga de Historias de Hospitalización (OVIDA) ---")
        
        # 1. Pedir archivo Excel y mapeo
        archivo_excel = filedialog.askopenfilename(
            parent=self.root, title="Seleccionar Excel con datos de hospitalización",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel:
            self.log("Operación cancelada. No se seleccionó archivo Excel.")
            return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja", f"Ingrese el nombre de la hoja a usar.\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return
            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            # Pedir mapeo de columnas
            letra_col_estudio = simpledialog.askstring("Columna 'Nro ESTUDIO'", f"Ingrese la letra de la columna 'Nro ESTUDIO'.\n(Opciones: {letras_rango})", parent=self.root)
            letra_col_ingreso = simpledialog.askstring("Columna 'FECHA DE INGRESO'", f"Ingrese la letra de la columna 'FECHA DE INGRESO'.\n(Opciones: {letras_rango})", parent=self.root)
            letra_col_egreso = simpledialog.askstring("Columna 'FECHA EGRESO'", f"Ingrese la letra de la columna 'FECHA EGRESO'.\n(Opciones: {letras_rango})", parent=self.root)
            letra_col_carpeta = simpledialog.askstring("Columna 'NOMBRE DE LAS CARPETAS'", f"Ingrese la letra de la columna 'NOMBRE DE LAS CARPETAS'.\n(Opciones: {letras_rango})", parent=self.root)

            if not all([letra_col_estudio, letra_col_ingreso, letra_col_egreso, letra_col_carpeta]):
                self.log("Operación cancelada, no se especificaron todas las columnas.")
                return
            
            col_map = {
                'estudio': df.columns[ord(letra_col_estudio.upper()) - ord('A')],
                'ingreso': df.columns[ord(letra_col_ingreso.upper()) - ord('A')],
                'egreso': df.columns[ord(letra_col_egreso.upper()) - ord('A')],
                'carpeta': df.columns[ord(letra_col_carpeta.upper()) - ord('A')],
            }

        except Exception as e:
            self.log(f"❌ Error procesando el archivo Excel: {e}")
            messagebox.showerror("Error de Excel", f"No se pudo leer el archivo o las columnas.\n\n{e}", parent=self.root)
            return

        # 2. Pedir carpeta de guardado
        save_path = filedialog.askdirectory(parent=self.root, title="Seleccione la carpeta BASE para guardar las historias")
        if not save_path:
            self.log("Operación cancelada. No se seleccionó carpeta de guardado.")
            return

        # 3. Iniciar Selenium y esperar login
        webdriver = self.selenium_deps["webdriver"]
        Service = self.selenium_deps["Service"]
        ChromeDriverManager = self.selenium_deps["ChromeDriverManager"]
        driver = None
        
        try:
            self.log("Abriendo navegador. Por favor, inicie sesión en OVIDA.")
            # Iniciar navegador visible
            options = webdriver.ChromeOptions()
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)
            driver.get("https://ovidazs.siesacloud.com/ZeusSalud/ips/iniciando.php")

            # Pausar y esperar que el usuario confirme
            if not messagebox.askokcancel(
                "Acción Requerida: Iniciar Sesión",
                "Se ha abierto una ventana de Chrome.\n\n"
                "1. Inicie sesión en la plataforma OVIDA.\n"
                "2. Seleccione el punto de atención si es necesario.\n\n"
                "Cuando esté en la página principal, haga clic en 'Aceptar' para comenzar las descargas.",
                parent=self.root):
                self.log("Operación cancelada por el usuario durante el inicio de sesión.")
                driver.quit()
                return

            self.log("Inicio de sesión confirmado por el usuario. Empezando descargas...")
            
            # 4. Bucle de procesamiento
            descargados, errores, conflictos = 0, 0, 0
            total_filas = len(df)
            
            progress_win, progress_bar, progress_label = self._create_progress_window("Descargando Historias...", total_filas)

            for index, row in df.iterrows():
                progress_bar['value'] = index + 1
                
                # Extraer y limpiar datos de la fila
                try:
                    nro_estudio = str(int(row[col_map['estudio']])).strip()
                    # Convertir fechas
                    fecha_ingreso_dt = pd.to_datetime(row[col_map['ingreso']])
                    fecha_ingreso = fecha_ingreso_dt.strftime('%Y/%m/%d')
                    fecha_egreso_dt = pd.to_datetime(row[col_map['egreso']])
                    fecha_egreso = fecha_egreso_dt.strftime('%Y/%m/%d')
                    nombre_carpeta = str(row[col_map['carpeta']]).strip()
                except Exception as ex:
                    self.log(f"⚠️ Fila {index+2} omitida por datos inválidos o vacíos (Estudio, Fechas, Carpeta). Error: {ex}")
                    errores += 1
                    continue
                
                if not all([nro_estudio, fecha_ingreso, fecha_egreso, nombre_carpeta]):
                    self.log(f"⚠️ Fila {index+2} omitida por tener celdas vacías.")
                    errores += 1
                    continue

                progress_label.config(text=f"Procesando estudio: {nro_estudio}")
                self.root.update_idletasks()

                # Construir URL
                base_url = "https://ovidazs.siesacloud.com/ZeusSalud/Reportes/Cliente//html/reporte_historia_general.php"
                params = {
                    'estudio': nro_estudio, 'fecha_inicio': fecha_ingreso, 'fecha_fin': fecha_egreso,
                    'verHC': 1, 'verEvo': 1, 'verPar': 1, 'ImprimirOrdenamiento': 1,
                    'ImprimirNotasPcte': 0, 'ImprimirSolOrdenesExt': 1, 'ImprimirGraficasHC': 1,
                    'ImprimirFormatos': 1, 'ImprimirRegistroAdmon': 1, 'ImprimirNovedad': 0,
                    'ImprimirRecomendaciones': 0, 'ImprimirDescripcionQX': 0, 'ImprimirNotasEnfermeria': 1,
                    'ImprimirSignosVitales': 0, 'ImprimirLog': 0, 'ImprimirEpicrisisSinHC': 0
                }
                
                full_url = f"{base_url}?{urllib.parse.urlencode(params)}"

                # Crear carpeta destino y verificar si el archivo ya existe
                dest_folder = os.path.join(save_path, nombre_carpeta)
                os.makedirs(dest_folder, exist_ok=True)
                final_file_path = os.path.join(dest_folder, f"HC_{nro_estudio}.pdf")

                if os.path.exists(final_file_path):
                    self.log(f"  -> Conflicto: El archivo '{os.path.basename(final_file_path)}' ya existe en '{nombre_carpeta}'. Omitido.")
                    conflictos += 1
                    continue

                # Descargar
                try:
                    self.log(f"  > Navegando para descargar estudio {nro_estudio}")
                    driver.get(full_url)
                    time.sleep(2) # Pequeña espera para asegurar que la página se renderiza
                    
                    pdf_b64 = driver.execute_cdp_cmd("Page.printToPDF", {
                        "landscape": False, "printBackground": True,
                        "paperWidth": 8.5, "paperHeight": 11,
                        "marginTop": 0.4, "marginBottom": 0.4, "marginLeft": 0.4, "marginRight": 0.4
                    })
                    
                    pdf_data = base64.b64decode(pdf_b64['data'])
                    
                    with open(final_file_path, 'wb') as f:
                        f.write(pdf_data)
                    
                    self.log(f"  ✅ Descargado: '{os.path.basename(final_file_path)}' en carpeta '{nombre_carpeta}'")
                    descargados += 1

                except Exception as e:
                    self.log(f"  ❌ Error al descargar el estudio {nro_estudio}: {str(e).split(chr(10))[0]}")
                    errores += 1
            
            # Finalización
            progress_win.destroy()
            resumen_msg = (f"Proceso finalizado.\n\n"
                           f"- Historias descargadas: {descargados}\n"
                           f"- Errores: {errores}\n"
                           f"- Conflictos (ya existían): {conflictos}")
            self.log(f"--- {resumen_msg.replace(chr(10)*2, ' ').replace(chr(10), ' | ')} ---")
            messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)

        except Exception as e:
            self.log(f"❌ Ocurrió un error inesperado durante el proceso: {e}")
            messagebox.showerror("Error Crítico", f"Ha ocurrido un error:\n\n{e}", parent=self.root)
        finally:
            if driver:
                driver.quit()
                self.log("Navegador cerrado.")

    def _proceso_descarga_en_hilo(self, df, col_tid, col_num, col_estudio, download_path):
        """Maneja el bucle de descarga y la barra de progreso."""
        descargados = 0
        errores = 0
        total_pacientes = len(df)
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Descargando Historias Clínicas...", total_pacientes)
        
        for index, row in df.iterrows():
            progress_bar['value'] = index + 1
            
            tipo_id = str(row[col_tid]).strip()
            numero_id = str(row[col_num]).strip()
            nro_estudio = str(row[col_estudio]).strip()
            
            progress_label.config(text=f"Procesando: {numero_id}")
            self.root.update_idletasks()

            if not all([tipo_id, numero_id, nro_estudio]):
                self.log(f"⚠️ Fila {index+2} omitida por datos faltantes.")
                errores += 1
                continue
            
            self.log(f"--- Intentando descargar historia para ID: {numero_id}, Estudio: {nro_estudio} ---")
            
            if self._worker_descargar_historia(tipo_id, numero_id, nro_estudio, download_path):
                descargados += 1
            else:
                errores += 1
        
        progress_win.destroy()
        resumen_msg = (
            f"Proceso de descarga de historias finalizado.\n\n"
            f"- Historias descargadas: {descargados}\n"
            f"- Registros con error: {errores}"
        )
        self.log(f"--- {resumen_msg.replace(chr(10)*2, ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)


    def _worker_descargar_historia(self, tipo_id, numero_id, nro_estudio, download_path):
        """Realiza la automatización web para un solo paciente."""
        # Extraer las dependencias de Selenium
        webdriver = self.selenium_deps["webdriver"]
        By = self.selenium_deps["By"]
        WebDriverWait = self.selenium_deps["WebDriverWait"]
        Select = self.selenium_deps["Select"]
        EC = self.selenium_deps["EC"]
        Service = self.selenium_deps["Service"]
        Keys = self.selenium_deps["Keys"]
        ChromeDriverManager = self.selenium_deps["ChromeDriverManager"]
        
        # Configurar opciones de Chrome para descarga automática
        options = webdriver.ChromeOptions()
        prefs = {
            "download.default_directory": download_path,
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "plugins.always_open_pdf_externally": True  # Importante para forzar la descarga
        }
        options.add_experimental_option("prefs", prefs)
        options.add_argument("--headless")  # Ejecutar en segundo plano
        options.add_argument("--log-level=3") # Suprimir logs de consola de selenium
        
        driver = None
        try:
            # Instalar/usar el driver de Chrome automáticamente
            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=options)
            wait = WebDriverWait(driver, 20)  # Aumentar tiempo de espera a 20s

            # PASO 2: Ingresar al link
            url = "https://ovidazs.siesacloud.com/ZeusSalud/ips/App/Vistas/Hc/Historia.php"
            driver.get(url)

            # Seleccionar tipo de documento
            tipo_doc_select_el = wait.until(EC.presence_of_element_located((By.ID, "tipo-documento")))
            select = Select(tipo_doc_select_el)
            select.select_by_value(tipo_id) # Usar el valor (ej: 'CC') en lugar del texto visible
            
            # Ingresar número y presionar Enter
            num_doc_input = driver.find_element(By.ID, "documento")
            num_doc_input.send_keys(numero_id)
            num_doc_input.send_keys(Keys.RETURN)
            self.log("  > Datos de paciente enviados.")

            # PASO 3: Desplegar panel derecho y seleccionar estudio
            panel_button = wait.until(EC.element_to_be_clickable((By.ID, "iconoFrameDer")))
            panel_button.click()
            
            # Esperar y hacer clic en el número de estudio
            estudio_link = wait.until(EC.element_to_be_clickable((By.XPATH, f"//div[@class='letraDisplay' and contains(text(), '{nro_estudio}')]")))
            estudio_link.click()
            self.log(f"  > Estudio '{nro_estudio}' seleccionado.")
            
            # PASO 4: Pestaña Imprimir y opciones
            imprimir_tab = wait.until(EC.element_to_be_clickable((By.XPATH, "//div[contains(@class, 'tabsOrdinarios')]/ul/li/a[contains(@onclick, 'Imprimir()')]")))
            imprimir_tab.click()

            # Cambiar al iframe del modal de impresión
            iframe = wait.until(EC.presence_of_element_located((By.XPATH, "//iframe[contains(@src, 'reportes/impresionHC.php')]")))
            driver.switch_to.frame(iframe)

            # Clic en botón 'General'
            general_button = wait.until(EC.element_to_be_clickable((By.NAME, "opcion")))
            general_button.click()
            
            # Marcar checkbox 'Imp. Notas de enfermería'
            # El input no tiene id, lo buscamos por el texto de su label
            notas_enf_check = wait.until(EC.element_to_be_clickable((By.XPATH, "//td[contains(text(), 'Imp. Notas de enfermería')]/input[@type='checkbox']")))
            if not notas_enf_check.is_selected():
                 notas_enf_check.click()
            self.log("  > Opciones de impresión seleccionadas.")

            # Preparar para la descarga
            files_before = os.listdir(download_path)
            
            # Clic en el icono de imprimir
            imprimir_icon = wait.until(EC.element_to_be_clickable((By.XPATH, "//img[@title='Imprimir']")))
            imprimir_icon.click()
            
            # Esperar a que la descarga finalice
            self.log("  > Descarga iniciada. Esperando finalización...")
            time_waited = 0
            download_completed = False
            while time_waited < 60: # Esperar máximo 60 segundos por la descarga
                files_after = os.listdir(download_path)
                new_files = [f for f in files_after if f not in files_before and not f.endswith('.crdownload')]
                if new_files:
                    downloaded_file = new_files[0]
                    # Renombrar el archivo
                    os.rename(
                        os.path.join(download_path, downloaded_file),
                        os.path.join(download_path, f"{numero_id}.pdf")
                    )
                    self.log(f"✅ Descarga completada y renombrada a '{numero_id}.pdf'")
                    download_completed = True
                    break
                time.sleep(1)
                time_waited += 1

            if not download_completed:
                self.log(f"❌ Error: La descarga no se completó en 60 segundos para el ID {numero_id}.")
                return False

            return True

        except Exception as e:
            # Captura cualquier error de Selenium (ej: elemento no encontrado) o del sistema de archivos
            self.log(f"❌ Error procesando ID {numero_id}: {type(e).__name__} - {str(e).split(chr(10))[0]}")
            return False
        
        finally:
            # Asegurarse de cerrar el navegador sin importar lo que pase
            if driver:
                driver.quit()
    # === FIN NUEVA FUNCIONALIDAD OVIDA ===

    def accion_consolidar_archivos_subcarpetas(self):
        base_path = filedialog.askdirectory(parent=self.root, title="Seleccionar la carpeta base que contiene las carpetas a procesar")
        if not base_path:
            self.log("Operación de consolidación cancelada.")
            return

        try:
            main_folders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not main_folders:
            self.log("No se encontraron carpetas para procesar.")
            return

        self.log(f"--- Iniciando consolidación de archivos para {len(main_folders)} carpetas en '{base_path}' ---")
        
        copiados, conflictos, errores = 0, 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Consolidando carpetas...", len(main_folders))

        try:
            for i, folder_name in enumerate(main_folders):
                main_folder_path = os.path.join(base_path, folder_name)
                
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                self.log(f"-> Procesando carpeta: '{folder_name}'")
                
                for sub_root, _, files in os.walk(main_folder_path):
                    if sub_root == main_folder_path:
                        continue

                    for file_name in files:
                        source_path = os.path.join(sub_root, file_name)
                        dest_path = os.path.join(main_folder_path, file_name)

                        try:
                            if os.path.exists(dest_path):
                                self.log(f"  ⚠️ Conflicto: El archivo '{file_name}' ya existe en la raíz de '{folder_name}'. Se omite la copia.")
                                conflictos += 1
                                continue
                            
                            shutil.copy2(source_path, dest_path)
                            self.log(f"  ✅ Copiado: '{file_name}' a la carpeta '{folder_name}'")
                            copiados += 1
                        
                        except Exception as e:
                            self.log(f"  ❌ Error al copiar '{file_name}': {e}")
                            errores += 1
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen = (f"Proceso de consolidación finalizado.\n\n"
                   f"- Archivos copiados: {copiados}\n"
                   f"- Conflictos (archivos omitidos): {conflictos}\n"
                   f"- Errores: {errores}")
        self.log(f"--- {resumen.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)

    def accion_crear_firma_digital(self):
        base_path = filedialog.askdirectory(
            parent=self.root,
            title="Seleccionar carpeta base con las carpetas a procesar"
        )
        if not base_path:
            self.log("Operación de creación de firma cancelada.")
            return

        font_path = filedialog.askopenfilename(
            parent=self.root,
            title="Selecciona un archivo de fuente para la firma (ej: .ttf, .otf)",
            filetypes=[("Archivos de Fuente", "*.ttf *.otf"), ("Todos los archivos", "*.*")]
        )
        if not font_path:
            self.log("No se seleccionó una fuente. Operación cancelada.")
            return

        font_size = simpledialog.askinteger(
            "Tamaño de Fuente", "Introduce el tamaño de la fuente para la firma:",
            parent=self.root, initialvalue=60, minvalue=10, maxvalue=500
        )
        if not font_size:
            self.log("No se especificó un tamaño de fuente. Operación cancelada.")
            return

        try:
            folders_to_process = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not folders_to_process:
            self.log("No se encontraron carpetas para procesar.")
            return

        self.log(f"--- Iniciando creación de firmas digitales para {len(folders_to_process)} carpetas ---")
        self.log(f"Usando fuente: {os.path.basename(font_path)}")
        
        creadas, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Creando firmas...", len(folders_to_process))

        try:
            font = ImageFont.truetype(font_path, font_size)

            for i, folder_name in enumerate(folders_to_process):
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()
                
                try:
                    text_to_draw = folder_name

                    temp_img = Image.new('RGB', (1, 1))
                    draw_temp = ImageDraw.Draw(temp_img)
                    bbox = draw_temp.textbbox((0, 0), text_to_draw, font=font)
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1]

                    padding = 20
                    img_width = text_width + (2 * padding)
                    img_height = text_height + (2 * padding)
                    
                    image = Image.new('RGB', (img_width, img_height), color='white')
                    draw = ImageDraw.Draw(image)

                    draw.text((padding, padding), text_to_draw, font=font, fill='black')
                    
                    main_folder_path = os.path.join(base_path, folder_name)
                    tipografia_folder = os.path.join(main_folder_path, "tipografia")
                    output_path = os.path.join(tipografia_folder, "firma.jpg")

                    os.makedirs(tipografia_folder, exist_ok=True)
                    
                    image.save(output_path, 'JPEG')
                    
                    self.log(f"✅ Firma creada para '{folder_name}' en '{output_path}'")
                    creadas += 1

                except Exception as e:
                    self.log(f"❌ Error creando firma en '{folder_name}': {e}")
                    errores += 1
        except IOError:
             self.log(f"❌ Error: No se pudo cargar el archivo de fuente '{font_path}'. Asegúrate de que es un archivo de fuente válido.")
             messagebox.showerror("Error de Fuente", "El archivo seleccionado no pudo ser cargado como una fuente. La operación ha sido cancelada.", parent=self.root)
             errores += len(folders_to_process)
        finally:
            if progress_win:
                progress_win.destroy()
        
        resumen = f"Proceso de creación de firmas finalizado.\n\n- Firmas creadas: {creadas}\n- Errores: {errores}"
        self.log(f"--- {resumen.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)
    
    def accion_firmar_docx_con_imagen(self):
        base_path = filedialog.askdirectory(
            parent=self.root,
            title="Seleccionar carpeta base con las carpetas a procesar"
        )
        if not base_path:
            self.log("Operación de firma de DOCX cancelada.")
            return

        docx_filename = simpledialog.askstring(
            "Nombre del Documento", "Introduce el nombre del archivo DOCX a firmar:",
            parent=self.root, initialvalue="plantilla.docx"
        )
        if not docx_filename: return

        signature_filename = simpledialog.askstring(
            "Nombre de la Firma", "Introduce el nombre del archivo de imagen de la firma (JPG, PNG):",
            parent=self.root, initialvalue="firma.jpg"
        )
        if not signature_filename: return

        try:
            folders_to_process = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]
        except Exception as e:
            self.log(f"❌ Error al leer el directorio base: {e}")
            messagebox.showerror("Error", f"No se pudo leer el contenido de la carpeta seleccionada:\n{e}", parent=self.root)
            return
        
        if not folders_to_process:
            self.log("No se encontraron carpetas para procesar.")
            return

        self.log(f"--- Iniciando reemplazo de firma en DOCX para {len(folders_to_process)} carpetas ---")
        
        procesados, errores = 0, 0
        progress_win, progress_bar, progress_label = self._create_progress_window("Modificando documentos DOCX...", len(folders_to_process))

        try:
            for i, folder_name in enumerate(folders_to_process):
                progress_label.config(text=f"Procesando: {folder_name}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                folder_path = os.path.join(base_path, folder_name)
                docx_path = os.path.join(folder_path, docx_filename)
                
                signature_path = None
                possible_sig_paths = [
                    os.path.join(folder_path, signature_filename),
                    os.path.join(folder_path, "tipografia", signature_filename)
                ]
                for path in possible_sig_paths:
                    if os.path.exists(path):
                        signature_path = path
                        break

                if not os.path.exists(docx_path):
                    self.log(f"⚠️ DOCX '{docx_filename}' no encontrado en '{folder_name}'. Se omite.")
                    errores += 1
                    continue
                
                if not signature_path:
                    self.log(f"⚠️ Firma '{signature_filename}' no encontrada en '{folder_name}' o sus subcarpetas. Se omite.")
                    errores += 1
                    continue
                
                if self._reemplazar_firma_en_docx_worker(docx_path, signature_path):
                    procesados += 1
                else:
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        resumen = f"Proceso de firma finalizado.\n\n- Documentos DOCX modificados: {procesados}\n- Errores/Omitidos: {errores}"
        self.log(f"--- {resumen.replace(chr(10)+chr(10), ' ').replace(chr(10), ' | ')} ---")
        messagebox.showinfo("Proceso Completado", resumen, parent=self.root)
    
    def _reemplazar_firma_en_docx_worker(self, docx_path, signature_path):
        try:
            doc = Document(docx_path)
            anchor_text = "Firma de Aceptacion"
            signature_p_index = -1

            for i, p in enumerate(doc.paragraphs):
                if anchor_text.lower() in p.text.lower():
                    target_index = i + 1
                    if target_index < len(doc.paragraphs):
                        signature_p_index = target_index
                    break
            
            if signature_p_index == -1:
                self.log(f"  ❌ ERROR: No se encontró el ancla '{anchor_text}' o el párrafo de la firma en '{os.path.basename(docx_path)}'.")
                return False

            signature_p = doc.paragraphs[signature_p_index]
            
            p_element = signature_p._p
            p_element.clear_content()

            run = signature_p.add_run()
            try:
                run.add_picture(signature_path, width=Inches(1.5))
                signature_p.alignment = WD_ALIGN_PARAGRAPH.LEFT 

            except FileNotFoundError:
                self.log(f"  ❌ ERROR: El archivo de imagen de firma no fue encontrado: {signature_path}")
                return False
            except Exception as img_e:
                self.log(f"  ❌ ERROR: No se pudo insertar la imagen '{os.path.basename(signature_path)}': {img_e}")
                return False

            doc.save(docx_path)
            self.log(f"  ✅ Firma reemplazada exitosamente en: '{os.path.basename(docx_path)}'")
            return True

        except Exception as e:
            self.log(f"❌ ERROR CRÍTICO al modificar '{os.path.basename(docx_path)}': {e}")
            return False
    
    def exportar_lista(self):
        if not self.source_path.get() or not os.path.isdir(self.source_path.get()):
            messagebox.showwarning("Advertencia", "Selecciona una carpeta de origen válida.", parent=self.root)
            return
        
        ruta_guardado = filedialog.asksaveasfilename(parent=self.root, defaultextension=".xlsx", filetypes=[("Archivo Excel", "*.xlsx")], title="Guardar análisis como")
        if not ruta_guardado: return
        
        self.log("Iniciando análisis de carpetas...")
        base_path = self.source_path.get()
        
        detalle_data = []
        resumen_data = []

        top_level_folders = [d for d in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, d))]

        for top_folder in top_level_folders:
            top_folder_path = os.path.join(base_path, top_folder)
            total_files_in_top_folder = 0
            has_any_files = False

            for root, dirs, files in os.walk(top_folder_path):
                total_files_in_top_folder += len(files)
                if files:
                    has_any_files = True
                    for file in files:
                        detalle_data.append({'nombre carpeta': top_folder, 'archivos que contiene': file})
            
            if not has_any_files:
                if not any(os.scandir(top_folder_path)):
                    detalle_data.append({'nombre carpeta': top_folder, 'archivos que contiene': 'carpeta vacia'})
            
            resumen_data.append({'Carpeta': top_folder, 'Cantidad de archivos': total_files_in_top_folder})

        if not detalle_data and not resumen_data:
            self.log("No se encontraron carpetas o archivos para analizar.")
            messagebox.showinfo("Análisis Vacío", "La carpeta seleccionada no contiene subcarpetas para analizar.")
            return

        df_detalle = pd.DataFrame(detalle_data, columns=['nombre carpeta', 'archivos que contiene'])
        df_resumen = pd.DataFrame(resumen_data)
        
        total_row = pd.DataFrame([{'Carpeta': 'TOTAL DE CARPETAS', 'Cantidad de archivos': len(top_level_folders)}])
        df_resumen = pd.concat([df_resumen, total_row], ignore_index=True)

        try:
            with pd.ExcelWriter(ruta_guardado, engine='xlsxwriter') as writer:
                df_detalle.to_excel(writer, sheet_name='Detalle', index=False)
                df_resumen.to_excel(writer, sheet_name='Resumen', index=False)
            
            self.log(f"✅ Análisis de carpetas exportado a: {ruta_guardado}")
            os.startfile(ruta_guardado)
            messagebox.showinfo("Éxito", f"La lista de archivos y carpetas se ha guardado en:\n{ruta_guardado}", parent=self.root)
        except Exception as e:
            self.log(f"❌ Error al exportar a Excel: {e}")
            messagebox.showerror("Error", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)
    
    def _extract_field(self, text, pattern, default="No encontrado"):
        """Busca un patrón regex en un texto y devuelve el primer grupo de captura limpio."""
        match = pattern.search(text)
        if match:
            # Limpiar el texto extraído: quitar espacios al inicio/final y reemplazar saltos de línea y espacios múltiples por uno solo.
            return " ".join(match.group(1).strip().split())
        return default

    def accion_analisis_historia_clinica(self):
        """
        Analiza masivamente los archivos PDF de los resultados de búsqueda para extraer
        datos de historias clínicas y exportarlos a un archivo Excel.
        """
        self.log("--- Iniciando Análisis de Historia Clínica desde PDFs ---")
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar el análisis.", parent=self.root)
            self.log("Análisis cancelado: no hay archivos en los resultados.")
            return

        archivos_pdf = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        if not archivos_pdf:
            messagebox.showinfo("Sin PDFs", "No se encontraron archivos PDF en los resultados de la búsqueda.", parent=self.root)
            self.log("Análisis cancelado: no se encontraron PDFs en los resultados.")
            return

        ruta_guardado = filedialog.asksaveasfilename(
            parent=self.root,
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx")],
            title="Guardar Análisis de Historias Clínicas como"
        )
        if not ruta_guardado:
            self.log("Análisis cancelado por el usuario (no se seleccionó archivo de guardado).")
            return

        # Expresiones regulares para cada campo solicitado
        patterns = {
            # Captura solo hasta el final de la línea para evitar incluir "Fecha".
            'Paciente': re.compile(r"Paciente:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            # Mantiene el patrón original que funciona bien.
            'Estrato': re.compile(r"Estrato:\s*(.*?)\s*Municipio:", re.IGNORECASE | re.DOTALL),
            # Captura solo hasta el final de la línea para evitar incluir "Ocupación".
            'Contrato': re.compile(r"Contrato:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            # Captura el texto entre "DATOS HISTORIA CLÍNICA" y "¿ES VICTIMA...", manejando saltos de línea.
            # También es más flexible con los acentos.
            'HISTORIA CLÍNICA': re.compile(r"DATOS HISTORIA CL[IÍ]NICA\s*(.*?)\s*¿ES V[IÍ]CTIMA DE VIOLENCIA\?", re.IGNORECASE | re.DOTALL)
        }
        
        extracted_data = []
        procesados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Analizando Historias Clínicas...", len(archivos_pdf))
        
        try:
            for i, pdf_path in enumerate(archivos_pdf):
                nombre_archivo = os.path.basename(pdf_path)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    full_text = ""
                    with fitz.open(pdf_path) as doc:
                        for page in doc:
                            full_text += page.get_text("text") + "\n"
                    
                    record = {'Archivo': nombre_archivo}
                    for key, pattern in patterns.items():
                        record[key] = self._extract_field(full_text, pattern)

                    extracted_data.append(record)
                    self.log(f"✅ Analizado: {nombre_archivo}")
                    procesados += 1

                except Exception as e:
                    self.log(f"❌ Error al procesar el archivo '{nombre_archivo}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        if not extracted_data:
            self.log("No se pudo extraer información de ningún archivo. No se generó el Excel.")
            messagebox.showinfo("Sin Datos", "No se pudo extraer datos de los archivos PDF seleccionados.", parent=self.root)
            return

        try:
            # Orden de columnas actualizado para coincidir con las claves del diccionario.
            column_order = ['Archivo', 'Paciente', 'Estrato', 'Contrato', 'HISTORIA CLÍNICA']
            df = pd.DataFrame(extracted_data, columns=column_order)
            df.to_excel(ruta_guardado, index=False)
            self.log(f"--- Análisis completado. Procesados: {procesados}, Errores: {errores} ---")
            self.log(f"📊 Datos guardados exitosamente en: {ruta_guardado}")
            messagebox.showinfo("Proceso Completado", f"El análisis ha finalizado.\n\n- Archivos procesados: {procesados}\n- Errores: {errores}\n\nEl archivo Excel ha sido guardado.", parent=self.root)
            os.startfile(ruta_guardado)
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo Excel: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)
        # --- FIN DE LA MODIFICACIÓN ---

    ### NUEVO ###
    def accion_analisis_autorizacion_nueva_eps(self):
        """
        Analiza masivamente archivos PDF de autorizaciones de Nueva EPS desde los resultados 
        de búsqueda, extrae datos clave y los exporta a un archivo Excel.
        """
        self.log("--- Iniciando Análisis de Autorización Nueva EPS desde PDFs ---")
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar el análisis.", parent=self.root)
            self.log("Análisis cancelado: no hay archivos en los resultados.")
            return

        archivos_pdf = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        if not archivos_pdf:
            messagebox.showinfo("Sin PDFs", "No se encontraron archivos PDF en los resultados de la búsqueda.", parent=self.root)
            self.log("Análisis cancelado: no se encontraron PDFs en los resultados.")
            return

        ruta_guardado = filedialog.asksaveasfilename(
            parent=self.root,
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx")],
            title="Guardar Análisis de Autorizaciones Nueva EPS como"
        )
        if not ruta_guardado:
            self.log("Análisis cancelado por el usuario (no se seleccionó archivo de guardado).")
            return

        # Expresiones regulares para cada campo solicitado
        patterns = {
            'Afiliado': re.compile(r"Afiliado:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            'N° Autorización': re.compile(r"N° Autorización:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            'Autorizada el': re.compile(r"Autorizada el:\s*(.*?)(?:\n|$)", re.IGNORECASE),
            'Descripción Servicio': re.compile(r"Descripción Servicio\s*\n\s*\d+\s+\d+\s+(.*?)(?:\n|$)", re.IGNORECASE | re.DOTALL),
            # --- INICIO DE LA MODIFICACIÓN ---
            # Se generaliza el patrón para que capture cualquier línea que comience con "Afiliado Cancela" o "Afiliado No Cancela".
            'Info de Pago': re.compile(r"(Afiliado (?:No )?Cancela.*?)(?:\n|$)", re.IGNORECASE)
            # --- FIN DE LA MODIFICACIÓN ---
        }
        
        extracted_data = []
        procesados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Analizando Autorizaciones...", len(archivos_pdf))
        
        try:
            for i, pdf_path in enumerate(archivos_pdf):
                nombre_archivo = os.path.basename(pdf_path)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    full_text = ""
                    with fitz.open(pdf_path) as doc:
                        for page in doc:
                            full_text += page.get_text("text") + "\n"
                    
                    record = {'Archivo': nombre_archivo}
                    for key, pattern in patterns.items():
                        record[key] = self._extract_field(full_text, pattern)

                    extracted_data.append(record)
                    self.log(f"✅ Analizado: {nombre_archivo}")
                    procesados += 1

                except Exception as e:
                    self.log(f"❌ Error al procesar el archivo '{nombre_archivo}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        if not extracted_data:
            self.log("No se pudo extraer información de ningún archivo. No se generó el Excel.")
            messagebox.showinfo("Sin Datos", "No se pudo extraer datos de los archivos PDF seleccionados.", parent=self.root)
            return

        try:
            column_order = ['Archivo', 'Afiliado', 'N° Autorización', 'Autorizada el', 'Descripción Servicio', 'Info de Pago']
            df = pd.DataFrame(extracted_data, columns=column_order)
            
            df.to_excel(ruta_guardado, index=False)
            
            self.log(f"--- Análisis de autorizaciones completado. Procesados: {procesados}, Errores: {errores} ---")
            self.log(f"📊 Datos guardados exitosamente en: {ruta_guardado}")
            messagebox.showinfo("Proceso Completado", f"El análisis ha finalizado.\n\n- Archivos procesados: {procesados}\n- Errores: {errores}\n\nEl archivo Excel ha sido guardado.", parent=self.root)
            os.startfile(ruta_guardado)
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo Excel: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)
    ### FIN NUEVO ###

    def accion_analisis_cargue_sanitas(self):
        """
        Analiza masivamente archivos PDF del cargue de Sanitas desde los resultados 
        de búsqueda, extrae datos clave (FEOV, Estado, Fecha) y los exporta a un archivo Excel.
        """
        self.log("--- Iniciando Análisis de Cargue Sanitas desde PDFs ---")
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Primero debes buscar archivos para poder ejecutar el análisis.", parent=self.root)
            self.log("Análisis cancelado: no hay archivos en los resultados.")
            return

        archivos_pdf = [ruta for ruta in self.resultados if ruta.lower().endswith('.pdf')]
        if not archivos_pdf:
            messagebox.showinfo("Sin PDFs", "No se encontraron archivos PDF en los resultados de la búsqueda.", parent=self.root)
            self.log("Análisis cancelado: no se encontraron PDFs en los resultados.")
            return

        ruta_guardado = filedialog.asksaveasfilename(
            parent=self.root,
            defaultextension=".xlsx",
            filetypes=[("Archivo Excel", "*.xlsx")],
            title="Guardar Análisis de Cargue Sanitas como"
        )
        if not ruta_guardado:
            self.log("Análisis cancelado por el usuario (no se seleccionó archivo de guardado).")
            return

        # Expresiones regulares para los campos resaltados en amarillo en la imagen:
        # 1. Factura (FEOVXXX): Busca la sigla FEOV seguida de números.
        # 2. Estado: Busca la palabra que sigue a "Estado" y está antes del final de la línea o de la fecha.
        # 3. Fecha: Busca la fecha y hora combinada.
        patterns = {
            'Factura (FEOV)': re.compile(r"FEOV(\d+)", re.IGNORECASE),
        #'Estado': re.compile(r"Estado\s*(.*?)(?:\s*-\s*\d{1,2}\s+\w+\s+\d{4})", re.IGNORECASE), # Captura lo que sigue a 'Estado' hasta la fecha o fin de línea
            'Fecha y hora de cargue': re.compile(r"(\d{1,2}\s+\w+\s+\d{4}\s*-\s*\d{1,2}:\d{2})", re.IGNORECASE) # Captura la fecha y hora
        }
        
        extracted_data = []
        procesados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Analizando Cargue Sanitas...", len(archivos_pdf))
        
        try:
            for i, pdf_path in enumerate(archivos_pdf):
                nombre_archivo = os.path.basename(pdf_path)
                progress_label.config(text=f"Procesando: {nombre_archivo}")
                progress_bar['value'] = i + 1
                self.root.update_idletasks()

                try:
                    full_text = ""
                    with fitz.open(pdf_path) as doc:
                        for page in doc:
                            full_text += page.get_text("text") + "\n"
                    
                    record = {'Archivo': nombre_archivo}
                    for key, pattern in patterns.items():
                        record[key] = self._extract_field(full_text, pattern)

                    # Extraer el valor numérico de la factura si el patrón 'FEOV' lo encuentra
                    factura_match = patterns['Factura (FEOV)'].search(full_text)
                    if factura_match:
                        record['Factura (FEOV)'] = factura_match.group(1) # Guardar solo los números

                    extracted_data.append(record)
                    self.log(f"✅ Analizado: {nombre_archivo}")
                    procesados += 1

                except Exception as e:
                    self.log(f"❌ Error al procesar el archivo '{nombre_archivo}': {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        if not extracted_data:
            self.log("No se pudo extraer información de ningún archivo. No se generó el Excel.")
            messagebox.showinfo("Sin Datos", "No se pudo extraer datos de los archivos PDF seleccionados.", parent=self.root)
            return

        try:
            # Definir el orden de las columnas para el DataFrame
            column_order = ['Archivo', 'Factura (FEOV)', 'Fecha y hora de cargue']
            df = pd.DataFrame(extracted_data, columns=column_order)
            
            # Limpiar el valor de 'Estado' para quitar posibles espacios extra o saltos de línea
            if 'Estado' in df.columns:
                df['Estado'] = df['Estado'].apply(lambda x: " ".join(str(x).strip().split()) if pd.notna(x) else "No encontrado")
            
            df.to_excel(ruta_guardado, index=False)
            
            self.log(f"--- Análisis de Cargue Sanitas completado. Procesados: {procesados}, Errores: {errores} ---")
            self.log(f"📊 Datos guardados exitosamente en: {ruta_guardado}")
            messagebox.showinfo("Proceso Completado", f"El análisis ha finalizado.\n\n- Archivos procesados: {procesados}\n- Errores: {errores}\n\nEl archivo Excel ha sido guardado.", parent=self.root)
            os.startfile(ruta_guardado)
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo Excel: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo Excel:\n{e}", parent=self.root)


    def exportar_para_renombrar(self):
        if not self.resultados:
            messagebox.showwarning("Sin resultados", "Busca archivos primero para poder exportarlos.", parent=self.root)
            return
        
        ruta_guardado = filedialog.asksaveasfilename(parent=self.root, defaultextension=".xlsx", filetypes=[("Archivo Excel", "*.xlsx")], title="Guardar para renombrar")
        if not ruta_guardado: return
        
        data = []
        for ruta in self.resultados:
            data.append({"Ruta actual": ruta, "Nuevo nombre": os.path.basename(ruta)})
        
        df = pd.DataFrame(data)
        try:
            df.to_excel(ruta_guardado, index=False)
            self.log(f"✅ Lista para renombrar exportada a: {ruta_guardado}")
            messagebox.showinfo("Éxito", f"El archivo para renombrar se ha guardado en:\n{ruta_guardado}\n\nEdita la columna 'Nuevo nombre' y luego usa la opción 'Aplicar renombrado desde Excel'.", parent=self.root)
        except Exception as e:
            self.log(f"❌ Error al exportar: {e}")

    def accion_anadir_sufijo_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(
            parent=self.root,
            title="Seleccionar archivo Excel con los sufijos",
            filetypes=[("Archivos Excel", "*.xlsx *.xls")]
        )
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring(
                "Seleccionar Hoja",
                f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}",
                parent=self.root,
                initialvalue=nombres_hojas[0] if nombres_hojas else ""
            )
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas:
                self.log("Operación cancelada o nombre de hoja no válido.")
                return

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_carpeta = simpledialog.askstring(
                "Columna de Carpetas",
                f"Ingrese la LETRA de la columna que contiene los NOMBRES DE CARPETA.\n(Opciones: {letras_rango})",
                parent=self.root
            )
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles:
                self.log(f"Letra de columna para carpetas no válida: '{letra_col_carpeta}'.")
                return
            
            letra_col_sufijo = simpledialog.askstring(
                "Columna de Sufijos",
                f"Ingrese la LETRA de la columna que contiene el SUFIJO (ej: _FEOV12345).\n(Opciones: {letras_rango})",
                parent=self.root
            )
            if not letra_col_sufijo or letra_col_sufijo.upper() not in letras_disponibles:
                self.log(f"Letra de columna para sufijos no válida: '{letra_col_sufijo}'.")
                return

            indice_col_carpeta = ord(letra_col_carpeta.upper()) - ord('A')
            indice_col_sufijo = ord(letra_col_sufijo.upper()) - ord('A')
            
            self.log(f"--- Iniciando proceso de añadir sufijos desde Excel ---")
            carpetas_procesadas, archivos_renombrados, errores = 0, 0, 0
            
            progress_win, progress_bar, progress_label = self._create_progress_window("Añadiendo sufijos...", len(df))
            try:
                for index, fila in df.iterrows():
                    nombre_carpeta = fila.iloc[indice_col_carpeta]
                    sufijo = fila.iloc[indice_col_sufijo]

                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()

                    if pd.isna(nombre_carpeta) or pd.isna(sufijo): continue

                    nombre_carpeta = str(nombre_carpeta).strip()
                    sufijo = str(sufijo).strip()
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")

                    if not nombre_carpeta or not sufijo: continue 

                    ruta_carpeta = os.path.join(carpeta_origen, nombre_carpeta)

                    if os.path.isdir(ruta_carpeta):
                        self.log(f"Procesando carpeta: '{nombre_carpeta}'")
                        carpetas_procesadas += 1
                        for nombre_archivo in os.listdir(ruta_carpeta):
                            ruta_archivo_actual = os.path.join(ruta_carpeta, nombre_archivo)
                            if os.path.isfile(ruta_archivo_actual):
                                nombre_base, extension = os.path.splitext(nombre_archivo)
                                nuevo_nombre_archivo = f"{nombre_base}{sufijo}{extension}"
                                nueva_ruta_archivo = os.path.join(ruta_carpeta, nuevo_nombre_archivo)

                                if ruta_archivo_actual == nueva_ruta_archivo: continue
                                
                                if os.path.exists(nueva_ruta_archivo):
                                    self.log(f"⚠️  ERROR: El archivo '{nuevo_nombre_archivo}' ya existe. No se renombró '{nombre_archivo}'.")
                                    errores += 1
                                    continue
                                
                                try:
                                    os.rename(ruta_archivo_actual, nueva_ruta_archivo)
                                    self.log(f"  ✅ Renombrado: '{nombre_archivo}' → '{nuevo_nombre_archivo}'")
                                    archivos_renombrados += 1
                                except Exception as e:
                                    self.log(f"  ❌ ERROR al renombrar '{nombre_archivo}': {e}")
                                    errores += 1
                    else:
                        self.log(f"⚠️  ADVERTENCIA: No se encontró la carpeta '{nombre_carpeta}' en la ruta de origen.")
                        errores += 1
            finally:
                if progress_win:
                    progress_win.destroy()
            
            self.log(f"--- Proceso Finalizado. Carpetas procesadas: {carpetas_procesadas}, Archivos renombrados: {archivos_renombrados}, Errores/Advertencias: {errores} ---")
            messagebox.showinfo("Proceso Completado", f"Proceso de renombrado con sufijo finalizado.\n\n- Archivos renombrados: {archivos_renombrados}\n- Carpetas procesadas: {carpetas_procesadas}\n- Errores: {errores}\n\nRevise el registro de actividad para más detalles.", parent=self.root)

        except Exception as e:
            self.log(f"❌ Ocurrió un error general: {e}")
            messagebox.showerror("Error", f"Ocurrió un error inesperado al procesar el archivo Excel:\n{e}", parent=self.root)

    def accion_autorizacion_docx_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona la carpeta de origen que contiene todas las subcarpetas.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar archivo Excel con los datos", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring("Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}", parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else "")
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_carpeta = simpledialog.askstring("Columna de Carpetas", f"Ingrese la LETRA de la columna que contiene los NOMBRES DE CARPETA.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return
            
            letra_col_auth = simpledialog.askstring("Columna de Autorización", f"Ingrese la LETRA de la columna que contiene el nuevo NÚMERO DE AUTORIZACIÓN.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_auth or letra_col_auth.upper() not in letras_disponibles: return

            indice_col_carpeta = ord(letra_col_carpeta.upper()) - ord('A')
            indice_col_auth = ord(letra_col_auth.upper()) - ord('A')
            
            self.log(f"--- Iniciando modificación de Autorización en DOCX (por carpeta) ---")
            
            modificados, errores_carpeta, errores_docx, errores_proceso = 0, 0, 0, 0
            docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)

            progress_win, progress_bar, progress_label = self._create_progress_window("Modificando DOCX...", len(df))
            try:
                for index, fila in df.iterrows():
                    nombre_carpeta = fila.iloc[indice_col_carpeta]
                    nueva_autorizacion = fila.iloc[indice_col_auth]
                    
                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()
                    
                    if pd.isna(nombre_carpeta) or pd.isna(nueva_autorizacion): continue

                    nombre_carpeta = str(nombre_carpeta).strip()
                    nueva_autorizacion = str(int(nueva_autorizacion)) if isinstance(nueva_autorizacion, (float, int)) else str(nueva_autorizacion).strip()
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")

                    if not nombre_carpeta or not nueva_autorizacion: continue
                    
                    ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
                    if not os.path.isdir(ruta_carpeta_especifica):
                        self.log(f"⚠️ ADVERTENCIA: No se encontró la carpeta '{nombre_carpeta}' en la ruta de origen.")
                        errores_carpeta += 1
                        continue
                    
                    ruta_docx_encontrada = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if docx_pattern.match(f)), None)

                    if not ruta_docx_encontrada:
                        self.log(f"❌ ERROR: No se encontró un archivo DOCX con el patrón 'CRC_..._FEOV.docx' dentro de la carpeta '{nombre_carpeta}'.")
                        errores_docx += 1
                        continue

                    try:
                        doc = Document(ruta_docx_encontrada)
                        fue_modificado = False
                        for p in doc.paragraphs:
                            if "AUTORIZACION:" in p.text.upper():
                                p.text = re.sub(r'(AUTORIZACION:)\s*.*', r'\1 ' + str(nueva_autorizacion), p.text, flags=re.IGNORECASE)
                                fue_modificado = True
                                break
                        
                        if fue_modificado:
                            doc.save(ruta_docx_encontrada)
                            self.log(f"✅ Modificado '{os.path.basename(ruta_docx_encontrada)}' en carpeta '{nombre_carpeta}' con autorización: '{nueva_autorizacion}'")
                            modificados += 1
                        else:
                            self.log(f"⚠️ ADVERTENCIA: Se encontró el archivo '{os.path.basename(ruta_docx_encontrada)}' pero no la línea 'AUTORIZACION:'.")
                            errores_proceso += 1
                    except Exception as e:
                        self.log(f"❌ ERROR CRÍTICO al procesar el archivo '{os.path.basename(ruta_docx_encontrada)}': {e}")
                        errores_proceso += 1
            finally:
                if progress_win:
                    progress_win.destroy()

            resumen_msg = (f"Proceso de modificación de autorización finalizado.\n\n" f"- Modificados: {modificados}\n" f"- Carpetas no encontradas: {errores_carpeta}\n" f"- DOCX no encontrados: {errores_docx}\n" f"- Errores de procesamiento: {errores_proceso}")
            self.log(f"--- Proceso Finalizado. Modificados: {modificados}, Errores: {errores_carpeta}/{errores_docx}/{errores_proceso} ---")
            messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
        except Exception as e:
            self.log(f"❌ Ocurrió un error general: {e}")
            messagebox.showerror("Error", f"Ocurrió un error inesperado al procesar el archivo Excel:\n{e}", parent=self.root)

    def accion_regimen_docx_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona la carpeta de origen que contiene todas las subcarpetas.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar archivo Excel con los datos", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring("Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}", parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else "")
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return

            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
            letras_disponibles = list(string.ascii_uppercase)[:len(df.columns)]
            letras_rango = f"A - {letras_disponibles[-1] if letras_disponibles else 'A'}"

            letra_col_carpeta = simpledialog.askstring("Columna de Carpetas", f"Ingrese la LETRA de la columna que contiene los NOMBRES DE CARPETA.\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_carpeta or letra_col_carpeta.upper() not in letras_disponibles: return
            
            letra_col_regimen = simpledialog.askstring("Columna de Régimen", f"Ingrese la LETRA de la columna que contiene el nuevo RÉGIMEN (ej: CONTRIBUTIVO).\n(Opciones: {letras_rango})", parent=self.root)
            if not letra_col_regimen or letra_col_regimen.upper() not in letras_disponibles: return

            indice_col_carpeta = ord(letra_col_carpeta.upper()) - ord('A')
            indice_col_regimen = ord(letra_col_regimen.upper()) - ord('A')
            
            self.log(f"--- Iniciando modificación de Régimen en DOCX (por carpeta) ---")
            
            modificados, errores_carpeta, errores_docx, errores_proceso = 0, 0, 0, 0
            docx_pattern = re.compile(r'CRC_.*_FEOV.*\.docx$', re.IGNORECASE)

            progress_win, progress_bar, progress_label = self._create_progress_window("Modificando DOCX...", len(df))
            try:
                for index, fila in df.iterrows():
                    nombre_carpeta = fila.iloc[indice_col_carpeta]
                    nuevo_regimen = fila.iloc[indice_col_regimen]
                    
                    progress_bar['value'] = index + 1
                    self.root.update_idletasks()
                    
                    if pd.isna(nombre_carpeta) or pd.isna(nuevo_regimen): continue

                    nombre_carpeta = str(nombre_carpeta).strip()
                    nuevo_regimen = str(nuevo_regimen).strip()
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")

                    if not nombre_carpeta or not nuevo_regimen: continue
                    
                    ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
                    if not os.path.isdir(ruta_carpeta_especifica):
                        self.log(f"⚠️ ADVERTENCIA: No se encontró la carpeta '{nombre_carpeta}' en la ruta de origen.")
                        errores_carpeta += 1
                        continue
                    
                    ruta_docx_encontrada = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if docx_pattern.match(f)), None)

                    if not ruta_docx_encontrada:
                        self.log(f"❌ ERROR: No se encontró un archivo DOCX con el patrón 'CRC_..._FEOV.docx' dentro de la carpeta '{nombre_carpeta}'.")
                        errores_docx += 1
                        continue

                    try:
                        doc = Document(ruta_docx_encontrada)
                        fue_modificado = False
                        for p in doc.paragraphs:
                            if "REGIMEN:" in p.text.upper():
                                p.text = re.sub(r'(REGIMEN:)\s*.*', r'\1 ' + str(nuevo_regimen), p.text, flags=re.IGNORECASE)
                                fue_modificado = True
                                break
                        
                        if fue_modificado:
                            doc.save(ruta_docx_encontrada)
                            self.log(f"✅ Modificado '{os.path.basename(ruta_docx_encontrada)}' en carpeta '{nombre_carpeta}' con régimen: '{nuevo_regimen}'")
                            modificados += 1
                        else:
                            self.log(f"⚠️ ADVERTENCIA: Se encontró el archivo '{os.path.basename(ruta_docx_encontrada)}' pero no la línea 'REGIMEN:'.")
                            errores_proceso += 1

                    except Exception as e:
                        self.log(f"❌ ERROR CRÍTICO al procesar el archivo '{os.path.basename(ruta_docx_encontrada)}': {e}")
                        errores_proceso += 1
            finally:
                if progress_win:
                    progress_win.destroy()

            resumen_msg = (f"Proceso de modificación de régimen finalizado.\n\n" f"- Modificados: {modificados}\n" f"- Carpetas no encontradas: {errores_carpeta}\n" f"- DOCX no encontrados: {errores_docx}\n" f"- Errores de procesamiento: {errores_proceso}")
            self.log(f"--- Proceso Finalizado. Modificados: {modificados}, Errores: {errores_carpeta}/{errores_docx}/{errores_proceso} ---")
            messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
        except Exception as e:
            self.log(f"❌ Ocurrió un error general: {e}")
            messagebox.showerror("Error", f"Ocurrió un error inesperado al procesar el archivo Excel:\n{e}", parent=self.root)
    
    def accion_modificar_docx_completo_desde_excel(self):
        carpeta_origen = self.source_path.get()
        if not carpeta_origen or not os.path.isdir(carpeta_origen):
            messagebox.showwarning("Advertencia", "Por favor, selecciona una carpeta de origen válida que contenga las subcarpetas de los pacientes.", parent=self.root)
            return

        archivo_excel = filedialog.askopenfilename(parent=self.root, title="Seleccionar archivo Excel con los datos", filetypes=[("Archivos Excel", "*.xlsx *.xls")])
        if not archivo_excel: return

        try:
            xls = pd.ExcelFile(archivo_excel)
            nombres_hojas = xls.sheet_names
            hoja_seleccionada = simpledialog.askstring("Seleccionar Hoja", f"Ingrese el nombre de la hoja de cálculo a usar.\n\nDisponibles: {', '.join(nombres_hojas)}", parent=self.root, initialvalue=nombres_hojas[0] if nombres_hojas else "")
            if not hoja_seleccionada or hoja_seleccionada not in nombres_hojas: return
            df = pd.read_excel(xls, sheet_name=hoja_seleccionada)
        except Exception as e:
            self.log(f"❌ Error al leer el archivo Excel: {e}")
            messagebox.showerror("Error de Lectura", f"No se pudo leer el archivo o la hoja de Excel:\n{e}", parent=self.root)
            return
        
        column_map, missing_cols = self._create_column_map_from_headers(df)
        
        if not column_map:
            error_msg = ("El archivo Excel no tiene el formato correcto.\n\n"
                         "Faltan las siguientes columnas obligatorias:\n- " + 
                         "\n- ".join(missing_cols))
            self.log(f"❌ Error de formato en Excel: {error_msg.replace(chr(10), ' ')}")
            messagebox.showerror("Error en Encabezados", error_msg, parent=self.root)
            return

        self.log(f"--- Iniciando modificación completa de DOCX desde Excel (Modo Automático) ---")
        self.log(f"Columnas encontradas y mapeadas correctamente desde la hoja '{hoja_seleccionada}'.")
        modificados, errores = 0, 0
        
        progress_win, progress_bar, progress_label = self._create_progress_window("Modificando DOCX...", len(df))
        try:
            for index, row in df.iterrows():
                progress_bar['value'] = index + 1
                datos = {}
                try:
                    datos = {key: str(row[col_name]).strip() if pd.notna(row[col_name]) else "" for key, col_name in column_map.items()}
                    nombre_carpeta = datos.get('folder')

                    if not nombre_carpeta:
                        self.log(f"⚠️ Fila {index+2} omitida: El nombre de la carpeta está vacío.")
                        continue
                    
                    progress_label.config(text=f"Procesando: {nombre_carpeta}")
                    self.root.update_idletasks()

                    ruta_carpeta_especifica = os.path.join(carpeta_origen, nombre_carpeta)
                    if not os.path.isdir(ruta_carpeta_especifica):
                        self.log(f"⚠️ No se encontró la carpeta '{nombre_carpeta}' (Fila {index+2}).")
                        errores += 1
                        continue
                    
                    ruta_docx = next((os.path.join(ruta_carpeta_especifica, f) for f in os.listdir(ruta_carpeta_especifica) if f.lower().endswith('.docx') and 'plantilla' in f.lower()), None)
                    if not ruta_docx:
                        self.log(f"❌ No se encontró archivo DOCX con 'plantilla' en el nombre dentro de la carpeta '{nombre_carpeta}'.")
                        errores += 1
                        continue
                    
                    doc = Document(ruta_docx)
                    
                    for p in doc.paragraphs:
                        if "Santiago de Cali, " in p.text:
                            p.text =(f"Santiago de Cali,  {datos['date']}")
                            continue

                        if "Yo " in p.text and "identificado con" in p.text:
                            p.text = (f"Yo {datos['full_name']} identificado con {datos['doc_type']}, "
                                      f"Numero {datos['doc_num']} en calidad de paciente, doy fé y acepto el "
                                      f"servicio de {datos['service']} brindado por la IPS OPORTUNIDAD DE VIDA S.A.S"
                                      f"   "
                                      f"   ")
                            continue

                        if "EPS:" in p.text:
                            p.text = re.sub(r'(EPS:)\s*.*', r'\1 ' + datos['eps'], p.text, count=1)
                        if "TIPO SERVICIO:" in p.text:
                            p.text = re.sub(r'(TIPO SERVICIO:)\s*.*', r'\1 ' + datos['tipo_servicio'], p.text, count=1)
                        if "REGIMEN:" in p.text:
                            p.text = re.sub(r'(REGIMEN:)\s*.*', r'\1 ' + datos['regimen'], p.text, count=1)
                        if "CATEGORIA:" in p.text:
                            p.text = re.sub(r'(CATEGORIA:)\s*.*', r'\1 ' + datos['categoria'], p.text, count=1)
                        if "VALOR CUOTA MODERADORA:" in p.text:
                            p.text = re.sub(r'(VALOR CUOTA MODERADORA:)\s*.*', r'\1 ' + datos['cuota'], p.text, count=1)
                        if "AUTORIZACION:" in p.text:
                            p.text = re.sub(r'(AUTORIZACION:)\s*.*', r'\1 ' + datos['auth'], p.text, count=1)
                        if "Fecha de Atención:" in p.text:
                            p.text = re.sub(r'(Fecha de Atención:)\s*.*', r'\1 ' + datos['fecha_atencion'], p.text, count=1)
                        if "Fecha de Finalización:" in p.text:
                            p.text = re.sub(r'(Fecha de Finalización:)\s*.*', r'\1 ' + datos['fecha_fin'], p.text, count=1)
                            continue

                    signature_line_index = -1
                    for i, p in enumerate(doc.paragraphs):
                        if "FIRMA DE ACEPTACION" in p.text.upper():
                            signature_line_index = i
                            break 

                    if signature_line_index != -1 and signature_line_index + 2 < len(doc.paragraphs):
                        name_paragraph = doc.paragraphs[signature_line_index + 2]
                        name_paragraph.text = datos['full_name'].upper()
                    else:
                        self.log(f"⚠️ No se pudo encontrar la posición del nombre para la firma en '{os.path.basename(ruta_docx)}'.")

                    doc.save(ruta_docx)
                    self.log(f"✅ Modificado: {os.path.basename(ruta_docx)} en carpeta '{nombre_carpeta}'")
                    modificados += 1

                except Exception as e:
                    self.log(f"❌ ERROR CRÍTICO al procesar la fila {index+2} (Carpeta: {datos.get('folder', 'N/A')}): {e}")
                    errores += 1
        finally:
            if progress_win:
                progress_win.destroy()

        resumen_msg = (f"Proceso de modificación de DOCX finalizado.\n\n" f"- Archivos modificados: {modificados}\n" f"- Errores/Advertencias: {errores}")
        self.log(f"--- Proceso Finalizado. Modificados: {modificados}, Errores: {errores} ---")
        messagebox.showinfo("Proceso Completado", resumen_msg, parent=self.root)
    
    def _create_column_map_from_headers(self, df):
        """
        Valida que el DataFrame contenga los encabezados necesarios y devuelve el mapa de columnas.
        Si faltan columnas, devuelve None y una lista de las columnas faltantes.
        """
        required_map = {
            'folder': 'Nombre Carpeta',
            'date': 'Ciudad y Fecha',
            'full_name': 'Nombre Completo',
            'doc_type': 'Tipo Documento',
            'doc_num': 'Numero Documento',
            'service': 'Servicio',
            'eps': 'EPS',
            'tipo_servicio': 'Tipo Servicio',
            'regimen': 'Regimen',
            'categoria': 'Categoria',
            'cuota': 'Valor Cuota Moderadora',
            'auth': 'Numero Autorizacion',
            'fecha_atencion': 'Fecha y Hora Atencion',
            'fecha_fin': 'Fecha Finalizacion'
        }

        excel_headers = df.columns
        missing_cols = [header for header in required_map.values() if header not in excel_headers]

        if missing_cols:
            return None, missing_cols

        return required_map, []

    #================================================================================#
    #========== INICIO DE LA SECCIÓN MODIFICADA: VISOR Y EDITOR JSON / XML ==========
    #================================================================================#
    
    def _setup_data_editor_tab(self, tab):
        """Configura la interfaz de usuario para la pestaña del editor de datos."""
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(2, weight=1)

        # --- Frame para controles de archivo ---
        top_frame = ttk.Frame(tab, padding=(10, 10))
        top_frame.grid(row=0, column=0, sticky="ew")
        
        ttk.Button(top_frame, text="📂 Abrir Archivo (JSON/XML)", command=self._open_data_file).pack(side="left", padx=(0, 10))
        self.save_button = ttk.Button(top_frame, text="💾 Guardar Cambios", command=self._save_data_file, state="disabled")
        self.save_button.pack(side="left")
        
        # --- Frame para búsqueda ---
        search_frame = ttk.Frame(tab, padding=(10, 0, 10, 10))
        search_frame.grid(row=1, column=0, sticky="ew")
        ttk.Label(search_frame, text="Buscar:").pack(side="left")
        search_entry = ttk.Entry(search_frame, textvariable=self.search_data_var)
        search_entry.pack(side="left", fill="x", expand=True, padx=5)
        search_entry.bind("<Return>", self._search_in_tree)
        ttk.Button(search_frame, text="Buscar y Resaltar", command=self._search_in_tree).pack(side="left")
        ttk.Button(search_frame, text="Limpiar", command=self._clear_search_highlight).pack(side="left", padx=5)

        # --- Frame para el Treeview ---
        tree_frame = ttk.LabelFrame(tab, text="Estructura del Archivo", padding=10)
        tree_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=5)
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.data_tree = ttk.Treeview(tree_frame, columns=("value",), show="tree headings")
        self.data_tree.heading("#0", text="Clave / Tag")
        self.data_tree.heading("value", text="Valor")
        self.data_tree.column("#0", width=350, stretch=tk.NO)
        self.data_tree.column("value", width=500, stretch=tk.YES)
        
        ysb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.data_tree.yview)
        xsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.data_tree.xview)
        self.data_tree.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)
        
        self.data_tree.grid(row=0, column=0, sticky="nsew")
        ysb.grid(row=0, column=1, sticky="ns")
        xsb.grid(row=1, column=0, sticky="ew")

        self._configure_tree_tags_and_menu()
        self.data_tree.bind("<Double-1>", self._on_tree_double_click)
        self.data_tree.bind("<Button-3>", self._show_context_menu)
        self.data_tree.bind("<Motion>", self._on_tree_motion)
        self.data_tree.bind("<Leave>", self._on_tree_leave)

    def _configure_tree_tags_and_menu(self):
        """Configura los estilos de color para el Treeview y el menú contextual."""
        self.data_tree.tag_configure('string', foreground='green')
        self.data_tree.tag_configure('number', foreground='blue')
        self.data_tree.tag_configure('boolean', foreground='purple')
        self.data_tree.tag_configure('null', foreground='gray')
        self.data_tree.tag_configure('dict', foreground='#a52a2a') # Brown
        self.data_tree.tag_configure('list', foreground='#a52a2a')
        self.data_tree.tag_configure('xml_tag', foreground='#00008B') # DarkBlue
        self.data_tree.tag_configure('xml_attr', foreground='#8B008B') # DarkMagenta
        
        self.data_tree.tag_configure('search_hit', background='yellow')

        self.data_context_menu = tk.Menu(self.root, tearoff=0)

    def _open_data_file(self):
        path = filedialog.askopenfilename(
            parent=self.root,
            title="Seleccionar archivo JSON o XML",
            filetypes=[("Archivos de Datos", "*.json *.xml"), ("Todos los archivos", "*.*")]
        )
        if not path:
            return

        self._reset_editor_state()
        self.root.title(f"📂 Organizador de Archivos - Editando: {os.path.basename(path)}")
        
        try:
            if path.lower().endswith(".json"):
                self.current_data_type = 'json'
                with open(path, 'r', encoding='utf-8') as f:
                    self.parsed_data = json.load(f)
                self._populate_tree_from_json(self.parsed_data, '')
                self.log(f"✅ Archivo JSON '{os.path.basename(path)}' cargado.")
            
            elif path.lower().endswith(".xml"):
                self.current_data_type = 'xml'
                tree = ET.parse(path)
                self.parsed_data = tree.getroot()
                self._populate_tree_from_xml(self.parsed_data, '')
                self.log(f"✅ Archivo XML '{os.path.basename(path)}' cargado.")
            else:
                raise ValueError("Tipo de archivo no soportado.")
            self.current_file_path = path
        
        except Exception as e:
            self.log(f"❌ Error abriendo o procesando el archivo: {e}")
            messagebox.showerror("Error de Archivo", f"No se pudo cargar o procesar el archivo.\n\n{e}", parent=self.root)
            self._reset_editor_state()

    def _reset_editor_state(self):
        """Restaura el estado del editor de datos a su valor inicial."""
        if self.data_tree:
            self.data_tree.delete(*self.data_tree.get_children())
        self.current_file_path = None
        self.parsed_data = None
        self.current_data_type = None
        self.is_data_modified = False
        self.embedded_xml_docs.clear()
        self.xml_element_map.clear()
        self.search_data_var.set("")
        self._clear_search_highlight()
        if self.save_button:
            self.save_button.config(state="disabled")
        self.root.title("📂 Organizador de Archivos v5.7 (Mapeo Parcial)")
    
    def _clean_tag(self, tag):
        """Elimina el namespace {....} de una etiqueta XML para hacerla más legible."""
        return re.sub(r'\{.*?\}', '', tag)
        
    def _get_value_tag(self, value):
        """Devuelve el nombre del tag de color según el tipo de dato."""
        if isinstance(value, str): return 'string'
        if isinstance(value, (int, float)): return 'number'
        if isinstance(value, bool): return 'boolean'
        if value is None: return 'null'
        return ''

    def _populate_tree_from_xml(self, element, parent_iid):
        """Recursivamente pobla el Treeview con datos de un elemento XML."""
        tag = self._clean_tag(element.tag)
        
        # Manejar el caso especial de XML incrustado dentro de un texto
        if element.text and element.text.strip().startswith('<?xml'):
            node_id = self.data_tree.insert(parent_iid, 'end', text=tag, open=True, tags=('xml_tag',))
            self.xml_element_map[node_id] = element # Guardar el elemento padre para referencia
            try:
                # Parsear el XML incrustado
                cdata_node_id = self.data_tree.insert(node_id, 'end', text="[Contenido XML Adjunto]", open=True, tags=('xml_tag',))
                inner_root = ET.fromstring(element.text.strip())
                self.embedded_xml_docs[cdata_node_id] = (element, inner_root) # Almacenar relación padre-hijo y raíz interna
                self.xml_element_map[cdata_node_id] = inner_root # Mapear el nodo del treeview a la raíz interna
                self._populate_tree_from_xml(inner_root, cdata_node_id) # Procesar recursivamente el XML incrustado
            except ET.ParseError as e:
                err_node_id = self.data_tree.insert(node_id, 'end', text="[Error de parseo]", values=(str(e),), tags=('null',))
                self.xml_element_map[err_node_id] = None # Indicar que hubo un error
            return

        # Manejar el texto normal del elemento XML
        value_text = element.text.strip() if element.text and element.text.strip() else ""
        node_id = self.data_tree.insert(parent_iid, 'end', text=tag, values=(value_text,), tags=('xml_tag', self._get_value_tag(value_text)))
        self.xml_element_map[node_id] = element # Mapear el nodo del treeview al elemento XML

        # Manejar atributos del elemento XML
        if element.attrib:
            self.data_tree.item(node_id, open=True) # Abrir el nodo si tiene atributos
            attr_parent_id = self.data_tree.insert(node_id, 'end', text="@Atributos", open=True, tags=('xml_attr',))
            self.xml_element_map[attr_parent_id] = element # Referencia al elemento padre
            for key, value in element.attrib.items():
                attr_node_id = self.data_tree.insert(attr_parent_id, 'end', text=f"@{key}", values=(value,), tags=('xml_attr', self._get_value_tag(value)))
                self.xml_element_map[attr_node_id] = (element, key) # Mapear al par (elemento, clave_atributo)

        # Procesar recursivamente los hijos del elemento XML
        for child in element:
            self._populate_tree_from_xml(child, node_id)

    def _populate_tree_from_json(self, data, parent_iid):
        """Recursivamente pobla el Treeview con datos de un objeto JSON."""
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, dict):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"{{}} {key}", values=('{...}'), tags=('dict',))
                    self._populate_tree_from_json(value, node)
                elif isinstance(value, list):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"[] {key}", values=('[...]'), tags=('list',))
                    self._populate_tree_from_json(value, node)
                else:
                    self.data_tree.insert(parent_iid, 'end', text=key, values=(self._format_value(value),), tags=(self._get_value_tag(value),))
        elif isinstance(data, list):
            for index, item in enumerate(data):
                key = f"[{index}]"
                if isinstance(item, dict):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"{{}} {key}", values=('{...}'), tags=('dict',))
                    self._populate_tree_from_json(item, node)
                elif isinstance(item, list):
                    node = self.data_tree.insert(parent_iid, 'end', text=f"[] {key}", values=('[...]'), tags=('list',))
                    self._populate_tree_from_json(item, node)
                else:
                    self.data_tree.insert(parent_iid, 'end', text=key, values=(self._format_value(item),), tags=(self._get_value_tag(item),))
    
    def _format_value(self, value):
        """Formatea un valor para mostrarlo en el Treeview (ej: añade comillas a strings)."""
        if isinstance(value, str):
            return f'"{value}"' # Añadir comillas a las cadenas de texto
        return value # Devolver otros tipos (números, booleanos, None) tal cual

    def _get_node_path(self, item_id):
        """Obtiene la ruta de un nodo en el Treeview como una lista de claves/índices."""
        path = []
        parent = item_id
        while parent:
            text = self.data_tree.item(parent, "text")
            # Limpiar prefijos de icono para obtener la clave/índice real
            if text.startswith('{} ') or text.startswith('[] '):
                text = text[3:]
            path.insert(0, text)
            parent = self.data_tree.parent(parent)
        return path

    def _on_tree_double_click(self, event):
        """Maneja la doble clic para editar el valor de un nodo."""
        item_id = self.data_tree.focus()
        # Evitar la edición si es un nodo padre que se puede expandir o un contenedor de atributos
        if not item_id or self.data_tree.get_children(item_id):
            is_parent_of_attributes = self.data_tree.item(item_id, "text") == "@Atributos"
            if not is_parent_of_attributes:
                return

        if self.current_data_type == 'json':
            path = self._get_node_path(item_id)
            raw_value = self.data_tree.item(item_id, "values")[0]
            # Si el valor es una cadena, quitar las comillas para edición más limpia
            if isinstance(raw_value, str) and raw_value.startswith('"') and raw_value.endswith('"'):
                current_value = raw_value[1:-1]
            else:
                current_value = raw_value

            new_value_str = simpledialog.askstring("Editar Valor JSON", f"Nuevo valor para '{path[-1]}':", initialvalue=current_value, parent=self.root)

            if new_value_str is not None:
                try:
                    # Intentar convertir el valor a su tipo de dato correcto (bool, int, float, str, None)
                    if new_value_str.lower() == 'true': typed_value = True
                    elif new_value_str.lower() == 'false': typed_value = False
                    elif new_value_str.lower() == 'null': typed_value = None
                    else:
                        try: typed_value = int(new_value_str)
                        except ValueError:
                            try: typed_value = float(new_value_str)
                            except ValueError: typed_value = new_value_str
                    
                    # Navegar por la estructura de datos para encontrar el nodo y actualizarlo
                    data_node = self.parsed_data
                    for i, key_or_index_str in enumerate(path):
                        if i == len(path) - 1: # Es el último elemento (el que se va a modificar)
                            if isinstance(data_node, dict):
                                data_node[key_or_index_str] = typed_value
                            elif isinstance(data_node, list):
                                # Convertir el índice a entero
                                index = int(re.search(r'\[(\d+)\]', key_or_index_str).group(1))
                                data_node[index] = typed_value
                        else:
                            # Navegar al siguiente nivel
                            if isinstance(data_node, dict):
                                data_node = data_node[key_or_index_str]
                            elif isinstance(data_node, list):
                                # Convertir el índice a entero
                                index = int(re.search(r'\[(\d+)\]', key_or_index_str).group(1))
                                data_node = data_node[index]
                    
                    # Actualizar la visualización en el Treeview
                    self.data_tree.item(item_id, values=(self._format_value(typed_value),), tags=(self._get_value_tag(typed_value),))
                    self.log(f"📝 Valor JSON actualizado para '{'/'.join(path)}'. No olvides guardar.")
                    self.is_data_modified = True
                    self.save_button.config(state="normal")
                except Exception as e:
                    self.log(f"❌ Error al actualizar el valor JSON: {e}")

        elif self.current_data_type == 'xml':
            target = self.xml_element_map.get(item_id)
            if target is None: # Si no se encontró el elemento mapeado
                return

            current_value = self.data_tree.item(item_id, "values")[0]
            
            if isinstance(target, tuple): # Es un atributo
                element, key = target
                new_value = simpledialog.askstring(f"Editar Atributo XML", f"Nuevo valor para el atributo '{key}':", initialvalue=current_value, parent=self.root)
                if new_value is not None:
                    element.set(key, new_value) # Modificar el atributo en el elemento XML
                    self.data_tree.item(item_id, values=(new_value,), tags=('xml_attr', self._get_value_tag(new_value)))
                    self.log(f"📝 Atributo '{key}' actualizado. No olvides guardar.")
                    self.is_data_modified = True
                    self.save_button.config(state="normal")

            elif isinstance(target, ET.Element): # Es el texto de un elemento
                element = target
                new_value = simpledialog.askstring(f"Editar Valor de Tag XML", f"Nuevo valor para <{self._clean_tag(element.tag)}>:", initialvalue=current_value, parent=self.root)
                if new_value is not None:
                    element.text = new_value # Modificar el texto del elemento XML
                    self.data_tree.item(item_id, values=(new_value,), tags=('xml_tag', self._get_value_tag(new_value)))
                    self.log(f"📝 Valor del tag '<{self._clean_tag(element.tag)}>' actualizado. No olvides guardar.")
                    self.is_data_modified = True
                    self.save_button.config(state="normal")
            
    def _save_data_file(self):
        """Guarda los cambios realizados en el archivo de datos."""
        if not self.is_data_modified or not self.current_file_path:
            return
        
        warning_msg = f"¿Guardar cambios en '{os.path.basename(self.current_file_path)}'?"
        if self.current_data_type == 'xml':
            warning_msg += "\n\nADVERTENCIA: Guardar un XML firmado lo INVALIDARÁ digitalmente. ¿Continuar?"

        if not messagebox.askyesno("Confirmar Guardado", warning_msg, parent=self.root):
            return

        try:
            # Si es XML, reconstruir el texto con los cambios y guardar
            if self.current_data_type == 'xml':
                # Reconstruir el XML incrustado si existe
                for parent_element, inner_root in self.embedded_xml_docs.values():
                    inner_xml_string = ET.tostring(inner_root, encoding='unicode', method='xml')
                    parent_element.text = f'{inner_xml_string}'
                
                tree = ET.ElementTree(self.parsed_data)
                tree.write(self.current_file_path, encoding='utf-8', xml_declaration=True)
            
            # Si es JSON, volcar la estructura de datos modificada a formato JSON
            elif self.current_data_type == 'json':
                with open(self.current_file_path, 'w', encoding='utf-8') as f:
                    json.dump(self.parsed_data, f, indent=4, ensure_ascii=False) # indent=4 para formato legible

            self.log(f"💾 Cambios guardados en '{os.path.basename(self.current_file_path)}'.")
            self.is_data_modified = False # Marcar como no modificado
            self.save_button.config(state="disabled") # Deshabilitar botón de guardar
        except Exception as e:
            self.log(f"❌ Error al guardar el archivo: {e}")
            messagebox.showerror("Error de Guardado", f"No se pudo guardar el archivo.\n\n{e}", parent=self.root)

    def _show_context_menu(self, event):
        """Muestra el menú contextual en la posición del cursor del ratón."""
        self.data_context_menu.delete(0, 'end') # Limpiar ítems anteriores
        item_id = self.data_tree.identify_row(event.y) # Identificar el item bajo el cursor
        if not item_id: return # Si no hay item, no hacer nada
        
        self.data_tree.selection_set(item_id) # Seleccionar el item
        
        item_text = self.data_tree.item(item_id, 'text')
        is_parent = bool(self.data_tree.get_children(item_id)) # ¿Tiene hijos?
        is_attribute_parent = item_text == '@Atributos' # ¿Es el nodo padre de atributos XML?
        is_root = not self.data_tree.parent(item_id) # ¿Es el nodo raíz?
        
        # Añadir opciones de copiar si pyperclip está disponible
        if self.pyperclip_available:
            self.data_context_menu.add_command(label="Copiar Clave/Tag", command=lambda: self._copy_to_clipboard(item_id, 'key'))
            # Solo permitir copiar valor si no es un nodo padre (que no tiene valor directo visible)
            if not is_parent or (is_parent and self.current_data_type == 'xml' and not is_attribute_parent):
                 self.data_context_menu.add_command(label="Copiar Valor", command=lambda: self._copy_to_clipboard(item_id, 'value'))
            self.data_context_menu.add_command(label="Copiar Ruta", command=lambda: self._copy_to_clipboard(item_id, 'path'))
            self.data_context_menu.add_separator()

        # Opción de editar si el nodo no es un contenedor y no es un padre de atributos
        if not is_parent and not is_attribute_parent:
             self.data_context_menu.add_command(label="Editar Valor...", command=lambda: self._on_tree_double_click(None))
        
        # Opción de añadir/eliminar (actualmente deshabilitada)
        self.data_context_menu.add_command(label="Añadir/Eliminar (En desarrollo)", state='disabled')

        # Opciones de expansión si el nodo es un padre
        if is_parent and not is_attribute_parent: # No expandir el nodo @Atributos
            self.data_context_menu.add_separator()
            self.data_context_menu.add_command(label="Expandir Todo", command=lambda: self._toggle_expand(item_id, True))
            self.data_context_menu.add_command(label="Contraer Todo", command=lambda: self._toggle_expand(item_id, False))
        
        # Eliminar opción si el nodo no es raíz y no es padre de atributos
        if not is_root and not is_attribute_parent:
            self.data_context_menu.add_command(label="Eliminar", command=lambda: self._delete_item(item_id), foreground='red')

        # Mostrar el menú en la posición del clic
        self.data_context_menu.post(event.x_root, event.y_root)

    def _copy_to_clipboard(self, item_id, what):
        """Copia la información seleccionada (clave, valor o ruta) al portapapeles."""
        if not self.pyperclip_available:
            self.log("La función de copiar no está disponible. Instala 'pyperclip'.")
            return
        try:
            if what == 'key': # Copiar clave/tag
                text = self.data_tree.item(item_id, 'text')
                if text.startswith('{} ') or text.startswith('[] '): # Quitar prefijos
                    text = text[3:]
                self.pyperclip.copy(text)
                self.log(f"📋 Clave '{text}' copiada al portapapeles.")
            elif what == 'value': # Copiar valor
                val_tuple = self.data_tree.item(item_id, 'values')
                if val_tuple:
                    val = str(val_tuple[0])
                    if val.startswith('"') and val.endswith('"'): # Quitar comillas de cadenas JSON
                        val = val[1:-1]
                    self.pyperclip.copy(val)
                    self.log(f"📋 Valor '{val}' copiado al portapapeles.")
            elif what == 'path': # Copiar ruta (jerarquía de claves/índices)
                path_list = self._get_node_path(item_id)
                # Formatear ruta como 'clave1.clave2[0].clave3'
                path_str = ".".join(p.replace(" ", "").replace("[", ".").replace("]", "").replace(".", "") if p.startswith('[') else p for p in path_list)

                self.pyperclip.copy(path_str)
                self.log(f"📋 Ruta '{path_str}' copiada al portapapeles.")
        except Exception as e:
            self.log(f"❌ Error al copiar al portapapeles: {e}")
            messagebox.showerror("Error de Portapapeles", f"No se pudo copiar el texto.\n{e}", parent=self.root)

    def _toggle_expand(self, item_id, expand=True):
        """Expande o contrae recursivamente todos los hijos de un nodo."""
        self.data_tree.item(item_id, open=expand) # Controlar el estado de expansión del nodo actual
        for child_id in self.data_tree.get_children(item_id): # Iterar sobre los hijos
            self._toggle_expand(child_id, expand) # Llamada recursiva
    
    def _delete_item(self, item_id):
        """Elimina un item tanto del modelo de datos como del Treeview."""
        if not messagebox.askyesno("Confirmar Eliminación", f"¿Seguro que quieres eliminar el elemento '{self.data_tree.item(item_id, 'text')}' y todo su contenido?", parent=self.root):
            return
        
        self.log(f"🔴 Funcionalidad de eliminación pendiente de implementación completa en el modelo de datos.")
        messagebox.showinfo("En Desarrollo", "La eliminación directa de nodos está en desarrollo.\n\nComo alternativa, puedes editar el valor a 'null' (JSON) o un texto vacío y guardar el archivo.", parent=self.root)
    
    def _search_in_tree(self, event=None):
        """Busca el término introducido en el campo de búsqueda y resalta coincidencias."""
        self._clear_search_highlight() # Limpiar resaltados anteriores
        search_term = self.search_data_var.get().lower() # Obtener término de búsqueda en minúsculas
        if not search_term: return # No hacer nada si la búsqueda está vacía

        self.search_results_iids = [] # Reiniciar la lista de resultados
        self._collect_search_results(search_term, '') # Recopilar todos los nodos coincidentes
        
        if not self.search_results_iids: # Si no hay resultados
            self.log(f"Búsqueda: No se encontraron resultados para '{search_term}'.")
            return

        self.log(f"Búsqueda: Se encontraron {len(self.search_results_iids)} coincidencias para '{search_term}'.")
        
        first_item = None # Guardar el primer item para enfocarlo
        for item_id in self.search_results_iids:
            if not first_item:
                first_item = item_id
            
            # Añadir la etiqueta 'search_hit' para el resaltado
            current_tags = list(self.data_tree.item(item_id, 'tags'))
            if 'search_hit' not in current_tags:
                current_tags.append('search_hit')
                self.data_tree.item(item_id, tags=tuple(current_tags))

        # Si se encontraron resultados, enfocar el primero
        if first_item:
            self._expand_to_item(first_item) # Asegurarse de que el nodo y sus padres estén expandidos
            self.data_tree.selection_set(first_item) # Seleccionar el nodo
            self.data_tree.focus(first_item) # Establecer foco
            self.data_tree.see(first_item) # Asegurarse de que sea visible

    def _clear_search_highlight(self):
        """Elimina el resaltado de búsqueda de todos los items."""
        if not self.search_results_iids: return # Si no hay resultados, salir
            
        for item_id in self.search_results_iids:
            current_tags = list(self.data_tree.item(item_id, 'tags'))
            if 'search_hit' in current_tags: # Si tiene la etiqueta de resaltado
                current_tags.remove('search_hit') # Quitarla
                self.data_tree.item(item_id, tags=tuple(current_tags)) # Aplicar cambios
        self.search_results_iids = [] # Limpiar la lista de IDs
        self.log("Resaltado de búsqueda limpiado.")

    def _collect_search_results(self, term, parent_iid):
        """Busca recursivamente un término en los nodos del Treeview y añade los IIDs coincidentes a la lista global."""
        for item_id in self.data_tree.get_children(parent_iid): # Iterar sobre los hijos del nodo actual
            item_text = self.data_tree.item(item_id, "text").lower() # Texto del nodo (clave/tag)
            values_tuple = self.data_tree.item(item_id, "values") # Valor del nodo
            item_value = str(values_tuple[0]).lower() if values_tuple else "" # Obtener valor o cadena vacía
            
            # Comprobar si el término de búsqueda está en el texto o valor del nodo
            if term in item_text or term in item_value:
                self.search_results_iids.append(item_id) # Añadir el ID del nodo a la lista de resultados
            
            # Si el nodo tiene hijos, llamar recursivamente para buscarlos
            if self.data_tree.get_children(item_id):
                self._collect_search_results(term, item_id)
            
    def _expand_to_item(self, item_id):
        """Expande todos los nodos padre de un item para hacerlo visible en el Treeview."""
        parent = self.data_tree.parent(item_id) # Obtener el padre del item
        while parent: # Mientras haya un padre
            self.data_tree.item(parent, open=True) # Expandir el padre
            parent = self.data_tree.parent(parent) # Moverse al siguiente nivel padre
            
    def _on_tree_motion(self, event):
        """Maneja el evento de movimiento del ratón sobre el Treeview para mostrar tooltips."""
        if self.tooltip_window: # Si ya hay una ventana de tooltip, destruirla
            self.tooltip_window.destroy()
            self.tooltip_window = None

        item_id = self.data_tree.identify_row(event.y) # Identificar la fila bajo el cursor
        # Comprobar si es un item válido y si el cursor está sobre la columna de valor (#2)
        if not item_id or self.data_tree.identify_column(event.x) != "#2": 
            return
        
        try:
            values_tuple = self.data_tree.item(item_id, "values")
            if not values_tuple or not values_tuple[0]: return # Si no hay valor, salir
            
            full_text = str(values_tuple[0]) # Obtener el texto completo del valor
            
            bbox = self.data_tree.bbox(item_id, column="value") # Obtener el cuadro delimitador de la celda
            if not bbox: return

            column_width = self.data_tree.column("value", "width") # Obtener el ancho de la columna
            
            # Estimación simple: si el texto es más largo que el ancho de la columna, mostrar tooltip
            if len(full_text) * 7 > column_width: # 7px por caracter es una estimación
                self.tooltip_window = tk.Toplevel(self.root)
                self.tooltip_window.wm_overrideredirect(True) # Sin bordes ni título
                # Posicionar el tooltip cerca del cursor
                self.tooltip_window.wm_geometry(f"+{event.x_root + 15}+{event.y_root + 10}")
                
                label = ttk.Label(self.tooltip_window, text=full_text,
                                  background="#ffffe0", relief="solid", borderwidth=1, # Estilo del tooltip
                                  wraplength=600, padding=5) # Limitar ancho y añadir padding
                label.pack()
        except (IndexError, AttributeError): # Manejar posibles errores al acceder a datos
            pass

    def _on_tree_leave(self, event):
        """Oculta el tooltip cuando el cursor sale del área del Treeview."""
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

# --- Bloque principal para iniciar la aplicación ---
if __name__ == "__main__":
    root = tk.Tk()
    app = MainApp(root) # Usamos la clase principal renombrada
    root.mainloop()